"""
Fylorra - Office Tools (pure Python where possible)
Focused on reliable conversions that don't need external apps.
"""

from __future__ import annotations

import csv
import re
import unicodedata
from dataclasses import dataclass
from pathlib import Path
from typing import Optional


@dataclass(frozen=True)
class OfficeOpResult:
    ok: bool
    message: str
    output_path: Optional[str] = None


def _parse_markdown_blocks(text: str) -> list[tuple[str, int, str]]:
    blocks: list[tuple[str, int, str]] = []
    para: list[str] = []
    in_code = False
    code_lines: list[str] = []

    def flush_para() -> None:
        if para:
            blocks.append(("p", 0, " ".join(para).strip()))
            para.clear()

    def flush_code() -> None:
        if code_lines:
            blocks.append(("code", 0, "\n".join(code_lines).strip("\n")))
            code_lines.clear()

    for raw in str(text or "").splitlines():
        line = raw.rstrip()
        stripped = line.strip()
        if stripped.startswith("```"):
            flush_para()
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if not stripped:
            flush_para()
            blocks.append(("spacer", 0, ""))
            continue
        if stripped in {"---", "***", "___"}:
            flush_para()
            blocks.append(("spacer", 0, ""))
            continue
        if stripped.startswith("#"):
            flush_para()
            level = len(stripped) - len(stripped.lstrip("#"))
            title = stripped[level:].strip()
            if title:
                blocks.append(("h", max(1, min(level, 6)), title))
            continue
        if stripped.startswith("- ") or stripped.startswith("* "):
            flush_para()
            blocks.append(("li", 0, stripped[2:].strip()))
            continue
        num_match = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if num_match:
            flush_para()
            blocks.append(("li_num", 0, num_match.group(2).strip()))
            continue
        para.append(stripped)

    flush_para()
    if in_code:
        flush_code()
    return blocks


def _strip_inline_markdown(text: str) -> str:
    line = str(text or "")
    if not line:
        return ""
    line = re.sub(r"`([^`]+)`", r"\1", line)
    line = re.sub(r"\*\*([^*]+)\*\*", r"\1", line)
    line = re.sub(r"\*([^*]+)\*", r"\1", line)
    line = re.sub(r"(?<!\w)__([^_]+)__(?!\w)", r"\1", line)
    line = re.sub(r"(?<!\w)_([^_]+)_(?!\w)", r"\1", line)
    line = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", line)
    return line.strip()


def _sanitize_pdf_text(text: str) -> str:
    if not text:
        return ""
    out = unicodedata.normalize("NFKC", str(text))
    replacements = {
        "’": "'",
        "‘": "'",
        "“": "\"",
        "”": "\"",
        "–": "-",
        "—": "-",
        "•": "-",
        "…": "...",
        " ": " ",
    }
    for k, v in replacements.items():
        out = out.replace(k, v)
    out = re.sub(r"[\x00-\x08\x0B-\x1F\x7F]", "", out)
    out = re.sub(r"\s+", " ", out).strip()
    return out


def xlsx_to_csv(
    input_path: Path,
    *,
    output_path: Path,
    sheet: str | int | None = None,
    overwrite: bool = False,
    delimiter: str = ",",
) -> OfficeOpResult:
    try:
        import openpyxl
    except Exception as e:
        raise RuntimeError("Excel→CSV requires 'openpyxl'. Install: pip install openpyxl") from e

    input_path = Path(input_path)
    output_path = Path(output_path)
    if not input_path.exists():
        return OfficeOpResult(ok=False, message="Input file not found.")
    if output_path.exists() and not overwrite:
        return OfficeOpResult(ok=False, message=f"Output already exists: {output_path.name}")

    wb = openpyxl.load_workbook(filename=str(input_path), read_only=True, data_only=True)
    try:
        if sheet is None:
            ws = wb.worksheets[0]
        elif isinstance(sheet, int):
            ws = wb.worksheets[int(sheet)]
        else:
            ws = wb[str(sheet)]

        output_path.parent.mkdir(parents=True, exist_ok=True)
        with output_path.open("w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f, delimiter=delimiter)
            for row in ws.iter_rows(values_only=True):
                writer.writerow(["" if v is None else v for v in row])
        return OfficeOpResult(ok=True, message="Converted.", output_path=str(output_path))
    finally:
        try:
            wb.close()
        except Exception:
            pass


def markdown_to_docx(
    markdown_text: str,
    *,
    output_path: Path,
    overwrite: bool = False,
) -> OfficeOpResult:
    try:
        from docx import Document  # type: ignore
        from docx.shared import Inches, Pt  # type: ignore
    except Exception as e:
        return OfficeOpResult(ok=False, message="Markdown→DOCX requires 'python-docx'. Install: pip install python-docx")

    output_path = Path(output_path)
    if output_path.exists() and not overwrite:
        return OfficeOpResult(ok=False, message=f"Output already exists: {output_path.name}")
    output_path.parent.mkdir(parents=True, exist_ok=True)

    try:
        doc = Document()
        try:
            sec = doc.sections[0]
            sec.top_margin = Inches(0.8)
            sec.bottom_margin = Inches(0.8)
            sec.left_margin = Inches(0.9)
            sec.right_margin = Inches(0.9)
        except Exception:
            pass

        try:
            normal = doc.styles["Normal"]
            normal.font.name = "Calibri"
            normal.font.size = Pt(11)
        except Exception:
            pass

        blocks = _parse_markdown_blocks(markdown_text)
        if not blocks:
            blocks = [("p", 0, str(markdown_text or "").strip())]

        for kind, level, content in blocks:
            content = _strip_inline_markdown(content)
            if kind == "h":
                doc.add_heading(content, level=min(max(int(level), 1), 3))
            elif kind == "li":
                try:
                    doc.add_paragraph(content, style="List Bullet")
                except Exception:
                    doc.add_paragraph(f"- {content}")
            elif kind == "li_num":
                try:
                    doc.add_paragraph(content, style="List Number")
                except Exception:
                    doc.add_paragraph(content)
            elif kind == "code":
                p = doc.add_paragraph()
                run = p.add_run(content)
                try:
                    run.font.name = "Consolas"
                    run.font.size = Pt(10)
                except Exception:
                    pass
            elif kind == "p":
                doc.add_paragraph(content)
            else:
                doc.add_paragraph("")

        doc.save(str(output_path))
        return OfficeOpResult(ok=True, message="Exported Markdown to DOCX.", output_path=str(output_path))
    except Exception as e:
        return OfficeOpResult(ok=False, message=f"DOCX export failed: {e}")


def markdown_to_pdf(
    markdown_text: str,
    *,
    output_path: Path,
    overwrite: bool = False,
) -> OfficeOpResult:
    output_path = Path(output_path)
    if output_path.exists() and not overwrite:
        return OfficeOpResult(ok=False, message=f"Output already exists: {output_path.name}")
    output_path.parent.mkdir(parents=True, exist_ok=True)

    try:
        from reportlab.lib.pagesizes import letter  # type: ignore
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle  # type: ignore
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer  # type: ignore
        from xml.sax.saxutils import escape
    except Exception:
        try:
            from reportlab.lib.pagesizes import letter  # type: ignore
            from reportlab.pdfgen import canvas  # type: ignore
        except Exception:
            return OfficeOpResult(ok=False, message="Markdown→PDF requires 'reportlab'. Install: pip install reportlab")
        try:
            c = canvas.Canvas(str(output_path), pagesize=letter)
            width, height = letter
            y = height - 72
            for line in str(markdown_text or "").splitlines():
                line = _sanitize_pdf_text(_strip_inline_markdown(line))
                if y < 72:
                    c.showPage()
                    y = height - 72
                c.drawString(72, y, line[:120])
                y -= 14
            c.save()
            return OfficeOpResult(ok=True, message="Exported Markdown to PDF.", output_path=str(output_path))
        except Exception as e:
            return OfficeOpResult(ok=False, message=f"PDF export failed: {e}")

    try:
        styles = getSampleStyleSheet()
        h1 = ParagraphStyle("MDHeading1", parent=styles["Heading1"], fontSize=17, leading=21, spaceAfter=8)
        h2 = ParagraphStyle("MDHeading2", parent=styles["Heading2"], fontSize=14, leading=18, spaceAfter=6)
        h3 = ParagraphStyle("MDHeading3", parent=styles["Heading3"], fontSize=12, leading=16, spaceAfter=4)
        body = ParagraphStyle("MDBody", parent=styles["BodyText"], fontSize=11, leading=15, spaceAfter=6)
        bullet = ParagraphStyle("MDBullet", parent=body, leftIndent=14, bulletIndent=4, spaceAfter=2)
        code = ParagraphStyle("MDCode", parent=body, fontName="Courier", fontSize=9.5, leading=12, leftIndent=10, spaceAfter=6)

        story = []
        blocks = _parse_markdown_blocks(markdown_text)
        if not blocks:
            blocks = [("p", 0, str(markdown_text or "").strip())]
        for kind, level, content in blocks:
            if kind == "code":
                txt = _sanitize_pdf_text(content).replace("\n", "<br/>")
                story.append(Paragraph(escape(txt), code))
                continue
            content = _sanitize_pdf_text(_strip_inline_markdown(content))
            if kind == "h":
                style = h1 if level <= 1 else h2 if level == 2 else h3
                story.append(Paragraph(escape(content), style))
            elif kind == "li":
                story.append(Paragraph(escape(f"- {content}"), bullet))
            elif kind == "li_num":
                story.append(Paragraph(escape(f"1. {content}"), bullet))
            elif kind == "p":
                story.append(Paragraph(escape(content), body))
            else:
                story.append(Spacer(1, 8))
        if not story:
            story = [Paragraph(escape(_sanitize_pdf_text(markdown_text)), body)]
        doc = SimpleDocTemplate(str(output_path), pagesize=letter, leftMargin=48, rightMargin=48, topMargin=54, bottomMargin=54)
        doc.build(story)
        return OfficeOpResult(ok=True, message="Exported Markdown to PDF.", output_path=str(output_path))
    except Exception as e:
        return OfficeOpResult(ok=False, message=f"PDF export failed: {e}")
