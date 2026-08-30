from __future__ import annotations

from pathlib import Path
import re
import unicodedata
from typing import Any

from core.pipeline.agent import AgentCapability, AgentResult, PipelineAgent
from core.pipeline.context import PipelineContext


def _parse_markdown_blocks(text: str) -> list[tuple[str, int, str]]:
    blocks: list[tuple[str, int, str]] = []
    para: list[str] = []

    def flush_para():
        if para:
            blocks.append(("p", 0, " ".join(para).strip()))
            para.clear()

    for raw in str(text or "").splitlines():
        line = raw.strip()
        if not line:
            flush_para()
            blocks.append(("spacer", 0, ""))
            continue
        if line in {"---", "***", "___"}:
            flush_para()
            blocks.append(("spacer", 0, ""))
            continue
        if line.startswith("#"):
            flush_para()
            level = len(line) - len(line.lstrip("#"))
            title = line[level:].strip()
            if title:
                blocks.append(("h", level, title))
            continue
        if line.startswith("- ") or line.startswith("* "):
            flush_para()
            blocks.append(("li", 0, line[2:].strip()))
            continue
        if "." in line:
            prefix, rest = line.split(".", 1)
            if prefix.isdigit() and rest.startswith(" "):
                flush_para()
                blocks.append(("li", 0, rest.strip()))
                continue
        para.append(line)

    flush_para()
    return blocks


def _sanitize_pdf_text(text: str) -> str:
    if not text:
        return ""
    out = unicodedata.normalize("NFKC", str(text))
    replacements = {
        "’": "'",
        "‘": "'",
        "“": "\"",
        "”": "\"",
        "–": "-",
        "—": "-",
        "•": "-",
        "…": "...",
        " ": " ",
    }
    for k, v in replacements.items():
        out = out.replace(k, v)
    # Keep accented letters and other valid Unicode; only remove control chars.
    out = re.sub(r"[\x00-\x08\x0B-\x1F\x7F]", "", out)
    out = re.sub(r"\s+", " ", out).strip()
    return out


def _strip_inline_markdown(text: str) -> str:
    line = str(text or "")
    if not line:
        return ""
    # Strip simple markdown emphasis/inline code markers for cleaner exports.
    line = re.sub(r"`([^`]+)`", r"\1", line)
    line = re.sub(r"\*\*([^*]+)\*\*", r"\1", line)
    line = re.sub(r"\*([^*]+)\*", r"\1", line)
    # Only strip underscore-based emphasis when underscores are standalone markers,
    # not when they're part of URLs/identifiers.
    line = re.sub(r"(?<!\w)__([^_]+)__(?!\w)", r"\1", line)
    line = re.sub(r"(?<!\w)_([^_]+)_(?!\w)", r"\1", line)
    return line.strip()


def _collapse_sources_sections(document: str) -> str:
    if not document:
        return document
    has_works_cited = bool(
        re.search(r"(?im)^\s*#{1,6}\s+(Works Cited|Where We Learned This)\b", document)
    )
    pattern = re.compile(r"(?ms)^\s*#+\s+Sources[^\n]*\n.*?(?=^\s*#+\s+\w|\Z)", re.IGNORECASE)
    matches = list(pattern.finditer(document))
    if not matches:
        return document
    cleaned = pattern.sub("", document).strip()
    if has_works_cited:
        return cleaned
    first_block = matches[0].group(0).strip()
    if cleaned:
        return cleaned.rstrip() + "\n\n" + first_block
    return first_block


def _strip_verification_report(document: str) -> str:
    if not document:
        return document
    lines = str(document).splitlines()
    start = -1
    start_level = 0
    for idx, line in enumerate(lines):
        m = re.match(r"^\s*(#{1,6})\s*Verification Report\b", line, flags=re.IGNORECASE)
        if m:
            start = idx
            start_level = len(m.group(1))
            break
    if start < 0:
        return str(document).strip()
    # Strip the entire verification block. The report contains internal headings like
    # "## Verified", so we must resume only at the next heading of the same or higher
    # level (<= start_level).
    for idx in range(start + 1, len(lines)):
        m = re.match(r"^\s*(#{1,6})\s+\S", lines[idx])
        if not m:
            continue
        if re.match(r"^\s*#{1,6}\s*Verification Report\b", lines[idx], flags=re.IGNORECASE):
            continue
        level = len(m.group(1))
        if level <= start_level:
            return "\n".join(lines[idx:]).strip()
    return ""


def _dedupe_section(document: str, title: str) -> str:
    if not document:
        return document
    pattern = re.compile(rf"(?ms)^\s*#+\s+{re.escape(title)}[^\n]*\n.*?(?=^\s*#+\s+\w|\Z)", re.IGNORECASE)
    matches = list(pattern.finditer(document))
    if len(matches) <= 1:
        return document
    kept = matches[0].group(0).strip()
    cleaned = pattern.sub("", document).strip()
    if cleaned:
        return cleaned.rstrip() + "\n\n" + kept
    return kept


def _strip_internal_reasoning(document: str) -> str:
    if not document:
        return document
    text = str(document)
    text = re.sub(r"(?is)<\s*(think|analysis|reasoning)[^>]*>.*?<\s*/\s*\1\s*>", "", text)
    text = re.sub(r"(?is)```(?:thinking|analysis|reasoning).*?```", "", text)
    lines = []
    for line in text.splitlines():
        s = line.strip().lower()
        if s in {"<think>", "</think>", "<analysis>", "</analysis>", "<reasoning>", "</reasoning>"}:
            continue
        lines.append(line)
    return "\n".join(lines).strip()


class ExportAgent(PipelineAgent):
    def __init__(self, config: dict[str, Any] | None = None):
        super().__init__(
            agent_id="export_agent",
            role="exporter",
            system_prompt="Export formatted documents to files.",
            capabilities=[
                AgentCapability(
                    capability_id="export",
                    name="Multi-Format Export",
                    description="Export markdown/plain text to PDF/DOCX/MD.",
                    input_types=["text"],
                    output_types=["file"],
                    requires_internet=False,
                    estimated_time_seconds=15,
                )
            ],
            config=config or {},
        )

    def execute(self, context: PipelineContext, inputs: dict[str, Any] | None = None) -> AgentResult:
        payload = dict(inputs or {})
        text = payload.get("document") or payload.get("text") or ""
        text = _strip_internal_reasoning(str(text))
        text = _strip_verification_report(str(text))
        text = _collapse_sources_sections(str(text))
        text = _dedupe_section(text, "Discussion Questions")
        text = _dedupe_section(text, "Key People")
        text = _dedupe_section(text, "Timeline")
        text = _dedupe_section(text, "Key Events")
        text = _dedupe_section(text, "Key Turning Points")
        out_folder = Path(str(payload.get("output_folder") or context.initial_parameters.get("output_folder") or context.working_directory))
        out_folder.mkdir(parents=True, exist_ok=True)

        name = str(payload.get("output_name") or context.initial_parameters.get("output_name") or "workflow_output").strip()
        safe = "".join(ch if ch.isalnum() or ch in {"-", "_"} else "_" for ch in name).strip("_")
        safe = safe or "workflow_output"

        formats = payload.get("formats") or context.initial_parameters.get("export_formats") or ["md"]
        if isinstance(formats, str):
            formats = [formats]

        written: list[str] = []
        for fmt in formats:
            fmt = str(fmt or "").lower().lstrip(".")
            if fmt in {"md", "markdown"}:
                p = out_folder / f"{safe}.md"
                p.write_text(str(text), encoding="utf-8")
                written.append(str(p))
            elif fmt in {"txt", "text"}:
                p = out_folder / f"{safe}.txt"
                p.write_text(str(text), encoding="utf-8")
                written.append(str(p))
            elif fmt == "docx":
                try:
                    import docx
                except Exception:
                    continue
                doc = docx.Document()
                blocks = _parse_markdown_blocks(str(text))
                for kind, level, content in blocks:
                    content = _strip_inline_markdown(content)
                    if kind == "h":
                        doc.add_heading(content, level=min(max(int(level), 1), 3))
                    elif kind == "li":
                        try:
                            doc.add_paragraph(content, style="List Bullet")
                        except Exception:
                            doc.add_paragraph(f"- {content}")
                    elif kind == "p":
                        doc.add_paragraph(content)
                    else:
                        doc.add_paragraph("")
                p = out_folder / f"{safe}.docx"
                doc.save(str(p))
                written.append(str(p))
            elif fmt == "pdf":
                p = out_folder / f"{safe}.pdf"
                try:
                    from reportlab.lib.pagesizes import letter
                    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                    from xml.sax.saxutils import escape
                except Exception:
                    try:
                        from reportlab.lib.pagesizes import letter
                        from reportlab.pdfgen import canvas
                    except Exception:
                        continue
                    c = canvas.Canvas(str(p), pagesize=letter)
                    width, height = letter
                    y = height - 72
                    for line in str(text).splitlines():
                        line = _sanitize_pdf_text(_strip_inline_markdown(line))
                        if y < 72:
                            c.showPage()
                            y = height - 72
                        c.drawString(72, y, line[:120])
                        y -= 14
                    c.save()
                    written.append(str(p))
                    continue

                styles = getSampleStyleSheet()
                h1 = ParagraphStyle("WFHeading1", parent=styles["Heading1"], fontSize=16, leading=20, spaceAfter=8)
                h2 = ParagraphStyle("WFHeading2", parent=styles["Heading2"], fontSize=14, leading=18, spaceAfter=6)
                h3 = ParagraphStyle("WFHeading3", parent=styles["Heading3"], fontSize=12, leading=16, spaceAfter=4)
                body = ParagraphStyle("WFBody", parent=styles["BodyText"], fontSize=11, leading=15, spaceAfter=6)
                bullet = ParagraphStyle("WFBullet", parent=body, leftIndent=14, bulletIndent=4, spaceAfter=2)

                story = []
                blocks = _parse_markdown_blocks(str(text))
                for kind, level, content in blocks:
                    content = _sanitize_pdf_text(_strip_inline_markdown(content))
                    if kind == "h":
                        style = h1 if level <= 1 else h2 if level == 2 else h3
                        story.append(Paragraph(escape(content), style))
                    elif kind == "li":
                        # Prefix with "-" to avoid odd control glyphs in some PDF text extractors.
                        story.append(Paragraph(escape(f"- {content}"), bullet))
                    elif kind == "p":
                        story.append(Paragraph(escape(content), body))
                    else:
                        story.append(Spacer(1, 8))
                if not story:
                    story = [Paragraph(escape(str(text)), body)]
                doc = SimpleDocTemplate(str(p), pagesize=letter, leftMargin=48, rightMargin=48, topMargin=54, bottomMargin=54)
                doc.build(story)
                written.append(str(p))
            else:
                continue

        if not written:
            return AgentResult(ok=False, message="No outputs created. Missing export libraries or formats.")
        return AgentResult(ok=True, data={"files": written})
