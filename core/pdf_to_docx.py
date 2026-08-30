"""
Fylorra - PDF -> DOCX (fallback)
Best-effort text extraction to DOCX when LibreOffice export isn't available/reliable.

This is not layout-perfect; it aims to be useful (searchable/editable text).
"""

from __future__ import annotations

from pathlib import Path


def pdf_to_docx_text(pdf_path: Path, *, output_path: Path) -> Path:
    pdf_path = Path(pdf_path)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    from docx import Document  # python-docx

    doc = Document()
    text_pages: list[str] = []

    # Prefer PyMuPDF when available (best text extraction quality).
    try:
        import fitz  # type: ignore

        with fitz.open(pdf_path) as pdf:
            for page_index in range(len(pdf)):
                page = pdf[page_index]
                text_pages.append(page.get_text("text") or "")
    except Exception:
        # Fallback: pypdf (works without PyMuPDF, but may extract less text on some PDFs).
        from pypdf import PdfReader

        reader = PdfReader(str(pdf_path))
        for page in reader.pages:
            try:
                text_pages.append(page.extract_text() or "")
            except Exception:
                text_pages.append("")

    for i, text in enumerate(text_pages):
        if text.strip():
            for line in text.splitlines():
                doc.add_paragraph(line)
        else:
            doc.add_paragraph("[No extractable text on this page]")
        if i < len(text_pages) - 1:
            doc.add_page_break()

    doc.save(output_path)
    return output_path
