from __future__ import annotations

import re
import html
import uuid
import queue
import threading
import time
from dataclasses import dataclass, field
from pathlib import Path

from PySide6.QtCore import Qt, QSize, QObject, Signal, QThread, QUrl, QEvent, QFileInfo, QRectF, QByteArray, QTimer, QDateTime, QPointF, QBuffer, QIODevice
from PySide6.QtGui import QDesktopServices, QAction, QPixmap, QIcon, QImage, QImageReader, QPainter, QPolygonF, QColor, QTextCursor, QIntValidator
from PySide6.QtWidgets import (
    QAbstractItemView,
    QApplication,
    QCheckBox,
    QComboBox,
    QButtonGroup,
    QDialog,
    QDialogButtonBox,
    QDateTimeEdit,
    QFormLayout,
    QFrame,
    QGridLayout,
    QGroupBox,
    QHBoxLayout,
    QHeaderView,
    QLabel,
    QListWidget,
    QListWidgetItem,
    QLineEdit,
    QMainWindow,
    QMessageBox,
    QMenu,
    QProgressBar,
    QPushButton,
    QRadioButton,
    QScrollArea,
    QSlider,
    QSizePolicy,
    QSpinBox,
    QStackedWidget,
    QTableWidget,
    QTableWidgetItem,
    QTextEdit,
    QTextBrowser,
    QToolButton,
    QVBoxLayout,
    QWidget,
    QFileDialog,
    QSplitter,
    QSystemTrayIcon,
    QTabWidget,
    QInputDialog,
    QFileSystemModel,
    QTreeView,
    QTreeWidget,
    QTreeWidgetItem,
    QFileIconProvider,
    QStyle,
)
from PySide6.QtSvg import QSvgRenderer

from qt_app.icons import QtIconLoader
from qt_app.pages import PageDef, build_placeholder_page
from qt_app.signals import MonitorSignals
from core.branding import APP_NAME, APP_WINDOW_TITLE, DEFAULT_LINKS_FOLDER_NAME


def _render_svg_pixmap(svg_path: Path, size: QSize, *, viewbox: str | None = None) -> QPixmap:
    pm = QPixmap(size)
    pm.fill(Qt.transparent)
    try:
        raw = Path(svg_path).read_bytes()
        if viewbox:
            text = raw.decode("utf-8", errors="ignore")
            text = re.sub(r'viewBox="[^"]+"', f'viewBox="{viewbox}"', text, count=1)
            raw = text.encode("utf-8")
        renderer = QSvgRenderer()
        renderer.load(QByteArray(raw))
        if not renderer.isValid():
            return QPixmap()
        painter = QPainter(pm)
        try:
            painter.setRenderHint(QPainter.Antialiasing, True)
            painter.setRenderHint(QPainter.SmoothPixmapTransform, True)
            renderer.render(painter, QRectF(0, 0, size.width(), size.height()))
        finally:
            painter.end()
        return pm
    except Exception:
        return QPixmap()


def _qt_modern_dialog_stylesheet() -> str:
    """Shared dialog stylesheet that follows the active app theme/accent."""
    mode = "dark"
    accent = "blue"
    try:
        app = QApplication.instance()
        if app:
            mode = str(app.property("fg_theme_mode") or "dark").strip().lower()
            accent = str(app.property("fg_color_theme") or "blue").strip().lower()
    except Exception:
        pass

    if mode == "light":
        bg = "#f4f6fb"
        card = "#ffffff"
        border = "#d8dee9"
        text = "#1f2937"
        muted = "#5f6b7a"
        input_bg = "#ffffff"
        input_text = "#1f2937"
        btn_bg = "#edf2f8"
        btn_hover = "#e3eaf5"
        btn_press = "#d9e3f2"
        tab_bg = "#edf2f8"
        tree_bg = "#fbfcff"
    elif mode == "black":
        bg = "#0a0a0a"
        card = "#101010"
        border = "#1f1f1f"
        text = "#efefef"
        muted = "#a0a0a0"
        input_bg = "#090909"
        input_text = "#f2f2f2"
        btn_bg = "#1b1b1b"
        btn_hover = "#252525"
        btn_press = "#151515"
        tab_bg = "#131313"
        tree_bg = "#0e0e0e"
    else:
        bg = "#0f1217"
        card = "#14171c"
        border = "#232730"
        text = "#c8ccd6"
        muted = "#9aa0a9"
        input_bg = "#101318"
        input_text = "#e6e9ef"
        btn_bg = "#2a2f38"
        btn_hover = "#343a45"
        btn_press = "#242a33"
        tab_bg = "#1a1e25"
        tree_bg = "#14171c"

    accent_map = {
        "blue": ("#1677ff", "#2a86ff", "#0f5bd6"),
        "teal": ("#14b8a6", "#12a392", "#0d7f72"),
        "green": ("#22c55e", "#16a34a", "#15803d"),
        "orange": ("#f59e0b", "#d97706", "#b45309"),
        "rose": ("#f43f5e", "#e11d48", "#be123c"),
        "violet": ("#8b5cf6", "#7c3aed", "#6d28d9"),
        "cyan": ("#06b6d4", "#0891b2", "#0e7490"),
    }
    ac, ac_hover, ac_press = accent_map.get(accent, accent_map["blue"])

    return (
        f"QDialog{{background:{bg}; color:{text};}}"
        f"#DialogHeader{{background:{card}; border:1px solid {border}; border-radius:12px;}}"
        f"#DialogTitle{{font-size:20px; font-weight:800; color:{text};}}"
        f"#DialogSubtitle{{color:{muted};}}"
        f"#DialogSummary{{color:{text}; font-weight:700; padding:4px 0;}}"
        f"#Card{{background:{card}; border:1px solid {border}; border-radius:12px;}}"
        f"QGroupBox{{color:{text}; border:1px solid {border}; border-radius:12px; margin-top:12px;}}"
        "QGroupBox::title{subcontrol-origin: margin; left: 12px; padding: 0 4px;}"
        f"QLabel{{color:{text};}}"
        f"QCheckBox{{color:{text}; spacing:10px;}}"
        f"QCheckBox::indicator{{width:16px; height:16px; border-radius:4px; border:1px solid {border}; background:{input_bg};}}"
        f"QCheckBox::indicator:checked{{background:{ac}; border:1px solid {ac};}}"
        f"QLineEdit,QComboBox,QSpinBox,QTextEdit{{background:{input_bg}; border:1px solid {border}; border-radius:10px; padding:8px; color:{input_text};}}"
        "QComboBox::drop-down{border:0; width:22px;}"
        f"QComboBox::down-arrow{{image:none; border-left:5px solid transparent; border-right:5px solid transparent; border-top:6px solid {muted}; margin-right:8px;}}"
        f"#PrimaryButton{{background:{ac}; color:#ffffff; border:0; border-radius:10px; padding:10px 14px; font-weight:700;}}"
        f"#PrimaryButton:hover{{background:{ac_hover};}}"
        f"#PrimaryButton:pressed{{background:{ac_press};}}"
        f"#SecondaryButton{{background:{btn_bg}; color:{text}; border:1px solid {border}; border-radius:10px; padding:9px 12px;}}"
        f"#SecondaryButton:hover{{background:{btn_hover};}}"
        f"#SecondaryButton:pressed{{background:{btn_press};}}"
        f"QPushButton{{background:{btn_bg}; color:{text}; border:1px solid {border}; border-radius:10px; padding:9px 12px;}}"
        f"QPushButton:hover{{background:{btn_hover};}}"
        f"QPushButton:pressed{{background:{btn_press};}}"
        f"QProgressBar{{background:{input_bg}; border:1px solid {border}; border-radius:10px; height:16px; text-align:center; color:{text};}}"
        f"QProgressBar::chunk{{background:{ac}; border-radius:10px;}}"
        f"#ResultsTree{{background:{tree_bg}; border:1px solid {border}; border-radius:10px; color:{text};}}"
        "QTreeWidget::item{padding:4px 2px;}"
        f"QHeaderView::section{{background:{input_bg}; color:{muted}; border:0; padding:6px;}}"
        f"QTabBar::tab{{background:{tab_bg}; color:{text}; padding:8px 12px; border-radius:8px; margin-right:6px;}}"
        f"QTabBar::tab:selected{{background:{ac}; color:#ffffff;}}"
        f"QListWidget{{background:{tree_bg}; border:1px solid {border}; border-radius:10px; color:{text};}}"
    )


def _ui_theme_mode() -> str:
    try:
        app = QApplication.instance()
        if app:
            mode = str(app.property("fg_theme_mode") or "dark").strip().lower()
            if mode in {"dark", "light", "black"}:
                return mode
    except Exception:
        pass
    return "dark"


def _ui_theme_tokens() -> dict[str, str]:
    mode = _ui_theme_mode()
    accent = "blue"
    try:
        app = QApplication.instance()
        if app:
            accent = str(app.property("fg_color_theme") or "blue").strip().lower()
    except Exception:
        accent = "blue"
    accent_map = {
        "blue": "#0d6efd",
        "teal": "#14b8a6",
        "green": "#22c55e",
        "orange": "#f59e0b",
        "rose": "#f43f5e",
        "violet": "#8b5cf6",
        "cyan": "#06b6d4",
    }
    acc = accent_map.get(accent, "#0d6efd")
    if mode == "light":
        return {
            "text": "#0f172a",
            "muted": "#475569",
            "subtle": "#64748b",
            "icon": "#334155",
            "icon_on_solid": "#ffffff",
            "accent": acc,
        }
    if mode == "black":
        return {
            "text": "#f2f2f2",
            "muted": "#a0a0a0",
            "subtle": "#808080",
            "icon": "#d8dee9",
            "icon_on_solid": "#ffffff",
            "accent": acc,
        }
    return {
        "text": "#e6e8ee",
        "muted": "#9aa0a9",
        "subtle": "#69707a",
        "icon": "#d8dee9",
        "icon_on_solid": "#ffffff",
        "accent": acc,
    }


class _ChatBubble(QFrame):
    def __init__(self, bg: str, radius: int = 18, parent: QWidget | None = None) -> None:
        super().__init__(parent)
        self._bg = QColor(bg)
        self._radius = radius
        self.setAttribute(Qt.WA_StyledBackground, True)
        self.setAutoFillBackground(False)

    def set_bg(self, bg: str) -> None:
        self._bg = QColor(bg)
        self.update()

    def set_radius(self, radius: int) -> None:
        self._radius = int(radius)
        self.update()

    def paintEvent(self, event) -> None:  # noqa: N802
        painter = QPainter(self)
        painter.setRenderHint(QPainter.Antialiasing, True)
        rect = QRectF(self.rect()).adjusted(0.5, 0.5, -0.5, -0.5)
        radius = min(self._radius, rect.width() / 2.0, rect.height() / 2.0)
        painter.setPen(Qt.NoPen)
        painter.setBrush(self._bg)
        painter.drawRoundedRect(rect, radius, radius)

@dataclass
class MediaJob:
    job_id: str
    kind: str
    title: str
    source: str
    dest: str
    fmt: str
    status: str = "queued"
    progress: float = 0.0
    current_file: str = ""
    message: str = ""
    profile: str = ""
    created_at: QDateTime = field(default_factory=QDateTime.currentDateTime)
    finished_at: QDateTime | None = None
    worker: QObject | None = None
    thread: QThread | None = None
    settings: dict = field(default_factory=dict)
    file_entries: dict = field(default_factory=dict)
    file_order: list = field(default_factory=list)


def _load_oriented_pixmap(path: str, *, max_size: QSize, smooth: bool = True) -> QPixmap | None:
    """
    Load an image with EXIF orientation applied (fixes sideways/rotated thumbnails).
    Uses QImageReader.setAutoTransform(True) (Qt handles EXIF orientation).
    """
    p = (path or "").strip()
    if not p:
        return None
    try:
        reader = QImageReader(p)
        reader.setAutoTransform(True)
        img = reader.read()
        if img.isNull():
            return None
        pm = QPixmap.fromImage(img)
        if pm.isNull():
            return None
        mode = Qt.SmoothTransformation if smooth else Qt.FastTransformation
        return pm.scaled(max_size, Qt.KeepAspectRatio, mode)
    except Exception:
        return None


class _QtNavItem(QFrame):
    """
    Custom navigation row for precise icon↔label spacing (matches the reference sidebar style).
    """

    def __init__(self, *, icon_pm: QPixmap, text: str, key: str, on_click):
        super().__init__()
        self.setObjectName("NavButton")
        self._key = key
        self._on_click = on_click
        self._collapsed = False
        self._base_icon_pm = QPixmap(icon_pm) if isinstance(icon_pm, QPixmap) else QPixmap()

        self.setCursor(Qt.PointingHandCursor)
        self.setProperty("active", False)
        self.setToolTip(text)
        self.setMinimumHeight(44)
        self.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)

        layout = QHBoxLayout(self)
        layout.setContentsMargins(12, 10, 12, 10)
        layout.setSpacing(18)

        self.icon = QLabel()
        self.icon.setFixedSize(22, 22)
        try:
            self.icon.setPixmap(icon_pm)
        except Exception:
            pass
        layout.addWidget(self.icon, 0, Qt.AlignVCenter)

        self.text = QLabel(text)
        self.text.setObjectName("NavText")
        layout.addWidget(self.text, 1, Qt.AlignVCenter)
        self._refresh_icon()

    def _theme_mode(self) -> str:
        try:
            app = QApplication.instance()
            if app:
                mode = str(app.property("fg_theme_mode") or "dark").strip().lower()
                if mode in {"dark", "light", "black"}:
                    return mode
        except Exception:
            pass
        return "dark"

    def _accent_color(self) -> str:
        try:
            app = QApplication.instance()
            if app:
                accent = str(app.property("fg_color_theme") or "blue").strip().lower()
                accent_map = {
                    "blue": "#0d6efd",
                    "teal": "#14b8a6",
                    "green": "#22c55e",
                    "orange": "#f59e0b",
                    "rose": "#f43f5e",
                    "violet": "#8b5cf6",
                    "cyan": "#06b6d4",
                }
                if accent in accent_map:
                    return accent_map[accent]
        except Exception:
            pass
        return "#0d6efd"

    def _icon_color(self) -> str:
        if bool(self.property("active")):
            return self._accent_color()
        mode = self._theme_mode()
        if mode == "light":
            return "#334155"
        if mode == "black":
            return "#cfd6e4"
        return "#d8dee9"

    def _tint_pixmap(self, pm: QPixmap, color: str) -> QPixmap:
        if not isinstance(pm, QPixmap) or pm.isNull():
            return QPixmap()
        out = QPixmap(pm.size())
        out.fill(Qt.transparent)
        p = QPainter(out)
        try:
            p.drawPixmap(0, 0, pm)
            p.setCompositionMode(QPainter.CompositionMode_SourceIn)
            p.fillRect(out.rect(), QColor(color))
        finally:
            p.end()
        return out

    def _refresh_icon(self) -> None:
        try:
            if isinstance(self._base_icon_pm, QPixmap) and not self._base_icon_pm.isNull():
                self.icon.setPixmap(self._tint_pixmap(self._base_icon_pm, self._icon_color()))
        except Exception:
            pass

    def mousePressEvent(self, event):  # noqa: N802
        try:
            if callable(self._on_click):
                self._on_click(self._key)
        except Exception:
            pass
        super().mousePressEvent(event)

    def set_active(self, active: bool) -> None:
        self.setProperty("active", bool(active))
        self.style().unpolish(self)
        self.style().polish(self)
        try:
            self.text.style().unpolish(self.text)
            self.text.style().polish(self.text)
        except Exception:
            pass
        self._refresh_icon()

    def refresh_theme(self) -> None:
        self._refresh_icon()

    def set_collapsed(self, collapsed: bool) -> None:
        collapsed = bool(collapsed)
        if self._collapsed == collapsed:
            return
        self._collapsed = collapsed
        try:
            self.text.setVisible(not collapsed)
        except Exception:
            pass
        try:
            lay: QHBoxLayout = self.layout()  # type: ignore[assignment]
            if collapsed:
                lay.setContentsMargins(10, 10, 10, 10)
                lay.setSpacing(0)
                lay.setAlignment(self.icon, Qt.AlignHCenter | Qt.AlignVCenter)
            else:
                lay.setContentsMargins(12, 10, 12, 10)
                lay.setSpacing(18)
        except Exception:
            pass


class _QtDropZone(QFrame):
    """
    Simple drag-and-drop target. Emits dropped local file paths.
    """

    files_dropped = Signal(list)

    def __init__(self, text: str):
        super().__init__()
        self.setAcceptDrops(True)
        self.setObjectName("PageCard")
        lay = QVBoxLayout(self)
        lay.setContentsMargins(16, 14, 16, 14)
        lay.setSpacing(8)
        lbl = QLabel(text)
        lbl.setWordWrap(True)
        lbl.setStyleSheet("color:#9aa0a9;")
        lay.addWidget(lbl)

    def dragEnterEvent(self, event):  # noqa: N802
        try:
            if event.mimeData().hasUrls():
                event.acceptProposedAction()
                return
        except Exception:
            pass
        event.ignore()

    def dropEvent(self, event):  # noqa: N802
        paths: list[str] = []
        try:
            for u in event.mimeData().urls():
                p = u.toLocalFile()
                if p:
                    paths.append(p)
        except Exception:
            paths = []
        if paths:
            try:
                self.files_dropped.emit(paths)
            except Exception:
                pass
        try:
            event.acceptProposedAction()
        except Exception:
            pass


class _QtWorkflowApprovalDialog(QDialog):
    def __init__(self, parent: QWidget, *, payload: dict):
        super().__init__(parent)
        self._payload = payload if isinstance(payload, dict) else {}
        self._output_ref: dict | None = None
        self._document_text: str = ""
        self._preview: QTextEdit | None = None
        self._title_combo: QComboBox | None = None
        self.setModal(True)
        self.setObjectName("Dialog")
        self.setWindowTitle("Workflow Approval")
        self.setStyleSheet(_qt_modern_dialog_stylesheet())
        self.resize(860, 640)

        root = QVBoxLayout(self)
        root.setContentsMargins(16, 14, 16, 14)
        root.setSpacing(12)

        header = QFrame()
        header.setObjectName("DialogHeader")
        h = QVBoxLayout(header)
        h.setContentsMargins(16, 12, 16, 12)
        h.setSpacing(6)
        title = QLabel(str(self._payload.get("stage_name") or "Approval"))
        title.setObjectName("DialogTitle")
        h.addWidget(title)
        subtitle = QLabel(str(self._payload.get("message") or "Review output before continuing."))
        subtitle.setObjectName("DialogSubtitle")
        subtitle.setWordWrap(True)
        h.addWidget(subtitle)
        root.addWidget(header)

        preview = QTextEdit()
        preview.setReadOnly(True)
        preview.setObjectName("ResultsTree")
        self._preview = preview
        text = ""
        title_options: list[str] = []
        recommended_index = 0
        try:
            data = self._payload.get("output") or {}
            if isinstance(data, dict):
                self._output_ref = data
            pipeline_id = str(self._payload.get("pipeline_id") or "")
            hide_verification_report = pipeline_id == "family_history_project"
            if isinstance(data, dict):
                self._document_text = str(
                    data.get("improved_text")
                    or data.get("document")
                    or data.get("text")
                    or data.get("summary")
                    or ""
                )
                text = self._document_text
                verification_report = data.get("verification_report")
                if verification_report and not hide_verification_report:
                    text = f"{verification_report}\n\n{text}".strip()
                issues = data.get("issues")
                if issues and isinstance(issues, list):
                    issue_lines = "\n".join(f"- {x}" for x in issues if x)
                    if issue_lines:
                        text = f"Issues:\n{issue_lines}\n\n{text}".strip()
                if not text:
                    try:
                        import json

                        text = json.dumps(data, indent=2)
                    except Exception:
                        text = str(data)
            else:
                text = str(data)
            raw_options = (data.get("title_options") if isinstance(data, dict) else None) or []
            if isinstance(raw_options, list):
                title_options = [str(x).strip() for x in raw_options if str(x).strip()]
            try:
                recommended_index = max(0, int((data.get("recommended_index") if isinstance(data, dict) else 1) or 1) - 1)
            except Exception:
                recommended_index = 0
        except Exception:
            text = ""
        preview.setPlainText(str(text or "No preview available."))
        root.addWidget(preview, 1)

        if title_options and self._document_text:
            card = QFrame()
            card.setObjectName("Card")
            card_lay = QVBoxLayout(card)
            card_lay.setContentsMargins(12, 10, 12, 10)
            card_lay.setSpacing(8)

            title_lbl = QLabel("Choose Final Title")
            title_lbl.setObjectName("DialogSummary")
            card_lay.addWidget(title_lbl)

            hint_lbl = QLabel("Select one of the 5 title options for the final export.")
            hint_lbl.setWordWrap(True)
            hint_lbl.setStyleSheet("color:#9aa0a9;")
            card_lay.addWidget(hint_lbl)

            combo = QComboBox()
            combo.addItems(title_options)
            combo.setCurrentIndex(max(0, min(recommended_index, len(title_options) - 1)))
            combo.currentTextChanged.connect(self._on_title_changed)
            self._title_combo = combo
            card_lay.addWidget(combo)
            root.addWidget(card)

            self._on_title_changed(combo.currentText())

        buttons = QDialogButtonBox()
        btn_continue = QPushButton("Continue")
        btn_continue.setObjectName("PrimaryButton")
        btn_cancel = QPushButton("Cancel Pipeline")
        btn_cancel.setObjectName("SecondaryButton")
        buttons.addButton(btn_continue, QDialogButtonBox.AcceptRole)
        buttons.addButton(btn_cancel, QDialogButtonBox.RejectRole)
        btn_continue.clicked.connect(self._on_continue)
        btn_cancel.clicked.connect(self.reject)
        root.addWidget(buttons)

    @staticmethod
    def _apply_title_to_markdown(document: str, title: str) -> str:
        text = str(document or "")
        chosen = str(title or "").strip()
        if not chosen:
            return text
        lines = text.splitlines()
        for idx, line in enumerate(lines):
            if re.match(r"^\s*#\s+.+$", line):
                lines[idx] = f"# {chosen}"
                return "\n".join(lines).strip() + "\n"
        body = text.strip()
        if body:
            return f"# {chosen}\n\n{body}\n"
        return f"# {chosen}\n"

    def _on_title_changed(self, title: str) -> None:
        if not self._preview:
            return
        if not self._document_text:
            return
        selected = str(title or "").strip()
        if not selected:
            return
        updated = self._apply_title_to_markdown(self._document_text, selected)
        self._preview.setPlainText(updated)

    def _on_continue(self) -> None:
        try:
            if self._title_combo and isinstance(self._output_ref, dict) and self._document_text:
                selected = str(self._title_combo.currentText() or "").strip()
                if selected:
                    updated = self._apply_title_to_markdown(self._document_text, selected)
                    for key in ("document", "improved_text", "text"):
                        if key in self._output_ref:
                            self._output_ref[key] = updated
                    if not any(k in self._output_ref for k in ("document", "improved_text", "text")):
                        self._output_ref["document"] = updated
                    self._output_ref["selected_title"] = selected
                    self._output_ref["selected_title_index"] = int(self._title_combo.currentIndex()) + 1
        except Exception:
            pass
        self.accept()


class _WorkflowUiSignals(QObject):
    progress = Signal(str, int, dict)
    completed = Signal(object)
    approval_request = Signal(object)


class _QtDeviceTransferSendWorker(QObject):
    progress = Signal(object)
    finished = Signal(object)
    error = Signal(str)

    def __init__(self, *, service, host: str, port: int, access_code: str, paths: list[str]):
        super().__init__()
        self.service = service
        self.host = str(host or "").strip()
        self.port = int(port or 0)
        self.access_code = str(access_code or "").strip()
        self.paths = list(paths or [])

    def run(self):
        try:
            result = self.service.send_paths(
                host=self.host,
                port=self.port,
                access_code=self.access_code,
                paths=self.paths,
                progress=lambda event: self.progress.emit(dict(event or {})),
            )
            self.finished.emit(result)
        except Exception as e:
            self.error.emit(str(e))


class _QtWorkspaceReviewDialog(QDialog):
    """
    Modern review dialog for Workspace actions that would change files.
    """

    def __init__(self, parent: QWidget, *, kind: str, payload: dict):
        super().__init__(parent)
        self.setModal(True)
        self.setObjectName("Dialog")
        self.kind = str(kind or "")
        self.payload = dict(payload or {})
        self.result_payload: dict = {"decision": "skip", "selected": []}

        self.setWindowTitle("Workspace Review")
        self.resize(860, 600)

        root = QVBoxLayout(self)
        root.setContentsMargins(16, 14, 16, 14)
        root.setSpacing(12)

        title = QLabel("Review changes")
        title.setObjectName("DialogTitle")
        root.addWidget(title)

        subtitle = QLabel("")
        subtitle.setStyleSheet("color:#9aa0a9;")
        subtitle.setWordWrap(True)
        root.addWidget(subtitle)

        self.table = QTableWidget(0, 4)
        self.table.setObjectName("ResultsTree")
        self.table.setShowGrid(False)
        self.table.verticalHeader().setVisible(False)
        self.table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.table.setSelectionMode(QAbstractItemView.SingleSelection)
        self.table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.table.horizontalHeader().setStretchLastSection(True)
        self.table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeToContents)
        self.table.horizontalHeader().setSectionResizeMode(1, QHeaderView.Stretch)
        self.table.horizontalHeader().setSectionResizeMode(2, QHeaderView.Stretch)
        self.table.horizontalHeader().setSectionResizeMode(3, QHeaderView.Stretch)
        root.addWidget(self.table, 1)

        btn_row = QHBoxLayout()
        btn_row.setSpacing(10)
        self.btn_all = QPushButton("Select All")
        self.btn_none = QPushButton("Select None")
        self.btn_all.clicked.connect(self._select_all)
        self.btn_none.clicked.connect(self._select_none)
        btn_row.addWidget(self.btn_all)
        btn_row.addWidget(self.btn_none)
        btn_row.addStretch(1)

        btn_cancel = QPushButton("Cancel Workflow")
        btn_cancel.setObjectName("SecondaryButton")
        btn_skip = QPushButton("Skip")
        btn_apply = QPushButton("Apply")
        btn_apply.setObjectName("PrimaryButton")

        btn_cancel.clicked.connect(self._cancel)
        btn_skip.clicked.connect(self._skip)
        btn_apply.clicked.connect(self._apply)

        btn_row.addWidget(btn_cancel)
        btn_row.addWidget(btn_skip)
        btn_row.addWidget(btn_apply)
        root.addLayout(btn_row)

        if self.kind == "rename":
            subtitle.setText("Smart Rename found possible new names. Review and apply the selected renames.")
            self.table.setHorizontalHeaderLabels(["", "File", "Suggested name", "Folder"])
            self._populate_rename(list(self.payload.get("items") or []))
        elif self.kind == "categorize":
            subtitle.setText("Auto-Categorize found placements. Review and apply the selected moves into category folders.")
            self.table.setHorizontalHeaderLabels(["", "File", "Category", "Folder"])
            self._populate_categorize(dict(self.payload.get("items") or {}))
        elif self.kind == "convert_images":
            subtitle.setText("Images → WebP will create converted copies. Select which files to convert.")
            self.table.setHorizontalHeaderLabels(["", "File", "Output", "Folder"])
            self._populate_convert(list(self.payload.get("items") or []), output_label=str(self.payload.get("output") or "webp"))
        elif self.kind == "convert_media":
            subtitle.setText("Media → MP4 will create converted copies. Select which files to convert.")
            self.table.setHorizontalHeaderLabels(["", "File", "Output", "Folder"])
            self._populate_convert(list(self.payload.get("items") or []), output_label=str(self.payload.get("output") or "mp4"))
        elif self.kind == "zip":
            subtitle.setText("Archive (ZIP) will create an archive in the target folder. Select which files to include.")
            self.table.setHorizontalHeaderLabels(["", "File", "In archive", "Folder"])
            self._populate_zip(list(self.payload.get("items") or []))
        elif self.kind == "index":
            # Indexing doesn't change files; show a compact confirmation UI.
            est = self.payload.get("estimated_files")
            subtitle.setText(
                f"Index the selected folder to speed up AI Search. "
                f"Folder: {self.payload.get('folder') or ''}"
                + (f" • Estimated files: {est}" if est is not None else "")
            )
            self.table.setVisible(False)
            self.btn_all.setVisible(False)
            self.btn_none.setVisible(False)
        else:
            subtitle.setText("Review the results.")
            self.table.setHorizontalHeaderLabels(["", "Item", "Info", ""])
        note = str(self.payload.get("note") or "").strip()
        if note:
            subtitle.setText((subtitle.text() + "\n" + note).strip())

    def _add_check(self, row: int, checked: bool = True) -> None:
        cb = QCheckBox()
        cb.setChecked(bool(checked))
        w = QWidget()
        lay = QHBoxLayout(w)
        lay.setContentsMargins(8, 0, 8, 0)
        lay.addWidget(cb, 0, Qt.AlignCenter)
        self.table.setCellWidget(row, 0, w)

    def _populate_rename(self, items: list[dict]) -> None:
        self.table.setRowCount(0)
        for it in items:
            path = str(it.get("path") or "")
            sug = str(it.get("suggested") or "")
            try:
                p = Path(path)
                folder = str(p.parent)
                name = p.name
            except Exception:
                folder = ""
                name = path
            r = self.table.rowCount()
            self.table.insertRow(r)
            self._add_check(r, True)
            it_name = QTableWidgetItem(name)
            it_name.setData(Qt.UserRole, path)
            self.table.setItem(r, 1, it_name)
            self.table.setItem(r, 2, QTableWidgetItem(sug))
            self.table.setItem(r, 3, QTableWidgetItem(folder))

    def _populate_categorize(self, items: dict[str, list[str]]) -> None:
        self.table.setRowCount(0)
        for cat_key, paths in items.items():
            for p_str in (paths or []):
                try:
                    p = Path(p_str)
                    folder = str(p.parent)
                    name = p.name
                except Exception:
                    folder = ""
                    name = p_str
                r = self.table.rowCount()
                self.table.insertRow(r)
                self._add_check(r, True)
                it_name = QTableWidgetItem(name)
                it_name.setData(Qt.UserRole, p_str)
                it_cat = QTableWidgetItem(str(cat_key))
                it_cat.setData(Qt.UserRole, str(cat_key))
                self.table.setItem(r, 1, it_name)
                self.table.setItem(r, 2, it_cat)
                self.table.setItem(r, 3, QTableWidgetItem(folder))

    def _populate_convert(self, items: list[dict], *, output_label: str) -> None:
        self.table.setRowCount(0)
        for it in items:
            path = str(it.get("path") or "")
            out_name = str(it.get("output_name") or "")
            try:
                p = Path(path)
                folder = str(p.parent)
                name = p.name
            except Exception:
                folder = ""
                name = path
            r = self.table.rowCount()
            self.table.insertRow(r)
            self._add_check(r, True)
            it_name = QTableWidgetItem(name)
            it_name.setData(Qt.UserRole, path)
            self.table.setItem(r, 1, it_name)
            self.table.setItem(r, 2, QTableWidgetItem(out_name or output_label))
            self.table.setItem(r, 3, QTableWidgetItem(folder))

    def _populate_zip(self, items: list[dict]) -> None:
        self.table.setRowCount(0)
        for it in items:
            path = str(it.get("path") or "")
            arc = str(it.get("arc") or "")
            try:
                p = Path(path)
                folder = str(p.parent)
                name = p.name
            except Exception:
                folder = ""
                name = path
            r = self.table.rowCount()
            self.table.insertRow(r)
            self._add_check(r, True)
            it_name = QTableWidgetItem(name)
            it_name.setData(Qt.UserRole, path)
            self.table.setItem(r, 1, it_name)
            self.table.setItem(r, 2, QTableWidgetItem(arc))
            self.table.setItem(r, 3, QTableWidgetItem(folder))

    def _selected_paths(self) -> list[str]:
        out: list[str] = []
        for r in range(self.table.rowCount()):
            w = self.table.cellWidget(r, 0)
            if not w:
                continue
            cb = w.findChild(QCheckBox)
            if not cb or not cb.isChecked():
                continue
            it = self.table.item(r, 1)
            if it is None:
                continue
            p = it.data(Qt.UserRole)
            if p:
                out.append(str(p))
        return out

    def _select_all(self):
        for r in range(self.table.rowCount()):
            w = self.table.cellWidget(r, 0)
            if w:
                cb = w.findChild(QCheckBox)
                if cb:
                    cb.setChecked(True)

    def _select_none(self):
        for r in range(self.table.rowCount()):
            w = self.table.cellWidget(r, 0)
            if w:
                cb = w.findChild(QCheckBox)
                if cb:
                    cb.setChecked(False)

    def _apply(self):
        self.result_payload = {"decision": "apply", "selected": self._selected_paths()}
        self.accept()

    def _skip(self):
        self.result_payload = {"decision": "skip", "selected": []}
        self.accept()

    def _cancel(self):
        self.result_payload = {"decision": "cancel", "selected": []}
        self.reject()


class _QtHeaderSvgLogo(QWidget):
    """
    Render an SVG in the header without stretching and with a small internal padding to avoid stroke clipping.
    QSvgWidget renders into the full widget rect (can look stretched and can clip strokes if the SVG viewBox is tight).
    """

    def __init__(self, svg_path: Path, *, padding: int = 2, y_offset: int = 0):
        super().__init__()
        self._renderer = None
        self._svg_path = Path(svg_path) if svg_path else None
        self._base_svg_bytes: bytes = b""
        self._theme_mode = "dark"
        self._accent = "blue"
        try:
            if svg_path and Path(svg_path).exists():
                self._base_svg_bytes = Path(svg_path).read_bytes()
        except Exception:
            self._base_svg_bytes = b""
        try:
            self.set_theme(theme="dark", accent="blue")
        except Exception:
            self._renderer = QSvgRenderer(str(svg_path)) if svg_path and Path(svg_path).exists() else None
        self._padding = int(padding)
        self._y_offset = int(y_offset)
        self.setAttribute(Qt.WA_TranslucentBackground, True)
        self.setSizePolicy(QSizePolicy.Fixed, QSizePolicy.Fixed)

    @staticmethod
    def _color_hex(color: QColor) -> str:
        return f"#{color.red():02X}{color.green():02X}{color.blue():02X}"

    @staticmethod
    def _with_lighter(color_hex: str, factor: int) -> str:
        c = QColor(str(color_hex))
        if not c.isValid():
            c = QColor("#0D6EFD")
        return _QtHeaderSvgLogo._color_hex(c.lighter(int(factor)))

    @staticmethod
    def _with_darker(color_hex: str, factor: int) -> str:
        c = QColor(str(color_hex))
        if not c.isValid():
            c = QColor("#0D6EFD")
        return _QtHeaderSvgLogo._color_hex(c.darker(int(factor)))

    @staticmethod
    def _accent_color(accent: str) -> str:
        return {
            "blue": "#0D6EFD",
            "teal": "#14B8A6",
            "green": "#22C55E",
            "orange": "#F59E0B",
            "rose": "#F43F5E",
            "violet": "#8B5CF6",
            "cyan": "#06B6D4",
        }.get(str(accent or "blue").strip().lower(), "#0D6EFD")

    @staticmethod
    def _palette_for(theme: str, accent: str) -> dict[str, str]:
        mode = str(theme or "dark").strip().lower()
        ac = _QtHeaderSvgLogo._accent_color(accent)
        ac_soft = _QtHeaderSvgLogo._with_lighter(ac, 128)
        ac_soft2 = _QtHeaderSvgLogo._with_lighter(ac, 156)
        ac_bright = _QtHeaderSvgLogo._with_lighter(ac, 176)
        ac_dark = _QtHeaderSvgLogo._with_darker(ac, 124)
        ac_deep = _QtHeaderSvgLogo._with_darker(ac, 168)
        ac_navy = _QtHeaderSvgLogo._with_darker(ac, 230)

        if mode == "light":
            txt_silver_1 = "#334155"
            txt_silver_2 = "#475569"
            txt_silver_3 = "#64748B"
            subtitle = "#64748B"
            return {
                "#2E73FF": ac_dark,
                "#38B6FF": ac,
                "#7BE3FF": ac_soft,
                "#E9F2FF": txt_silver_1,
                "#B7C8DC": txt_silver_2,
                "#8FA3BF": txt_silver_3,
                "#1E5BFF": ac_dark,
                "#2E5BB8": _QtHeaderSvgLogo._with_darker(ac, 150),
                "#1A3D7A": ac_deep,
                "#0D1F3D": ac_navy,
                "#6FB8FF": ac_soft2,
                "#3A7FD5": ac,
                "#1E4A8C": ac_dark,
                "#4A9EFF": ac_soft2,
                "#5BA3FF": ac_soft,
                "#7AB8FF": ac_soft2,
                "#5AA8FF": ac_soft,
                "#8FB1D9": subtitle,
                "#1E5BB8": ac_dark,
            }

        if mode == "black":
            return {
                "#2E73FF": ac,
                "#38B6FF": ac_soft,
                "#7BE3FF": ac_bright,
                "#E9F2FF": "#F1F5F9",
                "#B7C8DC": "#CBD5E1",
                "#8FA3BF": "#94A3B8",
                "#1E5BFF": ac,
                "#2E5BB8": ac_dark,
                "#1A3D7A": ac_deep,
                "#0D1F3D": ac_navy,
                "#6FB8FF": ac_soft2,
                "#3A7FD5": ac_soft,
                "#1E4A8C": ac_dark,
                "#4A9EFF": ac_soft2,
                "#5BA3FF": ac_soft,
                "#7AB8FF": ac_soft2,
                "#5AA8FF": ac_soft,
                "#8FB1D9": "#A7B9CF",
                "#1E5BB8": ac_dark,
            }

        # dark
        return {
            "#2E73FF": ac,
            "#38B6FF": ac_soft,
            "#7BE3FF": ac_bright,
            "#E9F2FF": "#E9F2FF",
            "#B7C8DC": "#B7C8DC",
            "#8FA3BF": "#8FA3BF",
            "#1E5BFF": ac_dark,
            "#2E5BB8": ac_dark,
            "#1A3D7A": ac_deep,
            "#0D1F3D": ac_navy,
            "#6FB8FF": ac_soft2,
            "#3A7FD5": ac_soft,
            "#1E4A8C": ac_dark,
            "#4A9EFF": ac_soft2,
            "#5BA3FF": ac_soft,
            "#7AB8FF": ac_soft2,
            "#5AA8FF": ac_soft,
            "#8FB1D9": "#8FB1D9",
            "#1E5BB8": ac_dark,
        }

    @staticmethod
    def _replace_svg_colors(svg_bytes: bytes, mapping: dict[str, str]) -> bytes:
        try:
            txt = svg_bytes.decode("utf-8", errors="ignore")
            out = txt
            for src, dst in (mapping or {}).items():
                out = re.sub(re.escape(str(src)), str(dst), out, flags=re.IGNORECASE)
            return out.encode("utf-8")
        except Exception:
            return svg_bytes

    def set_theme(self, *, theme: str, accent: str) -> None:
        self._theme_mode = str(theme or "dark").strip().lower()
        self._accent = str(accent or "blue").strip().lower()
        raw = bytes(self._base_svg_bytes or b"")
        if not raw and self._svg_path and self._svg_path.exists():
            try:
                raw = self._svg_path.read_bytes()
                self._base_svg_bytes = raw
            except Exception:
                raw = b""
        if not raw:
            return
        mapping = self._palette_for(self._theme_mode, self._accent)
        themed = self._replace_svg_colors(raw, mapping)
        themed = self._expand_svg_viewbox(themed, pad_lr=4.0, pad_top=4.0, pad_bottom=8.0)
        r = QSvgRenderer()
        r.load(QByteArray(themed))
        if r.isValid():
            self._renderer = r
            self.update()

    def has_svg(self) -> bool:
        return self._renderer is not None and self._renderer.isValid()

    def sizeHint(self) -> QSize:  # noqa: N802
        if not self.has_svg():
            return QSize(220, 40)
        w0, h0 = self._source_size()
        if h0 <= 0:
            return QSize(220, 40)
        h = 44
        w = int(round((w0 / h0) * h))
        return QSize(max(220, min(520, w)), h)

    def paintEvent(self, event):  # noqa: N802,ARG002
        if not self.has_svg():
            return
        r = self._renderer
        w0, h0 = self._source_size()
        if w0 <= 0 or h0 <= 0:
            return

        pad = self._padding
        rect = self.rect().adjusted(pad, pad, -pad, -pad)
        if rect.width() <= 2 or rect.height() <= 2:
            return

        scale = min(float(rect.width()) / float(w0), float(rect.height()) / float(h0))
        w = float(w0) * scale
        h = float(h0) * scale

        x = float(rect.x()) + (float(rect.width()) - w) / 2.0
        y = float(rect.y()) + (float(rect.height()) - h) / 2.0 + float(self._y_offset)

        painter = QPainter(self)
        painter.setRenderHint(QPainter.Antialiasing, True)
        painter.setRenderHint(QPainter.SmoothPixmapTransform, True)
        try:
            r.render(painter, QRectF(x, y, w, h))
        finally:
            painter.end()

    def _source_size(self) -> tuple[float, float]:
        """
        Prefer the SVG viewBox ratio over defaultSize().
        defaultSize() often reflects width/height attributes (which can differ from viewBox),
        causing subtle stretching when rendering into a constrained header area.
        """
        r = self._renderer
        if r is None:
            return (0.0, 0.0)
        try:
            vb = r.viewBoxF()
            if vb.isValid() and vb.width() > 0 and vb.height() > 0:
                return (float(vb.width()), float(vb.height()))
        except Exception:
            pass
        try:
            ds = r.defaultSize()
            if ds.isValid() and ds.width() > 0 and ds.height() > 0:
                return (float(ds.width()), float(ds.height()))
        except Exception:
            pass
        return (0.0, 0.0)

    @staticmethod
    def _expand_svg_viewbox(svg_bytes: bytes, *, pad_lr: float, pad_top: float, pad_bottom: float) -> bytes:
        """
        Expand the root <svg viewBox="..."> a bit so strokes near the edges don't get clipped.
        This is especially common with Inkscape exports where strokes extend slightly outside the viewBox.
        """
        try:
            text = svg_bytes.decode("utf-8", errors="ignore")
            m = re.search(r'viewBox="([^"]+)"', text)
            if not m:
                return svg_bytes
            parts = m.group(1).strip().split()
            if len(parts) != 4:
                return svg_bytes
            x, y, w, h = (float(parts[0]), float(parts[1]), float(parts[2]), float(parts[3]))
            x2 = x - pad_lr
            y2 = y - pad_top
            w2 = w + (pad_lr * 2.0)
            h2 = h + (pad_top + pad_bottom)
            repl = f'viewBox="{x2:g} {y2:g} {w2:g} {h2:g}"'
            text2 = re.sub(r'viewBox="[^"]+"', repl, text, count=1)
            return text2.encode("utf-8")
        except Exception:
            return svg_bytes


class FylorraQtMainWindow(QMainWindow):
    def __init__(self, *, backend):
        super().__init__()
        self.backend = backend
        self.setWindowTitle(APP_WINDOW_TITLE)
        self.setMinimumSize(1100, 720)
        self.resize(1300, 780)

        self.icons = QtIconLoader()
        self._wheel_guard_installed = False
        self._tray: QSystemTrayIcon | None = None
        self._force_quit = False
        self._sidebar_collapsed = False
        self._sidebar_expanded_width = 230
        self._sidebar_collapsed_width = 72

        central = QWidget()
        self.setCentralWidget(central)

        root_layout = QVBoxLayout(central)
        root_layout.setContentsMargins(0, 0, 0, 0)
        root_layout.setSpacing(0)

        self.topbar = self._build_topbar()
        root_layout.addWidget(self.topbar)

        body = QFrame()
        body_layout = QHBoxLayout(body)
        body_layout.setContentsMargins(0, 0, 0, 0)
        body_layout.setSpacing(0)

        self.sidebar = self._build_sidebar()
        body_layout.addWidget(self.sidebar)

        self.content_shell = QFrame()
        self.content_shell.setObjectName("ContentShell")
        content_layout = QVBoxLayout(self.content_shell)
        content_layout.setContentsMargins(18, 18, 18, 18)
        content_layout.setSpacing(0)

        # Rounded "content background" container (so corners are actually rounded).
        self.content_card = QFrame()
        self.content_card.setObjectName("ContentCard")
        try:
            self.content_card.setAttribute(Qt.WA_StyledBackground, True)
        except Exception:
            pass
        card_layout = QVBoxLayout(self.content_card)
        card_layout.setContentsMargins(0, 0, 0, 0)
        card_layout.setSpacing(0)

        self.stack = QStackedWidget()
        try:
            # Let the rounded ContentCard background show through.
            self.stack.setAttribute(Qt.WA_StyledBackground, False)
        except Exception:
            pass
        card_layout.addWidget(self.stack, 1)

        content_layout.addWidget(self.content_card, 1)
        body_layout.addWidget(self.content_shell, 1)

        root_layout.addWidget(body, 1)

        self.pages: list[PageDef] = [
            PageDef("monitors", "Monitors", "Intelligent folder monitoring and automation.", "monitor_folders"),
            PageDef("links", "Links", "Create and manage symbolic link paths.", "folder"),
            PageDef("device_transfer", "Device Transfer", "Send and receive files between trusted devices.", "export"),
            PageDef("ai_rules", "AI Rules", "Natural-language rules and scheduled tasks.", "brain"),
            PageDef("scheduled_tasks", "Scheduled Tasks", "Time-based automation tasks.", "analytics"),
            # Use explicit filename so we don't hit the cached vector icon when a new png is added.
            PageDef("cloud_sync", "Cloud Sync", "Connect OneDrive/Google Drive and sync files.", "cloud.png"),
            PageDef("ai_hub", "AI Hub", "Unified AI operations for folders.", "ai_hub"),
            PageDef("ai_command", "AI Command", "Natural language → workflow plan → run locally.", "ai"),
            PageDef("workflow_automation", "Workflow Automation", "Multi-agent pipelines for research, writing, and export.", "workflow_automation"),
            PageDef("writing_assistant", "Writing Assistant", "Offline writing help for school, work, and email.", "Writing_Assistant.png"),
            PageDef("ai_search", "AI Search", "Local semantic search over indexed files.", "search"),
            PageDef("file_tools", "File Tools", "Convert, ZIP/7z, PDF tools, batch operations.", "file_tools.png"),
            PageDef("media_editors", "Media Editors", "Audio & Video editors (Qt migration).", "media_editors"),
            PageDef("workspace", "Workspace", "Workspace actions and history.", "workspace"),
            PageDef("settings", "Settings", "App settings and downloads.", "settings"),
        ]

        self._nav_buttons: dict[str, QWidget] = {}
        self._page_index: dict[str, int] = {}
        self._monitor_cards: dict[str, QWidget] = {}
        self.signals = MonitorSignals()
        self._wf_ui_signals = _WorkflowUiSignals()
        self._wf_ui_signals.progress.connect(self._wf_on_progress)
        self._wf_ui_signals.completed.connect(self._wf_on_completed)
        self._wf_ui_signals.approval_request.connect(self._wf_on_approval_request)
        self._hook_backend_events()
        self._init_pages()
        self._polish_spinboxes(self)
        self._init_nav()
        self._active_page_key = ""
        self.set_active_page("monitors")
        self._init_task_refresh_timer()
        self._apply_inline_theme_overrides()
        try:
            self._refresh_brand_logo_theme()
            self._refresh_hamburger_icon()
            self._refresh_cloud_sync_theme()
            self._refresh_ai_hub_theme()
            self._ft_apply_home_theme()
            self._wa_apply_theme_styles(refresh_chat=False)
        except Exception:
            pass

        # System tray support and startup behaviors (parity with the Tk app).
        try:
            self._init_system_tray()
        except Exception:
            pass
        try:
            self._apply_startup_behaviors()
        except Exception:
            pass

    def _polish_spinboxes(self, root: QWidget):
        """
        Make QSpinBox/QDoubleSpinBox controls feel modern:
        - Use Up/Down buttons with custom +/− icons via QSS
        - Avoid accidental fast value churn while typing
        """
        try:
            from PySide6.QtWidgets import QAbstractSpinBox, QSpinBox, QDoubleSpinBox

            for sb in root.findChildren(QAbstractSpinBox):
                try:
                    sb.setButtonSymbols(QAbstractSpinBox.UpDownArrows)
                except Exception:
                    pass
                try:
                    sb.setKeyboardTracking(False)
                except Exception:
                    pass
                try:
                    sb.setAccelerated(True)
                except Exception:
                    pass
            # Some spinboxes may be created before parenting; ensure direct attrs too.
            for sb in root.findChildren(QSpinBox):
                try:
                    sb.setButtonSymbols(QAbstractSpinBox.UpDownArrows)
                except Exception:
                    pass
            for sb in root.findChildren(QDoubleSpinBox):
                try:
                    sb.setButtonSymbols(QAbstractSpinBox.UpDownArrows)
                except Exception:
                    pass
        except Exception:
            return

    def _init_system_tray(self):
        if not QSystemTrayIcon.isSystemTrayAvailable():
            return
        if self._tray is not None:
            return

        tray = QSystemTrayIcon(self)
        try:
            tray.setIcon(self.windowIcon())
        except Exception:
            pass
        tray.setToolTip(APP_NAME)

        menu = QMenu()
        act_show = QAction("Show", self)
        act_show.triggered.connect(self._tray_show_main)
        menu.addAction(act_show)

        act_hide = QAction("Hide", self)
        act_hide.triggered.connect(self._tray_hide_main)
        menu.addAction(act_hide)
        menu.addSeparator()

        act_quit = QAction("Quit", self)
        act_quit.triggered.connect(self._tray_quit)
        menu.addAction(act_quit)

        tray.setContextMenu(menu)
        tray.activated.connect(self._on_tray_activated)

        self._tray = tray

    def _tray_show_main(self):
        try:
            self.show()
            self.raise_()
            self.activateWindow()
        except Exception:
            pass

    def _tray_hide_main(self):
        try:
            self.hide()
        except Exception:
            pass

    def _tray_quit(self):
        self._force_quit = True
        try:
            if self._tray:
                self._tray.hide()
        except Exception:
            pass
        try:
            self.close()
        except Exception:
            pass

    def _on_tray_activated(self, reason):
        try:
            if reason == QSystemTrayIcon.Trigger:
                if self.isVisible():
                    self.hide()
                else:
                    self._tray_show_main()
        except Exception:
            pass

    def _settings_minimize_to_tray_enabled(self) -> bool:
        s = getattr(self.backend, "settings_manager", None)
        try:
            return bool(s.get_setting("minimize_to_tray", True)) if s else False
        except Exception:
            return False

    def _apply_startup_behaviors(self):
        s = getattr(self.backend, "settings_manager", None)
        if not s:
            return
        # Startup notification (best-effort)
        try:
            if bool(s.get_setting("show_startup_notification", True)):
                self._init_system_tray()
                if self._tray:
                    self._tray.show()
                    self._tray.showMessage(APP_NAME, f"{APP_NAME} is running.", QSystemTrayIcon.Information, 2500)
        except Exception:
            pass

    def closeEvent(self, event):  # noqa: N802
        if not self._force_quit and self._settings_minimize_to_tray_enabled():
            try:
                self._init_system_tray()
                if self._tray:
                    self._tray.show()
                    self.hide()
                    try:
                        self._tray.showMessage(
                            APP_NAME,
                            "Minimized to tray. Right-click the tray icon to quit.",
                            QSystemTrayIcon.Information,
                            2500,
                        )
                    except Exception:
                        pass
                event.ignore()
                return
            except Exception:
                pass
        try:
            if self._tray:
                self._tray.hide()
        except Exception:
            pass
        super().closeEvent(event)

    def _build_sidebar(self) -> QFrame:
        sidebar = QFrame()
        sidebar.setObjectName("Sidebar")
        sidebar.setFixedWidth(self._sidebar_expanded_width)

        layout = QVBoxLayout(sidebar)
        layout.setContentsMargins(12, 12, 12, 12)
        layout.setSpacing(10)

        # Nav host
        self.nav_host = QFrame()
        nav_layout = QVBoxLayout(self.nav_host)
        nav_layout.setContentsMargins(0, 0, 0, 0)
        nav_layout.setSpacing(6)
        layout.addWidget(self.nav_host, 1)

        return sidebar

    def _build_topbar(self) -> QFrame:
        top = QFrame()
        top.setObjectName("HeaderBar")
        top.setFixedHeight(58)

        layout = QHBoxLayout(top)
        # Left padding tuned so the hamburger lines up with the sidebar icon column.
        # Keep padding tight so the logo has enough room in a 58px header.
        layout.setContentsMargins(16, 6, 18, 6)
        layout.setSpacing(12)

        self._sidebar_menu_btn = QToolButton()
        self._sidebar_menu_btn.setObjectName("HamburgerButton")
        self._sidebar_menu_btn.setCursor(Qt.PointingHandCursor)
        self._sidebar_menu_btn.setIconSize(QSize(22, 22))
        self._sidebar_menu_btn.setToolTip("Collapse/expand sidebar")
        self._sidebar_menu_btn.clicked.connect(self._toggle_sidebar)
        self._sidebar_menu_btn.setFixedSize(40, 40)
        self._refresh_hamburger_icon()
        layout.addWidget(self._sidebar_menu_btn)

        layout.addStretch(1)

        # Brand mark: compact icon plus readable wordmark for the tight header.
        brand = QFrame()
        brand.setSizePolicy(QSizePolicy.Fixed, QSizePolicy.Fixed)
        brand_row = QHBoxLayout(brand)
        brand_row.setContentsMargins(0, 0, 0, 0)
        brand_row.setSpacing(2)

        icon = QLabel()
        icon.setObjectName("BrandIcon")
        icon.setFixedSize(44, 44)
        icon.setAlignment(Qt.AlignCenter)
        icon_path = Path(__file__).resolve().parents[1] / "assets" / "fylorra-icon.svg"
        if icon_path.exists():
            # Crop the icon SVG's internal whitespace so the visible hex sits close to the wordmark.
            pm = _render_svg_pixmap(icon_path, QSize(44, 44), viewbox="38 28 224 224")
            icon.setPixmap(pm if not pm.isNull() else QIcon(str(icon_path)).pixmap(QSize(44, 44)))
        self._brand_icon = icon
        self._brand_logo = icon

        text_col = QFrame()
        text_col.setSizePolicy(QSizePolicy.Fixed, QSizePolicy.Fixed)
        text_lay = QVBoxLayout(text_col)
        text_lay.setContentsMargins(0, 0, 0, 0)
        text_lay.setSpacing(0)

        name = QLabel(APP_NAME)
        name.setObjectName("BrandName")
        name.setFixedHeight(24)
        name.setAlignment(Qt.AlignLeft | Qt.AlignBottom)
        self._brand_name_label = name

        tagline = QLabel("Watch | Route | Verify")
        tagline.setObjectName("BrandTagline")
        tagline.setFixedHeight(13)
        tagline.setAlignment(Qt.AlignLeft | Qt.AlignTop)
        self._brand_tagline_label = tagline

        text_lay.addWidget(name)
        text_lay.addWidget(tagline)

        brand_row.addWidget(icon, 0, Qt.AlignRight | Qt.AlignVCenter)
        brand_row.addWidget(text_col, 0, Qt.AlignRight | Qt.AlignVCenter)
        layout.addWidget(brand, 0, Qt.AlignRight | Qt.AlignVCenter)
        self._refresh_brand_logo_theme()

        return top

    def _tint_pixmap(self, pm: QPixmap, color: str) -> QPixmap:
        if not isinstance(pm, QPixmap) or pm.isNull():
            return QPixmap()
        out = QPixmap(pm.size())
        out.fill(Qt.transparent)
        p = QPainter(out)
        try:
            p.drawPixmap(0, 0, pm)
            p.setCompositionMode(QPainter.CompositionMode_SourceIn)
            p.fillRect(out.rect(), QColor(color))
        finally:
            p.end()
        return out

    def _refresh_hamburger_icon(self) -> None:
        try:
            if not hasattr(self, "_sidebar_menu_btn"):
                return
            t = _ui_theme_tokens()
            mode = _ui_theme_mode()
            color = t["icon"] if mode == "light" else t["icon"]
            pm = self.icons.pixmap("hamburger-menu", 22)
            if pm.isNull():
                self._sidebar_menu_btn.setIcon(self.icons.icon("hamburger-menu"))
                return
            self._sidebar_menu_btn.setIcon(QIcon(self._tint_pixmap(pm, color)))
        except Exception:
            pass

    def _refresh_brand_logo_theme(self) -> None:
        try:
            if not any(hasattr(self, attr) for attr in ("_brand_logo", "_brand_name_label", "_brand_tagline_label")):
                return
            logo = getattr(self, "_brand_logo", None)
            app = QApplication.instance()
            mode = "dark"
            accent = "blue"
            if app:
                mode = str(app.property("fg_theme_mode") or "dark").strip().lower()
                accent = str(app.property("fg_color_theme") or "blue").strip().lower()
            if hasattr(logo, "set_theme"):
                logo.set_theme(theme=mode, accent=accent)  # type: ignore[attr-defined]
            if hasattr(self, "_brand_name_label"):
                name = getattr(self, "_brand_name_label", None)
                if name:
                    if mode == "light":
                        name.setStyleSheet("color:#172033; font:800 21px 'Segoe UI'; letter-spacing:0;")
                    else:
                        name.setStyleSheet("color:#e8f2ff; font:800 21px 'Segoe UI'; letter-spacing:0;")
            if hasattr(self, "_brand_tagline_label"):
                tagline = getattr(self, "_brand_tagline_label", None)
                if tagline:
                    if mode == "light":
                        tagline.setStyleSheet("color:#64748b; font:600 8px 'Segoe UI'; letter-spacing:0;")
                    else:
                        tagline.setStyleSheet("color:#9fb6d6; font:600 8px 'Segoe UI'; letter-spacing:0;")
        except Exception:
            pass

    def _cloud_icon(self, stem: str, *, size: int = 18, color: str | None = None) -> QIcon:
        try:
            icon_dir = Path(__file__).resolve().parents[1] / "assets" / "icons" / "Cloud_Sync"
            p = icon_dir / f"{str(stem or '').strip()}.png"
            if not p.exists():
                return QIcon()
            pm = QPixmap(str(p))
            if pm.isNull():
                return QIcon(str(p))
            if color:
                pm = self._tint_pixmap(pm, str(color))
            if int(size) > 0:
                pm = pm.scaled(int(size), int(size), Qt.KeepAspectRatio, Qt.SmoothTransformation)
            return QIcon(pm)
        except Exception:
            return QIcon()

    def _refresh_cloud_sync_theme(self) -> None:
        t = _ui_theme_tokens()
        try:
            for w in getattr(self, "_cloud_muted_widgets", []):
                try:
                    if isinstance(w, QToolButton):
                        w.setStyleSheet(f"color:{t['muted']};")
                    else:
                        w.setStyleSheet(f"color:{t['muted']};")
                except Exception:
                    pass
        except Exception:
            pass
        try:
            if hasattr(self, "_cloud_splitter"):
                if _ui_theme_mode() == "light":
                    self._cloud_splitter.setStyleSheet(
                        "QSplitter::handle{background:#d8e0ec; border:0px;}"
                    )
                else:
                    self._cloud_splitter.setStyleSheet(
                        "QSplitter::handle{background:#232730; border:0px;}"
                    )
        except Exception:
            pass

    def _refresh_ai_hub_theme(self) -> None:
        t = _ui_theme_tokens()
        try:
            if hasattr(self, "_ai_hub_ops_title"):
                self._ai_hub_ops_title.setStyleSheet(f"font-size:14px; font-weight:700; color:{t['text']};")
            if hasattr(self, "_ai_hub_ops_sub"):
                self._ai_hub_ops_sub.setStyleSheet(f"color:{t['muted']};")
        except Exception:
            pass
        try:
            for c in getattr(self, "_ai_hub_ops", {}).values():
                try:
                    if hasattr(c, "refresh_theme"):
                        c.refresh_theme()  # type: ignore[attr-defined]
                except Exception:
                    pass
        except Exception:
            pass

    def _apply_inline_theme_overrides(self) -> None:
        mode = _ui_theme_mode()
        light_map = {
            "#e6e8ee": "#0f172a",
            "#c8ccd6": "#1f2937",
            "#9aa0a9": "#475569",
            "#69707a": "#64748b",
            "#14171c": "#f8fbff",
            "#171a1f": "#ffffff",
            "#101318": "#ffffff",
            "#0f1217": "#f8fbff",
            "#232730": "#cfd7e4",
            "#2a303a": "#c5cedb",
            "#151a21": "#f3f6fc",
            "#313847": "#dfe6f2",
            "#232834": "#e8edf6",
            "#1b2433": "#e8eef8",
            "#1f3a5c": "#dbeafe",
        }

        def _map_style(base_ss: str) -> str:
            if mode != "light":
                return str(base_ss or "")
            out = str(base_ss or "")
            for src, dst in light_map.items():
                out = re.sub(re.escape(src), str(dst), out, flags=re.IGNORECASE)
            return out

        widgets: list[QWidget] = [self]
        try:
            widgets.extend(self.findChildren(QWidget))
        except Exception:
            pass

        for w in widgets:
            try:
                cur = str(w.styleSheet() or "")
                base_prop = w.property("fg_base_stylesheet")
                if base_prop is None:
                    w.setProperty("fg_base_stylesheet", cur)
                    base = cur
                else:
                    base = str(base_prop or "")
                mapped = _map_style(base)
                if mapped != cur:
                    w.setStyleSheet(mapped)
            except Exception:
                continue
        try:
            for b in getattr(self, "_cloud_icon_buttons", []):
                try:
                    stem = str(b.property("fg_cloud_icon_stem") or "").strip()
                    if not stem:
                        continue
                    primary = bool(b.property("fg_cloud_icon_primary"))
                    color = t["icon_on_solid"] if primary else t["icon"]
                    ic = self._cloud_icon(stem, size=18, color=color)
                    if not ic.isNull():
                        b.setIcon(ic)
                        b.setIconSize(QSize(18, 18))
                except Exception:
                    pass
        except Exception:
            pass

    def _init_pages(self) -> None:
        for p in self.pages:
            if p.key == "monitors":
                w = self._build_monitors_page()
            elif p.key == "links":
                w = self._build_links_page()
            elif p.key == "device_transfer":
                w = self._build_device_transfer_page()
            elif p.key == "scheduled_tasks":
                w = self._build_scheduled_tasks_page()
            elif p.key == "ai_rules":
                w = self._build_ai_rules_page()
            elif p.key == "cloud_sync":
                w = self._build_cloud_sync_page()
            elif p.key == "ai_hub":
                w = self._build_ai_hub_page()
            elif p.key == "ai_command":
                w = self._build_ai_command_page()
            elif p.key == "workflow_automation":
                w = self._build_workflow_automation_page()
            elif p.key == "writing_assistant":
                w = self._build_writing_assistant_page()
            elif p.key == "ai_search":
                w = self._build_ai_search_page()
            elif p.key == "file_tools":
                w = self._build_file_tools_page()
            elif p.key == "settings":
                w = self._build_settings_page()
            elif p.key == "workspace":
                w = self._build_workspace_page()
            elif p.key == "media_editors":
                w = self._build_media_editors_page()
            else:
                w = build_placeholder_page(p.title, p.subtitle)
            idx = self.stack.addWidget(w)
            self._page_index[p.key] = idx

    def _init_nav(self) -> None:
        nav_layout: QVBoxLayout = self.nav_host.layout()  # type: ignore[assignment]
        if nav_layout is None:
            nav_layout = QVBoxLayout(self.nav_host)

        for p in self.pages:
            item = _QtNavItem(
                icon_pm=self.icons.pixmap(p.icon, 22),
                text=p.title,
                key=p.key,
                on_click=self.set_active_page,
            )
            nav_layout.addWidget(item)
            self._nav_buttons[p.key] = item

        nav_layout.addStretch(1)

    def _toggle_sidebar(self) -> None:
        self._set_sidebar_collapsed(not self._sidebar_collapsed)

    def _set_sidebar_collapsed(self, collapsed: bool) -> None:
        collapsed = bool(collapsed)
        if self._sidebar_collapsed == collapsed:
            return
        self._sidebar_collapsed = collapsed

        try:
            self.sidebar.setFixedWidth(self._sidebar_collapsed_width if collapsed else self._sidebar_expanded_width)
        except Exception:
            pass

        # Nav buttons
        for key, b in self._nav_buttons.items():
            try:
                b.set_collapsed(collapsed)  # type: ignore[attr-defined]
            except Exception:
                pass

        # Keep the hamburger icon direction consistent (optional).
        try:
            self._sidebar_menu_btn.setToolTip("Expand sidebar" if collapsed else "Collapse sidebar")
        except Exception:
            pass

    def set_active_page(self, key: str) -> None:
        idx = self._page_index.get(key)
        if idx is None:
            return
        self.stack.setCurrentIndex(idx)
        self._active_page_key = str(key)
        for k, b in self._nav_buttons.items():
            try:
                b.set_active(k == key)  # type: ignore[attr-defined]
            except Exception:
                b.setProperty("active", k == key)
                b.style().unpolish(b)
                b.style().polish(b)
        # Ensure Scheduled Tasks reflect background executions.
        try:
            if key == "scheduled_tasks":
                self._reload_tasks()
        except Exception:
            pass

    def _init_task_refresh_timer(self):
        """
        Scheduled tasks run in a background thread; refresh the UI when tasks.json changes.
        """
        self._scheduled_tasks_mtime = 0.0
        self._tasks_refresh_timer = QTimer(self)
        self._tasks_refresh_timer.setInterval(1500)

        def _tick():
            try:
                if getattr(self, "_active_page_key", "") != "scheduled_tasks":
                    return
                settings = getattr(self.backend, "settings_manager", None)
                if not settings:
                    return
                p = Path(getattr(settings, "scheduled_tasks_file", ""))
                if not p.exists():
                    return
                m = float(p.stat().st_mtime)
                if m <= float(getattr(self, "_scheduled_tasks_mtime", 0.0)):
                    return
                self._scheduled_tasks_mtime = m
                try:
                    self.backend.monitor_manager.scheduled_tasks.reload()
                except Exception:
                    pass
                self._reload_tasks()
            except Exception:
                return

        self._tasks_refresh_timer.timeout.connect(_tick)
        self._tasks_refresh_timer.start()

    def _on_run(self) -> None:
        # Legacy quick-run bar (removed from header in Qt UI).
        if not hasattr(self, "command"):
            return
        txt = self.command.text().strip()
        if not txt:
            return
        self.set_active_page("ai_command")
        try:
            self._ai_cmd_text.setPlainText(txt)
            self._ai_cmd_text.setFocus()
        except Exception:
            pass
        self.command.clear()

    def _hook_backend_events(self) -> None:
        try:
            self.backend.monitor_manager.add_event_callback(self._on_backend_event)
        except Exception:
            pass
        self.signals.monitor_event.connect(self._on_monitor_event_ui)

    def _on_backend_event(self, monitor_id: str, event_type: str, src_path: str, dest_path=None):
        self.signals.monitor_event.emit(monitor_id, event_type, src_path, dest_path)

    def _on_monitor_event_ui(self, monitor_id: str, event_type: str, src_path: str, dest_path):
        card = self._monitor_cards.get(monitor_id)
        if not card:
            return
        try:
            card._add_event(event_type, src_path)  # type: ignore[attr-defined]
            card._refresh_stats()  # type: ignore[attr-defined]
        except Exception:
            pass

    def _build_monitors_page(self) -> QWidget:
        host = QFrame()
        host.setObjectName("PageHost")
        layout = QVBoxLayout(host)
        layout.setContentsMargins(18, 18, 18, 18)
        layout.setSpacing(14)

        header = QFrame()
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(0, 0, 0, 0)
        header_layout.setSpacing(10)

        title = QLabel("Monitors")
        title.setObjectName("PageTitle")
        header_layout.addWidget(title)

        header_layout.addStretch(1)

        add_folder = QPushButton("Folder Monitor")
        add_folder.setObjectName("PrimaryButton")
        add_folder.setIcon(self.icons.icon("add"))
        add_folder.setIconSize(QSize(18, 18))
        add_folder.clicked.connect(self._add_folder_monitor)
        header_layout.addWidget(add_folder)

        add_ftp = QPushButton("FTP Monitor")
        add_ftp.setObjectName("PrimaryButton")
        add_ftp.setIcon(self.icons.icon("ftp"))
        add_ftp.setIconSize(QSize(18, 18))
        add_ftp.clicked.connect(self._add_ftp_monitor)
        header_layout.addWidget(add_ftp)

        layout.addWidget(header)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QFrame.NoFrame)

        list_host = QFrame()
        list_layout = QVBoxLayout(list_host)
        list_layout.setContentsMargins(0, 0, 0, 0)
        list_layout.setSpacing(10)

        scroll.setWidget(list_host)
        layout.addWidget(scroll, 1)

        self._qt_monitors_list_host = list_host
        self._qt_monitors_list_layout = list_layout

        # Populate from backend
        try:
            loaded = list(self.backend.loaded_monitors or [])
        except Exception:
            loaded = []
        if not loaded:
            try:
                mgr = self.backend.monitor_manager
                if not getattr(mgr, "monitors", None) and not getattr(mgr, "ftp_manager", None):
                    loaded = list(mgr.load_monitors() or [])
                elif not mgr.monitors and not mgr.ftp_manager.ftp_monitors:
                    loaded = list(mgr.load_monitors() or [])
                else:
                    loaded = []
                    for mid, mon in mgr.monitors.items():
                        loaded.append({
                            "type": "folder",
                            "id": str(mid),
                            "path": getattr(mon, "path", ""),
                            "rules": getattr(mon, "rules", []) or [],
                            "auto_start": bool(getattr(mon, "is_running", False)),
                        })
                    for mid, mon in mgr.ftp_manager.ftp_monitors.items():
                        loaded.append({
                            "type": "ftp",
                            "id": str(mid),
                            "host": getattr(mon, "host", ""),
                            "remote_path": getattr(mon, "remote_path", ""),
                            "auto_start": bool(getattr(mon, "is_running", False)),
                        })
            except Exception:
                loaded = []
        auto_start_global = True
        try:
            auto_start_global = bool(self.backend.settings_manager.get_setting("auto_start_monitors", True))  # type: ignore[attr-defined]
        except Exception:
            auto_start_global = True
        for m in loaded:
            mtype = m.get("type")
            if mtype == "folder":
                self._add_monitor_card(m["id"], monitor_kind="folder", rules=m.get("rules") or [])
                if auto_start_global and m.get("auto_start"):
                    try:
                        self.backend.monitor_manager.start_monitor(m["id"])
                    except Exception:
                        pass
                    try:
                        self._monitor_cards[m["id"]]._refresh_stats()  # type: ignore[attr-defined]
                    except Exception:
                        pass
            elif mtype == "ftp":
                self._add_monitor_card(m["id"], monitor_kind="ftp", rules=[])
                if auto_start_global and m.get("auto_start"):
                    try:
                        self.backend.monitor_manager.ftp_manager.start_ftp_monitor(m["id"])
                    except Exception:
                        pass
                    try:
                        self._monitor_cards[m["id"]]._refresh_stats()  # type: ignore[attr-defined]
                    except Exception:
                        pass

        list_layout.addStretch(1)
        return host

    def _build_links_page(self) -> QWidget:
        from core.symlink_manager import SymlinkManager

        self._symlink_manager = SymlinkManager(self.backend.settings_manager)

        host = QFrame()
        host.setObjectName("PageHost")
        layout = QVBoxLayout(host)
        layout.setContentsMargins(18, 18, 18, 18)
        layout.setSpacing(14)

        header = QFrame()
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(0, 0, 0, 0)
        header_layout.setSpacing(10)

        title = QLabel("Links")
        title.setObjectName("PageTitle")
        header_layout.addWidget(title)

        subtitle = QLabel("Create shortcut-like paths that behave like real folders or files.")
        subtitle.setObjectName("PageSubTitle")
        subtitle.setWordWrap(True)
        header_layout.addWidget(subtitle, 1)

        refresh = QPushButton("Refresh")
        refresh.clicked.connect(self._links_refresh_table)
        header_layout.addWidget(refresh)
        layout.addWidget(header)

        create_card = QFrame()
        create_card.setObjectName("PageCard")
        create_layout = QVBoxLayout(create_card)
        create_layout.setContentsMargins(16, 14, 16, 14)
        create_layout.setSpacing(10)

        form = QFormLayout()
        form.setLabelAlignment(Qt.AlignRight)
        form.setVerticalSpacing(10)
        form.setHorizontalSpacing(12)

        target_row = QHBoxLayout()
        target_row.setSpacing(8)
        self._links_target = QLineEdit()
        self._links_target.setPlaceholderText("Existing file or folder to point at")
        target_row.addWidget(self._links_target, 1)
        target_file = QPushButton("File")
        target_file.clicked.connect(lambda: self._links_pick_target(folder=False))
        target_row.addWidget(target_file)
        target_folder = QPushButton("Folder")
        target_folder.clicked.connect(lambda: self._links_pick_target(folder=True))
        target_row.addWidget(target_folder)
        form.addRow("Target:", target_row)

        link_row = QHBoxLayout()
        link_row.setSpacing(8)
        self._links_link = QLineEdit()
        self._links_link.setPlaceholderText("New link path, for example D:/Work/Photos")
        link_row.addWidget(self._links_link, 1)
        link_parent = QPushButton("Parent")
        link_parent.clicked.connect(self._links_pick_link_parent)
        link_row.addWidget(link_parent)
        form.addRow("Link path:", link_row)

        self._links_type = QComboBox()
        self._links_type.addItem("Auto", "auto")
        self._links_type.addItem("File link", "file")
        self._links_type.addItem("Folder link", "directory")
        self._links_type.addItem("Windows junction", "junction")
        form.addRow("Type:", self._links_type)
        create_layout.addLayout(form)

        actions = QHBoxLayout()
        actions.addStretch(1)
        create_btn = QPushButton("Create Link")
        create_btn.setObjectName("PrimaryButton")
        create_btn.setIcon(self.icons.icon("add"))
        create_btn.setIconSize(QSize(18, 18))
        create_btn.clicked.connect(self._links_create)
        actions.addWidget(create_btn)
        create_layout.addLayout(actions)

        layout.addWidget(create_card)

        table_card = QFrame()
        table_card.setObjectName("PageCard")
        table_layout = QVBoxLayout(table_card)
        table_layout.setContentsMargins(16, 14, 16, 14)
        table_layout.setSpacing(10)

        tools = QHBoxLayout()
        tools.addWidget(QLabel("Managed Links"))
        tools.addStretch(1)
        open_btn = QPushButton("Open Link")
        open_btn.clicked.connect(lambda: self._links_open_selected("link"))
        tools.addWidget(open_btn)
        open_target_btn = QPushButton("Open Target")
        open_target_btn.clicked.connect(lambda: self._links_open_selected("target"))
        tools.addWidget(open_target_btn)
        remove_btn = QPushButton("Remove / Forget")
        remove_btn.clicked.connect(self._links_remove_selected)
        tools.addWidget(remove_btn)
        table_layout.addLayout(tools)

        self._links_table = QTableWidget(0, 4)
        self._links_table.setHorizontalHeaderLabels(["Link path", "Target", "Type", "Status"])
        self._links_table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self._links_table.setSelectionMode(QAbstractItemView.SingleSelection)
        self._links_table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self._links_table.verticalHeader().setVisible(False)
        self._links_table.horizontalHeader().setStretchLastSection(False)
        self._links_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.Stretch)
        self._links_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.Stretch)
        self._links_table.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeToContents)
        self._links_table.horizontalHeader().setSectionResizeMode(3, QHeaderView.ResizeToContents)
        table_layout.addWidget(self._links_table, 1)

        layout.addWidget(table_card, 1)
        self._links_refresh_table()
        return host

    def _dt_service(self):
        return self.backend.get_device_transfer_service()

    def _build_device_transfer_page(self) -> QWidget:
        service = self._dt_service()
        status = service.status()
        self._dt_selected_paths: list[str] = []
        self._dt_jobs: list[tuple[QThread, QObject]] = []

        def _dt_text_input(text: str = "", placeholder: str = "") -> QLineEdit:
            edit = QLineEdit(str(text or ""))
            if placeholder:
                edit.setPlaceholderText(str(placeholder))
            edit.setFixedHeight(40)
            edit.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
            return edit

        def _dt_port_input(value: int) -> QLineEdit:
            edit = _dt_text_input(str(int(value or 47832)))
            edit.setValidator(QIntValidator(1, 65535, edit))
            edit.setMinimumWidth(160)
            edit.setMaximumWidth(200)
            edit.setToolTip("Network port, 1-65535.")
            return edit

        def _dt_button(text: str, primary: bool = False, min_width: int = 86) -> QPushButton:
            btn = QPushButton(text)
            btn.setObjectName("PrimaryButton" if primary else "SecondaryButton")
            btn.setFixedHeight(40)
            btn.setMinimumWidth(int(min_width))
            return btn

        host = QFrame()
        host.setObjectName("PageHost")
        layout = QVBoxLayout(host)
        layout.setContentsMargins(18, 18, 18, 18)
        layout.setSpacing(14)

        header = QFrame()
        header_l = QHBoxLayout(header)
        header_l.setContentsMargins(0, 0, 0, 0)
        header_l.setSpacing(10)
        title = QLabel("Device Transfer")
        title.setObjectName("PageTitle")
        header_l.addWidget(title)
        subtitle = QLabel("Move files directly between trusted Fylorra devices on LAN, VPN, or a reachable remote address.")
        subtitle.setObjectName("PageSubTitle")
        subtitle.setWordWrap(True)
        header_l.addWidget(subtitle, 1)
        layout.addWidget(header)

        receiver = QFrame()
        receiver.setObjectName("PageCard")
        receiver_l = QVBoxLayout(receiver)
        receiver_l.setContentsMargins(16, 14, 16, 14)
        receiver_l.setSpacing(10)

        receiver_top = QHBoxLayout()
        receiver_top.addWidget(QLabel("Receive Files"))
        receiver_top.addStretch(1)
        self._dt_receive_status = QLabel("Off")
        self._dt_receive_status.setObjectName("PageSubTitle")
        receiver_top.addWidget(self._dt_receive_status)
        self._dt_start_btn = _dt_button("Start Receiving", primary=True)
        self._dt_start_btn.clicked.connect(self._dt_toggle_receiving)
        receiver_top.addWidget(self._dt_start_btn)
        receiver_l.addLayout(receiver_top)

        form = QFormLayout()
        form.setLabelAlignment(Qt.AlignRight)
        form.setVerticalSpacing(9)
        form.setHorizontalSpacing(12)
        form.setFieldGrowthPolicy(QFormLayout.AllNonFixedFieldsGrow)

        self._dt_device_name = _dt_text_input(str(status.get("device_name") or ""), "Device name shown to other Fylorra PCs")
        form.addRow("This device:", self._dt_device_name)

        inbox_row = QHBoxLayout()
        self._dt_inbox = _dt_text_input(str(status.get("inbox_dir") or ""))
        inbox_row.addWidget(self._dt_inbox, 1)
        inbox_pick = _dt_button("Browse")
        inbox_pick.clicked.connect(self._dt_browse_inbox)
        inbox_row.addWidget(inbox_pick)
        inbox_open = _dt_button("Open")
        inbox_open.clicked.connect(self._dt_open_inbox)
        inbox_row.addWidget(inbox_open)
        form.addRow("Inbox:", inbox_row)

        port_row = QHBoxLayout()
        self._dt_port = _dt_port_input(int(status.get("port") or 47832))
        port_row.addWidget(self._dt_port)
        port_row.addStretch(1)
        form.addRow("Port:", port_row)

        code_row = QHBoxLayout()
        self._dt_access_code = _dt_text_input(str(status.get("access_code") or ""))
        self._dt_access_code.setReadOnly(True)
        code_row.addWidget(self._dt_access_code, 1)
        copy_code = _dt_button("Copy Code")
        copy_code.clicked.connect(self._dt_copy_code)
        code_row.addWidget(copy_code)
        rotate = _dt_button("Rotate")
        rotate.clicked.connect(self._dt_rotate_code)
        code_row.addWidget(rotate)
        form.addRow("Access code:", code_row)

        self._dt_addresses = QLabel("")
        self._dt_addresses.setTextInteractionFlags(Qt.TextSelectableByMouse)
        self._dt_addresses.setWordWrap(True)
        form.addRow("Addresses:", self._dt_addresses)
        receiver_l.addLayout(form)

        steps = QLabel(
            "Receiver: start receiving, then share the address and access code. "
            "Sender: pick a discovered device or enter address/port, add files or folders, then send. "
            "For internet transfers, use a VPN, private tunnel, or port-forwarded address you trust."
        )
        steps.setObjectName("PageSubTitle")
        steps.setWordWrap(True)
        receiver_l.addWidget(steps)
        layout.addWidget(receiver)

        split = QSplitter(Qt.Horizontal)
        split.setChildrenCollapsible(False)

        peers_card = QFrame()
        peers_card.setObjectName("PageCard")
        peers_l = QVBoxLayout(peers_card)
        peers_l.setContentsMargins(16, 14, 16, 14)
        peers_l.setSpacing(10)
        peers_header = QHBoxLayout()
        peers_header.addWidget(QLabel("Discovered Devices"))
        peers_header.addStretch(1)
        refresh = _dt_button("Refresh")
        refresh.clicked.connect(self._dt_refresh)
        peers_header.addWidget(refresh)
        peers_l.addLayout(peers_header)

        self._dt_peers = QTableWidget(0, 4)
        self._dt_peers.setObjectName("ResultsTree")
        self._dt_peers.setHorizontalHeaderLabels(["Device", "Address", "Platform", "Last seen"])
        self._dt_peers.setSelectionBehavior(QAbstractItemView.SelectRows)
        self._dt_peers.setSelectionMode(QAbstractItemView.SingleSelection)
        self._dt_peers.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self._dt_peers.verticalHeader().setVisible(False)
        self._dt_peers.horizontalHeader().setStretchLastSection(False)
        self._dt_peers.horizontalHeader().setSectionResizeMode(0, QHeaderView.Stretch)
        self._dt_peers.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeToContents)
        self._dt_peers.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeToContents)
        self._dt_peers.horizontalHeader().setSectionResizeMode(3, QHeaderView.ResizeToContents)
        self._dt_peers.itemSelectionChanged.connect(self._dt_use_selected_peer)
        peers_l.addWidget(self._dt_peers, 1)
        split.addWidget(peers_card)

        send_card = QFrame()
        send_card.setObjectName("PageCard")
        send_card.setMinimumWidth(540)
        send_card.setMinimumHeight(305)
        send_l = QVBoxLayout(send_card)
        send_l.setContentsMargins(18, 16, 18, 16)
        send_l.setSpacing(14)
        send_l.addWidget(QLabel("Send Files"))

        def _dt_field_label(text: str) -> QLabel:
            label = QLabel(text)
            label.setFixedWidth(92)
            label.setAlignment(Qt.AlignRight | Qt.AlignVCenter)
            return label

        def _dt_field_row(label_text: str, field: QWidget) -> QFrame:
            row = QFrame()
            row.setFixedHeight(44)
            row_l = QHBoxLayout(row)
            row_l.setContentsMargins(0, 0, 0, 0)
            row_l.setSpacing(12)
            row_l.addWidget(_dt_field_label(label_text))
            row_l.addWidget(field, 1)
            return row

        self._dt_remote_host = _dt_text_input("", "IP, hostname, or VPN name")
        send_l.addWidget(_dt_field_row("Address:", self._dt_remote_host))

        port_wrap = QFrame()
        port_wrap_l = QHBoxLayout(port_wrap)
        port_wrap_l.setContentsMargins(0, 0, 0, 0)
        port_wrap_l.setSpacing(0)
        self._dt_remote_port = _dt_port_input(int(status.get("port") or 47832))
        port_wrap_l.addWidget(self._dt_remote_port)
        port_wrap_l.addStretch(1)
        send_l.addWidget(_dt_field_row("Port:", port_wrap))

        self._dt_remote_code = _dt_text_input("", "Access code from receiving device")
        send_l.addWidget(_dt_field_row("Access code:", self._dt_remote_code))

        file_tools = QHBoxLayout()
        file_tools.setContentsMargins(0, 8, 0, 0)
        file_tools.setSpacing(10)
        add_files = _dt_button("Add Files", min_width=90)
        add_files.clicked.connect(self._dt_add_files)
        file_tools.addWidget(add_files)
        add_folder = _dt_button("Add Folder", min_width=106)
        add_folder.clicked.connect(self._dt_add_folder)
        file_tools.addWidget(add_folder)
        clear_files = _dt_button("Clear", min_width=76)
        clear_files.clicked.connect(self._dt_clear_selection)
        file_tools.addWidget(clear_files)
        file_tools.addStretch(1)
        self._dt_send_btn = _dt_button("Send", primary=True, min_width=96)
        self._dt_send_btn.clicked.connect(self._dt_send_selected)
        file_tools.addWidget(self._dt_send_btn)
        send_l.addLayout(file_tools)

        self._dt_files = QListWidget()
        self._dt_files.setObjectName("ResultsTree")
        self._dt_files.setMinimumHeight(170)
        send_l.addWidget(self._dt_files, 1)
        split.addWidget(send_card)
        split.setStretchFactor(0, 1)
        split.setStretchFactor(1, 1)
        layout.addWidget(split, 1)

        activity_card = QFrame()
        activity_card.setObjectName("PageCard")
        activity_l = QVBoxLayout(activity_card)
        activity_l.setContentsMargins(16, 14, 16, 14)
        activity_l.setSpacing(8)
        activity_l.addWidget(QLabel("Transfer History"))
        self._dt_activity = QTextEdit()
        self._dt_activity.setObjectName("ResultsTree")
        self._dt_activity.setReadOnly(True)
        self._dt_activity.setMinimumHeight(120)
        activity_l.addWidget(self._dt_activity)
        layout.addWidget(activity_card)

        self._dt_timer = QTimer(self)
        self._dt_timer.timeout.connect(self._dt_refresh)
        self._dt_timer.start(2500)
        self._dt_refresh()
        return host

    def _dt_refresh(self) -> None:
        try:
            service = self._dt_service()
            status = service.status()
        except Exception as e:
            try:
                self._dt_receive_status.setText(f"Unavailable: {e}")
            except Exception:
                pass
            return
        running = bool(status.get("running"))
        try:
            self._dt_receive_status.setText("Receiving" if running else "Off")
            self._dt_start_btn.setText("Stop Receiving" if running else "Start Receiving")
            self._dt_access_code.setText(str(status.get("access_code") or ""))
            self._dt_addresses.setText(", ".join(status.get("local_addresses") or []))
        except Exception:
            pass

        table = getattr(self, "_dt_peers", None)
        if table is not None:
            peers = service.peers()
            table.setRowCount(0)
            for row, peer in enumerate(peers):
                table.insertRow(row)
                last_seen = max(0, int(time.time() - float(peer.get("last_seen") or time.time())))
                values = [
                    str(peer.get("device_name") or "Unknown device"),
                    f"{peer.get('host')}:{peer.get('port')}",
                    str(peer.get("platform") or ""),
                    f"{last_seen}s ago",
                ]
                for col, value in enumerate(values):
                    item = QTableWidgetItem(value)
                    item.setData(Qt.UserRole, dict(peer))
                    item.setToolTip(value)
                    table.setItem(row, col, item)

        activity = getattr(self, "_dt_activity", None)
        if activity is not None:
            lines = []
            for entry in service.activity()[:80]:
                ts = time.strftime("%H:%M:%S", time.localtime(float(entry.get("time") or time.time())))
                kind = str(entry.get("kind") or "info").upper()
                msg = str(entry.get("message") or "")
                lines.append(f"{ts} | {kind} | {msg}")
            activity.setPlainText("\n".join(lines))

    def _dt_port_value(self, widget, *, default: int = 47832) -> int:
        text_getter = getattr(widget, "text", None)
        raw = str(text_getter() if callable(text_getter) else "").strip()
        try:
            value = int(raw or default)
        except Exception as e:
            raise ValueError("Port must be a number from 1 to 65535.") from e
        if value < 1 or value > 65535:
            raise ValueError("Port must be a number from 1 to 65535.")
        return value

    def _dt_save_receiver_settings(self) -> None:
        service = self._dt_service()
        service.update_config(
            device_name=self._dt_device_name.text().strip() or "This device",
            inbox_dir=self._dt_inbox.text().strip(),
            port=self._dt_port_value(self._dt_port),
        )

    def _dt_toggle_receiving(self) -> None:
        service = self._dt_service()
        try:
            if service.status().get("running"):
                service.stop()
            else:
                self._dt_save_receiver_settings()
                service.start()
        except Exception as e:
            QMessageBox.critical(self, "Device Transfer", str(e))
        self._dt_refresh()

    def _dt_browse_inbox(self) -> None:
        path = QFileDialog.getExistingDirectory(self, "Select transfer inbox", self._dt_inbox.text().strip())
        if path:
            self._dt_inbox.setText(path)
            try:
                self._dt_save_receiver_settings()
            except Exception:
                pass
            self._dt_refresh()

    def _dt_open_inbox(self) -> None:
        path = self._dt_inbox.text().strip()
        if path:
            QDesktopServices.openUrl(QUrl.fromLocalFile(path))

    def _dt_copy_code(self) -> None:
        QApplication.clipboard().setText(self._dt_access_code.text().strip())

    def _dt_rotate_code(self) -> None:
        try:
            self._dt_service().rotate_access_code()
            self._dt_refresh()
        except Exception as e:
            QMessageBox.critical(self, "Device Transfer", str(e))

    def _dt_use_selected_peer(self) -> None:
        table = getattr(self, "_dt_peers", None)
        if table is None or table.currentRow() < 0:
            return
        item = table.item(table.currentRow(), 0)
        peer = item.data(Qt.UserRole) if item else None
        if not isinstance(peer, dict):
            return
        try:
            self._dt_remote_host.setText(str(peer.get("host") or ""))
            self._dt_remote_port.setText(str(int(peer.get("port") or self._dt_port_value(self._dt_remote_port))))
        except Exception:
            pass

    def _dt_add_files(self) -> None:
        files = QFileDialog.getOpenFileNames(self, "Choose files to send")[0]
        self._dt_add_paths(files)

    def _dt_add_folder(self) -> None:
        folder = QFileDialog.getExistingDirectory(self, "Choose folder to send")
        if folder:
            self._dt_add_paths([folder])

    def _dt_add_paths(self, paths: list[str]) -> None:
        changed = False
        for path in paths or []:
            p = str(path or "").strip()
            if p and p not in self._dt_selected_paths:
                self._dt_selected_paths.append(p)
                changed = True
        if changed:
            self._dt_files.clear()
            for p in self._dt_selected_paths:
                item = QListWidgetItem(p)
                item.setToolTip(p)
                self._dt_files.addItem(item)

    def _dt_clear_selection(self) -> None:
        self._dt_selected_paths = []
        self._dt_files.clear()

    def _dt_log_line(self, message: str) -> None:
        try:
            self._dt_service().record_activity("ui", message)
            self._dt_refresh()
        except Exception:
            pass

    def _dt_send_selected(self) -> None:
        if not self._dt_selected_paths:
            QMessageBox.information(self, "Device Transfer", "Add at least one file or folder to send.")
            return
        host = self._dt_remote_host.text().strip()
        try:
            port = self._dt_port_value(self._dt_remote_port)
        except ValueError as e:
            QMessageBox.warning(self, "Device Transfer", str(e))
            return
        code = self._dt_remote_code.text().strip()
        if not host or not code:
            QMessageBox.information(self, "Device Transfer", "Choose a device or enter address, port, and access code.")
            return

        service = self._dt_service()
        worker = _QtDeviceTransferSendWorker(
            service=service,
            host=host,
            port=port,
            access_code=code,
            paths=list(self._dt_selected_paths),
        )
        thread = QThread(self)
        worker.moveToThread(thread)
        thread.started.connect(worker.run)
        self._dt_jobs.append((thread, worker))
        self._dt_send_btn.setEnabled(False)
        self._dt_send_btn.setText("Sending...")

        def on_progress(event):
            ev = dict(event or {})
            name = Path(str(ev.get("file") or "")).name
            if ev.get("event") == "sending":
                self._dt_log_line(f"Sending {int(ev.get('index') or 0)}/{int(ev.get('total') or 0)}: {name}")
            elif ev.get("event") == "sent":
                self._dt_log_line(f"Sent: {name}")
            elif ev.get("event") == "skipped":
                self._dt_log_line(f"Skipped {name}: {ev.get('reason')}")
            elif ev.get("event") == "failed":
                self._dt_log_line(f"Failed {name}: {ev.get('error')}")

        def on_done(result):
            self._dt_send_btn.setEnabled(True)
            self._dt_send_btn.setText("Send")
            res = dict(result or {})
            failures = res.get("failed") or []
            msg = f"Sent {res.get('sent', 0)} of {res.get('total', 0)} item(s). Skipped: {res.get('skipped', 0)}."
            if failures:
                msg += f" Failed: {len(failures)}."
            self._dt_log_line(msg)
            try:
                thread.quit()
                thread.wait(500)
            except Exception:
                pass

        def on_error(message: str):
            self._dt_send_btn.setEnabled(True)
            self._dt_send_btn.setText("Send")
            QMessageBox.critical(self, "Device Transfer", str(message))
            self._dt_log_line(f"Transfer error: {message}")
            try:
                thread.quit()
                thread.wait(500)
            except Exception:
                pass

        worker.progress.connect(on_progress)
        worker.finished.connect(on_done)
        worker.error.connect(on_error)
        worker.finished.connect(worker.deleteLater)
        worker.error.connect(worker.deleteLater)
        thread.finished.connect(thread.deleteLater)
        thread.start()

    def _links_pick_target(self, *, folder: bool) -> None:
        if folder:
            path = QFileDialog.getExistingDirectory(self, "Select target folder")
        else:
            path = QFileDialog.getOpenFileName(self, "Select target file")[0]
        if not path:
            return
        self._links_target.setText(path)
        if not self._links_link.text().strip():
            try:
                self._links_link.setText(str(Path.home() / DEFAULT_LINKS_FOLDER_NAME / Path(path).name))
            except Exception:
                pass

    def _links_pick_link_parent(self) -> None:
        parent = QFileDialog.getExistingDirectory(self, "Select where the link should be created")
        if not parent:
            return
        name = ""
        try:
            target = Path(self._links_target.text().strip())
            name = target.name
        except Exception:
            name = ""
        self._links_link.setText(str(Path(parent) / (name or "New_Link")))

    def _links_create(self) -> None:
        target = self._links_target.text().strip()
        link = self._links_link.text().strip()
        link_type = self._links_type.currentData() or "auto"
        if not target or not link:
            QMessageBox.warning(self, "Create Link", "Choose both a target and a new link path.")
            return
        try:
            self._symlink_manager.create_link(target, link, str(link_type))
        except Exception as e:
            QMessageBox.critical(self, "Create Link Failed", str(e))
            return
        self._links_link.clear()
        self._links_refresh_table()

    def _links_refresh_table(self) -> None:
        table = getattr(self, "_links_table", None)
        manager = getattr(self, "_symlink_manager", None)
        if not table or not manager:
            return
        records = manager.list_links()
        table.setRowCount(0)
        for row, rec in enumerate(records):
            table.insertRow(row)
            values = [
                str(rec.get("link_path") or ""),
                str(rec.get("target_path") or ""),
                str(rec.get("link_type") or ""),
                str(rec.get("status") or ""),
            ]
            for col, value in enumerate(values):
                item = QTableWidgetItem(value)
                item.setToolTip(value)
                item.setData(Qt.UserRole, rec)
                table.setItem(row, col, item)

    def _links_selected_record(self) -> dict | None:
        table = getattr(self, "_links_table", None)
        if not table:
            return None
        row = table.currentRow()
        if row < 0:
            return None
        item = table.item(row, 0)
        if not item:
            return None
        rec = item.data(Qt.UserRole)
        return dict(rec) if isinstance(rec, dict) else None

    def _links_open_selected(self, which: str) -> None:
        rec = self._links_selected_record()
        if not rec:
            return
        path = rec.get("target_path") if which == "target" else rec.get("link_path")
        if not path:
            return
        QDesktopServices.openUrl(QUrl.fromLocalFile(str(path)))

    def _links_remove_selected(self) -> None:
        rec = self._links_selected_record()
        if not rec:
            return
        link = str(rec.get("link_path") or "")
        status = str(rec.get("status") or "")
        if not link:
            return
        if status in {"missing", "broken"}:
            msg = "Forget this saved link record?"
            forget_only = status == "missing"
        else:
            msg = "Remove this link path? The original target will not be deleted."
            forget_only = False
        if QMessageBox.question(self, "Remove Link", msg) != QMessageBox.Yes:
            return
        try:
            self._symlink_manager.remove_link(link, forget_only=forget_only)
        except Exception as e:
            QMessageBox.critical(self, "Remove Link Failed", str(e))
            return
        self._links_refresh_table()

    def _build_scheduled_tasks_page(self) -> QWidget:
        host = QFrame()
        host.setObjectName("PageHost")
        layout = QVBoxLayout(host)
        layout.setContentsMargins(18, 18, 18, 18)
        layout.setSpacing(14)

        header = QFrame()
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(0, 0, 0, 0)
        header_layout.setSpacing(10)

        title = QLabel("Scheduled Tasks")
        title.setObjectName("PageTitle")
        header_layout.addWidget(title)
        header_layout.addStretch(1)

        safe_temp_btn = QPushButton("Safe Temp Cleanup")
        safe_temp_btn.setIcon(self.icons.icon("shield"))
        safe_temp_btn.setIconSize(QSize(18, 18))
        safe_temp_btn.setToolTip("Create a safe scheduled cleanup for old Temp files. Active downloads and fresh files are skipped.")
        safe_temp_btn.clicked.connect(self._add_safe_temp_cleanup_task)
        header_layout.addWidget(safe_temp_btn)

        add_btn = QPushButton("Add Task")
        add_btn.setObjectName("PrimaryButton")
        add_btn.setIcon(self.icons.icon("add"))
        add_btn.setIconSize(QSize(18, 18))
        add_btn.clicked.connect(self._add_scheduled_task)
        header_layout.addWidget(add_btn)

        layout.addWidget(header)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QFrame.NoFrame)

        list_host = QFrame()
        list_layout = QVBoxLayout(list_host)
        list_layout.setContentsMargins(0, 0, 0, 0)
        list_layout.setSpacing(10)

        scroll.setWidget(list_host)
        layout.addWidget(scroll, 1)

        self._qt_tasks_list_host = list_host
        self._qt_tasks_list_layout = list_layout

        self._reload_tasks()
        return host

    def eventFilter(self, obj, event):  # noqa: N802
        # Prevent accidental value changes while scrolling the Settings page.
        if event.type() == QEvent.Wheel:
            try:
                if isinstance(obj, (QSlider, QComboBox, QSpinBox)) and not obj.hasFocus():
                    return True
            except Exception:
                pass
        if event.type() == QEvent.Resize:
            try:
                if hasattr(self, "_wa_chat_area") and obj is self._wa_chat_area.viewport():
                    self._wa_update_chat_widths()
            except Exception:
                pass
        return super().eventFilter(obj, event)

    def _ensure_ai_ready(self, *, title: str = "AI Model Required", kind: str | None = None) -> bool:
        ai = getattr(self.backend, "ai_manager", None)
        if not ai:
            QMessageBox.warning(self, title, "AI is not available in this build.")
            return False
        # Default: most app features use the vision model.
        requested_kind = (kind or "").strip().lower()
        if not requested_kind:
            requested_kind = "vision"
        if requested_kind not in ("vision", "text"):
            requested_kind = "vision"

        if getattr(ai, "is_ready", False):
            try:
                current_kind = ai.get_active_kind()
            except Exception:
                current_kind = "vision" if getattr(ai, "is_vision_model", False) else "text"
            if current_kind == requested_kind:
                return True
        try:
            ai.select_kind(requested_kind)
        except Exception:
            pass
        loader = _QtAIModelLoadDialog(self, ai_manager=ai, kind=requested_kind)
        return loader.exec() == QDialog.Accepted

    def _build_ai_search_page(self) -> QWidget:
        host = QFrame()
        host.setObjectName("PageHost")
        layout = QVBoxLayout(host)
        layout.setContentsMargins(18, 18, 18, 18)
        layout.setSpacing(14)

        header = QFrame()
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(0, 0, 0, 0)
        header_layout.setSpacing(10)

        title = QLabel("AI Search")
        title.setObjectName("PageTitle")
        header_layout.addWidget(title)

        subtitle = QLabel("Index a folder, then search by meaning (not filenames).")
        subtitle.setObjectName("PageSubTitle")
        header_layout.addWidget(subtitle, 1)

        layout.addWidget(header)

        card = QFrame()
        card.setObjectName("PageCard")
        card_layout = QVBoxLayout(card)
        card_layout.setContentsMargins(16, 14, 16, 14)
        card_layout.setSpacing(12)

        # Index controls
        index_row = QHBoxLayout()
        index_row.setSpacing(10)
        self._ai_search_folder = _DropFolderLineEdit()
        self._ai_search_folder.setPlaceholderText("Folder to index…")
        index_row.addWidget(self._ai_search_folder, 1)

        browse = QPushButton("Browse")
        browse.clicked.connect(self._ai_search_browse_folder)
        index_row.addWidget(browse)

        self._ai_search_index_btn = QPushButton("Index Now")
        self._ai_search_index_btn.setObjectName("PrimaryButton")
        self._ai_search_index_btn.setIcon(self.icons.icon("analytics"))
        self._ai_search_index_btn.setIconSize(QSize(18, 18))
        self._ai_search_index_btn.clicked.connect(self._ai_search_start_index)
        index_row.addWidget(self._ai_search_index_btn)

        self._ai_search_cancel_index_btn = QPushButton("Cancel")
        self._ai_search_cancel_index_btn.setEnabled(False)
        self._ai_search_cancel_index_btn.clicked.connect(self._ai_search_cancel_index)
        index_row.addWidget(self._ai_search_cancel_index_btn)

        card_layout.addLayout(index_row)

        opts = QHBoxLayout()
        opts.setSpacing(16)
        self._ai_search_include_sub = QCheckBox("Include subfolders")
        self._ai_search_include_sub.setChecked(True)
        opts.addWidget(self._ai_search_include_sub)

        self._ai_search_ai_summ = QCheckBox("AI summaries")
        opts.addWidget(self._ai_search_ai_summ)

        self._ai_search_ocr = QCheckBox("OCR scanned PDFs")
        opts.addWidget(self._ai_search_ocr)

        self._ai_search_extract_images = QCheckBox("AI vision captions (images)")
        opts.addWidget(self._ai_search_extract_images)

        self._ai_search_hashes = QCheckBox("Compute hashes (<=20MB)")
        opts.addWidget(self._ai_search_hashes)

        opts.addStretch(1)

        try:
            from core.library_index import LibraryIndex

            self._ai_search_fts = QLabel(f"FTS: {'ON' if LibraryIndex().fts_enabled else 'OFF'}")
        except Exception:
            self._ai_search_fts = QLabel("FTS: ?")
        self._ai_search_fts.setStyleSheet("color:#9aa0a9;")
        opts.addWidget(self._ai_search_fts)
        card_layout.addLayout(opts)

        self._ai_search_index_status = QLabel("Ready.")
        self._ai_search_index_status.setStyleSheet("color:#9aa0a9;")
        card_layout.addWidget(self._ai_search_index_status)

        self._ai_search_index_bar = QProgressBar()
        self._ai_search_index_bar.setRange(0, 1000)
        self._ai_search_index_bar.setValue(0)
        card_layout.addWidget(self._ai_search_index_bar)

        # Search controls
        search_box = QGroupBox("Search")
        search_layout = QVBoxLayout(search_box)
        search_layout.setContentsMargins(12, 10, 12, 10)
        search_layout.setSpacing(10)

        qrow = QHBoxLayout()
        qrow.setSpacing(10)
        self._ai_search_query = QLineEdit()
        self._ai_search_query.setPlaceholderText("Search (e.g. “invoices from Energy Company”, “wiring instructions PDF”)")
        self._ai_search_query.returnPressed.connect(self._ai_search_run)
        qrow.addWidget(self._ai_search_query, 1)

        self._ai_search_rerank = QCheckBox("AI rerank")
        self._ai_search_rerank.setChecked(True)
        qrow.addWidget(self._ai_search_rerank)

        self._ai_search_btn = QPushButton("Search")
        self._ai_search_btn.setObjectName("PrimaryButton")
        self._ai_search_btn.setIcon(self.icons.icon("search"))
        self._ai_search_btn.setIconSize(QSize(18, 18))
        self._ai_search_btn.clicked.connect(self._ai_search_run)
        qrow.addWidget(self._ai_search_btn)

        search_layout.addLayout(qrow)

        # Search status / meta (shown during search + after results)
        self._ai_search_search_status = QLabel("Ready.")
        self._ai_search_search_status.setStyleSheet("color:#9aa0a9;")
        search_layout.addWidget(self._ai_search_search_status)

        hint = QLabel("Tip: to search images by content (e.g. “beach”), enable “AI vision captions (images)” and re-index.")
        hint.setStyleSheet("color:#7f8793;")
        search_layout.addWidget(hint)

        splitter = QSplitter(Qt.Vertical)
        splitter.setChildrenCollapsible(False)

        # Results list (card-style widgets)
        res_wrap = QWidget()
        res_l = QVBoxLayout(res_wrap)
        res_l.setContentsMargins(0, 0, 0, 0)
        res_l.setSpacing(8)

        self._ai_search_results = QListWidget()
        self._ai_search_results.itemSelectionChanged.connect(self._ai_search_on_select)
        self._ai_search_results.itemDoubleClicked.connect(lambda _it: self._ai_search_open_selected())
        self._ai_search_results.setVerticalScrollMode(QAbstractItemView.ScrollPerPixel)
        self._ai_search_results.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self._ai_search_results.setSpacing(8)
        self._ai_search_results.setStyleSheet(
            "QListWidget{background:#14171c; border:1px solid #232730; border-radius:12px; padding:10px;}"
            "QListWidget::item{border:0; padding:0; margin:0;}"
            "QListWidget::item:selected{background:transparent;}"
        )
        res_l.addWidget(self._ai_search_results, 1)
        splitter.addWidget(res_wrap)

        # Preview panel
        prev_wrap = QFrame()
        prev_wrap.setStyleSheet("QFrame{background:#14171c; border:1px solid #232730; border-radius:12px;}")
        prev_l = QVBoxLayout(prev_wrap)
        prev_l.setContentsMargins(12, 10, 12, 10)
        prev_l.setSpacing(8)

        prev_head = QHBoxLayout()
        prev_head.setSpacing(10)
        left = QVBoxLayout()
        left.setSpacing(2)
        self._ai_search_prev_title = QLabel("No selection")
        self._ai_search_prev_title.setStyleSheet("color:#d7dbe5; font-weight:700;")
        self._ai_search_prev_path = QLabel("")
        self._ai_search_prev_path.setStyleSheet("color:#9aa0a9;")
        self._ai_search_prev_path.setTextInteractionFlags(Qt.TextSelectableByMouse)
        left.addWidget(self._ai_search_prev_title)
        left.addWidget(self._ai_search_prev_path)
        prev_head.addLayout(left, 1)

        self._ai_search_open_btn = QPushButton("Open")
        self._ai_search_open_btn.clicked.connect(self._ai_search_open_selected)
        prev_head.addWidget(self._ai_search_open_btn)

        self._ai_search_show_btn = QPushButton("Show in Folder")
        self._ai_search_show_btn.clicked.connect(self._ai_search_show_in_folder)
        prev_head.addWidget(self._ai_search_show_btn)

        self._ai_search_explain_btn = QPushButton("Explain Filename")
        self._ai_search_explain_btn.clicked.connect(self._ai_search_explain_selected)
        prev_head.addWidget(self._ai_search_explain_btn)

        prev_l.addLayout(prev_head)

        self._ai_search_preview = QTextBrowser()
        self._ai_search_preview.setOpenExternalLinks(False)
        self._ai_search_preview.setReadOnly(True)
        self._ai_search_preview.setPlaceholderText("Select a result to preview extracted text / AI summary…")
        self._ai_search_preview.setStyleSheet(
            "QTextBrowser{background:transparent; border:0; color:#c8ccd6;}"
        )
        prev_l.addWidget(self._ai_search_preview, 1)
        splitter.addWidget(prev_wrap)

        splitter.setSizes([520, 280])
        search_layout.addWidget(splitter, 1)

        card_layout.addWidget(search_box, 1)

        layout.addWidget(card, 1)
        return host

    def _build_workflow_automation_page(self) -> QWidget:
        host = QFrame()
        host.setObjectName("PageHost")
        layout = QVBoxLayout(host)
        layout.setContentsMargins(18, 18, 18, 18)
        layout.setSpacing(14)

        header = QFrame()
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(0, 0, 0, 0)
        header_layout.setSpacing(10)
        title = QLabel("Workflow Automation")
        title.setObjectName("PageTitle")
        header_layout.addWidget(title)
        sub = QLabel("Multi-agent pipelines for research, writing, validation, and export.")
        sub.setObjectName("PageSubTitle")
        header_layout.addWidget(sub, 1)
        layout.addWidget(header)

        settings = getattr(self.backend, "settings_manager", None)
        wf_settings = {}
        if settings:
            try:
                wf_settings = settings.get_workflow_settings()
            except Exception:
                wf_settings = {}

        body = QSplitter(Qt.Horizontal)
        body.setHandleWidth(6)
        body.setStyleSheet(
            "QSplitter::handle{background:#101318; border:1px solid #232730; border-radius:6px;}"
        )
        self._wf_body_splitter = body

        # Left panel (templates + history)
        left = QFrame()
        left.setObjectName("PageCard")
        left_layout = QVBoxLayout(left)
        left_layout.setContentsMargins(14, 12, 14, 12)
        left_layout.setSpacing(12)

        tmpl_title = QLabel("Templates")
        tmpl_title.setStyleSheet("color:#ffffff; font-weight:700;")
        self._wf_tmpl_title = tmpl_title
        left_layout.addWidget(tmpl_title)
        self._wf_template_list = QListWidget()
        self._wf_template_list.setStyleSheet(
            "QListWidget{background:#14171c; border:1px solid #232730; border-radius:10px;}"
            "QListWidget::item{padding:8px 10px;}"
        )
        self._wf_template_list.itemClicked.connect(self._wf_select_template)
        left_layout.addWidget(self._wf_template_list, 1)

        hist_header = QHBoxLayout()
        hist_label = QLabel("Execution History")
        hist_label.setStyleSheet("color:#ffffff; font-weight:700;")
        self._wf_hist_label = hist_label
        hist_header.addWidget(hist_label)
        hist_header.addStretch(1)
        self._wf_history_delete_btn = QToolButton()
        self._wf_history_delete_btn.setIcon(self.icons.icon("delete"))
        self._wf_history_delete_btn.setIconSize(QSize(16, 16))
        self._wf_history_delete_btn.setToolTip("Delete selected execution")
        self._wf_history_delete_btn.setStyleSheet(
            "QToolButton{background:#1b1f26; border:1px solid #2a303a; border-radius:8px; padding:4px;}"
            "QToolButton:hover{background:#232834;}"
        )
        self._wf_history_delete_btn.clicked.connect(self._wf_delete_selected_history)
        hist_header.addWidget(self._wf_history_delete_btn)
        left_layout.addLayout(hist_header)

        self._wf_history_list = QListWidget()
        self._wf_history_list.setStyleSheet(
            "QListWidget{background:#14171c; border:1px solid #232730; border-radius:10px;}"
            "QListWidget::item{padding:8px 10px;}"
        )
        self._wf_history_list.itemClicked.connect(self._wf_open_history_item)
        left_layout.addWidget(self._wf_history_list, 1)

        body.addWidget(left)

        # Right panel (details + run)
        right = QFrame()
        right.setObjectName("PageCard")
        right_layout = QVBoxLayout(right)
        right_layout.setContentsMargins(16, 14, 16, 14)
        right_layout.setSpacing(12)

        info_row = QHBoxLayout()
        self._wf_detail_title = QLabel("Select a template")
        self._wf_detail_title.setStyleSheet("color:#ffffff; font-weight:800; font-size:16px;")
        info_row.addWidget(self._wf_detail_title)
        info_row.addStretch(1)
        right_layout.addLayout(info_row)

        self._wf_detail_desc = QLabel("")
        self._wf_detail_desc.setWordWrap(True)
        self._wf_detail_desc.setStyleSheet("color:#9aa0a9;")
        right_layout.addWidget(self._wf_detail_desc)

        self._wf_stage_tree = QTreeWidget()
        self._wf_stage_tree.setHeaderLabels(["Stage", "Status"])
        self._wf_stage_tree.setRootIsDecorated(False)
        self._wf_stage_tree.setAlternatingRowColors(False)
        self._wf_stage_tree.setStyleSheet(
            "QTreeWidget{background:#14171c; border:1px solid #232730; border-radius:10px; color:#c8ccd6;}"
            "QHeaderView::section{background:#101318; color:#9aa0a9; border:0; padding:4px 6px;}"
            "QTreeWidget::item{padding:4px 6px;}"
        )
        self._wf_stage_tree.setMinimumHeight(140)
        self._wf_stage_tree.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        self._wf_stage_tree.itemSelectionChanged.connect(self._wf_on_stage_selected)

        form = QGridLayout()
        form.setHorizontalSpacing(12)
        form.setVerticalSpacing(10)
        form.setColumnStretch(1, 1)
        form.setColumnStretch(2, 1)
        form.setColumnStretch(3, 0)

        self._wf_request = QTextEdit()
        self._wf_request.setPlaceholderText("Request or input text...")
        self._wf_request.setMinimumHeight(80)
        self._wf_request.setStyleSheet("background:#101318; border:1px solid #232730; border-radius:10px; padding:8px; color:#e6e9ef;")
        form.addWidget(QLabel("Request / Input"), 0, 0)
        form.addWidget(self._wf_request, 0, 1, 1, 3)

        self._wf_source_drop = _QtDropZone("Drop files here (PDF, DOCX, images, text)")
        self._wf_source_drop.setStyleSheet(
            "QFrame{background:#0f1217; border:1px dashed #2a2f38; border-radius:10px;}"
            "QLabel{color:#7f8793;}"
        )
        self._wf_source_drop.files_dropped.connect(self._wf_on_files_dropped)

        self._wf_source_list = QListWidget()
        self._wf_source_list.setStyleSheet(
            "QListWidget{background:#101318; border:1px solid #232730; border-radius:10px;}"
            "QListWidget::item{padding:6px 8px;}"
        )
        self._wf_source_list.setMinimumHeight(90)

        src_container = QWidget()
        src_container_layout = QVBoxLayout(src_container)
        src_container_layout.setContentsMargins(0, 0, 0, 0)
        src_container_layout.setSpacing(6)

        self._wf_source_drop.setMinimumHeight(52)
        self._wf_source_drop.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)

        drop_row = QHBoxLayout()
        drop_row.setContentsMargins(0, 0, 0, 0)
        drop_row.setSpacing(10)

        src_btn_style = (
            "QToolButton{background:#1b1f26; border:1px solid #2a303a; border-radius:10px; padding:6px;}"
            "QToolButton:hover{background:#232834;}"
        )
        self._wf_add_sources_btn = QToolButton()
        self._wf_add_sources_btn.setIcon(self.icons.icon("add_Flow_Automation"))
        self._wf_add_sources_btn.setToolTip("Add")
        self._wf_add_sources_btn.setIconSize(QSize(24, 24))
        self._wf_add_sources_btn.setFixedSize(42, 42)
        self._wf_add_sources_btn.setStyleSheet(src_btn_style)
        self._wf_add_sources_btn.clicked.connect(self._wf_add_source_files)
        self._wf_remove_sources_btn = QToolButton()
        self._wf_remove_sources_btn.setIcon(self.icons.icon("remove_Flow_Automation"))
        self._wf_remove_sources_btn.setToolTip("Remove")
        self._wf_remove_sources_btn.setIconSize(QSize(24, 24))
        self._wf_remove_sources_btn.setFixedSize(42, 42)
        self._wf_remove_sources_btn.setStyleSheet(src_btn_style)
        self._wf_remove_sources_btn.clicked.connect(self._wf_remove_source_files)

        btn_wrap = QWidget()
        btn_row = QHBoxLayout(btn_wrap)
        btn_row.setContentsMargins(0, 0, 0, 0)
        btn_row.setSpacing(6)
        btn_row.addWidget(self._wf_add_sources_btn)
        btn_row.addWidget(self._wf_remove_sources_btn)

        drop_row.addWidget(self._wf_source_drop, 1)
        drop_row.addWidget(btn_wrap, 0, Qt.AlignVCenter)
        src_container_layout.addLayout(drop_row)

        src_container_layout.addWidget(self._wf_source_list, 1)

        form.addWidget(QLabel("Source files"), 1, 0)
        form.addWidget(src_container, 1, 1, 1, 3)

        self._wf_output_folder = QLineEdit()
        self._wf_output_folder.setPlaceholderText("Output folder...")
        last_folder = str((wf_settings or {}).get("last_output_folder") or "").strip()
        if last_folder:
            self._wf_output_folder.setText(last_folder)
        self._wf_output_folder.editingFinished.connect(self._wf_save_output_folder_field)
        out_browse = QPushButton("Browse")
        out_browse.clicked.connect(self._wf_pick_output_folder)
        form.addWidget(QLabel("Output folder"), 2, 0)
        form.addWidget(self._wf_output_folder, 2, 1, 1, 2)
        form.addWidget(out_browse, 2, 3)

        self._wf_output_name = QLineEdit("workflow_output")
        form.addWidget(QLabel("Output name"), 3, 0)
        form.addWidget(self._wf_output_name, 3, 1)

        self._wf_target_lang = QComboBox()
        self._wf_target_lang.addItems(["English", "Spanish", "French", "German", "Italian", "Portuguese"])
        form.addWidget(QLabel("Target language"), 3, 2)
        form.addWidget(self._wf_target_lang, 3, 3)

        fmt_row = QHBoxLayout()
        fmt_row.setSpacing(10)
        self._wf_fmt_md = QCheckBox("Markdown")
        self._wf_fmt_md.setChecked(True)
        self._wf_fmt_txt = QCheckBox("Text")
        self._wf_fmt_docx = QCheckBox("DOCX")
        self._wf_fmt_pdf = QCheckBox("PDF")
        for w in (self._wf_fmt_md, self._wf_fmt_txt, self._wf_fmt_docx, self._wf_fmt_pdf):
            fmt_row.addWidget(w)
        fmt_row.addStretch(1)
        form.addWidget(QLabel("Export formats"), 4, 0)
        form.addLayout(fmt_row, 4, 1, 1, 3)

        self._wf_model_pref = QComboBox()
        self._wf_model_pref.addItem("Auto (recommended)", "auto")
        self._wf_model_pref.addItem("Text model", "text")
        self._wf_model_pref.addItem("Vision model", "vision")
        try:
            cur_pref = str((wf_settings or {}).get("model_preference", "auto") or "auto").strip().lower()
        except Exception:
            cur_pref = "auto"
        if cur_pref not in {"auto", "text", "vision"}:
            cur_pref = "auto"
        idx = self._wf_model_pref.findData(cur_pref)
        if idx >= 0:
            self._wf_model_pref.setCurrentIndex(idx)
        self._wf_model_pref.setToolTip("Choose which AI model this workflow should use.")
        self._wf_model_pref.currentIndexChanged.connect(lambda _v: self._wf_save_model_preference())
        form.addWidget(QLabel("AI model"), 5, 0)
        form.addWidget(self._wf_model_pref, 5, 1, 1, 3)

        web_row = QHBoxLayout()
        self._wf_allow_web = QCheckBox("Allow web research")
        self._wf_allow_web.setChecked(bool(wf_settings.get("allow_web_research", False)))
        self._wf_allow_web.setToolTip("Allow web research (online) to gather sources.")
        self._wf_web_max = QSpinBox()
        self._wf_web_max.setRange(1, 10)
        self._wf_web_max.setValue(int(wf_settings.get("web_max_results", 5)))
        self._wf_web_max.setEnabled(self._wf_allow_web.isChecked())
        self._wf_web_max.setMinimumWidth(84)
        self._wf_web_max.setMaximumWidth(96)
        self._wf_web_max.setAlignment(Qt.AlignCenter)
        self._wf_web_max.setMinimumHeight(30)
        self._wf_web_max.setStyleSheet(
            "QSpinBox{background:#101318; border:1px solid #232730; border-radius:10px; padding:4px 8px; color:#e6e9ef;}"
            "QSpinBox:disabled{color:#8f96a3;}"
        )
        self._wf_allow_web.stateChanged.connect(lambda _v: self._wf_web_max.setEnabled(self._wf_allow_web.isChecked()))
        self._wf_allow_web.stateChanged.connect(lambda _v: self._wf_save_web_settings())
        self._wf_web_max.valueChanged.connect(lambda _v: self._wf_save_web_settings())
        web_row.addWidget(self._wf_allow_web)
        web_row.addSpacing(14)
        max_label = QLabel("Max sources")
        max_label.setToolTip("How many web sources to pull (1-10).")
        web_row.addWidget(max_label)
        web_row.addWidget(self._wf_web_max)
        web_row.addStretch(1)
        form.addWidget(QLabel("Web research"), 6, 0)
        form.addLayout(web_row, 6, 1, 1, 3)

        right_layout.addLayout(form)

        btn_row = QHBoxLayout()
        btn_row.addStretch(1)
        self._wf_run_btn = QPushButton("Run Pipeline")
        self._wf_run_btn.setObjectName("PrimaryButton")
        self._wf_run_btn.clicked.connect(self._wf_run_pipeline)
        btn_row.addWidget(self._wf_run_btn)
        self._wf_cancel_btn = QPushButton("Cancel")
        self._wf_cancel_btn.setEnabled(False)
        self._wf_cancel_btn.clicked.connect(self._wf_cancel_pipeline)
        btn_row.addWidget(self._wf_cancel_btn)
        right_layout.addLayout(btn_row)

        self._wf_progress = QProgressBar()
        self._wf_progress.setRange(0, 100)
        self._wf_progress.setValue(0)
        right_layout.addWidget(self._wf_progress)

        tabs = QTabWidget()
        tabs.setStyleSheet(
            "QTabWidget::pane{background:#14171c; border:1px solid #232730; border-radius:10px;}"
            "QTabBar::tab{background:#101318; color:#c8ccd6; border:1px solid #232730; "
            "padding:6px 12px; margin-right:6px; border-top-left-radius:8px; border-top-right-radius:8px;}"
            "QTabBar::tab:selected{background:#14171c; color:#ffffff;}"
            "QTabBar::tab:hover{background:#1b1f26;}"
        )
        status_tab = QWidget()
        status_tab.setStyleSheet("background:transparent;")
        status_layout = QVBoxLayout(status_tab)
        status_layout.setContentsMargins(8, 6, 8, 8)
        status_layout.setSpacing(8)
        status_header = QHBoxLayout()
        status_label = QLabel("Pipeline status")
        status_label.setStyleSheet("color:#c8ccd6;")
        self._wf_status_caption = status_label
        status_header.addWidget(status_label)
        status_header.addStretch(1)
        self._wf_status = QLabel("Ready.")
        self._wf_status.setStyleSheet("color:#c8ccd6;")
        status_header.addWidget(self._wf_status)
        status_layout.addLayout(status_header)
        status_layout.addWidget(self._wf_stage_tree, 1)
        tabs.addTab(status_tab, "Status")
        self._wf_output_view = QTextEdit()
        self._wf_output_view.setReadOnly(True)
        self._wf_output_view.setPlaceholderText("Stage output will appear here.")
        tabs.addTab(self._wf_output_view, "Output")
        self._wf_log_view = QTextEdit()
        self._wf_log_view.setReadOnly(True)
        self._wf_log_view.setPlaceholderText("Execution log...")
        tabs.addTab(self._wf_log_view, "Log")
        self._wf_tabs = tabs
        tabs.setCurrentWidget(status_tab)
        right_layout.addWidget(tabs, 1)

        body.addWidget(right)
        body.setStretchFactor(0, 1)
        body.setStretchFactor(1, 2)
        body.setSizes([280, 760])

        layout.addWidget(body, 1)

        # Runtime state
        self._wf_orchestrator = None
        self._wf_templates = []
        self._wf_template_map = {}
        self._wf_current_pipeline = None
        self._wf_current_exec_id = None

        self._wf_load_templates()
        self._wf_load_execution_history()
        self._wf_apply_theme_styles()

        return host

    def _wf_theme_tokens(self) -> dict[str, str]:
        mode = "dark"
        accent = "blue"
        try:
            app = QApplication.instance()
            if app:
                mode = str(app.property("fg_theme_mode") or "dark").strip().lower()
                accent = str(app.property("fg_color_theme") or "blue").strip().lower()
        except Exception:
            pass

        if mode == "light":
            base = {
                "text": "#1f2937",
                "muted": "#5f6b7a",
                "panel": "#fbfcff",
                "panel_alt": "#ffffff",
                "input": "#ffffff",
                "border": "#cfd7e4",
                "header": "#edf2f8",
                "hover": "#e8edf6",
                "disabled": "#98a3b4",
            }
        elif mode == "black":
            base = {
                "text": "#efefef",
                "muted": "#a0a0a0",
                "panel": "#0b0b0b",
                "panel_alt": "#101010",
                "input": "#090909",
                "border": "#1f1f1f",
                "header": "#121212",
                "hover": "#1a1a1a",
                "disabled": "#6f6f6f",
            }
        else:
            base = {
                "text": "#e6e8ee",
                "muted": "#9aa0a9",
                "panel": "#14171c",
                "panel_alt": "#171a1f",
                "input": "#101318",
                "border": "#232730",
                "header": "#101318",
                "hover": "#232834",
                "disabled": "#8f96a3",
            }

        accent_map = {
            "blue": {"main": "#1677ff", "active": "#1f3a5c"},
            "teal": {"main": "#14b8a6", "active": "#1f4d48"},
            "green": {"main": "#22c55e", "active": "#1f4f36"},
            "orange": {"main": "#f59e0b", "active": "#5c4320"},
            "rose": {"main": "#f43f5e", "active": "#5d2332"},
            "violet": {"main": "#8b5cf6", "active": "#40306a"},
            "cyan": {"main": "#06b6d4", "active": "#1d4f5a"},
        }
        acc = accent_map.get(accent, accent_map["blue"])
        base["accent"] = acc["main"]
        base["active"] = acc["active"]
        return base

    def _wf_apply_theme_styles(self):
        if not hasattr(self, "_wf_template_list"):
            return
        t = self._wf_theme_tokens()

        try:
            if hasattr(self, "_wf_body_splitter"):
                self._wf_body_splitter.setStyleSheet(
                    f"QSplitter::handle{{background:{t['header']}; border:1px solid {t['border']}; border-radius:6px;}}"
                )
        except Exception:
            pass

        try:
            if hasattr(self, "_wf_tmpl_title"):
                self._wf_tmpl_title.setStyleSheet(f"color:{t['text']}; font-weight:700;")
            if hasattr(self, "_wf_hist_label"):
                self._wf_hist_label.setStyleSheet(f"color:{t['text']}; font-weight:700;")
            if hasattr(self, "_wf_detail_title"):
                self._wf_detail_title.setStyleSheet(f"color:{t['text']}; font-weight:800; font-size:16px;")
            if hasattr(self, "_wf_detail_desc"):
                self._wf_detail_desc.setStyleSheet(f"color:{t['muted']};")
            if hasattr(self, "_wf_status_caption"):
                self._wf_status_caption.setStyleSheet(f"color:{t['text']};")
            if hasattr(self, "_wf_status"):
                self._wf_status.setStyleSheet(f"color:{t['text']};")
        except Exception:
            pass

        list_style = (
            f"QListWidget{{background:{t['panel']}; border:1px solid {t['border']}; border-radius:10px; color:{t['text']};}}"
            f"QListWidget::item{{padding:8px 10px;}}"
            f"QListWidget::item:selected{{background:{t['hover']}; color:{t['text']};}}"
        )
        src_list_style = (
            f"QListWidget{{background:{t['input']}; border:1px solid {t['border']}; border-radius:10px; color:{t['text']};}}"
            f"QListWidget::item{{padding:6px 8px;}}"
            f"QListWidget::item:selected{{background:{t['hover']}; color:{t['text']};}}"
        )
        tree_style = (
            f"QTreeWidget{{background:{t['panel']}; border:1px solid {t['border']}; border-radius:10px; color:{t['text']};}}"
            f"QHeaderView::section{{background:{t['header']}; color:{t['muted']}; border:0; padding:4px 6px;}}"
            "QTreeWidget::item{padding:4px 6px;}"
            f"QTreeWidget::item:selected{{background:{t['hover']}; color:{t['text']};}}"
        )
        request_style = (
            f"QTextEdit{{background:{t['input']}; border:1px solid {t['border']}; border-radius:10px; padding:8px; color:{t['text']};}}"
        )
        drop_style = (
            f"QFrame{{background:{t['panel_alt']}; border:1px dashed {t['border']}; border-radius:10px;}}"
            f"QLabel{{color:{t['muted']};}}"
        )
        small_btn_style = (
            f"QToolButton{{background:{t['input']}; border:1px solid {t['border']}; border-radius:10px; padding:6px; color:{t['text']};}}"
            f"QToolButton:hover{{background:{t['hover']};}}"
        )
        del_btn_style = (
            f"QToolButton{{background:{t['input']}; border:1px solid {t['border']}; border-radius:8px; padding:4px; color:{t['text']};}}"
            f"QToolButton:hover{{background:{t['hover']};}}"
        )
        spin_style = (
            f"QSpinBox{{background:{t['input']}; border:1px solid {t['border']}; border-radius:10px; padding:4px 8px; color:{t['text']};}}"
            f"QSpinBox:disabled{{color:{t['disabled']};}}"
        )
        tabs_style = (
            f"QTabWidget::pane{{background:{t['panel']}; border:1px solid {t['border']}; border-radius:10px;}}"
            f"QTabBar::tab{{background:{t['header']}; color:{t['text']}; border:1px solid {t['border']}; "
            "padding:6px 12px; margin-right:6px; border-top-left-radius:8px; border-top-right-radius:8px;}"
            f"QTabBar::tab:selected{{background:{t['panel']}; color:{t['text']};}}"
            f"QTabBar::tab:hover{{background:{t['hover']};}}"
        )
        output_style = (
            f"QTextEdit{{background:{t['panel']}; border:1px solid {t['border']}; border-radius:10px; color:{t['text']};}}"
        )

        try:
            self._wf_template_list.setStyleSheet(list_style)
            self._wf_history_list.setStyleSheet(list_style)
            self._wf_stage_tree.setStyleSheet(tree_style)
            self._wf_request.setStyleSheet(request_style)
            self._wf_source_drop.setStyleSheet(drop_style)
            self._wf_source_list.setStyleSheet(src_list_style)
            self._wf_history_delete_btn.setStyleSheet(del_btn_style)
            self._wf_add_sources_btn.setStyleSheet(small_btn_style)
            self._wf_remove_sources_btn.setStyleSheet(small_btn_style)
            self._wf_web_max.setStyleSheet(spin_style)
            if hasattr(self, "_wf_tabs"):
                self._wf_tabs.setStyleSheet(tabs_style)
            if hasattr(self, "_wf_output_view"):
                self._wf_output_view.setStyleSheet(output_style)
            if hasattr(self, "_wf_log_view"):
                self._wf_log_view.setStyleSheet(output_style)
        except Exception:
            pass

    def _wf_load_templates(self) -> None:
        try:
            from core.pipeline.storage import load_all_pipelines
        except Exception:
            load_all_pipelines = None

        self._wf_template_list.clear()
        self._wf_template_map = {}
        self._wf_templates = []

        settings = getattr(self.backend, "settings_manager", None)
        pipelines = []
        if load_all_pipelines:
            try:
                pipelines = load_all_pipelines(settings)
            except Exception:
                try:
                    pipelines = load_all_pipelines(None)
                except Exception:
                    pipelines = []

        hidden_templates = {
            "kids_history_project",
            "history_timeline_builder",
            "history_study_guide",
            "history_exhibit_pack",
        }
        visible_pipelines = [p for p in pipelines if str(p.metadata.pipeline_id) not in hidden_templates]
        self._wf_templates = list(visible_pipelines)
        for pipeline in visible_pipelines:
            if str(pipeline.metadata.pipeline_id) == "research_to_report":
                try:
                    from core.pipeline.stage import StageConfig

                    for stage in pipeline.stages:
                        if str(stage.stage_id) == "research":
                            cfg = stage.config
                            stage.config = StageConfig(
                                approval_required=cfg.approval_required,
                                approval_message=cfg.approval_message,
                                retry_on_failure=cfg.retry_on_failure,
                                max_retries=cfg.max_retries,
                                fallback_agent=cfg.fallback_agent,
                                timeout_seconds=cfg.timeout_seconds,
                                skip_if_previous_failed=True,
                                condition=cfg.condition,
                            )
                except Exception:
                    pass
            pid = pipeline.metadata.pipeline_id
            name = pipeline.metadata.name or pid
            desc = pipeline.metadata.description or ""
            item = QListWidgetItem(name)
            example = ""
            try:
                example = str((pipeline.global_config or {}).get("example_prompt") or "").strip()
            except Exception:
                example = ""
            tooltip = desc
            if example:
                if tooltip:
                    tooltip = f"{tooltip}\n\nExample:\n{example}"
                else:
                    tooltip = f"Example:\n{example}"
            if tooltip:
                item.setToolTip(tooltip)
            item.setData(Qt.UserRole, pid)
            self._wf_template_list.addItem(item)
            self._wf_template_map[pid] = pipeline

        if not pipelines:
            item = QListWidgetItem("No workflow templates available.")
            item.setFlags(Qt.NoItemFlags)
            self._wf_template_list.addItem(item)

    def _wf_select_template(self, item: QListWidgetItem) -> None:
        try:
            pid = item.data(Qt.UserRole)
        except Exception:
            pid = None
        pipeline = self._wf_template_map.get(pid)
        if not pipeline:
            return
        self._wf_current_pipeline = pipeline
        self._wf_detail_title.setText(pipeline.metadata.name or "Workflow")
        self._wf_detail_desc.setText(pipeline.metadata.description or "")
        self._wf_status.setText("Ready.")
        self._wf_update_stage_tree(pipeline)
        if self._wf_output_name and not self._wf_output_name.text().strip():
            self._wf_output_name.setText(pipeline.metadata.pipeline_id)

    def _wf_update_stage_tree(self, pipeline=None) -> None:
        self._wf_stage_tree.clear()
        self._wf_stage_items = {}
        self._wf_stage_items_by_id = {}
        if not pipeline:
            return
        for stage in pipeline.stages:
            item = QTreeWidgetItem([stage.name, "Pending"])
            item.setData(0, Qt.UserRole, stage.stage_id)
            self._wf_stage_tree.addTopLevelItem(item)
            self._wf_stage_items[str(stage.name)] = item
            self._wf_stage_items_by_id[str(stage.stage_id)] = item

    def _wf_set_stage_status(self, *, stage_name: str | None = None, stage_id: str | None = None, status: str) -> None:
        item = None
        if stage_id and stage_id in getattr(self, "_wf_stage_items_by_id", {}):
            item = self._wf_stage_items_by_id.get(stage_id)
        if not item and stage_name and stage_name in getattr(self, "_wf_stage_items", {}):
            item = self._wf_stage_items.get(stage_name)
        if item is not None:
            try:
                item.setText(1, status)
            except Exception:
                pass

    def _wf_add_source_files(self) -> None:
        try:
            files, _ = QFileDialog.getOpenFileNames(self, "Select source files")
        except Exception:
            files = []
        if not files:
            return
        self._wf_add_source_paths(files)

    def _wf_on_files_dropped(self, paths: list) -> None:
        self._wf_add_source_paths(paths)

    def _wf_add_source_paths(self, paths: list) -> None:
        if not paths:
            return
        existing = set()
        for i in range(self._wf_source_list.count()):
            existing.add(self._wf_source_list.item(i).text())
        for f in paths:
            if f and f not in existing:
                self._wf_source_list.addItem(str(f))
                existing.add(str(f))

    def _wf_remove_source_files(self) -> None:
        for item in list(self._wf_source_list.selectedItems()):
            row = self._wf_source_list.row(item)
            if row >= 0:
                self._wf_source_list.takeItem(row)

    def _wf_pick_output_folder(self) -> None:
        start_dir = ""
        try:
            start_dir = (self._wf_output_folder.text() or "").strip()
        except Exception:
            start_dir = ""
        try:
            path = QFileDialog.getExistingDirectory(self, "Select output folder", start_dir or None)
        except Exception:
            path = ""
        if path:
            self._wf_output_folder.setText(path)
            self._wf_save_output_folder_setting(path)

    def _wf_save_output_folder_field(self) -> None:
        try:
            path = (self._wf_output_folder.text() or "").strip()
        except Exception:
            path = ""
        if path:
            self._wf_save_output_folder_setting(path)

    def _wf_save_output_folder_setting(self, path: str) -> None:
        settings = getattr(self.backend, "settings_manager", None)
        if not settings or not path:
            return
        try:
            wf_settings = settings.get_workflow_settings()
        except Exception:
            wf_settings = {}
        wf_settings = dict(wf_settings or {})
        wf_settings["last_output_folder"] = str(path)
        try:
            settings.save_workflow_settings(wf_settings)
        except Exception:
            pass

    def _wf_collect_source_files(self) -> list[str]:
        files = []
        for i in range(self._wf_source_list.count()):
            try:
                text = self._wf_source_list.item(i).text()
            except Exception:
                continue
            if text:
                files.append(text)
        return files

    def _wf_collect_export_formats(self) -> list[str]:
        formats: list[str] = []
        if self._wf_fmt_md.isChecked():
            formats.append("md")
        if self._wf_fmt_txt.isChecked():
            formats.append("txt")
        if self._wf_fmt_docx.isChecked():
            formats.append("docx")
        if self._wf_fmt_pdf.isChecked():
            formats.append("pdf")
        if not formats:
            formats.append("md")
        return formats

    def _wf_save_web_settings(self) -> None:
        settings = getattr(self.backend, "settings_manager", None)
        if not settings:
            return
        try:
            wf_settings = settings.get_workflow_settings()
        except Exception:
            wf_settings = {}
        wf_settings = dict(wf_settings or {})
        if hasattr(self, "_wf_allow_web"):
            wf_settings["allow_web_research"] = bool(self._wf_allow_web.isChecked())
        if hasattr(self, "_wf_web_max"):
            wf_settings["web_max_results"] = int(self._wf_web_max.value())
        try:
            settings.save_workflow_settings(wf_settings)
        except Exception:
            pass

    def _wf_save_model_preference(self) -> None:
        settings = getattr(self.backend, "settings_manager", None)
        if not settings or not hasattr(self, "_wf_model_pref"):
            return
        try:
            wf_settings = settings.get_workflow_settings()
        except Exception:
            wf_settings = {}
        wf_settings = dict(wf_settings or {})
        try:
            pref = str(self._wf_model_pref.currentData() or "auto").strip().lower()
        except Exception:
            pref = "auto"
        if pref not in {"auto", "text", "vision"}:
            pref = "auto"
        wf_settings["model_preference"] = pref
        try:
            settings.save_workflow_settings(wf_settings)
        except Exception:
            pass

    def _wf_resolve_model_kind(self, source_files: list[str]) -> str:
        pref = "auto"
        try:
            if hasattr(self, "_wf_model_pref"):
                pref = str(self._wf_model_pref.currentData() or "auto").strip().lower()
        except Exception:
            pref = "auto"
        if pref in {"text", "vision"}:
            return pref

        # Auto: prefer text unless images are involved (vision model is better for image-based sources).
        image_exts = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".gif", ".tif", ".tiff", ".heic"}
        try:
            from pathlib import Path

            for p in source_files or []:
                try:
                    if Path(p).suffix.lower() in image_exts:
                        return "vision"
                except Exception:
                    continue
        except Exception:
            pass
        return "text"

    def _wf_append_log(self, text: str) -> None:
        if not text:
            return
        try:
            self._wf_log_view.append(text)
        except Exception:
            pass

    def _wf_set_busy(self, busy: bool) -> None:
        self._wf_run_btn.setEnabled(not busy)
        self._wf_cancel_btn.setEnabled(busy)
        self._wf_template_list.setEnabled(not busy)
        self._wf_history_list.setEnabled(not busy)

    def _wf_request_approval(self, payload: dict) -> bool:
        import threading

        approved = {"ok": False}
        evt = threading.Event()
        self._wf_ui_signals.approval_request.emit({"payload": payload, "event": evt, "holder": approved})
        evt.wait()
        return bool(approved["ok"])

    def _wf_render_output(self, outputs: dict) -> str:
        if not outputs:
            return ""
        last_key = list(outputs.keys())[-1]
        data = outputs.get(last_key)
        return self._wf_stringify_output(data)

    def _wf_stringify_output(self, data) -> str:
        if isinstance(data, list):
            lines = [str(x) for x in data if x is not None]
            return "\n".join(lines)
        if isinstance(data, dict):
            for key in ("document", "text", "summary", "improved_text"):
                val = data.get(key)
                if val:
                    return str(val)
            files = data.get("files")
            if isinstance(files, list) and files:
                return "Files:\n" + "\n".join(str(f) for f in files)
            try:
                import json

                return json.dumps(data, indent=2)
            except Exception:
                return str(data)
        return "" if data is None else str(data)

    def _wf_run_pipeline(self) -> None:
        pipeline = self._wf_current_pipeline
        if not pipeline:
            QMessageBox.information(self, "Workflow", "Select a template before running.")
            return
        request_text = (self._wf_request.toPlainText() or "").strip()
        source_files = self._wf_collect_source_files()
        if not request_text and not source_files:
            QMessageBox.warning(self, "Workflow", "Add a request or source files.")
            return

        model_kind = self._wf_resolve_model_kind(list(source_files))
        if not self._ensure_ai_ready(title="AI Required for Workflow", kind=model_kind):
            return

        settings = getattr(self.backend, "settings_manager", None)
        output_folder = (self._wf_output_folder.text() or "").strip()
        if not output_folder:
            base = Path(getattr(settings, "app_folder", Path.home() / ".fylorra"))
            output_folder = str(base / "automation_outputs")
            self._wf_output_folder.setText(output_folder)
        self._wf_save_output_folder_setting(output_folder)

        output_name = (self._wf_output_name.text() or "workflow_output").strip()
        export_formats = self._wf_collect_export_formats()
        target_language = self._wf_target_lang.currentText() or "English"
        workflow_max_tokens = 0
        if settings:
            try:
                wf_settings = settings.get_workflow_settings()
                workflow_max_tokens = int((wf_settings or {}).get("max_output_tokens", 0) or 0)
            except Exception:
                workflow_max_tokens = 0

        initial_params = {
            "user_request": request_text,
            "source_text": request_text if request_text and not source_files else "",
            "notes": request_text,
            "document": request_text,
            "text": request_text,
            "source_files": list(source_files),
            "workflow_model_kind": model_kind,
            "output_folder": output_folder,
            "output_name": output_name,
            "export_formats": export_formats,
            "target_language": target_language,
            "allow_web_research": bool(self._wf_allow_web.isChecked()) if hasattr(self, "_wf_allow_web") else False,
            "web_max_results": int(self._wf_web_max.value()) if hasattr(self, "_wf_web_max") else 5,
        }
        if workflow_max_tokens > 0:
            initial_params["workflow_max_tokens"] = workflow_max_tokens

        try:
            from core.pipeline.orchestrator import PipelineOrchestrator
        except Exception:
            QMessageBox.critical(self, "Workflow", "Pipeline system is not available in this build.")
            return

        if not self._wf_orchestrator:
            ai = getattr(self.backend, "ai_manager", None)
            self._wf_orchestrator = PipelineOrchestrator(ai, settings)

        self._wf_progress.setValue(0)
        self._wf_output_view.clear()
        self._wf_log_view.clear()
        try:
            import sys

            in_venv = getattr(sys, "base_prefix", sys.prefix) != sys.prefix
            self._wf_append_log(f"Python: {sys.executable} (venv={in_venv})")
        except Exception:
            pass
        self._wf_status.setText("Starting pipeline...")
        self._wf_update_stage_tree(pipeline)
        self._wf_stage_status = {}
        self._wf_last_outputs = {}
        self._wf_last_stage_name = None
        self._wf_active_pipeline = pipeline
        self._wf_active_request = request_text
        self._wf_active_settings = settings
        self._wf_save_web_settings()
        self._wf_set_busy(True)

        def progress_callback(message: str, pct: int, meta: dict):
            try:
                self._wf_ui_signals.progress.emit(str(message or ""), int(pct), dict(meta or {}))
            except Exception:
                pass

        def completion_callback(result):
            try:
                self._wf_ui_signals.completed.emit(result)
            except Exception:
                pass

        self._wf_current_exec_id = self._wf_orchestrator.execute_pipeline(
            pipeline,
            initial_params,
            progress_callback,
            self._wf_request_approval,
            completion_callback,
        )

    def _wf_on_progress(self, message: str, pct: int, meta: dict) -> None:
        msg = str(message or "")
        self._wf_status.setText(msg or "Running...")
        self._wf_progress.setValue(max(0, min(100, int(pct))))
        if msg:
            self._wf_append_log(msg)
        stage_name = (meta or {}).get("stage")
        if stage_name:
            if self._wf_last_stage_name and self._wf_last_stage_name != stage_name:
                if self._wf_stage_status.get(self._wf_last_stage_name) == "Running":
                    self._wf_set_stage_status(stage_name=self._wf_last_stage_name, status="Done")
                    self._wf_stage_status[self._wf_last_stage_name] = "Done"
            if meta.get("skipped"):
                self._wf_set_stage_status(stage_name=stage_name, status="Skipped")
                self._wf_stage_status[stage_name] = "Skipped"
            else:
                self._wf_set_stage_status(stage_name=stage_name, status="Running")
                self._wf_stage_status[stage_name] = "Running"
            self._wf_last_stage_name = stage_name

    def _wf_on_completed(self, result) -> None:
        status = "Completed."
        if getattr(result, "aborted", False):
            status = f"Cancelled: {getattr(result, 'abort_reason', None) or 'User cancelled'}"
        elif not getattr(result, "success", False):
            status = "Completed with errors."
        self._wf_status.setText(status)
        if getattr(result, "success", False):
            self._wf_progress.setValue(100)
        self._wf_append_log(status)

        for err in list(getattr(result, "errors", []) or []):
            stage = err.get("stage") if isinstance(err, dict) else None
            if stage:
                self._wf_set_stage_status(stage_name=stage, status="Failed")
                self._wf_stage_status[stage] = "Failed"

        if self._wf_last_stage_name and self._wf_stage_status.get(self._wf_last_stage_name) == "Running":
            if getattr(result, "aborted", False):
                self._wf_set_stage_status(stage_name=self._wf_last_stage_name, status="Cancelled")
                self._wf_stage_status[self._wf_last_stage_name] = "Cancelled"
            else:
                self._wf_set_stage_status(stage_name=self._wf_last_stage_name, status="Done")
                self._wf_stage_status[self._wf_last_stage_name] = "Done"

        outputs = dict(getattr(result, "outputs", {}) or {})
        self._wf_last_outputs = outputs
        self._wf_output_view.setPlainText(self._wf_render_output(outputs))
        self._wf_set_busy(False)
        self._wf_current_exec_id = None

        pipeline = getattr(self, "_wf_active_pipeline", None) or self._wf_current_pipeline
        request_text = getattr(self, "_wf_active_request", "") or ""
        settings = getattr(self, "_wf_active_settings", None) or getattr(self.backend, "settings_manager", None)

        if pipeline:
            payload = {
                "execution_id": getattr(result, "execution_id", None),
                "pipeline_id": pipeline.metadata.pipeline_id,
                "pipeline_name": pipeline.metadata.name,
                "description": pipeline.metadata.description,
                "completed_at": getattr(result, "completed_at", None),
                "success": getattr(result, "success", False),
                "aborted": getattr(result, "aborted", False),
                "abort_reason": getattr(result, "abort_reason", None),
                "errors": getattr(result, "errors", []),
                "outputs": outputs,
                "request": request_text,
                "stage_status": dict(self._wf_stage_status),
                "log": self._wf_log_view.toPlainText().splitlines(),
            }
            try:
                from core.pipeline.storage import save_execution

                save_execution(str(payload.get("execution_id") or ""), payload, settings)
            except Exception:
                pass
            self._wf_load_execution_history()

    def _wf_on_approval_request(self, payload: dict) -> None:
        try:
            data = dict(payload or {})
            evt = data.get("event")
            holder = data.get("holder")
            req = data.get("payload") or {}
        except Exception:
            return
        try:
            self._wf_status.setText("Waiting for approval...")
            dialog = _QtWorkflowApprovalDialog(self, payload=req)
            approved = dialog.exec() == QDialog.Accepted
            if isinstance(holder, dict):
                holder["ok"] = bool(approved)
        except Exception:
            if isinstance(holder, dict):
                holder["ok"] = False
        try:
            if evt:
                evt.set()
        except Exception:
            pass

    def _wf_cancel_pipeline(self) -> None:
        if not self._wf_orchestrator or not self._wf_current_exec_id:
            return
        self._wf_orchestrator.cancel_pipeline(self._wf_current_exec_id)
        self._wf_status.setText("Cancel requested.")

    def _wf_load_execution_history(self) -> None:
        try:
            import json
            from core.pipeline.storage import ensure_pipeline_dirs
        except Exception:
            self._wf_history_list.clear()
            return

        self._wf_history_list.clear()
        settings = getattr(self.backend, "settings_manager", None)
        dirs = ensure_pipeline_dirs(settings)
        exec_dir = dirs.get("executions")
        if not exec_dir or not exec_dir.exists():
            item = QListWidgetItem("No executions yet.")
            item.setFlags(Qt.NoItemFlags)
            self._wf_history_list.addItem(item)
            return

        entries = []
        for path in exec_dir.glob("*.json"):
            try:
                payload = json.loads(path.read_text(encoding="utf-8"))
            except Exception:
                continue
            try:
                mtime = path.stat().st_mtime
            except Exception:
                mtime = 0
            entries.append((mtime, payload, path))

        entries.sort(key=lambda x: x[0], reverse=True)
        for _, payload, path in entries:
            name = str(payload.get("pipeline_name") or payload.get("pipeline_id") or "Workflow")
            completed = str(payload.get("completed_at") or "")
            label = name
            if completed:
                label = f"{name} - {completed}"
            item = QListWidgetItem(label)
            item.setData(Qt.UserRole, payload)
            item.setData(Qt.UserRole + 1, str(path))
            self._wf_history_list.addItem(item)

        if not entries:
            item = QListWidgetItem("No executions yet.")
            item.setFlags(Qt.NoItemFlags)
            self._wf_history_list.addItem(item)

    def _wf_delete_selected_history(self) -> None:
        if not hasattr(self, "_wf_history_list"):
            return
        item = self._wf_history_list.currentItem()
        if not item:
            return
        try:
            payload = item.data(Qt.UserRole)
        except Exception:
            payload = None
        if not isinstance(payload, dict):
            return
        try:
            raw_path = item.data(Qt.UserRole + 1)
        except Exception:
            raw_path = ""
        try:
            path = Path(str(raw_path)) if raw_path else None
        except Exception:
            path = None
        if path and path.exists():
            try:
                path.unlink()
            except Exception:
                return
        self._wf_load_execution_history()

    def _wf_open_history_item(self, item: QListWidgetItem) -> None:
        try:
            payload = item.data(Qt.UserRole)
        except Exception:
            payload = None
        if not isinstance(payload, dict):
            return

        pipeline_id = payload.get("pipeline_id")
        pipeline = self._wf_template_map.get(pipeline_id)
        if pipeline:
            self._wf_current_pipeline = pipeline
            self._wf_detail_title.setText(pipeline.metadata.name or "Workflow")
            self._wf_detail_desc.setText(pipeline.metadata.description or "")
            self._wf_update_stage_tree(pipeline)
        else:
            self._wf_detail_title.setText(str(payload.get("pipeline_name") or "Workflow"))
            self._wf_detail_desc.setText(str(payload.get("description") or ""))
            self._wf_stage_tree.clear()

        self._wf_status.setText("History loaded.")
        outputs = dict(payload.get("outputs") or {})
        self._wf_last_outputs = outputs
        self._wf_output_view.setPlainText(self._wf_render_output(outputs))
        log_lines = payload.get("log") or []
        if isinstance(log_lines, list):
            self._wf_log_view.setPlainText("\n".join(str(x) for x in log_lines))
        else:
            self._wf_log_view.setPlainText(str(log_lines))

        stage_status = payload.get("stage_status") or {}
        if isinstance(stage_status, dict):
            for stage_name, status in stage_status.items():
                self._wf_set_stage_status(stage_name=str(stage_name), status=str(status))

    def _wf_on_stage_selected(self) -> None:
        if not hasattr(self, "_wf_last_outputs"):
            return
        items = self._wf_stage_tree.selectedItems()
        if not items:
            return
        item = items[0]
        try:
            stage_id = item.data(0, Qt.UserRole)
        except Exception:
            stage_id = None
        outputs = getattr(self, "_wf_last_outputs", {}) or {}
        if not stage_id or stage_id not in outputs:
            self._wf_output_view.setPlainText("No output available for this stage yet.")
            return
        self._wf_output_view.setPlainText(self._wf_stringify_output(outputs.get(stage_id)))

    def _ai_search_browse_folder(self):
        try:
            path = QFileDialog.getExistingDirectory(self, "Select folder to index")
            if path:
                self._ai_search_folder.setText(path)
        except Exception:
            pass

    def _ai_search_start_index(self):
        folder = (self._ai_search_folder.text() or "").strip()
        if not folder:
            QMessageBox.warning(self, "Index", "Pick a folder to index.")
            return

        # Auto-load AI only if user enabled AI-dependent index options.
        wants_ai = bool(self._ai_search_ai_summ.isChecked() or self._ai_search_ocr.isChecked() or self._ai_search_extract_images.isChecked())
        if wants_ai and not self._ensure_ai_ready(title="AI Required for Indexing"):
            return
        ai = getattr(self.backend, "ai_manager", None)
        ai_ready = bool(ai and getattr(ai, "is_ready", False))
        ai_summ = bool(self._ai_search_ai_summ.isChecked() and ai_ready)
        extract_images = bool(self._ai_search_extract_images.isChecked() and ai_ready)
        ocr_scanned = bool(self._ai_search_ocr.isChecked() and ai_ready)

        self._ai_search_index_btn.setEnabled(False)
        self._ai_search_cancel_index_btn.setEnabled(True)
        self._ai_search_index_bar.setValue(0)
        self._ai_search_index_status.setText("Starting…")

        worker = _QtLibraryIndexWorker(
            backend=self.backend,
            folder=folder,
            include_subfolders=bool(self._ai_search_include_sub.isChecked()),
            ai_summarize=ai_summ,
            ocr_scanned_pdfs=ocr_scanned,
            extract_images=extract_images,
            compute_hashes=bool(self._ai_search_hashes.isChecked()),
        )
        th = QThread(self)
        worker.moveToThread(th)
        th.started.connect(worker.run)
        worker.status.connect(self._ai_search_index_status.setText)
        worker.progress.connect(lambda p: self._ai_search_index_bar.setValue(int(max(0.0, min(1.0, p)) * 1000)))
        worker.finished.connect(lambda n: self._ai_search_index_done(n))
        worker.error.connect(lambda msg: self._ai_search_index_failed(msg))
        worker.finished.connect(th.quit)
        worker.error.connect(th.quit)
        th.finished.connect(th.deleteLater)
        th.finished.connect(worker.deleteLater)
        self._ai_search_index_worker = worker
        self._ai_search_index_thread = th
        th.start()

    def _ai_search_cancel_index(self):
        try:
            w = getattr(self, "_ai_search_index_worker", None)
            if w:
                w.cancel()
        except Exception:
            pass

    def _ai_search_index_done(self, n: int):
        self._ai_search_index_btn.setEnabled(True)
        self._ai_search_cancel_index_btn.setEnabled(False)
        self._ai_search_index_bar.setValue(1000)
        self._ai_search_index_status.setText(f"Index complete. Scanned {int(n)} file(s).")

    def _ai_search_index_failed(self, msg: str):
        self._ai_search_index_btn.setEnabled(True)
        self._ai_search_cancel_index_btn.setEnabled(False)
        self._ai_search_index_status.setText("Failed.")
        QMessageBox.critical(self, "Index Failed", msg)

    def _ai_search_run(self):
        q = (self._ai_search_query.text() or "").strip()
        if not q:
            return

        scope_folder = (self._ai_search_folder.text() or "").strip()
        if not scope_folder:
            QMessageBox.warning(self, "Search", "Pick a folder (same folder you indexed), then search.")
            return

        # If the query looks like a visual search and the user hasn't enabled vision captions,
        # prompt them to enable + reindex (otherwise results will be filename-only).
        try:
            looks_visual = any(w in q.lower() for w in ("image", "photo", "picture", "screenshot", "show me", "find all"))
            looks_visual = looks_visual or any(w in q.lower() for w in ("girl", "boy", "teen", "dog", "cat", "car", "food", "invoice"))
            if looks_visual and hasattr(self, "_ai_search_extract_images") and not self._ai_search_extract_images.isChecked():
                QMessageBox.information(
                    self,
                    "Image Search Tip",
                    "To search images by content (vision), enable “AI vision captions (images)” and click “Index Now” to update the index.\n\n"
                    "Otherwise search will only match filenames/visible text.",
                )
        except Exception:
            pass

        if bool(self._ai_search_rerank.isChecked()):
            if not self._ensure_ai_ready(title="AI Required for Rerank"):
                return
        self._ai_search_last_scope_folder = scope_folder

        self._ai_search_btn.setEnabled(False)
        self._ai_search_results.clear()
        self._ai_search_preview.clear()
        self._ai_search_prev_title.setText("Searching…")
        self._ai_search_prev_path.setText("")
        self._ai_search_search_status.setText("Searching…")
        self._ai_search_preview.setPlainText("Searching…")

        worker = _QtLibrarySearchWorker(
            backend=self.backend,
            query=q,
            rerank=bool(self._ai_search_rerank.isChecked()),
            folder=scope_folder,
        )
        th = QThread(self)
        worker.moveToThread(th)
        th.started.connect(worker.run)
        worker.status.connect(lambda msg: self._ai_search_search_status.setText(str(msg)))
        worker.finished.connect(lambda results: self._ai_search_results_ready(results))
        worker.error.connect(lambda msg: self._ai_search_results_failed(msg))
        worker.finished.connect(th.quit)
        worker.error.connect(th.quit)
        th.finished.connect(th.deleteLater)
        th.finished.connect(worker.deleteLater)

        # Keep references to avoid GC while the thread is running.
        self._ai_search_search_worker = worker
        self._ai_search_search_thread = th
        th.finished.connect(lambda: setattr(self, "_ai_search_search_worker", None))
        th.finished.connect(lambda: setattr(self, "_ai_search_search_thread", None))

        th.start()

    def _ai_search_results_ready(self, results):
        self._ai_search_btn.setEnabled(True)
        self._ai_search_search_status.setText("")
        self._ai_search_preview.setPlainText("")
        self._ai_search_results.clear()
        for r in results or []:
            path = getattr(r.item, "path", "")
            name = getattr(r.item, "name", "") or Path(path).name
            it = QListWidgetItem()
            it.setData(Qt.UserRole, r)
            it.setSizeHint(QSize(10, 74))
            self._ai_search_results.addItem(it)
            w = _QtAISearchResultCard(r)
            self._ai_search_results.setItemWidget(it, w)
        if self._ai_search_results.count() == 0:
            self._ai_search_preview.setPlainText(
                "No results.\n\nTip: click “Index Now” for this folder first, then try broader terms."
            )
            self._ai_search_prev_title.setText("No results")
            self._ai_search_prev_path.setText("")
            self._ai_search_search_status.setText("No results.")
        else:
            try:
                first = results[0]
                scope = ""
                try:
                    scope = (getattr(self, "_ai_search_last_scope_folder", "") or "").strip()
                except Exception:
                    scope = ""
                scope_txt = f"   •   Scope: {scope}" if scope else ""
                self._ai_search_search_status.setText(
                    f"Results: {len(results)}   •   Query used: {first.matched_query}   •   AI rewrite: {'ON' if first.used_ai else 'OFF'}   •   Rerank: {'ON' if first.used_rerank else 'OFF'}{scope_txt}"
                )
            except Exception:
                self._ai_search_search_status.setText(f"Results: {self._ai_search_results.count()}")
            self._ai_search_results.setCurrentRow(0)

    def _ai_search_results_failed(self, msg: str):
        self._ai_search_btn.setEnabled(True)
        QMessageBox.critical(self, "Search Failed", msg)
        try:
            self._ai_search_search_status.setText("Search failed.")
            self._ai_search_prev_title.setText("Search failed")
            self._ai_search_prev_path.setText("")
        except Exception:
            pass
        self._ai_search_preview.setPlainText("Search failed.\n\n" + (msg or ""))

    def _ai_search_on_select(self):
        # Update card highlight state.
        try:
            for i in range(int(self._ai_search_results.count())):
                it = self._ai_search_results.item(i)
                w = self._ai_search_results.itemWidget(it)
                if w and hasattr(w, "set_selected"):
                    w.set_selected(bool(it.isSelected()))
        except Exception:
            pass
        items = self._ai_search_results.selectedItems()
        if not items:
            return
        r = items[0].data(Qt.UserRole)
        try:
            txt = (r.item.ai_summary or r.item.extracted_text or "").strip()
        except Exception:
            txt = ""
        if not txt:
            txt = "No extracted text/summary available for this file."
        try:
            p = str(r.item.path or "")
            self._ai_search_prev_title.setText(str(r.item.name or Path(p).name or "Selected file"))
            self._ai_search_prev_path.setText(p)
        except Exception:
            pass
        self._ai_search_preview.setPlainText(txt)

    def _ai_search_selected_path(self) -> str:
        items = self._ai_search_results.selectedItems()
        if not items:
            return ""
        r = items[0].data(Qt.UserRole)
        try:
            return str(r.item.path)
        except Exception:
            return ""

    def _ai_search_open_selected(self):
        path = self._ai_search_selected_path()
        if not path:
            return
        try:
            p = Path(str(path))
        except Exception:
            p = None
        if p and not p.exists():
            QMessageBox.warning(self, "Open", "File no longer exists.")
            return
        # Prefer native path open on Windows to avoid URL encoding issues (e.g. apostrophes).
        try:
            import os
            import sys

            if sys.platform.startswith("win"):
                os.startfile(str(p or path))  # type: ignore[attr-defined]
                return
        except Exception:
            pass
        try:
            ok = QDesktopServices.openUrl(QUrl.fromLocalFile(str(p or path)))
            if not ok:
                raise RuntimeError("ShellExecute failed")
        except Exception as e:
            QMessageBox.warning(self, "Open", f"Could not open file.\n\n{e}")

    def _ai_search_show_in_folder(self):
        p = self._ai_search_selected_path()
        if not p:
            return
        try:
            import subprocess

            subprocess.Popen(f'explorer /select,"{p}"')
        except Exception:
            try:
                QDesktopServices.openUrl(QUrl.fromLocalFile(str(Path(p).parent)))
            except Exception:
                pass

    def _ai_search_explain_selected(self):
        p = self._ai_search_selected_path()
        if not p:
            return
        ai = getattr(self.backend, "ai_manager", None)
        # If model isn't loaded, we can still explain using undo DB evidence.
        try:
            from core.filename_explainer import explain_filename

            text = explain_filename(Path(p), ai_manager=ai)
        except Exception as e:
            text = str(e)
        dlg = QDialog(self)
        dlg.setWindowTitle("Explain Filename")
        dlg.setModal(True)
        dlg.setMinimumWidth(720)
        dlg.setMinimumHeight(420)
        v = QVBoxLayout(dlg)
        v.setContentsMargins(14, 14, 14, 14)
        v.setSpacing(10)
        view = QTextEdit()
        view.setReadOnly(True)
        view.setPlainText(text)
        v.addWidget(view, 1)
        buttons = QDialogButtonBox(QDialogButtonBox.Ok)
        try:
            ok_btn = buttons.button(QDialogButtonBox.Ok)
            if ok_btn:
                ok_btn.setObjectName("PrimaryButton")
        except Exception:
            pass
        buttons.accepted.connect(dlg.accept)
        v.addWidget(buttons)
        dlg.exec()

    def _build_file_tools_page(self) -> QWidget:
        host = QFrame()
        host.setObjectName("PageHost")
        layout = QVBoxLayout(host)
        layout.setContentsMargins(18, 18, 18, 18)
        layout.setSpacing(14)

        header = QFrame()
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(0, 0, 0, 0)
        header_layout.setSpacing(10)
        title = QLabel("File Tools")
        title.setObjectName("PageTitle")
        header_layout.addWidget(title)
        sub = QLabel("Convert media/images, extract text (OCR), create/extract archives, and run batch operations with progress.")
        sub.setObjectName("PageSubTitle")
        header_layout.addWidget(sub, 1)
        layout.addWidget(header)

        card = QFrame()
        card.setObjectName("PageCard")
        card_layout = QVBoxLayout(card)
        card_layout.setContentsMargins(16, 14, 16, 14)
        card_layout.setSpacing(12)

        # Target folder row
        top = QHBoxLayout()
        top.setSpacing(10)
        self._ft_target = QLineEdit()
        self._ft_target.setPlaceholderText("Target folder…")
        top.addWidget(self._ft_target, 1)
        browse = QPushButton("Browse")
        browse.clicked.connect(self._ft_browse_target)
        top.addWidget(browse)

        ai_btn = QPushButton("Use AI")
        ai_btn.setObjectName("PrimaryButton")
        ai_btn.setIcon(self.icons.icon("ai"))
        ai_btn.setIconSize(QSize(18, 18))
        ai_btn.clicked.connect(self._ft_open_ai_command)
        top.addWidget(ai_btn)
        self._ft_target_row = QWidget()
        self._ft_target_row.setLayout(top)
        card_layout.addWidget(self._ft_target_row)

        # In-page navigation: Home -> Tool pages (more space, less clutter)
        nav = QHBoxLayout()
        nav.setSpacing(10)
        self._ft_back = QToolButton()
        self._ft_back.setText("← Back")
        self._ft_back.setCursor(Qt.PointingHandCursor)
        self._ft_back.clicked.connect(self._ft_show_home)
        self._ft_back.setVisible(False)
        nav.addWidget(self._ft_back)

        self._ft_media_tab_host = QWidget()
        tab_row = QHBoxLayout(self._ft_media_tab_host)
        tab_row.setContentsMargins(0, 0, 0, 0)
        tab_row.setSpacing(6)
        self._ft_media_tab_group = QButtonGroup(self)
        self._ft_media_tab_group.setExclusive(True)
        tab_names = ["Convert", "Converting", "Converted"]
        for idx, name in enumerate(tab_names):
            btn = QToolButton()
            btn.setText(name)
            btn.setCheckable(True)
            btn.setObjectName("NavButton")
            self._ft_media_tab_group.addButton(btn, idx)
            tab_row.addWidget(btn)
        try:
            self._ft_media_tab_group.button(0).setChecked(True)
        except Exception:
            pass
        self._ft_media_tab_group.buttonClicked.connect(self._ft_media_tab_changed)
        self._ft_media_tab_host.setVisible(False)
        nav.addWidget(self._ft_media_tab_host)

        self._ft_tool_title = QLabel("Choose a tool")
        self._ft_tool_title.setStyleSheet("color:#9aa0a9; font-weight:600;")
        nav.addWidget(self._ft_tool_title)
        nav.addStretch(1)
        card_layout.addLayout(nav)

        self._ft_stack = QStackedWidget()

        # Build pages
        self._ft_home_page = self._build_file_tools_home_page()
        self._ft_media_page = self._build_file_tools_media_page()
        self._ft_images_page = self._build_file_tools_images_page()
        self._ft_ocr_page = self._build_file_tools_ocr_page()
        self._ft_arch_page = self._build_file_tools_archives_page()
        self._ft_pdf_page = self._build_file_tools_pdf_page()
        self._ft_office_page = self._build_file_tools_office_page()
        self._ft_stack.addWidget(self._ft_home_page)   # 0
        self._ft_stack.addWidget(self._ft_media_page)  # 1
        self._ft_stack.addWidget(self._ft_images_page)  # 2
        self._ft_stack.addWidget(self._ft_ocr_page)    # 3
        self._ft_stack.addWidget(self._ft_arch_page)   # 4
        self._ft_stack.addWidget(self._ft_pdf_page)    # 5
        self._ft_stack.addWidget(self._ft_office_page)  # 6
        self._ft_stack.setCurrentIndex(0)

        # Progress area (shared)
        self._ft_status = QLabel("Ready.")
        self._ft_status.setStyleSheet("color:#9aa0a9;")
        self._ft_bar = QProgressBar()
        self._ft_bar.setRange(0, 1000)
        self._ft_bar.setValue(0)
        self._ft_cancel = QPushButton("Cancel")
        self._ft_cancel.setEnabled(False)
        self._ft_cancel.clicked.connect(self._ft_cancel_running)

        progress_row = QHBoxLayout()
        progress_row.setSpacing(10)
        progress_row.addWidget(self._ft_status, 1)
        progress_row.addWidget(self._ft_cancel)

        card_layout.addWidget(self._ft_stack, 1)
        self._ft_progress_host = QWidget()
        prog_layout = QVBoxLayout(self._ft_progress_host)
        prog_layout.setContentsMargins(0, 0, 0, 0)
        prog_layout.setSpacing(8)
        prog_layout.addWidget(self._ft_bar)
        prog_layout.addLayout(progress_row)
        card_layout.addWidget(self._ft_progress_host)

        layout.addWidget(card, 1)
        return host

    def _build_file_tools_home_page(self) -> QWidget:
        page = QWidget()
        outer = QVBoxLayout(page)
        outer.setContentsMargins(0, 0, 0, 0)
        outer.setSpacing(12)
        self._ft_home_tile_parts = []

        tip = QLabel("Pick a tool below. Use the Back button to return here.")
        self._ft_home_tip = tip
        outer.addWidget(tip)

        grid_host = QWidget()
        grid = QGridLayout(grid_host)
        grid.setContentsMargins(0, 0, 0, 0)
        grid.setHorizontalSpacing(12)
        grid.setVerticalSpacing(12)
        grid.setColumnStretch(0, 1)
        grid.setColumnStretch(1, 1)

        def tile(title: str, desc: str, *, icon_key: str | None, open_fn):
            frame = QFrame()
            frame.setObjectName("ToolTile")
            v = QVBoxLayout(frame)
            v.setContentsMargins(14, 12, 14, 12)
            v.setSpacing(10)

            head = QHBoxLayout()
            head.setSpacing(10)
            badge = QFrame()
            badge.setObjectName("IconBadge")
            badge_l = QHBoxLayout(badge)
            badge_l.setContentsMargins(0, 0, 0, 0)
            badge_l.setSpacing(0)
            badge_l.setAlignment(Qt.AlignCenter)
            icon = QLabel()
            if icon_key:
                try:
                    pm = self.icons.pixmap(icon_key, 28)
                    if pm.isNull():
                        pm = self.icons.pixmap("file_tools", 28)
                    if not pm.isNull():
                        icon.setPixmap(pm)
                except Exception:
                    pass
            icon.setFixedSize(28, 28)
            badge_l.addWidget(icon)
            badge.setFixedSize(44, 44)
            head.addWidget(badge)
            t = QLabel(title)
            head.addWidget(t, 1)
            head.addStretch(1)
            v.addLayout(head)

            d = QLabel(desc)
            d.setWordWrap(True)
            v.addWidget(d, 1)
            self._ft_home_tile_parts.append(
                {
                    "title": t,
                    "desc": d,
                    "icon": icon,
                    "icon_key": icon_key,
                }
            )

            btn = QPushButton("Open")
            btn.setObjectName("PrimaryButton")
            btn.clicked.connect(open_fn)
            v.addWidget(btn)
            return frame

        grid.addWidget(
            tile(
                "Media Converter",
                "Convert audio/video in batch or single-file mode (GPU if available).",
                icon_key="File_Tools/Media_Converter.png",
                open_fn=lambda: self._ft_show_tool("media"),
            ),
            0,
            0,
        )
        grid.addWidget(
            tile(
                "Images",
                "Batch convert images (WebP/PNG/JPG) with optional subfolders.",
                icon_key="File_Tools/Images.png",
                open_fn=lambda: self._ft_show_tool("images"),
            ),
            0,
            1,
        )
        grid.addWidget(
            tile(
                "Text Extractor",
                "Extract text from PDFs, images, and office files (OCR + markdown/export).",
                icon_key="File_Tools/OCR.png",
                open_fn=lambda: self._ft_show_tool("ocr"),
            ),
            1,
            0,
        )
        grid.addWidget(
            tile(
                "PDF Tools",
                "Merge, split, extract pages, rotate, and split by bookmarks.",
                icon_key="File_Tools/PDF_Tools.png",
                open_fn=lambda: self._ft_show_tool("pdf"),
            ),
            1,
            1,
        )
        grid.addWidget(
            tile(
                "Archives",
                "Create/extract ZIP, 7z, tar.* archives and split into parts.",
                icon_key="File_Tools/Archives.png",
                open_fn=lambda: self._ft_show_tool("archives"),
            ),
            2,
            0,
        )
        grid.addWidget(
            tile(
                "Office Convert",
                "Convert documents/spreadsheets/presentations (and some PDFs) using LibreOffice headless.",
                icon_key="File_Tools/Office_Convert.png",
                open_fn=lambda: self._ft_show_tool("office"),
            ),
            2,
            1,
        )

        outer.addWidget(grid_host, 1)
        self._ft_apply_home_theme()
        return page

    def _ft_apply_home_theme(self) -> None:
        t = _ui_theme_tokens()
        mode = _ui_theme_mode()
        try:
            if hasattr(self, "_ft_home_tip"):
                self._ft_home_tip.setStyleSheet(f"color:{t['muted']};")
        except Exception:
            pass
        for part in list(getattr(self, "_ft_home_tile_parts", []) or []):
            try:
                tl = part.get("title")
                if isinstance(tl, QLabel):
                    tl.setStyleSheet(f"font-size:14px; font-weight:800; color:{t['text']};")
            except Exception:
                pass
            try:
                ds = part.get("desc")
                if isinstance(ds, QLabel):
                    ds.setStyleSheet(f"color:{t['muted']};")
            except Exception:
                pass
            try:
                ic = part.get("icon")
                key = part.get("icon_key")
                if isinstance(ic, QLabel):
                    icon_key = str(key or "file_tools")
                    pm = self.icons.pixmap(icon_key, 28)
                    if pm.isNull():
                        pm = self.icons.pixmap("file_tools", 28)
                    if not pm.isNull() and mode == "light":
                        pm = self._tint_pixmap(pm, t["icon"])
                    if not pm.isNull():
                        ic.setPixmap(pm)
            except Exception:
                pass

    def _ft_show_home(self):
        try:
            self._ft_stack.setCurrentIndex(0)
            self._ft_back.setVisible(False)
            self._ft_tool_title.setText("Choose a tool")
            try:
                self._ft_media_tab_host.setVisible(False)
            except Exception:
                pass
            try:
                self._ft_progress_host.setVisible(True)
            except Exception:
                pass
            try:
                self._ft_target_row.setVisible(True)
            except Exception:
                pass
        except Exception:
            pass

    def _ft_show_tool(self, which: str):
        which = (which or "").strip().lower()
        idx = 0
        title = "Choose a tool"
        if which == "media":
            idx, title = 1, "Media Converter"
        elif which == "images":
            idx, title = 2, "Images"
        elif which == "ocr":
            idx, title = 3, "Text Extractor"
        elif which == "archives":
            idx, title = 4, "Archives"
        elif which == "pdf":
            idx, title = 5, "PDF Tools"
        elif which == "office":
            idx, title = 6, "Office Convert"
        try:
            self._ft_stack.setCurrentIndex(idx)
            self._ft_back.setVisible(idx != 0)
            self._ft_tool_title.setText(title)
            try:
                self._ft_target_row.setVisible(idx != 1)
            except Exception:
                pass
            try:
                self._ft_media_tab_host.setVisible(idx == 1)
                self._ft_progress_host.setVisible(idx != 1)
                if idx == 1:
                    try:
                        self._ft_media_tab_group.button(0).setChecked(True)
                    except Exception:
                        pass
                    try:
                        self._ft_media_tab_stack.setCurrentIndex(0)
                    except Exception:
                        pass
                    try:
                        if hasattr(self, "_ft_media_autostart_toggle"):
                            self._ft_media_autostart_toggle.setChecked(self._ft_media_autostart_enabled())
                    except Exception:
                        pass
            except Exception:
                pass
        except Exception:
            pass

    def _ft_media_tab_changed(self, btn):
        try:
            idx = self._ft_media_tab_group.id(btn)
            self._ft_media_tab_stack.setCurrentIndex(int(idx))
        except Exception:
            pass

    def _ft_require_libreoffice(self) -> bool:
        try:
            from core.lo_converter import LibreOfficeConverter

            conv = LibreOfficeConverter()
            if conv.is_available():
                return True
        except Exception:
            pass

        box = QMessageBox(self)
        box.setIcon(QMessageBox.Warning)
        box.setWindowTitle("LibreOffice Required")
        box.setText("Office conversions require LibreOffice (soffice).")
        box.setInformativeText("Download LibreOffice, or open Settings to configure an existing install.")
        dl = box.addButton("Download LibreOffice", QMessageBox.AcceptRole)
        st = box.addButton("Open Settings", QMessageBox.ActionRole)
        box.addButton(QMessageBox.Cancel)
        box.exec()
        clicked = box.clickedButton()
        if clicked == dl:
            try:
                self._settings_download_lo()
            except Exception:
                pass
        elif clicked == st:
            try:
                self.set_active_page("settings")
            except Exception:
                pass
        return False

    def _build_file_tools_office_page(self) -> QWidget:
        page = QWidget()
        outer = QVBoxLayout(page)
        outer.setContentsMargins(0, 0, 0, 0)
        outer.setSpacing(12)

        # Status row
        status_row = QHBoxLayout()
        status_row.setSpacing(10)
        self._ft_office_status = QLabel("LibreOffice: checking…")
        self._ft_office_status.setStyleSheet("color:#9aa0a9;")
        status_row.addWidget(self._ft_office_status, 1)
        btn_dl = QPushButton("Download LibreOffice")
        btn_dl.setObjectName("PrimaryButton")
        btn_dl.setIcon(self.icons.icon("download"))
        btn_dl.setIconSize(QSize(18, 18))
        btn_dl.clicked.connect(self._settings_download_lo)
        status_row.addWidget(btn_dl)
        outer.addLayout(status_row)

        # Mode switch
        mode_row = QHBoxLayout()
        mode_row.setSpacing(10)
        mode_row.addWidget(QLabel("Mode:"), 0)
        self._ft_off_mode = QButtonGroup(page)
        self._ft_off_mode.setExclusive(True)
        b_batch = QToolButton()
        b_batch.setText("Batch")
        b_batch.setCheckable(True)
        b_batch.setObjectName("NavButton")
        b_single = QToolButton()
        b_single.setText("Single")
        b_single.setCheckable(True)
        b_single.setObjectName("NavButton")
        b_md = QToolButton()
        b_md.setText("Markdown")
        b_md.setCheckable(True)
        b_md.setObjectName("NavButton")
        self._ft_off_mode.addButton(b_batch, 0)
        self._ft_off_mode.addButton(b_single, 1)
        self._ft_off_mode.addButton(b_md, 2)
        mode_row.addWidget(b_batch)
        mode_row.addWidget(b_single)
        mode_row.addWidget(b_md)
        mode_row.addStretch(1)
        outer.addLayout(mode_row)

        self._ft_off_stack = QStackedWidget()
        self._ft_off_stack.addWidget(self._build_file_tools_office_batch())
        self._ft_off_stack.addWidget(self._build_file_tools_office_single())
        self._ft_off_stack.addWidget(self._build_file_tools_office_markdown())
        outer.addWidget(self._ft_off_stack, 1)

        def on_mode(btn):
            try:
                self._ft_off_stack.setCurrentIndex(self._ft_off_mode.id(btn))
            except Exception:
                pass

        self._ft_off_mode.buttonClicked.connect(on_mode)
        b_batch.setChecked(True)
        self._ft_off_stack.setCurrentIndex(0)

        self._ft_refresh_office_status()
        return page

    def _ft_refresh_office_status(self):
        try:
            from core.lo_converter import LibreOfficeConverter

            conv = LibreOfficeConverter()
            self._ft_office_status.setText(f"LibreOffice: {'OK' if conv.is_available() else 'missing'}")
        except Exception:
            self._ft_office_status.setText("LibreOffice: ?")

    def _build_file_tools_office_batch(self) -> QWidget:
        w = QWidget()
        outer = QVBoxLayout(w)
        outer.setContentsMargins(0, 0, 0, 0)
        outer.setSpacing(12)

        box = QGroupBox("Batch Convert (uses Target folder)")
        bwrap = QVBoxLayout(box)
        bwrap.setContentsMargins(12, 10, 12, 10)
        bwrap.setSpacing(8)

        g = QGridLayout()
        g.setContentsMargins(0, 0, 0, 0)
        g.setHorizontalSpacing(12)
        g.setVerticalSpacing(10)
        g.setColumnStretch(1, 1)

        self._ft_off_src_sub = QLineEdit()
        self._ft_off_src_sub.setPlaceholderText("Source subfolder (optional, relative under target)")
        lbl_src = QLabel("Source subfolder:")
        lbl_src.setAlignment(Qt.AlignRight | Qt.AlignVCenter)
        g.addWidget(lbl_src, 0, 0)
        g.addWidget(self._ft_off_src_sub, 0, 1)

        self._ft_off_include_sub = QCheckBox("Include subfolders")
        self._ft_off_include_sub.setChecked(True)
        g.addWidget(self._ft_off_include_sub, 1, 1)

        self._ft_off_out_sub = QLineEdit("Converted_Office")
        lbl_outsub = QLabel("Output subfolder:")
        lbl_outsub.setAlignment(Qt.AlignRight | Qt.AlignVCenter)
        g.addWidget(lbl_outsub, 2, 0)
        g.addWidget(self._ft_off_out_sub, 2, 1)

        fmt_row = QHBoxLayout()
        fmt_row.setSpacing(10)
        self._ft_off_fmt = QComboBox()
        self._ft_off_fmt.addItems(["pdf", "docx", "odt", "xlsx", "ods", "csv", "pptx", "odp", "txt", "html"])
        self._ft_off_root = QComboBox()
        self._ft_off_root.addItems(["target", "source"])
        fmt_row.addWidget(QLabel("Format:"))
        fmt_row.addWidget(self._ft_off_fmt)
        fmt_row.addWidget(QLabel("Output:"))
        fmt_row.addWidget(self._ft_off_root)
        fmt_host = QWidget()
        fmt_host.setLayout(fmt_row)
        g.addWidget(QLabel(""), 3, 0)
        g.addWidget(fmt_host, 3, 1)

        self._ft_off_overwrite = QCheckBox("Overwrite outputs")
        g.addWidget(QLabel(""), 4, 0)
        g.addWidget(self._ft_off_overwrite, 4, 1)

        bwrap.addLayout(g)

        run = QPushButton("Convert Documents (Batch)")
        run.setObjectName("PrimaryButton")
        run.clicked.connect(self._ft_run_office_batch)
        bwrap.addWidget(run)

        note = QLabel("Tip: Some conversions (PDF→DOCX/TXT) may use text-extraction fallback if LibreOffice output is empty.")
        note.setStyleSheet("color:#9aa0a9;")
        note.setWordWrap(True)
        bwrap.addWidget(note)

        outer.addWidget(box)
        outer.addStretch(1)
        return w

    def _build_file_tools_office_single(self) -> QWidget:
        w = QWidget()
        outer = QVBoxLayout(w)
        outer.setContentsMargins(0, 0, 0, 0)
        outer.setSpacing(12)

        box = QGroupBox("Single File Convert")
        f = QFormLayout(box)
        f.setLabelAlignment(Qt.AlignRight)
        f.setVerticalSpacing(10)
        f.setHorizontalSpacing(12)

        row = QHBoxLayout()
        self._ft_off_single_in = QLineEdit()
        self._ft_off_single_in.setPlaceholderText("Pick a document…")
        row.addWidget(self._ft_off_single_in, 1)
        pick = QPushButton("Browse")
        pick.clicked.connect(self._ft_pick_office_single)
        row.addWidget(pick)
        host = QWidget()
        host.setLayout(row)
        f.addRow("Input:", host)

        self._ft_off_single_out_sub = QLineEdit("Converted_Office")
        f.addRow("Output subfolder:", self._ft_off_single_out_sub)

        fmt_row = QHBoxLayout()
        fmt_row.setSpacing(10)
        self._ft_off_single_fmt = QComboBox()
        self._ft_off_single_fmt.addItems(["pdf", "docx", "odt", "xlsx", "ods", "csv", "pptx", "odp", "txt", "html"])
        self._ft_off_single_root = QComboBox()
        self._ft_off_single_root.addItems(["target", "source"])
        fmt_row.addWidget(self._ft_off_single_fmt)
        fmt_row.addWidget(self._ft_off_single_root)
        fmt_host = QWidget()
        fmt_host.setLayout(fmt_row)
        f.addRow("Format / output:", fmt_host)

        self._ft_off_single_overwrite = QCheckBox("Overwrite")
        f.addRow("", self._ft_off_single_overwrite)

        self._ft_off_single_preview = QLabel("")
        self._ft_off_single_preview.setStyleSheet("color:#9aa0a9;")
        self._ft_off_single_preview.setWordWrap(True)
        f.addRow("Preview:", self._ft_off_single_preview)

        run = QPushButton("Convert Document (Single)")
        run.setObjectName("PrimaryButton")
        run.clicked.connect(self._ft_run_office_single)
        f.addRow("", run)

        self._ft_off_single_in.textChanged.connect(lambda _t: self._ft_update_office_single_preview())
        self._ft_off_single_fmt.currentTextChanged.connect(lambda _t: self._ft_update_office_single_preview())
        self._ft_off_single_root.currentTextChanged.connect(lambda _t: self._ft_update_office_single_preview())
        self._ft_off_single_out_sub.textChanged.connect(lambda _t: self._ft_update_office_single_preview())

        outer.addWidget(box)
        outer.addStretch(1)
        return w

    def _build_file_tools_office_markdown(self) -> QWidget:
        w = QWidget()
        outer = QVBoxLayout(w)
        outer.setContentsMargins(0, 0, 0, 0)
        outer.setSpacing(12)

        box = QGroupBox("Markdown Studio")
        box_wrap = QVBoxLayout(box)
        box_wrap.setContentsMargins(12, 10, 12, 10)
        box_wrap.setSpacing(10)

        form = QGridLayout()
        form.setContentsMargins(0, 0, 0, 0)
        form.setHorizontalSpacing(12)
        form.setVerticalSpacing(10)
        form.setColumnStretch(1, 1)

        self._ft_md_name = QLineEdit("markdown_document")
        name_lbl = QLabel("File name:")
        name_lbl.setAlignment(Qt.AlignRight | Qt.AlignVCenter)
        form.addWidget(name_lbl, 0, 0)
        form.addWidget(self._ft_md_name, 0, 1)

        self._ft_md_out_sub = QLineEdit("Converted_Office")
        sub_lbl = QLabel("Output subfolder:")
        sub_lbl.setAlignment(Qt.AlignRight | Qt.AlignVCenter)
        form.addWidget(sub_lbl, 1, 0)
        form.addWidget(self._ft_md_out_sub, 1, 1)

        fmt_row = QHBoxLayout()
        fmt_row.setSpacing(10)
        self._ft_md_fmt = QComboBox()
        self._ft_md_fmt.addItems(["pdf", "docx", "pdf + docx"])
        self._ft_md_overwrite = QCheckBox("Overwrite outputs")
        fmt_row.addWidget(QLabel("Export format:"))
        fmt_row.addWidget(self._ft_md_fmt)
        fmt_row.addWidget(self._ft_md_overwrite)
        fmt_row.addStretch(1)
        fmt_host = QWidget()
        fmt_host.setLayout(fmt_row)
        form.addWidget(QLabel(""), 2, 0)
        form.addWidget(fmt_host, 2, 1)

        box_wrap.addLayout(form)

        btns = QHBoxLayout()
        btns.setSpacing(8)
        load_btn = QPushButton("Load .md/.txt")
        load_btn.clicked.connect(self._ft_pick_markdown_source)
        clear_btn = QPushButton("Clear")
        clear_btn.clicked.connect(lambda: self._ft_md_editor.setPlainText(""))
        export_btn = QPushButton("Export Markdown")
        export_btn.setObjectName("PrimaryButton")
        export_btn.clicked.connect(self._ft_run_office_markdown)
        btns.addWidget(load_btn)
        btns.addWidget(clear_btn)
        btns.addStretch(1)
        btns.addWidget(export_btn)
        box_wrap.addLayout(btns)

        self._ft_md_editor = QTextEdit()
        self._ft_md_editor.setPlaceholderText(
            "# Your Title\n\n"
            "Paste markdown here.\n\n"
            "## Section\n"
            "- Bullet one\n"
            "- Bullet two\n\n"
            "1. Numbered item\n"
            "2. Another item"
        )
        self._ft_md_editor.setMinimumHeight(320)
        box_wrap.addWidget(self._ft_md_editor, 1)

        self._ft_md_preview = QLabel("")
        self._ft_md_preview.setStyleSheet("color:#9aa0a9;")
        self._ft_md_preview.setWordWrap(True)
        box_wrap.addWidget(self._ft_md_preview)

        note = QLabel("Markdown export uses built-in formatting (no LibreOffice required).")
        note.setStyleSheet("color:#9aa0a9;")
        note.setWordWrap(True)
        box_wrap.addWidget(note)

        self._ft_md_name.textChanged.connect(lambda _t: self._ft_update_office_markdown_preview())
        self._ft_md_out_sub.textChanged.connect(lambda _t: self._ft_update_office_markdown_preview())
        self._ft_md_fmt.currentTextChanged.connect(lambda _t: self._ft_update_office_markdown_preview())

        outer.addWidget(box)
        outer.addStretch(1)
        self._ft_update_office_markdown_preview()
        return w

    def _ft_pick_markdown_source(self):
        try:
            f, _ = QFileDialog.getOpenFileName(self, "Load markdown", filter="Markdown/Text (*.md *.markdown *.txt);;All files (*.*)")
            if not f:
                return
            p = Path(f)
            text = ""
            for enc in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
                try:
                    text = p.read_text(encoding=enc)
                    break
                except Exception:
                    continue
            if text:
                self._ft_md_editor.setPlainText(text)
            stem = p.stem.strip()
            if stem:
                self._ft_md_name.setText(stem)
            self._ft_update_office_markdown_preview()
        except Exception:
            pass

    def _ft_update_office_markdown_preview(self):
        try:
            folder = (self._ft_target.text() or "").strip()
            out_sub = (self._ft_md_out_sub.text() or "Converted_Office").strip() or "Converted_Office"
            raw_name = (self._ft_md_name.text() or "markdown_document").strip() or "markdown_document"
            safe = re.sub(r"[^A-Za-z0-9_.-]+", "_", raw_name).strip("._") or "markdown_document"
            fmt = (self._ft_md_fmt.currentText() or "pdf").strip().lower()
            if not folder:
                self._ft_md_preview.setText("Set a target folder to preview output path.")
                return
            out_dir = Path(folder) / out_sub
            if fmt == "pdf + docx":
                self._ft_md_preview.setText(f"Output: {out_dir / (safe + '.pdf')} and {out_dir / (safe + '.docx')}")
            else:
                self._ft_md_preview.setText(f"Output: {out_dir / (safe + '.' + fmt)}")
        except Exception:
            self._ft_md_preview.setText("")

    def _ft_run_office_markdown(self):
        folder = (self._ft_target.text() or "").strip()
        if not folder:
            QMessageBox.warning(self, "Markdown Export", "Pick a target folder first.")
            return
        md_text = (self._ft_md_editor.toPlainText() or "").strip()
        if not md_text:
            QMessageBox.warning(self, "Markdown Export", "Paste or load markdown first.")
            return
        out_sub = (self._ft_md_out_sub.text() or "Converted_Office").strip() or "Converted_Office"
        raw_name = (self._ft_md_name.text() or "markdown_document").strip() or "markdown_document"
        safe = re.sub(r"[^A-Za-z0-9_.-]+", "_", raw_name).strip("._") or "markdown_document"

        worker = _QtMarkdownExportWorker(
            folder=folder,
            markdown_text=md_text,
            output_subfolder=out_sub,
            output_format=str(self._ft_md_fmt.currentText() or "pdf"),
            output_name=safe,
            overwrite=bool(self._ft_md_overwrite.isChecked()),
        )
        self._ft_start_worker(worker)

    def _ft_pick_office_single(self):
        try:
            f, _ = QFileDialog.getOpenFileName(self, "Select document", filter="Documents (*.pdf *.doc *.docx *.odt *.rtf *.txt *.html *.htm *.xls *.xlsx *.ods *.csv *.tsv *.ppt *.pptx *.odp);;All files (*.*)")
            if f:
                self._ft_off_single_in.setText(f)
        except Exception:
            pass

    def _ft_update_office_single_preview(self):
        try:
            inp = (self._ft_off_single_in.text() or "").strip()
            if not inp:
                self._ft_off_single_preview.setText("")
                return
            out_sub = (self._ft_off_single_out_sub.text() or "Converted_Office").strip() or "Converted_Office"
            out_fmt = (self._ft_off_single_fmt.currentText() or "pdf").strip().lower().lstrip(".")
            out_root = (self._ft_off_single_root.currentText() or "target").strip().lower()
            target = Path((self._ft_target.text() or "").strip() or Path(inp).parent)
            out_base = target if out_root == "target" else Path(inp).parent
            out_path = out_base / out_sub / (Path(inp).stem + "." + out_fmt)
            self._ft_off_single_preview.setText(f"Output: {out_path}")
        except Exception:
            self._ft_off_single_preview.setText("")

    def _ft_run_office_batch(self):
        folder = (self._ft_target.text() or "").strip()
        if not folder:
            QMessageBox.warning(self, "Office Convert", "Pick a target folder first.")
            return
        if not self._ft_require_libreoffice():
            return

        worker = _QtOfficeBatchWorker(
            folder=folder,
            source_sub=(self._ft_off_src_sub.text() or "").strip(),
            include_subfolders=bool(self._ft_off_include_sub.isChecked()),
            output_subfolder=(self._ft_off_out_sub.text() or "Converted_Office").strip() or "Converted_Office",
            output_root=str(self._ft_off_root.currentText() or "target"),
            output_format=str(self._ft_off_fmt.currentText() or "pdf"),
            overwrite=bool(self._ft_off_overwrite.isChecked()),
        )
        self._ft_start_worker(worker)

    def _ft_run_office_single(self):
        folder = (self._ft_target.text() or "").strip()
        inp = (self._ft_off_single_in.text() or "").strip()
        if not folder:
            QMessageBox.warning(self, "Office Convert", "Pick a target folder first.")
            return
        if not inp:
            QMessageBox.warning(self, "Office Convert", "Pick a document first.")
            return
        if not self._ft_require_libreoffice():
            return

        worker = _QtOfficeSingleWorker(
            folder=folder,
            input_path=inp,
            output_subfolder=(self._ft_off_single_out_sub.text() or "Converted_Office").strip() or "Converted_Office",
            output_root=str(self._ft_off_single_root.currentText() or "target"),
            output_format=str(self._ft_off_single_fmt.currentText() or "pdf"),
            overwrite=bool(self._ft_off_single_overwrite.isChecked()),
        )
        self._ft_start_worker(worker)

    def _ft_browse_target(self):
        try:
            path = QFileDialog.getExistingDirectory(self, "Select target folder")
            if path:
                self._ft_target.setText(path)
        except Exception:
            pass

    def _ft_open_ai_command(self):
        self.set_active_page("ai_command")
        try:
            self._ai_cmd_target.setText(self._ft_target.text())
        except Exception:
            pass
        try:
            self._ai_cmd_text.setFocus()
        except Exception:
            pass

    def _ft_cancel_running(self):
        try:
            w = getattr(self, "_ft_worker", None)
            if w:
                w.cancel()
        except Exception:
            pass
        self._ft_cancel.setEnabled(False)
        self._ft_status.setText("Cancelling…")

    def _ft_start_worker(self, worker: QObject):
        th = QThread(self)
        worker.moveToThread(th)
        th.started.connect(worker.run)  # type: ignore[attr-defined]
        worker.status.connect(self._ft_status.setText)  # type: ignore[attr-defined]
        worker.progress.connect(lambda p: self._ft_bar.setValue(int(max(0.0, min(1.0, p)) * 1000)))  # type: ignore[attr-defined]
        worker.finished.connect(self._ft_worker_done)  # type: ignore[attr-defined]
        worker.error.connect(self._ft_worker_failed)  # type: ignore[attr-defined]
        worker.finished.connect(th.quit)  # type: ignore[attr-defined]
        worker.error.connect(th.quit)  # type: ignore[attr-defined]
        th.finished.connect(th.deleteLater)
        th.finished.connect(worker.deleteLater)
        self._ft_worker = worker
        self._ft_thread = th
        self._ft_cancel.setEnabled(True)
        self._ft_bar.setValue(0)
        th.start()

    def _ft_worker_done(self, msg: str):
        self._ft_cancel.setEnabled(False)
        self._ft_bar.setValue(1000)
        self._ft_status.setText(msg or "Done.")
        QMessageBox.information(self, "File Tools", msg or "Done.")

    def _ft_worker_failed(self, msg: str):
        self._ft_cancel.setEnabled(False)
        self._ft_status.setText("Failed.")
        QMessageBox.critical(self, "File Tools", msg or "Failed.")

    def _build_file_tools_media_convert_page(self) -> QWidget:
        page = QWidget()
        outer = QVBoxLayout(page)
        outer.setContentsMargins(0, 0, 0, 0)
        outer.setSpacing(12)

        media_formats = [
            "mp3",
            "wav",
            "flac",
            "m4a",
            "aac",
            "ogg",
            "opus",
            "wma",
            "aiff",
            "alac",
            "mp4",
            "mkv",
            "webm",
            "mov",
            "avi",
            "wmv",
            "flv",
            "m4v",
            "mpg",
            "mpeg",
        ]
        self._ft_media_profiles = self._ft_media_load_profiles()
        self._ft_media_profile_sync = False

        splitter = QSplitter(Qt.Horizontal)
        splitter.setChildrenCollapsible(False)

        # Batch
        batch = QGroupBox("Convert Folder")
        bwrap = QVBoxLayout(batch)
        bwrap.setContentsMargins(12, 10, 12, 10)
        bwrap.setSpacing(10)
        b = QGridLayout()
        b.setContentsMargins(0, 0, 0, 0)
        b.setHorizontalSpacing(12)
        b.setVerticalSpacing(8)
        b.setColumnStretch(1, 1)
        b.setColumnStretch(3, 0)

        self._ft_media_profile = QComboBox()
        self._ft_media_profile_manage = QPushButton("Manage")
        self._ft_media_profile_manage.clicked.connect(self._ft_media_manage_profiles)
        prof_row = QHBoxLayout()
        prof_row.setSpacing(8)
        prof_row.addWidget(self._ft_media_profile, 1)
        prof_row.addWidget(self._ft_media_profile_manage)
        prof_host = QWidget()
        prof_host.setLayout(prof_row)
        lbl_prof = QLabel("Profile:")
        lbl_prof.setAlignment(Qt.AlignRight | Qt.AlignVCenter)
        b.addWidget(lbl_prof, 0, 0)
        b.addWidget(prof_host, 0, 1, 1, 3)

        self._ft_media_source = QLineEdit()
        self._ft_media_source.setPlaceholderText("Source folder (where the media are)")
        try:
            self._ft_media_source.setText(self._ft_target.text())
        except Exception:
            pass
        self._ft_media_syncing = False

        def _sync_source_to_target(text: str) -> None:
            if getattr(self, "_ft_media_syncing", False):
                return
            self._ft_media_syncing = True
            try:
                self._ft_target.setText(text)
            except Exception:
                pass
            self._ft_media_syncing = False

        def _sync_target_to_source(text: str) -> None:
            if getattr(self, "_ft_media_syncing", False):
                return
            self._ft_media_syncing = True
            try:
                self._ft_media_source.setText(text)
            except Exception:
                pass
            self._ft_media_syncing = False

        self._ft_media_source.textChanged.connect(_sync_source_to_target)
        self._ft_target.textChanged.connect(_sync_target_to_source)
        src_pick = QPushButton("Browse")
        src_pick.setIcon(self.icons.icon("folder"))
        src_pick.clicked.connect(self._ft_pick_media_source)
        src_row = QHBoxLayout()
        src_row.addWidget(self._ft_media_source, 1)
        src_row.addWidget(src_pick)
        src_host = QWidget()
        src_host.setLayout(src_row)
        lbl_source = QLabel("Source:")
        lbl_source.setAlignment(Qt.AlignRight | Qt.AlignVCenter)
        b.addWidget(lbl_source, 1, 0)
        b.addWidget(src_host, 1, 1, 1, 3)

        self._ft_media_dest = QLineEdit()
        self._ft_media_dest.setPlaceholderText("Destination folder (converted files)")
        dest_pick = QPushButton("Browse")
        dest_pick.setIcon(self.icons.icon("folder"))
        dest_pick.clicked.connect(self._ft_pick_media_dest)
        dest_row = QHBoxLayout()
        dest_row.addWidget(self._ft_media_dest, 1)
        dest_row.addWidget(dest_pick)
        dest_host = QWidget()
        dest_host.setLayout(dest_row)
        lbl_dest = QLabel("Output:")
        lbl_dest.setAlignment(Qt.AlignRight | Qt.AlignVCenter)
        b.addWidget(lbl_dest, 2, 0)
        b.addWidget(dest_host, 2, 1, 1, 3)

        self._ft_media_format = QComboBox()
        self._ft_media_format.addItems(media_formats)
        lbl_fmt = QLabel("Output format:")
        lbl_fmt.setAlignment(Qt.AlignRight | Qt.AlignVCenter)
        self._ft_media_include_sub = QCheckBox("Include subfolders")
        self._ft_media_include_sub.setChecked(True)
        b.addWidget(lbl_fmt, 3, 0)
        b.addWidget(self._ft_media_format, 3, 1)
        b.addWidget(self._ft_media_include_sub, 3, 2, 1, 2)

        self._ft_media_out_root = QComboBox()
        self._ft_media_out_root.addItems(["custom"])
        self._ft_media_out_root.setCurrentText("custom")
        self._ft_media_out_root.setVisible(False)
        self._ft_media_out_sub = QLineEdit("")
        self._ft_media_out_sub.setVisible(False)
        self._ft_media_preserve_sub = QCheckBox()
        self._ft_media_preserve_sub.setChecked(True)
        self._ft_media_preserve_sub.setVisible(False)

        self._ft_media_audio_bitrate = QComboBox()
        self._ft_media_audio_bitrate.addItems(["(auto)", "320k", "256k", "192k", "160k", "128k", "96k", "64k"])
        self._ft_media_audio_codec = QComboBox()
        self._ft_media_audio_codec.addItems(["(auto)", "aac", "mp3", "opus", "vorbis", "copy"])
        self._ft_media_codec = QComboBox()
        self._ft_media_codec.addItems(["(auto)", "h264", "h265", "vp9", "copy"])
        self._ft_media_scale = QComboBox()
        self._ft_media_scale.addItems(["(keep)", "480", "720", "1080"])
        self._ft_media_overwrite = QCheckBox("Overwrite outputs")
        self._ft_media_metadata = QCheckBox("Preserve metadata")
        self._ft_media_metadata.setChecked(True)
        self._ft_media_cover = QCheckBox("Preserve cover art")
        self._ft_media_cover.setChecked(True)

        # Advanced collapsible
        adv_btn = QToolButton()
        adv_btn.setText("Advanced ▸")
        adv_btn.setCheckable(True)
        adv_btn.setChecked(False)
        adv_btn.setCursor(Qt.PointingHandCursor)

        adv = QFrame()
        adv.setVisible(False)
        adv_l = QGridLayout(adv)
        adv_l.setContentsMargins(0, 6, 0, 0)
        adv_l.setHorizontalSpacing(12)
        adv_l.setVerticalSpacing(10)

        lbl_br = QLabel("Audio bitrate:")
        lbl_br.setAlignment(Qt.AlignRight | Qt.AlignVCenter)
        adv_l.addWidget(lbl_br, 0, 0)
        adv_l.addWidget(self._ft_media_audio_bitrate, 0, 1)

        lbl_acodec = QLabel("Audio codec:")
        lbl_acodec.setAlignment(Qt.AlignRight | Qt.AlignVCenter)
        adv_l.addWidget(lbl_acodec, 0, 2)
        adv_l.addWidget(self._ft_media_audio_codec, 0, 3)

        lbl_codec = QLabel("Video codec:")
        lbl_codec.setAlignment(Qt.AlignRight | Qt.AlignVCenter)
        adv_l.addWidget(lbl_codec, 1, 0)
        adv_l.addWidget(self._ft_media_codec, 1, 1)

        lbl_scale = QLabel("Resolution:")
        lbl_scale.setAlignment(Qt.AlignRight | Qt.AlignVCenter)
        adv_l.addWidget(lbl_scale, 1, 2)
        adv_l.addWidget(self._ft_media_scale, 1, 3)

        adv_l.addWidget(self._ft_media_metadata, 2, 0, 1, 2)
        adv_l.addWidget(self._ft_media_cover, 2, 2, 1, 2)
        adv_l.addWidget(self._ft_media_overwrite, 3, 0, 1, 2)

        run_queue = QPushButton("Add to Queue")
        run_queue.setObjectName("SecondaryButton")
        run_queue.clicked.connect(self._ft_media_queue_batch)
        run = QPushButton("Convert Folder")
        run.setIcon(self.icons.icon("export"))
        run.setIconSize(QSize(18, 18))
        run.setObjectName("PrimaryButton")
        run.clicked.connect(self._ft_run_media_batch)
        run_row = QHBoxLayout()
        run_row.addStretch(1)
        run_row.addWidget(run_queue)
        run_row.addWidget(run)
        tip = QLabel("Tip: Include subfolders mirrors the same folder structure in the destination.")
        tip.setStyleSheet("color:#9aa0a9;")

        def _toggle_adv(on: bool):
            adv.setVisible(bool(on))
            adv_btn.setText("Advanced ▾" if on else "Advanced ▸")

        adv_btn.toggled.connect(_toggle_adv)

        bwrap.addLayout(b)
        bwrap.addWidget(tip)
        bwrap.addWidget(adv_btn, 0, Qt.AlignLeft)
        bwrap.addWidget(adv)
        bwrap.addLayout(run_row)

        # Single
        single = QGroupBox("Convert Single File")
        swrap = QVBoxLayout(single)
        swrap.setContentsMargins(12, 10, 12, 10)
        swrap.setSpacing(10)
        s = QGridLayout()
        s.setContentsMargins(0, 0, 0, 0)
        s.setHorizontalSpacing(12)
        s.setVerticalSpacing(8)
        s.setColumnStretch(1, 1)
        s.setColumnStretch(3, 1)

        self._ft_media_single_profile = QComboBox()
        self._ft_media_single_profile_manage = QPushButton("Manage")
        self._ft_media_single_profile_manage.clicked.connect(self._ft_media_manage_profiles)
        sprof_row = QHBoxLayout()
        sprof_row.setSpacing(8)
        sprof_row.addWidget(self._ft_media_single_profile, 1)
        sprof_row.addWidget(self._ft_media_single_profile_manage)
        sprof_host = QWidget()
        sprof_host.setLayout(sprof_row)
        lbl_sprof = QLabel("Profile:")
        lbl_sprof.setAlignment(Qt.AlignRight | Qt.AlignVCenter)
        s.addWidget(lbl_sprof, 0, 0)
        s.addWidget(sprof_host, 0, 1, 1, 3)

        file_row = QHBoxLayout()
        self._ft_media_single_in = QLineEdit()
        self._ft_media_single_in.setPlaceholderText("Pick a media file…")
        file_row.addWidget(self._ft_media_single_in, 1)
        pick = QPushButton("Browse")
        pick.setIcon(self.icons.icon("folder"))
        pick.clicked.connect(self._ft_pick_media_file)
        file_row.addWidget(pick)
        file_host = QWidget()
        file_host.setLayout(file_row)
        lbl_in = QLabel("Source:")
        lbl_in.setAlignment(Qt.AlignRight | Qt.AlignVCenter)
        s.addWidget(lbl_in, 1, 0)
        s.addWidget(file_host, 1, 1, 1, 3)

        self._ft_media_single_dest = QLineEdit()
        self._ft_media_single_dest.setPlaceholderText("Destination folder (converted files)")
        self._ft_media_single_dest.setToolTip("Converted file will be saved here.")
        dest_pick_single = QPushButton("Browse")
        dest_pick_single.setIcon(self.icons.icon("folder"))
        dest_pick_single.clicked.connect(self._ft_pick_media_single_dest)
        dest_row_single = QHBoxLayout()
        dest_row_single.addWidget(self._ft_media_single_dest, 1)
        dest_row_single.addWidget(dest_pick_single)
        dest_host_single = QWidget()
        dest_host_single.setLayout(dest_row_single)
        lbl_dest_single = QLabel("Output:")
        lbl_dest_single.setAlignment(Qt.AlignRight | Qt.AlignVCenter)
        s.addWidget(lbl_dest_single, 2, 0)
        s.addWidget(dest_host_single, 2, 1, 1, 3)

        self._ft_media_single_fmt = QComboBox()
        self._ft_media_single_fmt.addItems(media_formats)
        lbl_of = QLabel("Output format:")
        lbl_of.setAlignment(Qt.AlignRight | Qt.AlignVCenter)
        s.addWidget(lbl_of, 3, 0)
        s.addWidget(self._ft_media_single_fmt, 3, 1)

        self._ft_media_single_name = QLineEdit("Edited_Media")
        lbl_on = QLabel("Output name:")
        lbl_on.setAlignment(Qt.AlignRight | Qt.AlignVCenter)
        name_row = QHBoxLayout()
        name_row.setSpacing(8)
        name_row.addWidget(lbl_on)
        name_row.addWidget(self._ft_media_single_name, 1)
        name_host = QWidget()
        name_host.setLayout(name_row)
        s.addWidget(name_host, 3, 2, 1, 2)

        self._ft_media_single_bitrate = QComboBox()
        self._ft_media_single_bitrate.addItems(["(auto)", "320k", "256k", "192k", "160k", "128k", "96k", "64k"])

        def _set_bitrate_options(combo: QComboBox, fmt: str, audio_codec: str | None = None) -> None:
            fmt = (fmt or "").strip().lower()
            lossy_audio = {"mp3", "m4a", "aac", "ogg", "opus", "wma"}
            video_out = {"mp4", "mkv", "webm", "mov", "avi", "wmv", "flv", "m4v", "mpg", "mpeg"}
            is_lossy = fmt in lossy_audio or fmt in video_out
            if (audio_codec or "").strip().lower() == "copy":
                is_lossy = False
            items = ["(auto)", "320k", "256k", "192k", "160k", "128k", "96k", "64k"] if is_lossy else ["(auto)"]
            prev = combo.currentText()
            if [combo.itemText(i) for i in range(combo.count())] != items:
                combo.blockSignals(True)
                combo.clear()
                combo.addItems(items)
                combo.blockSignals(False)
            if prev in items:
                combo.setCurrentText(prev)
            else:
                combo.setCurrentIndex(0)
            combo.setEnabled(is_lossy)

        def _apply_bitrate_rules() -> None:
            _set_bitrate_options(
                self._ft_media_audio_bitrate,
                self._ft_media_format.currentText(),
                self._ft_media_audio_codec.currentText() if hasattr(self, "_ft_media_audio_codec") else None,
            )
            _set_bitrate_options(
                self._ft_media_single_bitrate,
                self._ft_media_single_fmt.currentText(),
                self._ft_media_single_audio_codec.currentText() if hasattr(self, "_ft_media_single_audio_codec") else None,
            )

        self._ft_media_format.currentTextChanged.connect(lambda _val=None: _apply_bitrate_rules())
        self._ft_media_single_fmt.currentTextChanged.connect(lambda _val=None: _apply_bitrate_rules())
        try:
            self._ft_media_audio_codec.currentTextChanged.connect(lambda _val=None: _apply_bitrate_rules())
            self._ft_media_single_audio_codec.currentTextChanged.connect(lambda _val=None: _apply_bitrate_rules())
        except Exception:
            pass
        _apply_bitrate_rules()

        try:
            self._ft_media_single_dest.setEnabled(True)
            dest_pick_single.setEnabled(True)
        except Exception:
            pass

        self._ft_media_single_audio_codec = QComboBox()
        self._ft_media_single_audio_codec.addItems(["(auto)", "aac", "mp3", "opus", "vorbis", "copy"])
        self._ft_media_single_codec = QComboBox()
        self._ft_media_single_codec.addItems(["(auto)", "h264", "h265", "vp9", "copy"])
        self._ft_media_single_scale = QComboBox()
        self._ft_media_single_scale.addItems(["(keep)", "480", "720", "1080"])

        self._ft_media_single_metadata = QCheckBox("Metadata")
        self._ft_media_single_metadata.setChecked(True)
        self._ft_media_single_cover = QCheckBox("Cover Art")
        self._ft_media_single_cover.setChecked(True)
        self._ft_media_single_overwrite = QCheckBox("Overwrite output")

        adv_single = QFrame()
        adv_single.setVisible(False)
        adv_s = QGridLayout(adv_single)
        adv_s.setContentsMargins(0, 6, 0, 0)
        adv_s.setHorizontalSpacing(12)
        adv_s.setVerticalSpacing(10)

        lbl_ob = QLabel("Audio bitrate:")
        lbl_ob.setAlignment(Qt.AlignRight | Qt.AlignVCenter)
        adv_s.addWidget(lbl_ob, 0, 0)
        adv_s.addWidget(self._ft_media_single_bitrate, 0, 1)

        lbl_acodec = QLabel("Audio codec:")
        lbl_acodec.setAlignment(Qt.AlignRight | Qt.AlignVCenter)
        adv_s.addWidget(lbl_acodec, 0, 2)
        adv_s.addWidget(self._ft_media_single_audio_codec, 0, 3)

        lbl_codec = QLabel("Video codec:")
        lbl_codec.setAlignment(Qt.AlignRight | Qt.AlignVCenter)
        adv_s.addWidget(lbl_codec, 1, 0)
        adv_s.addWidget(self._ft_media_single_codec, 1, 1)

        lbl_scale = QLabel("Resolution:")
        lbl_scale.setAlignment(Qt.AlignRight | Qt.AlignVCenter)
        adv_s.addWidget(lbl_scale, 1, 2)
        adv_s.addWidget(self._ft_media_single_scale, 1, 3)

        adv_s.addWidget(self._ft_media_single_metadata, 2, 0, 1, 2)
        adv_s.addWidget(self._ft_media_single_cover, 2, 2, 1, 2)
        adv_s.addWidget(self._ft_media_single_overwrite, 3, 0, 1, 2)

        adv_btn_s = QToolButton()
        adv_btn_s.setText("Advanced ▸")
        adv_btn_s.setCheckable(True)
        adv_btn_s.setChecked(False)
        adv_btn_s.setCursor(Qt.PointingHandCursor)

        def _toggle_adv_single(on: bool):
            adv_single.setVisible(bool(on))
            adv_btn_s.setText("Advanced ▾" if on else "Advanced ▸")

        adv_btn_s.toggled.connect(_toggle_adv_single)

        run_s_queue = QPushButton("Add to Queue")
        run_s_queue.setObjectName("SecondaryButton")
        run_s_queue.clicked.connect(self._ft_media_queue_single)
        run_s = QPushButton("Convert Single")
        run_s.setIcon(self.icons.icon("export"))
        run_s.setIconSize(QSize(18, 18))
        run_s.setObjectName("PrimaryButton")
        run_s.clicked.connect(self._ft_run_media_single)
        run_s_row = QHBoxLayout()
        run_s_row.addStretch(1)
        run_s_row.addWidget(run_s_queue)
        run_s_row.addWidget(run_s)

        swrap.addLayout(s)
        swrap.addWidget(adv_btn_s, 0, Qt.AlignLeft)
        swrap.addWidget(adv_single)
        swrap.addLayout(run_s_row)
        swrap.addStretch(1)

        self._ft_media_apply_bitrate_rules = _apply_bitrate_rules
        try:
            self._ft_media_profile.currentTextChanged.connect(lambda name: self._ft_media_apply_profile(name, target="batch"))
            self._ft_media_single_profile.currentTextChanged.connect(lambda name: self._ft_media_apply_profile(name, target="single"))
        except Exception:
            pass
        self._ft_media_refresh_profile_menus()
        try:
            self._ft_media_apply_profile(self._ft_media_profile.currentText(), target="batch")
            self._ft_media_apply_profile(self._ft_media_single_profile.currentText(), target="single")
        except Exception:
            pass

        splitter.addWidget(batch)
        splitter.addWidget(single)
        splitter.setSizes([520, 520])

        outer.addWidget(splitter, 1)
        return page

    def _build_file_tools_media_page(self) -> QWidget:
        page = QWidget()
        outer = QVBoxLayout(page)
        outer.setContentsMargins(0, 0, 0, 0)
        outer.setSpacing(12)

        self._ft_media_tab_stack = QStackedWidget()
        self._ft_media_convert_page = self._build_file_tools_media_convert_page()
        self._ft_media_converting_page = self._build_file_tools_media_converting_page()
        self._ft_media_converted_page = self._build_file_tools_media_converted_page()
        self._ft_media_tab_stack.addWidget(self._ft_media_convert_page)
        self._ft_media_tab_stack.addWidget(self._ft_media_converting_page)
        self._ft_media_tab_stack.addWidget(self._ft_media_converted_page)
        self._ft_media_tab_stack.setCurrentIndex(0)

        outer.addWidget(self._ft_media_tab_stack, 1)

        self._ft_media_jobs: list[MediaJob] = []
        self._ft_media_history: list[MediaJob] = []
        self._ft_media_queue_paused = False
        self._ft_media_detail_job_id: str | None = None
        return page

    def _build_file_tools_media_converting_page(self) -> QWidget:
        page = QWidget()
        outer = QVBoxLayout(page)
        outer.setContentsMargins(0, 0, 0, 0)
        outer.setSpacing(10)

        head = QHBoxLayout()
        head.setSpacing(10)
        title = QLabel("Queue")
        title.setStyleSheet("font-weight:700;")
        head.addWidget(title)
        head.addStretch(1)

        self._ft_media_queue_start = QPushButton("Start Queue")
        self._ft_media_queue_start.setObjectName("PrimaryButton")
        self._ft_media_queue_start.clicked.connect(self._ft_media_start_queue)
        head.addWidget(self._ft_media_queue_start)

        self._ft_media_queue_pause = QPushButton("Pause")
        self._ft_media_queue_pause.setCheckable(True)
        self._ft_media_queue_pause.clicked.connect(self._ft_media_toggle_pause)
        head.addWidget(self._ft_media_queue_pause)

        self._ft_media_autostart_toggle = QCheckBox("Auto-start")
        self._ft_media_autostart_toggle.setChecked(self._ft_media_autostart_enabled())
        self._ft_media_autostart_toggle.toggled.connect(self._ft_media_set_autostart)
        head.addWidget(self._ft_media_autostart_toggle)

        self._ft_media_queue_clear = QPushButton("Clear Pending")
        self._ft_media_queue_clear.clicked.connect(self._ft_media_clear_pending)
        head.addWidget(self._ft_media_queue_clear)

        outer.addLayout(head)

        self._ft_media_converting_table = QTableWidget(0, 8)
        self._ft_media_converting_table.setHorizontalHeaderLabels(
            ["Status", "Type", "Source", "Destination", "Format", "Progress", "File", "Action"]
        )
        self._ft_media_converting_table.verticalHeader().setVisible(False)
        self._ft_media_converting_table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self._ft_media_converting_table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self._ft_media_converting_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeToContents)
        self._ft_media_converting_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeToContents)
        self._ft_media_converting_table.horizontalHeader().setSectionResizeMode(2, QHeaderView.Stretch)
        self._ft_media_converting_table.horizontalHeader().setSectionResizeMode(3, QHeaderView.Stretch)
        self._ft_media_converting_table.horizontalHeader().setSectionResizeMode(4, QHeaderView.ResizeToContents)
        self._ft_media_converting_table.horizontalHeader().setSectionResizeMode(5, QHeaderView.ResizeToContents)
        self._ft_media_converting_table.horizontalHeader().setSectionResizeMode(6, QHeaderView.Stretch)
        self._ft_media_converting_table.horizontalHeader().setSectionResizeMode(7, QHeaderView.ResizeToContents)
        self._ft_media_converting_table.itemSelectionChanged.connect(self._ft_media_show_selected_details)

        split = QSplitter(Qt.Vertical)
        split.addWidget(self._ft_media_converting_table)

        detail = QFrame()
        detail.setObjectName("PageCard")
        d = QVBoxLayout(detail)
        d.setContentsMargins(12, 10, 12, 10)
        d.setSpacing(8)

        dhead = QHBoxLayout()
        dhead.setSpacing(8)
        self._ft_media_detail_title = QLabel("Job Details")
        self._ft_media_detail_title.setStyleSheet("font-weight:700;")
        dhead.addWidget(self._ft_media_detail_title)
        dhead.addStretch(1)
        d.addLayout(dhead)

        self._ft_media_detail_hint = QLabel("Select a job to see file progress.")
        self._ft_media_detail_hint.setStyleSheet("color:#9aa0a9;")
        d.addWidget(self._ft_media_detail_hint)

        info = QFormLayout()
        info.setLabelAlignment(Qt.AlignRight | Qt.AlignVCenter)
        info.setHorizontalSpacing(10)
        info.setVerticalSpacing(6)
        self._ft_media_detail_source = QLabel("-")
        self._ft_media_detail_dest = QLabel("-")
        self._ft_media_detail_profile = QLabel("-")
        self._ft_media_detail_status = QLabel("-")
        info.addRow("Source:", self._ft_media_detail_source)
        info.addRow("Destination:", self._ft_media_detail_dest)
        info.addRow("Profile:", self._ft_media_detail_profile)
        info.addRow("Status:", self._ft_media_detail_status)
        d.addLayout(info)

        self._ft_media_detail_files = QTableWidget(0, 3)
        self._ft_media_detail_files.setHorizontalHeaderLabels(["File", "Progress", "Status"])
        self._ft_media_detail_files.verticalHeader().setVisible(False)
        self._ft_media_detail_files.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self._ft_media_detail_files.setSelectionBehavior(QAbstractItemView.SelectRows)
        self._ft_media_detail_files.horizontalHeader().setSectionResizeMode(0, QHeaderView.Stretch)
        self._ft_media_detail_files.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeToContents)
        self._ft_media_detail_files.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeToContents)
        d.addWidget(self._ft_media_detail_files, 1)

        split.addWidget(detail)
        split.setSizes([360, 220])
        outer.addWidget(split, 1)
        return page

    def _build_file_tools_media_converted_page(self) -> QWidget:
        page = QWidget()
        outer = QVBoxLayout(page)
        outer.setContentsMargins(0, 0, 0, 0)
        outer.setSpacing(10)

        head = QHBoxLayout()
        head.setSpacing(10)
        title = QLabel("Converted")
        title.setStyleSheet("font-weight:700;")
        head.addWidget(title)
        head.addStretch(1)
        self._ft_media_history_clear = QPushButton("Clear History")
        self._ft_media_history_clear.clicked.connect(self._ft_media_clear_history)
        head.addWidget(self._ft_media_history_clear)
        outer.addLayout(head)

        self._ft_media_converted_table = QTableWidget(0, 7)
        self._ft_media_converted_table.setHorizontalHeaderLabels(
            ["Status", "Type", "Source", "Destination", "Format", "Message", "Finished"]
        )
        self._ft_media_converted_table.verticalHeader().setVisible(False)
        self._ft_media_converted_table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self._ft_media_converted_table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self._ft_media_converted_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeToContents)
        self._ft_media_converted_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeToContents)
        self._ft_media_converted_table.horizontalHeader().setSectionResizeMode(2, QHeaderView.Stretch)
        self._ft_media_converted_table.horizontalHeader().setSectionResizeMode(3, QHeaderView.Stretch)
        self._ft_media_converted_table.horizontalHeader().setSectionResizeMode(4, QHeaderView.ResizeToContents)
        self._ft_media_converted_table.horizontalHeader().setSectionResizeMode(5, QHeaderView.Stretch)
        self._ft_media_converted_table.horizontalHeader().setSectionResizeMode(6, QHeaderView.ResizeToContents)
        outer.addWidget(self._ft_media_converted_table, 1)
        return page

    def _ft_media_default_profiles(self) -> list[dict]:
        return [
            {"name": "Custom (manual)", "manual": True},
            {
                "name": "Audio - MP3 320k",
                "output_format": "mp3",
                "audio_bitrate": "320k",
                "audio_codec": "mp3",
                "preserve_metadata": True,
                "preserve_cover_art": True,
            },
            {
                "name": "Audio - AAC (M4A)",
                "output_format": "m4a",
                "audio_bitrate": "256k",
                "audio_codec": "aac",
                "preserve_metadata": True,
                "preserve_cover_art": True,
            },
            {
                "name": "Audio - FLAC (Lossless)",
                "output_format": "flac",
                "audio_bitrate": None,
                "preserve_metadata": True,
                "preserve_cover_art": True,
            },
            {
                "name": "Video - H.264 1080p",
                "output_format": "mp4",
                "audio_bitrate": "192k",
                "audio_codec": "aac",
                "video_codec": "h264",
                "scale_height": 1080,
                "preserve_metadata": True,
            },
            {
                "name": "Video - H.265 720p",
                "output_format": "mp4",
                "audio_bitrate": "160k",
                "audio_codec": "aac",
                "video_codec": "h265",
                "scale_height": 720,
                "preserve_metadata": True,
            },
            {
                "name": "WebM VP9 720p",
                "output_format": "webm",
                "audio_bitrate": "160k",
                "audio_codec": "opus",
                "video_codec": "vp9",
                "scale_height": 720,
                "preserve_metadata": True,
            },
            {
                "name": "Remux (Copy Streams)",
                "output_format": "mkv",
                "audio_codec": "copy",
                "video_codec": "copy",
                "preserve_metadata": True,
                "preserve_cover_art": True,
            },
        ]

    def _ft_media_preset_packs(self) -> dict:
        return {
            "HandBrake - General": [
                {
                    "name": "Fast 1080p30",
                    "output_format": "mp4",
                    "audio_bitrate": "160k",
                    "audio_codec": "aac",
                    "video_codec": "h264",
                    "scale_height": 1080,
                    "preserve_metadata": True,
                },
                {
                    "name": "Fast 720p30",
                    "output_format": "mp4",
                    "audio_bitrate": "160k",
                    "audio_codec": "aac",
                    "video_codec": "h264",
                    "scale_height": 720,
                    "preserve_metadata": True,
                },
                {
                    "name": "HQ 1080p30",
                    "output_format": "mp4",
                    "audio_bitrate": "320k",
                    "audio_codec": "aac",
                    "video_codec": "h264",
                    "scale_height": 1080,
                    "preserve_metadata": True,
                },
                {
                    "name": "Super HQ 1080p30",
                    "output_format": "mp4",
                    "audio_bitrate": "320k",
                    "audio_codec": "aac",
                    "video_codec": "h265",
                    "scale_height": 1080,
                    "preserve_metadata": True,
                },
            ],
            "HandBrake - Web": [
                {
                    "name": "Web 1080p (VP9)",
                    "output_format": "webm",
                    "audio_bitrate": "160k",
                    "audio_codec": "opus",
                    "video_codec": "vp9",
                    "scale_height": 1080,
                    "preserve_metadata": True,
                },
                {
                    "name": "Web 720p (VP9)",
                    "output_format": "webm",
                    "audio_bitrate": "128k",
                    "audio_codec": "opus",
                    "video_codec": "vp9",
                    "scale_height": 720,
                    "preserve_metadata": True,
                },
            ],
            "HandBrake - Devices": [
                {
                    "name": "iPhone 1080p",
                    "output_format": "mp4",
                    "audio_bitrate": "160k",
                    "audio_codec": "aac",
                    "video_codec": "h264",
                    "scale_height": 1080,
                    "preserve_metadata": True,
                },
                {
                    "name": "Android 720p",
                    "output_format": "mp4",
                    "audio_bitrate": "128k",
                    "audio_codec": "aac",
                    "video_codec": "h264",
                    "scale_height": 720,
                    "preserve_metadata": True,
                },
            ],
        }

    def _ft_media_load_profiles(self) -> list[dict]:
        settings = getattr(self.backend, "settings_manager", None)
        profiles = None
        if settings:
            try:
                profiles = settings.get_setting("media_profiles", None)
            except Exception:
                profiles = None
        if not isinstance(profiles, list) or not profiles:
            profiles = self._ft_media_default_profiles()
        clean: list[dict] = []
        for prof in profiles:
            if not isinstance(prof, dict):
                continue
            name = (prof.get("name") or "").strip()
            if not name:
                continue
            clean.append(dict(prof))
        if not any((p.get("name") == "Custom (manual)") for p in clean):
            clean.insert(0, {"name": "Custom (manual)", "manual": True})
        return clean

    def _ft_media_save_profiles(self, profiles: list[dict]) -> None:
        settings = getattr(self.backend, "settings_manager", None)
        if not settings:
            return
        try:
            settings.set_setting("media_profiles", profiles)
        except Exception:
            pass

    def _ft_media_profile_by_name(self, name: str) -> dict | None:
        for prof in self._ft_media_profiles or []:
            if str(prof.get("name")) == str(name):
                return prof
        return None

    def _ft_media_refresh_profile_menus(self) -> None:
        profiles = self._ft_media_profiles or []
        names = [p.get("name") for p in profiles if p.get("name")]
        if not names:
            names = ["Custom (manual)"]
        settings = getattr(self.backend, "settings_manager", None)
        batch_last = None
        single_last = None
        if settings:
            try:
                batch_last = settings.get_setting("media_profile_last_batch", names[0])
                single_last = settings.get_setting("media_profile_last_single", names[0])
            except Exception:
                batch_last = names[0]
                single_last = names[0]
        self._ft_media_profile_sync = True
        try:
            self._ft_media_profile.blockSignals(True)
            self._ft_media_profile.clear()
            self._ft_media_profile.addItems(names)
            if batch_last in names:
                self._ft_media_profile.setCurrentText(batch_last)
        except Exception:
            pass
        try:
            self._ft_media_single_profile.blockSignals(True)
            self._ft_media_single_profile.clear()
            self._ft_media_single_profile.addItems(names)
            if single_last in names:
                self._ft_media_single_profile.setCurrentText(single_last)
        except Exception:
            pass
        try:
            self._ft_media_profile.blockSignals(False)
            self._ft_media_single_profile.blockSignals(False)
        except Exception:
            pass
        self._ft_media_profile_sync = False

    def _ft_media_apply_profile(self, name: str, *, target: str) -> None:
        if getattr(self, "_ft_media_profile_sync", False):
            return
        prof = self._ft_media_profile_by_name(name)
        if not prof or prof.get("manual"):
            return

        def _pick_combo(combo: QComboBox, value: str | None, fallback: str | None = None) -> None:
            if value:
                combo.setCurrentText(str(value))
                return
            if fallback is not None:
                combo.setCurrentText(str(fallback))

        fmt = prof.get("output_format")
        ab = prof.get("audio_bitrate")
        ac = prof.get("audio_codec")
        vc = prof.get("video_codec")
        scale = prof.get("scale_height")
        meta = prof.get("preserve_metadata")
        cover = prof.get("preserve_cover_art")
        overwrite = prof.get("overwrite")

        if target == "batch":
            _pick_combo(self._ft_media_format, fmt, self._ft_media_format.currentText())
            _pick_combo(self._ft_media_audio_bitrate, ab, self._ft_media_audio_bitrate.currentText())
            _pick_combo(self._ft_media_audio_codec, ac, self._ft_media_audio_codec.currentText())
            _pick_combo(self._ft_media_codec, vc, self._ft_media_codec.currentText())
            if scale:
                self._ft_media_scale.setCurrentText(str(scale))
            elif scale is None:
                self._ft_media_scale.setCurrentText("(keep)")
            if meta is not None:
                self._ft_media_metadata.setChecked(bool(meta))
            if cover is not None:
                self._ft_media_cover.setChecked(bool(cover))
            if overwrite is not None:
                self._ft_media_overwrite.setChecked(bool(overwrite))
        else:
            _pick_combo(self._ft_media_single_fmt, fmt, self._ft_media_single_fmt.currentText())
            _pick_combo(self._ft_media_single_bitrate, ab, self._ft_media_single_bitrate.currentText())
            _pick_combo(self._ft_media_single_audio_codec, ac, self._ft_media_single_audio_codec.currentText())
            _pick_combo(self._ft_media_single_codec, vc, self._ft_media_single_codec.currentText())
            if scale:
                self._ft_media_single_scale.setCurrentText(str(scale))
            elif scale is None:
                self._ft_media_single_scale.setCurrentText("(keep)")
            if meta is not None:
                self._ft_media_single_metadata.setChecked(bool(meta))
            if cover is not None:
                self._ft_media_single_cover.setChecked(bool(cover))
            if overwrite is not None:
                self._ft_media_single_overwrite.setChecked(bool(overwrite))

        if hasattr(self, "_ft_media_apply_bitrate_rules"):
            try:
                self._ft_media_apply_bitrate_rules()
            except Exception:
                pass

        settings = getattr(self.backend, "settings_manager", None)
        if settings:
            try:
                key = "media_profile_last_batch" if target == "batch" else "media_profile_last_single"
                settings.set_setting(key, name)
            except Exception:
                pass

    def _ft_media_manage_profiles(self) -> None:
        profiles = [dict(p) for p in (self._ft_media_profiles or [])]

        dlg = QDialog(self)
        dlg.setWindowTitle("Media Profiles")
        dlg.setModal(True)
        dlg.setMinimumWidth(820)
        dlg.setStyleSheet(_qt_modern_dialog_stylesheet())

        root = QVBoxLayout(dlg)
        root.setContentsMargins(16, 14, 16, 14)
        root.setSpacing(12)

        header = QFrame()
        header.setObjectName("DialogHeader")
        h = QVBoxLayout(header)
        h.setContentsMargins(14, 12, 14, 12)
        title = QLabel("Conversion Profiles")
        title.setObjectName("DialogTitle")
        subtitle = QLabel("Create presets for common media conversions.")
        subtitle.setObjectName("DialogSubtitle")
        h.addWidget(title)
        h.addWidget(subtitle)
        root.addWidget(header)

        body = QHBoxLayout()
        body.setSpacing(12)

        list_box = QFrame()
        list_box.setObjectName("Card")
        list_layout = QVBoxLayout(list_box)
        list_layout.setContentsMargins(12, 10, 12, 10)
        list_layout.setSpacing(8)
        prof_list = QListWidget()
        list_layout.addWidget(prof_list, 1)
        list_btns = QHBoxLayout()
        pack_btn = QToolButton()
        pack_btn.setText("Add Pack")
        pack_btn.setPopupMode(QToolButton.MenuButtonPopup)
        add_btn = QPushButton("Add")
        del_btn = QPushButton("Delete")
        list_btns.addWidget(pack_btn)
        list_btns.addWidget(add_btn)
        list_btns.addWidget(del_btn)
        list_layout.addLayout(list_btns)
        body.addWidget(list_box, 1)

        form_box = QFrame()
        form_box.setObjectName("Card")
        form_layout = QFormLayout(form_box)
        form_layout.setLabelAlignment(Qt.AlignRight | Qt.AlignVCenter)
        form_layout.setHorizontalSpacing(10)
        form_layout.setVerticalSpacing(10)

        name_edit = QLineEdit()
        fmt_combo = QComboBox()
        fmt_items: list[str] = []
        try:
            fmt_items = [self._ft_media_format.itemText(i) for i in range(self._ft_media_format.count())]
        except Exception:
            fmt_items = []
        if not fmt_items:
            fmt_items = ["mp3", "mp4", "mkv", "wav"]
        fmt_combo.addItems(fmt_items)
        bitrate_combo = QComboBox()
        bitrate_combo.addItems(["(auto)", "320k", "256k", "192k", "160k", "128k", "96k", "64k"])
        audio_codec_combo = QComboBox()
        audio_codec_combo.addItems(["(auto)", "aac", "mp3", "opus", "vorbis", "copy"])
        video_codec_combo = QComboBox()
        video_codec_combo.addItems(["(auto)", "h264", "h265", "vp9", "copy"])
        scale_combo = QComboBox()
        scale_combo.addItems(["(keep)", "480", "720", "1080"])
        remux_cb = QCheckBox("Remux (copy streams)")
        meta_cb = QCheckBox("Preserve metadata")
        cover_cb = QCheckBox("Preserve cover art")
        overwrite_cb = QCheckBox("Overwrite output")

        form_layout.addRow("Name:", name_edit)
        form_layout.addRow("Output format:", fmt_combo)
        form_layout.addRow("Audio bitrate:", bitrate_combo)
        form_layout.addRow("Audio codec:", audio_codec_combo)
        form_layout.addRow("Video codec:", video_codec_combo)
        form_layout.addRow("Scale height:", scale_combo)
        form_layout.addRow("", remux_cb)
        form_layout.addRow("", meta_cb)
        form_layout.addRow("", cover_cb)
        form_layout.addRow("", overwrite_cb)

        body.addWidget(form_box, 2)
        root.addLayout(body, 1)

        btn_row = QHBoxLayout()
        btn_row.addStretch(1)
        save_btn = QPushButton("Save Profiles")
        save_btn.setObjectName("PrimaryButton")
        close_btn = QPushButton("Close")
        btn_row.addWidget(save_btn)
        btn_row.addWidget(close_btn)
        root.addLayout(btn_row)

        def refresh_list(selected: int | None = None):
            prof_list.clear()
            for p in profiles:
                prof_list.addItem(str(p.get("name", "Profile")))
            if selected is not None and 0 <= selected < prof_list.count():
                prof_list.setCurrentRow(selected)
            elif prof_list.count() > 0:
                prof_list.setCurrentRow(0)

        def apply_remux_state(on: bool):
            if on:
                video_codec_combo.setCurrentText("copy")
                audio_codec_combo.setCurrentText("copy")
                bitrate_combo.setCurrentText("(auto)")
                scale_combo.setCurrentText("(keep)")
            video_codec_combo.setEnabled(not on)
            audio_codec_combo.setEnabled(not on)
            bitrate_combo.setEnabled(not on)
            scale_combo.setEnabled(not on)

        def load_profile(idx: int):
            if idx < 0 or idx >= len(profiles):
                return
            p = profiles[idx]
            name_edit.setText(str(p.get("name", "")))
            fmt_combo.setCurrentText(str(p.get("output_format", "mp4")))
            bitrate_combo.setCurrentText(str(p.get("audio_bitrate") or "(auto)"))
            audio_codec_combo.setCurrentText(str(p.get("audio_codec") or "(auto)"))
            video_codec_combo.setCurrentText(str(p.get("video_codec") or "(auto)"))
            scale_val = p.get("scale_height")
            scale_combo.setCurrentText(str(scale_val) if scale_val else "(keep)")
            meta_cb.setChecked(bool(p.get("preserve_metadata", True)))
            cover_cb.setChecked(bool(p.get("preserve_cover_art", True)))
            overwrite_cb.setChecked(bool(p.get("overwrite", False)))
            remux_cb.setChecked(bool(p.get("video_codec") == "copy" and p.get("audio_codec") == "copy"))
            apply_remux_state(remux_cb.isChecked())

        def store_profile(idx: int):
            if idx < 0 or idx >= len(profiles):
                return
            p = profiles[idx]
            p["name"] = (name_edit.text() or "Profile").strip()
            p["output_format"] = str(fmt_combo.currentText() or "").strip()
            p["audio_bitrate"] = None if bitrate_combo.currentText().startswith("(") else bitrate_combo.currentText()
            p["audio_codec"] = None if audio_codec_combo.currentText().startswith("(") else audio_codec_combo.currentText()
            p["video_codec"] = None if video_codec_combo.currentText().startswith("(") else video_codec_combo.currentText()
            p["scale_height"] = None if scale_combo.currentText().startswith("(") else int(scale_combo.currentText())
            p["preserve_metadata"] = bool(meta_cb.isChecked())
            p["preserve_cover_art"] = bool(cover_cb.isChecked())
            p["overwrite"] = bool(overwrite_cb.isChecked())
            if remux_cb.isChecked():
                p["audio_codec"] = "copy"
                p["video_codec"] = "copy"

        prev_idx = {"v": -1}

        def on_select(idx: int):
            if prev_idx["v"] >= 0:
                store_profile(prev_idx["v"])
            prev_idx["v"] = idx
            load_profile(idx)

        def on_add():
            profiles.append({"name": "New Profile"})
            refresh_list(len(profiles) - 1)

        def on_delete():
            idx = prof_list.currentRow()
            if idx < 0:
                return
            profiles.pop(idx)
            refresh_list(max(0, idx - 1))

        def add_pack(pack_name: str):
            idx = prof_list.currentRow()
            if idx >= 0:
                store_profile(idx)
            packs = self._ft_media_preset_packs()
            items = packs.get(pack_name, [])
            if not items:
                return
            existing = {str(p.get("name")) for p in profiles if p.get("name")}
            for item in items:
                base = str(item.get("name") or "Preset")
                name = base
                suffix = 2
                while name in existing:
                    name = f"{base} ({suffix})"
                    suffix += 1
                new_item = dict(item)
                new_item["name"] = name
                profiles.append(new_item)
                existing.add(name)
            refresh_list(len(profiles) - len(items))

        def on_save():
            idx = prof_list.currentRow()
            if idx >= 0:
                store_profile(idx)
            self._ft_media_profiles = profiles or self._ft_media_default_profiles()
            self._ft_media_save_profiles(self._ft_media_profiles)
            self._ft_media_refresh_profile_menus()
            try:
                self._ft_media_apply_profile(self._ft_media_profile.currentText(), target="batch")
                self._ft_media_apply_profile(self._ft_media_single_profile.currentText(), target="single")
            except Exception:
                pass
            dlg.accept()

        remux_cb.toggled.connect(apply_remux_state)
        prof_list.currentRowChanged.connect(on_select)
        add_btn.clicked.connect(on_add)
        del_btn.clicked.connect(on_delete)
        pack_menu = QMenu(pack_btn)
        for pack_name in self._ft_media_preset_packs().keys():
            act = pack_menu.addAction(pack_name)
            act.triggered.connect(lambda _=False, n=pack_name: add_pack(n))
        pack_btn.setMenu(pack_menu)
        save_btn.clicked.connect(on_save)
        close_btn.clicked.connect(dlg.reject)

        refresh_list(0)
        dlg.exec()

    def _ft_media_start_queue(self):
        self._ft_media_queue_paused = False
        try:
            self._ft_media_queue_pause.setChecked(False)
        except Exception:
            pass
        self._ft_media_maybe_start_jobs()

    def _ft_media_toggle_pause(self):
        self._ft_media_queue_paused = bool(self._ft_media_queue_pause.isChecked())
        if not self._ft_media_queue_paused:
            self._ft_media_maybe_start_jobs()

    def _ft_media_clear_pending(self):
        pending = [j for j in self._ft_media_jobs if j.status == "queued"]
        for job in pending:
            self._ft_media_remove_converting_row(job.job_id)
            self._ft_media_jobs.remove(job)
        if self._ft_media_detail_job_id and not self._ft_media_find_job(self._ft_media_detail_job_id):
            self._ft_media_detail_job_id = None
            self._ft_media_clear_details()

    def _ft_media_clear_history(self):
        self._ft_media_history.clear()
        self._ft_media_converted_table.setRowCount(0)

    def _ft_media_queue_batch(self):
        self._ft_media_queue_batch_job(start_immediately=False)

    def _ft_media_queue_single(self):
        self._ft_media_queue_single_job(start_immediately=False)

    def _ft_media_queue_batch_job(self, *, start_immediately: bool):
        job = self._ft_media_build_batch_job()
        if not job:
            return
        self._ft_media_add_job(job, start_immediately=start_immediately)

    def _ft_media_queue_single_job(self, *, start_immediately: bool):
        job = self._ft_media_build_single_job()
        if not job:
            return
        self._ft_media_add_job(job, start_immediately=start_immediately)

    def _ft_media_add_job(self, job: MediaJob, *, start_immediately: bool):
        self._ft_media_jobs.append(job)
        self._ft_media_add_converting_row(job)
        if start_immediately or self._ft_media_autostart_enabled():
            self._ft_media_maybe_start_jobs()

    def _ft_media_autostart_enabled(self) -> bool:
        settings = getattr(self.backend, "settings_manager", None)
        if settings:
            try:
                return bool(settings.get_setting("media_queue_autostart", False))
            except Exception:
                return False
        return False

    def _ft_media_set_autostart(self, on: bool):
        settings = getattr(self.backend, "settings_manager", None)
        if settings:
            try:
                settings.set_setting("media_queue_autostart", bool(on))
            except Exception:
                pass
        if on:
            self._ft_media_maybe_start_jobs()

    def _ft_media_max_parallel(self) -> int:
        settings = getattr(self.backend, "settings_manager", None)
        if settings:
            try:
                return max(1, int(settings.get_setting("media_max_parallel", 1) or 1))
            except Exception:
                return 1
        return 1

    def _ft_media_maybe_start_jobs(self):
        if self._ft_media_queue_paused:
            return
        max_parallel = self._ft_media_max_parallel()
        running = [j for j in self._ft_media_jobs if j.status == "running"]
        pending = [j for j in self._ft_media_jobs if j.status == "queued"]
        while len(running) < max_parallel and pending:
            job = pending.pop(0)
            self._ft_media_start_job(job)
            running.append(job)

    def _ft_media_start_job(self, job: MediaJob):
        job.status = "running"
        self._ft_media_update_converting_row(job)

        th = QThread(self)
        worker = None
        if job.kind == "batch":
            worker = _QtMediaBatchWorker(**job.settings)
        else:
            worker = _QtMediaSingleWorker(**job.settings)

        worker.moveToThread(th)
        th.started.connect(worker.run)  # type: ignore[attr-defined]
        worker.status.connect(lambda msg, jid=job.job_id: self._ft_media_job_status(jid, msg))  # type: ignore[attr-defined]
        worker.progress.connect(lambda p, jid=job.job_id: self._ft_media_job_progress(jid, p))  # type: ignore[attr-defined]
        try:
            worker.file_progress.connect(lambda p, f, jid=job.job_id: self._ft_media_job_file_progress(jid, p, f))  # type: ignore[attr-defined]
        except Exception:
            pass
        worker.finished.connect(lambda msg, jid=job.job_id: self._ft_media_job_done(jid, True, msg))  # type: ignore[attr-defined]
        worker.error.connect(lambda msg, jid=job.job_id: self._ft_media_job_done(jid, False, msg))  # type: ignore[attr-defined]
        worker.finished.connect(th.quit)  # type: ignore[attr-defined]
        worker.error.connect(th.quit)  # type: ignore[attr-defined]
        th.finished.connect(th.deleteLater)
        th.finished.connect(worker.deleteLater)
        job.thread = th
        job.worker = worker
        th.start()

    def _ft_media_job_status(self, job_id: str, msg: str):
        job = self._ft_media_find_job(job_id)
        if not job:
            return
        job.message = msg
        if not job.current_file:
            self._ft_media_update_converting_row(job)
        if self._ft_media_detail_job_id == job.job_id:
            self._ft_media_render_details(job)

    def _ft_media_job_progress(self, job_id: str, frac: float):
        job = self._ft_media_find_job(job_id)
        if not job:
            return
        job.progress = max(0.0, min(1.0, float(frac)))
        self._ft_media_update_converting_row(job)
        if self._ft_media_detail_job_id == job.job_id:
            self._ft_media_render_details(job)

    def _ft_media_job_file_progress(self, job_id: str, path: str, frac: float):
        job = self._ft_media_find_job(job_id)
        if not job:
            return
        frac = max(0.0, min(1.0, float(frac)))
        name = Path(path).name if path else ""
        pct = int(frac * 100)
        job.current_file = f"{name} ({pct}%)" if name else ""
        key = str(path) if path else name
        if key:
            entry = job.file_entries.get(key)
            if entry is None:
                job.file_entries[key] = {"progress": frac, "status": "processing"}
                job.file_order.append(key)
                if len(job.file_order) > 200:
                    drop = job.file_order.pop(0)
                    job.file_entries.pop(drop, None)
            else:
                entry["progress"] = frac
                entry["status"] = "done" if frac >= 1.0 else "processing"
        job.message = ""
        self._ft_media_update_converting_row(job)
        if self._ft_media_detail_job_id == job.job_id:
            self._ft_media_refresh_detail_files(job)

    def _ft_media_job_done(self, job_id: str, ok: bool, msg: str):
        job = self._ft_media_find_job(job_id)
        if not job:
            return
        if ok:
            job.status = "done"
        else:
            job.status = "cancelled" if "cancel" in (msg or "").lower() else "failed"
        job.message = msg
        job.finished_at = QDateTime.currentDateTime()
        if job.file_entries:
            final_status = "done" if job.status == "done" else job.status
            for entry in job.file_entries.values():
                entry["status"] = final_status
        self._ft_media_remove_converting_row(job.job_id)
        self._ft_media_add_converted_row(job)
        if job in self._ft_media_jobs:
            self._ft_media_jobs.remove(job)
        self._ft_media_history.append(job)
        if self._ft_media_detail_job_id == job.job_id:
            self._ft_media_detail_job_id = None
            self._ft_media_clear_details()
        else:
            self._ft_media_show_selected_details()
        self._ft_media_maybe_start_jobs()

    def _ft_media_cancel_job(self, job_id: str):
        job = self._ft_media_find_job(job_id)
        if not job:
            return
        if job.status == "queued":
            self._ft_media_remove_converting_row(job.job_id)
            if job in self._ft_media_jobs:
                self._ft_media_jobs.remove(job)
            if self._ft_media_detail_job_id == job.job_id:
                self._ft_media_detail_job_id = None
                self._ft_media_clear_details()
            return
        if job.worker:
            try:
                job.worker.cancel()
            except Exception:
                pass

    def _ft_media_find_job(self, job_id: str) -> MediaJob | None:
        for job in self._ft_media_jobs:
            if job.job_id == job_id:
                return job
        return None

    def _ft_media_find_row(self, table: QTableWidget, job_id: str) -> int:
        for row in range(table.rowCount()):
            item = table.item(row, 0)
            if item and item.data(Qt.UserRole) == job_id:
                return row
        return -1

    def _ft_media_add_converting_row(self, job: MediaJob):
        row = self._ft_media_converting_table.rowCount()
        self._ft_media_converting_table.insertRow(row)
        status_item = QTableWidgetItem(job.status.capitalize())
        status_item.setData(Qt.UserRole, job.job_id)
        self._ft_media_converting_table.setItem(row, 0, status_item)
        self._ft_media_converting_table.setItem(row, 1, QTableWidgetItem(job.kind.capitalize()))
        self._ft_media_converting_table.setItem(row, 2, QTableWidgetItem(job.source))
        self._ft_media_converting_table.setItem(row, 3, QTableWidgetItem(job.dest))
        self._ft_media_converting_table.setItem(row, 4, QTableWidgetItem(job.fmt))
        prog = QProgressBar()
        prog.setRange(0, 100)
        prog.setValue(int(job.progress * 100))
        self._ft_media_converting_table.setCellWidget(row, 5, prog)
        self._ft_media_converting_table.setItem(row, 6, QTableWidgetItem(job.current_file or job.message or ""))
        action = QPushButton("Remove" if job.status == "queued" else "Cancel")
        action.clicked.connect(lambda _=None, jid=job.job_id: self._ft_media_cancel_job(jid))
        self._ft_media_converting_table.setCellWidget(row, 7, action)

    def _ft_media_update_converting_row(self, job: MediaJob):
        row = self._ft_media_find_row(self._ft_media_converting_table, job.job_id)
        if row < 0:
            return
        status_item = self._ft_media_converting_table.item(row, 0)
        if status_item:
            status_item.setText(job.status.capitalize())
        prog = self._ft_media_converting_table.cellWidget(row, 5)
        if isinstance(prog, QProgressBar):
            prog.setValue(int(job.progress * 100))
        file_item = self._ft_media_converting_table.item(row, 6)
        if file_item:
            file_item.setText(job.current_file or job.message or "")
        action = self._ft_media_converting_table.cellWidget(row, 7)
        if isinstance(action, QPushButton):
            action.setText("Cancel" if job.status == "running" else "Remove")

    def _ft_media_remove_converting_row(self, job_id: str):
        row = self._ft_media_find_row(self._ft_media_converting_table, job_id)
        if row >= 0:
            self._ft_media_converting_table.removeRow(row)

    def _ft_media_add_converted_row(self, job: MediaJob):
        row = self._ft_media_converted_table.rowCount()
        self._ft_media_converted_table.insertRow(row)
        self._ft_media_converted_table.setItem(row, 0, QTableWidgetItem(job.status.capitalize()))
        self._ft_media_converted_table.setItem(row, 1, QTableWidgetItem(job.kind.capitalize()))
        self._ft_media_converted_table.setItem(row, 2, QTableWidgetItem(job.source))
        self._ft_media_converted_table.setItem(row, 3, QTableWidgetItem(job.dest))
        self._ft_media_converted_table.setItem(row, 4, QTableWidgetItem(job.fmt))
        self._ft_media_converted_table.setItem(row, 5, QTableWidgetItem(job.message or ""))
        finished = job.finished_at.toString("yyyy-MM-dd HH:mm:ss") if job.finished_at else ""
        self._ft_media_converted_table.setItem(row, 6, QTableWidgetItem(finished))

    def _ft_media_show_selected_details(self):
        try:
            rows = self._ft_media_converting_table.selectionModel().selectedRows()
        except Exception:
            rows = []
        if not rows:
            self._ft_media_detail_job_id = None
            self._ft_media_clear_details()
            return
        row = rows[0].row()
        item = self._ft_media_converting_table.item(row, 0)
        job_id = item.data(Qt.UserRole) if item else None
        job = self._ft_media_find_job(str(job_id)) if job_id else None
        if not job:
            self._ft_media_detail_job_id = None
            self._ft_media_clear_details()
            return
        self._ft_media_detail_job_id = job.job_id
        self._ft_media_render_details(job)

    def _ft_media_clear_details(self):
        try:
            self._ft_media_detail_hint.setVisible(True)
        except Exception:
            pass
        try:
            self._ft_media_detail_source.setText("-")
            self._ft_media_detail_dest.setText("-")
            self._ft_media_detail_profile.setText("-")
            self._ft_media_detail_status.setText("-")
            self._ft_media_detail_files.setRowCount(0)
        except Exception:
            pass

    def _ft_media_render_details(self, job: MediaJob):
        try:
            self._ft_media_detail_hint.setVisible(False)
        except Exception:
            pass
        try:
            self._ft_media_detail_source.setText(job.source or "-")
            self._ft_media_detail_dest.setText(job.dest or "-")
            self._ft_media_detail_profile.setText(job.profile or "Custom")
            status_txt = job.status.capitalize()
            if job.message:
                status_txt = f"{status_txt} - {job.message}"
            self._ft_media_detail_status.setText(status_txt)
        except Exception:
            pass
        self._ft_media_refresh_detail_files(job)

    def _ft_media_refresh_detail_files(self, job: MediaJob):
        try:
            self._ft_media_detail_files.setRowCount(0)
        except Exception:
            return
        if not job.file_order:
            return
        for p in job.file_order:
            entry = job.file_entries.get(p, {})
            row = self._ft_media_detail_files.rowCount()
            self._ft_media_detail_files.insertRow(row)
            self._ft_media_detail_files.setItem(row, 0, QTableWidgetItem(Path(p).name))
            prog = QProgressBar()
            prog.setRange(0, 100)
            prog.setValue(int(entry.get("progress", 0.0) * 100))
            self._ft_media_detail_files.setCellWidget(row, 1, prog)
            self._ft_media_detail_files.setItem(row, 2, QTableWidgetItem(str(entry.get("status", ""))))

    def _ft_media_collect_candidates(
        self,
        folder: Path,
        *,
        include_subfolders: bool,
        output_format: str,
        destination: str,
        limit: int = 200,
    ) -> list[str]:
        exts = {
            ".mp4",
            ".mkv",
            ".mov",
            ".avi",
            ".webm",
            ".wmv",
            ".flv",
            ".m4v",
            ".mpg",
            ".mpeg",
            ".mp3",
            ".wav",
            ".m4a",
            ".aac",
            ".flac",
            ".ogg",
            ".opus",
            ".wma",
            ".aiff",
            ".alac",
        }
        out_ext = (output_format or "").strip().lower().lstrip(".")
        if out_ext:
            exts.discard("." + out_ext)

        dest_path = Path(destination) if destination else None
        if dest_path and not dest_path.is_absolute():
            dest_path = folder / dest_path
        dest_resolved = None
        if dest_path:
            try:
                dest_resolved = dest_path.resolve()
            except Exception:
                dest_resolved = dest_path

        files: list[str] = []
        it = folder.rglob("*") if include_subfolders else folder.glob("*")
        for p in it:
            if len(files) >= limit:
                break
            try:
                if not p.is_file():
                    continue
            except Exception:
                continue
            if p.suffix.lower() not in exts:
                continue
            if dest_resolved:
                try:
                    if dest_resolved in p.resolve().parents:
                        continue
                except Exception:
                    pass
            files.append(str(p))
        return files

    def _ft_media_build_batch_job(self) -> MediaJob | None:
        folder = (self._ft_media_source.text() or self._ft_target.text() or "").strip()
        if not folder:
            QMessageBox.warning(self, "Media Convert", "Pick a source folder first.")
            return None
        dest = (self._ft_media_dest.text() or "").strip()
        if not dest:
            QMessageBox.warning(self, "Media Convert", "Pick a destination folder.")
            return None
        include_sub = bool(self._ft_media_include_sub.isChecked())
        settings = getattr(self.backend, "settings_manager", None)
        use_gpu = bool(settings.get_setting("media_use_gpu", False)) if settings else False
        audio_codec = None if self._ft_media_audio_codec.currentText().startswith("(") else self._ft_media_audio_codec.currentText()
        audio_bitrate = None if self._ft_media_audio_bitrate.currentText().startswith("(") else self._ft_media_audio_bitrate.currentText()
        if audio_codec == "copy":
            audio_bitrate = None
        video_codec = None if self._ft_media_codec.currentText().startswith("(") else self._ft_media_codec.currentText()
        scale_height = None if self._ft_media_scale.currentText().startswith("(") else int(self._ft_media_scale.currentText())

        job = MediaJob(
            job_id=str(uuid.uuid4()),
            kind="batch",
            title="Convert Folder",
            source=folder,
            dest=dest,
            fmt=str(self._ft_media_format.currentText()),
            profile=str(self._ft_media_profile.currentText()),
            status="queued",
            settings=dict(
                folder=folder,
                source_subfolder=None,
                include_subfolders=include_sub,
                output_format=str(self._ft_media_format.currentText()),
                output_subfolder="",
                output_root="custom",
                output_directory=dest,
                preserve_subfolders=include_sub,
                overwrite=bool(self._ft_media_overwrite.isChecked()),
                audio_bitrate=audio_bitrate,
                preserve_metadata=bool(self._ft_media_metadata.isChecked()),
                preserve_cover_art=bool(self._ft_media_cover.isChecked()),
                video_codec=video_codec,
                scale_height=scale_height,
                use_gpu=use_gpu,
                audio_codec=audio_codec,
            ),
        )
        try:
            candidates = self._ft_media_collect_candidates(
                Path(folder),
                include_subfolders=include_sub,
                output_format=str(self._ft_media_format.currentText()),
                destination=dest,
                limit=200,
            )
            for path in candidates:
                job.file_entries[path] = {"progress": 0.0, "status": "queued"}
                job.file_order.append(path)
        except Exception:
            pass
        return job

    def _ft_media_build_single_job(self) -> MediaJob | None:
        folder = (self._ft_target.text() or "").strip()
        inp = (self._ft_media_single_in.text() or "").strip()
        if not inp:
            QMessageBox.warning(self, "Media Convert", "Pick an input file.")
            return None
        out_fmt = str(self._ft_media_single_fmt.currentText())
        out_name = (self._ft_media_single_name.text() or "Edited_Media").strip()
        dest = (self._ft_media_single_dest.text() or "").strip()
        if not dest:
            QMessageBox.warning(self, "Media Convert", "Pick a destination folder.")
            return None
        settings = getattr(self.backend, "settings_manager", None)
        use_gpu = bool(settings.get_setting("media_use_gpu", False)) if settings else False
        audio_codec = None if self._ft_media_single_audio_codec.currentText().startswith("(") else self._ft_media_single_audio_codec.currentText()
        audio_bitrate = None if self._ft_media_single_bitrate.currentText().startswith("(") else self._ft_media_single_bitrate.currentText()
        if audio_codec == "copy":
            audio_bitrate = None
        video_codec = None if self._ft_media_single_codec.currentText().startswith("(") else self._ft_media_single_codec.currentText()
        scale_height = None if self._ft_media_single_scale.currentText().startswith("(") else int(self._ft_media_single_scale.currentText())

        job = MediaJob(
            job_id=str(uuid.uuid4()),
            kind="single",
            title="Convert Single",
            source=inp,
            dest=dest,
            fmt=out_fmt,
            profile=str(self._ft_media_single_profile.currentText()),
            status="queued",
            settings=dict(
                target_folder=folder,
                input_path=inp,
                output_format=out_fmt,
                output_name=out_name,
                output_root="custom",
                output_directory=dest,
                overwrite=bool(self._ft_media_single_overwrite.isChecked()),
                audio_bitrate=audio_bitrate,
                preserve_metadata=bool(self._ft_media_single_metadata.isChecked()),
                preserve_cover_art=bool(self._ft_media_single_cover.isChecked()),
                use_gpu=use_gpu,
                video_codec=video_codec,
                scale_height=scale_height,
                audio_codec=audio_codec,
            ),
        )
        job.file_entries[str(inp)] = {"progress": 0.0, "status": "queued"}
        job.file_order.append(str(inp))
        return job

    def _ft_pick_media_file(self):
        try:
            f, _ = QFileDialog.getOpenFileName(self, "Select media file")
            if f:
                self._ft_media_single_in.setText(f)
                try:
                    if not (self._ft_media_single_dest.text() or "").strip():
                        self._ft_media_single_dest.setText(str(Path(f).parent))
                except Exception:
                    pass
        except Exception:
            pass

    def _ft_pick_media_source(self):
        try:
            d = QFileDialog.getExistingDirectory(self, "Select source folder")
            if d:
                self._ft_media_source.setText(d)
        except Exception:
            pass

    def _ft_pick_media_dest(self):
        try:
            d = QFileDialog.getExistingDirectory(self, "Select destination folder")
            if d:
                self._ft_media_dest.setText(d)
        except Exception:
            pass

    def _ft_pick_media_single_dest(self):
        try:
            d = QFileDialog.getExistingDirectory(self, "Select destination folder")
            if d:
                self._ft_media_single_dest.setText(d)
        except Exception:
            pass

    def _ft_run_media_batch(self):
        self._ft_media_queue_batch_job(start_immediately=True)

    def _ft_run_media_single(self):
        self._ft_media_queue_single_job(start_immediately=True)

    def _build_file_tools_images_page(self) -> QWidget:
        page = QWidget()
        outer = QVBoxLayout(page)
        outer.setContentsMargins(0, 0, 0, 0)
        outer.setSpacing(12)

        box = QGroupBox("Convert Images")
        f = QFormLayout(box)
        f.setLabelAlignment(Qt.AlignRight)
        f.setVerticalSpacing(10)
        f.setHorizontalSpacing(12)

        self._ft_img_include_sub = QCheckBox("Include subfolders")
        self._ft_img_include_sub.setChecked(True)
        f.addRow("", self._ft_img_include_sub)

        self._ft_img_format = QComboBox()
        self._ft_img_format.addItems(["webp", "png", "jpg"])
        f.addRow("Output format:", self._ft_img_format)

        self._ft_img_sub = QLineEdit("Converted_Images")
        f.addRow("Output subfolder:", self._ft_img_sub)

        self._ft_img_overwrite = QCheckBox("Overwrite outputs")
        f.addRow("", self._ft_img_overwrite)

        run = QPushButton("Convert Images")
        run.setObjectName("PrimaryButton")
        run.clicked.connect(self._ft_run_images)
        f.addRow("", run)

        outer.addWidget(box)
        outer.addStretch(1)
        return page

    def _ft_run_images(self):
        folder = (self._ft_target.text() or "").strip()
        if not folder:
            QMessageBox.warning(self, "Image Convert", "Pick a target folder first.")
            return
        worker = _QtImagesWorker(
            folder=folder,
            include_subfolders=bool(self._ft_img_include_sub.isChecked()),
            output_format=str(self._ft_img_format.currentText()),
            output_subfolder=(self._ft_img_sub.text() or "Converted_Images").strip(),
            overwrite=bool(self._ft_img_overwrite.isChecked()),
        )
        self._ft_start_worker(worker)

    def _build_file_tools_ocr_page(self) -> QWidget:
        page = QWidget()
        outer = QVBoxLayout(page)
        outer.setContentsMargins(0, 0, 0, 0)
        outer.setSpacing(12)

        split = QSplitter(Qt.Horizontal)
        split.setChildrenCollapsible(False)

        # Left: inputs + options (scrollable so options never clip)
        left = QGroupBox("Text Extractor")
        left.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)
        lv = QVBoxLayout(left)
        lv.setContentsMargins(12, 10, 12, 10)
        lv.setSpacing(10)

        drop = _QtDropZone("Drop PDFs, images, or office files here.")
        drop.setFixedHeight(90)
        drop.files_dropped.connect(self._ft_ocr_add_files)
        lv.addWidget(drop)

        btn_row = QHBoxLayout()
        btn_row.setSpacing(8)
        add_btn = QPushButton("Add Files")
        add_btn.clicked.connect(self._ft_ocr_pick_files)
        btn_row.addWidget(add_btn)
        rm_btn = QPushButton("Remove")
        rm_btn.clicked.connect(self._ft_ocr_remove_selected)
        btn_row.addWidget(rm_btn)
        clr_btn = QPushButton("Clear")
        clr_btn.clicked.connect(self._ft_ocr_clear_files)
        btn_row.addWidget(clr_btn)
        btn_row.addStretch(1)
        lv.addLayout(btn_row)

        self._ft_ocr_list = QListWidget()
        self._ft_ocr_list.setMinimumHeight(140)
        lv.addWidget(self._ft_ocr_list, 1)

        opts = QGroupBox("Options")
        f = QFormLayout(opts)
        f.setLabelAlignment(Qt.AlignRight)
        f.setVerticalSpacing(10)
        f.setHorizontalSpacing(12)

        self._ft_ocr_mode = QComboBox()
        self._ft_ocr_mode.addItems(["Auto (text + OCR if needed)", "Force OCR", "Text only"])
        f.addRow("OCR mode:", self._ft_ocr_mode)

        self._ft_ocr_engine = QComboBox()
        self._ft_ocr_engine.addItems(["Auto", "RapidOCR", "AI Vision"])
        f.addRow("OCR engine:", self._ft_ocr_engine)

        self._ft_ocr_lang = QLineEdit("eng")
        self._ft_ocr_lang.setToolTip("Language hint for AI OCR (e.g., eng, spa, deu).")
        f.addRow("OCR language:", self._ft_ocr_lang)

        self._ft_ocr_pages = QSpinBox()
        self._ft_ocr_pages.setRange(0, 5000)
        self._ft_ocr_pages.setValue(5)
        self._ft_ocr_pages.setToolTip("Max PDF pages to OCR. 0 = all pages.")
        f.addRow("OCR pages:", self._ft_ocr_pages)

        self._ft_ocr_format = QComboBox()
        self._ft_ocr_format.addItems(["Plain text (.txt)", "Markdown (.md)", "DOCX (.docx)"])
        f.addRow("Output format:", self._ft_ocr_format)

        out_row = QHBoxLayout()
        self._ft_ocr_out = QLineEdit()
        self._ft_ocr_out.setPlaceholderText("Destination folder (optional)")
        out_row.addWidget(self._ft_ocr_out, 1)
        out_btn = QPushButton("Browse")
        out_btn.clicked.connect(self._ft_ocr_pick_output_dir)
        out_row.addWidget(out_btn)
        out_host = QWidget()
        out_host.setLayout(out_row)
        f.addRow("Output folder:", out_host)

        self._ft_ocr_normalize = QCheckBox("Clean spacing / remove extra blank lines")
        self._ft_ocr_normalize.setChecked(True)
        f.addRow("", self._ft_ocr_normalize)

        self._ft_ocr_headers = QCheckBox("Add filename headers in output")
        self._ft_ocr_headers.setChecked(True)
        f.addRow("", self._ft_ocr_headers)

        self._ft_ocr_ai = QCheckBox("Enhance with AI (cleanup/structure)")
        f.addRow("", self._ft_ocr_ai)

        tr_row = QHBoxLayout()
        self._ft_ocr_translate = QCheckBox("Translate (AI)")
        tr_row.addWidget(self._ft_ocr_translate)
        self._ft_ocr_translate_lang = QComboBox()
        self._ft_ocr_translate_lang.addItems(
            ["English", "Spanish", "French", "German", "Italian", "Portuguese", "Japanese", "Korean", "Chinese"]
        )
        self._ft_ocr_translate_lang.setEnabled(False)
        self._ft_ocr_translate.toggled.connect(self._ft_ocr_translate_lang.setEnabled)
        tr_row.addWidget(self._ft_ocr_translate_lang, 1)
        tr_host = QWidget()
        tr_host.setLayout(tr_row)
        f.addRow("", tr_host)

        run = QPushButton("Extract Text")
        run.setObjectName("PrimaryButton")
        run.clicked.connect(self._ft_run_ocr)
        f.addRow("", run)

        lv.addWidget(opts)

        # Right: preview
        right = QGroupBox("Preview")
        rv = QVBoxLayout(right)
        rv.setContentsMargins(12, 10, 12, 10)
        rv.setSpacing(8)
        self._ft_ocr_preview = QTextEdit()
        self._ft_ocr_preview.setReadOnly(False)
        self._ft_ocr_preview.setPlaceholderText("Extracted text will appear here…")
        rv.addWidget(self._ft_ocr_preview, 1)
        self._ft_ocr_last_text = ""
        self._ft_ocr_text_map = {}

        pv_row = QHBoxLayout()
        pv_row.setSpacing(8)
        copy_btn = QPushButton("Copy Text")
        copy_btn.clicked.connect(self._ft_ocr_copy_preview)
        pv_row.addWidget(copy_btn)
        reset_btn = QPushButton("Reset")
        reset_btn.clicked.connect(self._ft_ocr_reset_preview)
        pv_row.addWidget(reset_btn)
        compare_btn = QPushButton("Compare")
        compare_btn.clicked.connect(self._ft_ocr_compare)
        pv_row.addWidget(compare_btn)
        save_btn = QPushButton("Save As…")
        save_btn.clicked.connect(self._ft_ocr_save_preview)
        pv_row.addWidget(save_btn)
        pv_row.addStretch(1)
        rv.addLayout(pv_row)

        left_scroll = QScrollArea()
        left_scroll.setWidgetResizable(True)
        left_scroll.setFrameShape(QFrame.NoFrame)
        left_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        left_scroll.setWidget(left)
        split.addWidget(left_scroll)
        split.addWidget(right)
        split.setSizes([520, 520])
        outer.addWidget(split, 1)
        return page

    def _ft_ocr_pick_files(self):
        try:
            filt = (
                "Documents (*.pdf *.png *.jpg *.jpeg *.webp *.bmp *.tif *.tiff *.gif *.docx *.pptx *.xlsx "
                "*.txt *.md *.csv *.json *.xml);;All files (*.*)"
            )
            files, _ = QFileDialog.getOpenFileNames(self, "Select files", filter=filt)
            if files:
                self._ft_ocr_add_files(files)
        except Exception:
            pass

    def _ft_ocr_add_files(self, paths: list):
        if not hasattr(self, "_ft_ocr_files"):
            self._ft_ocr_files = []
        try:
            from core import text_extractor
        except Exception:
            text_extractor = None
        allowed = {".pdf"}
        if text_extractor:
            allowed = set(
                text_extractor.TEXT_EXTS
                | text_extractor.IMAGE_EXTS
                | text_extractor.DOC_EXTS
                | text_extractor.PPT_EXTS
                | text_extractor.SHEET_EXTS
                | {".pdf"}
            )
        added = 0
        for p in paths:
            try:
                path = Path(p)
                if path.is_dir():
                    continue
                if path.suffix.lower() not in allowed:
                    continue
                sp = str(path)
                if sp in self._ft_ocr_files:
                    continue
                item = QListWidgetItem(path.name)
                item.setData(Qt.UserRole, sp)
                self._ft_ocr_list.addItem(item)
                self._ft_ocr_files.append(sp)
                added += 1
            except Exception:
                pass
        if added:
            try:
                if not (self._ft_ocr_out.text() or "").strip():
                    self._ft_ocr_out.setText(str(Path(self._ft_ocr_files[-1]).parent))
            except Exception:
                pass

    def _ft_ocr_remove_selected(self):
        try:
            items = self._ft_ocr_list.selectedItems()
            for it in items:
                p = it.data(Qt.UserRole)
                if hasattr(self, "_ft_ocr_files") and p in self._ft_ocr_files:
                    self._ft_ocr_files.remove(p)
                row = self._ft_ocr_list.row(it)
                self._ft_ocr_list.takeItem(row)
        except Exception:
            pass

    def _ft_ocr_clear_files(self):
        try:
            self._ft_ocr_list.clear()
        except Exception:
            pass
        self._ft_ocr_files = []

    def _ft_ocr_pick_output_dir(self):
        try:
            d = QFileDialog.getExistingDirectory(self, "Select destination folder")
            if d:
                self._ft_ocr_out.setText(d)
        except Exception:
            pass

    def _ft_ocr_copy_preview(self):
        try:
            text = self._ft_ocr_preview.toPlainText()
            if text:
                QApplication.clipboard().setText(text)
        except Exception:
            pass

    def _ft_ocr_reset_preview(self):
        try:
            self._ft_ocr_preview.setPlainText(self._ft_ocr_last_text or "")
        except Exception:
            pass

    def _ft_ocr_save_preview(self):
        text = self._ft_ocr_preview.toPlainText()
        if not text.strip():
            QMessageBox.information(self, "Save", "No extracted text to save.")
            return
        fmt = str(self._ft_ocr_format.currentText() or "Plain text (.txt)")
        ext = ".txt" if "Plain" in fmt else ".md" if "Markdown" in fmt else ".docx"
        try:
            out, _ = QFileDialog.getSaveFileName(self, "Save extracted text", filter=f"Text (*{ext})")
            if not out:
                return
            out_path = Path(out)
            if out_path.suffix.lower() != ext:
                out_path = out_path.with_suffix(ext)
            if ext == ".docx":
                try:
                    import docx
                except Exception:
                    QMessageBox.warning(self, "Save", "DOCX export requires python-docx.")
                    return
                doc = docx.Document()
                for line in text.splitlines():
                    doc.add_paragraph(line)
                doc.save(out_path)
            else:
                out_path.write_text(text, encoding="utf-8", errors="replace")
        except Exception as e:
            QMessageBox.warning(self, "Save", str(e))

    def _ft_ocr_set_preview(self, text: str):
        self._ft_ocr_last_text = text or ""
        try:
            self._ft_ocr_preview.setPlainText(self._ft_ocr_last_text)
        except Exception:
            pass

    def _ft_ocr_store_text(self, file_path: str, text: str):
        try:
            if not hasattr(self, "_ft_ocr_text_map") or self._ft_ocr_text_map is None:
                self._ft_ocr_text_map = {}
            self._ft_ocr_text_map[str(file_path)] = text or ""
        except Exception:
            pass

    def _ft_ocr_selected_file(self) -> str | None:
        try:
            items = self._ft_ocr_list.selectedItems()
            if items:
                return items[0].data(Qt.UserRole)
        except Exception:
            pass
        try:
            if hasattr(self, "_ft_ocr_files") and self._ft_ocr_files:
                return self._ft_ocr_files[-1]
        except Exception:
            pass
        return None

    def _ft_ocr_compare(self):
        p = self._ft_ocr_selected_file()
        if not p:
            QMessageBox.information(self, "Compare", "Select a file to compare.")
            return
        text = ""
        try:
            text = (self._ft_ocr_text_map or {}).get(str(p), "")
        except Exception:
            text = ""
        if not text:
            text = self._ft_ocr_preview.toPlainText()
        dlg = _QtOCRCompareDialog(
            self,
            file_path=str(p),
            extracted_text=text,
            ocr_engine=str(self._ft_ocr_engine.currentText() or "Auto"),
        )
        dlg.exec()

    def _ft_run_ocr(self):
        files = []
        try:
            for i in range(self._ft_ocr_list.count()):
                item = self._ft_ocr_list.item(i)
                p = item.data(Qt.UserRole)
                if p:
                    files.append(p)
        except Exception:
            files = []
        if not files:
            QMessageBox.warning(self, "Text Extractor", "Add at least one file.")
            return

        mode_map = {
            "Auto (text + OCR if needed)": "auto",
            "Force OCR": "force",
            "Text only": "text_only",
        }
        ocr_mode = mode_map.get(str(self._ft_ocr_mode.currentText()), "auto")
        engine_map = {"Auto": "auto", "RapidOCR": "rapidocr", "AI Vision": "ai"}
        ocr_engine = engine_map.get(str(self._ft_ocr_engine.currentText()), "auto")
        out_fmt = str(self._ft_ocr_format.currentText() or "Plain text (.txt)")
        out_dir = (self._ft_ocr_out.text() or "").strip()
        normalize = bool(self._ft_ocr_normalize.isChecked())
        add_headers = bool(self._ft_ocr_headers.isChecked())
        use_ai = bool(self._ft_ocr_ai.isChecked())
        translate = bool(self._ft_ocr_translate.isChecked())
        lang = (self._ft_ocr_lang.text() or "eng").strip()
        max_pages = int(self._ft_ocr_pages.value())

        needs_ai = use_ai or translate or ocr_engine == "ai"
        try:
            import importlib.util

            has_rapid = importlib.util.find_spec("rapidocr_onnxruntime") is not None
        except Exception:
            has_rapid = False
        if ocr_engine == "rapidocr" and not has_rapid:
            QMessageBox.warning(self, "Text Extractor", "RapidOCR is not installed.")
            return
        if ocr_engine == "auto" and not has_rapid:
            needs_ai = True
        required_kind: str | None = None
        if needs_ai:
            # OCR may run in three ways:
            # - RapidOCR (onnxruntime): no model needed
            # - AI Vision OCR: requires vision model
            # - Text-only features (e.g., translate on extracted text): text model is fine
            if ocr_engine == "ai" or use_ai or (ocr_engine == "auto" and not has_rapid and ocr_mode != "text_only"):
                required_kind = "vision"
            elif translate:
                required_kind = "text"

        if needs_ai and not self._ensure_ai_ready(title="AI Required for OCR", kind=required_kind):
            return

        worker = _QtTextExtractWorker(
            files=files,
            ocr_mode=ocr_mode,
            ocr_engine=ocr_engine,
            ocr_lang=lang,
            max_ocr_pages=max_pages,
            output_format=out_fmt,
            output_dir=out_dir or None,
            normalize=normalize,
            add_headers=add_headers,
            use_ai=use_ai,
            translate=translate,
            translate_lang=str(self._ft_ocr_translate_lang.currentText() or "English"),
            ai_manager=getattr(self.backend, "ai_manager", None),
        )
        self._ft_ocr_text_map = {}
        self._ft_ocr_last_text = ""
        worker.text_ready.connect(self._ft_ocr_set_preview)
        worker.per_file_ready.connect(self._ft_ocr_store_text)
        self._ft_start_worker(worker)

    def _build_file_tools_archives_page(self) -> QWidget:
        page = QWidget()
        outer = QVBoxLayout(page)
        outer.setContentsMargins(0, 0, 0, 0)
        outer.setSpacing(12)

        splitter = QSplitter(Qt.Horizontal)
        splitter.setChildrenCollapsible(False)

        create = QGroupBox("Create Archive")
        c = QFormLayout(create)
        c.setLabelAlignment(Qt.AlignRight)
        c.setVerticalSpacing(10)
        c.setHorizontalSpacing(12)

        self._ft_arc_fmt = QComboBox()
        self._ft_arc_fmt.addItems(["zip", "7z", "tar.gz", "tar.xz", "tar.bz2"])
        c.addRow("Format:", self._ft_arc_fmt)

        self._ft_arc_name = QLineEdit("Archive")
        c.addRow("Output name:", self._ft_arc_name)

        self._ft_arc_overwrite = QCheckBox("Overwrite")
        c.addRow("", self._ft_arc_overwrite)

        self._ft_arc_parts_mb = QSpinBox()
        self._ft_arc_parts_mb.setRange(0, 50_000)
        self._ft_arc_parts_mb.setValue(0)
        self._ft_arc_parts_mb.setToolTip("Split into parts of this size (MB). 0 disables splitting.")
        c.addRow("Split parts (MB):", self._ft_arc_parts_mb)

        run = QPushButton("Create Archive")
        run.setObjectName("PrimaryButton")
        run.clicked.connect(self._ft_run_archive_create)
        c.addRow("", run)

        extract = QGroupBox("Extract Archive")
        e = QFormLayout(extract)
        e.setLabelAlignment(Qt.AlignRight)
        e.setVerticalSpacing(10)
        e.setHorizontalSpacing(12)

        arow = QHBoxLayout()
        self._ft_arc_in = QLineEdit()
        self._ft_arc_in.setPlaceholderText("Pick an archive file…")
        arow.addWidget(self._ft_arc_in, 1)
        pick = QPushButton("Browse")
        pick.clicked.connect(self._ft_pick_archive)
        arow.addWidget(pick)
        ahost = QWidget()
        ahost.setLayout(arow)
        e.addRow("Archive:", ahost)

        self._ft_arc_extract_dest = QComboBox()
        self._ft_arc_extract_dest.addItems(["target", "archive folder"])
        e.addRow("Output folder:", self._ft_arc_extract_dest)

        self._ft_arc_extract_overwrite = QCheckBox("Overwrite")
        e.addRow("", self._ft_arc_extract_overwrite)

        note = QLabel("Tip: RAR extraction may require an unrar backend (WinRAR, unrar, or 7-Zip).")
        note.setStyleSheet("color:#9aa0a9;")
        note.setWordWrap(True)
        e.addRow("", note)

        run2 = QPushButton("Extract")
        run2.setObjectName("PrimaryButton")
        run2.clicked.connect(self._ft_run_archive_extract)
        e.addRow("", run2)

        splitter.addWidget(create)
        splitter.addWidget(extract)
        splitter.setSizes([520, 520])

        outer.addWidget(splitter, 1)
        return page

    def _ft_pick_archive(self):
        try:
            f, _ = QFileDialog.getOpenFileName(self, "Select archive", filter="Archives (*.zip *.7z *.rar *.tar.gz *.tar.xz *.tar.bz2 *.tar)")
            if f:
                self._ft_arc_in.setText(f)
        except Exception:
            pass

    def _ft_run_archive_create(self):
        folder = (self._ft_target.text() or "").strip()
        if not folder:
            QMessageBox.warning(self, "Archive", "Pick a target folder first.")
            return
        fmt = str(self._ft_arc_fmt.currentText())
        name = (self._ft_arc_name.text() or "Archive").strip()
        ext = ".zip" if fmt == "zip" else ".7z" if fmt == "7z" else ".tar.gz" if fmt == "tar.gz" else ".tar.xz" if fmt == "tar.xz" else ".tar.bz2"
        out_path = str(Path(folder) / (name + ext))
        worker = _QtArchiveCreateWorker(
            folder=folder,
            archive_path=out_path,
            fmt=fmt,
            overwrite=bool(self._ft_arc_overwrite.isChecked()),
            part_size_mb=int(self._ft_arc_parts_mb.value()),
        )
        self._ft_start_worker(worker)

    def _ft_run_archive_extract(self):
        folder = (self._ft_target.text() or "").strip()
        arc = (self._ft_arc_in.text() or "").strip()
        if not arc:
            QMessageBox.warning(self, "Extract", "Pick an archive file.")
            return
        out_root = str(self._ft_arc_extract_dest.currentText())
        out_dir = (Path(folder) if out_root == "target" else Path(arc).parent) / ("Extracted_" + Path(arc).stem)
        worker = _QtArchiveExtractWorker(
            archive_path=arc,
            output_dir=str(out_dir),
            overwrite=bool(self._ft_arc_extract_overwrite.isChecked()),
        )
        self._ft_start_worker(worker)

    def _build_file_tools_pdf_page(self) -> QWidget:
        page = QWidget()
        outer = QVBoxLayout(page)
        outer.setContentsMargins(0, 0, 0, 0)
        outer.setSpacing(12)

        # Input PDF (shared)
        inp = QGroupBox("Input PDF")
        f = QFormLayout(inp)
        f.setLabelAlignment(Qt.AlignRight)
        f.setVerticalSpacing(10)
        f.setHorizontalSpacing(12)

        row = QHBoxLayout()
        self._ft_pdf_input = QLineEdit()
        self._ft_pdf_input.setPlaceholderText("Pick a PDF file…")
        row.addWidget(self._ft_pdf_input, 1)
        pick = QPushButton("Browse")
        pick.clicked.connect(self._ft_pick_pdf)
        row.addWidget(pick)
        host = QWidget()
        host.setLayout(row)
        f.addRow("PDF:", host)

        outer.addWidget(inp)

        split = QSplitter(Qt.Horizontal)
        split.setChildrenCollapsible(False)

        # Merge
        merge = QGroupBox("Merge PDFs (Target folder)")
        m = QFormLayout(merge)
        m.setLabelAlignment(Qt.AlignRight)
        m.setVerticalSpacing(10)
        m.setHorizontalSpacing(12)

        self._ft_pdf_merge_include = QCheckBox("Include subfolders")
        self._ft_pdf_merge_include.setChecked(True)
        m.addRow("", self._ft_pdf_merge_include)

        self._ft_pdf_merge_name = QLineEdit("Merged.pdf")
        m.addRow("Output name:", self._ft_pdf_merge_name)

        self._ft_pdf_merge_overwrite = QCheckBox("Overwrite")
        m.addRow("", self._ft_pdf_merge_overwrite)

        run_merge = QPushButton("Merge PDFs")
        run_merge.setObjectName("PrimaryButton")
        run_merge.clicked.connect(self._ft_run_pdf_merge)
        m.addRow("", run_merge)

        # Split / Extract / Rotate / Bookmarks
        ops = QGroupBox("Split / Extract / Rotate")
        owrap = QVBoxLayout(ops)
        owrap.setContentsMargins(12, 10, 12, 10)
        owrap.setSpacing(10)
        o = QGridLayout()
        o.setContentsMargins(0, 0, 0, 0)
        o.setHorizontalSpacing(12)
        o.setVerticalSpacing(10)
        o.setColumnStretch(1, 1)
        o.setColumnStretch(3, 1)

        self._ft_pdf_pages = QLineEdit("all")
        self._ft_pdf_pages.setPlaceholderText("all  |  1-3,5,7-9")
        lbl_pages = QLabel("Pages:")
        lbl_pages.setAlignment(Qt.AlignRight | Qt.AlignVCenter)
        o.addWidget(lbl_pages, 0, 0)
        o.addWidget(self._ft_pdf_pages, 0, 1, 1, 3)

        self._ft_pdf_out_folder = QComboBox()
        self._ft_pdf_out_folder.addItems(["target", "pdf folder"])
        lbl_out = QLabel("Output:")
        lbl_out.setAlignment(Qt.AlignRight | Qt.AlignVCenter)
        o.addWidget(lbl_out, 1, 0)
        o.addWidget(self._ft_pdf_out_folder, 1, 1)

        self._ft_pdf_out_sub = QLineEdit("PDF_Output")
        lbl_sub = QLabel("Subfolder:")
        lbl_sub.setAlignment(Qt.AlignRight | Qt.AlignVCenter)
        o.addWidget(lbl_sub, 1, 2)
        o.addWidget(self._ft_pdf_out_sub, 1, 3)

        self._ft_pdf_overwrite = QCheckBox("Overwrite")
        o.addWidget(QLabel(""), 2, 0)
        o.addWidget(self._ft_pdf_overwrite, 2, 1)

        btns = QHBoxLayout()
        btns.setSpacing(10)
        b_split = QPushButton("Split")
        b_split.setObjectName("PrimaryButton")
        b_split.clicked.connect(self._ft_run_pdf_split_pages)
        btns.addWidget(b_split)
        b_extract = QPushButton("Extract")
        b_extract.setObjectName("PrimaryButton")
        b_extract.clicked.connect(self._ft_run_pdf_extract_one)
        btns.addWidget(b_extract)
        btn_host = QWidget()
        btn_host.setLayout(btns)
        o.addWidget(QLabel(""), 3, 0)
        o.addWidget(btn_host, 3, 1, 1, 3)

        rot_row = QHBoxLayout()
        rot_row.setSpacing(10)
        self._ft_pdf_rotate_deg = QComboBox()
        self._ft_pdf_rotate_deg.addItems(["90", "180", "270"])
        rot_row.addWidget(self._ft_pdf_rotate_deg)
        self._ft_pdf_rotate_name = QLineEdit("Rotated.pdf")
        rot_row.addWidget(self._ft_pdf_rotate_name, 1)
        b_rotate = QPushButton("Rotate")
        b_rotate.setObjectName("PrimaryButton")
        b_rotate.clicked.connect(self._ft_run_pdf_rotate)
        rot_row.addWidget(b_rotate)
        rot_host = QWidget()
        rot_host.setLayout(rot_row)
        lbl_rot = QLabel("Rotate:")
        lbl_rot.setAlignment(Qt.AlignRight | Qt.AlignVCenter)
        o.addWidget(lbl_rot, 4, 0)
        o.addWidget(rot_host, 4, 1, 1, 3)

        bm_row = QHBoxLayout()
        bm_row.setSpacing(10)
        self._ft_pdf_bm_min_pages = QSpinBox()
        self._ft_pdf_bm_min_pages.setRange(1, 5000)
        self._ft_pdf_bm_min_pages.setValue(1)
        bm_row.addWidget(self._ft_pdf_bm_min_pages)
        b_bm = QPushButton("Split by Bookmarks")
        b_bm.setObjectName("PrimaryButton")
        b_bm.clicked.connect(self._ft_run_pdf_split_bookmarks)
        bm_row.addWidget(b_bm)
        bm_host = QWidget()
        bm_host.setLayout(bm_row)
        lbl_bm = QLabel("Bookmarks:")
        lbl_bm.setAlignment(Qt.AlignRight | Qt.AlignVCenter)
        o.addWidget(lbl_bm, 5, 0)
        o.addWidget(bm_host, 5, 1, 1, 3)

        owrap.addLayout(o)
        owrap.addStretch(1)

        split.addWidget(merge)
        split.addWidget(ops)
        split.setSizes([520, 520])

        outer.addWidget(split, 1)
        return page

    def _ft_pick_pdf(self):
        try:
            f, _ = QFileDialog.getOpenFileName(self, "Select PDF", filter="PDF (*.pdf)")
            if f:
                self._ft_pdf_input.setText(f)
        except Exception:
            pass

    def _ft_pdf_output_dir(self, input_pdf: str) -> Path:
        folder = (self._ft_target.text() or "").strip()
        out_root = str(self._ft_pdf_out_folder.currentText())
        sub = (self._ft_pdf_out_sub.text() or "PDF_Output").strip()
        base = Path(folder) if out_root == "target" else Path(input_pdf).parent
        return base / sub

    def _ft_run_pdf_merge(self):
        folder = (self._ft_target.text() or "").strip()
        if not folder:
            QMessageBox.warning(self, "PDF", "Pick a target folder first.")
            return
        out_name = (self._ft_pdf_merge_name.text() or "Merged.pdf").strip()
        if not out_name.lower().endswith(".pdf"):
            out_name += ".pdf"
        out = str(Path(folder) / out_name)
        worker = _QtPdfMergeWorker(
            folder=folder,
            include_subfolders=bool(self._ft_pdf_merge_include.isChecked()),
            output_pdf=out,
            overwrite=bool(self._ft_pdf_merge_overwrite.isChecked()),
        )
        self._ft_start_worker(worker)

    def _ft_run_pdf_split_pages(self):
        inp = (self._ft_pdf_input.text() or "").strip()
        if not inp:
            QMessageBox.warning(self, "PDF", "Pick an input PDF.")
            return
        out_dir = self._ft_pdf_output_dir(inp)
        worker = _QtPdfSplitPagesWorker(
            input_pdf=inp,
            output_dir=str(out_dir),
            page_ranges=(self._ft_pdf_pages.text() or "all").strip(),
            overwrite=bool(self._ft_pdf_overwrite.isChecked()),
        )
        self._ft_start_worker(worker)

    def _ft_run_pdf_extract_one(self):
        inp = (self._ft_pdf_input.text() or "").strip()
        if not inp:
            QMessageBox.warning(self, "PDF", "Pick an input PDF.")
            return
        out_dir = self._ft_pdf_output_dir(inp)
        out_dir.mkdir(parents=True, exist_ok=True)
        out = str(out_dir / f"{Path(inp).stem}_extracted.pdf")
        worker = _QtPdfExtractWorker(
            input_pdf=inp,
            output_pdf=out,
            page_ranges=(self._ft_pdf_pages.text() or "all").strip(),
            overwrite=bool(self._ft_pdf_overwrite.isChecked()),
        )
        self._ft_start_worker(worker)

    def _ft_run_pdf_rotate(self):
        inp = (self._ft_pdf_input.text() or "").strip()
        if not inp:
            QMessageBox.warning(self, "PDF", "Pick an input PDF.")
            return
        out_dir = self._ft_pdf_output_dir(inp)
        out_dir.mkdir(parents=True, exist_ok=True)
        out_name = (self._ft_pdf_rotate_name.text() or "Rotated.pdf").strip()
        if not out_name.lower().endswith(".pdf"):
            out_name += ".pdf"
        out = str(out_dir / out_name)
        worker = _QtPdfRotateWorker(
            input_pdf=inp,
            output_pdf=out,
            rotation_degrees=int(self._ft_pdf_rotate_deg.currentText() or "90"),
            page_ranges=(self._ft_pdf_pages.text() or "all").strip(),
            overwrite=bool(self._ft_pdf_overwrite.isChecked()),
        )
        self._ft_start_worker(worker)

    def _ft_run_pdf_split_bookmarks(self):
        inp = (self._ft_pdf_input.text() or "").strip()
        if not inp:
            QMessageBox.warning(self, "PDF", "Pick an input PDF.")
            return
        out_dir = self._ft_pdf_output_dir(inp)
        worker = _QtPdfSplitBookmarksWorker(
            input_pdf=inp,
            output_dir=str(out_dir),
            overwrite=bool(self._ft_pdf_overwrite.isChecked()),
            min_pages=int(self._ft_pdf_bm_min_pages.value() or 1),
        )
        self._ft_start_worker(worker)

    def _build_settings_page(self) -> QWidget:
        host = QFrame()
        host.setObjectName("PageHost")
        layout = QVBoxLayout(host)
        layout.setContentsMargins(18, 18, 18, 18)
        layout.setSpacing(14)

        header = QFrame()
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(0, 0, 0, 0)
        header_layout.setSpacing(10)
        title = QLabel("Settings")
        title.setObjectName("PageTitle")
        header_layout.addWidget(title)
        sub = QLabel("Manage AI and optional tools (ffmpeg, LibreOffice).")
        sub.setObjectName("PageSubTitle")
        header_layout.addWidget(sub, 1)
        layout.addWidget(header)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QFrame.NoFrame)
        inner = QFrame()
        inner_layout = QVBoxLayout(inner)
        inner_layout.setContentsMargins(0, 0, 0, 0)
        inner_layout.setSpacing(12)
        scroll.setWidget(inner)
        layout.addWidget(scroll, 1)

        # AI
        ai_box = QGroupBox("AI Model")
        a = QVBoxLayout(ai_box)
        a.setContentsMargins(12, 10, 12, 10)
        a.setSpacing(10)

        self._st_ai_status = QLabel("AI: unavailable")
        self._st_ai_status.setStyleSheet("color:#9aa0a9;")
        a.addWidget(self._st_ai_status)

        # Settings accessor (defined early so model pickers don't fall back due to UnboundLocalError).
        settings = getattr(self.backend, "settings_manager", None)
        get_setting = (settings.get_setting if settings else (lambda _k, d=None: d))

        # Model pickers (catalog-driven). Locked by default to avoid accidental changes while scrolling.
        self._st_ai_vision_model = QComboBox()
        self._st_ai_text_model = QComboBox()
        try:
            from core.ai_model_catalog import get_models_by_kind, DEFAULT_MODEL_ID, DEFAULT_TEXT_MODEL_ID

            for mid, spec in get_models_by_kind("vision").items():
                self._st_ai_vision_model.addItem(spec.name, mid)
            for mid, spec in get_models_by_kind("text").items():
                self._st_ai_text_model.addItem(spec.name, mid)

            cur_vision = str(get_setting("ai_vision_model_id", get_setting("ai_model_id", DEFAULT_MODEL_ID)) or DEFAULT_MODEL_ID)
            v_idx = self._st_ai_vision_model.findData(cur_vision)
            if v_idx >= 0:
                self._st_ai_vision_model.setCurrentIndex(v_idx)

            cur_text = str(get_setting("ai_text_model_id", DEFAULT_TEXT_MODEL_ID) or DEFAULT_TEXT_MODEL_ID)
            t_idx = self._st_ai_text_model.findData(cur_text)
            if t_idx >= 0:
                self._st_ai_text_model.setCurrentIndex(t_idx)
        except Exception:
            self._st_ai_vision_model.addItem("Vision model (default)", "")
            self._st_ai_text_model.addItem("Text model (default)", "")

        model_row_v = QHBoxLayout()
        model_row_v.setSpacing(10)
        model_row_v.addWidget(QLabel("Vision model:"))
        model_row_v.addWidget(self._st_ai_vision_model, 1)
        a.addLayout(model_row_v)

        model_row_t = QHBoxLayout()
        model_row_t.setSpacing(10)
        model_row_t.addWidget(QLabel("Text model:"))
        model_row_t.addWidget(self._st_ai_text_model, 1)
        a.addLayout(model_row_t)

        row = QHBoxLayout()
        row.setSpacing(10)
        self._st_ai_target = QComboBox()
        self._st_ai_target.addItem("Load vision", "vision")
        self._st_ai_target.addItem("Load text", "text")
        row.addWidget(self._st_ai_target)

        self._st_ai_load = QPushButton("Download Model")
        self._st_ai_load.setObjectName("PrimaryButton")
        self._st_ai_load.setIcon(self.icons.icon("download"))
        self._st_ai_load.setIconSize(QSize(18, 18))
        self._st_ai_load.setToolTip("Download the selected model files only. This does not load the model into memory.")
        self._st_ai_load.clicked.connect(self._settings_load_ai)
        row.addWidget(self._st_ai_load)

        self._st_ai_load_model = QPushButton("Load Model")
        self._st_ai_load_model.setIcon(self.icons.icon("ai"))
        self._st_ai_load_model.setIconSize(QSize(18, 18))
        self._st_ai_load_model.setToolTip("Load the downloaded model into memory. Use CPU-safe settings first unless you have configured GPU support.")
        self._st_ai_load_model.clicked.connect(self._settings_load_ai_into_memory)
        row.addWidget(self._st_ai_load_model)

        self._st_ai_unload = QPushButton("Unload")
        self._st_ai_unload.clicked.connect(self._settings_unload_ai)
        row.addWidget(self._st_ai_unload)

        self._st_ai_delete = QPushButton("Delete Model")
        self._st_ai_delete.clicked.connect(self._settings_delete_ai_files)
        row.addWidget(self._st_ai_delete)

        # Lock/unlock AI settings editing to prevent accidental wheel changes while scrolling.
        self._st_ai_edit = QPushButton("Edit")
        self._st_ai_edit.setCheckable(True)
        self._st_ai_edit.setChecked(False)
        self._st_ai_edit.clicked.connect(lambda _v: self._settings_set_ai_editable(bool(self._st_ai_edit.isChecked())))
        row.addWidget(self._st_ai_edit)
        row.addStretch(1)
        a.addLayout(row)

        # Performance + behavior settings (ported from the original Tk UI).

        import os

        perf_box = QGroupBox("Performance")
        perf = QFormLayout(perf_box)
        perf.setContentsMargins(10, 10, 10, 10)
        perf.setHorizontalSpacing(14)
        perf.setVerticalSpacing(10)

        # GPU layers
        self._st_ai_gpu_layers = QSlider(Qt.Horizontal)
        self._st_ai_gpu_layers.setRange(0, 0)
        self._st_ai_gpu_layers.setValue(int(get_setting("ai_gpu_layers", 0) or 0))
        gpu_value = QLabel(str(self._st_ai_gpu_layers.value()))
        gpu_row = QHBoxLayout()
        gpu_row.addWidget(self._st_ai_gpu_layers, 1)
        gpu_row.addWidget(gpu_value)
        self._st_ai_gpu_layers.valueChanged.connect(lambda v: gpu_value.setText(str(int(v))))
        perf.addRow("GPU Layers (0 = CPU only)", gpu_row)

        # CPU threads
        max_threads = int(os.cpu_count() or 8)
        self._st_ai_threads = QSlider(Qt.Horizontal)
        self._st_ai_threads.setRange(1, max_threads)
        self._st_ai_threads.setValue(int(get_setting("ai_threads", max_threads) or max_threads))
        th_value = QLabel(str(self._st_ai_threads.value()))
        th_row = QHBoxLayout()
        th_row.addWidget(self._st_ai_threads, 1)
        th_row.addWidget(th_value)
        self._st_ai_threads.valueChanged.connect(lambda v: th_value.setText(str(int(v))))
        perf.addRow(f"CPU Threads (1–{max_threads})", th_row)

        # Context size / batch / image size
        self._st_ai_ctx = QComboBox()
        self._settings_refresh_ai_context_presets()
        try:
            self._st_ai_vision_model.currentIndexChanged.connect(self._settings_refresh_ai_context_presets)
            self._st_ai_text_model.currentIndexChanged.connect(self._settings_refresh_ai_context_presets)
        except Exception:
            pass
        try:
            self._st_ai_ctx.currentTextChanged.connect(self._settings_update_workflow_warning)
        except Exception:
            pass
        perf.addRow("Context Size", self._st_ai_ctx)

        self._st_ai_batch = QComboBox()
        for v in ("128", "256", "512", "1024"):
            self._st_ai_batch.addItem(v)
        self._st_ai_batch.setCurrentText(str(get_setting("ai_batch_size", 512) or 512))
        perf.addRow("Batch Size", self._st_ai_batch)

        self._st_ai_img = QComboBox()
        for v in ("256", "384", "512", "768", "1024"):
            self._st_ai_img.addItem(v)
        self._st_ai_img.setCurrentText(str(get_setting("ai_image_size", 512) or 512))
        perf.addRow("Image Size", self._st_ai_img)

        # FlashAttention (llama-cpp)
        self._st_ai_flash_attn = QComboBox()
        self._st_ai_flash_attn.addItem("Auto", "auto")
        self._st_ai_flash_attn.addItem("Enabled", "enabled")
        self._st_ai_flash_attn.addItem("Disabled", "disabled")
        cur_fa = get_setting("ai_flash_attn_type", "disabled")
        try:
            if isinstance(cur_fa, int):
                cur_fa = "auto" if int(cur_fa) == -1 else ("enabled" if int(cur_fa) == 1 else "disabled")
            cur_fa = str(cur_fa or "auto").strip().lower()
        except Exception:
            cur_fa = "auto"
        idx = self._st_ai_flash_attn.findData(cur_fa)
        if idx >= 0:
            self._st_ai_flash_attn.setCurrentIndex(idx)
        perf.addRow("Flash Attention", self._st_ai_flash_attn)

        a.addWidget(perf_box)

        rename_box = QGroupBox("Smart Rename")
        rb = QFormLayout(rename_box)
        rb.setContentsMargins(10, 10, 10, 10)
        rb.setHorizontalSpacing(14)
        rb.setVerticalSpacing(10)

        self._st_ai_rename_max = QSpinBox()
        self._st_ai_rename_max.setRange(1, 12)
        self._st_ai_rename_max.setValue(int(get_setting("ai_rename_max_keywords", 8) or 8))
        rb.addRow("Max keywords", self._st_ai_rename_max)

        self._st_ai_rename_prompt = QTextEdit()
        self._st_ai_rename_prompt.setMinimumHeight(110)
        self._st_ai_rename_prompt.setPlaceholderText("Optional custom rename prompt…")
        self._st_ai_rename_prompt.setText(str(get_setting("ai_rename_prompt", "") or ""))
        rb.addRow("Prompt", self._st_ai_rename_prompt)

        a.addWidget(rename_box)

        btns = QHBoxLayout()
        btns.setSpacing(10)
        self._st_ai_save = QPushButton("Save AI Settings")
        self._st_ai_save.setObjectName("PrimaryButton")
        self._st_ai_save.clicked.connect(self._settings_save_ai_params)
        btns.addWidget(self._st_ai_save)
        self._st_ai_reload = QPushButton("Reload Model")
        self._st_ai_reload.clicked.connect(self._settings_reload_ai)
        btns.addWidget(self._st_ai_reload)
        btns.addStretch(1)
        a.addLayout(btns)

        # Install wheel guards (only changes values if the control is focused).
        if not self._wheel_guard_installed:
            for w in (
                self._st_ai_gpu_layers,
                self._st_ai_threads,
                self._st_ai_ctx,
                self._st_ai_batch,
                self._st_ai_img,
                self._st_ai_rename_max,
            ):
                try:
                    w.installEventFilter(self)
                except Exception:
                    pass
            self._wheel_guard_installed = True

        # Start in read-only mode.
        self._settings_set_ai_editable(False)

        inner_layout.addWidget(ai_box)

        # Automation Workflow
        wf_box = QGroupBox("Automation Workflow")
        wf_layout = QVBoxLayout(wf_box)
        wf_layout.setContentsMargins(12, 10, 12, 10)
        wf_layout.setSpacing(10)

        wf_hint = QLabel("Defaults for workflow automation outputs.")
        wf_hint.setStyleSheet("color:#9aa0a9;")
        wf_hint.setWordWrap(True)
        wf_layout.addWidget(wf_hint)

        wf_form = QFormLayout()
        wf_form.setContentsMargins(0, 0, 0, 0)
        wf_form.setHorizontalSpacing(14)
        wf_form.setVerticalSpacing(10)

        wf_settings = {}
        if settings:
            try:
                wf_settings = settings.get_workflow_settings()
            except Exception:
                wf_settings = {}

        self._st_wf_max_tokens = QSpinBox()
        self._st_wf_max_tokens.setRange(256, 256000)
        self._st_wf_max_tokens.setSingleStep(256)
        self._st_wf_max_tokens.setValue(int((wf_settings or {}).get("max_output_tokens", 4096) or 4096))
        wf_form.addRow("Max output tokens", self._st_wf_max_tokens)

        wf_layout.addLayout(wf_form)

        self._st_wf_warn = QLabel("")
        self._st_wf_warn.setStyleSheet("color:#e6b35c;")
        self._st_wf_warn.setWordWrap(True)
        wf_layout.addWidget(self._st_wf_warn)
        try:
            self._st_wf_max_tokens.valueChanged.connect(self._settings_update_workflow_warning)
        except Exception:
            pass
        self._settings_update_workflow_warning()

        wf_btns = QHBoxLayout()
        wf_btns.addStretch(1)
        self._st_wf_save = QPushButton("Save Workflow Settings")
        self._st_wf_save.setObjectName("PrimaryButton")
        self._st_wf_save.clicked.connect(self._settings_save_workflow)
        wf_btns.addWidget(self._st_wf_save)
        wf_layout.addLayout(wf_btns)

        inner_layout.addWidget(wf_box)

        # External tools
        tools_box = QGroupBox("External Tools")
        t = QVBoxLayout(tools_box)
        t.setContentsMargins(12, 10, 12, 10)
        t.setSpacing(12)

        self._st_ffmpeg = QLabel("ffmpeg: checking…")
        self._st_ffmpeg.setStyleSheet("color:#9aa0a9;")
        t.addWidget(self._st_ffmpeg)

        ff_row = QHBoxLayout()
        ff_row.setSpacing(10)
        self._st_ffmpeg_path = QLineEdit()
        self._st_ffmpeg_path.setPlaceholderText("Optional custom ffmpeg.exe path…")
        ff_row.addWidget(self._st_ffmpeg_path, 1)
        ff_pick = QPushButton("Browse")
        ff_pick.clicked.connect(self._settings_pick_ffmpeg)
        ff_row.addWidget(ff_pick)
        ff_save = QPushButton("Save")
        ff_save.setObjectName("PrimaryButton")
        ff_save.clicked.connect(self._settings_save_ffmpeg)
        ff_row.addWidget(ff_save)
        ff_clear = QPushButton("Clear")
        ff_clear.clicked.connect(self._settings_clear_ffmpeg)
        ff_row.addWidget(ff_clear)
        t.addLayout(ff_row)

        # ffprobe (for duration/progress) + ffplay (preview) are optional but recommended.
        fp_row = QHBoxLayout()
        fp_row.setSpacing(10)
        self._st_ffprobe_path = QLineEdit()
        self._st_ffprobe_path.setPlaceholderText("Optional custom ffprobe.exe path…")
        fp_row.addWidget(self._st_ffprobe_path, 1)
        fp_pick = QPushButton("Browse")
        fp_pick.clicked.connect(self._settings_pick_ffprobe)
        fp_row.addWidget(fp_pick)
        fp_save = QPushButton("Save")
        fp_save.setObjectName("PrimaryButton")
        fp_save.clicked.connect(self._settings_save_ffprobe)
        fp_row.addWidget(fp_save)
        fp_clear = QPushButton("Clear")
        fp_clear.clicked.connect(self._settings_clear_ffprobe)
        fp_row.addWidget(fp_clear)
        t.addLayout(fp_row)

        dl_ff_row = QHBoxLayout()
        dl_ff_row.setSpacing(10)
        self._st_ff_dl = QPushButton("Download FFmpeg Tools (ffmpeg+ffprobe+ffplay)")
        self._st_ff_dl.setObjectName("PrimaryButton")
        self._st_ff_dl.setIcon(self.icons.icon("download"))
        self._st_ff_dl.setIconSize(QSize(18, 18))
        self._st_ff_dl.clicked.connect(self._settings_download_ffmpeg_tools)
        dl_ff_row.addWidget(self._st_ff_dl)
        dl_ff_row.addStretch(1)
        t.addLayout(dl_ff_row)

        self._st_soffice = QLabel("LibreOffice: checking…")
        self._st_soffice.setStyleSheet("color:#9aa0a9;")
        t.addWidget(self._st_soffice)

        lo_row = QHBoxLayout()
        lo_row.setSpacing(10)
        self._st_soffice_path = QLineEdit()
        self._st_soffice_path.setPlaceholderText("Optional custom soffice.exe path…")
        lo_row.addWidget(self._st_soffice_path, 1)
        lo_pick = QPushButton("Browse")
        lo_pick.clicked.connect(self._settings_pick_soffice)
        lo_row.addWidget(lo_pick)
        lo_save = QPushButton("Save")
        lo_save.setObjectName("PrimaryButton")
        lo_save.clicked.connect(self._settings_save_soffice)
        lo_row.addWidget(lo_save)
        lo_clear = QPushButton("Clear")
        lo_clear.clicked.connect(self._settings_clear_soffice)
        lo_row.addWidget(lo_clear)
        t.addLayout(lo_row)

        dl_row = QHBoxLayout()
        dl_row.setSpacing(10)
        self._st_lo_dl = QPushButton("Download LibreOffice (Windows x64)")
        self._st_lo_dl.setObjectName("PrimaryButton")
        self._st_lo_dl.setIcon(self.icons.icon("download"))
        self._st_lo_dl.setIconSize(QSize(18, 18))
        self._st_lo_dl.clicked.connect(self._settings_download_lo)
        dl_row.addWidget(self._st_lo_dl)
        dl_row.addStretch(1)
        t.addLayout(dl_row)

        inner_layout.addWidget(tools_box)

        # Notifications + Behavior (ported from backup Settings dialog)
        app_box = QGroupBox("App")
        ab = QVBoxLayout(app_box)
        ab.setContentsMargins(12, 10, 12, 10)
        ab.setSpacing(12)

        settings = getattr(self.backend, "settings_manager", None)
        get_setting = (settings.get_setting if settings else (lambda _k, d=None: d))

        theme_form = QFormLayout()
        theme_form.setContentsMargins(0, 0, 0, 0)
        theme_form.setHorizontalSpacing(10)
        theme_form.setVerticalSpacing(8)

        self._st_theme_mode = QComboBox()
        self._st_theme_mode.addItem("Dark (Default)", "dark")
        self._st_theme_mode.addItem("Black", "black")
        self._st_theme_mode.addItem("Light", "light")
        cur_theme = str(get_setting("theme", "dark") or "dark").strip().lower()
        idx_theme = self._st_theme_mode.findData(cur_theme)
        if idx_theme >= 0:
            self._st_theme_mode.setCurrentIndex(idx_theme)
        theme_form.addRow("Theme", self._st_theme_mode)

        self._st_color_theme = QComboBox()
        self._st_color_theme.addItem("Blue", "blue")
        self._st_color_theme.addItem("Teal", "teal")
        self._st_color_theme.addItem("Green", "green")
        self._st_color_theme.addItem("Orange", "orange")
        self._st_color_theme.addItem("Rose", "rose")
        self._st_color_theme.addItem("Violet", "violet")
        self._st_color_theme.addItem("Cyan", "cyan")
        cur_accent = str(get_setting("color_theme", "blue") or "blue").strip().lower()
        idx_accent = self._st_color_theme.findData(cur_accent)
        if idx_accent >= 0:
            self._st_color_theme.setCurrentIndex(idx_accent)
        theme_form.addRow("Accent color", self._st_color_theme)

        ab.addLayout(theme_form)
        try:
            self._st_theme_mode.currentIndexChanged.connect(self._settings_preview_theme)
            self._st_color_theme.currentIndexChanged.connect(self._settings_preview_theme)
        except Exception:
            pass

        ab.addSpacing(2)

        self._st_notif_enabled = QCheckBox("Enable Windows notifications")
        self._st_notif_enabled.setChecked(bool(get_setting("notifications_enabled", True)))
        ab.addWidget(self._st_notif_enabled)

        self._st_notif_sound = QCheckBox("Enable notification sound")
        self._st_notif_sound.setChecked(bool(get_setting("notification_sound", True)))
        ab.addWidget(self._st_notif_sound)

        ab.addSpacing(6)

        self._st_min_to_tray = QCheckBox("Minimize to system tray when closing")
        self._st_min_to_tray.setChecked(bool(get_setting("minimize_to_tray", True)))
        ab.addWidget(self._st_min_to_tray)

        self._st_startup_toast = QCheckBox("Show startup notification")
        self._st_startup_toast.setChecked(bool(get_setting("show_startup_notification", True)))
        ab.addWidget(self._st_startup_toast)

        self._st_start_with_windows = QCheckBox("Start with Windows")
        self._st_start_with_windows.setChecked(bool(get_setting("start_with_windows", False)))
        ab.addWidget(self._st_start_with_windows)

        self._st_auto_start_monitors = QCheckBox("Auto-start monitors on application launch")
        self._st_auto_start_monitors.setChecked(bool(get_setting("auto_start_monitors", True)))
        ab.addWidget(self._st_auto_start_monitors)

        self._st_media_use_gpu = QCheckBox("Use GPU (NVENC) for media conversion by default")
        self._st_media_use_gpu.setChecked(bool(get_setting("media_use_gpu", False)))
        ab.addWidget(self._st_media_use_gpu)

        self._st_media_autostart = QCheckBox("Auto-start media queue when jobs are added")
        self._st_media_autostart.setChecked(bool(get_setting("media_queue_autostart", False)))
        ab.addWidget(self._st_media_autostart)

        media_row = QHBoxLayout()
        media_row.setSpacing(10)
        media_row.addWidget(QLabel("Media conversions in parallel:"))
        self._st_media_parallel = QSpinBox()
        self._st_media_parallel.setRange(1, 8)
        try:
            self._st_media_parallel.setValue(int(get_setting("media_max_parallel", 1) or 1))
        except Exception:
            self._st_media_parallel.setValue(1)
        media_row.addWidget(self._st_media_parallel)
        media_row.addStretch(1)
        ab.addLayout(media_row)

        save_row = QHBoxLayout()
        save_row.addStretch(1)
        self._st_app_save = QPushButton("Save")
        self._st_app_save.setObjectName("PrimaryButton")
        self._st_app_save.clicked.connect(self._settings_save_app)
        save_row.addWidget(self._st_app_save)
        ab.addLayout(save_row)

        inner_layout.addWidget(app_box)

        # Email Notifications (parity with Tk backup)
        email_box = QGroupBox("Email Notifications")
        el = QVBoxLayout(email_box)
        el.setContentsMargins(12, 10, 12, 10)
        el.setSpacing(10)

        email_hint = QLabel("Configure SMTP settings for email notifications (optional).")
        email_hint.setStyleSheet("color:#9aa0a9;")
        email_hint.setWordWrap(True)
        el.addWidget(email_hint)

        form = QFormLayout()
        form.setLabelAlignment(Qt.AlignRight | Qt.AlignVCenter)
        form.setFormAlignment(Qt.AlignLeft | Qt.AlignTop)
        form.setHorizontalSpacing(10)
        form.setVerticalSpacing(10)

        self._st_smtp_server = QLineEdit()
        self._st_smtp_server.setPlaceholderText("e.g., smtp.gmail.com")
        self._st_smtp_server.setText(str(get_setting("smtp_server", "") or ""))
        form.addRow("SMTP Server", self._st_smtp_server)

        self._st_smtp_port = QSpinBox()
        self._st_smtp_port.setRange(1, 65535)
        try:
            self._st_smtp_port.setValue(int(get_setting("smtp_port", 587) or 587))
        except Exception:
            self._st_smtp_port.setValue(587)
        form.addRow("SMTP Port", self._st_smtp_port)

        self._st_smtp_username = QLineEdit()
        self._st_smtp_username.setPlaceholderText("your-email@example.com")
        self._st_smtp_username.setText(str(get_setting("smtp_username", "") or ""))
        form.addRow("SMTP Username", self._st_smtp_username)

        self._st_smtp_password = QLineEdit()
        self._st_smtp_password.setEchoMode(QLineEdit.Password)
        self._st_smtp_password.setPlaceholderText("SMTP password / app password")
        self._st_smtp_password.setText(str(get_setting("smtp_password", "") or ""))
        form.addRow("SMTP Password", self._st_smtp_password)

        self._st_sender_email = QLineEdit()
        self._st_sender_email.setPlaceholderText("sender@example.com")
        self._st_sender_email.setText(str(get_setting("sender_email", "") or ""))
        form.addRow("Sender Email", self._st_sender_email)

        el.addLayout(form)

        btn_row = QHBoxLayout()
        btn_row.addStretch(1)
        self._st_smtp_test = QPushButton("Test SMTP Connection")
        self._st_smtp_test.setObjectName("PrimaryButton")
        self._st_smtp_test.clicked.connect(self._settings_test_smtp)
        btn_row.addWidget(self._st_smtp_test)

        self._st_smtp_save = QPushButton("Save")
        self._st_smtp_save.setObjectName("PrimaryButton")
        self._st_smtp_save.clicked.connect(self._settings_save_email)
        btn_row.addWidget(self._st_smtp_save)

        self._st_smtp_clear = QPushButton("Clear")
        self._st_smtp_clear.clicked.connect(self._settings_clear_email)
        btn_row.addWidget(self._st_smtp_clear)

        el.addLayout(btn_row)
        inner_layout.addWidget(email_box)

        inner_layout.addStretch(1)

        self._settings_refresh_tools()
        self._settings_refresh_ai()
        return host

    def _settings_save_app(self):
        s = getattr(self.backend, "settings_manager", None)
        if not s:
            return
        try:
            theme_mode = "dark"
            accent = "blue"
            try:
                if hasattr(self, "_st_theme_mode") and isinstance(getattr(self, "_st_theme_mode"), QComboBox):
                    theme_mode = str(self._st_theme_mode.currentData() or self._st_theme_mode.currentText() or "dark").strip().lower()
            except Exception:
                theme_mode = "dark"
            try:
                if hasattr(self, "_st_color_theme") and isinstance(getattr(self, "_st_color_theme"), QComboBox):
                    accent = str(self._st_color_theme.currentData() or self._st_color_theme.currentText() or "blue").strip().lower()
            except Exception:
                accent = "blue"
            s.set_setting("theme", theme_mode)
            s.set_setting("color_theme", accent)
            s.set_setting("notifications_enabled", bool(self._st_notif_enabled.isChecked()))
            s.set_setting("notification_sound", bool(self._st_notif_sound.isChecked()))
            s.set_setting("minimize_to_tray", bool(self._st_min_to_tray.isChecked()))
            s.set_setting("show_startup_notification", bool(self._st_startup_toast.isChecked()))
            s.set_setting("start_with_windows", bool(self._st_start_with_windows.isChecked()))
            s.set_setting("auto_start_monitors", bool(self._st_auto_start_monitors.isChecked()))
            s.set_setting("media_use_gpu", bool(self._st_media_use_gpu.isChecked()))
            s.set_setting("media_queue_autostart", bool(self._st_media_autostart.isChecked()))
            s.set_setting("media_max_parallel", int(self._st_media_parallel.value()))
            try:
                from core.windows_autostart import set_start_with_windows

                set_start_with_windows(enabled=bool(self._st_start_with_windows.isChecked()))
            except Exception:
                # Non-fatal; still persist the setting.
                pass
            self._settings_preview_theme()
        except Exception as e:
            QMessageBox.critical(self, "Settings", str(e))
            return
        QMessageBox.information(self, "Settings", "Saved.")

    def _settings_preview_theme(self, *_args):
        try:
            from qt_app.styles import apply_app_theme
        except Exception:
            return
        try:
            theme_mode = "dark"
            accent = "blue"
            if hasattr(self, "_st_theme_mode") and isinstance(getattr(self, "_st_theme_mode"), QComboBox):
                theme_mode = str(self._st_theme_mode.currentData() or self._st_theme_mode.currentText() or "dark").strip().lower()
            if hasattr(self, "_st_color_theme") and isinstance(getattr(self, "_st_color_theme"), QComboBox):
                accent = str(self._st_color_theme.currentData() or self._st_color_theme.currentText() or "blue").strip().lower()
            app = QApplication.instance()
            if app:
                apply_app_theme(app, theme=theme_mode, accent=accent)
            try:
                self._apply_inline_theme_overrides()
            except Exception:
                pass
            try:
                self._refresh_brand_logo_theme()
            except Exception:
                pass
            try:
                self._refresh_hamburger_icon()
            except Exception:
                pass
            try:
                active_key = str(getattr(self, "_active_page_key", "") or "")
                for k, b in getattr(self, "_nav_buttons", {}).items():
                    try:
                        if hasattr(b, "set_active"):
                            b.set_active(str(k) == active_key)  # type: ignore[attr-defined]
                        elif hasattr(b, "refresh_theme"):
                            b.refresh_theme()  # type: ignore[attr-defined]
                    except Exception:
                        pass
            except Exception:
                pass
            try:
                for c in getattr(self, "_monitor_cards", {}).values():
                    try:
                        if hasattr(c, "refresh_theme"):
                            c.refresh_theme()  # type: ignore[attr-defined]
                    except Exception:
                        pass
            except Exception:
                pass
            try:
                self._refresh_cloud_sync_theme()
            except Exception:
                pass
            try:
                self._refresh_ai_hub_theme()
            except Exception:
                pass
            try:
                self._ft_apply_home_theme()
            except Exception:
                pass
            try:
                self._wa_apply_theme_styles(refresh_chat=True)
            except Exception:
                pass
            try:
                self._reload_tasks()
            except Exception:
                pass
            try:
                self._reload_ai_rules()
            except Exception:
                pass
            try:
                self._wf_apply_theme_styles()
            except Exception:
                pass
        except Exception:
            pass

    def _settings_save_workflow(self):
        s = getattr(self.backend, "settings_manager", None)
        if not s:
            return
        try:
            wf_settings = s.get_workflow_settings()
        except Exception:
            wf_settings = {}
        wf_settings = dict(wf_settings or {})
        try:
            wf_settings["max_output_tokens"] = int(getattr(self, "_st_wf_max_tokens").value())
        except Exception:
            pass
        try:
            s.save_workflow_settings(wf_settings)
        except Exception as e:
            QMessageBox.critical(self, "Settings", str(e))
            return
        QMessageBox.information(self, "Settings", "Workflow settings saved.")

    def _settings_update_workflow_warning(self, *_args):
        if not hasattr(self, "_st_wf_warn"):
            return
        try:
            ctx = 0
            if hasattr(self, "_st_ai_ctx") and isinstance(getattr(self, "_st_ai_ctx"), QComboBox):
                ctx = int(str(self._st_ai_ctx.currentText()).strip() or "0")
            max_tokens = 0
            if hasattr(self, "_st_wf_max_tokens"):
                max_tokens = int(getattr(self, "_st_wf_max_tokens").value())
            if ctx and max_tokens > ctx:
                self._st_wf_warn.setText(
                    f"Max output tokens exceeds current AI context size ({ctx}). "
                    "Output may truncate. Reload the model after changing context size."
                )
            else:
                self._st_wf_warn.setText("")
        except Exception:
            try:
                self._st_wf_warn.setText("")
            except Exception:
                pass

    def _settings_save_email(self):
        s = getattr(self.backend, "settings_manager", None)
        if not s:
            return
        try:
            s.set_setting("smtp_server", (self._st_smtp_server.text() or "").strip())
            s.set_setting("smtp_port", int(self._st_smtp_port.value()))
            s.set_setting("smtp_username", (self._st_smtp_username.text() or "").strip())
            s.set_setting("smtp_password", (self._st_smtp_password.text() or "").strip())
            s.set_setting("sender_email", (self._st_sender_email.text() or "").strip())
        except Exception as e:
            QMessageBox.critical(self, "Email Settings", str(e))
            return
        QMessageBox.information(self, "Email Settings", "Saved.")

    def _settings_clear_email(self):
        try:
            self._st_smtp_server.setText("")
            self._st_smtp_port.setValue(587)
            self._st_smtp_username.setText("")
            self._st_smtp_password.setText("")
            self._st_sender_email.setText("")
            self._settings_save_email()
        except Exception:
            pass

    def _settings_test_smtp(self):
        smtp_settings = {
            "smtp_server": (self._st_smtp_server.text() or "").strip(),
            "smtp_port": int(self._st_smtp_port.value()),
            "smtp_username": (self._st_smtp_username.text() or "").strip(),
            "smtp_password": (self._st_smtp_password.text() or "").strip(),
            "sender_email": (self._st_sender_email.text() or "").strip(),
        }
        if not all([smtp_settings["smtp_server"], smtp_settings["smtp_username"], smtp_settings["smtp_password"], smtp_settings["sender_email"]]):
            QMessageBox.warning(self, "Email Settings", "Please fill in all SMTP fields before testing.")
            return

        self._st_smtp_test.setEnabled(False)
        self._st_smtp_test.setText("Testing…")

        dlg = QDialog(self)
        dlg.setWindowTitle("SMTP Test")
        dlg.setModal(True)
        dlg.setMinimumWidth(520)
        lay = QVBoxLayout(dlg)
        lay.setContentsMargins(16, 14, 16, 14)
        lay.setSpacing(10)
        st = QLabel("Testing SMTP connection…")
        st.setWordWrap(True)
        st.setStyleSheet("color:#9aa0a9;")
        lay.addWidget(st)
        bar = QProgressBar()
        bar.setRange(0, 0)
        lay.addWidget(bar)
        btns = QDialogButtonBox(QDialogButtonBox.Close)
        btns.button(QDialogButtonBox.Close).setEnabled(False)
        lay.addWidget(btns)

        worker = _QtSmtpTestWorker(smtp_settings)
        thread = QThread(dlg)
        worker.moveToThread(thread)
        thread.started.connect(worker.run)
        worker.finished.connect(lambda ok, msg: self._settings_on_smtp_test_done(dlg, btns, st, ok, msg))
        worker.finished.connect(thread.quit)
        worker.finished.connect(worker.deleteLater)
        thread.finished.connect(thread.deleteLater)
        thread.start()
        dlg.exec()

        self._st_smtp_test.setEnabled(True)
        self._st_smtp_test.setText("Test SMTP Connection")

    def _settings_on_smtp_test_done(self, dlg: QDialog, btns: QDialogButtonBox, st: QLabel, ok: bool, msg: str):
        st.setText(("Connection successful.\n\n" if ok else "Connection failed.\n\n") + (msg or ""))
        try:
            btns.button(QDialogButtonBox.Close).setEnabled(True)
        except Exception:
            pass

    def _settings_set_ai_editable(self, editable: bool):
        """
        Prevent accidental changes while scrolling:
        - Default: read-only (controls disabled)
        - User clicks Edit to enable controls
        """
        editable = bool(editable)
        for attr in (
            "_st_ai_vision_model",
            "_st_ai_text_model",
            "_st_ai_target",
            "_st_ai_gpu_layers",
            "_st_ai_threads",
            "_st_ai_ctx",
            "_st_ai_batch",
            "_st_ai_img",
            "_st_ai_flash_attn",
            "_st_ai_rename_max",
            "_st_ai_rename_prompt",
        ):
            w = getattr(self, attr, None)
            if w is None:
                continue
            try:
                w.setEnabled(editable)
            except Exception:
                pass

        # Save should only be clickable in edit mode.
        try:
            if hasattr(self, "_st_ai_save"):
                self._st_ai_save.setEnabled(editable)
        except Exception:
            pass

        # Keep reload available (users may want to reload after changing model files outside the UI).
        try:
            if hasattr(self, "_st_ai_edit"):
                self._st_ai_edit.setText("Lock" if editable else "Edit")
                self._st_ai_edit.setChecked(editable)
        except Exception:
            pass

    def _settings_refresh_ai_context_presets(self, *_args):
        if not hasattr(self, "_st_ai_ctx") or not isinstance(getattr(self, "_st_ai_ctx"), QComboBox):
            return

        settings = getattr(self.backend, "settings_manager", None)
        get_setting = (settings.get_setting if settings else (lambda _k, d=None: d))

        vision_id = ""
        text_id = ""
        try:
            if hasattr(self, "_st_ai_vision_model") and isinstance(getattr(self, "_st_ai_vision_model"), QComboBox):
                vision_id = str(self._st_ai_vision_model.currentData() or "").strip()
        except Exception:
            vision_id = ""
        try:
            if hasattr(self, "_st_ai_text_model") and isinstance(getattr(self, "_st_ai_text_model"), QComboBox):
                text_id = str(self._st_ai_text_model.currentData() or "").strip()
        except Exception:
            text_id = ""

        try:
            from core.ai_model_catalog import get_model_max_context

            v_ctx = int(get_model_max_context(vision_id or None) or 65536)
            t_ctx = int(get_model_max_context(text_id or None) or 65536) if text_id else v_ctx
            max_ctx = int(min(v_ctx, t_ctx) or 65536)
        except Exception:
            max_ctx = 65536

        presets = [
            512,
            1024,
            2048,
            4096,
            8192,
            12288,
            16384,
            32768,
            65536,
            80000,
            96000,
            112000,
            128000,
            160000,
            192000,
            224000,
            256000,
        ]
        presets = [p for p in presets if p <= max_ctx]
        if not presets:
            presets = [max_ctx]

        def _to_int(w, default: int):
            try:
                if isinstance(w, QComboBox):
                    return int(str(w.currentText()).strip())
            except Exception:
                pass
            return int(default)

        default_ctx = int(get_setting("ai_context_size", 2048) or 2048)
        current = _to_int(self._st_ai_ctx, default_ctx)
        if current not in presets:
            lower = [p for p in presets if p <= current]
            current = max(lower) if lower else min(presets)

        self._st_ai_ctx.blockSignals(True)
        self._st_ai_ctx.clear()
        for p in presets:
            self._st_ai_ctx.addItem(str(p))
        self._st_ai_ctx.setCurrentText(str(current))
        self._st_ai_ctx.blockSignals(False)
        try:
            self._settings_update_workflow_warning()
        except Exception:
            pass

    def _settings_apply_ai_params_to_manager(self):
        settings = getattr(self.backend, "settings_manager", None)
        if not settings:
            return

        def _to_int(w, default: int):
            try:
                if isinstance(w, QComboBox):
                    return int(str(w.currentText()).strip())
                if isinstance(w, QSlider):
                    return int(w.value())
                if isinstance(w, QSpinBox):
                    return int(w.value())
            except Exception:
                pass
            return int(default)

        gpu_layers = _to_int(getattr(self, "_st_ai_gpu_layers", None), int(settings.get_setting("ai_gpu_layers", 0) or 0))
        threads = _to_int(getattr(self, "_st_ai_threads", None), int(settings.get_setting("ai_threads", 8) or 8))
        ctx = _to_int(getattr(self, "_st_ai_ctx", None), int(settings.get_setting("ai_context_size", 2048) or 2048))
        batch = _to_int(getattr(self, "_st_ai_batch", None), int(settings.get_setting("ai_batch_size", 512) or 512))
        img = _to_int(getattr(self, "_st_ai_img", None), int(settings.get_setting("ai_image_size", 512) or 512))
        flash_attn = str(settings.get_setting("ai_flash_attn_type", "disabled") or "disabled").strip().lower()
        try:
            if hasattr(self, "_st_ai_flash_attn") and isinstance(getattr(self, "_st_ai_flash_attn"), QComboBox):
                flash_attn = str(self._st_ai_flash_attn.currentData() or self._st_ai_flash_attn.currentText() or "disabled").strip().lower()
        except Exception:
            pass

        # Stability guard for packaged desktop builds. The downloader must never crash
        # because older settings selected GPU layers or FlashAttention automatically.
        gpu_layers = 0
        flash_attn = "disabled"
        try:
            if hasattr(self, "_st_ai_gpu_layers") and isinstance(getattr(self, "_st_ai_gpu_layers"), QSlider):
                self._st_ai_gpu_layers.setValue(0)
            if hasattr(self, "_st_ai_flash_attn") and isinstance(getattr(self, "_st_ai_flash_attn"), QComboBox):
                idx = self._st_ai_flash_attn.findData("disabled")
                if idx >= 0:
                    self._st_ai_flash_attn.setCurrentIndex(idx)
        except Exception:
            pass

        try:
            # Model selection (catalog-driven, separate text + vision slots)
            vision_id = ""
            text_id = ""
            if hasattr(self, "_st_ai_vision_model") and isinstance(getattr(self, "_st_ai_vision_model"), QComboBox):
                vision_id = str(self._st_ai_vision_model.currentData() or "").strip()
            if hasattr(self, "_st_ai_text_model") and isinstance(getattr(self, "_st_ai_text_model"), QComboBox):
                text_id = str(self._st_ai_text_model.currentData() or "").strip()

            try:
                from core.ai_model_catalog import get_model_spec, get_model_max_context

                if vision_id:
                    v_spec = get_model_spec(vision_id)
                    # New vision keys
                    settings.set_setting("ai_vision_model_id", vision_id)
                    settings.set_setting("ai_vision_model_repo", v_spec.repo)
                    settings.set_setting("ai_vision_model_file", v_spec.model_file)
                    settings.set_setting("ai_vision_mmproj_file", v_spec.mmproj_file or "")
                    settings.set_setting("ai_vision_chat_format", v_spec.chat_format or "")
                    # Legacy single-slot keys (treat as vision)
                    settings.set_setting("ai_model_id", vision_id)
                    settings.set_setting("ai_model_repo", v_spec.repo)
                    settings.set_setting("ai_model_file", v_spec.model_file)
                    settings.set_setting("ai_mmproj_file", v_spec.mmproj_file or "")
                    settings.set_setting("ai_chat_format", v_spec.chat_format or "")

                if text_id:
                    t_spec = get_model_spec(text_id)
                    settings.set_setting("ai_text_model_id", text_id)
                    settings.set_setting("ai_text_model_repo", t_spec.repo)
                    settings.set_setting("ai_text_model_file", t_spec.model_file)
                    settings.set_setting("ai_text_mmproj_file", "")
                    settings.set_setting("ai_text_chat_format", t_spec.chat_format or "")

                # Shared n_ctx must be valid for BOTH selected models.
                v_ctx = int(get_model_max_context(vision_id or None) or 65536)
                t_ctx = int(get_model_max_context(text_id or None) or 65536) if text_id else v_ctx
                max_ctx = int(min(v_ctx, t_ctx) or 65536)
                if ctx > max_ctx:
                    ctx = max_ctx
            except Exception:
                pass

            try:
                if hasattr(self, "_st_ai_ctx") and isinstance(getattr(self, "_st_ai_ctx"), QComboBox):
                    if str(self._st_ai_ctx.currentText()) != str(ctx):
                        self._st_ai_ctx.setCurrentText(str(ctx))
            except Exception:
                pass

            settings.set_setting("ai_gpu_layers", gpu_layers)
            settings.set_setting("ai_threads", threads)
            settings.set_setting("ai_context_size", ctx)
            settings.set_setting("ai_batch_size", batch)
            settings.set_setting("ai_image_size", img)
            settings.set_setting("ai_flash_attn_type", flash_attn)
        except Exception:
            pass

        # Smart rename settings
        try:
            settings.set_setting("ai_rename_max_keywords", int(getattr(self, "_st_ai_rename_max", QSpinBox()).value()))
        except Exception:
            pass
        try:
            prompt = ""
            if hasattr(self, "_st_ai_rename_prompt"):
                prompt = str(self._st_ai_rename_prompt.toPlainText() or "")
            settings.set_setting("ai_rename_prompt", prompt)
        except Exception:
            pass

        # Apply to the live AI manager (takes effect on next load; reload to apply immediately).
        ai = getattr(self.backend, "ai_manager", None)
        if ai:
            try:
                try:
                    if hasattr(ai, "refresh_profiles_from_settings"):
                        ai.refresh_profiles_from_settings()
                except Exception:
                    pass

                ai.n_gpu_layers = gpu_layers
                ai.n_threads = threads
                ai.n_ctx = ctx
                ai.n_batch = batch
                ai.image_size = img
                try:
                    ai.flash_attn_type = flash_attn
                except Exception:
                    pass
                ai.rename_max_keywords = int(settings.get_setting("ai_rename_max_keywords", 8) or 8)
                ai.rename_prompt = str(settings.get_setting("ai_rename_prompt", "") or ai.rename_prompt)
            except Exception:
                pass

        try:
            settings.save_settings()
        except Exception:
            pass

    def _settings_save_ai_params(self):
        self._settings_apply_ai_params_to_manager()
        QMessageBox.information(self, "Settings", "AI settings saved.")

    def _settings_reload_ai(self):
        self._settings_apply_ai_params_to_manager()
        ai = getattr(self.backend, "ai_manager", None)
        if not ai:
            return
        try:
            if getattr(ai, "is_ready", False):
                ai.unload_model()
        except Exception:
            pass
        if self._ensure_ai_ready(title="Reload AI Model"):
            self._settings_refresh_ai()

    def _settings_refresh_ai(self):
        ai = getattr(self.backend, "ai_manager", None)
        if not ai:
            self._st_ai_status.setText("AI: unavailable")
            self._st_ai_load.setEnabled(False)
            self._st_ai_unload.setEnabled(False)
            if hasattr(self, "_st_ai_load_model"):
                self._st_ai_load_model.setEnabled(False)
            if hasattr(self, "_st_ai_reload"):
                self._st_ai_reload.setEnabled(False)
            return
        ready = bool(getattr(ai, "is_ready", False))
        err = str(getattr(ai, "load_error", "") or "").strip()
        model_label = ""
        try:
            mf = getattr(ai, "model_file", "") or getattr(ai, "MODEL_FILE", "")
            if mf:
                model_label = f" • {mf}"
        except Exception:
            model_label = ""
        if ready:
            self._st_ai_status.setText("AI: ready" + model_label)
        else:
            self._st_ai_status.setText("AI: not loaded" + model_label + (f" • {err}" if err else ""))
        self._st_ai_load.setEnabled(True)
        if hasattr(self, "_st_ai_load_model"):
            self._st_ai_load_model.setEnabled(True)
        self._st_ai_unload.setEnabled(bool(ready))
        if hasattr(self, "_st_ai_reload"):
            self._st_ai_reload.setEnabled(True)

    def _settings_selected_ai_kind(self) -> str:
        try:
            if hasattr(self, "_st_ai_target") and isinstance(getattr(self, "_st_ai_target"), QComboBox):
                kind = str(self._st_ai_target.currentData() or "vision").strip().lower() or "vision"
                return kind if kind in ("vision", "text") else "vision"
        except Exception:
            pass
        return "vision"

    def _settings_prepare_ai_kind(self) -> str:
        self._settings_apply_ai_params_to_manager()
        kind = self._settings_selected_ai_kind()
        ai = getattr(self.backend, "ai_manager", None)
        try:
            if ai and hasattr(ai, "select_kind"):
                ai.select_kind(kind)
        except Exception:
            pass
        return kind

    def _settings_load_ai(self):
        kind = self._settings_prepare_ai_kind()
        ai = getattr(self.backend, "ai_manager", None)
        if not ai:
            QMessageBox.warning(self, "AI Model", "AI is not available in this build.")
            return
        loader = _QtAIModelDownloadDialog(self, ai_manager=ai, kind=kind)
        if loader.exec() == QDialog.Accepted:
            QMessageBox.information(self, "AI Model", "Model files are downloaded and ready. Use Load Model only when you want to start AI features now.")
        self._settings_refresh_ai()

    def _settings_load_ai_into_memory(self):
        kind = self._settings_prepare_ai_kind()
        if self._ensure_ai_ready(title="Load AI Model", kind=kind):
            self._settings_refresh_ai()

    def _settings_unload_ai(self):
        ai = getattr(self.backend, "ai_manager", None)
        if not ai:
            return
        try:
            ai.unload_model()
        except Exception:
            pass
        self._settings_refresh_ai()

    def _settings_delete_ai_files(self):
        # Apply model selection to AI manager first (so we delete the selected model's files).
        try:
            self._settings_apply_ai_params_to_manager()
        except Exception:
            pass

        ai = getattr(self.backend, "ai_manager", None)
        if not ai:
            QMessageBox.warning(self, "AI Model", "AI manager not available.")
            return
        try:
            kind = "vision"
            if hasattr(self, "_st_ai_target") and isinstance(getattr(self, "_st_ai_target"), QComboBox):
                kind = str(self._st_ai_target.currentData() or "vision").strip().lower() or "vision"
            if hasattr(ai, "select_kind"):
                ai.select_kind(kind)
        except Exception:
            pass

        model_file = str(getattr(ai, "model_file", "") or getattr(ai, "MODEL_FILE", "") or "").strip()
        mmproj_file = str(getattr(ai, "mmproj_file", "") or getattr(ai, "MMPROJ_FILE", "") or "").strip()
        models_dir = getattr(ai, "models_folder", None)
        model_id = str(getattr(ai, "model_id", "") or "").strip()
        if not model_file or not models_dir:
            QMessageBox.warning(self, "AI Model", "Could not determine model files to delete.")
            return

        try:
            from pathlib import Path
            import re
            import shutil

            model_root = Path(models_dir)
            safe = re.sub(r"[^A-Za-z0-9_.-]+", "_", model_id or Path(model_file).stem)[:80] or "default"
            model_dir = model_root / safe

            candidates: list[Path] = []
            if model_dir.exists():
                candidates.append(model_dir)
            legacy_model = model_root / model_file
            if legacy_model.exists():
                candidates.append(legacy_model)
            if mmproj_file:
                legacy_mmproj = model_root / mmproj_file
                if legacy_mmproj.exists():
                    candidates.append(legacy_mmproj)

            if not candidates:
                QMessageBox.information(self, "AI Model", "No downloaded files found for this model.")
                return
        except Exception:
            QMessageBox.warning(self, "AI Model", "Could not resolve model paths.")
            return

        msg = "Delete downloaded AI model?\n\n" + "\n".join(str(p) for p in candidates)
        if QMessageBox.question(self, "Delete AI Model Files", msg) != QMessageBox.Yes:
            return

        try:
            if getattr(ai, "is_ready", False):
                ai.unload_model()
        except Exception:
            pass

        deleted = 0
        errors: list[str] = []
        for p in candidates:
            try:
                if p.is_dir():
                    shutil.rmtree(p)
                    deleted += 1
                else:
                    p.unlink(missing_ok=True)  # type: ignore[arg-type]
                    deleted += 1
            except Exception as e:
                errors.append(f"{p}: {e}")

        self._settings_refresh_ai()
        if errors:
            QMessageBox.warning(self, "AI Model", f"Deleted {deleted} file(s), but some failed:\n\n" + "\n".join(errors))
        else:
            QMessageBox.information(self, "AI Model", f"Deleted {deleted} file(s).")

    def _settings_refresh_tools(self):
        try:
            from core.ffmpeg_manager import get_ffmpeg_exe, get_ffprobe_exe
            from core.tool_manager import ToolManager
            from core.lo_converter import LibreOfficeConverter

            tm = ToolManager()
            ff = get_ffmpeg_exe()
            fp = get_ffprobe_exe()
            self._st_ffmpeg.setText(f"ffmpeg: {'OK' if ff else 'missing'} • ffprobe: {'OK' if fp else 'missing'}")
            self._st_ffmpeg_path.setText(tm._paths.ffmpeg or "")  # type: ignore[attr-defined]
            try:
                self._st_ffprobe_path.setText(getattr(tm._paths, "ffprobe", "") or "")  # type: ignore[attr-defined]
            except Exception:
                self._st_ffprobe_path.setText("")
            try:
                if hasattr(self, "_st_ff_dl") and self._st_ff_dl:
                    # If ffprobe is present, FFmpeg tools are effectively ready.
                    self._st_ff_dl.setEnabled(not bool(fp))
                    self._st_ff_dl.setText("FFmpeg Tools Ready" if fp else "Download FFmpeg Tools (ffmpeg+ffprobe+ffplay)")
            except Exception:
                pass

            conv = LibreOfficeConverter()
            self._st_soffice.setText(f"LibreOffice: {'OK' if conv.is_available() else 'missing'}")
            self._st_soffice_path.setText((tm._paths.soffice or ""))  # type: ignore[attr-defined]
        except Exception:
            self._st_ffmpeg.setText("ffmpeg: ?")
            self._st_soffice.setText("LibreOffice: ?")

    def _settings_pick_ffmpeg(self):
        f, _ = QFileDialog.getOpenFileName(self, "Select ffmpeg executable", filter="ffmpeg (ffmpeg.exe)")
        if f:
            self._st_ffmpeg_path.setText(f)

    def _settings_save_ffmpeg(self):
        try:
            from core.tool_manager import ToolManager
            p = (self._st_ffmpeg_path.text() or "").strip()
            ToolManager().set_ffmpeg(p or None)
            self._settings_refresh_tools()
        except Exception as e:
            QMessageBox.critical(self, "Save ffmpeg", str(e))

    def _settings_clear_ffmpeg(self):
        try:
            from core.tool_manager import ToolManager
            ToolManager().set_ffmpeg(None)
            self._st_ffmpeg_path.setText("")
            self._settings_refresh_tools()
        except Exception:
            pass

    def _settings_pick_ffprobe(self):
        f, _ = QFileDialog.getOpenFileName(self, "Select ffprobe executable", filter="ffprobe (ffprobe.exe)")
        if f:
            self._st_ffprobe_path.setText(f)

    def _settings_save_ffprobe(self):
        try:
            from core.tool_manager import ToolManager
            p = (self._st_ffprobe_path.text() or "").strip()
            ToolManager().set_ffprobe(p or None)
            self._settings_refresh_tools()
        except Exception as e:
            QMessageBox.critical(self, "Save ffprobe", str(e))

    def _settings_clear_ffprobe(self):
        try:
            from core.tool_manager import ToolManager
            ToolManager().set_ffprobe(None)
            self._st_ffprobe_path.setText("")
            self._settings_refresh_tools()
        except Exception:
            pass

    def _settings_download_ffmpeg_tools(self):
        dlg = _QtFfmpegToolsDownloadDialog(self)
        dlg.exec()
        self._settings_refresh_tools()

    def _settings_pick_soffice(self):
        f, _ = QFileDialog.getOpenFileName(self, "Select LibreOffice soffice executable", filter="soffice (soffice.exe)")
        if f:
            self._st_soffice_path.setText(f)

    def _settings_save_soffice(self):
        try:
            from core.tool_manager import ToolManager
            p = (self._st_soffice_path.text() or "").strip()
            ToolManager().set_soffice(p or None)
            self._settings_refresh_tools()
        except Exception as e:
            QMessageBox.critical(self, "Save LibreOffice", str(e))

    def _settings_clear_soffice(self):
        try:
            from core.tool_manager import ToolManager
            ToolManager().set_soffice(None)
            self._st_soffice_path.setText("")
            self._settings_refresh_tools()
        except Exception:
            pass

    def _settings_download_lo(self):
        dlg = _QtLibreOfficeDownloadDialog(self)
        dlg.exec()
        self._settings_refresh_tools()

    def _build_cloud_sync_page(self) -> QWidget:
        host = QFrame()
        host.setObjectName("PageHost")
        layout = QVBoxLayout(host)
        layout.setContentsMargins(18, 18, 18, 18)
        layout.setSpacing(14)
        t = _ui_theme_tokens()
        cloud_muted = t["muted"]
        self._cloud_muted_widgets = []
        self._cloud_icon_buttons = []

        def _cloud_set_btn_icon(btn: QPushButton, stem: str, *, primary: bool = False) -> None:
            try:
                btn.setProperty("fg_cloud_icon_stem", str(stem))
                btn.setProperty("fg_cloud_icon_primary", bool(primary))
                self._cloud_icon_buttons.append(btn)
                ic = self._cloud_icon(
                    stem,
                    size=18,
                    color=t["icon_on_solid"] if primary else t["icon"],
                )
                if not ic.isNull():
                    btn.setIcon(ic)
                    btn.setIconSize(QSize(18, 18))
            except Exception:
                pass

        def _cs_btn(
            icon_stem: str,
            tooltip: str,
            *,
            text: str | None = None,
            primary: bool = False,
            fixed: int | None = None,
        ) -> QPushButton:
            b = QPushButton("" if text is None else str(text))
            b.setObjectName("PrimaryButton" if primary else "SecondaryButton")
            _cloud_set_btn_icon(b, icon_stem, primary=primary)
            if fixed is not None:
                b.setFixedWidth(int(fixed))
            b.setMinimumHeight(34)
            if tooltip:
                b.setToolTip(str(tooltip))
            return b

        def _cs_icon(stem: str) -> QIcon:
            return self._cloud_icon(stem, size=18, color=t["icon"])

        header = QFrame()
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(0, 0, 0, 0)
        header_layout.setSpacing(10)
        title = QLabel("Cloud Sync")
        title.setObjectName("PageTitle")
        header_layout.addWidget(title)
        sub = QLabel("Connect OneDrive / Google Drive, test access, and upload/download files (MVP).")
        sub.setObjectName("PageSubTitle")
        header_layout.addWidget(sub, 1)
        layout.addWidget(header)

        card = QFrame()
        card.setObjectName("PageCard")
        c = QVBoxLayout(card)
        c.setContentsMargins(16, 14, 16, 14)
        c.setSpacing(12)

        tabs = QTabWidget()
        tabs.setDocumentMode(True)

        # Jobs (queued operations with progress/cancel).
        jobs_box = QGroupBox("Jobs")
        jobs_l = QVBoxLayout(jobs_box)
        jobs_l.setContentsMargins(10, 10, 10, 10)
        jobs_l.setSpacing(8)

        self._cloud_jobs_table = QTableWidget(0, 4)
        self._cloud_jobs_table.setObjectName("ResultsTree")
        self._cloud_jobs_table.setHorizontalHeaderLabels(["Action", "Status", "Progress", ""])
        self._cloud_jobs_table.verticalHeader().setVisible(False)
        self._cloud_jobs_table.setShowGrid(False)
        self._cloud_jobs_table.setAlternatingRowColors(True)
        self._cloud_jobs_table.setWordWrap(False)
        self._cloud_jobs_table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self._cloud_jobs_table.setSelectionMode(QAbstractItemView.SingleSelection)
        self._cloud_jobs_table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self._cloud_jobs_table.horizontalHeader().setStretchLastSection(False)
        self._cloud_jobs_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.Stretch)
        self._cloud_jobs_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.Stretch)
        self._cloud_jobs_table.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeToContents)
        self._cloud_jobs_table.horizontalHeader().setSectionResizeMode(3, QHeaderView.ResizeToContents)
        self._cloud_jobs_table.setMinimumHeight(140)
        jobs_l.addWidget(self._cloud_jobs_table, 1)

        jobs_btn_row = QHBoxLayout()
        jobs_btn_row.setSpacing(10)
        jobs_btn_row.addStretch(1)
        btn_clear_jobs = QPushButton("Clear Finished")
        btn_clear_jobs.setObjectName("SecondaryButton")
        _cloud_set_btn_icon(btn_clear_jobs, "clear")
        jobs_btn_row.addWidget(btn_clear_jobs)
        jobs_l.addLayout(jobs_btn_row)

        def _clear_finished_jobs():
            tbl = self._cloud_jobs_table
            # Remove rows that are done/error/cancelled.
            for r in range(tbl.rowCount() - 1, -1, -1):
                st = tbl.item(r, 1).text() if tbl.item(r, 1) else ""
                if st.lower().startswith(("done", "error", "cancelled")):
                    tbl.removeRow(r)

        btn_clear_jobs.clicked.connect(_clear_finished_jobs)
        # Splitter: tabs (top) + jobs (bottom). Jobs is collapsible/hidden by default.
        splitter = QSplitter(Qt.Vertical)
        splitter.setObjectName("CloudSyncSplitter")
        splitter.setChildrenCollapsible(True)
        splitter.setHandleWidth(6)
        self._cloud_splitter = splitter
        splitter.addWidget(tabs)
        splitter.addWidget(jobs_box)
        splitter.setStretchFactor(0, 1)
        splitter.setStretchFactor(1, 0)
        c.addWidget(splitter, 1)
        try:
            # Hide jobs by default, but keep splitter handle for quick expansion.
            splitter.setSizes([1000, 0])
        except Exception:
            pass

        # Status + Jobs toggle in the tab bar corner (saves vertical space).
        self._cloud_status = QLabel("Ready.")
        self._cloud_status.setStyleSheet(f"color:{cloud_muted};")
        self._cloud_muted_widgets.append(self._cloud_status)

        # Use a normal QPushButton here (QToolButton + arrow styling can overlap text in tight corner widgets).
        self._cloud_jobs_toggle = QPushButton("Jobs ▸")
        self._cloud_jobs_toggle.setObjectName("SecondaryButton")
        self._cloud_jobs_toggle.setCheckable(True)
        self._cloud_jobs_toggle.setChecked(False)  # hidden by default
        self._cloud_jobs_toggle.setCursor(Qt.PointingHandCursor)
        self._cloud_jobs_toggle.setMinimumWidth(86)
        try:
            self._cloud_jobs_toggle.setFixedHeight(30)
        except Exception:
            pass

        corner = QFrame()
        corner_l = QHBoxLayout(corner)
        corner_l.setContentsMargins(0, 0, 0, 0)
        corner_l.setSpacing(10)
        corner_l.addStretch(1)
        corner_l.addWidget(self._cloud_status, 0, Qt.AlignRight | Qt.AlignVCenter)
        corner_l.addWidget(self._cloud_jobs_toggle, 0, Qt.AlignRight | Qt.AlignVCenter)
        tabs.setCornerWidget(corner, Qt.TopRightCorner)

        def _set_jobs_visible(vis: bool):
            vis = bool(vis)
            try:
                self._cloud_jobs_toggle.setText("Jobs ▾" if vis else "Jobs ▸")
            except Exception:
                pass
            try:
                sizes = splitter.sizes()
                if vis:
                    if len(sizes) >= 2 and sizes[1] < 20:
                        splitter.setSizes([max(1, sizes[0] - 220), 220])
                else:
                    splitter.setSizes([sum(sizes), 0] if len(sizes) >= 2 else [1000, 0])
            except Exception:
                pass

        self._cloud_jobs_toggle.toggled.connect(_set_jobs_visible)

        # Keep thread refs alive for the lifetime of each job (avoid GC edge-cases).
        if not hasattr(self, "_cloud_jobs"):
            self._cloud_jobs: list[tuple[QThread, QObject]] = []

        # Lazily create providers (keeps optional deps import failures scoped to this page).
        def _ensure_providers():
            if getattr(self, "_cloud_mgr", None) is not None and getattr(self, "_onedrive", None) is not None:
                return
            # Share a single manager across the app so other pages can reuse tokens and providers.
            self._cloud_mgr = self.backend.get_cloud_sync_manager()
            self._cloud_tokens = self._cloud_mgr.token_store
            self._onedrive = self._cloud_mgr.provider("onedrive")
            self._gdrive = self._cloud_mgr.provider("gdrive")

        def _open_cloud_sync_help():
            try:
                p = Path(__file__).resolve().parents[1] / "docs" / "CLOUD_SYNC.md"
                if p.exists():
                    QDesktopServices.openUrl(QUrl.fromLocalFile(str(p)))
                    return
            except Exception:
                pass
            QMessageBox.information(
                self,
                "Cloud Sync Help",
                "Open `docs/CLOUD_SYNC.md` for setup details.\n\n"
                "OneDrive/Google Drive require OAuth sign-in. Fylorra does not (and should not) collect your email/password.",
            )

        def _why_no_passwords():
            QMessageBox.information(
                self,
                "Why not email/password?",
                "OneDrive and Google Drive require OAuth for third‑party apps.\n\n"
                "Fylorra cannot safely take your email/password and sign in on your behalf.\n"
                "Instead, you click Connect and sign in using Microsoft/Google’s official secure login.\n\n"
                "After you connect once, Fylorra stores refresh tokens locally and you won't have to sign in every time.",
            )

        def _refresh_status():
            try:
                _ensure_providers()
            except Exception as e:
                self._cloud_status.setText(f"Cloud Sync unavailable: {e}")
                return
            try:
                od = "Connected" if self._onedrive.is_connected() else "Not connected"
                gd = "Connected" if self._gdrive.is_connected() else "Not connected"
                self._od_status.setText(f"OneDrive: {od}")
                self._gd_status.setText(f"Google Drive: {gd}")
            except Exception:
                pass

        def _run_cloud_job(*, action: str, fn, on_ok=None, on_err=None, status_sink=None, progress_sink=None, show_in_jobs: bool = True):
            worker = _QtCloudWorker(fn=fn)
            thread = QThread(self)
            worker.moveToThread(thread)
            thread.started.connect(worker.run)

            # Keep references until completion.
            try:
                self._cloud_jobs.append((thread, worker))
            except Exception:
                pass

            job_row = None
            job_progress: QProgressBar | None = None
            job_cancel: QPushButton | None = None

            if show_in_jobs:
                tbl = self._cloud_jobs_table
                job_row = tbl.rowCount()
                tbl.insertRow(job_row)
                tbl.setItem(job_row, 0, QTableWidgetItem(str(action or "Cloud Job")))
                tbl.setItem(job_row, 1, QTableWidgetItem("Starting…"))

                job_progress = QProgressBar()
                # Start indeterminate (some operations do an initial scan/count pass).
                job_progress.setRange(0, 0)
                job_progress.setValue(0)
                job_progress.setFormat("%p%")
                tbl.setCellWidget(job_row, 2, job_progress)

                job_cancel = QPushButton("Cancel")
                job_cancel.setObjectName("SecondaryButton")
                tbl.setCellWidget(job_row, 3, job_cancel)
                job_cancel.clicked.connect(lambda: (job_cancel.setEnabled(False), tbl.item(job_row, 1).setText("Cancelling…") if tbl.item(job_row, 1) else None, worker.cancel()))

            def _set_job_status(text: str):
                if job_row is None:
                    return
                try:
                    it = self._cloud_jobs_table.item(job_row, 1)
                    if it is not None:
                        it.setText(text)
                except Exception:
                    pass

            worker.status.connect(lambda s: (self._cloud_status.setText(s), _set_job_status(str(s)[:240] if s else "Working…")))
            if status_sink is not None:
                try:
                    worker.status.connect(status_sink)
                except Exception:
                    pass

            _last_prog_ts = 0.0
            _last_prog_val = -1

            def _on_prog(done: int, total: int, msg: str):
                nonlocal _last_prog_ts, _last_prog_val
                if msg:
                    self._cloud_status.setText(msg)
                    _set_job_status(msg[:240])
                if job_progress is None:
                    return
                try:
                    from time import monotonic

                    if not total or total <= 0:
                        job_progress.setRange(0, 0)
                        return

                    val = int(max(0, min(1.0, float(done) / float(total))) * 1000.0)
                    now = monotonic()
                    if done < total and now - _last_prog_ts < 0.12 and val == _last_prog_val:
                        return
                    if done < total and now - _last_prog_ts < 0.12 and (val - _last_prog_val) < 2:
                        return

                    job_progress.setRange(0, 1000)
                    job_progress.setValue(val)
                    _last_prog_ts = now
                    _last_prog_val = val
                except Exception:
                    pass

            worker.progress.connect(_on_prog)
            if progress_sink is not None:
                try:
                    worker.progress.connect(progress_sink)
                except Exception:
                    pass

            def _finish_ok(res):
                if job_progress is not None:
                    try:
                        job_progress.setRange(0, 1000)
                        job_progress.setValue(1000)
                    except Exception:
                        pass
                _set_job_status("Done.")
                if job_cancel is not None:
                    job_cancel.setEnabled(False)
                if callable(on_ok):
                    on_ok(res)

            def _finish_err(msg: str):
                if "cancelled" in (msg or "").lower():
                    _set_job_status("Cancelled.")
                else:
                    _set_job_status(f"Error: {msg}"[:240])
                if job_cancel is not None:
                    job_cancel.setEnabled(False)
                if callable(on_err):
                    on_err(msg)
                else:
                    QMessageBox.critical(self, "Cloud Sync", msg)

            worker.finished.connect(_finish_ok)
            worker.error.connect(_finish_err)

            worker.finished.connect(thread.quit)
            worker.error.connect(thread.quit)
            thread.finished.connect(thread.deleteLater)
            thread.finished.connect(worker.deleteLater)
            # Drop refs when done.
            def _cleanup():
                try:
                    self._cloud_jobs = [(t, w) for (t, w) in self._cloud_jobs if t is not thread]
                except Exception:
                    pass
            thread.finished.connect(_cleanup)
            thread.start()

        def _schedule_cloud_sync_upload_only(*, provider: str, local_folder: str, remote_base: str, include_subfolders: bool, dry_run: bool):
            local_folder = str(local_folder or "").strip()
            if not local_folder:
                QMessageBox.warning(self, "Schedule Cloud Sync", "Pick a local folder first.")
                return

            dlg = _EditScheduledTaskDialog(self)
            try:
                dlg.setWindowTitle("Schedule Cloud Sync")
                dlg.action_type.setCurrentText("cloud_sync_upload_only")
                dlg.action_type.setEnabled(False)
                dlg.target.setText(local_folder)
                dlg.target.setEnabled(False)
                dlg.title.setText(f"Cloud Sync ({provider}): {Path(local_folder).name}")
                if not dlg.time.text().strip():
                    dlg.time.setText("00:00")
            except Exception:
                pass

            if dlg.exec() != QDialog.Accepted:
                return

            try:
                task = dlg.task_dict()
                task["action_type"] = "cloud_sync_upload_only"
                task["target_path"] = local_folder
                task["action_params"] = {
                    "provider": str(provider or "").strip().lower(),
                    "remote_base": str(remote_base or "Fylorra Sync").strip(),
                    "include_subfolders": bool(include_subfolders),
                    "dry_run": bool(dry_run),
                }
                ok = bool(self.backend.monitor_manager.add_scheduled_task(task))
            except Exception as e:
                QMessageBox.critical(self, "Schedule Cloud Sync", str(e))
                return

            if not ok:
                QMessageBox.critical(self, "Schedule Cloud Sync", "Could not save scheduled task.")
                return

            try:
                self.backend.monitor_manager.scheduled_tasks.reload()
            except Exception:
                pass
            try:
                self.backend.monitor_manager.start_scheduled_tasks()
            except Exception:
                pass
            try:
                self._reload_tasks()
            except Exception:
                pass
            try:
                self._reload_ai_rules()
            except Exception:
                pass
            QMessageBox.information(self, "Schedule Cloud Sync", "Scheduled. Manage it under Scheduled Tasks.")

        def _schedule_cloud_sync_full(
            *,
            action_type: str,
            provider: str,
            local_folder: str,
            remote_base: str,
            include_subfolders: bool,
            dry_run: bool,
            delete_policy: str = "ignore",
            conflict_policy: str = "keep_both",
        ):
            local_folder = str(local_folder or "").strip()
            if not local_folder:
                QMessageBox.warning(self, "Schedule Cloud Sync", "Pick a local folder first.")
                return

            at = str(action_type or "").strip() or "cloud_sync_two_way"
            if at not in ("cloud_sync_download_only", "cloud_sync_two_way"):
                at = "cloud_sync_two_way"

            dlg = _EditScheduledTaskDialog(self)
            try:
                dlg.setWindowTitle("Schedule Cloud Sync")
                dlg.action_type.setCurrentText(at)
                dlg.action_type.setEnabled(False)
                dlg.target.setText(local_folder)
                dlg.target.setEnabled(False)
                dlg.title.setText(f"Cloud Sync ({provider}): {Path(local_folder).name}")
                if not dlg.time.text().strip():
                    dlg.time.setText("00:00")
            except Exception:
                pass

            if dlg.exec() != QDialog.Accepted:
                return

            try:
                task = dlg.task_dict()
                task["action_type"] = at
                task["target_path"] = local_folder
                task["action_params"] = {
                    "provider": str(provider or "").strip().lower(),
                    "remote_base": str(remote_base or "Fylorra Sync").strip(),
                    "include_subfolders": bool(include_subfolders),
                    "dry_run": bool(dry_run),
                    "delete_policy": str(delete_policy or "ignore").strip().lower(),
                    "conflict_policy": str(conflict_policy or "keep_both").strip().lower(),
                }
                ok = bool(self.backend.monitor_manager.add_scheduled_task(task))
            except Exception as e:
                QMessageBox.critical(self, "Schedule Cloud Sync", str(e))
                return

            if not ok:
                QMessageBox.critical(self, "Schedule Cloud Sync", "Could not save scheduled task.")
                return

            try:
                self.backend.monitor_manager.scheduled_tasks.reload()
            except Exception:
                pass
            try:
                self.backend.monitor_manager.start_scheduled_tasks()
            except Exception:
                pass
            try:
                self._reload_tasks()
            except Exception:
                pass
            try:
                self._reload_ai_rules()
            except Exception:
                pass
            QMessageBox.information(self, "Schedule Cloud Sync", "Scheduled. Manage it under Scheduled Tasks.")

        # OneDrive tab
        od_tab = QWidget()
        od_layout = QVBoxLayout(od_tab)
        od_layout.setContentsMargins(10, 10, 10, 10)
        od_layout.setSpacing(10)

        od_hint = QHBoxLayout()
        od_hint.setSpacing(10)
        od_hint_lbl = QLabel("Uses Microsoft secure sign-in (OAuth). Fylorra never sees your password.")
        od_hint_lbl.setStyleSheet(f"color:{cloud_muted};")
        self._cloud_muted_widgets.append(od_hint_lbl)
        od_hint.addWidget(od_hint_lbl, 1)
        # od_help = _cs_btn("browse", "Help / setup instructions", text=None, fixed=38)
        # od_help.clicked.connect(_open_cloud_sync_help)
        # od_hint.addWidget(od_help)
        # od_why = _cs_btn("search", "Why OAuth is required", text="?", fixed=38)
        # od_why.clicked.connect(_why_no_passwords)
        # od_hint.addWidget(od_why)
        od_layout.addLayout(od_hint)

        od_msg = QTextBrowser()
        od_msg.setPlaceholderText("OneDrive device login instructions will appear here during Connect.")
        od_msg.setFixedHeight(70)
        od_msg.setVisible(False)
        self._od_device_prompt_shown = False

        od_header = QHBoxLayout()
        od_header.setSpacing(10)
        self._od_status = QLabel("OneDrive: Not connected")
        self._od_status.setStyleSheet(f"color:{cloud_muted};")
        self._cloud_muted_widgets.append(self._od_status)
        od_header.addWidget(self._od_status, 1)
        btn = _cs_btn("connect", "Connect OneDrive", text=None, primary=True, fixed=38)

        def _od_connect():
            try:
                od_msg.setVisible(True)
                od_msg.setText("Starting OneDrive sign-in…")
                self._cloud_status.setText("Starting OneDrive sign-in…")
            except Exception:
                pass
            try:
                from core.cloud_sync.app_credentials import get_onedrive_client_id

                cid = str(get_onedrive_client_id(self.backend.settings_manager) or "").strip()
            except Exception:
                cid = str(self.backend.settings_manager.get_setting("onedrive_client_id", "") or "").strip()
            if not cid:
                QMessageBox.information(
                    self,
                    "OneDrive setup required",
                    "To connect OneDrive, Fylorra needs an Azure App Client ID.\n\n"
                    "If you are using the official Fylorra build, this should already be configured.\n"
                    "Otherwise, click Help and follow the setup, then paste the Client ID in Advanced settings.",
                )
                _open_cloud_sync_help()
                return
            _run_cloud_job(
                action="OneDrive: Connect",
                fn=lambda status, progress: (  # noqa: ARG005
                    _ensure_providers(),
                    self._onedrive.connect(
                        status_cb=lambda m: status(m or "Open https://microsoft.com/devicelogin and enter the code.")
                    ),
                )[-1],
                on_ok=lambda _res: (_refresh_status(), self._cloud_status.setText("OneDrive connected.")),
                status_sink=lambda s: (od_msg.setVisible(True), od_msg.setText(str(s)), _od_maybe_prompt(str(s))),
            )

        btn.clicked.connect(_od_connect)
        od_header.addWidget(btn)

        btn_disc = _cs_btn("disconnect", "Disconnect OneDrive", text=None, primary=False, fixed=38)
        btn_test = _cs_btn("test", "Test OneDrive access", text=None, primary=False, fixed=38)
        od_header.addWidget(btn_disc)
        od_header.addWidget(btn_test)

        # Auto popup device-code prompt when the message includes Code/URL.
        def _od_maybe_prompt(txt: str):
            if self._od_device_prompt_shown:
                return
            s = str(txt or "")
            if "Code:" not in s or "URL:" not in s:
                return
            try:
                code = ""
                url = ""
                for line in s.splitlines():
                    if line.strip().lower().startswith("code:"):
                        code = line.split(":", 1)[1].strip()
                    if line.strip().lower().startswith("url:"):
                        url = line.split(":", 1)[1].strip()
                if not code and not url:
                    return
                self._od_device_prompt_shown = True
                dlg = _QtDeviceCodeDialog(
                    title="Sign in to OneDrive",
                    url=url or "https://microsoft.com/devicelogin",
                    code=code,
                    message=s,
                    parent=self,
                )
                dlg.exec()
            except Exception:
                pass

        # (status_sink above calls _od_maybe_prompt)

        btn_disc = QPushButton("Disconnect")
        _cloud_set_btn_icon(btn_disc, "disconnect")
        btn_disc.clicked.connect(
            lambda: _run_cloud_job(
                action="OneDrive: Disconnect",
                fn=lambda status, progress: (  # noqa: ARG005
                    _ensure_providers(),
                    self._onedrive.disconnect(),
                    "Disconnected",
                )[-1],
                on_ok=lambda _res: (_refresh_status(), self._cloud_status.setText("OneDrive disconnected.")),
            )
        )
        od_header.addWidget(btn_disc)

        btn_test = QPushButton("Test")
        _cloud_set_btn_icon(btn_test, "test")
        btn_test.clicked.connect(
            lambda: _run_cloud_job(
                action="OneDrive: Test",
                fn=lambda status, progress: (  # noqa: ARG005
                    _ensure_providers(),
                    self._onedrive.test_connection(),
                )[-1],
                on_ok=lambda res: QMessageBox.information(self, "OneDrive", str(res)),
            )
        )
        od_header.addWidget(btn_test)
        od_layout.addLayout(od_header)

        def _od_disconnect():
            _run_cloud_job(
                action="OneDrive: Disconnect",
                fn=lambda status, progress: (  # noqa: ARG005
                    _ensure_providers(),
                    self._onedrive.disconnect(),
                )[-1],
                on_ok=lambda _r: (_refresh_status(), self._cloud_status.setText("OneDrive disconnected.")),
            )

        def _od_test():
            _run_cloud_job(
                action="OneDrive: Test",
                fn=lambda status, progress: (  # noqa: ARG005
                    _ensure_providers(),
                    self._onedrive.test_access(),
                )[-1],
                on_ok=lambda _r: QMessageBox.information(self, "OneDrive", "Access OK."),
                on_err=lambda m: QMessageBox.critical(self, "OneDrive", m or "Access failed."),
            )

        btn_disc.clicked.connect(_od_disconnect)
        btn_test.clicked.connect(_od_test)

        # Advanced (admin) config: hide by default to keep UX simple.
        od_adv_toggle = QToolButton()
        od_adv_toggle.setText("Advanced")
        od_adv_toggle.setCheckable(True)
        od_adv_toggle.setChecked(False)
        od_adv_toggle.setToolButtonStyle(Qt.ToolButtonTextBesideIcon)
        od_adv_toggle.setArrowType(Qt.RightArrow)
        od_layout.addWidget(od_adv_toggle, 0, Qt.AlignLeft)

        od_adv = QFrame()
        od_adv.setVisible(False)
        od_adv_l = QVBoxLayout(od_adv)
        od_adv_l.setContentsMargins(0, 0, 0, 0)
        od_adv_l.setSpacing(8)

        od_adv_note = QLabel("Admin setup: required only once to enable OneDrive OAuth.")
        od_adv_note.setStyleSheet(f"color:{cloud_muted};")
        self._cloud_muted_widgets.append(od_adv_note)
        od_adv_l.addWidget(od_adv_note)

        od_form = QFormLayout()
        od_form.setHorizontalSpacing(12)
        od_form.setVerticalSpacing(8)
        self._od_client_id = QLineEdit(str(self.backend.settings_manager.get_setting("onedrive_client_id", "") or ""))
        self._od_client_id.setPlaceholderText("Azure App Client ID (Application ID)")
        od_form.addRow("Client ID", self._od_client_id)
        self._od_tenant = QLineEdit(str(self.backend.settings_manager.get_setting("onedrive_tenant", "common") or "common"))
        self._od_tenant.setPlaceholderText("common (recommended)")
        od_form.addRow("Tenant", self._od_tenant)
        od_adv_l.addLayout(od_form)

        od_save = QPushButton("Save Advanced Settings")
        od_save.clicked.connect(
            lambda: (
                self.backend.settings_manager.set_setting("onedrive_client_id", self._od_client_id.text().strip()),
                self.backend.settings_manager.set_setting("onedrive_tenant", self._od_tenant.text().strip() or "common"),
                self._cloud_status.setText("Saved OneDrive settings."),
            )
        )
        od_adv_l.addWidget(od_save, 0, Qt.AlignLeft)
        od_layout.addWidget(od_adv)

        def _od_set_adv(vis: bool):
            od_adv.setVisible(bool(vis))
            od_adv_toggle.setArrowType(Qt.DownArrow if vis else Qt.RightArrow)

        od_adv_toggle.toggled.connect(_od_set_adv)

        # Auto-hide Advanced section if publisher creds are available (end-user experience).
        try:
            from core.cloud_sync.app_credentials import get_onedrive_client_id

            if str(get_onedrive_client_id(self.backend.settings_manager) or "").strip():
                od_adv_toggle.setVisible(False)
        except Exception:
            pass

        od_layout.addWidget(od_msg)

        # Remote browser (OneDrive)
        od_browser = QGroupBox("Remote Browser")
        ob = QVBoxLayout(od_browser)
        ob.setContentsMargins(10, 10, 10, 10)
        ob.setSpacing(8)

        od_addr = QHBoxLayout()
        od_addr.setSpacing(8)
        self._od_remote_path = QLineEdit(self._od_sync_remote.text() if hasattr(self, "_od_sync_remote") else "Fylorra Sync")
        self._od_remote_path.setPlaceholderText("Remote path under OneDrive root (e.g. Fylorra Sync/Pictures)")
        od_addr.addWidget(self._od_remote_path, 1)
        btn_up = _cs_btn("up", "Go to parent folder", text=None, fixed=38)
        btn_refresh = _cs_btn("refresh", "Refresh list", text=None, fixed=38)
        btn_use = _cs_btn("sync", "Use for Sync Base", text=None, primary=True, fixed=38)
        od_addr.addWidget(btn_up)
        od_addr.addWidget(btn_refresh)
        od_addr.addWidget(btn_use)
        ob.addLayout(od_addr)

        self._od_tree = QTreeWidget()
        self._od_tree.setObjectName("ResultsTree")
        self._od_tree.setColumnCount(3)
        self._od_tree.setHeaderLabels(["Name", "Type", "Size"])
        try:
            self._od_tree.header().setSectionResizeMode(0, QHeaderView.Stretch)
            self._od_tree.header().setSectionResizeMode(1, QHeaderView.ResizeToContents)
            self._od_tree.header().setSectionResizeMode(2, QHeaderView.ResizeToContents)
        except Exception:
            pass
        self._od_tree.setMinimumHeight(190)
        ob.addWidget(self._od_tree, 1)

        def _od_refresh():
            path = (self._od_remote_path.text() or "").strip().strip("/").strip("\\")
            _run_cloud_job(
                action="OneDrive: Browse",
                fn=lambda status, progress: (  # noqa: ARG005
                    _ensure_providers(),
                    self._onedrive.list_folder(path, limit=200),
                )[-1],
                on_ok=lambda items: _od_populate(items),
            )

        def _od_populate(items):
            self._od_tree.clear()
            for it in items or []:
                typ = "Folder" if it.is_folder else "File"
                size = "" if it.is_folder else (f"{(it.size or 0)/1024/1024:.2f} MB" if it.size is not None else "")
                node = QTreeWidgetItem([it.name, typ, size])
                node.setData(0, Qt.UserRole, {"id": it.id, "is_folder": bool(it.is_folder), "name": it.name})
                self._od_tree.addTopLevelItem(node)

        def _od_enter(item: QTreeWidgetItem, _col: int):
            try:
                meta = item.data(0, Qt.UserRole) or {}
                if not meta.get("is_folder"):
                    return
                cur = (self._od_remote_path.text() or "").strip().strip("/").strip("\\")
                name = str(meta.get("name") or "").strip()
                nxt = f"{cur}/{name}".strip("/") if cur else name
                self._od_remote_path.setText(nxt)
                _od_refresh()
            except Exception:
                pass

        def _od_up():
            cur = (self._od_remote_path.text() or "").strip().strip("/").strip("\\")
            if not cur:
                return
            parent = str(Path(cur).parent).replace("\\", "/")
            if parent == ".":
                parent = ""
            self._od_remote_path.setText(parent)
            _od_refresh()

        def _od_use_sync_base():
            base = (self._od_remote_path.text() or "").strip().strip("/").strip("\\")
            if hasattr(self, "_od_sync_remote"):
                self._od_sync_remote.setText(base or "Fylorra Sync")
            try:
                self._od_up_remote.setText(base)
            except Exception:
                pass
            self._cloud_status.setText(f"Remote base set to: {base or 'Fylorra Sync'}")

        btn_refresh.clicked.connect(_od_refresh)
        btn_up.clicked.connect(_od_up)
        btn_use.clicked.connect(_od_use_sync_base)
        self._od_tree.itemDoubleClicked.connect(_od_enter)

        od_layout.addWidget(od_browser)

        # Advanced operations (collapsed; most users use the Explorer tab).
        od_ops_toggle_row = QHBoxLayout()
        od_ops_toggle_row.setContentsMargins(0, 0, 0, 0)
        od_ops_toggle_row.setSpacing(8)
        od_ops_toggle = QToolButton()
        od_ops_toggle.setCheckable(True)
        od_ops_toggle.setChecked(False)
        od_ops_toggle.setArrowType(Qt.RightArrow)
        od_ops_toggle.setText("Advanced operations")
        od_ops_toggle.setToolButtonStyle(Qt.ToolButtonTextBesideIcon)
        od_ops_toggle.setStyleSheet(f"color:{cloud_muted};")
        self._cloud_muted_widgets.append(od_ops_toggle)
        od_ops_toggle.setToolTip("Show optional upload/sync controls.")
        od_ops_toggle_row.addWidget(od_ops_toggle, 0, Qt.AlignLeft)
        od_ops_toggle_row.addStretch(1)
        od_layout.addLayout(od_ops_toggle_row)

        od_ops_box = QWidget()
        od_ops_box.setVisible(False)
        od_ops_l = QVBoxLayout(od_ops_box)
        od_ops_l.setContentsMargins(0, 0, 0, 0)
        od_ops_l.setSpacing(10)

        od_upload = QGroupBox("Upload File")
        ou = QFormLayout(od_upload)
        ou.setHorizontalSpacing(12)
        ou.setVerticalSpacing(8)
        self._od_up_local = QLineEdit()
        self._od_up_local.setPlaceholderText("Pick a local file…")
        pick = _cs_btn("browse", "Pick a local file", text=None, fixed=38)
        pick.clicked.connect(lambda: self._od_up_local.setText(QFileDialog.getOpenFileName(self, "Select file")[0] or ""))
        row_local = QHBoxLayout()
        row_local.addWidget(self._od_up_local, 1)
        row_local.addWidget(pick)
        ou.addRow("Local file", row_local)
        self._od_up_remote = QLineEdit()
        self._od_up_remote.setPlaceholderText("Remote folder under root (optional), e.g. Fylorra")
        ou.addRow("Remote folder", self._od_up_remote)
        self._od_up_bar = QProgressBar()
        self._od_up_bar.setRange(0, 1000)
        self._od_up_bar.setValue(0)
        ou.addRow("Progress", self._od_up_bar)
        btn_up = QPushButton("Upload")
        btn_up.setObjectName("PrimaryButton")
        _cloud_set_btn_icon(btn_up, "upload", primary=True)
        btn_up.clicked.connect(
            lambda: _run_cloud_job(
                action="OneDrive: Upload File",
                fn=lambda status, progress: (  # noqa: ARG005
                    _ensure_providers(),
                    self._onedrive.upload_file(
                        Path(self._od_up_local.text().strip()),
                        remote_folder=self._od_up_remote.text().strip() or None,
                        progress_cb=lambda d, t, m: progress(int(d), int(t or 0), str(m or "")),
                    ),
                )[-1],
                on_ok=lambda it: QMessageBox.information(self, "OneDrive Upload", f"Uploaded: {getattr(it, 'name', '')}"),
                progress_sink=lambda done, total, _msg: self._od_up_bar.setValue(int(1000 * (float(done) / float(total))) if total else 0),
            )
        )
        ou.addRow("", btn_up)
        od_ops_l.addWidget(od_upload)

        # Folder sync (upload-only) for OneDrive
        od_sync = QGroupBox("Folder Sync (upload-only)")
        osy = QFormLayout(od_sync)
        osy.setHorizontalSpacing(12)
        osy.setVerticalSpacing(8)
        self._od_sync_local = QLineEdit()
        self._od_sync_local.setPlaceholderText("Pick a local folder to upload…")
        pick_folder = _cs_btn("browse", "Pick a local folder", text=None, fixed=38)
        pick_folder.clicked.connect(lambda: self._od_sync_local.setText(QFileDialog.getExistingDirectory(self, "Select folder") or ""))
        row_f = QHBoxLayout()
        row_f.addWidget(self._od_sync_local, 1)
        row_f.addWidget(pick_folder)
        osy.addRow("Local folder", row_f)
        self._od_sync_remote = QLineEdit("Fylorra Sync")
        self._od_sync_remote.setPlaceholderText("Remote base folder under OneDrive root")
        osy.addRow("Remote base", self._od_sync_remote)
        self._od_sync_sub = QCheckBox("Include subfolders")
        self._od_sync_sub.setChecked(True)
        osy.addRow("", self._od_sync_sub)
        self._od_sync_dry = QCheckBox("Dry run (no uploads)")
        self._od_sync_dry.setChecked(False)
        osy.addRow("", self._od_sync_dry)
        self._od_sync_bar = QProgressBar()
        self._od_sync_bar.setRange(0, 1000)
        self._od_sync_bar.setValue(0)
        osy.addRow("Progress", self._od_sync_bar)
        btn_sync = QPushButton("Sync to OneDrive")
        btn_sync.setObjectName("PrimaryButton")
        _cloud_set_btn_icon(btn_sync, "sync", primary=True)
        def _od_run_upload_only():
            local = self._od_sync_local.text().strip()
            if not local:
                QMessageBox.warning(self, "OneDrive Sync", "Pick a local folder first.")
                return
            _run_cloud_job(
                action="OneDrive: Sync Folder",
                fn=lambda status, progress: (  # noqa: ARG005
                    _ensure_providers(),
                    __import__("core.cloud_sync.sync_engine", fromlist=["sync_folder_upload_only"]).sync_folder_upload_only(  # noqa: E501
                        self._onedrive,
                        local_root=Path(local),
                        remote_base=self._od_sync_remote.text().strip() or "Fylorra Sync",
                        include_subfolders=bool(self._od_sync_sub.isChecked()),
                        dry_run=bool(self._od_sync_dry.isChecked()),
                        status_cb=status,
                        progress_cb=progress,
                    ),
                )[-1],
                on_ok=lambda st: QMessageBox.information(
                    self,
                    "OneDrive Sync",
                    f"Done.\nScanned: {st.scanned}\nUploaded: {st.uploaded}\nSkipped: {st.skipped}",
                ),
                progress_sink=lambda done, total, msg: self._od_sync_bar.setValue(
                    int(1000 * (float(done) / float(total))) if total else 0
                ),
            )

        btn_sync.clicked.connect(_od_run_upload_only)
        btn_schedule = QPushButton("Schedule…")
        _cloud_set_btn_icon(btn_schedule, "sync")
        btn_schedule.clicked.connect(
            lambda: _schedule_cloud_sync_upload_only(
                provider="onedrive",
                local_folder=self._od_sync_local.text().strip(),
                remote_base=self._od_sync_remote.text().strip() or "Fylorra Sync",
                include_subfolders=bool(self._od_sync_sub.isChecked()),
                dry_run=bool(self._od_sync_dry.isChecked()),
            )
        )

        btn_two_way = QPushButton("Two-way Sync")
        btn_two_way.setObjectName("PrimaryButton")
        _cloud_set_btn_icon(btn_two_way, "sync", primary=True)
        def _od_run_two_way():
            local = self._od_sync_local.text().strip()
            if not local:
                QMessageBox.warning(self, "OneDrive Two-way Sync", "Pick a local folder first.")
                return
            _run_cloud_job(
                action="OneDrive: Two-way Sync",
                fn=lambda status, progress: (  # noqa: ARG005
                    _ensure_providers(),
                    __import__("core.cloud_sync.sync_engine", fromlist=["sync_folder_two_way"]).sync_folder_two_way(  # noqa: E501
                        self._onedrive,
                        local_root=Path(local),
                        remote_base=self._od_sync_remote.text().strip() or "Fylorra Sync",
                        include_subfolders=bool(self._od_sync_sub.isChecked()),
                        dry_run=bool(self._od_sync_dry.isChecked()),
                        delete_policy="ignore",
                        conflict_policy="keep_both",
                        status_cb=status,
                        progress_cb=progress,
                    ),
                )[-1],
                on_ok=lambda st: QMessageBox.information(
                    self,
                    "OneDrive Two-way Sync",
                    f"Done.\nUploaded: {st.uploaded}\nDownloaded: {st.downloaded}\nConflicts: {st.conflicts}\nSkipped: {st.skipped}",
                ),
                progress_sink=lambda done, total, msg: self._od_sync_bar.setValue(
                    int(1000 * (float(done) / float(total))) if total else 0
                ),
            )

        btn_two_way.clicked.connect(_od_run_two_way)

        btn_two_way_schedule = QPushButton("Schedule…")
        _cloud_set_btn_icon(btn_two_way_schedule, "sync")
        btn_two_way_schedule.clicked.connect(
            lambda: _schedule_cloud_sync_full(
                action_type="cloud_sync_two_way",
                provider="onedrive",
                local_folder=self._od_sync_local.text().strip(),
                remote_base=self._od_sync_remote.text().strip() or "Fylorra Sync",
                include_subfolders=bool(self._od_sync_sub.isChecked()),
                dry_run=bool(self._od_sync_dry.isChecked()),
            )
        )
        sync_row = QHBoxLayout()
        sync_row.setSpacing(10)
        sync_row.addWidget(btn_sync, 1)
        sync_row.addWidget(btn_schedule)
        sync_host = QWidget()
        sync_host.setLayout(sync_row)
        osy.addRow("", sync_host)

        tw_row = QHBoxLayout()
        tw_row.setSpacing(10)
        tw_row.addWidget(btn_two_way, 1)
        tw_row.addWidget(btn_two_way_schedule)
        tw_host = QWidget()
        tw_host.setLayout(tw_row)
        osy.addRow("", tw_host)
        od_ops_l.addWidget(od_sync)

        def _od_sync_update_enabled():
            ok = bool(self._od_sync_local.text().strip())
            try:
                btn_sync.setEnabled(ok)
                btn_schedule.setEnabled(ok)
                btn_two_way.setEnabled(ok)
                btn_two_way_schedule.setEnabled(ok)
            except Exception:
                pass

        self._od_sync_local.textChanged.connect(lambda _t: _od_sync_update_enabled())
        _od_sync_update_enabled()

        od_layout.addWidget(od_ops_box)

        def _od_ops_set(vis: bool):
            od_ops_box.setVisible(bool(vis))
            od_ops_toggle.setArrowType(Qt.DownArrow if vis else Qt.RightArrow)

        od_ops_toggle.toggled.connect(_od_ops_set)

        od_scroll = QScrollArea()
        od_scroll.setWidgetResizable(True)
        od_scroll.setFrameShape(QFrame.NoFrame)
        od_scroll.setWidget(od_tab)
        tabs.addTab(od_scroll, "OneDrive")

        # Google Drive tab
        gd_tab = QWidget()
        gd_layout = QVBoxLayout(gd_tab)
        gd_layout.setContentsMargins(10, 10, 10, 10)
        gd_layout.setSpacing(10)

        gd_hint = QHBoxLayout()
        gd_hint.setSpacing(10)
        gd_hint_lbl = QLabel("Uses Google secure sign-in (OAuth). Fylorra never sees your password.")
        gd_hint_lbl.setStyleSheet(f"color:{cloud_muted};")
        self._cloud_muted_widgets.append(gd_hint_lbl)
        gd_hint.addWidget(gd_hint_lbl, 1)
        # gd_help = _cs_btn("browse", "Help / setup instructions", text=None, fixed=38)
        # gd_help.clicked.connect(_open_cloud_sync_help)
        # gd_hint.addWidget(gd_help)
        # gd_why = _cs_btn("search", "Why OAuth is required", text="?", fixed=38)
        # gd_why.clicked.connect(_why_no_passwords)
        # gd_hint.addWidget(gd_why)
        gd_layout.addLayout(gd_hint)

        gd_header = QHBoxLayout()
        gd_header.setSpacing(10)
        self._gd_status = QLabel("Google Drive: Not connected")
        self._gd_status.setStyleSheet(f"color:{cloud_muted};")
        self._cloud_muted_widgets.append(self._gd_status)
        gd_header.addWidget(self._gd_status, 1)
        gbtn = _cs_btn("connect", "Connect Google Drive", text=None, primary=True, fixed=38)

        def _gd_connect():
            try:
                self._cloud_status.setText("Starting Google Drive sign-in…")
            except Exception:
                pass
            try:
                from core.cloud_sync.app_credentials import get_gdrive_client_secrets_path

                sp = str(get_gdrive_client_secrets_path(self.backend.settings_manager) or "").strip()
            except Exception:
                sp = str(self.backend.settings_manager.get_setting("gdrive_client_secrets_path", "") or "").strip()
            if not sp:
                QMessageBox.information(
                    self,
                    "Google Drive setup required",
                    "To connect Google Drive, Fylorra needs an OAuth client secrets JSON.\n\n"
                    "If you are using the official Fylorra build, this should already be configured.\n"
                    "Otherwise, click Help and select the JSON in Advanced settings.",
                )
                _open_cloud_sync_help()
                return
            _run_cloud_job(
                action="Google Drive: Connect",
                fn=lambda status, progress: (  # noqa: ARG005
                    _ensure_providers(),
                    self._gdrive.connect(),
                )[-1],
                on_ok=lambda _res: (_refresh_status(), self._cloud_status.setText("Google Drive connected.")),
            )

        gbtn.clicked.connect(_gd_connect)
        gd_header.addWidget(gbtn)

        gbtn_disc = _cs_btn("disconnect", "Disconnect Google Drive", text=None, primary=False, fixed=38)
        gbtn_test = _cs_btn("test", "Test Google Drive access", text=None, primary=False, fixed=38)
        gd_header.addWidget(gbtn_disc)
        gd_header.addWidget(gbtn_test)
        gdisc = QPushButton("Disconnect")
        _cloud_set_btn_icon(gdisc, "disconnect")
        gdisc.clicked.connect(
            lambda: _run_cloud_job(
                action="Google Drive: Disconnect",
                fn=lambda status, progress: (  # noqa: ARG005
                    _ensure_providers(),
                    self._gdrive.disconnect(),
                    "Disconnected",
                )[-1],
                on_ok=lambda _res: (_refresh_status(), self._cloud_status.setText("Google Drive disconnected.")),
            )
        )
        gd_header.addWidget(gdisc)
        gtest = QPushButton("Test")
        _cloud_set_btn_icon(gtest, "test")
        gtest.clicked.connect(
            lambda: _run_cloud_job(
                action="Google Drive: Test",
                fn=lambda status, progress: (  # noqa: ARG005
                    _ensure_providers(),
                    self._gdrive.test_connection(),
                )[-1],
                on_ok=lambda res: QMessageBox.information(self, "Google Drive", str(res)),
            )
        )
        gd_header.addWidget(gtest)
        gd_layout.addLayout(gd_header)

        def _gd_disconnect():
            _run_cloud_job(
                action="Google Drive: Disconnect",
                fn=lambda status, progress: (  # noqa: ARG005
                    _ensure_providers(),
                    self._gdrive.disconnect(),
                )[-1],
                on_ok=lambda _r: (_refresh_status(), self._cloud_status.setText("Google Drive disconnected.")),
            )

        def _gd_test():
            _run_cloud_job(
                action="Google Drive: Test",
                fn=lambda status, progress: (  # noqa: ARG005
                    _ensure_providers(),
                    self._gdrive.test_access(),
                )[-1],
                on_ok=lambda _r: QMessageBox.information(self, "Google Drive", "Access OK."),
                on_err=lambda m: QMessageBox.critical(self, "Google Drive", m or "Access failed."),
            )

        gbtn_disc.clicked.connect(_gd_disconnect)
        gbtn_test.clicked.connect(_gd_test)

        gd_adv_toggle = QToolButton()
        gd_adv_toggle.setText("Advanced")
        gd_adv_toggle.setCheckable(True)
        gd_adv_toggle.setChecked(False)
        gd_adv_toggle.setToolButtonStyle(Qt.ToolButtonTextBesideIcon)
        gd_adv_toggle.setArrowType(Qt.RightArrow)
        gd_layout.addWidget(gd_adv_toggle, 0, Qt.AlignLeft)

        gd_adv = QFrame()
        gd_adv.setVisible(False)
        gd_adv_l = QVBoxLayout(gd_adv)
        gd_adv_l.setContentsMargins(0, 0, 0, 0)
        gd_adv_l.setSpacing(8)

        gd_adv_note = QLabel("Admin setup: select your Google OAuth client secrets JSON.")
        gd_adv_note.setStyleSheet(f"color:{cloud_muted};")
        self._cloud_muted_widgets.append(gd_adv_note)
        gd_adv_l.addWidget(gd_adv_note)

        gd_form = QFormLayout()
        gd_form.setHorizontalSpacing(12)
        gd_form.setVerticalSpacing(8)
        self._gd_secrets = QLineEdit(str(self.backend.settings_manager.get_setting("gdrive_client_secrets_path", "") or ""))
        self._gd_secrets.setPlaceholderText("OAuth client secrets JSON file (download from Google Cloud Console)")
        pick_sec = QPushButton("Browse")
        pick_sec.clicked.connect(lambda: self._gd_secrets.setText(QFileDialog.getOpenFileName(self, "Select Google OAuth JSON", filter="*.json")[0] or ""))
        sec_row = QHBoxLayout()
        sec_row.addWidget(self._gd_secrets, 1)
        sec_row.addWidget(pick_sec)
        gd_form.addRow("Client secrets", sec_row)
        gd_adv_l.addLayout(gd_form)

        gd_save = QPushButton("Save Advanced Settings")
        gd_save.clicked.connect(
            lambda: (
                self.backend.settings_manager.set_setting("gdrive_client_secrets_path", self._gd_secrets.text().strip()),
                self._cloud_status.setText("Saved Google Drive settings."),
            )
        )
        gd_adv_l.addWidget(gd_save, 0, Qt.AlignLeft)
        gd_layout.addWidget(gd_adv)

        def _gd_set_adv(vis: bool):
            gd_adv.setVisible(bool(vis))
            gd_adv_toggle.setArrowType(Qt.DownArrow if vis else Qt.RightArrow)

        gd_adv_toggle.toggled.connect(_gd_set_adv)

        # Auto-hide Advanced section if publisher creds are available (end-user experience).
        try:
            from core.cloud_sync.app_credentials import get_gdrive_client_secrets_path

            if str(get_gdrive_client_secrets_path(self.backend.settings_manager) or "").strip():
                gd_adv_toggle.setVisible(False)
        except Exception:
            pass

        # Remote browser (Google Drive)
        gd_browser = QGroupBox("Remote Browser")
        gb = QVBoxLayout(gd_browser)
        gb.setContentsMargins(10, 10, 10, 10)
        gb.setSpacing(8)

        self._gd_folder_id = "root"
        self._gd_path_parts: list[str] = []

        gd_addr = QHBoxLayout()
        gd_addr.setSpacing(8)
        self._gd_remote_path = QLineEdit("Fylorra Sync")
        self._gd_remote_path.setPlaceholderText("Remote folder path (under Drive root), e.g. Fylorra Sync/Pictures")
        gd_addr.addWidget(self._gd_remote_path, 1)
        gbtn_up = _cs_btn("up", "Go to parent folder", text=None, fixed=38)
        gbtn_refresh = _cs_btn("refresh", "Refresh list", text=None, fixed=38)
        gbtn_use = _cs_btn("sync", "Use for Sync Base", text=None, primary=True, fixed=38)
        gd_addr.addWidget(gbtn_up)
        gd_addr.addWidget(gbtn_refresh)
        gd_addr.addWidget(gbtn_use)
        gb.addLayout(gd_addr)

        self._gd_tree = QTreeWidget()
        self._gd_tree.setObjectName("ResultsTree")
        self._gd_tree.setColumnCount(3)
        self._gd_tree.setHeaderLabels(["Name", "Type", "Size"])
        try:
            self._gd_tree.header().setSectionResizeMode(0, QHeaderView.Stretch)
            self._gd_tree.header().setSectionResizeMode(1, QHeaderView.ResizeToContents)
            self._gd_tree.header().setSectionResizeMode(2, QHeaderView.ResizeToContents)
        except Exception:
            pass
        self._gd_tree.setMinimumHeight(190)
        gb.addWidget(self._gd_tree, 1)

        def _gd_current_path_str() -> str:
            return "/".join([p for p in (self._gd_path_parts or []) if p])

        def _gd_populate(items):
            self._gd_tree.clear()
            for it in items or []:
                typ = "Folder" if it.is_folder else "File"
                size = "" if it.is_folder else (f"{(it.size or 0)/1024/1024:.2f} MB" if it.size is not None else "")
                node = QTreeWidgetItem([it.name, typ, size])
                node.setData(0, Qt.UserRole, {"id": it.id, "is_folder": bool(it.is_folder), "name": it.name})
                self._gd_tree.addTopLevelItem(node)

        def _gd_refresh():
            # If user typed a path, resolve/create it and browse into it.
            typed = (self._gd_remote_path.text() or "").strip().strip("/").strip("\\")
            def _do(status, progress):  # noqa: ARG001
                _ensure_providers()
                if typed:
                    folder_id = self._gdrive.ensure_folder_path(typed)
                    self._gd_folder_id = folder_id
                    self._gd_path_parts = [p for p in typed.replace("\\", "/").split("/") if p]
                    return self._gdrive.list_folder(folder_id, limit=200)
                self._gd_folder_id = "root"
                self._gd_path_parts = []
                return self._gdrive.list_folder("root", limit=200)
            _run_cloud_job(action="Google Drive: Browse", fn=_do, on_ok=_gd_populate)

        def _gd_enter(item: QTreeWidgetItem, _col: int):
            try:
                meta = item.data(0, Qt.UserRole) or {}
                if not meta.get("is_folder"):
                    return
                fid = str(meta.get("id") or "").strip()
                name = str(meta.get("name") or "").strip()
                if not fid or not name:
                    return
                self._gd_folder_id = fid
                self._gd_path_parts.append(name)
                self._gd_remote_path.setText(_gd_current_path_str())
                _gd_refresh()
            except Exception:
                pass

        def _gd_up():
            if not self._gd_path_parts:
                self._gd_remote_path.setText("")
                _gd_refresh()
                return
            try:
                self._gd_path_parts.pop()
            except Exception:
                self._gd_path_parts = []
            self._gd_remote_path.setText(_gd_current_path_str())
            _gd_refresh()

        def _gd_use_sync_base():
            base = (self._gd_remote_path.text() or "").strip().strip("/").strip("\\")
            if hasattr(self, "_gd_sync_remote"):
                self._gd_sync_remote.setText(base or "Fylorra Sync")
            try:
                # For uploads we can use the current folder ID directly
                self._gd_up_folder.setText(str(self._gd_folder_id or "root"))
            except Exception:
                pass
            self._cloud_status.setText(f"Remote base set to: {base or 'Fylorra Sync'}")

        gbtn_refresh.clicked.connect(_gd_refresh)
        gbtn_up.clicked.connect(_gd_up)
        gbtn_use.clicked.connect(_gd_use_sync_base)
        self._gd_tree.itemDoubleClicked.connect(_gd_enter)

        gd_layout.addWidget(gd_browser)

        # Advanced operations (collapsed; most users use the Explorer tab).
        gd_ops_toggle_row = QHBoxLayout()
        gd_ops_toggle_row.setContentsMargins(0, 0, 0, 0)
        gd_ops_toggle_row.setSpacing(8)
        gd_ops_toggle = QToolButton()
        gd_ops_toggle.setCheckable(True)
        gd_ops_toggle.setChecked(False)
        gd_ops_toggle.setArrowType(Qt.RightArrow)
        gd_ops_toggle.setText("Advanced operations")
        gd_ops_toggle.setToolButtonStyle(Qt.ToolButtonTextBesideIcon)
        gd_ops_toggle.setStyleSheet(f"color:{cloud_muted};")
        self._cloud_muted_widgets.append(gd_ops_toggle)
        gd_ops_toggle.setToolTip("Show optional upload/sync controls.")
        gd_ops_toggle_row.addWidget(gd_ops_toggle, 0, Qt.AlignLeft)
        gd_ops_toggle_row.addStretch(1)
        gd_layout.addLayout(gd_ops_toggle_row)

        gd_ops_box = QWidget()
        gd_ops_box.setVisible(False)
        gd_ops_l = QVBoxLayout(gd_ops_box)
        gd_ops_l.setContentsMargins(0, 0, 0, 0)
        gd_ops_l.setSpacing(10)

        gd_upload = QGroupBox("Upload File")
        gu = QFormLayout(gd_upload)
        gu.setHorizontalSpacing(12)
        gu.setVerticalSpacing(8)
        self._gd_up_local = QLineEdit()
        self._gd_up_local.setPlaceholderText("Pick a local file…")
        gpick = _cs_btn("browse", "Pick a local file", text=None, fixed=38)
        gpick.clicked.connect(lambda: self._gd_up_local.setText(QFileDialog.getOpenFileName(self, "Select file")[0] or ""))
        row_local_g = QHBoxLayout()
        row_local_g.addWidget(self._gd_up_local, 1)
        row_local_g.addWidget(gpick)
        gu.addRow("Local file", row_local_g)
        self._gd_up_folder = QLineEdit()
        self._gd_up_folder.setPlaceholderText("Remote folder ID (optional)")
        gu.addRow("Remote folder ID", self._gd_up_folder)
        self._gd_up_bar = QProgressBar()
        self._gd_up_bar.setRange(0, 1000)
        self._gd_up_bar.setValue(0)
        gu.addRow("Progress", self._gd_up_bar)
        gup = QPushButton("Upload")
        gup.setObjectName("PrimaryButton")
        _cloud_set_btn_icon(gup, "upload", primary=True)
        gup.clicked.connect(
            lambda: _run_cloud_job(
                action="Google Drive: Upload File",
                fn=lambda status, progress: (  # noqa: ARG005
                    _ensure_providers(),
                    self._gdrive.upload_file(
                        Path(self._gd_up_local.text().strip()),
                        remote_folder=self._gd_up_folder.text().strip() or None,
                        progress_cb=lambda d, t, m: progress(int(d), int(t or 0), str(m or "")),
                    ),
                )[-1],
                on_ok=lambda it: QMessageBox.information(self, "Google Drive Upload", f"Uploaded: {getattr(it, 'name', '')}"),
                progress_sink=lambda done, total, _msg: self._gd_up_bar.setValue(int(1000 * (float(done) / float(total))) if total else 0),
            )
        )
        gu.addRow("", gup)
        gd_ops_l.addWidget(gd_upload)

        # Folder sync (upload-only) for Google Drive
        gd_sync = QGroupBox("Folder Sync (upload-only)")
        gsy = QFormLayout(gd_sync)
        gsy.setHorizontalSpacing(12)
        gsy.setVerticalSpacing(8)
        self._gd_sync_local = QLineEdit()
        self._gd_sync_local.setPlaceholderText("Pick a local folder to upload…")
        pick_folder_g = _cs_btn("browse", "Pick a local folder", text=None, fixed=38)
        pick_folder_g.clicked.connect(lambda: self._gd_sync_local.setText(QFileDialog.getExistingDirectory(self, "Select folder") or ""))
        row_fg = QHBoxLayout()
        row_fg.addWidget(self._gd_sync_local, 1)
        row_fg.addWidget(pick_folder_g)
        gsy.addRow("Local folder", row_fg)
        self._gd_sync_remote = QLineEdit("Fylorra Sync")
        self._gd_sync_remote.setPlaceholderText("Remote base folder under Drive root")
        gsy.addRow("Remote base", self._gd_sync_remote)
        self._gd_sync_sub = QCheckBox("Include subfolders")
        self._gd_sync_sub.setChecked(True)
        gsy.addRow("", self._gd_sync_sub)
        self._gd_sync_dry = QCheckBox("Dry run (no uploads)")
        self._gd_sync_dry.setChecked(False)
        gsy.addRow("", self._gd_sync_dry)
        self._gd_sync_bar = QProgressBar()
        self._gd_sync_bar.setRange(0, 1000)
        self._gd_sync_bar.setValue(0)
        gsy.addRow("Progress", self._gd_sync_bar)
        btn_sync_g = QPushButton("Sync to Google Drive")
        btn_sync_g.setObjectName("PrimaryButton")
        _cloud_set_btn_icon(btn_sync_g, "sync", primary=True)
        def _gd_run_upload_only():
            local = self._gd_sync_local.text().strip()
            if not local:
                QMessageBox.warning(self, "Google Drive Sync", "Pick a local folder first.")
                return
            _run_cloud_job(
                action="Google Drive: Sync Folder",
                fn=lambda status, progress: (  # noqa: ARG005
                    _ensure_providers(),
                    __import__("core.cloud_sync.sync_engine", fromlist=["sync_folder_upload_only"]).sync_folder_upload_only(  # noqa: E501
                        self._gdrive,
                        local_root=Path(local),
                        remote_base=self._gd_sync_remote.text().strip() or "Fylorra Sync",
                        include_subfolders=bool(self._gd_sync_sub.isChecked()),
                        dry_run=bool(self._gd_sync_dry.isChecked()),
                        status_cb=status,
                        progress_cb=progress,
                    ),
                )[-1],
                on_ok=lambda st: QMessageBox.information(
                    self,
                    "Google Drive Sync",
                    f"Done.\nScanned: {st.scanned}\nUploaded: {st.uploaded}\nSkipped: {st.skipped}",
                ),
                progress_sink=lambda done, total, msg: self._gd_sync_bar.setValue(
                    int(1000 * (float(done) / float(total))) if total else 0
                ),
            )

        btn_sync_g.clicked.connect(_gd_run_upload_only)
        btn_schedule_g = QPushButton("Schedule…")
        _cloud_set_btn_icon(btn_schedule_g, "sync")
        btn_schedule_g.clicked.connect(
            lambda: _schedule_cloud_sync_upload_only(
                provider="gdrive",
                local_folder=self._gd_sync_local.text().strip(),
                remote_base=self._gd_sync_remote.text().strip() or "Fylorra Sync",
                include_subfolders=bool(self._gd_sync_sub.isChecked()),
                dry_run=bool(self._gd_sync_dry.isChecked()),
            )
        )

        btn_two_way_g = QPushButton("Two-way Sync")
        btn_two_way_g.setObjectName("PrimaryButton")
        _cloud_set_btn_icon(btn_two_way_g, "sync", primary=True)
        def _gd_run_two_way():
            local = self._gd_sync_local.text().strip()
            if not local:
                QMessageBox.warning(self, "Google Drive Two-way Sync", "Pick a local folder first.")
                return
            _run_cloud_job(
                action="Google Drive: Two-way Sync",
                fn=lambda status, progress: (  # noqa: ARG005
                    _ensure_providers(),
                    __import__("core.cloud_sync.sync_engine", fromlist=["sync_folder_two_way"]).sync_folder_two_way(  # noqa: E501
                        self._gdrive,
                        local_root=Path(local),
                        remote_base=self._gd_sync_remote.text().strip() or "Fylorra Sync",
                        include_subfolders=bool(self._gd_sync_sub.isChecked()),
                        dry_run=bool(self._gd_sync_dry.isChecked()),
                        delete_policy="ignore",
                        conflict_policy="keep_both",
                        status_cb=status,
                        progress_cb=progress,
                    ),
                )[-1],
                on_ok=lambda st: QMessageBox.information(
                    self,
                    "Google Drive Two-way Sync",
                    f"Done.\nUploaded: {st.uploaded}\nDownloaded: {st.downloaded}\nConflicts: {st.conflicts}\nSkipped: {st.skipped}",
                ),
                progress_sink=lambda done, total, msg: self._gd_sync_bar.setValue(
                    int(1000 * (float(done) / float(total))) if total else 0
                ),
            )

        btn_two_way_g.clicked.connect(_gd_run_two_way)

        btn_two_way_schedule_g = QPushButton("Schedule…")
        _cloud_set_btn_icon(btn_two_way_schedule_g, "sync")
        btn_two_way_schedule_g.clicked.connect(
            lambda: _schedule_cloud_sync_full(
                action_type="cloud_sync_two_way",
                provider="gdrive",
                local_folder=self._gd_sync_local.text().strip(),
                remote_base=self._gd_sync_remote.text().strip() or "Fylorra Sync",
                include_subfolders=bool(self._gd_sync_sub.isChecked()),
                dry_run=bool(self._gd_sync_dry.isChecked()),
            )
        )
        sync_row_g = QHBoxLayout()
        sync_row_g.setSpacing(10)
        sync_row_g.addWidget(btn_sync_g, 1)
        sync_row_g.addWidget(btn_schedule_g)
        sync_host_g = QWidget()
        sync_host_g.setLayout(sync_row_g)
        gsy.addRow("", sync_host_g)

        tw_row_g = QHBoxLayout()
        tw_row_g.setSpacing(10)
        tw_row_g.addWidget(btn_two_way_g, 1)
        tw_row_g.addWidget(btn_two_way_schedule_g)
        tw_host_g = QWidget()
        tw_host_g.setLayout(tw_row_g)
        gsy.addRow("", tw_host_g)
        gd_ops_l.addWidget(gd_sync)

        def _gd_sync_update_enabled():
            ok = bool(self._gd_sync_local.text().strip())
            try:
                btn_sync_g.setEnabled(ok)
                btn_schedule_g.setEnabled(ok)
                btn_two_way_g.setEnabled(ok)
                btn_two_way_schedule_g.setEnabled(ok)
            except Exception:
                pass

        self._gd_sync_local.textChanged.connect(lambda _t: _gd_sync_update_enabled())
        _gd_sync_update_enabled()

        gd_layout.addWidget(gd_ops_box)

        def _gd_ops_set(vis: bool):
            gd_ops_box.setVisible(bool(vis))
            gd_ops_toggle.setArrowType(Qt.DownArrow if vis else Qt.RightArrow)

        gd_ops_toggle.toggled.connect(_gd_ops_set)

        gd_scroll = QScrollArea()
        gd_scroll.setWidgetResizable(True)
        gd_scroll.setFrameShape(QFrame.NoFrame)
        gd_scroll.setWidget(gd_tab)
        tabs.addTab(gd_scroll, "Google Drive")

        # Explorer (two-pane: local + remote) for a more professional UX (SyncThing/AirExplorer style).
        exp_tab = QWidget()
        exp_layout = QVBoxLayout(exp_tab)
        exp_layout.setContentsMargins(6, 6, 6, 6)
        exp_layout.setSpacing(8)

        exp_top = QHBoxLayout()
        exp_top.setSpacing(8)
        exp_provider = QComboBox()
        exp_provider.addItems(["OneDrive", "Google Drive"])
        try:
            exp_provider.setItemIcon(0, _cs_icon("onedrive"))
            exp_provider.setItemIcon(1, _cs_icon("google_drive"))
        except Exception:
            pass
        exp_provider.setMinimumWidth(200)
        exp_provider.setMaximumWidth(240)
        exp_top.addWidget(exp_provider)
        btn_switch = _cs_btn("refresh", "Switch provider", text="Switch")
        exp_top.addWidget(btn_switch)
        exp_top.addStretch(1)
        # btn_help = _cs_btn("browse", "Help / setup instructions", text="?", fixed=38)
        # btn_help.clicked.connect(_open_cloud_sync_help)
        # exp_top.addWidget(btn_help)
        exp_layout.addLayout(exp_top)

        # Cloud-app style layout: tabs for Local/Cloud plus a compact action bar.
        exp_mode_tabs = QTabWidget()
        exp_mode_tabs.setDocumentMode(True)
        exp_layout.addWidget(exp_mode_tabs, 1)

        # Local tab (fast, uses QFileSystemModel).
        local_tab = QWidget()
        lt = QVBoxLayout(local_tab)
        lt.setContentsMargins(0, 0, 0, 0)
        lt.setSpacing(8)

        local_row = QHBoxLayout()
        local_row.setSpacing(8)
        local_root = QLineEdit(str(Path.home()))
        local_root.setPlaceholderText("Local folder…")
        local_row.addWidget(local_root, 1)
        local_browse = _cs_btn("browse", "Browse local folder", text=None, fixed=38)
        local_browse.clicked.connect(lambda: local_root.setText(QFileDialog.getExistingDirectory(self, "Select local folder") or local_root.text()))
        local_row.addWidget(local_browse)
        local_search = QLineEdit()
        local_search.setPlaceholderText("Search…")
        local_search.setMaximumWidth(280)
        local_row.addWidget(local_search)
        btn_local_search = _cs_btn("search", "Search in this folder", text=None, fixed=38)
        local_row.addWidget(btn_local_search)
        lt.addLayout(local_row)

        local_model = QFileSystemModel()
        local_model.setRootPath(local_root.text())
        local_view = QTreeView()
        local_view.setModel(local_model)
        local_view.setRootIndex(local_model.index(local_root.text()))
        local_view.setSortingEnabled(True)
        local_view.setSelectionMode(QAbstractItemView.ExtendedSelection)
        local_view.setSelectionBehavior(QAbstractItemView.SelectRows)
        local_view.setAlternatingRowColors(True)
        try:
            local_view.header().setSectionResizeMode(0, QHeaderView.Stretch)
            local_view.header().setSectionResizeMode(1, QHeaderView.ResizeToContents)
            local_view.header().setSectionResizeMode(2, QHeaderView.ResizeToContents)
            local_view.header().setSectionResizeMode(3, QHeaderView.ResizeToContents)
        except Exception:
            pass
        lt.addWidget(local_view, 1)

        def _set_local_root():
            p = local_root.text().strip()
            if not p:
                return
            local_model.setRootPath(p)
            local_view.setRootIndex(local_model.index(p))

        local_root.editingFinished.connect(_set_local_root)

        def _local_search_apply():
            q = (local_search.text() or "").strip().lower()
            if not q:
                return
            root_index = local_view.rootIndex()
            for r in range(local_model.rowCount(root_index)):
                idx = local_model.index(r, 0, root_index)
                name = str(local_model.data(idx) or "").lower()
                if q in name:
                    local_view.setCurrentIndex(idx)
                    local_view.scrollTo(idx)
                    break

        btn_local_search.clicked.connect(_local_search_apply)
        local_search.returnPressed.connect(_local_search_apply)

        exp_mode_tabs.addTab(local_tab, "Local Files")
        try:
            exp_mode_tabs.setTabIcon(0, _cs_icon("local_files"))
        except Exception:
            pass

        # Cloud tab
        cloud_tab = QWidget()
        ct = QVBoxLayout(cloud_tab)
        ct.setContentsMargins(0, 0, 0, 0)
        ct.setSpacing(8)

        remote_row = QHBoxLayout()
        remote_row.setSpacing(8)
        remote_path = QLineEdit("Fylorra Sync")
        remote_path.setPlaceholderText("Cloud folder…")
        remote_row.addWidget(remote_path, 1)
        btn_remote_up = _cs_btn("up", "Go to parent folder", text=None, fixed=38)
        btn_remote_refresh = _cs_btn("refresh", "Refresh list", text=None, fixed=38)
        remote_row.addWidget(btn_remote_up)
        remote_row.addWidget(btn_remote_refresh)
        remote_search = QLineEdit()
        remote_search.setPlaceholderText("Search…")
        remote_search.setMaximumWidth(280)
        remote_row.addWidget(remote_search)
        btn_remote_search = _cs_btn("search", "Filter current list", text=None, fixed=38)
        remote_row.addWidget(btn_remote_search)
        ct.addLayout(remote_row)

        remote_tree = QTreeWidget()
        remote_tree.setObjectName("ResultsTree")
        remote_tree.setColumnCount(3)
        remote_tree.setHeaderLabels(["Name", "Type", "Size"])
        try:
            remote_tree.header().setSectionResizeMode(0, QHeaderView.Stretch)
            remote_tree.header().setSectionResizeMode(1, QHeaderView.ResizeToContents)
            remote_tree.header().setSectionResizeMode(2, QHeaderView.ResizeToContents)
        except Exception:
            pass
        ct.addWidget(remote_tree, 1)

        def _remote_filter_apply():
            q = (remote_search.text() or "").strip().lower()
            for i in range(remote_tree.topLevelItemCount()):
                it = remote_tree.topLevelItem(i)
                name = (it.text(0) or "").lower()
                it.setHidden(bool(q) and (q not in name))

        btn_remote_search.clicked.connect(_remote_filter_apply)
        remote_search.returnPressed.connect(_remote_filter_apply)

        exp_mode_tabs.addTab(cloud_tab, "Cloud Storage")
        try:
            exp_mode_tabs.setTabIcon(1, _cs_icon("upload-to-cloud"))
        except Exception:
            pass

        # Bottom action bar (compact).
        exp_action_bar = QHBoxLayout()
        exp_action_bar.setSpacing(8)
        exp_action_bar.addStretch(1)
        exp_upload = _cs_btn("upload", "Upload to current cloud folder", text=None, primary=True, fixed=42)
        exp_download = _cs_btn("download", "Download selected cloud items", text=None, primary=False, fixed=42)
        btn_sync = _cs_btn("sync", "Sync local folder to cloud folder", text=None, primary=False, fixed=42)
        btn_new_folder = _cs_btn("new-folder", "Create new folder (cloud)", text=None, primary=False, fixed=42)

        sep = QFrame()
        sep.setFrameShape(QFrame.VLine)
        sep.setFrameShadow(QFrame.Sunken)
        sep.setStyleSheet("color:#2b2f36;")
        sep.setFixedHeight(26)

        btn_rename = _cs_btn("rename", "Rename selected item", text=None, primary=False, fixed=42)
        btn_delete = _cs_btn("delete", "Delete selected item(s)", text=None, primary=False, fixed=42)
        btn_copy = _cs_btn("copy", "Copy selected item(s)", text=None, primary=False, fixed=42)
        btn_cut = _cs_btn("cut", "Cut (move) selected item(s)", text=None, primary=False, fixed=42)
        btn_paste = _cs_btn("paste", "Paste into current folder", text=None, primary=False, fixed=42)
        btn_send = _cs_btn("upload-to-cloud", "Send selected file to the other cloud", text=None, primary=False, fixed=42)

        exp_action_bar.addWidget(exp_upload)
        exp_action_bar.addWidget(exp_download)
        exp_action_bar.addWidget(btn_sync)
        exp_action_bar.addWidget(btn_new_folder)
        exp_action_bar.addWidget(sep)
        exp_action_bar.addWidget(btn_rename)
        exp_action_bar.addWidget(btn_delete)
        exp_action_bar.addWidget(btn_copy)
        exp_action_bar.addWidget(btn_cut)
        exp_action_bar.addWidget(btn_paste)
        exp_action_bar.addWidget(btn_send)
        exp_layout.addLayout(exp_action_bar)

        # Remote navigation state for Google Drive.
        self._exp_gd_path_parts = []
        self._exp_gd_folder_id = "root"

        def _exp_provider_key() -> str:
            return "onedrive" if exp_provider.currentIndex() == 0 else "gdrive"

        def _exp_switch_provider():
            exp_provider.setCurrentIndex(1 if exp_provider.currentIndex() == 0 else 0)

        btn_switch.clicked.connect(_exp_switch_provider)

        def _exp_connected() -> bool:
            try:
                _ensure_providers()
                if _exp_provider_key() == "onedrive":
                    return bool(self._onedrive.is_connected())
                return bool(self._gdrive.is_connected())
            except Exception:
                return False

        def _exp_populate(items):
            remote_tree.clear()
            for it in items or []:
                typ = "Folder" if it.is_folder else "File"
                size = "" if it.is_folder else (f"{(it.size or 0)/1024/1024:.2f} MB" if it.size is not None else "")
                node = QTreeWidgetItem([it.name, typ, size])
                node.setData(0, Qt.UserRole, {"id": it.id, "is_folder": bool(it.is_folder), "name": it.name})
                remote_tree.addTopLevelItem(node)

        def _exp_refresh_remote():
            if not _exp_connected():
                QMessageBox.information(self, "Cloud Sync", "Please connect this provider first (OneDrive/Google Drive tab).")
                return
            key = _exp_provider_key()
            typed = (remote_path.text() or "").strip().strip("/").strip("\\")

            def _do(status, progress, cancel_cb):  # noqa: ARG001
                cancel_cb()
                _ensure_providers()
                if key == "onedrive":
                    return self._onedrive.list_folder(typed, limit=200)
                # Google Drive: resolve/create path under root → folder id, then list
                if typed:
                    fid = self._gdrive.ensure_folder_path(typed)
                    self._exp_gd_folder_id = fid
                    self._exp_gd_path_parts = [p for p in typed.replace("\\", "/").split("/") if p]
                    return self._gdrive.list_folder(fid, limit=200)
                self._exp_gd_folder_id = "root"
                self._exp_gd_path_parts = []
                return self._gdrive.list_folder("root", limit=200)

            _run_cloud_job(action=f"Explorer: Browse {key}", fn=_do, on_ok=_exp_populate)

        def _exp_up():
            key = _exp_provider_key()
            if key == "onedrive":
                cur = (remote_path.text() or "").strip().strip("/").strip("\\")
                if not cur:
                    return
                parts = [p for p in cur.replace("\\", "/").split("/") if p]
                if parts:
                    parts.pop()
                remote_path.setText("/".join(parts))
                _exp_refresh_remote()
                return
            # gdrive
            if not self._exp_gd_path_parts:
                remote_path.setText("")
                _exp_refresh_remote()
                return
            try:
                self._exp_gd_path_parts.pop()
            except Exception:
                self._exp_gd_path_parts = []
            remote_path.setText("/".join(self._exp_gd_path_parts))
            _exp_refresh_remote()

        def _exp_enter(item: QTreeWidgetItem, _col: int):
            meta = item.data(0, Qt.UserRole) or {}
            if not meta.get("is_folder"):
                return
            key = _exp_provider_key()
            name = str(meta.get("name") or "").strip()
            if not name:
                return
            if key == "onedrive":
                cur = (remote_path.text() or "").strip().strip("/").strip("\\")
                nxt = f"{cur}/{name}".strip("/") if cur else name
                remote_path.setText(nxt)
                _exp_refresh_remote()
                return
            fid = str(meta.get("id") or "").strip()
            if not fid:
                return
            self._exp_gd_folder_id = fid
            self._exp_gd_path_parts.append(name)
            remote_path.setText("/".join(self._exp_gd_path_parts))
            _exp_refresh_remote()

        remote_tree.itemDoubleClicked.connect(_exp_enter)
        btn_remote_refresh.clicked.connect(_exp_refresh_remote)
        btn_remote_up.clicked.connect(_exp_up)
        exp_provider.currentIndexChanged.connect(lambda _i: _exp_refresh_remote())

        # Remote clipboard (copy/cut/paste between folders; can also transfer across clouds for files).
        if not hasattr(self, "_cloud_remote_clipboard"):
            self._cloud_remote_clipboard = None  # type: ignore[assignment]

        def _current_remote_folder_ref():
            key = _exp_provider_key()
            typed = (remote_path.text() or "").strip().strip("/").strip("\\")
            if key == "onedrive":
                return typed  # path under root
            # gdrive folder id for current view
            return str(self._exp_gd_folder_id or "root")

        def _remote_selected_meta():
            it = remote_tree.currentItem()
            if it is None:
                return None
            return it.data(0, Qt.UserRole) or {}

        def _remote_selected_metas() -> list[dict]:
            items = remote_tree.selectedItems() or []
            metas: list[dict] = []
            for it in items:
                try:
                    meta = it.data(0, Qt.UserRole) or {}
                    if meta:
                        metas.append(dict(meta))
                except Exception:
                    continue
            if metas:
                return metas
            one = _remote_selected_meta()
            return [dict(one)] if one else []

        def _cloud_mgr():
            _ensure_providers()
            return self._cloud_mgr

        def _exp_new_folder():
            if not _exp_connected():
                QMessageBox.information(self, "Cloud Sync", "Please connect this provider first.")
                return
            name, ok = QInputDialog.getText(self, "New Folder", "Folder name:")
            if not ok:
                return
            name = (name or "").strip()
            if not name:
                return
            key = _exp_provider_key()
            parent_ref = _current_remote_folder_ref()

            def _do(status, progress, cancel_cb):  # noqa: ARG001
                cancel_cb()
                mgr = _cloud_mgr()
                if key == "onedrive":
                    return mgr.provider("onedrive").create_folder(parent_path=parent_ref, name=name)  # type: ignore[attr-defined]
                return mgr.provider("gdrive").create_folder(parent_id=str(parent_ref or "root"), name=name)  # type: ignore[attr-defined]

            _run_cloud_job(action=f"Explorer: New Folder ({key})", fn=_do, on_ok=lambda _it: _exp_refresh_remote())

        def _exp_rename():
            meta = _remote_selected_meta()
            if not meta:
                QMessageBox.information(self, "Rename", "Select a remote item.")
                return
            old = str(meta.get("name") or "")
            new, ok = QInputDialog.getText(self, "Rename", "New name:", text=old)
            if not ok:
                return
            new = (new or "").strip()
            if not new or new == old:
                return
            key = _exp_provider_key()
            item_id = str(meta.get("id") or "").strip()

            def _do(status, progress, cancel_cb):  # noqa: ARG001
                cancel_cb()
                mgr = _cloud_mgr()
                if key == "onedrive":
                    return mgr.provider("onedrive").rename_item(item_id, new)  # type: ignore[attr-defined]
                return mgr.provider("gdrive").rename_item(item_id, new)  # type: ignore[attr-defined]

            _run_cloud_job(action=f"Explorer: Rename ({key})", fn=_do, on_ok=lambda _it: _exp_refresh_remote())

        def _exp_delete():
            meta = _remote_selected_meta()
            if not meta:
                QMessageBox.information(self, "Delete", "Select a remote item.")
                return
            name = str(meta.get("name") or "")
            item_id = str(meta.get("id") or "").strip()
            if not item_id:
                return
            if QMessageBox.question(self, "Delete", f"Delete '{name}' from cloud?") != QMessageBox.Yes:
                return
            key = _exp_provider_key()

            def _do(status, progress, cancel_cb):  # noqa: ARG001
                cancel_cb()
                mgr = _cloud_mgr()
                if key == "onedrive":
                    mgr.provider("onedrive").delete_item(item_id)  # type: ignore[attr-defined]
                else:
                    mgr.provider("gdrive").delete_item(item_id)  # type: ignore[attr-defined]
                return True

            _run_cloud_job(action=f"Explorer: Delete ({key})", fn=_do, on_ok=lambda _res: _exp_refresh_remote())

        def _exp_copy(cut: bool = False):
            metas = _remote_selected_metas()
            if not metas:
                QMessageBox.information(self, "Copy/Cut", "Select a remote item.")
                return
            parent_ref = _current_remote_folder_ref()
            self._cloud_remote_clipboard = {
                "provider": _exp_provider_key(),
                "items": [
                    {
                        "id": str(m.get("id") or ""),
                        "name": str(m.get("name") or ""),
                        "is_folder": bool(m.get("is_folder")),
                        "parent_ref": str(parent_ref or ""),
                    }
                    for m in metas
                    if str(m.get("id") or "").strip()
                ],
                "op": "cut" if cut else "copy",
            }
            n = len(self._cloud_remote_clipboard.get("items") or [])
            label = (metas[0].get("name") if metas else "") if n == 1 else f"{n} items"
            self._cloud_status.setText(("Cut" if cut else "Copied") + f": {label}")

        def _exp_paste():
            clip = getattr(self, "_cloud_remote_clipboard", None) or None
            if not clip:
                QMessageBox.information(self, "Paste", "Nothing to paste yet.")
                return
            if not _exp_connected():
                QMessageBox.information(self, "Cloud Sync", "Please connect this provider first.")
                return
            cur_key = _exp_provider_key()
            dest_ref = _current_remote_folder_ref()

            def _do(status, progress, cancel_cb):
                cancel_cb()
                mgr = _cloud_mgr()
                src_key = str(clip.get("provider") or "")
                items = clip.get("items") or []
                op = str(clip.get("op") or "copy")

                if not items:
                    raise RuntimeError("Clipboard is empty.")

                def _within_one(item_id: str, *, is_folder: bool) -> None:
                    if cur_key == "onedrive":
                        if op == "cut":
                            mgr.provider("onedrive").move_item_to_path(item_id, dest_parent_path=str(dest_ref or ""))  # type: ignore[attr-defined]
                            return
                        mgr.provider("onedrive").copy_item_to_path(item_id, dest_parent_path=str(dest_ref or ""), new_name=None)  # type: ignore[attr-defined]
                        return
                    # gdrive
                    if op == "cut":
                        mgr.provider("gdrive").move_item(item_id, dest_parent_id=str(dest_ref or "root"))  # type: ignore[attr-defined]
                        return
                    if is_folder:
                        raise RuntimeError("Google Drive: folder copy is not supported yet. Use Cut/Paste to move folders.")
                    mgr.provider("gdrive").copy_item(item_id, dest_parent_id=str(dest_ref or "root"))  # type: ignore[attr-defined]

                if src_key == cur_key:
                    for i, it in enumerate(items, start=1):
                        cancel_cb()
                        item_id = str(it.get("id") or "").strip()
                        if not item_id:
                            continue
                        is_folder = bool(it.get("is_folder"))
                        if progress:
                            progress(i, len(items), f"{'Moving' if op=='cut' else 'Copying'}: {it.get('name')}")
                        _within_one(item_id, is_folder=is_folder)
                    return True

                # cross-cloud: download + upload
                other = "onedrive" if cur_key == "gdrive" else "gdrive"
                # src_key is the other side; ensure it is connected too
                if not mgr.provider(src_key).is_connected():  # type: ignore[arg-type]
                    raise RuntimeError(f"Source provider '{src_key}' is not connected.")

                # If destination is gdrive and dest_ref is a path string, resolve it.
                dest_folder = dest_ref
                if cur_key == "gdrive":
                    typed = (remote_path.text() or "").strip().strip("/").strip("\\")
                    if typed:
                        dest_folder = mgr.provider("gdrive").ensure_folder_path(typed)  # type: ignore[attr-defined]
                # Cross-cloud: files only (folders not supported yet).
                for i, it in enumerate(items, start=1):
                    cancel_cb()
                    item_id = str(it.get("id") or "").strip()
                    item_name = str(it.get("name") or "file")
                    if progress:
                        progress(i, len(items), f"Transferring: {item_name}")
                    if bool(it.get("is_folder")):
                        # Resolve folder ref for OneDrive (path) vs Google Drive (id).
                        src_folder_ref = str(item_id or "")
                        if src_key == "onedrive":
                            parent_path = str(it.get("parent_ref") or "").strip().strip("/").strip("\\")
                            name = str(item_name or "").strip().strip("/").strip("\\")
                            src_folder_ref = f"{parent_path}/{name}".strip("/").strip("\\") if parent_path else name
                        mgr.transfer_folder_between_providers(
                            src=src_key,  # type: ignore[arg-type]
                            src_folder_ref=str(src_folder_ref or ""),
                            src_folder_name=item_name,
                            dest=cur_key,  # type: ignore[arg-type]
                            dest_folder=str(dest_folder or ""),
                            progress_cb=progress,
                            cancel_cb=cancel_cb,
                        )
                    else:
                        mgr.transfer_file_between_providers(
                            src=src_key,  # type: ignore[arg-type]
                            src_item_id=item_id,
                            src_name=item_name,
                            dest=cur_key,  # type: ignore[arg-type]
                            dest_folder=str(dest_folder or ""),
                            progress_cb=progress,
                            cancel_cb=cancel_cb,
                        )
                    if op == "cut":
                        try:
                            if src_key == "onedrive":
                                mgr.provider("onedrive").delete_item(item_id)  # type: ignore[attr-defined]
                            elif src_key == "gdrive":
                                mgr.provider("gdrive").delete_item(item_id)  # type: ignore[attr-defined]
                        except Exception:
                            pass
                return True

            _run_cloud_job(action=f"Explorer: Paste → {cur_key}", fn=_do, on_ok=lambda _r: _exp_refresh_remote())

        def _exp_send_to_other():
            meta = _remote_selected_meta()
            if not meta:
                QMessageBox.information(self, "Send to Other Cloud", "Select a remote file.")
                return
            if not _exp_connected():
                QMessageBox.information(self, "Cloud Sync", "Please connect this provider first.")
                return
            cur_key = _exp_provider_key()
            other_key = "gdrive" if cur_key == "onedrive" else "onedrive"
            try:
                _ensure_providers()
                if not self._cloud_mgr.provider(other_key).is_connected():
                    QMessageBox.information(self, "Send to Other Cloud", f"Connect {('Google Drive' if other_key=='gdrive' else 'OneDrive')} first.")
                    return
            except Exception:
                pass
            dest_path_str, ok = QInputDialog.getText(
                self,
                "Destination Folder",
                f"Destination folder on {('Google Drive' if other_key=='gdrive' else 'OneDrive')} (path under root):",
                text="Fylorra Sync",
            )
            if not ok:
                return
            dest_path_str = (dest_path_str or "").strip().strip("/").strip("\\")
            item_id = str(meta.get("id") or "")
            item_name = str(meta.get("name") or "file")
            is_folder = bool(meta.get("is_folder"))

            def _resolve_src_folder_ref() -> str:
                if cur_key != "onedrive":
                    return str(item_id or "")
                parent_path = str(_current_remote_folder_ref() or "").strip().strip("/").strip("\\")
                name = str(item_name or "").strip().strip("/").strip("\\")
                if not name:
                    return parent_path
                return f"{parent_path}/{name}".strip("/").strip("\\") if parent_path else name

            def _do(status, progress, cancel_cb):
                cancel_cb()
                mgr = _cloud_mgr()
                dest_folder = dest_path_str
                if other_key == "gdrive":
                    dest_folder = mgr.provider("gdrive").ensure_folder_path(dest_path_str or "Fylorra Sync")  # type: ignore[attr-defined]
                if is_folder:
                    return mgr.transfer_folder_between_providers(
                        src=cur_key,  # type: ignore[arg-type]
                        src_folder_ref=_resolve_src_folder_ref(),
                        src_folder_name=item_name,
                        dest=other_key,  # type: ignore[arg-type]
                        dest_folder=str(dest_folder or ""),
                        progress_cb=progress,
                        cancel_cb=cancel_cb,
                    )
                return mgr.transfer_file_between_providers(
                    src=cur_key,  # type: ignore[arg-type]
                    src_item_id=item_id,
                    src_name=item_name,
                    dest=other_key,  # type: ignore[arg-type]
                    dest_folder=str(dest_folder or ""),
                    progress_cb=progress,
                    cancel_cb=cancel_cb,
                )

            _run_cloud_job(action=f"Explorer: Send {cur_key} → {other_key}", fn=_do, on_ok=lambda _r: QMessageBox.information(self, "Send to Other Cloud", "Done."))

        btn_new_folder.clicked.connect(_exp_new_folder)
        btn_rename.clicked.connect(_exp_rename)
        btn_delete.clicked.connect(_exp_delete)
        btn_copy.clicked.connect(lambda: _exp_copy(False))
        btn_cut.clicked.connect(lambda: _exp_copy(True))
        btn_paste.clicked.connect(_exp_paste)
        btn_send.clicked.connect(_exp_send_to_other)

        # Context menus
        def _remote_menu(pos):
            m = QMenu(self)
            a_new = m.addAction("New Folder…")
            a_ren = m.addAction("Rename…")
            a_del = m.addAction("Delete")
            m.addSeparator()
            a_copy = m.addAction("Copy")
            a_cut = m.addAction("Cut")
            a_paste = m.addAction("Paste")
            m.addSeparator()
            a_send = m.addAction("Send to Other Cloud…")
            act = m.exec(remote_tree.viewport().mapToGlobal(pos))
            if act == a_new:
                _exp_new_folder()
            elif act == a_ren:
                _exp_rename()
            elif act == a_del:
                _exp_delete()
            elif act == a_copy:
                _exp_copy(False)
            elif act == a_cut:
                _exp_copy(True)
            elif act == a_paste:
                _exp_paste()
            elif act == a_send:
                _exp_send_to_other()

        remote_tree.setContextMenuPolicy(Qt.CustomContextMenu)
        remote_tree.customContextMenuRequested.connect(_remote_menu)

        def _local_menu(pos):
            m = QMenu(self)
            a_up = m.addAction("Upload Selected (to current cloud folder)")
            a_up_to = m.addAction("Upload to…")
            act = m.exec(local_view.viewport().mapToGlobal(pos))
            if act == a_up:
                _exp_upload()
            elif act == a_up_to:
                _exp_upload_to()

        local_view.setContextMenuPolicy(Qt.CustomContextMenu)
        local_view.customContextMenuRequested.connect(_local_menu)

        def _selected_local_paths() -> list[Path]:
            try:
                sel = local_view.selectionModel()
                if sel is not None:
                    idxs = sel.selectedRows(0)
                else:
                    idxs = []
            except Exception:
                idxs = []

            out: list[Path] = []
            for idx in idxs:
                try:
                    if not idx.isValid():
                        continue
                    p = Path(local_model.filePath(idx))
                    if p.exists():
                        out.append(p)
                except Exception:
                    continue

            if out:
                return out

            # Fallback to current index.
            try:
                idx = local_view.currentIndex()
                if idx.isValid():
                    p = Path(local_model.filePath(idx))
                    if p.exists():
                        return [p]
            except Exception:
                pass
            return []

        def _selected_local_path() -> Path | None:
            paths = _selected_local_paths()
            return paths[0] if paths else None

        def _selected_remote_item():
            it = remote_tree.currentItem()
            if it is None:
                return None
            return it.data(0, Qt.UserRole) or {}

        def _exp_sync():
            lp = _selected_local_path()
            if lp is None or not lp.exists() or not lp.is_dir():
                QMessageBox.information(self, "Sync", "Select a local folder to sync.")
                return
            if not _exp_connected():
                QMessageBox.information(self, "Cloud Sync", "Please connect this provider first.")
                return
            key = _exp_provider_key()
            typed = (remote_path.text() or "").strip().strip("/").strip("\\")

            def _do(status, progress, cancel_cb):
                cancel_cb()
                _ensure_providers()
                remote_base = f"{typed}/{lp.name}".strip("/") if typed else lp.name
                se = __import__("core.cloud_sync.sync_engine", fromlist=["sync_folder_upload_only"])
                prov = self._onedrive if key == "onedrive" else self._gdrive
                return se.sync_folder_upload_only(
                    prov,
                    local_root=lp,
                    remote_base=remote_base or "Fylorra Sync",
                    include_subfolders=True,
                    dry_run=False,
                    status_cb=status,
                    progress_cb=progress,
                    cancel_cb=cancel_cb,
                )

            _run_cloud_job(
                action=f"Explorer: Sync Folder ({key})",
                fn=_do,
                on_ok=lambda _it: (_exp_refresh_remote(), QMessageBox.information(self, "Sync", f"Synced: {lp.name}")),
            )

        def _delete_existing_same_name(*, provider_key: str, dest_path: str, filename: str) -> None:
            if not filename:
                return
            dest_path = str(dest_path or "").strip().strip("/").strip("\\")
            try:
                if provider_key == "onedrive":
                    items = self._onedrive.list_folder(dest_path or None, limit=200)
                else:
                    fid = self._gdrive.ensure_folder_path(dest_path) if dest_path else "root"
                    items = self._gdrive.list_folder(fid, limit=200)
                for it in items or []:
                    if (it.name or "") == filename and not it.is_folder:
                        if provider_key == "onedrive":
                            self._onedrive.delete_item(it.id)
                        else:
                            self._gdrive.delete_item(it.id)
                        return
            except Exception:
                return

        def _pick_cloud_folder_dialog(*, title_text: str) -> dict | None:
            """
            Returns:
              {"provider": "onedrive"|"gdrive", "path": "A/B/C", "keep_structure": bool, "overwrite": bool}
            where path is under the remote root.
            """

            dlg = QDialog(self)
            dlg.setWindowTitle(title_text)
            dlg.setStyleSheet(_qt_modern_dialog_stylesheet())
            dlg.setModal(True)

            root_l = QVBoxLayout(dlg)
            root_l.setContentsMargins(16, 14, 16, 14)
            root_l.setSpacing(12)

            header = QFrame()
            header.setObjectName("DialogHeader")
            hl = QVBoxLayout(header)
            hl.setContentsMargins(14, 12, 14, 12)
            hl.setSpacing(2)
            t = QLabel(title_text)
            t.setObjectName("DialogTitle")
            s = QLabel("Choose a destination folder, then click OK.")
            s.setObjectName("DialogSubtitle")
            hl.addWidget(t)
            hl.addWidget(s)
            root_l.addWidget(header)

            row = QHBoxLayout()
            row.setSpacing(10)
            provider_cb = QComboBox()
            provider_cb.addItems(["OneDrive", "Google Drive"])
            try:
                provider_cb.setItemIcon(0, _cs_icon("onedrive"))
                provider_cb.setItemIcon(1, _cs_icon("google_drive"))
            except Exception:
                pass
            provider_cb.setMinimumWidth(220)
            row.addWidget(provider_cb)

            path_edit = QLineEdit()
            path_edit.setPlaceholderText("Cloud folder (under root).")
            row.addWidget(path_edit, 1)

            btn_up = _cs_btn("up", "Go up", text=None, fixed=38)
            btn_refresh = _cs_btn("refresh", "Refresh", text=None, fixed=38)
            btn_newf = _cs_btn("new-folder", "New folder", text=None, fixed=38)
            row.addWidget(btn_up)
            row.addWidget(btn_refresh)
            row.addWidget(btn_newf)
            root_l.addLayout(row)

            search_row = QHBoxLayout()
            search_row.setSpacing(10)
            search = QLineEdit()
            search.setPlaceholderText("Search folders.")
            search_row.addWidget(search, 1)
            btn_search = _cs_btn("search", "Filter folders", text=None, fixed=38)
            search_row.addWidget(btn_search)
            root_l.addLayout(search_row)

            tree = QTreeWidget()
            tree.setObjectName("ResultsTree")
            tree.setColumnCount(1)
            tree.setHeaderLabels(["Folders"])
            try:
                tree.header().setSectionResizeMode(0, QHeaderView.Stretch)
            except Exception:
                pass
            root_l.addWidget(tree, 1)

            opts = QFrame()
            opts.setObjectName("Card")
            ol = QHBoxLayout(opts)
            ol.setContentsMargins(12, 10, 12, 10)
            ol.setSpacing(14)
            keep_structure = QCheckBox("Keep folder structure")
            keep_structure.setChecked(True)
            overwrite = QCheckBox("Overwrite existing files")
            overwrite.setChecked(False)
            ol.addWidget(keep_structure)
            ol.addWidget(overwrite)
            ol.addStretch(1)
            root_l.addWidget(opts)

            buttons = QDialogButtonBox(QDialogButtonBox.Cancel | QDialogButtonBox.Ok)
            try:
                buttons.button(QDialogButtonBox.Ok).setObjectName("PrimaryButton")
                buttons.button(QDialogButtonBox.Cancel).setObjectName("SecondaryButton")
            except Exception:
                pass
            buttons.accepted.connect(dlg.accept)
            buttons.rejected.connect(dlg.reject)
            root_l.addWidget(buttons)

            state = {"path": "", "gd_folder_id": "root"}

            def _provider_key() -> str:
                return "onedrive" if provider_cb.currentIndex() == 0 else "gdrive"

            def _ensure_current_folder_for_gdrive() -> str:
                if not state["path"]:
                    state["gd_folder_id"] = "root"
                    return "root"
                fid = self._gdrive.ensure_folder_path(state["path"])
                state["gd_folder_id"] = fid
                return fid

            def _provider_connected() -> bool:
                try:
                    _ensure_providers()
                    return bool(self._onedrive.is_connected()) if _provider_key() == "onedrive" else bool(self._gdrive.is_connected())
                except Exception:
                    return False

            def _populate():
                tree.clear()
                if not _provider_connected():
                    return
                key = _provider_key()
                typed = (path_edit.text() or "").strip().strip("/").strip("\\")
                state["path"] = typed
                items = []
                try:
                    _ensure_providers()
                    if key == "onedrive":
                        items = self._onedrive.list_folder(typed or None, limit=200)
                    else:
                        fid = _ensure_current_folder_for_gdrive()
                        items = self._gdrive.list_folder(fid, limit=200)
                except Exception:
                    items = []
                q = (search.text() or "").strip().lower()
                for it in items or []:
                    if not it.is_folder:
                        continue
                    if q and (q not in (it.name or "").lower()):
                        continue
                    node = QTreeWidgetItem([it.name])
                    node.setData(0, Qt.UserRole, {"name": it.name})
                    tree.addTopLevelItem(node)

            def _go_into_selected():
                it = tree.currentItem()
                if it is None:
                    return
                name = (it.data(0, Qt.UserRole) or {}).get("name") or it.text(0)
                name = str(name or "").strip()
                if not name:
                    return
                cur = (path_edit.text() or "").strip().strip("/").strip("\\")
                path_edit.setText(f"{cur}/{name}".strip("/") if cur else name)
                _populate()

            def _go_up():
                cur = (path_edit.text() or "").strip().strip("/").strip("\\")
                if not cur:
                    return
                parts = [p for p in cur.replace("\\", "/").split("/") if p]
                if not parts:
                    return
                parts.pop()
                path_edit.setText("/".join(parts))
                _populate()

            def _new_folder():
                if not _provider_connected():
                    QMessageBox.information(self, "Cloud Sync", "Please connect the provider first.")
                    return
                name, ok = QInputDialog.getText(dlg, "New Folder", "Folder name:")
                if not ok:
                    return
                name = str(name or "").strip()
                if not name:
                    return
                key = _provider_key()
                typed = (path_edit.text() or "").strip().strip("/").strip("\\")
                try:
                    _ensure_providers()
                    if key == "onedrive":
                        self._onedrive.create_folder(parent_path=typed or None, name=name)
                    else:
                        parent_id = _ensure_current_folder_for_gdrive()
                        self._gdrive.create_folder(parent_id=parent_id, name=name)
                    _populate()
                except Exception as e:
                    QMessageBox.critical(dlg, "New Folder", str(e))

            btn_up.clicked.connect(_go_up)
            btn_refresh.clicked.connect(_populate)
            btn_newf.clicked.connect(_new_folder)
            tree.itemDoubleClicked.connect(lambda *_: _go_into_selected())
            btn_search.clicked.connect(_populate)
            search.returnPressed.connect(_populate)
            path_edit.returnPressed.connect(_populate)
            provider_cb.currentIndexChanged.connect(_populate)

            try:
                provider_cb.setCurrentIndex(0 if _exp_provider_key() == "onedrive" else 1)
            except Exception:
                pass
            path_edit.setText((remote_path.text() or "").strip())
            _populate()

            if dlg.exec() != QDialog.Accepted:
                return None
            return {
                "provider": _provider_key(),
                "path": (path_edit.text() or "").strip().strip("/").strip("\\"),
                "keep_structure": bool(keep_structure.isChecked()),
                "overwrite": bool(overwrite.isChecked()),
            }

        def _exp_upload():
            paths = _selected_local_paths()
            if not paths:
                QMessageBox.information(self, "Upload", "Select a local file or folder to upload.")
                return
            if not _exp_connected():
                QMessageBox.information(self, "Cloud Sync", "Please connect this provider first.")
                return
            key = _exp_provider_key()
            typed = (remote_path.text() or "").strip().strip("/").strip("\\")

            def _do(status, progress, cancel_cb):
                cancel_cb()
                _ensure_providers()
                se = __import__("core.cloud_sync.sync_engine", fromlist=["sync_folder_upload_only"])
                prov = self._onedrive if key == "onedrive" else self._gdrive
                for i, lp in enumerate(paths):
                    cancel_cb()
                    status(f"Uploading {i + 1}/{len(paths)}: {lp.name}")
                    if lp.is_dir():
                        remote_base = f"{typed}/{lp.name}".strip("/") if typed else lp.name
                        se.sync_folder_upload_only(
                            prov,
                            local_root=lp,
                            remote_base=remote_base or "Fylorra Sync",
                            include_subfolders=True,
                            dry_run=False,
                            status_cb=status,
                            progress_cb=progress,
                            cancel_cb=cancel_cb,
                        )
                        continue
                    if key == "onedrive":
                        self._onedrive.upload_file(lp, remote_folder=typed or None, progress_cb=progress)
                        continue
                    folder_id = self._exp_gd_folder_id
                    if typed:
                        folder_id = self._gdrive.ensure_folder_path(typed)
                        self._exp_gd_folder_id = folder_id
                    self._gdrive.upload_file(lp, remote_folder=folder_id, progress_cb=progress)
                return True

            _run_cloud_job(
                action=f"Explorer: Upload ({key})",
                fn=_do,
                on_ok=lambda _it: (_exp_refresh_remote(), QMessageBox.information(self, "Upload", f"Uploaded: {len(paths)} item(s).")),
            )

        def _exp_upload_to():
            paths = _selected_local_paths()
            if not paths:
                QMessageBox.information(self, "Upload", "Select a local file or folder to upload.")
                return
            if not _exp_connected():
                QMessageBox.information(self, "Cloud Sync", "Please connect this provider first.")
                return

            picked = _pick_cloud_folder_dialog(title_text="Upload To")
            if not picked:
                return

            provider_key = str(picked.get("provider") or "").strip()
            dest_path = str(picked.get("path") or "").strip().strip("/").strip("\\")
            keep_structure = bool(picked.get("keep_structure"))
            overwrite = bool(picked.get("overwrite"))

            # Reflect selection in the Explorer UI.
            try:
                exp_provider.setCurrentIndex(0 if provider_key == "onedrive" else 1)
            except Exception:
                pass
            try:
                remote_path.setText(dest_path or "Fylorra Sync")
            except Exception:
                pass

            def _do(status, progress, cancel_cb):
                cancel_cb()
                _ensure_providers()
                se = __import__("core.cloud_sync.sync_engine", fromlist=["sync_folder_upload_only"])
                prov = self._onedrive if provider_key == "onedrive" else self._gdrive

                for i, lp in enumerate(paths):
                    cancel_cb()
                    status(f"Uploading {i + 1}/{len(paths)}: {lp.name}")
                    if lp.is_dir():
                        remote_base = dest_path
                        if keep_structure:
                            remote_base = f"{dest_path}/{lp.name}".strip("/") if dest_path else lp.name
                        se.sync_folder_upload_only(
                            prov,
                            local_root=lp,
                            remote_base=(remote_base or "Fylorra Sync"),
                            include_subfolders=True,
                            dry_run=False,
                            status_cb=status,
                            progress_cb=progress,
                            cancel_cb=cancel_cb,
                        )
                        continue

                    if overwrite:
                        _delete_existing_same_name(provider_key=provider_key, dest_path=dest_path, filename=lp.name)

                    if provider_key == "onedrive":
                        self._onedrive.upload_file(lp, remote_folder=dest_path or None, progress_cb=progress)
                    else:
                        folder_id = self._gdrive.ensure_folder_path(dest_path) if dest_path else "root"
                        self._gdrive.upload_file(lp, remote_folder=folder_id, progress_cb=progress)
                return True

            _run_cloud_job(
                action=f"Explorer: Upload To ({provider_key})",
                fn=_do,
                on_ok=lambda _it: (_exp_refresh_remote(), QMessageBox.information(self, "Upload", f"Uploaded: {len(paths)} item(s).")),
            )

        def _exp_download():
            meta = _selected_remote_item()
            if not meta:
                QMessageBox.information(self, "Download", "Select a remote file to download.")
                return
            if bool(meta.get("is_folder")):
                QMessageBox.information(self, "Download", "Select a file (not a folder).")
                return
            if not _exp_connected():
                QMessageBox.information(self, "Cloud Sync", "Please connect this provider first.")
                return
            key = _exp_provider_key()
            item_id = str(meta.get("id") or "").strip()
            name = str(meta.get("name") or "download")
            dest_dir = Path(local_root.text().strip() or str(Path.home()))
            dest_dir.mkdir(parents=True, exist_ok=True)
            dest = dest_dir / name

            def _do(status, progress, cancel_cb):  # noqa: ARG001
                cancel_cb()
                _ensure_providers()
                if key == "onedrive":
                    return self._onedrive.download_file(item_id, dest, progress_cb=progress)
                return self._gdrive.download_file(item_id, dest, progress_cb=progress)

            _run_cloud_job(
                action=f"Explorer: Download ← {key}",
                fn=_do,
                on_ok=lambda _p: QMessageBox.information(self, "Download", f"Saved to: {dest}"),
            )

        exp_upload.clicked.connect(_exp_upload)
        exp_download.clicked.connect(_exp_download)
        btn_sync.clicked.connect(_exp_sync)

        tabs.addTab(exp_tab, "Explorer")

        layout.addWidget(card, 1)
        _refresh_status()
        self._refresh_cloud_sync_theme()
        return host

    def _build_workspace_page(self) -> QWidget:
        host = QFrame()
        host.setObjectName("PageHost")
        layout = QVBoxLayout(host)
        layout.setContentsMargins(18, 18, 18, 18)
        layout.setSpacing(14)

        header = QFrame()
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(0, 0, 0, 0)
        header_layout.setSpacing(10)
        title = QLabel("Workspace")
        title.setObjectName("PageTitle")
        header_layout.addWidget(title)
        sub = QLabel("Pick a folder once, then run multiple tools together in a single workflow.")
        sub.setObjectName("PageSubTitle")
        header_layout.addWidget(sub, 1)
        layout.addWidget(header)

        card = QFrame()
        card.setObjectName("PageCard")
        c = QVBoxLayout(card)
        c.setContentsMargins(16, 14, 16, 14)
        c.setSpacing(12)

        # Target folder
        row = QHBoxLayout()
        row.setSpacing(10)
        self._ws_target = QLineEdit()
        self._ws_target.setPlaceholderText("Target folder…")
        row.addWidget(self._ws_target, 1)
        self._ws_include_sub = QCheckBox("Include subfolders")
        self._ws_include_sub.setChecked(True)
        row.addWidget(self._ws_include_sub)
        b = QPushButton("Browse")
        b.clicked.connect(lambda: self._ft_browse_target_into(self._ws_target))
        row.addWidget(b)
        c.addLayout(row)

        # Actions grid
        actions_box = QGroupBox("Actions")
        box_l = QVBoxLayout(actions_box)
        box_l.setContentsMargins(12, 10, 12, 10)
        box_l.setSpacing(10)

        top_row = QHBoxLayout()
        top_row.setSpacing(10)
        top_row.addStretch(1)
        sel_all = QPushButton("Select All")
        sel_none = QPushButton("Clear")
        sel_all.clicked.connect(lambda: [cb.setChecked(True) for cb in self._ws_actions.values()])
        sel_none.clicked.connect(lambda: [cb.setChecked(False) for cb in self._ws_actions.values()])
        top_row.addWidget(sel_all)
        top_row.addWidget(sel_none)
        self._ws_run = QPushButton("Run Selected")
        self._ws_run.setObjectName("PrimaryButton")
        self._ws_run.setIcon(self.icons.icon("play"))
        self._ws_run.setIconSize(QSize(18, 18))
        self._ws_run.clicked.connect(self._ws_run_selected)
        top_row.addWidget(self._ws_run)
        box_l.addLayout(top_row)

        ag_host = QWidget()
        ag = QGridLayout(ag_host)
        ag.setContentsMargins(0, 0, 0, 0)
        ag.setHorizontalSpacing(14)
        ag.setVerticalSpacing(12)

        self._ws_actions = {}

        def add_action(key: str, text: str, desc: str, i: int, default: bool = False):
            card = QFrame()
            card.setObjectName("PageCard")
            card.setFixedHeight(86)
            try:
                card.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
            except Exception:
                pass
            cl = QVBoxLayout(card)
            cl.setContentsMargins(12, 10, 12, 10)
            cl.setSpacing(6)

            cb = QCheckBox(text)
            cb.setChecked(bool(default))
            cb.setToolTip(desc)
            self._ws_actions[key] = cb
            cl.addWidget(cb)

            d = QLabel(desc)
            d.setWordWrap(True)
            d.setStyleSheet("color:#9aa0a9; font-size:11px;")
            cl.addWidget(d)

            r = i // 2
            col = i % 2
            ag.addWidget(card, r, col)

        add_action("index_folder", "Index Folder", "Build/update the local search index for this folder.", 0)
        add_action("convert_images_webp", "Images → WebP", "Convert images to WebP in Converted_Images.", 1)
        add_action("convert_media_mp4", "Media → MP4", "Convert media to MP4 in Converted_Media.", 2)
        add_action("zip_folder", "Archive (ZIP)", "Create Archive.zip in the target folder.", 3)
        add_action("ai_hub_rename", "Smart Rename (AI)", "Run AI Smart Rename (report-only by default).", 4)
        add_action("ai_hub_categorize", "Auto-Categorize (AI)", "Run AI Auto-Categorize (report-only by default).", 5)
        add_action("ai_hub_security", "Security Scan (AI)", "Scan images for sensitive content and write report.", 6)
        add_action("ai_hub_content", "Content Analysis (AI)", "Analyze documents and write report (bounded).", 7)

        # Put the grid into a scroll area so opening Jobs doesn't squash the cards.
        ag_scroll = QScrollArea()
        ag_scroll.setWidgetResizable(True)
        ag_scroll.setFrameShape(QFrame.NoFrame)
        ag_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        ag_scroll.setWidget(ag_host)
        try:
            # Keep the grid's "natural" height so rows don't compress when viewport is small.
            ag_host.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.Minimum)
            import math

            rows = max(1, int(math.ceil(len(self._ws_actions) / 2.0)))
            min_h = rows * 86 + max(0, rows - 1) * int(ag.verticalSpacing()) + 16
            ag_host.setMinimumHeight(max(min_h, ag_host.sizeHint().height(), 1))
        except Exception:
            pass
        box_l.addWidget(ag_scroll, 1)
        # Give Actions the flexible space (prevents child cards from getting vertically squashed).
        c.addWidget(actions_box, 1)

        # Progress/log
        self._ws_status = QLabel("Ready.")
        self._ws_status.setStyleSheet("color:#9aa0a9;")
        self._ws_bar = QProgressBar()
        self._ws_bar.setRange(0, 1000)
        self._ws_bar.setValue(0)
        self._ws_cancel = QPushButton("Cancel")
        self._ws_cancel.setEnabled(False)
        self._ws_cancel.clicked.connect(self._ws_cancel_run)
        self._ws_log = QListWidget()
        self._ws_log.setStyleSheet(
            "QListWidget{background:#14171c; border:1px solid #232730; border-radius:10px; color:#c8ccd6;}"
        )
        self._ws_log.setMinimumHeight(110)
        self._ws_log.setMaximumHeight(160)

        c.addWidget(self._ws_bar)
        plw = QWidget()
        pl = QGridLayout(plw)
        pl.setContentsMargins(0, 0, 0, 0)
        pl.setHorizontalSpacing(10)
        pl.setVerticalSpacing(0)

        # Keep the Jobs toggle visually fixed (center column) so it doesn't shift when status text changes.
        self._ws_jobs_toggle = QToolButton()
        self._ws_jobs_toggle.setCheckable(True)
        self._ws_jobs_toggle.setChecked(False)  # closed by default
        self._ws_jobs_toggle.setAutoRaise(True)
        self._ws_jobs_toggle.setText("Jobs ▸")
        self._ws_jobs_toggle.setToolTip("Show jobs")
        self._ws_jobs_toggle.setStyleSheet(
            "QToolButton{background:#1a1e25; color:#c8ccd6; border:1px solid #232730; border-radius:12px; padding:6px 12px;}"
            "QToolButton:hover{background:#222833;}"
            "QToolButton:checked{background:#243040; color:#ffffff; border:1px solid #2a3b52;}"
        )

        pl.addWidget(self._ws_status, 0, 0, 1, 1, Qt.AlignVCenter | Qt.AlignLeft)
        pl.addWidget(self._ws_jobs_toggle, 0, 1, 1, 1, Qt.AlignVCenter | Qt.AlignHCenter)
        pl.addWidget(self._ws_cancel, 0, 2, 1, 1, Qt.AlignVCenter | Qt.AlignRight)
        pl.setColumnStretch(0, 1)
        pl.setColumnStretch(1, 0)
        pl.setColumnStretch(2, 0)
        c.addWidget(plw)

        # Collapsible Jobs/Log box (hidden by default to save space).
        self._ws_jobs_frame = QFrame()
        self._ws_jobs_frame.setObjectName("PageCard")
        jl = QVBoxLayout(self._ws_jobs_frame)
        jl.setContentsMargins(12, 10, 12, 10)
        jl.setSpacing(8)

        jl.addWidget(self._ws_log)
        c.addWidget(self._ws_jobs_frame)

        def _set_jobs_open(open_: bool):
            open_ = bool(open_)
            try:
                self._ws_jobs_frame.setVisible(open_)
            except Exception:
                pass
            try:
                self._ws_jobs_toggle.setText("Jobs ▾" if open_ else "Jobs ▸")
                self._ws_jobs_toggle.setToolTip("Hide jobs" if open_ else "Show jobs")
            except Exception:
                pass

        self._ws_jobs_toggle.toggled.connect(_set_jobs_open)
        _set_jobs_open(False)

        layout.addWidget(card, 1)
        return host

    def _ft_browse_target_into(self, line_edit: QLineEdit):
        try:
            path = QFileDialog.getExistingDirectory(self, "Select folder")
            if path:
                line_edit.setText(path)
        except Exception:
            pass

    def _ws_cancel_run(self):
        try:
            w = getattr(self, "_ws_worker", None)
            if w:
                w.cancel()
        except Exception:
            pass
        self._ws_cancel.setEnabled(False)
        self._ws_status.setText("Cancelling…")

    def _ws_run_selected(self):
        folder = (self._ws_target.text() or "").strip()
        if not folder:
            QMessageBox.warning(self, "Workspace", "Pick a target folder.")
            return

        selected = [k for k, cb in self._ws_actions.items() if cb.isChecked()]
        if not selected:
            QMessageBox.information(self, "Workspace", "Select at least one action.")
            return

        needs_ai = any(k.startswith("ai_hub_") for k in selected)
        if needs_ai and not self._ensure_ai_ready(title="AI Required for Workspace"):
            return

        self._ws_run.setEnabled(False)
        self._ws_cancel.setEnabled(True)
        self._ws_bar.setValue(0)
        self._ws_log.clear()
        self._ws_status.setText("Starting…")

        worker = _QtWorkspaceWorker(
            backend=self.backend,
            folder=folder,
            include_subfolders=bool(self._ws_include_sub.isChecked()),
            actions=selected,
        )
        th = QThread(self)
        worker.moveToThread(th)
        th.started.connect(worker.run)
        worker.status.connect(self._ws_status.setText)
        worker.log.connect(lambda m: self._ws_log.insertItem(0, QListWidgetItem(m)))
        worker.progress.connect(lambda p: self._ws_bar.setValue(int(max(0.0, min(1.0, p)) * 1000)))
        worker.review_requested.connect(self._ws_review_requested)
        worker.finished.connect(self._ws_done)
        worker.error.connect(self._ws_failed)
        worker.finished.connect(th.quit)
        worker.error.connect(th.quit)
        th.finished.connect(th.deleteLater)
        th.finished.connect(worker.deleteLater)
        self._ws_worker = worker
        self._ws_thread = th
        th.start()

    def _ws_review_requested(self, kind: str, payload: dict):
        """
        Called when a Workspace action wants user approval before applying changes.
        Runs in the UI thread (Qt signal).
        """
        try:
            dlg = _QtWorkspaceReviewDialog(self, kind=str(kind), payload=dict(payload or {}))
            res = dlg.exec()
            decision = dict(dlg.result_payload or {})
            if res != QDialog.Accepted:
                decision = {"decision": "cancel", "selected": []}
        except Exception:
            decision = {"decision": "cancel", "selected": []}

        try:
            w = getattr(self, "_ws_worker", None)
            if w:
                w.submit_review_result(decision)
        except Exception:
            pass

    def _ws_done(self, msg: str):
        self._ws_run.setEnabled(True)
        self._ws_cancel.setEnabled(False)
        self._ws_bar.setValue(1000)
        self._ws_status.setText(msg or "Done.")
        self._ws_log.insertItem(0, QListWidgetItem("Done."))

    def _ws_failed(self, msg: str):
        self._ws_run.setEnabled(True)
        self._ws_cancel.setEnabled(False)
        self._ws_status.setText("Failed.")
        QMessageBox.critical(self, "Workspace", msg or "Failed.")

    def _build_media_editors_page(self) -> QWidget:
        host = QFrame()
        host.setObjectName("PageHost")
        layout = QVBoxLayout(host)
        layout.setContentsMargins(18, 18, 18, 18)
        layout.setSpacing(14)

        header = QFrame()
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(0, 0, 0, 0)
        header_layout.setSpacing(10)
        title = QLabel("Media Editors")
        title.setObjectName("PageTitle")
        header_layout.addWidget(title)
        sub = QLabel("Open Audio or Video editor directly (no File Tools needed).")
        sub.setObjectName("PageSubTitle")
        header_layout.addWidget(sub, 1)

        open_any = QPushButton("Open File…")
        open_any.setObjectName("PrimaryButton")
        open_any.setIcon(self.icons.icon("folder"))
        open_any.setIconSize(QSize(18, 18))
        open_any.clicked.connect(self._me_open_any_file)
        header_layout.addWidget(open_any)

        layout.addWidget(header)

        body = QFrame()
        body_layout = QGridLayout(body)
        body_layout.setContentsMargins(0, 0, 0, 0)
        body_layout.setHorizontalSpacing(14)
        body_layout.setVerticalSpacing(14)

        video = QFrame()
        video.setObjectName("PageCard")
        vl = QVBoxLayout(video)
        vl.setContentsMargins(16, 14, 16, 14)
        vl.setSpacing(10)
        vl.addWidget(QLabel("Video Editor"))
        vd = QLabel("Edit videos, images (slideshow), audio lanes and render with presets.")
        vd.setWordWrap(True)
        vd.setStyleSheet("color:#9aa0a9;")
        vl.addWidget(vd)
        vrow = QHBoxLayout()
        vrow.setSpacing(10)
        vb1 = QPushButton("Open Video Editor")
        vb1.setObjectName("PrimaryButton")
        vb1.setIcon(self.icons.icon("edit"))
        vb1.setIconSize(QSize(18, 18))
        vb1.clicked.connect(lambda: self._launch_legacy_editor("video"))
        vrow.addWidget(vb1)
        vb2 = QPushButton("Import Video / Images…")
        vb2.setIcon(self.icons.icon("add"))
        vb2.setIconSize(QSize(18, 18))
        vb2.clicked.connect(self._me_import_video_or_images)
        vrow.addWidget(vb2)
        vrow.addStretch(1)
        vl.addLayout(vrow)

        audio = QFrame()
        audio.setObjectName("PageCard")
        al = QVBoxLayout(audio)
        al.setContentsMargins(16, 14, 16, 14)
        al.setSpacing(10)
        al.addWidget(QLabel("Audio Editor"))
        ad = QLabel("Waveform editing: select range, cut, normalize, fade in/out, export presets.")
        ad.setWordWrap(True)
        ad.setStyleSheet("color:#9aa0a9;")
        al.addWidget(ad)
        arow = QHBoxLayout()
        arow.setSpacing(10)
        ab1 = QPushButton("Open Audio Editor")
        ab1.setObjectName("PrimaryButton")
        ab1.setIcon(self.icons.icon("edit"))
        ab1.setIconSize(QSize(18, 18))
        ab1.clicked.connect(lambda: self._launch_legacy_editor("audio"))
        arow.addWidget(ab1)
        ab2 = QPushButton("Import Audio…")
        ab2.setIcon(self.icons.icon("add"))
        ab2.setIconSize(QSize(18, 18))
        ab2.clicked.connect(self._me_import_audio)
        arow.addWidget(ab2)
        arow.addStretch(1)
        al.addLayout(arow)

        drop = _QtDropZone("Drop an audio/video/image file here to open the right editor automatically.")
        drop.files_dropped.connect(self._me_on_files_dropped)

        body_layout.addWidget(video, 0, 0)
        body_layout.addWidget(audio, 0, 1)
        body_layout.addWidget(drop, 1, 0, 1, 2)

        layout.addWidget(body)
        layout.addStretch(1)
        return host

    def _me_open_any_file(self):
        f, _ = QFileDialog.getOpenFileName(self, "Open a media file", filter="Media files (*.*)")
        if not f:
            return
        self._me_open_path(Path(f))

    def _me_import_video_or_images(self):
        f, _ = QFileDialog.getOpenFileName(
            self,
            "Import video/images",
            filter="Video/Images (*.mp4 *.mkv *.mov *.webm *.avi *.m4v *.png *.jpg *.jpeg *.webp *.bmp *.gif *.tif *.tiff);;All files (*.*)",
        )
        if not f:
            return
        self._launch_legacy_editor("video", initial_file=Path(f))

    def _me_import_audio(self):
        f, _ = QFileDialog.getOpenFileName(
            self,
            "Import audio",
            filter="Audio (*.mp3 *.wav *.flac *.m4a *.aac *.ogg *.opus);;All files (*.*)",
        )
        if not f:
            return
        self._launch_legacy_editor("audio", initial_file=Path(f))

    def _me_on_files_dropped(self, paths: list):
        if not paths:
            return
        self._me_open_path(Path(paths[0]))

    def _me_open_path(self, p: Path):
        if self._me_is_audio(p):
            self._launch_legacy_editor("audio", initial_file=p)
        elif self._me_is_video(p) or self._me_is_image(p):
            self._launch_legacy_editor("video", initial_file=p)
        else:
            QMessageBox.information(self, "Media Editors", "Unsupported file type.")

    @staticmethod
    def _me_is_audio(p: Path) -> bool:
        return p.suffix.lower() in {".mp3", ".wav", ".flac", ".m4a", ".aac", ".ogg", ".opus"}

    @staticmethod
    def _me_is_video(p: Path) -> bool:
        return p.suffix.lower() in {".mp4", ".mkv", ".mov", ".webm", ".avi", ".m4v"}

    @staticmethod
    def _me_is_image(p: Path) -> bool:
        return p.suffix.lower() in {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".gif", ".tif", ".tiff"}

    def _launch_legacy_editor(self, which: str, *, initial_file: Path | None = None):
        try:
            # Audio editor is now Qt-based; open it in-process for better UX and to share
            # the already-loaded AI manager/settings with the main app.
            if which == "audio":
                from gui.audio_editor_dialog import AudioEditorDialog

                w = AudioEditorDialog(
                    parent=None,
                    ai_manager=getattr(self.backend, "ai_manager", None),
                    initial_file=str(initial_file) if initial_file else None,
                )
                if not hasattr(self, "_aux_windows"):
                    self._aux_windows = []
                self._aux_windows.append(w)
                w.show()
                w.raise_()
                w.activateWindow()
                return
            if which == "video":
                from gui.video_editor_dialog import VideoEditorDialog

                w = VideoEditorDialog(
                    parent=None,
                    ai_manager=getattr(self.backend, "ai_manager", None),
                    initial_file=initial_file if initial_file else None,
                )
                if not hasattr(self, "_aux_windows"):
                    self._aux_windows = []
                self._aux_windows.append(w)
                w.show()
                w.raise_()
                w.activateWindow()
                return

            import subprocess, sys
            import json
            root = Path(__file__).resolve().parents[1]
            p = str(initial_file) if initial_file else ""
            p_lit = json.dumps(p)
            code = (
                "import tkinter as tk; from core.settings_manager import SettingsManager; "
                "from core.ai_manager import AIManager; from gui.media_editors_dialog import MediaEditorsDialog; "
                "root=tk.Tk(); root.withdraw(); s=SettingsManager(); ai=AIManager(s.app_folder, s); "
                "MediaEditorsDialog(root, ai); root.mainloop()"
            )
            subprocess.Popen([sys.executable, "-c", code], cwd=str(root))
        except Exception as e:
            QMessageBox.critical(self, "Launch Editor", str(e))

    def _reload_tasks(self):
        lay: QVBoxLayout = getattr(self, "_qt_tasks_list_layout", None)
        if lay is None:
            return
        t = _ui_theme_tokens()
        while lay.count():
            item = lay.takeAt(0)
            w = item.widget()
            if w:
                w.deleteLater()
        try:
            tasks = self.backend.monitor_manager.scheduled_tasks.list_tasks()
        except Exception:
            tasks = []
        if not tasks:
            empty = QLabel("No scheduled tasks yet.")
            empty.setStyleSheet(f"color:{t['muted']};")
            lay.addWidget(empty)
            lay.addStretch(1)
            return
        for t in tasks:
            card = _QtScheduledTaskCard(
                task=t,
                icons=self.icons,
                on_edit=lambda task=t: self._edit_scheduled_task(task),
                on_delete=lambda task_id=t.task_id: self._delete_scheduled_task(task_id),
                on_run=lambda task_id=t.task_id: self._run_scheduled_task(task_id),
                on_toggle=lambda enabled, task=t: self._toggle_scheduled_task(task, enabled),
            )
            lay.addWidget(card)
        lay.addStretch(1)

    def _add_scheduled_task(self):
        dlg = _EditScheduledTaskDialog(self)
        if dlg.exec() != QDialog.Accepted:
            return
        try:
            ok = self.backend.monitor_manager.add_scheduled_task(dlg.task_dict())
        except Exception as e:
            QMessageBox.critical(self, "Save Failed", str(e))
            return
        if not ok:
            QMessageBox.critical(self, "Save Failed", "Could not save task.")
            return
        self._reload_tasks()

    def _add_safe_temp_cleanup_task(self):
        try:
            import tempfile

            temp_dir = str(Path(tempfile.gettempdir()))
        except Exception:
            temp_dir = str(Path.home() / "AppData" / "Local" / "Temp")

        try:
            tasks = list(self.backend.monitor_manager.scheduled_tasks.list_tasks() or [])
        except Exception:
            tasks = []
        target_norm = str(Path(temp_dir)).lower().replace("/", "\\")
        for task in tasks:
            try:
                same_target = str(getattr(task, "target_path", "") or "").lower().replace("/", "\\") == target_norm
                same_action = str(getattr(task, "action_type", "") or "").strip().lower() == "clean_folder"
                if same_target and same_action:
                    QMessageBox.information(self, "Safe Temp Cleanup", "A Temp cleanup task already exists.")
                    return
            except Exception:
                continue

        message = (
            f"Create a safe daily cleanup for:\n{temp_dir}\n\n"
            "It will run at 3:00 AM, use Recycle Bin/app trash when available, skip active download files, "
            "and only remove files older than 7 days."
        )
        if QMessageBox.question(self, "Safe Temp Cleanup", message) != QMessageBox.Yes:
            return

        task = {
            "title": "Safe Temp Cleanup",
            "schedule": {"type": "daily", "time": "3:00 AM"},
            "action_type": "clean_folder",
            "action_params": {
                "include_subfolders": True,
                "use_recycle_bin": True,
                "skip_active_downloads": True,
                "min_age_seconds": 604800,
            },
            "target_path": temp_dir,
            "enabled": True,
        }
        try:
            ok = bool(self.backend.monitor_manager.add_scheduled_task(task))
        except Exception as e:
            QMessageBox.critical(self, "Save Failed", str(e))
            return
        if not ok:
            QMessageBox.critical(self, "Save Failed", "Could not save the Safe Temp Cleanup task.")
            return
        try:
            self.backend.monitor_manager.scheduled_tasks.reload()
            self.backend.monitor_manager.start_scheduled_tasks()
        except Exception:
            pass
        self._reload_tasks()

    def _edit_scheduled_task(self, task):
        dlg = _EditScheduledTaskDialog(self, task=task)
        if dlg.exec() != QDialog.Accepted:
            return
        try:
            ok = self.backend.monitor_manager.add_scheduled_task(dlg.task_dict())
        except Exception as e:
            QMessageBox.critical(self, "Save Failed", str(e))
            return
        if not ok:
            QMessageBox.critical(self, "Save Failed", "Could not save task.")
            return
        try:
            self.backend.monitor_manager.scheduled_tasks.reload()
        except Exception:
            pass
        self._reload_tasks()

    def _delete_scheduled_task(self, task_id: str):
        if QMessageBox.question(self, "Delete Task", "Delete this scheduled task?") != QMessageBox.Yes:
            return
        try:
            self.backend.monitor_manager.scheduled_tasks.delete_task(task_id)
        except Exception as e:
            QMessageBox.critical(self, "Delete Failed", str(e))
            return
        self._reload_tasks()

    def _run_scheduled_task(self, task_id: str):
        try:
            ok = bool(self.backend.monitor_manager.scheduled_tasks.run_now(task_id))
        except Exception as e:
            QMessageBox.critical(self, "Run Failed", str(e))
            return
        QMessageBox.information(self, "Task Run", f"Completed. ok={ok}")
        self._reload_tasks()

    def _toggle_scheduled_task(self, task, enabled: bool):
        try:
            task.enabled = bool(enabled)
            self.backend.monitor_manager.scheduled_tasks.upsert_task(task)
        except Exception:
            pass
        self._reload_tasks()

    def _build_ai_rules_page(self) -> QWidget:
        host = QFrame()
        host.setObjectName("PageHost")
        layout = QVBoxLayout(host)
        layout.setContentsMargins(18, 18, 18, 18)
        layout.setSpacing(14)

        header = QFrame()
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(0, 0, 0, 0)
        header_layout.setSpacing(10)

        title = QLabel("AI Rules")
        title.setObjectName("PageTitle")
        header_layout.addWidget(title)
        sub = QLabel("Natural-language rules and scheduled tasks.")
        sub.setObjectName("PageSubTitle")
        header_layout.addWidget(sub, 1)

        build_btn = QPushButton("Build Rule")
        build_btn.setObjectName("PrimaryButton")
        build_btn.setIcon(self.icons.icon("brain"))
        build_btn.setIconSize(QSize(18, 18))
        build_btn.clicked.connect(self._toggle_ai_rule_builder_panel)
        header_layout.addWidget(build_btn)

        tasks_btn = QPushButton("Scheduled Tasks")
        tasks_btn.setIcon(self.icons.icon("analytics"))
        tasks_btn.setIconSize(QSize(18, 18))
        tasks_btn.clicked.connect(lambda: self.set_active_page("scheduled_tasks"))
        header_layout.addWidget(tasks_btn)

        layout.addWidget(header)

        # Embedded (non-dialog) AI Rule Builder
        self._ai_rules_builder_panel = _QtAiRuleBuilderPanel(
            backend=self.backend,
            icons=self.icons,
            ensure_ai_ready_cb=lambda **kw: self._ensure_ai_ready(**kw),
        )
        self._ai_rules_builder_panel.setVisible(False)
        self._ai_rules_builder_panel.rule_added.connect(self._on_ai_rule_added_from_panel)
        try:
            self._ai_rules_builder_panel.btn_hide.clicked.connect(lambda: self._ai_rules_builder_panel.setVisible(False))
            self._ai_rules_builder_panel.btn_popout.clicked.connect(self._open_ai_rule_builder)
        except Exception:
            pass
        layout.addWidget(self._ai_rules_builder_panel)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QFrame.NoFrame)

        list_host = QFrame()
        list_layout = QVBoxLayout(list_host)
        list_layout.setContentsMargins(0, 0, 0, 0)
        list_layout.setSpacing(10)
        scroll.setWidget(list_host)
        layout.addWidget(scroll, 1)

        self._qt_ai_rules_list_host = list_host
        self._qt_ai_rules_list_layout = list_layout
        self._reload_ai_rules()
        return host

    def _toggle_ai_rule_builder_panel(self) -> None:
        panel = getattr(self, "_ai_rules_builder_panel", None)
        if not panel:
            self._open_ai_rule_builder()
            return
        new_vis = not panel.isVisible()
        panel.setVisible(new_vis)
        if new_vis:
            try:
                panel._populate_monitors()  # type: ignore[attr-defined]
                panel.input_text.setFocus()  # type: ignore[attr-defined]
            except Exception:
                pass

    def _on_ai_rule_added_from_panel(self) -> None:
        try:
            self._reload_ai_rules()
        finally:
            try:
                panel = getattr(self, "_ai_rules_builder_panel", None)
                if panel:
                    panel._populate_monitors()  # type: ignore[attr-defined]
            except Exception:
                pass

    def _reload_ai_rules(self) -> None:
        host = getattr(self, "_qt_ai_rules_list_host", None)
        layout: QVBoxLayout | None = getattr(self, "_qt_ai_rules_list_layout", None)
        if not host or not layout:
            return
        theme = _ui_theme_tokens()

        while layout.count():
            it = layout.takeAt(0)
            w = it.widget()
            if w:
                w.deleteLater()

        # Scheduled tasks (time-based)
        tasks = []
        try:
            tasks = list(self.backend.monitor_manager.scheduled_tasks.list_tasks() or [])
        except Exception:
            tasks = []

        if tasks:
            hdr = QLabel("⏰ Scheduled Tasks")
            hdr.setStyleSheet(f"color:{theme['text']}; font-weight:700;")
            layout.addWidget(hdr)
            for task_obj in tasks[:10]:
                layout.addWidget(_QtScheduledTaskMiniCard(task=task_obj, icons=self.icons))
            if len(tasks) > 10:
                more = QLabel(f"+ {len(tasks) - 10} more…")
                more.setStyleSheet(f"color:{theme['muted']};")
                layout.addWidget(more)

        # AI rules (event-based)
        rules: list[tuple[str, dict]] = []
        try:
            for _mid, mon in self.backend.monitor_manager.monitors.items():
                for r in getattr(mon, "rules", []) or []:
                    if _is_ai_rule(r):
                        rules.append((getattr(mon, "path", ""), r))
        except Exception:
            rules = []

        if rules:
            hdr = QLabel("🤖 AI Rules (Event‑Based)")
            hdr.setStyleSheet(f"color:{theme['text']}; font-weight:700;")
            layout.addWidget(hdr)
            for path, rule in rules:
                layout.addWidget(_QtAiRuleCard(path=path, rule=rule))

        if not tasks and not rules:
            empty = QLabel("No AI rules or scheduled tasks yet. Click “Build Rule” to create one.")
            empty.setStyleSheet(f"color:{theme['muted']};")
            layout.addWidget(empty)

        layout.addStretch(1)

    def _open_ai_rule_builder(self) -> None:
        dlg = _QtAiRuleBuilderDialog(self, backend=self.backend, icons=self.icons)
        if dlg.exec() == QDialog.Accepted:
            self._reload_ai_rules()

    def _build_ai_hub_page(self) -> QWidget:
        host = QFrame()
        host.setObjectName("PageHost")
        root = QVBoxLayout(host)
        root.setContentsMargins(18, 18, 18, 18)
        root.setSpacing(14)

        header = QFrame()
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(0, 0, 0, 0)
        header_layout.setSpacing(12)

        icon = QLabel()
        icon.setPixmap(self.icons.pixmap("ai_hub", 28))
        icon.setFixedSize(28, 28)
        header_layout.addWidget(icon)

        title_col = QVBoxLayout()
        title_col.setContentsMargins(0, 0, 0, 0)
        title_col.setSpacing(2)

        title = QLabel("Fylorra AI Hub")
        title.setObjectName("PageTitle")
        sub = QLabel("Unified AI Operations Center • Process folders with multiple AI features")
        sub.setObjectName("PageSubTitle")
        title_col.addWidget(title)
        title_col.addWidget(sub)
        header_layout.addLayout(title_col, 1)

        root.addWidget(header)

        top_row = QFrame()
        top_layout = QHBoxLayout(top_row)
        top_layout.setContentsMargins(0, 0, 0, 0)
        top_layout.setSpacing(12)

        # Target folder
        target_box = QGroupBox("Target Folder")
        target_layout = QVBoxLayout(target_box)
        target_layout.setContentsMargins(12, 12, 12, 12)
        target_layout.setSpacing(10)

        tf_row = QHBoxLayout()
        self._ai_hub_target = QLineEdit()
        self._ai_hub_target.setPlaceholderText("Select a folder to process…")
        tf_row.addWidget(self._ai_hub_target, 1)
        browse = QPushButton("Browse")
        browse.setObjectName("PrimaryButton")
        browse.setIcon(self.icons.icon("folder"))
        browse.setIconSize(QSize(18, 18))
        browse.clicked.connect(self._ai_hub_browse_folder)
        tf_row.addWidget(browse)
        target_layout.addLayout(tf_row)

        self._ai_hub_recursive = QCheckBox("Include subfolders (recursive)")
        self._ai_hub_recursive.setChecked(True)
        target_layout.addWidget(self._ai_hub_recursive)

        top_layout.addWidget(target_box, 1)

        # Filters
        filter_box = QGroupBox("File Filters")
        filter_layout = QVBoxLayout(filter_box)
        filter_layout.setContentsMargins(12, 12, 12, 12)
        filter_layout.setSpacing(8)

        self._ai_hub_filter_group = QButtonGroup(host)
        grid = QGridLayout()
        grid.setHorizontalSpacing(14)
        grid.setVerticalSpacing(8)

        def _add_radio(text: str, key: str, r: int, c: int, checked: bool = False):
            rb = QRadioButton(text)
            if checked:
                rb.setChecked(True)
            self._ai_hub_filter_group.addButton(rb)
            rb.setProperty("filter_key", key)
            grid.addWidget(rb, r, c)
            return rb

        _add_radio("All Files", "all", 0, 0, checked=True)
        _add_radio("Images Only", "images", 0, 1)
        _add_radio("Videos Only", "videos", 0, 2)
        _add_radio("Documents Only", "documents", 1, 0)
        _add_radio("Code Files Only", "code", 1, 1)

        filter_layout.addLayout(grid)
        top_layout.addWidget(filter_box, 1)

        root.addWidget(top_row)

        ops_outer = QFrame()
        ops_outer.setObjectName("PageCard")
        ops_outer_layout = QVBoxLayout(ops_outer)
        ops_outer_layout.setContentsMargins(14, 12, 14, 14)
        ops_outer_layout.setSpacing(10)

        ops_title_row = QHBoxLayout()
        ops_title = QLabel("AI Operations")
        self._ai_hub_ops_title = ops_title
        ops_sub = QLabel("Select one or more options")
        self._ai_hub_ops_sub = ops_sub
        ops_title_row.addWidget(ops_title)
        ops_title_row.addSpacing(10)
        ops_title_row.addWidget(ops_sub)
        ops_title_row.addStretch(1)
        ops_outer_layout.addLayout(ops_title_row)

        opts_row = QHBoxLayout()
        opts_row.setSpacing(14)
        self._ai_hub_apply = QCheckBox("Apply changes (rename/move files)")
        self._ai_hub_apply.setChecked(False)
        self._ai_hub_include_other = QCheckBox("Include 'Other' (move unknown formats)")
        self._ai_hub_include_other.setChecked(True)
        self._ai_hub_use_vision = QCheckBox("Use AI vision (images)")
        self._ai_hub_use_vision.setChecked(False)
        self._ai_hub_use_ai_docs = QCheckBox("Use AI for documents (PDF/DOCX)")
        self._ai_hub_use_ai_docs.setChecked(False)
        opts_row.addWidget(self._ai_hub_apply)
        opts_row.addWidget(self._ai_hub_include_other)
        opts_row.addWidget(self._ai_hub_use_vision)
        opts_row.addWidget(self._ai_hub_use_ai_docs)
        opts_row.addStretch(1)
        opts_host = QWidget()
        opts_host.setLayout(opts_row)
        ops_outer_layout.addWidget(opts_host)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QFrame.NoFrame)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        scroll.setVerticalScrollBarPolicy(Qt.ScrollBarAlwaysOff)  # scroll with wheel, no bar

        grid_host = QFrame()
        grid_layout = QGridLayout(grid_host)
        grid_layout.setContentsMargins(0, 0, 0, 0)
        grid_layout.setHorizontalSpacing(12)
        grid_layout.setVerticalSpacing(12)

        self._ai_hub_ops: dict[str, _QtAiOpCard] = {}
        cards = [
            ("smart_rename", "Smart Rename", "Rename files using AI vision analysis and smart rules", "edit", True),
            ("auto_categorize", "Auto-Categorize", "Organize files into 51 comprehensive categories", "grid", True),
            ("duplicate_detection", "Duplicate Detection", "Find exact duplicate files (content hash) and generate a report", "search", True),
            ("content_analysis", "Content Analysis", "Analyze and classify content with AI understanding", "analytics", True),
            ("security_scan", "Security Scan", "Scan folders for suspicious files and risks", "shield", True),
        ]
        for idx, (key, name, desc, icon_name, enabled) in enumerate(cards):
            card = _QtAiOpCard(
                key=key,
                title=name,
                description=desc,
                icon=self.icons.icon(icon_name),
                enabled=enabled,
            )
            self._ai_hub_ops[key] = card
            grid_layout.addWidget(card, idx // 3, idx % 3)
            try:
                card.check.stateChanged.connect(lambda v, k=key: self._ai_hub_on_op_toggled(k, bool(v)))
            except Exception:
                pass

        scroll.setWidget(grid_host)
        ops_outer_layout.addWidget(scroll, 1)
        root.addWidget(ops_outer, 1)
        self._refresh_ai_hub_theme()

        bottom = QFrame()
        bottom_layout = QHBoxLayout(bottom)
        bottom_layout.setContentsMargins(0, 0, 0, 0)
        bottom_layout.setSpacing(12)
        bottom_layout.addStretch(1)

        cancel = QPushButton("Cancel")
        cancel.clicked.connect(lambda: self._ai_hub_reset())
        bottom_layout.addWidget(cancel)

        start = QPushButton("Start AI Operations")
        start.setObjectName("PrimaryButton")
        start.setIcon(self.icons.icon("play"))
        start.setIconSize(QSize(18, 18))
        start.clicked.connect(self._ai_hub_start)
        bottom_layout.addWidget(start)

        root.addWidget(bottom)
        return host

    def _build_ai_command_page(self) -> QWidget:
        host = QFrame()
        host.setObjectName("PageHost")
        root = QVBoxLayout(host)
        root.setContentsMargins(18, 18, 18, 18)
        root.setSpacing(14)

        header = QFrame()
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(0, 0, 0, 0)
        header_layout.setSpacing(12)

        title_col = QVBoxLayout()
        title_col.setContentsMargins(0, 0, 0, 0)
        title_col.setSpacing(2)

        title = QLabel("AI Command Center")
        title.setObjectName("PageTitle")
        sub = QLabel("Natural language → workflow plan → run (local)")
        sub.setObjectName("PageSubTitle")
        title_col.addWidget(title)
        title_col.addWidget(sub)
        header_layout.addLayout(title_col, 1)

        help_btn = QPushButton()
        help_btn.setFixedSize(36, 36)
        help_btn.setIcon(self.icons.icon("ai"))
        help_btn.setIconSize(QSize(18, 18))
        help_btn.setToolTip("Examples")
        help_btn.clicked.connect(self._ai_command_show_examples)
        header_layout.addWidget(help_btn)

        root.addWidget(header)

        card = QFrame()
        card.setObjectName("PageCard")
        card_layout = QVBoxLayout(card)
        card_layout.setContentsMargins(14, 12, 14, 14)
        card_layout.setSpacing(12)

        folder_row = QHBoxLayout()
        folder_row.setSpacing(10)
        self._ai_cmd_target = QLineEdit()
        self._ai_cmd_target.setPlaceholderText("Target folder")
        folder_row.addWidget(self._ai_cmd_target, 1)
        browse = QPushButton("Browse")
        browse.setObjectName("PrimaryButton")
        browse.setIcon(self.icons.icon("folder"))
        browse.setIconSize(QSize(18, 18))
        browse.clicked.connect(self._ai_command_browse_target)
        folder_row.addWidget(browse)
        card_layout.addLayout(folder_row)

        self._ai_cmd_text = QTextEdit()
        self._ai_cmd_text.setPlaceholderText("What do you want to do?")
        self._ai_cmd_text.setMinimumHeight(110)
        card_layout.addWidget(self._ai_cmd_text)

        btns = QHBoxLayout()
        btns.setSpacing(10)

        self._ai_cmd_generate = QPushButton("Generate Plan")
        self._ai_cmd_generate.setObjectName("PrimaryButton")
        self._ai_cmd_generate.setIcon(self.icons.icon("brain"))
        self._ai_cmd_generate.setIconSize(QSize(18, 18))
        self._ai_cmd_generate.clicked.connect(self._ai_command_generate_plan)
        btns.addWidget(self._ai_cmd_generate)

        self._ai_cmd_run = QPushButton("Run Plan")
        self._ai_cmd_run.setObjectName("PrimaryButton")
        self._ai_cmd_run.setIcon(self.icons.icon("play"))
        self._ai_cmd_run.setIconSize(QSize(18, 18))
        self._ai_cmd_run.setEnabled(False)
        self._ai_cmd_run.clicked.connect(self._ai_command_run_plan)
        btns.addWidget(self._ai_cmd_run)

        btns.addStretch(1)
        card_layout.addLayout(btns)

        self._ai_cmd_status = QLabel("No plan yet.")
        self._ai_cmd_status.setStyleSheet("color:#9aa0a9;")
        card_layout.addWidget(self._ai_cmd_status)

        self._ai_cmd_plan_view = QTextEdit()
        self._ai_cmd_plan_view.setReadOnly(True)
        self._ai_cmd_plan_view.setMinimumHeight(220)
        card_layout.addWidget(self._ai_cmd_plan_view, 1)

        root.addWidget(card, 1)
        self._ai_cmd_plan = None
        return host

    def _build_writing_assistant_page(self) -> QWidget:
        host = QFrame()
        host.setObjectName("PageHost")
        root = QVBoxLayout(host)
        root.setContentsMargins(18, 18, 18, 18)
        root.setSpacing(16)

        header = QFrame()
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(0, 0, 0, 0)
        header_layout.setSpacing(12)

        title_col = QVBoxLayout()
        title_col.setContentsMargins(0, 0, 0, 0)
        title_col.setSpacing(2)

        title = QLabel("Writing Assistant")
        title.setObjectName("PageTitle")
        sub = QLabel("Offline writing help for school, work, and email.")
        sub.setObjectName("PageSubTitle")
        title_col.addWidget(title)
        title_col.addWidget(sub)
        header_layout.addLayout(title_col, 1)

        root.addWidget(header)

        splitter = QSplitter(Qt.Horizontal)
        splitter.setChildrenCollapsible(True)

        left_col = QFrame()
        left_col.setMinimumWidth(0)
        left_col.setSizePolicy(QSizePolicy.Minimum, QSizePolicy.Expanding)
        left_layout = QVBoxLayout(left_col)
        left_layout.setContentsMargins(0, 0, 0, 0)
        left_layout.setSpacing(12)

        tools_card = QFrame()
        tools_card.setObjectName("PageCard")
        tools_layout = QVBoxLayout(tools_card)
        tools_layout.setContentsMargins(16, 16, 16, 16)
        tools_layout.setSpacing(12)

        tools_title = QLabel("Writing tools")
        self._wa_tools_title = tools_title
        tools_layout.addWidget(tools_title)

        tools_div = QFrame()
        tools_div.setFixedHeight(1)
        self._wa_tools_div = tools_div
        tools_layout.addWidget(tools_div)

        form = QFormLayout()
        form.setContentsMargins(0, 0, 0, 0)
        form.setHorizontalSpacing(12)
        form.setVerticalSpacing(12)

        self._wa_goal = QComboBox()
        self._wa_goal.addItems(
            [
                "Improve writing",
                "Fix grammar",
                "Make concise",
                "More formal",
                "More friendly",
                "Rewrite as email",
                "Summarize",
                "Bullet points",
            ]
        )
        form.addRow("Goal", self._wa_goal)

        self._wa_tone = QComboBox()
        self._wa_tone.addItems(["Neutral", "Academic", "Professional", "Friendly", "Persuasive"])
        form.addRow("Tone", self._wa_tone)

        self._wa_length = QComboBox()
        self._wa_length.addItems(["Shorter", "Same length", "Longer"])
        form.addRow("Length", self._wa_length)

        self._wa_include_notes = QCheckBox("Include brief suggestions")
        self._wa_include_notes.setChecked(False)
        form.addRow("Notes", self._wa_include_notes)

        self._wa_model_pref = QComboBox()
        self._wa_model_pref.addItem("Auto (recommended)", "auto")
        self._wa_model_pref.addItem("Text model", "text")
        self._wa_model_pref.addItem("Vision model", "vision")
        cur_pref = "text"
        try:
            settings = getattr(self.backend, "settings_manager", None)
            if settings:
                cur_pref = str(settings.get_setting("writing_assistant_model_preference", "text") or "text").strip().lower()
        except Exception:
            cur_pref = "text"
        if cur_pref not in {"auto", "text", "vision"}:
            cur_pref = "text"
        idx = self._wa_model_pref.findData(cur_pref)
        if idx >= 0:
            self._wa_model_pref.setCurrentIndex(idx)
        self._wa_model_pref.setToolTip("Choose which AI model to use for Writing Assistant.")
        self._wa_model_pref.currentIndexChanged.connect(lambda _v: self._wa_save_model_preference())
        form.addRow("AI model", self._wa_model_pref)

        tools_layout.addLayout(form)

        tip = QLabel("Tip: Paste text and ask for edits, or describe what you need.")
        self._wa_tip_label = tip
        tip.setWordWrap(True)
        tools_layout.addWidget(tip)

        left_layout.addWidget(tools_card, 0)

        history_card = QFrame()
        history_card.setObjectName("PageCard")
        history_layout = QVBoxLayout(history_card)
        history_layout.setContentsMargins(16, 16, 16, 16)
        history_layout.setSpacing(10)

        history_header = QHBoxLayout()
        history_header.setSpacing(8)
        hist_title = QLabel("Conversation history")
        self._wa_hist_title = hist_title
        history_header.addWidget(hist_title)
        history_header.addStretch(1)
        self._wa_history_new_btn = QToolButton()
        self._wa_history_new_btn.setIcon(self.icons.icon("add"))
        self._wa_history_new_btn.setIconSize(QSize(14, 14))
        self._wa_history_new_btn.setToolTip("New conversation")
        self._wa_history_new_btn.clicked.connect(self._wa_new_conversation)
        history_header.addWidget(self._wa_history_new_btn)
        self._wa_history_delete_btn = QToolButton()
        self._wa_history_delete_btn.setIcon(self.icons.icon("delete"))
        self._wa_history_delete_btn.setIconSize(QSize(16, 16))
        self._wa_history_delete_btn.setToolTip("Delete selected conversation")
        self._wa_history_delete_btn.clicked.connect(self._wa_delete_selected_history)
        history_header.addWidget(self._wa_history_delete_btn)
        history_layout.addLayout(history_header)

        hist_div = QFrame()
        hist_div.setFixedHeight(1)
        self._wa_hist_div = hist_div
        history_layout.addWidget(hist_div)

        self._wa_history_list = QListWidget()
        self._wa_history_list.setSelectionMode(QAbstractItemView.SingleSelection)
        self._wa_history_list.setSelectionBehavior(QAbstractItemView.SelectRows)
        self._wa_history_list.setFocusPolicy(Qt.StrongFocus)
        self._wa_history_list.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        self._wa_history_list.itemClicked.connect(self._wa_open_history_item)
        history_layout.addWidget(self._wa_history_list, 1)

        left_layout.addWidget(history_card, 1)

        splitter.addWidget(left_col)

        chat = QFrame()
        chat.setObjectName("PageCard")
        chat_layout = QVBoxLayout(chat)
        chat_layout.setContentsMargins(16, 16, 16, 16)
        chat_layout.setSpacing(12)

        chat_header = QHBoxLayout()
        chat_header.setSpacing(10)
        chat_title = QLabel("Conversation")
        self._wa_chat_title = chat_title
        chat_header.addWidget(chat_title)
        chat_header.addStretch(1)
        self._wa_status = QLabel("Ready.")
        chat_header.addWidget(self._wa_status)
        chat_layout.addLayout(chat_header)

        self._wa_chat_area = QScrollArea()
        self._wa_chat_area.setObjectName("WritingAssistantChat")
        self._wa_chat_area.setFrameShape(QFrame.NoFrame)
        self._wa_chat_area.setWidgetResizable(True)
        self._wa_chat_area.setMinimumHeight(240)

        self._wa_chat_view = QWidget()
        self._wa_chat_view.setObjectName("WritingAssistantChatView")
        self._wa_chat_layout = QVBoxLayout(self._wa_chat_view)
        self._wa_chat_layout.setContentsMargins(10, 10, 10, 10)
        self._wa_chat_layout.setSpacing(12)
        self._wa_chat_layout.addStretch(1)
        self._wa_chat_rows = []
        icons_dir = Path(__file__).resolve().parents[1] / "assets" / "icons"
        self._wa_user_icon = QPixmap(str(icons_dir / "user.png"))
        self._wa_bot_icon = QPixmap(str(icons_dir / "bot_assistant.png"))
        self._wa_chat_area.setWidget(self._wa_chat_view)
        self._wa_chat_area.viewport().installEventFilter(self)

        composer = QFrame()
        self._wa_composer = composer
        composer_layout = QVBoxLayout(composer)
        composer_layout.setContentsMargins(12, 12, 12, 12)
        composer_layout.setSpacing(10)

        self._wa_input = QTextEdit()
        self._wa_input.setPlaceholderText("Ask for edits or paste your text here...")
        self._wa_input.setMinimumHeight(90)
        self._wa_input.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        composer_layout.addWidget(self._wa_input)

        btn_row = QHBoxLayout()
        btn_row.setSpacing(10)

        self._wa_send_btn = QPushButton("Send")
        self._wa_send_btn.setObjectName("PrimaryButton")
        self._wa_send_btn.setIcon(self.icons.icon("ai"))
        self._wa_send_btn.setIconSize(QSize(18, 18))
        self._wa_send_btn.clicked.connect(self._wa_send_message)
        btn_row.addWidget(self._wa_send_btn)

        self._wa_copy_btn = QPushButton("Copy last")
        self._wa_copy_btn.clicked.connect(self._wa_copy_last)
        btn_row.addWidget(self._wa_copy_btn)

        self._wa_clear_btn = QPushButton("Clear chat")
        self._wa_clear_btn.clicked.connect(self._wa_clear_chat)
        btn_row.addWidget(self._wa_clear_btn)

        btn_row.addStretch(1)
        composer_layout.addLayout(btn_row)
        self._wa_chat_splitter = QSplitter(Qt.Vertical)
        self._wa_chat_splitter.setHandleWidth(6)
        self._wa_chat_splitter.addWidget(self._wa_chat_area)
        self._wa_chat_splitter.addWidget(composer)
        self._wa_chat_splitter.setStretchFactor(0, 3)
        self._wa_chat_splitter.setStretchFactor(1, 1)
        self._wa_chat_splitter.setSizes([520, 200])
        chat_layout.addWidget(self._wa_chat_splitter, 1)

        splitter.addWidget(chat)
        splitter.setSizes([300, 820])
        root.addWidget(splitter, 1)

        self._wa_sessions: list[dict] = []
        self._wa_current_session_id: str | None = None
        self._wa_messages: list[dict] = []
        self._wa_last_response = ""
        self._wa_busy = False
        self._wa_worker = None
        self._wa_thread = None
        self._wa_load_history()
        if self._wa_sessions:
            self._wa_open_history_session(self._wa_sessions[0].get("id"))
        else:
            self._wa_refresh_history_list()

        self._wa_apply_theme_styles(refresh_chat=False)

        return host

    def _wa_apply_theme_styles(self, *, refresh_chat: bool = True) -> None:
        t = _ui_theme_tokens()
        mode = _ui_theme_mode()
        if mode == "light":
            panel_bg = "#f8fbff"
            panel_border = "#cfd7e4"
            input_bg = "#ffffff"
            input_border = "#c5cedb"
            selected_bg = "#dbeafe"
            selected_border = "#93c5fd"
        elif mode == "black":
            panel_bg = "#101010"
            panel_border = "#1f1f1f"
            input_bg = "#090909"
            input_border = "#222222"
            selected_bg = "#1a2230"
            selected_border = "#2a3a52"
        else:
            panel_bg = "#101318"
            panel_border = "#232730"
            input_bg = "#101318"
            input_border = "#2a303a"
            selected_bg = "#1f3a5c"
            selected_border = "#2b445f"
        try:
            if hasattr(self, "_wa_tools_title"):
                self._wa_tools_title.setStyleSheet(f"color:{t['text']}; font-weight:700; font-size:13px;")
            if hasattr(self, "_wa_hist_title"):
                self._wa_hist_title.setStyleSheet(f"color:{t['text']}; font-weight:700; font-size:13px;")
            if hasattr(self, "_wa_chat_title"):
                self._wa_chat_title.setStyleSheet(f"color:{t['text']}; font-weight:700; font-size:13px;")
            if hasattr(self, "_wa_tip_label"):
                self._wa_tip_label.setStyleSheet(f"color:{t['muted']};")
            if hasattr(self, "_wa_tools_div"):
                self._wa_tools_div.setStyleSheet(f"background:{panel_border};")
            if hasattr(self, "_wa_hist_div"):
                self._wa_hist_div.setStyleSheet(f"background:{panel_border};")
        except Exception:
            pass
        try:
            if hasattr(self, "_wa_history_new_btn"):
                self._wa_history_new_btn.setStyleSheet(
                    "QToolButton{"
                    f"background:{input_bg}; border:1px solid {input_border}; border-radius:8px; padding:4px; color:{t['text']};"
                    "}"
                    f"QToolButton:hover{{background:{panel_bg};}}"
                )
            if hasattr(self, "_wa_history_delete_btn"):
                self._wa_history_delete_btn.setStyleSheet(
                    "QToolButton{"
                    f"background:{input_bg}; border:1px solid {input_border}; border-radius:8px; padding:4px; color:{t['text']};"
                    "}"
                    f"QToolButton:hover{{background:{panel_bg};}}"
                )
        except Exception:
            pass
        try:
            if hasattr(self, "_wa_history_list"):
                self._wa_history_list.setStyleSheet(
                    "QListWidget{"
                    f"background:{input_bg}; border:1px solid {panel_border}; border-radius:12px; color:{t['text']};"
                    "}"
                    "QListWidget::item{"
                    f"padding:8px 10px; margin:4px 6px; border-radius:8px; border:1px solid {panel_border};"
                    "}"
                    "QListWidget::item:selected{"
                    f"background:{selected_bg}; color:{t['text']}; border-color:{selected_border};"
                    "}"
                )
            if hasattr(self, "_wa_status"):
                self._wa_status.setStyleSheet(
                    "QLabel{"
                    f"background:{panel_bg}; color:{t['muted']}; border:1px solid {input_border}; border-radius:10px; "
                    "padding:4px 10px; font-weight:600;}"
                )
            if hasattr(self, "_wa_chat_area"):
                self._wa_chat_area.setStyleSheet(
                    "QScrollArea#WritingAssistantChat{"
                    f"background:{panel_bg}; border:1px solid {panel_border}; border-radius:12px;"
                    "}"
                    "QScrollArea#WritingAssistantChat > QWidget{background:transparent;}"
                )
            if hasattr(self, "_wa_chat_view"):
                self._wa_chat_view.setStyleSheet("QWidget#WritingAssistantChatView{background:transparent;}")
            if hasattr(self, "_wa_composer"):
                self._wa_composer.setStyleSheet(
                    f"QFrame{{background:{panel_bg}; border:1px solid {panel_border}; border-radius:12px;}}"
                )
            if hasattr(self, "_wa_input"):
                self._wa_input.setStyleSheet(
                    f"QTextEdit{{background:{input_bg}; border:1px solid {input_border}; border-radius:10px; padding:8px; color:{t['text']};}}"
                )
            if hasattr(self, "_wa_chat_splitter"):
                self._wa_chat_splitter.setStyleSheet(
                    f"QSplitter::handle{{background:{input_border}; border-radius:3px;}}"
                )
        except Exception:
            pass
        if not refresh_chat:
            return
        try:
            msgs = list(getattr(self, "_wa_messages", []) or [])
            self._wa_clear_chat_view()
            for msg in msgs:
                self._wa_append_message(str(msg.get("role") or "assistant"), str(msg.get("content") or ""))
        except Exception:
            pass

    def _wa_system_prompt(self) -> str:
        goal = self._wa_goal.currentText() if hasattr(self, "_wa_goal") else "Improve writing"
        tone = self._wa_tone.currentText() if hasattr(self, "_wa_tone") else "Neutral"
        length = self._wa_length.currentText() if hasattr(self, "_wa_length") else "Same length"
        include_notes = bool(getattr(self, "_wa_include_notes", None) and self._wa_include_notes.isChecked())
        prompt = (
            "You are a precise writing assistant for school, work, and office communication. "
            "Rewrite or answer with clear, correct, and natural writing. "
            f"Goal: {goal}. Tone: {tone}. Length: {length}. "
            "Preserve meaning and facts. "
            "If the user asks for a rewrite, return only the improved text unless notes are requested. "
            "If the user asks a question, respond with guidance and an example when helpful."
        )
        goal_key = goal.lower()
        if "email" in goal_key:
            prompt += " Include a subject line, greeting, and closing in the email."
        if "summarize" in goal_key:
            prompt += " Return a concise summary."
        if "bullet" in goal_key:
            prompt += " Return bullet points only."
        if include_notes:
            prompt += " After the final text, add a short Notes section with up to 3 bullets."
        return prompt

    def _wa_save_model_preference(self) -> None:
        settings = getattr(self.backend, "settings_manager", None)
        if not settings or not hasattr(self, "_wa_model_pref"):
            return
        try:
            pref = str(self._wa_model_pref.currentData() or "text").strip().lower()
        except Exception:
            pref = "text"
        if pref not in {"auto", "text", "vision"}:
            pref = "text"
        try:
            settings.set_setting("writing_assistant_model_preference", pref)
        except Exception:
            pass

    def _wa_resolve_model_kind(self) -> str:
        pref = "text"
        try:
            if hasattr(self, "_wa_model_pref"):
                pref = str(self._wa_model_pref.currentData() or "text").strip().lower()
        except Exception:
            pref = "text"
        if pref in {"text", "vision"}:
            return pref
        # Auto: Writing Assistant is text-first by default.
        return "text"

    def _wa_append_message(self, role: str, text: str) -> None:
        if not hasattr(self, "_wa_chat_layout"):
            return
        msg = (text or "").strip()
        if not msg:
            return
        mode = _ui_theme_mode()
        if mode == "light":
            bg = "#dbeafe" if role == "user" else "#edf2f7"
            text_color = "#0f172a"
            meta_color = "#64748b"
        elif mode == "black":
            bg = "#141b24" if role == "user" else "#0f141b"
            text_color = "#f2f2f2"
            meta_color = "#a0a0a0"
        else:
            bg = "#1f2b3a" if role == "user" else "#182028"
            text_color = "#e6e8ee"
            meta_color = "#9aa0a9"
        role_label_text = "You" if role == "user" else "Assistant"

        bubble = _ChatBubble(bg, radius=18)
        bubble.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Minimum)
        bubble_layout = QVBoxLayout(bubble)
        bubble_layout.setContentsMargins(12, 8, 12, 8)
        bubble_layout.setSpacing(4)

        label = QLabel(msg)
        label.setWordWrap(True)
        label.setTextFormat(Qt.PlainText)
        label.setTextInteractionFlags(Qt.TextSelectableByMouse)
        label.setStyleSheet(f"color:{text_color}; font-size:14px;")
        bubble_layout.addWidget(label)

        role_label = QLabel(role_label_text)
        role_label.setStyleSheet(f"color:{meta_color}; font-size:11px; font-weight:600;")
        role_label.setAlignment(Qt.AlignVCenter | Qt.AlignLeft)

        icon_label = QLabel()
        icon_pixmap = None
        if role == "user" and hasattr(self, "_wa_user_icon"):
            icon_pixmap = self._wa_user_icon
        elif role != "user" and hasattr(self, "_wa_bot_icon"):
            icon_pixmap = self._wa_bot_icon
        if isinstance(icon_pixmap, QPixmap) and not icon_pixmap.isNull():
            icon_label.setPixmap(icon_pixmap.scaled(16, 16, Qt.KeepAspectRatio, Qt.SmoothTransformation))
        icon_label.setFixedSize(18, 18)
        icon_label.setAlignment(Qt.AlignVCenter | Qt.AlignLeft)

        meta = QWidget()
        meta_layout = QHBoxLayout(meta)
        meta_layout.setContentsMargins(0, 0, 0, 0)
        meta_layout.setSpacing(6)
        meta_layout.addWidget(icon_label, 0)
        meta_layout.addWidget(role_label, 0)

        row_container = QWidget()
        row_container.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Minimum)
        row_container_layout = QHBoxLayout(row_container)
        row_container_layout.setContentsMargins(4, 0, 4, 0)
        row_container_layout.setSpacing(6)
        row_container_layout.addWidget(meta, 0, Qt.AlignVCenter | Qt.AlignLeft)
        row_container_layout.addWidget(bubble, 1)

        row_outer = QWidget()
        row_outer_layout = QHBoxLayout(row_outer)
        row_outer_layout.setContentsMargins(8, 0, 8, 0)
        row_outer_layout.setSpacing(0)
        row_outer_layout.addStretch(1)
        row_outer_layout.addWidget(row_container)
        row_outer_layout.addStretch(1)

        if hasattr(self, "_wa_chat_rows"):
            self._wa_chat_rows.append(row_container)
            self._wa_update_chat_widths()

        insert_index = max(0, self._wa_chat_layout.count() - 1)
        self._wa_chat_layout.insertWidget(insert_index, row_outer)
        self._wa_scroll_chat_to_bottom()

    def _wa_clear_chat_view(self) -> None:
        if not hasattr(self, "_wa_chat_layout"):
            return
        while self._wa_chat_layout.count():
            item = self._wa_chat_layout.takeAt(0)
            widget = item.widget()
            if widget is not None:
                widget.deleteLater()
        self._wa_chat_layout.addStretch(1)
        if hasattr(self, "_wa_chat_rows"):
            self._wa_chat_rows = []

    def _wa_scroll_chat_to_bottom(self) -> None:
        if not hasattr(self, "_wa_chat_area"):
            return
        bar = self._wa_chat_area.verticalScrollBar()
        bar.setValue(bar.maximum())

    def _wa_update_chat_widths(self) -> None:
        if not hasattr(self, "_wa_chat_area") or not hasattr(self, "_wa_chat_rows"):
            return
        try:
            viewport_width = int(self._wa_chat_area.viewport().width())
        except Exception:
            return
        if viewport_width <= 0:
            return
        max_width = max(420, int(viewport_width * 0.9))
        max_width = min(max_width, 980)
        for row in list(self._wa_chat_rows):
            try:
                row.setMinimumWidth(max_width)
                row.setMaximumWidth(max_width)
            except Exception:
                continue

    def _wa_trim_history(self, max_messages: int = 12) -> None:
        if len(self._wa_messages) > max_messages:
            del self._wa_messages[:-max_messages]

    def _wa_history_path(self) -> Path:
        base = None
        try:
            base = getattr(self.backend, "settings_manager", None).app_folder  # type: ignore[attr-defined]
        except Exception:
            base = None
        if not base:
            base = Path.home() / ".fylorra"
        return Path(base) / "writing_assistant_history.json"

    def _wa_load_history(self) -> None:
        self._wa_sessions = []
        try:
            import json

            path = self._wa_history_path()
            if not path.exists():
                return
            data = json.loads(path.read_text(encoding="utf-8"))
            if not isinstance(data, list):
                return
            for raw in data:
                if not isinstance(raw, dict):
                    continue
                sid = str(raw.get("id") or f"wa_{uuid.uuid4().hex[:8]}")
                title = str(raw.get("title") or "New chat")
                created_raw = raw.get("created")
                created = QDateTime.currentDateTime()
                if isinstance(created_raw, str) and created_raw:
                    dt = QDateTime.fromString(created_raw, Qt.ISODate)
                    if dt.isValid():
                        created = dt
                messages = []
                for msg in raw.get("messages") or []:
                    if not isinstance(msg, dict):
                        continue
                    role = str(msg.get("role") or "assistant")
                    if role not in {"user", "assistant", "system"}:
                        role = "assistant"
                    content = str(msg.get("content") or "")
                    if content:
                        messages.append({"role": role, "content": content})
                self._wa_sessions.append(
                    {
                        "id": sid,
                        "title": title,
                        "created": created,
                        "messages": messages,
                    }
                )
            self._wa_sessions = self._wa_sessions[:50]
        except Exception:
            self._wa_sessions = []

    def _wa_save_history(self) -> None:
        try:
            import json

            out: list[dict] = []
            for session in self._wa_sessions[:50]:
                created = session.get("created")
                created_str = ""
                if isinstance(created, QDateTime):
                    created_str = created.toString(Qt.ISODate)
                out.append(
                    {
                        "id": session.get("id"),
                        "title": session.get("title"),
                        "created": created_str,
                        "messages": list(session.get("messages") or []),
                    }
                )
            path = self._wa_history_path()
            path.write_text(json.dumps(out, indent=2), encoding="utf-8")
        except Exception:
            pass

    def _wa_get_session(self, session_id: str | None) -> dict | None:
        if not session_id:
            return None
        for session in self._wa_sessions:
            if session.get("id") == session_id:
                return session
        return None

    def _wa_refresh_history_list(self, select_id: str | None = None) -> None:
        if not hasattr(self, "_wa_history_list"):
            return
        self._wa_history_list.blockSignals(True)
        self._wa_history_list.clear()
        for session in self._wa_sessions:
            title = str(session.get("title") or "New chat")
            created = session.get("created")
            stamp = ""
            try:
                if isinstance(created, QDateTime):
                    stamp = created.toString("MMM d · HH:mm")
            except Exception:
                stamp = ""
            label = f"{title}" + (f"  ·  {stamp}" if stamp else "")
            item = QListWidgetItem(label)
            item.setData(Qt.UserRole, session.get("id"))
            self._wa_history_list.addItem(item)
            if select_id and session.get("id") == select_id:
                self._wa_history_list.setCurrentItem(item)
        self._wa_history_list.blockSignals(False)

    def _wa_new_conversation(self) -> None:
        session_id = f"wa_{uuid.uuid4().hex[:8]}"
        session = {
            "id": session_id,
            "title": "New chat",
            "messages": [],
            "created": QDateTime.currentDateTime(),
        }
        self._wa_sessions.insert(0, session)
        self._wa_current_session_id = session_id
        self._wa_messages = session["messages"]
        self._wa_last_response = ""
        try:
            self._wa_clear_chat_view()
        except Exception:
            pass
        self._wa_refresh_history_list(select_id=session_id)
        self._wa_save_history()
        if hasattr(self, "_wa_input"):
            self._wa_input.setFocus()

    def _wa_start_session_if_needed(self) -> None:
        if self._wa_current_session_id:
            return
        session_id = f"wa_{uuid.uuid4().hex[:8]}"
        session = {
            "id": session_id,
            "title": "New chat",
            "messages": [],
            "created": QDateTime.currentDateTime(),
        }
        self._wa_sessions.insert(0, session)
        self._wa_current_session_id = session_id
        self._wa_messages = session["messages"]
        self._wa_refresh_history_list(select_id=session_id)

    def _wa_record_message(self, role: str, text: str) -> None:
        self._wa_start_session_if_needed()
        self._wa_messages.append({"role": role, "content": text})
        if role == "user":
            session = self._wa_get_session(self._wa_current_session_id)
            if session and str(session.get("title") or "").strip() in ("", "New chat"):
                first = (text or "").strip().splitlines()[0] if text else "New chat"
                session["title"] = (first[:48] + "…") if len(first) > 48 else first
                self._wa_refresh_history_list(select_id=session.get("id"))
        session = self._wa_get_session(self._wa_current_session_id)
        if session:
            try:
                self._wa_sessions.remove(session)
            except Exception:
                pass
            self._wa_sessions.insert(0, session)
            self._wa_refresh_history_list(select_id=session.get("id"))
        self._wa_save_history()

    def _wa_open_history_session(self, session_id: str | None) -> None:
        session = self._wa_get_session(str(session_id or ""))
        if not session:
            return
        self._wa_current_session_id = str(session.get("id") or "")
        self._wa_messages = session.get("messages") or []
        try:
            self._wa_clear_chat_view()
        except Exception:
            pass
        for msg in list(self._wa_messages):
            self._wa_append_message(str(msg.get("role") or "assistant"), str(msg.get("content") or ""))
        last = ""
        for msg in reversed(self._wa_messages):
            if str(msg.get("role")) == "assistant":
                last = str(msg.get("content") or "")
                break
        self._wa_last_response = last
        self._wa_refresh_history_list(select_id=self._wa_current_session_id)

    def _wa_open_history_item(self, item: QListWidgetItem) -> None:
        session_id = item.data(Qt.UserRole)
        self._wa_open_history_session(str(session_id or ""))

    def _wa_delete_selected_history(self) -> None:
        if not hasattr(self, "_wa_history_list"):
            return
        item = self._wa_history_list.currentItem()
        if not item:
            return
        session_id = str(item.data(Qt.UserRole) or "")
        if not session_id:
            return
        session = self._wa_get_session(session_id)
        if not session:
            return
        try:
            self._wa_sessions.remove(session)
        except Exception:
            pass
        if self._wa_current_session_id == session_id:
            self._wa_current_session_id = None
            self._wa_messages = []
            try:
                self._wa_clear_chat_view()
            except Exception:
                pass
        self._wa_refresh_history_list()
        self._wa_save_history()

    def _wa_set_busy(self, busy: bool) -> None:
        self._wa_busy = bool(busy)
        if hasattr(self, "_wa_send_btn"):
            self._wa_send_btn.setEnabled(not busy)
        if hasattr(self, "_wa_input"):
            self._wa_input.setEnabled(not busy)
        if hasattr(self, "_wa_status"):
            self._wa_status.setText("Writing..." if busy else "Ready.")

    def _wa_send_message(self) -> None:
        if getattr(self, "_wa_busy", False):
            return
        kind = "text"
        try:
            kind = self._wa_resolve_model_kind()
        except Exception:
            kind = "text"
        if not self._ensure_ai_ready(title="AI Required for Writing Assistant", kind=kind):
            return
        text = (self._wa_input.toPlainText() if hasattr(self, "_wa_input") else "").strip()
        if not text:
            return

        self._wa_append_message("user", text)
        self._wa_record_message("user", text)
        self._wa_trim_history()
        try:
            self._wa_input.clear()
        except Exception:
            pass

        sys_prompt = self._wa_system_prompt()
        messages = [{"role": "system", "content": sys_prompt}] + list(self._wa_messages)

        worker = _QtWritingAssistantWorker(
            ai_manager=getattr(self.backend, "ai_manager", None),
            messages=messages,
            temperature=0.2,
            max_tokens=1200,
        )
        thread = QThread(self)
        worker.moveToThread(thread)
        thread.started.connect(worker.run)
        worker.finished.connect(self._wa_on_response)
        worker.error.connect(self._wa_on_error)
        worker.finished.connect(thread.quit)
        worker.error.connect(thread.quit)
        worker.finished.connect(worker.deleteLater)
        worker.error.connect(worker.deleteLater)
        thread.finished.connect(thread.deleteLater)

        self._wa_worker = worker
        self._wa_thread = thread
        self._wa_set_busy(True)
        thread.start()

    def _wa_on_response(self, text: str) -> None:
        reply = (text or "").strip()
        if reply:
            self._wa_append_message("assistant", reply)
            self._wa_record_message("assistant", reply)
            self._wa_trim_history()
            self._wa_last_response = reply
        self._wa_set_busy(False)

    def _wa_on_error(self, msg: str) -> None:
        err = (msg or "AI request failed.").strip()
        self._wa_append_message("assistant", f"Error: {err}")
        self._wa_set_busy(False)

    def _wa_clear_chat(self) -> None:
        try:
            self._wa_clear_chat_view()
        except Exception:
            pass
        self._wa_messages = []
        self._wa_current_session_id = None
        self._wa_last_response = ""
        try:
            self._wa_history_list.clearSelection()
        except Exception:
            pass
        self._wa_set_busy(False)

    def _wa_copy_last(self) -> None:
        text = str(getattr(self, "_wa_last_response", "") or "")
        if not text:
            return
        try:
            QApplication.clipboard().setText(text)
            if hasattr(self, "_wa_status"):
                self._wa_status.setText("Copied last response.")
        except Exception:
            pass

    def _ai_command_show_examples(self):
        examples = [
            "Convert all .flac in this folder to mp3 320k into MP3 Music",
            "Convert all images to webp into Converted_Images",
            "Index this folder, then find invoices from Duke Energy",
            "Split report.pdf by bookmarks",
            "Cut video.mp4 from 03:47 to 04:10",
            "Convert resume.docx to PDF",
            "Zip folder_rel=MP3/Album as Album.zip",
        ]
        QMessageBox.information(self, "AI Command Examples", "Try:\n- " + "\n- ".join(examples))

    def _ai_command_browse_target(self):
        folder = QFileDialog.getExistingDirectory(self, "Select target folder")
        if folder:
            self._ai_cmd_target.setText(folder)

    def _ai_command_generate_plan(self):
        folder = self._ai_cmd_target.text().strip()
        text = self._ai_cmd_text.toPlainText().strip()
        if not folder or not text:
            QMessageBox.information(self, "AI Command", "Pick a target folder and enter an instruction.")
            return

        # Auto-load AI model on-demand (no extra manual steps).
        try:
            ai = getattr(self.backend, "ai_manager", None)
            has_files = bool(getattr(ai, "model_files_exist", lambda: False)()) if ai else False
            if ai and has_files and not getattr(ai, "is_ready", False):
                _QtAIModelLoadDialog(self, ai_manager=ai).exec()
        except Exception:
            pass

        dlg = _QtAICommandPlanDialog(self, target_folder=folder, instruction=text)
        if dlg.exec() != QDialog.Accepted:
            return
        self._ai_cmd_plan = dlg.plan
        self._ai_cmd_plan_view.setPlainText(dlg.plan_text)
        self._ai_cmd_status.setText("Plan ready.")
        self._ai_cmd_run.setEnabled(True)

    def _ai_command_run_plan(self):
        if not self._ai_cmd_plan:
            return
        folder = self._ai_cmd_target.text().strip()
        dlg = _QtAICommandRunDialog(self, target_folder=folder, plan=self._ai_cmd_plan)
        dlg.exec()

    def _ai_hub_browse_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "Select folder")
        if folder:
            self._ai_hub_target.setText(folder)

    def _ai_hub_reset(self):
        try:
            self._ai_hub_target.setText("")
        except Exception:
            pass
        try:
            self._ai_hub_recursive.setChecked(True)
        except Exception:
            pass
        try:
            for c in self._ai_hub_ops.values():
                c.set_checked(False)
        except Exception:
            pass

    def _ai_hub_start(self):
        folder = self._ai_hub_target.text().strip()
        if not folder:
            QMessageBox.information(self, "AI Hub", "Pick a target folder first.")
            return
        selected = [k for k, c in self._ai_hub_ops.items() if c.is_checked()]
        if not selected:
            QMessageBox.information(self, "AI Hub", "Select at least one AI operation.")
            return
        filter_key = "all"
        try:
            b = self._ai_hub_filter_group.checkedButton()
            if b:
                filter_key = str(b.property("filter_key") or "all")
        except Exception:
            filter_key = "all"

        opts = {
            "include_subfolders": bool(self._ai_hub_recursive.isChecked()),
            "filter_key": filter_key,
            "apply": bool(getattr(self, "_ai_hub_apply", None) and self._ai_hub_apply.isChecked()),
            "include_other": bool(getattr(self, "_ai_hub_include_other", None) and self._ai_hub_include_other.isChecked()),
            "use_vision": bool(getattr(self, "_ai_hub_use_vision", None) and self._ai_hub_use_vision.isChecked()),
            "use_ai_docs": bool(getattr(self, "_ai_hub_use_ai_docs", None) and self._ai_hub_use_ai_docs.isChecked()),
        }
        needs_loaded_ai = bool(opts.get("use_vision")) or bool(opts.get("use_ai_docs")) or any(
            op in {"security_scan", "content_analysis"} for op in selected
        )
        if needs_loaded_ai and not self._ensure_ai_ready(title="Prepare AI Model", kind="vision"):
            return
        # Reduce dialog spam: if a single operation has a dedicated full UI, open it directly.
        if selected == ["auto_categorize"]:
            dlg = _QtAutoCategorizeDialog(
                self,
                backend=self.backend,
                target_folder=folder,
                autorun=True,
                initial_options={
                    "include_subfolders": opts.get("include_subfolders", True),
                    "smart_scope": True,
                    "include_other": opts.get("include_other", False),
                    "use_ai_vision": opts.get("use_vision", False),
                    "use_ai_documents": opts.get("use_ai_docs", False),
                    "move_files": opts.get("apply", False),
                },
            )
            dlg.exec()
            return
        if selected == ["smart_rename"]:
            dlg = _QtSmartRenameDialog(
                self,
                backend=self.backend,
                target_folder=folder,
                autorun=True,
                initial_options={
                    "include_subfolders": opts.get("include_subfolders", True),
                    "filter_key": filter_key,
                },
            )
            dlg.exec()
            return
        if selected == ["content_analysis"]:
            dlg = _QtContentAnalysisDialog(
                self,
                backend=self.backend,
                target_folder=folder,
                autorun=True,
                initial_options={
                    "include_subfolders": opts.get("include_subfolders", True),
                    "max_files": 500,
                },
            )
            dlg.exec()
            return

        dlg = _QtAiHubRunDialog(self, target_folder=folder, operations=selected, options=opts)
        dlg.exec()

    def _ai_hub_on_op_toggled(self, key: str, checked: bool):
        """
        UX: for operations that are primarily about organizing/moving files, default to "Apply changes"
        once the user explicitly selects the operation.
        """
        if not checked:
            return
        try:
            if key in {"auto_categorize", "duplicate_detection"} and hasattr(self, "_ai_hub_apply"):
                if not self._ai_hub_apply.isChecked():
                    self._ai_hub_apply.setChecked(True)
        except Exception:
            pass

    def _add_folder_monitor(self):
        folder = QFileDialog.getExistingDirectory(self, "Select folder to monitor")
        if not folder:
            return
        import uuid

        monitor_id = f"monitor_{uuid.uuid4().hex[:10]}"
        ok = False
        try:
            ok = self.backend.monitor_manager.add_monitor(monitor_id, folder, rules=[])
        except Exception as e:
            QMessageBox.critical(self, "Add Monitor Failed", str(e))
            return
        if not ok:
            msg = str(getattr(self.backend.monitor_manager, "last_error", "") or "Could not add monitor.")
            QMessageBox.critical(self, "Add Monitor Failed", msg)
            return
        self._add_monitor_card(monitor_id, monitor_kind="folder", rules=[])

        # Immediately open Edit dialog so the user can configure events/filters/rules (original behavior).
        try:
            mon = self.backend.monitor_manager.get_monitor(monitor_id)
        except Exception:
            mon = None
        if mon:
            dlg = _EditFolderMonitorDialog(self, monitor=mon, default_running=True)
            if dlg.exec() == QDialog.Accepted:
                try:
                    dlg.apply_to_monitor()
                except Exception as e:
                    QMessageBox.critical(self, "Save Failed", str(e))
                try:
                    if dlg.desired_running():
                        started = bool(self.backend.monitor_manager.start_monitor(monitor_id))
                        if not started:
                            msg = str(getattr(self.backend.monitor_manager, "last_error", "") or "Could not start monitor.")
                            QMessageBox.critical(self, "Start Monitor Failed", msg)
                        self._monitor_cards[monitor_id]._refresh_stats()  # type: ignore[attr-defined]
                except Exception:
                    pass
                try:
                    self.backend.monitor_manager.save_monitors()
                except Exception:
                    pass
            else:
                # User cancelled configuration -> remove the monitor we just created.
                try:
                    self._remove_monitor(monitor_id)
                except Exception:
                    pass
                return

        try:
            self.backend.monitor_manager.save_monitors()
        except Exception:
            pass

    def _add_ftp_monitor(self):
        dlg = _AddFtpMonitorDialog(self)
        if dlg.exec() != QDialog.Accepted:
            return
        values = dlg.values()
        import uuid

        monitor_id = f"ftp_{uuid.uuid4().hex[:10]}"
        try:
            ok = self.backend.monitor_manager.ftp_manager.add_ftp_monitor(
                monitor_id,
                values["host"],
                values["username"],
                values["password"],
                values["remote_path"],
                values["port"],
                values["use_tls"],
                values["poll_interval"],
                self.backend.monitor_manager._on_monitor_event,
                local_sync_dir=values.get("local_sync_dir"),
                download_on_created=bool(values.get("download_on_created", True)),
                download_on_modified=bool(values.get("download_on_modified", True)),
                delete_local_on_deleted=bool(values.get("delete_local_on_deleted", False)),
                overwrite_local=bool(values.get("overwrite_local", False)),
                allowed_extensions=values.get("allowed_extensions") or None,
                passive_mode=bool(values.get("passive_mode", True)),
                tls_implicit=bool(values.get("tls_implicit", False)),
                encoding=str(values.get("encoding", "utf-8") or "utf-8"),
                two_way_sync=bool(values.get("two_way_sync", False)),
                sync_subfolders=bool(values.get("sync_subfolders", False)),
            )
        except Exception as e:
            QMessageBox.critical(self, "Add FTP Monitor Failed", str(e))
            return
        if not ok:
            msg = str(getattr(self.backend.monitor_manager.ftp_manager, "last_error", "") or "Could not add FTP monitor.")
            QMessageBox.critical(self, "Add FTP Monitor Failed", msg)
            return

        self._add_monitor_card(monitor_id, monitor_kind="ftp", rules=[])
        if values.get("auto_start"):
            started = True
            try:
                started = bool(self.backend.monitor_manager.ftp_manager.start_ftp_monitor(monitor_id))
            except Exception:
                started = False
            if not started:
                msg = str(getattr(self.backend.monitor_manager.ftp_manager, "last_error", "") or "Could not start FTP monitor.")
                QMessageBox.critical(self, "Start FTP Monitor Failed", msg)
            try:
                self._monitor_cards[monitor_id]._refresh_stats()  # type: ignore[attr-defined]
            except Exception:
                pass
        try:
            self.backend.monitor_manager.save_monitors()
        except Exception:
            pass

    def _add_monitor_card(self, monitor_id: str, *, monitor_kind: str, rules: list[dict] | None):
        monitor = None
        if monitor_kind == "folder":
            try:
                monitor = self.backend.monitor_manager.get_monitor(monitor_id)
            except Exception:
                monitor = None
        elif monitor_kind == "ftp":
            try:
                monitor = self.backend.monitor_manager.ftp_manager.get_ftp_monitor(monitor_id)
            except Exception:
                monitor = None
        if not monitor:
            return
        try:
            if monitor_kind == "folder":
                rules = list(getattr(monitor, "rules", []) or [])
        except Exception:
            pass

        card = _QtMonitorCard(
            monitor_id=monitor_id,
            monitor_kind=monitor_kind,
            monitor=monitor,
            rules=rules or [],
            icons=self.icons,
            on_start=lambda: self.backend.monitor_manager.start_monitor(monitor_id),
            on_stop=lambda: self.backend.monitor_manager.stop_monitor(monitor_id),
            on_remove=lambda: self._remove_monitor(monitor_id),
            on_start_ftp=lambda: self.backend.monitor_manager.ftp_manager.start_ftp_monitor(monitor_id),
            on_stop_ftp=lambda: self.backend.monitor_manager.ftp_manager.stop_ftp_monitor(monitor_id),
            on_remove_ftp=lambda: self._remove_ftp_monitor(monitor_id),
            on_poll_ftp=lambda: self.backend.monitor_manager.ftp_manager.poll_once(monitor_id),
            on_build_rule=lambda: self._open_rule_builder_for_monitor(monitor_id),
            on_open_ai_hub=lambda op: self._open_ai_hub_for_monitor(monitor_id, op),
        )
        self._monitor_cards[monitor_id] = card
        # Insert before the stretch spacer (last item).
        lay: QVBoxLayout = self._qt_monitors_list_layout
        idx = max(0, lay.count() - 1)
        lay.insertWidget(idx, card)

    def _open_rule_builder_for_monitor(self, monitor_id: str) -> None:
        self.set_active_page("ai_rules")
        panel = getattr(self, "_ai_rules_builder_panel", None)
        if not panel:
            self._open_ai_rule_builder()
            return
        try:
            panel.setVisible(True)
        except Exception:
            pass
        try:
            panel._populate_monitors()  # type: ignore[attr-defined]
        except Exception:
            pass
        try:
            idx = panel.monitor_combo.findData(str(monitor_id))  # type: ignore[attr-defined]
            if idx >= 0:
                panel.monitor_combo.setCurrentIndex(idx)  # type: ignore[attr-defined]
        except Exception:
            pass
        try:
            panel.input_text.setFocus()  # type: ignore[attr-defined]
        except Exception:
            pass

    def _open_ai_hub_for_monitor(self, monitor_id: str, op_key: str) -> None:
        mon = None
        try:
            mon = self.backend.monitor_manager.get_monitor(monitor_id)
        except Exception:
            mon = None
        folder = ""
        try:
            folder = str(getattr(mon, "path", "") or "")
        except Exception:
            folder = ""
        if not folder:
            return
        self.set_active_page("ai_hub")
        try:
            self._ai_hub_target.setText(folder)
        except Exception:
            pass
        try:
            # Reset selection but keep folder
            for c in self._ai_hub_ops.values():
                c.set_checked(False)
        except Exception:
            pass
        try:
            card = self._ai_hub_ops.get(str(op_key))
            if card:
                card.set_checked(True)
        except Exception:
            pass

    def _remove_monitor(self, monitor_id: str):
        try:
            self.backend.monitor_manager.remove_monitor(monitor_id)
            self.backend.monitor_manager.save_monitors()
        except Exception:
            pass
        w = self._monitor_cards.pop(monitor_id, None)
        if w:
            try:
                w.setParent(None)
                w.deleteLater()
            except Exception:
                pass

    def _remove_ftp_monitor(self, monitor_id: str):
        try:
            self.backend.monitor_manager.ftp_manager.remove_ftp_monitor(monitor_id)
            self.backend.monitor_manager.save_monitors()
        except Exception:
            pass
        w = self._monitor_cards.pop(monitor_id, None)
        if w:
            try:
                w.setParent(None)
                w.deleteLater()
            except Exception:
                pass


def _qt_normalize_extensions(exts_raw: str) -> list[str]:
    raw = str(exts_raw or "").strip()
    if not raw or raw in {"*", ".*"}:
        return ["*"]

    normalized = []
    seen = set()
    for token in re.split(r"[,;\s]+", raw):
        cleaned = token.strip().lower()
        if not cleaned:
            continue
        if cleaned in {"*", ".*"}:
            return ["*"]
        if cleaned.startswith("*."):
            cleaned = cleaned[2:]
        cleaned = cleaned.lstrip(".")
        if not cleaned:
            continue
        ext = f".{cleaned}"
        if ext not in seen:
            seen.add(ext)
            normalized.append(ext)
    return normalized or ["*"]


def _qt_extensions_to_input(exts: list[str]) -> str:
    if not exts or "*" in exts or ".*" in exts:
        return "*"
    return ", ".join(str(ext).lstrip(".") for ext in exts)


def _qt_is_file_routing_action(action: str) -> bool:
    return str(action or "").strip().lower() in {"move", "organize"}


def _qt_rule_action_preview(rule: dict) -> str:
    try:
        action = str(rule.get("action_type") or "").strip().lower()
        params = dict(rule.get("action_params") or {})
        exts = list(rule.get("file_extensions") or ["*"])
        events = list(rule.get("event_types") or ["created"])
        event_text = ", ".join(str(e) for e in events) if events else "manual/scheduled"
        ext_text = "all files" if "*" in exts or ".*" in exts else ", ".join(str(e).lstrip(".") for e in exts)
        regex = str(rule.get("name_pattern") or "").strip()
        filter_bits = [f"events: {event_text}", f"file types: {ext_text}"]
        if regex:
            filter_bits.append(f"name regex: {regex}")
        prefix = "Order: monitor filters first, then this rule checks " + "; ".join(filter_bits) + "."
        if _qt_is_file_routing_action(action) and set(events) != {"created"}:
            prefix += " For move/organize rules, Created only is the safer default because Modified can fire while files are still downloading."

        dest = str(params.get("destination") or "").strip()
        if action == "copy":
            duplicate = str(params.get("handle_duplicates") or "rename")
            return f"{prefix}\nAction: copy each matching file to {dest or '[choose destination]'}. Original files stay where they are. Duplicates: {duplicate}."
        if action == "move":
            duplicate = str(params.get("handle_duplicates") or "rename")
            return f"{prefix}\nAction: move each matching file to {dest or '[choose destination]'}. The original leaves the monitored folder. Duplicates: {duplicate}."
        if action == "organize":
            by = str(params.get("organize_by") or "extension")
            if by == "extension":
                example = "Example: report.pdf -> destination/pdf/report.pdf"
            elif by == "date":
                example = "Example: photo.jpg -> destination/2026/08/photo.jpg"
            else:
                example = "Example: song.mp3 -> destination/audio/song.mp3"
            return f"{prefix}\nAction: move matching files into subfolders under {dest or '[choose destination]'} by {by}. {example}"
        if action == "rename":
            pattern = str(params.get("pattern") or "{name}")
            return f"{prefix}\nAction: rename matching files in the same folder using: {pattern}. Example placeholders: {{name}}, {{date}}, {{time}}, {{timestamp}}."
        if action == "archive":
            name = str(params.get("archive_name") or "archive_{date}.zip")
            return f"{prefix}\nAction: add matching files to zip archive {name} in {dest or '[choose destination]'}."
        if action == "delete":
            return f"{prefix}\nAction: delete matching files using the Recycle Bin/app trash when available."
        if action == "execute":
            cmd = str(params.get("command") or "").strip()
            return f"{prefix}\nAction: run this command for each matching file: {cmd or '[enter command]'}. Use {{path}} for the full file path."
        if action == "cloud_upload":
            provider = str(params.get("provider") or "cloud")
            base = str(params.get("remote_base") or "Fylorra Sync")
            return f"{prefix}\nAction: upload matching files to {provider} under {base}."
        return f"{prefix}\nAction: choose what should happen to matching files."
    except Exception:
        return "Preview unavailable for this rule."


class _QtSmtpTestWorker(QObject):
    finished = Signal(bool, str)  # ok, message

    def __init__(self, smtp_settings: dict):
        super().__init__()
        self.smtp_settings = dict(smtp_settings or {})

    def run(self):
        try:
            from utils.email_notifier import EmailNotifier

            notifier = EmailNotifier(self.smtp_settings)
            ok, msg = notifier.test_connection()
            self.finished.emit(bool(ok), str(msg))
        except Exception as e:
            self.finished.emit(False, str(e))


class _QtMonitorCard(QFrame):
    def __init__(
        self,
        *,
        monitor_id: str,
        monitor_kind: str,
        monitor,
        rules: list[dict],
        icons: QtIconLoader,
        on_start,
        on_stop,
        on_remove,
        on_start_ftp,
        on_stop_ftp,
        on_remove_ftp,
        on_poll_ftp=None,
        on_build_rule=None,
        on_open_ai_hub=None,
    ):
        super().__init__()
        self.setObjectName("PageCard")
        self.monitor_id = monitor_id
        self.monitor_kind = monitor_kind
        self.monitor = monitor
        self.rules = rules
        self.icons = icons
        self._on_start = on_start
        self._on_stop = on_stop
        self._on_remove = on_remove
        self._on_start_ftp = on_start_ftp
        self._on_stop_ftp = on_stop_ftp
        self._on_remove_ftp = on_remove_ftp
        self._on_poll_ftp = on_poll_ftp
        self._on_build_rule = on_build_rule
        self._on_open_ai_hub = on_open_ai_hub

        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 14, 16, 14)
        layout.setSpacing(10)

        top = QHBoxLayout()
        top.setContentsMargins(0, 0, 0, 0)
        top.setSpacing(10)

        icon_lbl = QLabel("▣" if self.monitor_kind == "folder" else "◉")
        icon_lbl.setFixedWidth(18)
        self._type_icon_label = icon_lbl
        icon_lbl.setStyleSheet("font-size:16px; font-weight:700;")
        top.addWidget(icon_lbl)

        if self.monitor_kind == "folder":
            display_path = getattr(self.monitor, "path", "")
        else:
            try:
                display_path = self.monitor.get_connection_info()
            except Exception:
                display_path = ""
        path = QLabel(display_path)
        path.setStyleSheet("font-weight:600;")
        path.setTextInteractionFlags(Qt.TextSelectableByMouse)
        self._path_label = path
        top.addWidget(path, 1)

        controls = QHBoxLayout()
        controls.setSpacing(6)

        self.btn_start = self._make_icon_btn(
            icon_name="play",
            bg="#2fa572",
            hover="#106a43",
            tooltip="Start monitor",
            on_click=self._toggle_running,
        )
        controls.addWidget(self.btn_start)

        edit = self._make_icon_btn(
            icon_name="edit",
            bg="#FF9800",
            hover="#F57C00",
            tooltip="Edit monitor",
            on_click=self._edit_monitor,
        )
        controls.addWidget(edit)

        rm = self._make_icon_btn(
            icon_name="delete",
            bg="#d32f2f",
            hover="#9a0007",
            tooltip="Remove monitor",
            on_click=self._remove_clicked,
        )
        controls.addWidget(rm)

        if self.monitor_kind == "folder":
            open_folder = self._make_icon_btn(
                icon_name="folder",
                bg="#424242",
                hover="#2f2f2f",
                tooltip="Open folder",
                on_click=self._open_folder,
            )
            controls.addWidget(open_folder)

            build_rule = self._make_icon_btn(
                icon_name="add",
                bg="#1f3a5c",
                hover="#15304e",
                tooltip="Build an AI rule for this monitor",
                on_click=self._build_rule_for_monitor,
            )
            controls.addWidget(build_rule)

            run_rules = self._make_icon_btn(
                icon_name="ai",
                bg="#9C27B0",
                hover="#7B1FA2",
                tooltip="Run automation rules now",
                on_click=self._run_rules_now,
            )
            controls.addWidget(run_rules)

            ai_rename = self._make_icon_btn(
                icon_name="brain",
                bg="#1976D2",
                hover="#115293",
                tooltip="Smart Rename (AI Hub)",
                on_click=lambda: self._open_ai_hub("smart_rename"),
            )
            controls.addWidget(ai_rename)

            ai_cat = self._make_icon_btn(
                icon_name="grid",
                bg="#673AB7",
                hover="#4E2A91",
                tooltip="Auto-Categorize (AI Hub)",
                on_click=lambda: self._open_ai_hub("auto_categorize"),
            )
            controls.addWidget(ai_cat)

            ai_sec = self._make_icon_btn(
                icon_name="shield",
                bg="#F44336",
                hover="#D32F2F",
                tooltip="Security Scan (AI Hub)",
                on_click=lambda: self._open_ai_hub("security_scan"),
            )
            controls.addWidget(ai_sec)

            ai_doc = self._make_icon_btn(
                icon_name="ai_2",
                bg="#00BCD4",
                hover="#0097A7",
                tooltip="Content Analysis (AI Hub)",
                on_click=lambda: self._open_ai_hub("content_analysis"),
            )
            controls.addWidget(ai_doc)
        else:
            sync_now = self._make_icon_btn(
                icon_name="download",
                bg="#1976D2",
                hover="#115293",
                tooltip="Sync now (poll FTP + optionally download)",
                on_click=self._poll_ftp_now,
                icon_color="#FFFFFF",
            )
            controls.addWidget(sync_now)

        controls_widget = QWidget()
        controls_widget.setLayout(controls)
        top.addWidget(controls_widget)

        layout.addLayout(top)

        self.stats = QLabel("")
        self.stats.setStyleSheet("")
        layout.addWidget(self.stats)

        self.rules_label = QLabel("")
        self.rules_label.setStyleSheet("")
        layout.addWidget(self.rules_label)

        # Collapsible details show the operational timeline for each monitor.
        self._details_open = self.monitor_kind == "folder"
        self._details_anchor_height = None
        self.details_frame = QFrame()
        self.details_frame.setObjectName("MonitorDetails")
        df = QVBoxLayout(self.details_frame)
        df.setContentsMargins(0, 0, 0, 0)
        df.setSpacing(8)

        self.workflow_summary = QLabel("")
        self.workflow_summary.setWordWrap(True)
        df.addWidget(self.workflow_summary)

        self.activity = QListWidget()
        self.activity.setFixedHeight(190 if self.monitor_kind == "folder" else 120)
        df.addWidget(self.activity)
        layout.addWidget(self.details_frame)
        self.details_frame.setVisible(self._details_open)

        # Small non-intrusive toggle button on the right edge (like a drawer handle).
        self.details_btn = QToolButton(self)
        self.details_btn.setCheckable(True)
        self.details_btn.setChecked(self._details_open)
        self.details_btn.setAutoRaise(False)
        self.details_btn.setCursor(Qt.PointingHandCursor)
        self.details_btn.setToolTip("Show details")
        self.details_btn.setText("‹" if self._details_open else "›")
        self.details_btn.clicked.connect(self._toggle_details)
        self._position_details_button()

        self._countdown_timer = QTimer(self)
        self._countdown_timer.setInterval(1000)
        self._countdown_timer.timeout.connect(self._refresh_stats)
        self._countdown_timer.start()

        self._apply_theme_styles()
        self._refresh_stats()

    def resizeEvent(self, event):  # noqa: N802
        super().resizeEvent(event)
        self._position_details_button()

    def _position_details_button(self):
        try:
            w = 24
            h = 46
            x = max(0, self.width() - w - 4)
            base_h = self._details_anchor_height or self.height()
            # Place the drawer handle clearly below center so it stays away from the
            # top action toolbar (matches the "lower right" expectation).
            y = int(base_h * 0.72 - h / 2)
            y = max(10, min(self.height() - h - 10, y))
            self.details_btn.setFixedSize(w, h)
            self.details_btn.move(x, y)
            self.details_btn.raise_()
        except Exception:
            pass

    def _toggle_details(self):
        self._details_open = bool(self.details_btn.isChecked())
        if self._details_open:
            self._details_anchor_height = self.height()
        else:
            self._details_anchor_height = None
        try:
            self.details_frame.setVisible(self._details_open)
        except Exception:
            pass
        try:
            self.details_btn.setToolTip("Hide details" if self._details_open else "Show details")
        except Exception:
            pass
        try:
            self.details_btn.setText("‹" if self._details_open else "›")
        except Exception:
            pass
        self._position_details_button()

    def _make_icon_btn(self, *, icon_name: str, bg: str, hover: str, tooltip: str, on_click, icon_color: str | None = None):
        btn = QPushButton()
        btn.setFixedSize(36, 36)
        if icon_color:
            btn.setIcon(self._tinted_icon(icon_name, 18, icon_color))
        else:
            btn.setIcon(self.icons.icon(icon_name))
        btn.setIconSize(QSize(18, 18))
        btn.setToolTip(tooltip)
        btn.setStyleSheet(
            f"QPushButton{{background:{bg}; border:1px solid {bg}; border-radius:10px;}}"
            f"QPushButton:hover{{background:{hover}; border:1px solid {hover};}}"
        )
        btn.clicked.connect(on_click)
        return btn

    def _theme_mode(self) -> str:
        try:
            app = QApplication.instance()
            if app:
                mode = str(app.property("fg_theme_mode") or "dark").strip().lower()
                if mode in {"dark", "light", "black"}:
                    return mode
        except Exception:
            pass
        return "dark"

    def _accent_name(self) -> str:
        try:
            app = QApplication.instance()
            if app:
                value = str(app.property("fg_color_theme") or "blue").strip().lower()
                if value in {"blue", "teal", "green", "orange", "rose", "violet", "cyan"}:
                    return value
        except Exception:
            pass
        return "blue"

    def _theme_tokens(self) -> dict[str, str]:
        mode = self._theme_mode()
        accent = self._accent_name()
        accent_main = {
            "blue": "#0d6efd",
            "teal": "#14b8a6",
            "green": "#22c55e",
            "orange": "#f59e0b",
            "rose": "#f43f5e",
            "violet": "#8b5cf6",
            "cyan": "#06b6d4",
        }.get(accent, "#0d6efd")
        if mode == "light":
            return {
                "path": "#0f172a",
                "meta": "#475569",
                "panel_bg": "#f8fbff",
                "panel_border": "#cfd7e4",
                "panel_fg": "#1f2937",
                "handle_bg": "#eef3fa",
                "handle_hover": "#e2eaf6",
                "handle_checked": "#dbeafe",
                "handle_fg": "#475569",
                "handle_fg_active": "#0f172a",
                "handle_border": "#c8d3e3",
                "accent": accent_main,
            }
        if mode == "black":
            return {
                "path": "#f2f2f2",
                "meta": "#a0a0a0",
                "panel_bg": "#090909",
                "panel_border": "#1f1f1f",
                "panel_fg": "#e5e7eb",
                "handle_bg": "#0f0f0f",
                "handle_hover": "#1a1a1a",
                "handle_checked": "#151515",
                "handle_fg": "#a0a0a0",
                "handle_fg_active": "#ffffff",
                "handle_border": "#2a2a2a",
                "accent": accent_main,
            }
        return {
            "path": "#e6e8ee",
            "meta": "#9aa0a9",
            "panel_bg": "#14171c",
            "panel_border": "#232730",
            "panel_fg": "#c8ccd6",
            "handle_bg": "#171a1f",
            "handle_hover": "#1f242d",
            "handle_checked": "#1f2b3a",
            "handle_fg": "#9aa0a9",
            "handle_fg_active": "#ffffff",
            "handle_border": "#232730",
            "accent": accent_main,
        }

    def _apply_theme_styles(self) -> None:
        t = self._theme_tokens()
        try:
            self._path_label.setStyleSheet(f"color:{t['path']}; font-weight:700;")
        except Exception:
            pass
        try:
            self.stats.setStyleSheet(f"color:{t['meta']};")
        except Exception:
            pass
        try:
            self.rules_label.setStyleSheet(f"color:{t['meta']};")
        except Exception:
            pass
        try:
            self.workflow_summary.setStyleSheet(f"color:{t['panel_fg']};")
        except Exception:
            pass
        try:
            icon_color = t["accent"] if self.monitor_kind == "folder" else "#a855f7"
            self._type_icon_label.setStyleSheet(f"color:{icon_color}; font-size:16px; font-weight:700;")
        except Exception:
            pass
        try:
            self.activity.setStyleSheet(
                "QListWidget{"
                f"background:{t['panel_bg']}; border:1px solid {t['panel_border']}; border-radius:10px; color:{t['panel_fg']};"
                "}"
                "QListWidget::item{padding:4px 8px;}"
            )
        except Exception:
            pass
        try:
            self.details_btn.setStyleSheet(
                "QToolButton{"
                f"background:{t['handle_bg']}; color:{t['handle_fg']}; border:1px solid {t['handle_border']};"
                "border-radius:8px; font-weight:700; font-size:24px; padding-top:0px; padding-bottom:6px;"
                "}"
                f"QToolButton:hover{{background:{t['handle_hover']}; color:{t['handle_fg_active']}; border-color:{t['accent']};}}"
                f"QToolButton:checked{{background:{t['handle_checked']}; color:{t['handle_fg_active']}; border-color:{t['accent']};}}"
            )
        except Exception:
            pass

    def refresh_theme(self) -> None:
        self._apply_theme_styles()

    def _tinted_icon(self, icon_name: str, size: int, color: str) -> QIcon:
        pm = self.icons.pixmap(icon_name, int(size))
        if pm.isNull():
            return self.icons.icon(icon_name)
        img = QImage(pm.size(), QImage.Format_ARGB32)
        img.fill(Qt.transparent)
        p = QPainter(img)
        try:
            p.drawPixmap(0, 0, pm)
            p.setCompositionMode(QPainter.CompositionMode_SourceIn)
            p.fillRect(img.rect(), QColor(color))
        finally:
            p.end()
        return QIcon(QPixmap.fromImage(img))

    def _remove_clicked(self):
        if self.monitor_kind == "ftp":
            self._on_remove_ftp()
        else:
            self._on_remove()

    def _edit_monitor(self):
        if self.monitor_kind == "ftp":
            dlg = _EditFtpMonitorDialog(self, monitor=self.monitor)
        else:
            dlg = _EditFolderMonitorDialog(self, monitor=self.monitor)
        if dlg.exec() != QDialog.Accepted:
            return
        try:
            dlg.apply_to_monitor()
        except Exception as e:
            QMessageBox.critical(self, "Save Failed", str(e))
            return
        try:
            top = self.window()
            backend = getattr(top, "backend", None)
            if backend:
                backend.monitor_manager.save_monitors()
        except Exception:
            pass
        try:
            if self.monitor_kind == "folder":
                wants_running = bool(dlg.desired_running())
                is_running = bool(getattr(self.monitor, "is_running", False))
                if wants_running and not is_running:
                    self._on_start()
                elif is_running and not wants_running:
                    self._on_stop()
        except Exception:
            pass
        self._refresh_stats()

    def _toggle_running(self):
        try:
            running = bool(getattr(self.monitor, "is_running", False))
        except Exception:
            running = False
        if running:
            try:
                if self.monitor_kind == "ftp":
                    self._on_stop_ftp()
                else:
                    self._on_stop()
            except Exception:
                pass
        else:
            try:
                if self.monitor_kind == "ftp":
                    ok = self._on_start_ftp()
                else:
                    ok = self._on_start()
                if ok is False:
                    self._show_monitor_action_error("Start Monitor Failed")
            except Exception:
                self._show_monitor_action_error("Start Monitor Failed")
        self._refresh_stats()

    def _show_monitor_action_error(self, title: str) -> None:
        msg = "The monitor action failed."
        try:
            top = self.window()
            backend = getattr(top, "backend", None)
            manager = getattr(backend, "monitor_manager", None)
            if manager is not None:
                if self.monitor_kind == "ftp":
                    ftp_manager = getattr(manager, "ftp_manager", None)
                    msg = str(getattr(ftp_manager, "last_error", "") or msg)
                else:
                    msg = str(getattr(manager, "last_error", "") or msg)
        except Exception:
            pass
        QMessageBox.critical(self, title, msg)

    def _open_folder(self):
        try:
            from PySide6.QtCore import QUrl
            from PySide6.QtGui import QDesktopServices

            p = getattr(self.monitor, "path", "")
            if p:
                QDesktopServices.openUrl(QUrl.fromLocalFile(str(p)))
        except Exception:
            pass

    def _poll_ftp_now(self):
        if self.monitor_kind != "ftp":
            return
        try:
            if callable(self._on_poll_ftp):
                ok = self._on_poll_ftp()
                if ok is False:
                    self._show_monitor_action_error("FTP Sync Failed")
        except Exception:
            self._show_monitor_action_error("FTP Sync Failed")
        self._refresh_stats()

    def _build_rule_for_monitor(self):
        try:
            if callable(self._on_build_rule):
                self._on_build_rule()
        except Exception:
            pass

    def _open_ai_hub(self, op_key: str):
        try:
            if callable(self._on_open_ai_hub):
                self._on_open_ai_hub(str(op_key))
        except Exception:
            pass

    def _refresh_stats(self):
        try:
            running = bool(getattr(self.monitor, "is_running", False))
            stats = getattr(self.monitor, "stats", {}) or {}
            status = "Running" if running else "Stopped"
            if self.monitor_kind == "ftp":
                last_status = str(stats.get("last_status") or "Never polled")
                last_success = self._short_time(stats.get("last_success"))
                last_poll = self._short_time(stats.get("last_poll"))
                uploaded = int(stats.get("files_uploaded", 0) or 0)
                errors = int(stats.get("connection_errors", 0) or 0) + int(stats.get("download_errors", 0) or 0) + int(stats.get("upload_errors", 0) or 0)
                self.stats.setText(
                    f"Status: {status} | Sync: {last_status} | Files: {stats.get('remote_files', 0)} | Remote: +{stats.get('files_created',0)} / "
                    f"~{stats.get('files_modified',0)} / -{stats.get('files_deleted',0)} | "
                    f"Downloaded: {stats.get('files_downloaded',0)} | Uploaded: {uploaded} | Errors: {errors} | "
                    f"Last OK: {last_success or '-'} | Last poll: {last_poll or '-'}"
                )
            else:
                self.stats.setText(
                    f"Status: {status} | Created: {stats.get('files_created',0)} | Modified: {stats.get('files_modified',0)} | "
                    f"Deleted: {stats.get('files_deleted',0)} | Actions: {stats.get('actions_executed',0)}"
                )
            self.btn_start.setIcon(self.icons.icon("pause" if running else "play"))
        except Exception:
            pass
        try:
            if self.monitor_kind == "folder":
                n = len(self.rules or [])
                filters = self._folder_filter_summary()
                suffix = f" | Filters: {filters}" if filters else " | Filters: none"
                self.rules_label.setText(f"⚙️ {n} automation rule{'s' if n != 1 else ''} configured{suffix}")
                self._refresh_folder_activity()
            else:
                sync_dir = str(getattr(self.monitor, "local_sync_dir", "") or "").strip()
                mode = "Two-way sync" if bool(getattr(self.monitor, "two_way_sync", False)) else "Download sync"
                subfolders = "including subfolders" if bool(getattr(self.monitor, "sync_subfolders", False)) else "top folder only"
                ext_filter = list(getattr(self.monitor, "allowed_extensions", []) or [])
                ext_text = ", ".join(ext_filter[:5]) if ext_filter else "all files"
                if sync_dir:
                    self.rules_label.setText(f"{mode}: {subfolders} | Filter: {ext_text} | Local: {sync_dir}")
                else:
                    self.rules_label.setText(f"Watch only: {subfolders} | Filter: {ext_text}")
        except Exception:
            pass
        try:
            if self.monitor_kind == "ftp":
                self._refresh_ftp_activity()
        except Exception:
            pass

    def _folder_filter_summary(self) -> str:
        parts = []
        try:
            min_size = getattr(self.monitor, "min_size_kb", None)
            max_size = getattr(self.monitor, "max_size_kb", None)
            if min_size is not None:
                parts.append(f">= {min_size} KB")
            if max_size is not None:
                parts.append(f"<= {max_size} KB")
            days = getattr(self.monitor, "modified_within_days", None)
            if days:
                parts.append(f"modified in {days} day(s)")
            excludes = list(getattr(self.monitor, "exclude_patterns", []) or [])
            if excludes:
                parts.append(f"{len(excludes)} exclude pattern(s)")
            regex = str(getattr(self.monitor, "filename_regex", "") or "").strip()
            if regex:
                parts.append("filename regex")
        except Exception:
            return ""
        return ", ".join(parts)

    def _short_time(self, value) -> str:
        text = str(value or "").strip()
        if not text:
            return ""
        if "T" in text:
            try:
                return text.split("T", 1)[1].split(".", 1)[0][:8]
            except Exception:
                return text
        return text[:19]

    def _refresh_ftp_activity(self) -> None:
        entries = list(getattr(self.monitor, "recent_activity", []) or [])
        self.activity.clear()
        if not entries:
            self.activity.addItem(QListWidgetItem("No FTP activity yet. Use Sync now or start the monitor."))
            return
        for entry in entries[:50]:
            try:
                ts = str(entry.get("time") or "")
                kind = str(entry.get("kind") or "event").capitalize()
                msg = str(entry.get("message") or "")
                level = str(entry.get("level") or "info").lower()
            except Exception:
                ts, kind, msg, level = "", "Event", str(entry), "info"
            item = QListWidgetItem(f"{ts}  {kind}: {msg}".strip())
            if level == "error":
                item.setForeground(QColor("#ef4444"))
            elif level == "warning":
                item.setForeground(QColor("#f59e0b"))
            self.activity.addItem(item)

    def _refresh_folder_activity(self) -> None:
        try:
            pending = list((getattr(self.monitor, "pending_actions", {}) or {}).values())
            if pending:
                lines = []
                now = time.time()
                for entry in pending[:3]:
                    name = str(entry.get("file_name") or Path(str(entry.get("path") or "")).name or "file")
                    due_at = float(entry.get("due_at") or now)
                    remaining = max(0, int(round(due_at - now)))
                    stable = int(round(float(entry.get("stable_wait_seconds") or 0)))
                    status = str(entry.get("status") or "waiting")
                    matched = int(entry.get("matched_rule_count") or 0)
                    match_text = f"{matched} rule{'s' if matched != 1 else ''}"
                    if remaining > 0:
                        lines.append(f"{name}: {match_text} matched, {status}, {remaining}s remaining")
                    elif stable > 0:
                        lines.append(f"{name}: {match_text} matched, checking stability ({stable}s)")
                    else:
                        lines.append(f"{name}: {match_text} matched, ready to run")
                more = f" +{len(pending) - 3} more" if len(pending) > 3 else ""
                self.workflow_summary.setText("Pending actions: " + "; ".join(lines) + more)
            else:
                delay = int(round(float(getattr(self.monitor, "action_delay_seconds", 0) or 0)))
                stable = int(round(float(getattr(self.monitor, "action_stability_seconds", 0) or 0)))
                parts = ["Watching for new files"]
                if delay:
                    parts.append(f"delay {delay}s")
                if stable:
                    parts.append(f"stable check {stable}s")
                parts.append("existing files run only by command")
                self.workflow_summary.setText(" | ".join(parts))
        except Exception:
            pass

        entries = list(getattr(self.monitor, "recent_activity", []) or [])
        self.activity.clear()
        if not entries:
            self.activity.addItem(QListWidgetItem("No folder activity yet. New files will show detection, matching, delay, and action result here."))
            return
        for entry in entries[:80]:
            try:
                ts = str(entry.get("time") or "")
                kind = str(entry.get("kind") or "Event")
                msg = str(entry.get("message") or "")
                level = str(entry.get("level") or "info").lower()
            except Exception:
                ts, kind, msg, level = "", "Event", str(entry), "info"
            item = QListWidgetItem(f"{ts}  {kind}: {msg}".strip())
            if level == "error":
                item.setForeground(QColor("#ef4444"))
            elif level == "warning":
                item.setForeground(QColor("#f59e0b"))
            elif kind.lower() in {"completed", "rule matched", "planned"}:
                item.setForeground(QColor("#22c55e"))
            self.activity.addItem(item)

    def _add_event(self, event_type: str, src_path: str):
        if hasattr(self.monitor, "recent_activity"):
            return
        try:
            from pathlib import Path

            name = Path(src_path).name
        except Exception:
            name = src_path
        item = QListWidgetItem(f"{event_type}: {name}")
        self.activity.insertItem(0, item)
        while self.activity.count() > 50:
            self.activity.takeItem(self.activity.count() - 1)

    def _run_rules_now(self):
        if self.monitor_kind != "folder":
            return
        if not self.rules:
            QMessageBox.information(self, "No Rules", "This monitor has no rules configured.")
            return
        if not self._confirm_run_rules_now():
            return
        dlg = _RunRulesNowDialog(self, monitor=self.monitor, rules=self.rules)
        dlg.exec()

    def _confirm_run_rules_now(self) -> bool:
        try:
            from pathlib import Path

            folder_path = Path(str(getattr(self.monitor, "path", "") or ""))
            if not folder_path.exists():
                QMessageBox.warning(self, "Run Rules Now", "The monitored folder does not exist.")
                return False

            engine = getattr(self.monitor, "action_engine", None)
            if not engine or not hasattr(engine, "preview_action"):
                return QMessageBox.question(
                    self,
                    "Run Rules Now",
                    "Run configured rules on existing files in this monitor?",
                ) == QMessageBox.Yes

            examples = []
            planned = 0
            scanned = 0
            max_scan = 3000
            passes_filters = getattr(self.monitor, "_passes_filters", None)

            for file_path in folder_path.rglob("*"):
                if scanned >= max_scan:
                    break
                if not file_path.is_file():
                    continue
                scanned += 1
                if callable(passes_filters):
                    try:
                        if not passes_filters(str(file_path)):
                            continue
                    except Exception:
                        pass

                for rule in self.rules:
                    events = rule.get("event_types") or []
                    if isinstance(events, str):
                        events = [events]
                    if events and "created" not in [str(e).strip().lower() for e in events]:
                        continue
                    try:
                        if not self.monitor._rule_matches(rule, "created", file_path.name, file_path.suffix.lower()):
                            continue
                    except Exception:
                        continue
                    preview = engine.preview_action(
                        str(rule.get("action_type") or ""),
                        str(file_path),
                        dict(rule.get("action_params") or {}),
                    )
                    planned += 1
                    if len(examples) < 8:
                        examples.append(str(preview.get("summary") or "Action planned."))

            if planned <= 0:
                QMessageBox.information(
                    self,
                    "Run Rules Now",
                    "No existing files match the current monitor filters and rule filters.",
                )
                return False

            more = "\n..." if planned > len(examples) else ""
            capped = "\nOnly the first 3,000 files were checked for this preview." if scanned >= max_scan else ""
            message = (
                f"Fylorra found {planned} action(s) that would run on existing files.\n\n"
                "Examples:\n"
                + "\n".join(f"- {x}" for x in examples)
                + more
                + capped
                + "\n\nContinue and apply these actions?"
            )
            return QMessageBox.question(self, "Confirm Run Rules", message) == QMessageBox.Yes
        except Exception as e:
            QMessageBox.warning(self, "Run Rules Now", f"Could not build preview:\n{e}")
            return False


class _AddFtpMonitorDialog(QDialog):
    def __init__(self, parent: QWidget):
        super().__init__(parent)
        self._main = parent
        self.setWindowTitle("Add FTP Monitor")
        self.setModal(True)
        self.setWindowFlag(Qt.WindowContextHelpButtonHint, False)
        self.setMinimumWidth(560)
        self.setMinimumHeight(520)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(18, 16, 18, 16)
        layout.setSpacing(12)

        title = QLabel("Add FTP Monitor")
        title.setObjectName("PageTitle")
        subtitle = QLabel("Monitor a remote FTP/FTPS folder and optionally sync new/changed files to a local folder.")
        subtitle.setObjectName("PageSubTitle")
        subtitle.setWordWrap(True)
        layout.addWidget(title)
        layout.addWidget(subtitle)

        card = QFrame()
        card.setObjectName("PageCard")
        card_layout = QVBoxLayout(card)
        card_layout.setContentsMargins(16, 14, 16, 14)
        card_layout.setSpacing(12)

        connection_box = QGroupBox("Connection")
        conn_form = QFormLayout(connection_box)
        conn_form.setLabelAlignment(Qt.AlignRight)
        conn_form.setVerticalSpacing(10)
        conn_form.setHorizontalSpacing(12)

        self.host = QLineEdit()
        self.host.setPlaceholderText("ftp.example.com")
        conn_form.addRow("Host:", self.host)

        self.port = QSpinBox()
        self.port.setRange(1, 65535)
        self.port.setValue(21)
        conn_form.addRow("Port:", self.port)

        remote_row = QHBoxLayout()
        self.remote_path = QLineEdit()
        self.remote_path.setPlaceholderText("/path/on/server")
        self.remote_path.setText("/")
        remote_row.addWidget(self.remote_path, 1)
        remote_browse = QPushButton("Browse")
        remote_browse.clicked.connect(self._browse_remote_path)
        remote_row.addWidget(remote_browse)
        conn_form.addRow("Remote path:", remote_row)

        self.poll_interval = QSpinBox()
        self.poll_interval.setRange(5, 3600)
        self.poll_interval.setValue(30)
        conn_form.addRow("Poll interval (s):", self.poll_interval)

        self.use_tls = QCheckBox("Use TLS (FTPS)")
        self.use_tls.toggled.connect(self._toggle_tls_options)
        conn_form.addRow("", self.use_tls)

        self.tls_implicit = QCheckBox("Implicit FTPS (port 990)")
        self.tls_implicit.setEnabled(False)
        self.tls_implicit.toggled.connect(self._maybe_set_implicit_port)
        conn_form.addRow("", self.tls_implicit)

        self.passive_mode = QCheckBox("Passive mode (recommended)")
        self.passive_mode.setChecked(True)
        conn_form.addRow("", self.passive_mode)

        self.ftp_encoding = QLineEdit("utf-8")
        self.ftp_encoding.setPlaceholderText("Server encoding (e.g. utf-8)")
        conn_form.addRow("Encoding:", self.ftp_encoding)

        card_layout.addWidget(connection_box)

        auth_box = QGroupBox("Authentication")
        auth_form = QFormLayout(auth_box)
        auth_form.setLabelAlignment(Qt.AlignRight)
        auth_form.setVerticalSpacing(10)
        auth_form.setHorizontalSpacing(12)

        self.username = QLineEdit()
        auth_form.addRow("Username:", self.username)

        self.password = QLineEdit()
        self.password.setEchoMode(QLineEdit.Password)
        auth_form.addRow("Password:", self.password)

        card_layout.addWidget(auth_box)

        sync_box = QGroupBox("Local Sync (optional)")
        sync_form = QFormLayout(sync_box)
        sync_form.setLabelAlignment(Qt.AlignRight)
        sync_form.setVerticalSpacing(10)
        sync_form.setHorizontalSpacing(12)

        sync_row = QHBoxLayout()
        self.local_sync_dir = QLineEdit()
        self.local_sync_dir.setPlaceholderText("Optional local folder to download new/changed files")
        sync_row.addWidget(self.local_sync_dir, 1)
        browse = QPushButton("Browse")
        browse.clicked.connect(self._browse_local_dir)
        sync_row.addWidget(browse)
        sync_form.addRow("Local folder:", sync_row)

        self.two_way_sync = QCheckBox("Two-way sync (upload + download)")
        self.two_way_sync.setChecked(False)
        sync_form.addRow("", self.two_way_sync)

        self.sync_subfolders = QCheckBox("Sync subfolders")
        self.sync_subfolders.setChecked(False)
        sync_form.addRow("", self.sync_subfolders)

        self.download_on_created = QCheckBox("Download new files")
        self.download_on_created.setChecked(True)
        sync_form.addRow("", self.download_on_created)

        self.download_on_modified = QCheckBox("Download modified files")
        self.download_on_modified.setChecked(True)
        sync_form.addRow("", self.download_on_modified)

        self.delete_local_on_deleted = QCheckBox("Delete local file when remote is deleted")
        self.delete_local_on_deleted.setChecked(False)
        sync_form.addRow("", self.delete_local_on_deleted)

        self.overwrite_local = QCheckBox("Overwrite local files")
        self.overwrite_local.setChecked(False)
        sync_form.addRow("", self.overwrite_local)

        self.allowed_exts = QLineEdit()
        self.allowed_exts.setPlaceholderText("Optional extensions: jpg,png,pdf,mp3")
        sync_form.addRow("Extensions:", self.allowed_exts)

        card_layout.addWidget(sync_box)

        self.auto_start = QCheckBox("Start immediately")
        self.auto_start.setChecked(True)
        card_layout.addWidget(self.auto_start)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QFrame.NoFrame)
        scroll.setWidget(card)
        layout.addWidget(scroll, 1)

        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        try:
            ok_btn = buttons.button(QDialogButtonBox.Ok)
            if ok_btn:
                ok_btn.setObjectName("PrimaryButton")
        except Exception:
            pass
        buttons.accepted.connect(self._validate_and_accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)

    def _validate_and_accept(self):
        top = getattr(self, "_main", None) or self.parent()
        backend = getattr(top, "backend", None)
        ftp_manager = getattr(getattr(backend, "monitor_manager", None), "ftp_manager", None)
        if ftp_manager and hasattr(ftp_manager, "validate_ftp_config"):
            values = self.values()
            problems = ftp_manager.validate_ftp_config(
                host=values["host"],
                username=values["username"],
                password=values["password"],
                remote_path=values["remote_path"],
                port=values["port"],
                poll_interval=values["poll_interval"],
                local_sync_dir=values.get("local_sync_dir"),
                allowed_extensions=values.get("allowed_extensions"),
            )
            if problems:
                QMessageBox.warning(self, "Invalid FTP Monitor", "\n".join(problems))
                return
        elif not self.host.text().strip() or not self.remote_path.text().strip():
            QMessageBox.warning(self, "Invalid FTP Monitor", "Host and remote path are required.")
            return
        self.accept()

    def values(self) -> dict:
        return {
            "host": self.host.text().strip(),
            "port": int(self.port.value()),
            "username": self.username.text().strip(),
            "password": self.password.text(),
            "remote_path": self.remote_path.text().strip(),
            "use_tls": bool(self.use_tls.isChecked()),
            "tls_implicit": bool(self.tls_implicit.isChecked()),
            "passive_mode": bool(self.passive_mode.isChecked()),
            "encoding": (self.ftp_encoding.text() or "").strip() or "utf-8",
            "poll_interval": int(self.poll_interval.value()),
            "two_way_sync": bool(self.two_way_sync.isChecked()),
            "sync_subfolders": bool(self.sync_subfolders.isChecked()),
            "local_sync_dir": self.local_sync_dir.text().strip() or None,
            "download_on_created": bool(self.download_on_created.isChecked()),
            "download_on_modified": bool(self.download_on_modified.isChecked()),
            "delete_local_on_deleted": bool(self.delete_local_on_deleted.isChecked()),
            "overwrite_local": bool(self.overwrite_local.isChecked()),
            "allowed_extensions": [e.strip().lower().lstrip(".") for e in (self.allowed_exts.text() or "").split(",") if e.strip()],
            "auto_start": bool(self.auto_start.isChecked()),
        }

    def _browse_local_dir(self):
        folder = QFileDialog.getExistingDirectory(self, "Select local sync folder")
        if folder:
            self.local_sync_dir.setText(folder)

    def _browse_remote_path(self):
        top = getattr(self, "_main", None) or self.parent()
        backend = getattr(top, "backend", None)
        ftp_mgr = getattr(getattr(getattr(backend, "monitor_manager", None), "ftp_manager", None), "list_remote_dirs", None)
        if not backend or not ftp_mgr:
            QMessageBox.warning(self, "FTP", "FTP remote browser is not available.")
            return

        host = self.host.text().strip()
        if not host:
            QMessageBox.warning(self, "FTP", "Enter Host first.")
            return

        dlg = _FtpRemoteBrowserDialog(
            self,
            ftp_manager=backend.monitor_manager.ftp_manager,
            host=host,
            username=self.username.text().strip(),
            password=self.password.text(),
            remote_path=self.remote_path.text().strip() or "/",
            port=int(self.port.value() or 21),
            use_tls=bool(self.use_tls.isChecked()),
            tls_implicit=bool(self.tls_implicit.isChecked()),
            passive_mode=bool(self.passive_mode.isChecked()),
            encoding=(self.ftp_encoding.text() or "").strip() or "utf-8",
        )
        if dlg.exec() == QDialog.Accepted:
            self.remote_path.setText(dlg.selected_path())

    def _toggle_tls_options(self, enabled: bool):
        enabled = bool(enabled)
        try:
            self.tls_implicit.setEnabled(enabled)
            if not enabled:
                self.tls_implicit.setChecked(False)
        except Exception:
            pass

    def _maybe_set_implicit_port(self, enabled: bool):
        if not enabled:
            return
        try:
            if int(self.port.value()) == 21:
                self.port.setValue(990)
        except Exception:
            pass


class _RunRulesWorker(QObject):
    progress = Signal(float, int, int)  # pct, current, total
    status = Signal(str)
    finished = Signal(int, int, bool)  # processed, actions, cancelled
    error = Signal(str)

    def __init__(self, *, monitor, rules: list[dict]):
        super().__init__()
        self.monitor = monitor
        self.rules = rules
        self.cancelled = False

    def cancel(self):
        self.cancelled = True

    def run(self):
        from pathlib import Path
        import re

        processed = 0
        actions_ok = 0
        was_running = False
        try:
            folder_path = Path(getattr(self.monitor, "path", ""))
            if not folder_path.exists():
                self.error.emit("Folder does not exist.")
                return

            self.status.emit("Scanning folder…")
            was_running = bool(getattr(self.monitor, "is_running", False))
            if was_running:
                try:
                    self.monitor.stop()
                except Exception:
                    pass

            destination_paths: set[Path] = set()
            for rule in self.rules:
                if rule.get("action_type") in ["move", "copy"]:
                    dest = (rule.get("action_params") or {}).get("destination", "")
                    if dest:
                        try:
                            destination_paths.add(Path(dest).resolve())
                        except Exception:
                            pass

            all_files: list[Path] = []
            for item in folder_path.rglob("*"):
                if self.cancelled:
                    break
                if not item.is_file():
                    continue
                skip = False
                try:
                    item_resolved = item.resolve()
                    for dest_path in destination_paths:
                        if dest_path in item_resolved.parents or item_resolved.parent == dest_path:
                            skip = True
                            break
                except Exception:
                    pass
                if not skip:
                    passes_filters = getattr(self.monitor, "_passes_filters", None)
                    if callable(passes_filters):
                        try:
                            if not passes_filters(str(item)):
                                continue
                        except Exception:
                            pass
                    all_files.append(item)

            total_files = len(all_files)
            if total_files == 0:
                self.status.emit("No files found to process.")
                self.finished.emit(0, 0, self.cancelled)
                return

            self.status.emit(f"Found {total_files} files. Running rules…")

            for idx, file_path in enumerate(all_files):
                if self.cancelled:
                    break
                processed = idx + 1
                pct = processed / total_files
                self.progress.emit(pct, processed, total_files)

                for rule in self.rules:
                    if self.cancelled:
                        break
                    try:
                        events = rule.get("event_types") or []
                        if isinstance(events, str):
                            events = [events]
                        if events and "created" not in [str(e).strip().lower() for e in events]:
                            continue
                        exts = rule.get("file_extensions", ["*"]) or ["*"]
                        if "*" not in exts:
                            normalized_exts = []
                            for e in exts:
                                value = str(e or "").strip().lower()
                                if value and value not in {"*", ".*"} and not value.startswith("."):
                                    value = "." + value
                                if value:
                                    normalized_exts.append(value)
                            if file_path.suffix.lower() not in normalized_exts:
                                continue
                        pat = rule.get("name_pattern")
                        if pat:
                            try:
                                if not re.search(pat, file_path.name, re.IGNORECASE):
                                    continue
                            except re.error:
                                continue
                        ok = self.monitor.action_engine.execute_action(
                            rule["action_type"], str(file_path), rule.get("action_params", {}) or {}
                        )
                        if ok:
                            actions_ok += 1
                    except Exception:
                        continue

            self.finished.emit(processed, actions_ok, self.cancelled)

        except Exception as e:
            self.error.emit(str(e))
        finally:
            if was_running:
                try:
                    self.monitor.start()
                except Exception:
                    pass


class _RunRulesNowDialog(QDialog):
    def __init__(self, parent: QWidget, *, monitor, rules: list[dict]):
        super().__init__(parent)
        self.setWindowTitle("Run Rules Now")
        self.setModal(True)
        self.setMinimumWidth(520)
        self._thread: QThread | None = None

        layout = QVBoxLayout(self)
        layout.setContentsMargins(14, 14, 14, 14)
        layout.setSpacing(10)

        title = QLabel("Running rules on existing files")
        title.setStyleSheet("font-size:16px; font-weight:700; color:#ffffff;")
        layout.addWidget(title)

        self.status = QLabel("Starting…")
        self.status.setStyleSheet("color:#9aa0a9;")
        layout.addWidget(self.status)

        self.bar = QProgressBar()
        self.bar.setRange(0, 1000)
        self.bar.setValue(0)
        layout.addWidget(self.bar)

        self.details = QLabel("")
        self.details.setStyleSheet("color:#9aa0a9;")
        layout.addWidget(self.details)

        btn_row = QHBoxLayout()
        btn_row.addStretch(1)
        self.btn_cancel = QPushButton("Cancel")
        self.btn_cancel.clicked.connect(self._cancel)
        btn_row.addWidget(self.btn_cancel)
        layout.addLayout(btn_row)

        self.worker = _RunRulesWorker(monitor=monitor, rules=rules)
        self._start_worker()

    def _start_worker(self):
        thread = QThread(self)
        self._thread = thread
        self.worker.moveToThread(thread)
        thread.started.connect(self.worker.run)
        self.worker.status.connect(self._on_status)
        self.worker.progress.connect(self._on_progress)
        self.worker.finished.connect(self._on_finished)
        self.worker.error.connect(self._on_error)
        self.worker.finished.connect(thread.quit)
        self.worker.error.connect(thread.quit)
        thread.finished.connect(thread.deleteLater)
        thread.start()

    def _cancel(self):
        self.btn_cancel.setEnabled(False)
        try:
            self.worker.cancel()
        except Exception:
            pass
        self.status.setText("Cancelling…")

    def _on_status(self, text: str):
        self.status.setText(text)

    def _on_progress(self, pct: float, cur: int, total: int):
        self.bar.setValue(int(max(0.0, min(1.0, pct)) * 1000))
        self.details.setText(f"Processed {cur}/{total} files…")

    def _on_finished(self, processed: int, actions_ok: int, cancelled: bool):
        if cancelled:
            self.status.setText("Cancelled.")
        else:
            self.status.setText("Complete.")
            self.bar.setValue(1000)
        self.details.setText(f"Processed {processed} files • Applied {actions_ok} actions")
        self.btn_cancel.setText("Close")
        self.btn_cancel.setEnabled(True)
        try:
            self.btn_cancel.clicked.disconnect()
        except Exception:
            pass
        self.btn_cancel.clicked.connect(self.accept)

    def _on_error(self, msg: str):
        QMessageBox.critical(self, "Run Failed", msg)
        self.reject()


class _EditFolderMonitorDialog(QDialog):
    def __init__(self, parent: QWidget, *, monitor, default_running: bool | None = None):
        super().__init__(parent)
        self.monitor = monitor
        self.setWindowTitle("Edit Folder Monitor")
        self.setModal(True)
        self.setWindowFlag(Qt.WindowContextHelpButtonHint, False)
        self.setMinimumWidth(640)
        self.setMinimumHeight(520)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(18, 16, 18, 16)
        layout.setSpacing(12)

        title = QLabel("Edit Folder Monitor")
        title.setObjectName("PageTitle")
        subtitle = QLabel("Set the folder, global filters, notification alerts, and automation rules. Processing order is: monitor filters, rule filters, then action.")
        subtitle.setObjectName("PageSubTitle")
        subtitle.setWordWrap(True)
        layout.addWidget(title)
        layout.addWidget(subtitle)

        card = QFrame()
        card.setObjectName("PageCard")
        card_layout = QVBoxLayout(card)
        card_layout.setContentsMargins(16, 14, 16, 14)
        card_layout.setSpacing(12)

        path_row = QHBoxLayout()
        path_lbl = QLabel("Path:")
        path_lbl.setStyleSheet("color:#9aa0a9; font-weight:600;")
        self.path = QLineEdit(str(getattr(self.monitor, "path", "")))
        self.path.setReadOnly(True)
        path_row.addWidget(path_lbl)
        path_row.addWidget(self.path, 1)
        card_layout.addLayout(path_row)

        events_box = QGroupBox("Notification Alerts")
        events_layout = QVBoxLayout(events_box)
        events_layout.setContentsMargins(12, 10, 12, 10)
        events_layout.setSpacing(8)
        event_help = QLabel("These checkboxes only control desktop/email alerts. File movement is controlled by each Automation Rule below.")
        event_help.setObjectName("PageSubTitle")
        event_help.setWordWrap(True)
        events_layout.addWidget(event_help)
        ev_row = QHBoxLayout()
        ev_row.setSpacing(16)

        self.notify_created = QCheckBox("Created alerts")
        self.notify_created.setChecked(bool(getattr(self.monitor, "notify_created", True)))
        self.notify_modified = QCheckBox("Modified alerts")
        self.notify_modified.setChecked(bool(getattr(self.monitor, "notify_modified", True)))
        self.notify_deleted = QCheckBox("Deleted alerts")
        self.notify_deleted.setChecked(bool(getattr(self.monitor, "notify_deleted", True)))
        self.notify_moved = QCheckBox("Moved alerts")
        self.notify_moved.setChecked(bool(getattr(self.monitor, "notify_moved", True)))
        for cb in (self.notify_created, self.notify_modified, self.notify_deleted, self.notify_moved):
            ev_row.addWidget(cb)
        ev_row.addStretch(1)
        events_layout.addLayout(ev_row)
        card_layout.addWidget(events_box)

        run_box = QGroupBox("Run State")
        run_layout = QVBoxLayout(run_box)
        run_layout.setContentsMargins(12, 10, 12, 10)
        if default_running is None:
            default_running = bool(getattr(self.monitor, "is_running", False))
        self.run_after_save = QCheckBox("Run this monitor after saving")
        self.run_after_save.setChecked(bool(default_running))
        run_layout.addWidget(self.run_after_save)
        card_layout.addWidget(run_box)

        notify_box = QGroupBox("Notifications")
        notify_form = QFormLayout(notify_box)
        notify_form.setLabelAlignment(Qt.AlignRight)
        notify_form.setVerticalSpacing(10)
        notify_form.setHorizontalSpacing(12)

        self.email = QLineEdit(str(getattr(self.monitor, "email_recipient", "")))
        self.email.setPlaceholderText("Optional email recipient")
        notify_form.addRow("Email:", self.email)
        card_layout.addWidget(notify_box)

        filters_box = QGroupBox("File Filters (optional)")
        form = QFormLayout(filters_box)
        form.setLabelAlignment(Qt.AlignRight)
        form.setVerticalSpacing(10)
        form.setHorizontalSpacing(12)

        filter_help = QLabel(
            "Filters run before notifications and rules. Leave a value at 0 or blank to disable it.\n"
            "Use Exclude for folders/files to ignore. Use Filename regex only when you need an advanced name match."
        )
        filter_help.setObjectName("PageSubTitle")
        filter_help.setWordWrap(True)
        form.addRow("", filter_help)

        self.min_size = QSpinBox()
        self.min_size.setRange(0, 10_000_000)
        self.min_size.setSpecialValueText("Disabled")
        self.min_size.setSuffix(" KB")
        self.min_size.setValue(int(getattr(self.monitor, "min_size_kb", 0) or 0))
        self.min_size.setToolTip("Only process files this size or larger. 0 disables the minimum size filter.")
        form.addRow("Minimum size:", self.min_size)

        self.max_size = QSpinBox()
        self.max_size.setRange(0, 10_000_000)
        self.max_size.setSpecialValueText("Disabled")
        self.max_size.setSuffix(" KB")
        self.max_size.setValue(int(getattr(self.monitor, "max_size_kb", 0) or 0))
        self.max_size.setToolTip("Only process files this size or smaller. 0 disables the maximum size filter.")
        form.addRow("Maximum size:", self.max_size)

        self.modified_days = QSpinBox()
        self.modified_days.setRange(0, 36500)
        self.modified_days.setSpecialValueText("Disabled")
        self.modified_days.setSuffix(" days")
        self.modified_days.setValue(int(getattr(self.monitor, "modified_within_days", 0) or 0))
        self.modified_days.setToolTip("Only process files modified within the last N days. 0 disables the date filter.")
        form.addRow("Modified in last:", self.modified_days)

        raw_exclude = getattr(self.monitor, "exclude_patterns", "") or ""
        if isinstance(raw_exclude, (list, tuple, set)):
            exclude_value = "\n".join(str(x) for x in raw_exclude if str(x).strip())
        else:
            exclude_value = str(raw_exclude)
        self.exclude = QTextEdit()
        self.exclude.setFixedHeight(86)
        self.exclude.setPlainText(exclude_value)
        self.exclude.setPlaceholderText("One per line or comma-separated: *.tmp, node_modules/, .git/, */cache/*")
        self.exclude.setToolTip("Files or folders matching these patterns are ignored before notifications and rules run.")
        form.addRow("Exclude patterns:", self.exclude)

        exclude_help = QLabel("Examples: *.tmp ignores temp files; node_modules/ ignores that folder anywhere; */cache/* ignores cache paths.")
        exclude_help.setStyleSheet("color:#9aa0a9;")
        exclude_help.setWordWrap(True)
        form.addRow("", exclude_help)

        self.regex = QLineEdit(str(getattr(self.monitor, "filename_regex", "") or ""))
        self.regex.setPlaceholderText(r"Advanced: ^invoice_.*\.pdf$ or ^client_[0-9]+\.docx$")
        self.regex.setToolTip("Advanced full filename match. Most users should leave this blank and use rule extensions instead.")
        form.addRow("Filename regex:", self.regex)

        regex_help = QLabel("Regex matches the full filename only, not the folder path. Leave blank unless you know you need it.")
        regex_help.setStyleSheet("color:#9aa0a9;")
        regex_help.setWordWrap(True)
        form.addRow("", regex_help)

        card_layout.addWidget(filters_box)

        timing_box = QGroupBox("Timing")
        timing_form = QFormLayout(timing_box)
        timing_form.setLabelAlignment(Qt.AlignRight)
        timing_form.setVerticalSpacing(10)
        timing_form.setHorizontalSpacing(12)

        self.action_delay = QSpinBox()
        self.action_delay.setRange(0, 1440)
        self.action_delay.setSuffix(" min")
        try:
            cur_delay = float(getattr(self.monitor, "action_delay_seconds", 0) or 0)
        except Exception:
            cur_delay = 0.0
        self.action_delay.setValue(int(round(cur_delay / 60.0)))
        timing_form.addRow("Delay actions:", self.action_delay)

        self.stable_wait = QSpinBox()
        self.stable_wait.setRange(0, 3600)
        self.stable_wait.setSuffix(" s")
        try:
            cur_stable = float(getattr(self.monitor, "action_stability_seconds", 0) or 0)
        except Exception:
            cur_stable = 0.0
        self.stable_wait.setValue(int(round(cur_stable)))
        timing_form.addRow("Wait for stable file:", self.stable_wait)

        delay_hint = QLabel("0 = run immediately. Useful for waiting until files finish writing.")
        delay_hint.setStyleSheet("color:#9aa0a9;")
        delay_hint.setWordWrap(True)
        timing_form.addRow("", delay_hint)

        card_layout.addWidget(timing_box)

        # Automation rules (same engine as the original app)
        rules_box = QGroupBox("Automation Rules")
        rules_layout = QVBoxLayout(rules_box)
        rules_layout.setContentsMargins(12, 10, 12, 10)
        rules_layout.setSpacing(10)

        rules_help = QLabel("Rules apply to new file events. To organize files already in this folder, use Run automation rules now from the monitor card.")
        rules_help.setObjectName("PageSubTitle")
        rules_help.setWordWrap(True)
        rules_layout.addWidget(rules_help)

        self.rules_list = QListWidget()
        self.rules_list.setMinimumHeight(170)
        self.rules_list.setStyleSheet(
            "QListWidget{background:#14171c; border:1px solid #232730; border-radius:10px; color:#c8ccd6;}"
        )
        rules_layout.addWidget(self.rules_list)

        btn_row = QHBoxLayout()
        btn_row.setSpacing(10)
        self.btn_rule_preset = QPushButton("Office Organizer")
        self.btn_rule_preset.setToolTip("Create one safe rule that sorts new files into type folders.")
        self.btn_rule_preset.clicked.connect(self._add_office_organizer_clicked)
        btn_row.addWidget(self.btn_rule_preset)
        self.btn_rule_add = QPushButton("Add Rule")
        self.btn_rule_add.clicked.connect(self._add_rule_clicked)
        btn_row.addWidget(self.btn_rule_add)
        self.btn_rule_edit = QPushButton("Edit")
        self.btn_rule_edit.clicked.connect(self._edit_rule_clicked)
        btn_row.addWidget(self.btn_rule_edit)
        self.btn_rule_del = QPushButton("Delete")
        self.btn_rule_del.clicked.connect(self._delete_rule_clicked)
        btn_row.addWidget(self.btn_rule_del)
        self.btn_rule_ai = QPushButton("Add with AI…")
        self.btn_rule_ai.setObjectName("PrimaryButton")
        self.btn_rule_ai.clicked.connect(self._add_rule_ai_clicked)
        btn_row.addWidget(self.btn_rule_ai)
        btn_row.addStretch(1)
        rules_layout.addLayout(btn_row)

        card_layout.addWidget(rules_box)
        self._refresh_rules_list()

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QFrame.NoFrame)
        scroll.setWidget(card)
        layout.addWidget(scroll, 1)

        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        try:
            ok_btn = buttons.button(QDialogButtonBox.Ok)
            if ok_btn:
                ok_btn.setObjectName("PrimaryButton")
        except Exception:
            pass
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)

    def apply_to_monitor(self):
        self.monitor.notify_created = bool(self.notify_created.isChecked())
        self.monitor.notify_modified = bool(self.notify_modified.isChecked())
        self.monitor.notify_deleted = bool(self.notify_deleted.isChecked())
        self.monitor.notify_moved = bool(self.notify_moved.isChecked())
        self.monitor.email_recipient = self.email.text().strip()
        self.monitor.min_size_kb = int(self.min_size.value()) or None
        self.monitor.max_size_kb = int(self.max_size.value()) or None
        self.monitor.modified_within_days = int(self.modified_days.value()) or None
        self.monitor.exclude_patterns = self._exclude_patterns()
        self.monitor.filename_regex = self.regex.text().strip() or None
        try:
            self.monitor.action_delay_seconds = int(self.action_delay.value()) * 60
        except Exception:
            self.monitor.action_delay_seconds = 0
        try:
            self.monitor.action_stability_seconds = int(self.stable_wait.value())
        except Exception:
            self.monitor.action_stability_seconds = 0

    def desired_running(self) -> bool:
        return bool(self.run_after_save.isChecked())

    def _exclude_patterns(self) -> list[str]:
        try:
            text = self.exclude.toPlainText()
        except Exception:
            text = ""
        return [
            part.strip()
            for part in str(text or "").replace("\r", "\n").replace(",", "\n").split("\n")
            if part.strip() and not part.strip().startswith("#")
        ]

    def accept(self):  # noqa: N802
        try:
            top = self.window()
            backend = getattr(top, "backend", None)
            manager = getattr(backend, "monitor_manager", None)
            if manager is not None:
                problems = manager.validate_monitor_config(
                    str(getattr(self.monitor, "path", "") or ""),
                    list(getattr(self.monitor, "rules", []) or []),
                    filename_regex=self.regex.text().strip() or None,
                    min_size_kb=int(self.min_size.value()) or None,
                    max_size_kb=int(self.max_size.value()) or None,
                    exclude_patterns=self._exclude_patterns(),
                )
                if problems:
                    QMessageBox.critical(self, "Monitor Needs Attention", "\n".join(problems[:8]))
                    return
        except Exception:
            pass
        super().accept()

    def _refresh_rules_list(self):
        try:
            self.rules_list.clear()
        except Exception:
            return
        rules = list(getattr(self.monitor, "rules", []) or [])
        for idx, r in enumerate(rules):
            action = (r.get("action_type") or "").strip()
            ev = ",".join(r.get("event_types") or [])
            exts = r.get("file_extensions") or []
            ext_txt = ",".join(exts) if exts else "*"
            interp = (r.get("interpretation") or "").strip()
            line = interp if interp else f"{action} • {ext_txt} • {ev or 'any'}"
            it = QListWidgetItem(line)
            it.setData(Qt.UserRole, idx)
            self.rules_list.addItem(it)

    def _selected_rule_index(self) -> int | None:
        it = self.rules_list.currentItem()
        if not it:
            return None
        try:
            return int(it.data(Qt.UserRole))
        except Exception:
            return None

    def _add_rule_clicked(self):
        dlg = _QtRuleEditorDialog(self, rule=None)
        if dlg.exec() != QDialog.Accepted:
            return
        try:
            self.monitor.rules.append(dlg.rule_dict())
        except Exception:
            return
        self._refresh_rules_list()

    def _edit_rule_clicked(self):
        idx = self._selected_rule_index()
        if idx is None:
            return
        try:
            rule = dict(self.monitor.rules[idx])
        except Exception:
            return
        dlg = _QtRuleEditorDialog(self, rule=rule)
        if dlg.exec() != QDialog.Accepted:
            return
        try:
            self.monitor.rules[idx] = dlg.rule_dict()
        except Exception:
            return
        self._refresh_rules_list()

    def _delete_rule_clicked(self):
        idx = self._selected_rule_index()
        if idx is None:
            return
        if QMessageBox.question(self, "Delete Rule", "Delete the selected rule?") != QMessageBox.Yes:
            return
        try:
            del self.monitor.rules[idx]
        except Exception:
            return
        self._refresh_rules_list()

    def _add_office_organizer_clicked(self):
        dest = QFileDialog.getExistingDirectory(self, "Choose where organized files should go")
        if not dest:
            return
        try:
            monitor_root = Path(str(getattr(self.monitor, "path", "") or "")).resolve()
            dest_root = Path(dest).resolve()
            if dest_root == monitor_root or monitor_root in dest_root.parents:
                QMessageBox.warning(
                    self,
                    "Choose Another Destination",
                    "Choose a destination outside the monitored folder. This prevents files from being moved into a folder that is also being watched.",
                )
                return
        except Exception:
            pass

        rule = {
            "event_types": ["created"],
            "file_extensions": ["*"],
            "action_type": "organize",
            "action_params": {
                "destination": dest,
                "organize_by": "type",
                "handle_duplicates": "rename",
            },
            "interpretation": "Office Organizer: sort new files by type into the selected destination",
        }

        rules = list(getattr(self.monitor, "rules", []) or [])
        if rules:
            choice = QMessageBox.question(
                self,
                "Office Organizer",
                "Replace the current rules with one clean Office Organizer rule?\n\nChoose No to add it without removing existing rules.",
                QMessageBox.Yes | QMessageBox.No | QMessageBox.Cancel,
                QMessageBox.Yes,
            )
            if choice == QMessageBox.Cancel:
                return
            if choice == QMessageBox.Yes:
                self.monitor.rules = [rule]
            else:
                self.monitor.rules.append(rule)
        else:
            self.monitor.rules.append(rule)
        self._refresh_rules_list()

    def _add_rule_ai_clicked(self):
        top = self.window()
        backend = getattr(top, "backend", None)
        icons = getattr(top, "icons", None)
        if not backend or not icons:
            QMessageBox.warning(self, "AI", "AI Rule Builder is not available.")
            return
        mid = getattr(self.monitor, "monitor_id", None)
        dlg = _QtAiRuleBuilderDialog(top, backend=backend, icons=icons, initial_monitor_id=str(mid) if mid else None)
        if dlg.exec() == QDialog.Accepted:
            self._refresh_rules_list()


class _EditFtpMonitorDialog(QDialog):
    def __init__(self, parent: QWidget, *, monitor):
        super().__init__(parent)
        self.monitor = monitor
        self._main = parent
        self.setWindowTitle("Edit FTP Monitor")
        self.setModal(True)
        self.setWindowFlag(Qt.WindowContextHelpButtonHint, False)
        self.setMinimumWidth(560)
        self.setMinimumHeight(520)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(18, 16, 18, 16)
        layout.setSpacing(12)

        title = QLabel("Edit FTP Monitor")
        title.setObjectName("PageTitle")
        subtitle = QLabel("Update connection details and polling interval. Changes apply immediately.")
        subtitle.setObjectName("PageSubTitle")
        subtitle.setWordWrap(True)
        layout.addWidget(title)
        layout.addWidget(subtitle)

        card = QFrame()
        card.setObjectName("PageCard")
        card_layout = QVBoxLayout(card)
        card_layout.setContentsMargins(16, 14, 16, 14)
        card_layout.setSpacing(12)

        connection_box = QGroupBox("Connection")
        form = QFormLayout(connection_box)
        form.setLabelAlignment(Qt.AlignRight)
        form.setVerticalSpacing(10)
        form.setHorizontalSpacing(12)

        self.host = QLineEdit(str(getattr(self.monitor, "host", "")))
        form.addRow("Host:", self.host)

        self.port = QSpinBox()
        self.port.setRange(1, 65535)
        self.port.setValue(int(getattr(self.monitor, "port", 21) or 21))
        form.addRow("Port:", self.port)

        self.username = QLineEdit(str(getattr(self.monitor, "username", "")))

        self.password = QLineEdit(str(getattr(self.monitor, "password", "")))
        self.password.setEchoMode(QLineEdit.Password)

        remote_row = QHBoxLayout()
        self.remote_path = QLineEdit(str(getattr(self.monitor, "remote_path", "/") or "/"))
        remote_row.addWidget(self.remote_path, 1)
        remote_browse = QPushButton("Browse")
        remote_browse.clicked.connect(self._browse_remote_path)
        remote_row.addWidget(remote_browse)
        form.addRow("Remote path:", remote_row)

        self.use_tls = QCheckBox("Use TLS (FTPS)")
        self.use_tls.setChecked(bool(getattr(self.monitor, "use_tls", False)))
        self.use_tls.toggled.connect(self._toggle_tls_options)
        form.addRow("", self.use_tls)

        self.tls_implicit = QCheckBox("Implicit FTPS (port 990)")
        self.tls_implicit.setChecked(bool(getattr(self.monitor, "tls_implicit", False)))
        self.tls_implicit.setEnabled(bool(self.use_tls.isChecked()))
        self.tls_implicit.toggled.connect(self._maybe_set_implicit_port)
        form.addRow("", self.tls_implicit)

        self.passive_mode = QCheckBox("Passive mode (recommended)")
        self.passive_mode.setChecked(bool(getattr(self.monitor, "passive_mode", True)))
        form.addRow("", self.passive_mode)

        self.ftp_encoding = QLineEdit(str(getattr(self.monitor, "encoding", "utf-8") or "utf-8"))
        self.ftp_encoding.setPlaceholderText("Server encoding (e.g. utf-8)")
        form.addRow("Encoding:", self.ftp_encoding)

        self.poll_interval = QSpinBox()
        self.poll_interval.setRange(5, 3600)
        self.poll_interval.setValue(int(getattr(self.monitor, "poll_interval", 30) or 30))
        form.addRow("Poll interval (s):", self.poll_interval)

        card_layout.addWidget(connection_box)

        auth_box = QGroupBox("Authentication")
        auth_form = QFormLayout(auth_box)
        auth_form.setLabelAlignment(Qt.AlignRight)
        auth_form.setVerticalSpacing(10)
        auth_form.setHorizontalSpacing(12)

        auth_form.addRow("Username:", self.username)
        auth_form.addRow("Password:", self.password)
        card_layout.addWidget(auth_box)

        sync_box = QGroupBox("Local Sync (optional)")
        sync_form = QFormLayout(sync_box)
        sync_form.setLabelAlignment(Qt.AlignRight)
        sync_form.setVerticalSpacing(10)
        sync_form.setHorizontalSpacing(12)

        sync_row = QHBoxLayout()
        self.local_sync_dir = QLineEdit(str(getattr(self.monitor, "local_sync_dir", "") or ""))
        self.local_sync_dir.setPlaceholderText("Optional local folder to download new/changed files")
        sync_row.addWidget(self.local_sync_dir, 1)
        browse = QPushButton("Browse")
        browse.clicked.connect(self._browse_local_dir)
        sync_row.addWidget(browse)
        sync_form.addRow("Local folder:", sync_row)

        self.two_way_sync = QCheckBox("Two-way sync (upload + download)")
        self.two_way_sync.setChecked(bool(getattr(self.monitor, "two_way_sync", False)))
        sync_form.addRow("", self.two_way_sync)

        self.sync_subfolders = QCheckBox("Sync subfolders")
        self.sync_subfolders.setChecked(bool(getattr(self.monitor, "sync_subfolders", False)))
        sync_form.addRow("", self.sync_subfolders)

        self.download_on_created = QCheckBox("Download new files")
        self.download_on_created.setChecked(bool(getattr(self.monitor, "download_on_created", True)))
        sync_form.addRow("", self.download_on_created)

        self.download_on_modified = QCheckBox("Download modified files")
        self.download_on_modified.setChecked(bool(getattr(self.monitor, "download_on_modified", True)))
        sync_form.addRow("", self.download_on_modified)

        self.delete_local_on_deleted = QCheckBox("Delete local file when remote is deleted")
        self.delete_local_on_deleted.setChecked(bool(getattr(self.monitor, "delete_local_on_deleted", False)))
        sync_form.addRow("", self.delete_local_on_deleted)

        self.overwrite_local = QCheckBox("Overwrite local files")
        self.overwrite_local.setChecked(bool(getattr(self.monitor, "overwrite_local", False)))
        sync_form.addRow("", self.overwrite_local)

        self.allowed_exts = QLineEdit(",".join(list(getattr(self.monitor, "allowed_extensions", []) or [])))
        self.allowed_exts.setPlaceholderText("Optional extensions: jpg,png,pdf,mp3")
        sync_form.addRow("Extensions:", self.allowed_exts)

        card_layout.addWidget(sync_box)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QFrame.NoFrame)
        scroll.setWidget(card)
        layout.addWidget(scroll, 1)

        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        try:
            ok_btn = buttons.button(QDialogButtonBox.Ok)
            if ok_btn:
                ok_btn.setObjectName("PrimaryButton")
        except Exception:
            pass
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)

    def accept(self):
        top = getattr(self, "_main", None) or self.parent()
        backend = getattr(top, "backend", None)
        ftp_manager = getattr(getattr(backend, "monitor_manager", None), "ftp_manager", None)
        if ftp_manager and hasattr(ftp_manager, "validate_ftp_config"):
            problems = ftp_manager.validate_ftp_config(
                host=self.host.text().strip(),
                username=self.username.text().strip(),
                password=self.password.text(),
                remote_path=self.remote_path.text().strip(),
                port=int(self.port.value()),
                poll_interval=int(self.poll_interval.value()),
                local_sync_dir=self.local_sync_dir.text().strip() or None,
                allowed_extensions=[
                    e.strip().lower().lstrip(".")
                    for e in (self.allowed_exts.text() or "").split(",")
                    if e.strip()
                ],
            )
            if problems:
                QMessageBox.warning(self, "Invalid FTP Monitor", "\n".join(problems))
                return
        super().accept()

    def apply_to_monitor(self):
        was_running = bool(getattr(self.monitor, "is_running", False))
        if was_running:
            try:
                self.monitor.stop()
            except Exception:
                pass
        self.monitor.host = self.host.text().strip()
        self.monitor.port = int(self.port.value())
        self.monitor.username = self.username.text().strip()
        self.monitor.password = self.password.text()
        self.monitor.remote_path = self.remote_path.text().strip()
        self.monitor.use_tls = bool(self.use_tls.isChecked())
        self.monitor.tls_implicit = bool(self.tls_implicit.isChecked())
        self.monitor.passive_mode = bool(self.passive_mode.isChecked())
        self.monitor.encoding = (self.ftp_encoding.text() or "").strip() or "utf-8"
        self.monitor.poll_interval = int(self.poll_interval.value())
        self.monitor.local_sync_dir = self.local_sync_dir.text().strip() or None
        self.monitor.two_way_sync = bool(self.two_way_sync.isChecked())
        self.monitor.sync_subfolders = bool(self.sync_subfolders.isChecked())
        self.monitor.download_on_created = bool(self.download_on_created.isChecked())
        self.monitor.download_on_modified = bool(self.download_on_modified.isChecked())
        self.monitor.delete_local_on_deleted = bool(self.delete_local_on_deleted.isChecked())
        self.monitor.overwrite_local = bool(self.overwrite_local.isChecked())
        self.monitor.allowed_extensions = [
            e.strip().lower().lstrip(".") for e in (self.allowed_exts.text() or "").split(",") if e.strip()
        ]
        if was_running:
            try:
                self.monitor.start()
            except Exception:
                pass

    def _browse_local_dir(self):
        folder = QFileDialog.getExistingDirectory(self, "Select local sync folder")
        if folder:
            self.local_sync_dir.setText(folder)

    def _browse_remote_path(self):
        top = getattr(self, "_main", None) or self.parent()
        backend = getattr(top, "backend", None)
        ftp_mgr = getattr(getattr(getattr(backend, "monitor_manager", None), "ftp_manager", None), "list_remote_dirs", None)
        if not backend or not ftp_mgr:
            QMessageBox.warning(self, "FTP", "FTP remote browser is not available.")
            return

        host = self.host.text().strip()
        if not host:
            QMessageBox.warning(self, "FTP", "Enter Host first.")
            return

        dlg = _FtpRemoteBrowserDialog(
            self,
            ftp_manager=backend.monitor_manager.ftp_manager,
            host=host,
            username=self.username.text().strip(),
            password=self.password.text(),
            remote_path=self.remote_path.text().strip() or "/",
            port=int(self.port.value() or 21),
            use_tls=bool(self.use_tls.isChecked()),
            tls_implicit=bool(self.tls_implicit.isChecked()),
            passive_mode=bool(self.passive_mode.isChecked()),
            encoding=(self.ftp_encoding.text() or "").strip() or "utf-8",
        )
        if dlg.exec() == QDialog.Accepted:
            self.remote_path.setText(dlg.selected_path())

    def _toggle_tls_options(self, enabled: bool):
        enabled = bool(enabled)
        try:
            self.tls_implicit.setEnabled(enabled)
            if not enabled:
                self.tls_implicit.setChecked(False)
        except Exception:
            pass

    def _maybe_set_implicit_port(self, enabled: bool):
        if not enabled:
            return
        try:
            if int(self.port.value()) == 21:
                self.port.setValue(990)
        except Exception:
            pass


class _FtpRemoteBrowserDialog(QDialog):
    class _Worker(QObject):
        finished = Signal(list)
        error = Signal(str)

        def __init__(
            self,
            *,
            ftp_manager,
            host: str,
            username: str,
            password: str,
            remote_path: str,
            port: int,
            use_tls: bool,
            tls_implicit: bool,
            passive_mode: bool,
            encoding: str,
        ):
            super().__init__()
            self.ftp_manager = ftp_manager
            self.host = host
            self.username = username
            self.password = password
            self.remote_path = remote_path
            self.port = port
            self.use_tls = use_tls
            self.tls_implicit = tls_implicit
            self.passive_mode = passive_mode
            self.encoding = encoding

        def run(self):
            try:
                dirs = self.ftp_manager.list_remote_dirs(
                    self.host,
                    self.username,
                    self.password,
                    self.remote_path,
                    port=self.port,
                    use_tls=self.use_tls,
                    tls_implicit=self.tls_implicit,
                    passive_mode=self.passive_mode,
                    encoding=self.encoding,
                )
                self.finished.emit(list(dirs or []))
            except Exception as e:
                self.error.emit(str(e))

    def __init__(
        self,
        parent: QWidget,
        *,
        ftp_manager,
        host: str,
        username: str,
        password: str,
        remote_path: str,
        port: int,
        use_tls: bool,
        tls_implicit: bool,
        passive_mode: bool,
        encoding: str,
    ):
        super().__init__(parent)
        self.setWindowTitle("Browse FTP Folder")
        self.setModal(True)
        self.setWindowFlag(Qt.WindowContextHelpButtonHint, False)
        self.setMinimumWidth(520)
        try:
            self.setStyleSheet(_qt_modern_dialog_stylesheet())
        except Exception:
            pass

        self._ftp_manager = ftp_manager
        self._host = host
        self._username = username
        self._password = password
        self._port = int(port or 21)
        self._use_tls = bool(use_tls)
        self._tls_implicit = bool(tls_implicit)
        self._passive_mode = bool(passive_mode)
        self._encoding = (encoding or "").strip() or "utf-8"

        self._selected = "/"
        self._thread: QThread | None = None
        self._worker: _FtpRemoteBrowserDialog._Worker | None = None
        self._closing = False
        self._refresh_seq = 0

        layout = QVBoxLayout(self)
        layout.setContentsMargins(18, 16, 18, 16)
        layout.setSpacing(12)

        header = QFrame()
        header.setObjectName("DialogHeader")
        header_l = QVBoxLayout(header)
        header_l.setContentsMargins(16, 14, 16, 14)
        header_l.setSpacing(8)

        title = QLabel("Select Remote Folder")
        title.setObjectName("DialogTitle")
        tls_label = "FTP"
        if self._use_tls:
            tls_label = "FTPS (Implicit)" if self._tls_implicit else "FTPS (Explicit)"
        subtitle = QLabel(f"{tls_label} • {self._host}:{self._port}")
        subtitle.setObjectName("DialogSubtitle")
        header_l.addWidget(title)
        header_l.addWidget(subtitle)

        path_row = QHBoxLayout()
        self.path_edit = QLineEdit()
        self.path_edit.setReadOnly(True)
        path_row.addWidget(self.path_edit, 1)
        self.btn_up = QPushButton("Up")
        self.btn_up.clicked.connect(self._go_up)
        path_row.addWidget(self.btn_up)
        self.btn_refresh = QPushButton("Refresh")
        self.btn_refresh.clicked.connect(self._refresh)
        path_row.addWidget(self.btn_refresh)
        header_l.addLayout(path_row)

        layout.addWidget(header)

        self.list = QListWidget()
        self.list.setSelectionMode(QAbstractItemView.SingleSelection)
        self.list.itemDoubleClicked.connect(lambda _it: self._enter_selected())
        layout.addWidget(self.list, 1)

        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        try:
            ok_btn = buttons.button(QDialogButtonBox.Ok)
            if ok_btn:
                ok_btn.setObjectName("PrimaryButton")
        except Exception:
            pass
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)

        self._set_path(remote_path)
        self._refresh()

    def closeEvent(self, event):  # noqa: N802
        self._closing = True
        self._stop_thread()
        super().closeEvent(event)

    def reject(self):  # noqa: N802
        self._closing = True
        self._stop_thread()
        super().reject()

    def accept(self):  # noqa: N802
        self._closing = True
        self._stop_thread()
        super().accept()

    def selected_path(self) -> str:
        return str(self._selected or "/")

    def _norm_path(self, p: str) -> str:
        p = str(p or "/").replace("\\", "/").strip() or "/"
        while "//" in p:
            p = p.replace("//", "/")
        if not p.startswith("/"):
            p = "/" + p
        if len(p) > 1 and p.endswith("/"):
            p = p[:-1]
        return p or "/"

    def _set_path(self, p: str):
        self._selected = self._norm_path(p)
        self.path_edit.setText(self._selected)
        self.btn_up.setEnabled(self._selected != "/")

    def _go_up(self):
        cur = self._selected or "/"
        if cur == "/":
            return
        parent = "/".join(cur.split("/")[:-1]) or "/"
        self._set_path(parent)
        self._refresh()

    def _enter_selected(self):
        it = self.list.currentItem()
        if not it:
            return
        name = str(it.text() or "").strip()
        if not name or name.startswith("("):
            return
        cur = self._selected or "/"
        nxt = f"/{name}" if cur == "/" else f"{cur}/{name}"
        self._set_path(nxt)
        self._refresh()

    def _stop_thread(self):
        th = getattr(self, "_thread", None)
        if not th:
            return
        try:
            th.quit()
        except Exception:
            pass
        try:
            th.wait(250)
        except Exception:
            pass
        self._thread = None
        self._worker = None

    def _refresh(self):
        if self._closing:
            return
        self.list.clear()
        self.list.addItem(QListWidgetItem("Loading..."))
        self.btn_up.setEnabled(False)
        self.btn_refresh.setEnabled(False)
        self.path_edit.setEnabled(False)

        self._stop_thread()

        self._refresh_seq += 1
        seq = int(self._refresh_seq)

        thread = QThread(self)
        worker = _FtpRemoteBrowserDialog._Worker(
            ftp_manager=self._ftp_manager,
            host=self._host,
            username=self._username,
            password=self._password,
            remote_path=self._selected,
            port=self._port,
            use_tls=self._use_tls,
            tls_implicit=self._tls_implicit,
            passive_mode=self._passive_mode,
            encoding=self._encoding,
        )
        worker.moveToThread(thread)
        self._thread = thread
        self._worker = worker

        def _done(dirs: list):
            if self._closing or seq != self._refresh_seq:
                return
            self.list.clear()
            for d in dirs or []:
                self.list.addItem(QListWidgetItem(str(d)))
            if not dirs:
                self.list.addItem(QListWidgetItem("(no subfolders)"))
            self.btn_up.setEnabled(self._selected != "/")
            self.btn_refresh.setEnabled(True)
            self.path_edit.setEnabled(True)

        def _err(msg: str):
            if self._closing or seq != self._refresh_seq:
                return
            self.list.clear()
            self.list.addItem(QListWidgetItem("(failed to load folders)"))
            self.btn_up.setEnabled(self._selected != "/")
            self.btn_refresh.setEnabled(True)
            self.path_edit.setEnabled(True)
            try:
                QMessageBox.critical(self, "FTP Browse", msg)
            except Exception:
                pass

        worker.finished.connect(_done)
        worker.error.connect(_err)
        thread.started.connect(worker.run)
        worker.finished.connect(thread.quit)
        worker.error.connect(thread.quit)
        thread.finished.connect(thread.deleteLater)
        thread.finished.connect(worker.deleteLater)
        thread.finished.connect(lambda: self._cleanup_thread(seq))
        thread.start()

    def _cleanup_thread(self, seq: int):
        if seq != self._refresh_seq:
            return
        self._thread = None
        self._worker = None


class _QtRuleEditorDialog(QDialog):
    def __init__(self, parent: QWidget, *, rule: dict | None):
        super().__init__(parent)
        self._rule = dict(rule) if rule else {}
        self.setWindowTitle("Automation Rule")
        self.setModal(True)
        self.setWindowFlag(Qt.WindowContextHelpButtonHint, False)
        self.setMinimumWidth(620)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(14, 14, 14, 14)
        layout.setSpacing(10)

        form_card = QFrame()
        form_card.setObjectName("PageCard")
        card_layout = QVBoxLayout(form_card)
        card_layout.setContentsMargins(14, 12, 14, 12)
        card_layout.setSpacing(12)

        row = QHBoxLayout()
        row.setSpacing(14)

        events_box = QGroupBox("Rule Triggers")
        ev_outer = QVBoxLayout(events_box)
        ev_outer.setContentsMargins(10, 8, 10, 8)
        ev_outer.setSpacing(8)
        ev_help = QLabel("These events start this action. For moving or organizing files, use Created only.")
        ev_help.setStyleSheet("color:#9aa0a9;")
        ev_help.setWordWrap(True)
        ev_outer.addWidget(ev_help)
        ev = QHBoxLayout()
        ev.setSpacing(10)
        self.ev_created = QCheckBox("Created")
        self.ev_modified = QCheckBox("Modified")
        self.ev_deleted = QCheckBox("Deleted")
        self.ev_moved = QCheckBox("Moved")
        for cb in (self.ev_created, self.ev_modified, self.ev_deleted, self.ev_moved):
            ev.addWidget(cb)
        self.btn_created_only = QPushButton("Use Created Only")
        self.btn_created_only.setToolTip("Recommended for move and organize rules.")
        self.btn_created_only.clicked.connect(self._set_created_only)
        ev.addWidget(self.btn_created_only)
        ev.addStretch(1)
        ev_outer.addLayout(ev)
        row.addWidget(events_box, 1)

        filt_box = QGroupBox("Filters")
        fl = QFormLayout(filt_box)
        fl.setLabelAlignment(Qt.AlignRight)
        rule_filter_help = QLabel(
            "These filters apply only to this rule. The monitor's File Filters run first, then this rule checks events, extensions, and name regex."
        )
        rule_filter_help.setStyleSheet("color:#9aa0a9;")
        rule_filter_help.setWordWrap(True)
        fl.addRow("", rule_filter_help)
        self.exts = QLineEdit()
        self.exts.setPlaceholderText("Examples: pdf, docx, jpg or * for all files")
        fl.addRow("Extensions:", self.exts)
        self.pattern = QLineEdit()
        self.pattern.setPlaceholderText(r"Advanced: ^invoice_.*\.pdf$ or ^client_[0-9]+\.docx$")
        fl.addRow("Name regex:", self.pattern)
        row.addWidget(filt_box, 1)

        card_layout.addLayout(row)

        action_box = QGroupBox("Action")
        act_form = QFormLayout(action_box)
        act_form.setLabelAlignment(Qt.AlignRight)
        self._action_rows = {}
        self.action_type = QComboBox()
        self.action_type.addItems(["copy", "move", "organize", "rename", "delete", "archive", "execute", "cloud_upload"])
        act_form.addRow("Type:", self.action_type)

        dest_widget = QWidget()
        dest_row = QHBoxLayout(dest_widget)
        dest_row.setContentsMargins(0, 0, 0, 0)
        self.destination = QLineEdit()
        self.destination.setPlaceholderText("Destination folder (for move/copy/organize/archive)")
        dest_row.addWidget(self.destination, 1)
        self.btn_pick_destination = QPushButton("Browse")
        self.btn_pick_destination.clicked.connect(self._browse_dest)
        dest_row.addWidget(self.btn_pick_destination)
        self._add_action_row(act_form, "destination", "Destination:", dest_widget)

        self.dup = QComboBox()
        self.dup.addItems(["rename", "skip", "overwrite"])
        self._add_action_row(act_form, "duplicates", "Duplicates:", self.dup)

        self.organize_by = QComboBox()
        self.organize_by.addItems(["extension", "date", "type"])
        self._add_action_row(act_form, "organize_by", "Organize by:", self.organize_by)

        self.rename_pattern = QLineEdit()
        self.rename_pattern.setPlaceholderText("{name}_{date} (extension is appended automatically)")
        self._add_action_row(act_form, "rename_pattern", "Rename pattern:", self.rename_pattern)

        self.archive_name = QLineEdit()
        self.archive_name.setPlaceholderText("archive_{date}.zip")
        self._add_action_row(act_form, "archive_name", "Archive name:", self.archive_name)

        self.exec_cmd = QLineEdit()
        self.exec_cmd.setPlaceholderText("Command to execute (use {path} placeholder)")
        self._add_action_row(act_form, "command", "Command:", self.exec_cmd)

        # Cloud Upload (monitor → cloud)
        self.cloud_provider = QComboBox()
        self.cloud_provider.addItems(["onedrive", "gdrive"])
        self._add_action_row(act_form, "cloud_provider", "Cloud provider:", self.cloud_provider)

        self.cloud_remote_base = QLineEdit("Managed File Sync")
        self.cloud_remote_base.setPlaceholderText("Remote base folder, for example Managed File Sync/Backups")
        self._add_action_row(act_form, "cloud_remote_base", "Remote base:", self.cloud_remote_base)

        self.cloud_include_sub = QCheckBox("Preserve subfolders (relative to monitor root)")
        self.cloud_include_sub.setChecked(True)
        self._add_action_row(act_form, "cloud_include_sub", "", self.cloud_include_sub)

        card_layout.addWidget(action_box)

        preview_box = QGroupBox("What Will Happen")
        preview_layout = QVBoxLayout(preview_box)
        preview_layout.setContentsMargins(12, 10, 12, 10)
        self.preview_text = QLabel("")
        self.preview_text.setWordWrap(True)
        self.preview_text.setStyleSheet("color:#c8ccd6;")
        preview_layout.addWidget(self.preview_text)
        card_layout.addWidget(preview_box)

        layout.addWidget(form_card)

        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        try:
            ok_btn = buttons.button(QDialogButtonBox.Ok)
            if ok_btn:
                ok_btn.setObjectName("PrimaryButton")
        except Exception:
            pass
        buttons.accepted.connect(self._validate_and_accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)

        self._load_rule_into_ui()
        self.action_type.currentTextChanged.connect(lambda _t: self._sync_action_fields())
        for cb in (self.ev_created, self.ev_modified, self.ev_deleted, self.ev_moved):
            cb.stateChanged.connect(lambda _state=None: self._sync_event_controls())
        for edit in (self.exts, self.pattern, self.destination, self.rename_pattern, self.archive_name, self.exec_cmd, self.cloud_remote_base):
            edit.textChanged.connect(lambda _text=None: self._update_rule_preview())
        self.dup.currentTextChanged.connect(lambda _text=None: self._update_rule_preview())
        self.organize_by.currentTextChanged.connect(lambda _text=None: self._update_rule_preview())
        self.cloud_provider.currentTextChanged.connect(lambda _text=None: self._update_rule_preview())
        self.cloud_include_sub.stateChanged.connect(lambda _state=None: self._update_rule_preview())
        self._sync_action_fields()
        self._update_rule_preview()

    def _add_action_row(self, form: QFormLayout, key: str, label_text: str, widget: QWidget):
        label = QLabel(label_text)
        form.addRow(label, widget)
        self._action_rows[key] = (label, widget)

    def _set_action_row_visible(self, key: str, visible: bool):
        row = self._action_rows.get(key)
        if not row:
            return
        label, widget = row
        label.setVisible(visible)
        widget.setVisible(visible)

    def _set_created_only(self):
        self.ev_created.setChecked(True)
        self.ev_modified.setChecked(False)
        self.ev_deleted.setChecked(False)
        self.ev_moved.setChecked(False)
        self._update_rule_preview()

    def _sync_event_controls(self):
        routing_action = _qt_is_file_routing_action(self.action_type.currentText())
        if routing_action:
            for cb, checked in (
                (self.ev_created, True),
                (self.ev_modified, False),
                (self.ev_deleted, False),
                (self.ev_moved, False),
            ):
                old = cb.blockSignals(True)
                cb.setChecked(checked)
                cb.blockSignals(old)
        for cb in (self.ev_modified, self.ev_deleted, self.ev_moved):
            cb.setEnabled(not routing_action)
            cb.setToolTip("Move/organize rules use Created only to avoid repeated moves while files are changing." if routing_action else "")
        self.ev_created.setEnabled(True)
        self.btn_created_only.setVisible(routing_action)
        self._update_rule_preview()

    def _browse_dest(self):
        folder = QFileDialog.getExistingDirectory(self, "Select destination folder")
        if folder:
            self.destination.setText(folder)

    def _load_rule_into_ui(self):
        evs = set(self._rule.get("event_types") or [])
        self.ev_created.setChecked("created" in evs)
        self.ev_modified.setChecked("modified" in evs)
        self.ev_deleted.setChecked("deleted" in evs)
        self.ev_moved.setChecked("moved" in evs)

        exts = self._rule.get("file_extensions") or []
        self.exts.setText(_qt_extensions_to_input(exts if isinstance(exts, list) else [str(exts)]))
        self.pattern.setText(self._rule.get("name_pattern") or "")

        self.action_type.setCurrentText(self._rule.get("action_type") or "move")
        params = dict(self._rule.get("action_params") or {})
        self.destination.setText(str(params.get("destination") or ""))
        self.dup.setCurrentText(str(params.get("handle_duplicates") or "rename"))
        self.organize_by.setCurrentText(str(params.get("organize_by") or "extension"))
        self.rename_pattern.setText(str(params.get("pattern") or "{name}"))
        self.archive_name.setText(str(params.get("archive_name") or "archive_{date}.zip"))
        self.exec_cmd.setText(str(params.get("command") or ""))
        self.cloud_provider.setCurrentText(str(params.get("provider") or "onedrive").strip().lower() or "onedrive")
        self.cloud_remote_base.setText(str(params.get("remote_base") or "Managed File Sync"))
        try:
            self.cloud_include_sub.setChecked(bool(params.get("include_subfolders", True)))
        except Exception:
            self.cloud_include_sub.setChecked(True)

    def _sync_action_fields(self):
        t = self.action_type.currentText()
        needs_dest = t in {"copy", "move", "organize", "archive"}
        visible_rows = {
            "destination": needs_dest,
            "duplicates": t in {"copy", "move"},
            "organize_by": t == "organize",
            "rename_pattern": t == "rename",
            "archive_name": t == "archive",
            "command": t == "execute",
            "cloud_provider": t == "cloud_upload",
            "cloud_remote_base": t == "cloud_upload",
            "cloud_include_sub": t == "cloud_upload",
        }
        for key, visible in visible_rows.items():
            self._set_action_row_visible(key, visible)
        self.destination.setEnabled(needs_dest)
        self.btn_pick_destination.setEnabled(needs_dest)
        self.dup.setEnabled(t in {"copy", "move"})
        self.organize_by.setEnabled(t == "organize")
        self.rename_pattern.setEnabled(t == "rename")
        self.archive_name.setEnabled(t == "archive")
        self.exec_cmd.setEnabled(t == "execute")
        self.cloud_provider.setEnabled(t == "cloud_upload")
        self.cloud_remote_base.setEnabled(t == "cloud_upload")
        self.cloud_include_sub.setEnabled(t == "cloud_upload")
        self._sync_event_controls()
        self._update_rule_preview()

    def _draft_rule_dict(self) -> dict:
        evs = []
        if self.ev_created.isChecked():
            evs.append("created")
        if self.ev_modified.isChecked():
            evs.append("modified")
        if self.ev_deleted.isChecked():
            evs.append("deleted")
        if self.ev_moved.isChecked():
            evs.append("moved")

        exts_raw = (self.exts.text() or "").strip()
        exts = _qt_normalize_extensions(exts_raw)

        rule = {
            "event_types": ["created"] if _qt_is_file_routing_action(self.action_type.currentText()) else (evs or ["created"]),
            "file_extensions": exts,
            "action_type": self.action_type.currentText(),
            "action_params": {},
        }
        pat = (self.pattern.text() or "").strip()
        if pat:
            rule["name_pattern"] = pat

        params = {}
        t = rule["action_type"]
        if t in {"copy", "move"}:
            params["destination"] = self.destination.text().strip()
            params["handle_duplicates"] = self.dup.currentText()
        elif t == "organize":
            params["destination"] = self.destination.text().strip()
            params["organize_by"] = self.organize_by.currentText()
            params["handle_duplicates"] = "rename"
        elif t == "rename":
            params["pattern"] = (self.rename_pattern.text() or "{name}").strip()
        elif t == "delete":
            params["use_recycle_bin"] = True
            params["silent"] = True
        elif t == "archive":
            params["destination"] = self.destination.text().strip()
            params["archive_name"] = (self.archive_name.text() or "archive_{date}.zip").strip()
        elif t == "execute":
            params["command"] = self.exec_cmd.text().strip()
        elif t == "cloud_upload":
            params["provider"] = self.cloud_provider.currentText().strip().lower()
            params["remote_base"] = (self.cloud_remote_base.text() or "Managed File Sync").strip()
            params["include_subfolders"] = bool(self.cloud_include_sub.isChecked())
        rule["action_params"] = params
        return rule

    def _update_rule_preview(self):
        try:
            self.preview_text.setText(_qt_rule_action_preview(self._draft_rule_dict()))
        except Exception:
            pass

    def _validate_and_accept(self):
        t = self.action_type.currentText()
        pat = (self.pattern.text() or "").strip()
        if pat:
            try:
                re.compile(pat)
            except Exception as e:
                QMessageBox.warning(self, "Invalid Name Regex", f"Filename regex is invalid:\n{e}")
                return
        exts = _qt_normalize_extensions((self.exts.text() or "").strip())
        if "*" not in exts and ".*" not in exts:
            for ext in exts:
                clean_ext = str(ext).lstrip(".")
                if any(ch in clean_ext for ch in "/\\:*?\"<>|\r\n"):
                    QMessageBox.warning(self, "Invalid Extension", f"Extension filter is invalid: {clean_ext}")
                    return
            self.exts.setText(_qt_extensions_to_input(exts))
        if t in {"copy", "move", "organize", "archive"} and not self.destination.text().strip():
            QMessageBox.warning(self, "Missing Destination", "Destination folder is required for this action.")
            return
        if t == "execute" and not self.exec_cmd.text().strip():
            QMessageBox.warning(self, "Missing Command", "Command is required for execute.")
            return
        if t == "cloud_upload" and not self.cloud_remote_base.text().strip():
            QMessageBox.warning(self, "Missing Remote Base", "Remote base folder is required for cloud upload.")
            return
        self.accept()

    def rule_dict(self) -> dict:
        rule = self._draft_rule_dict()
        t = rule["action_type"]
        exts = rule.get("file_extensions") or ["*"]
        if not rule.get("interpretation"):
            rule["interpretation"] = f"{t} ({', '.join(exts)})"
        return rule


def _is_ai_rule(rule: dict) -> bool:
    if "ai_prompt" in rule:
        return True
    if rule.get("name_pattern"):
        return True
    if rule.get("action_type") == "organize" and "organize_by" in (rule.get("action_params") or {}):
        return True
    if rule.get("action_type") in ["move", "copy"] and "handle_duplicates" in (rule.get("action_params") or {}):
        return True
    return False


class _QtAiRuleCard(QFrame):
    def __init__(self, *, path: str, rule: dict):
        super().__init__()
        self.setObjectName("PageCard")
        self._head: QLabel | None = None
        self._prompt: QLabel | None = None
        self._desc: QLabel | None = None
        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 14, 16, 14)
        layout.setSpacing(6)

        from pathlib import Path

        head = QLabel(f"📁 {Path(path).name}")
        self._head = head
        layout.addWidget(head)

        prompt = rule.get("ai_prompt") or ""
        if prompt:
            p = QLabel(f'💬 "{prompt}"')
            p.setWordWrap(True)
            self._prompt = p
            layout.addWidget(p)

        action = (rule.get("action_type") or "unknown").upper()
        exts = rule.get("file_extensions", ["*"]) or ["*"]
        ext_txt = ", ".join(exts[:6]) + ("…" if len(exts) > 6 else "")
        dest = (rule.get("action_params") or {}).get("destination", "N/A")
        desc = QLabel(f"⚙️ {action}: {ext_txt} → {Path(dest).name if dest and dest != 'N/A' else dest}")
        desc.setWordWrap(True)
        self._desc = desc
        layout.addWidget(desc)
        self.refresh_theme()

    def refresh_theme(self) -> None:
        t = _ui_theme_tokens()
        try:
            if self._head:
                self._head.setStyleSheet(f"color:{t['text']}; font-weight:700;")
            if self._prompt:
                self._prompt.setStyleSheet("color:#9C27B0; font-style:italic;")
            if self._desc:
                self._desc.setStyleSheet(f"color:{t['muted']};")
        except Exception:
            pass


class _QtScheduledTaskMiniCard(QFrame):
    def __init__(self, *, task, icons: QtIconLoader):
        super().__init__()
        self.setObjectName("PageCard")
        self._title: QLabel | None = None
        self._meta: QLabel | None = None
        self._last: QLabel | None = None
        self._next: QLabel | None = None
        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 12, 16, 12)
        layout.setSpacing(6)

        title = QLabel(f"⏰ {getattr(task, 'title', 'Scheduled Task')}")
        self._title = title
        layout.addWidget(title)

        schedule = getattr(task, "schedule", {}) or {}
        times_val = schedule.get("times") or schedule.get("time") or ""
        if isinstance(times_val, (list, tuple)):
            times_txt = ", ".join(str(t) for t in times_val if str(t).strip())
        else:
            times_txt = str(times_val).strip()
        sched_txt = f"{schedule.get('type','')} {times_txt}".strip()
        action = (getattr(task, "action_type", "") or "").upper()
        target = getattr(task, "target_path", "") or ""
        meta = QLabel(f"⚙️ {action} • ⏱ {sched_txt or 'unscheduled'} • 📁 {Path(target).name if target else target}")
        meta.setWordWrap(True)
        self._meta = meta
        layout.addWidget(meta)

        last = getattr(task, "last_run_iso", None) or ""
        if last:
            last_lbl = QLabel(f"Last run: {last}")
            self._last = last_lbl
            layout.addWidget(last_lbl)

        try:
            nr = task.next_run()  # type: ignore[attr-defined]
        except Exception:
            nr = None
        if nr:
            try:
                nr_txt = nr.strftime("%Y-%m-%d %H:%M")
            except Exception:
                nr_txt = str(nr)
            next_lbl = QLabel(f"Next run: {nr_txt}")
            self._next = next_lbl
            layout.addWidget(next_lbl)
        self.refresh_theme()

    def refresh_theme(self) -> None:
        t = _ui_theme_tokens()
        try:
            if self._title:
                self._title.setStyleSheet(f"color:{t['text']}; font-weight:700;")
            if self._meta:
                self._meta.setStyleSheet(f"color:{t['muted']};")
            if self._last:
                self._last.setStyleSheet(f"color:{t['subtle']};")
            if self._next:
                self._next.setStyleSheet(f"color:{t['subtle']};")
        except Exception:
            pass


class _QtScheduledTaskCard(QFrame):
    def __init__(self, *, task, icons: QtIconLoader, on_edit, on_delete, on_run, on_toggle):
        super().__init__()
        self.setObjectName("PageCard")
        self.task = task
        self.icons = icons
        self._on_edit = on_edit
        self._on_delete = on_delete
        self._on_run = on_run
        self._on_toggle = on_toggle
        self._title: QLabel | None = None
        self._meta: QLabel | None = None
        self._last: QLabel | None = None
        self._next: QLabel | None = None

        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 14, 16, 14)
        layout.setSpacing(10)

        top = QHBoxLayout()
        top.setSpacing(10)

        title = QLabel(task.title)
        self._title = title
        top.addWidget(title, 1)

        enabled = QCheckBox("Enabled")
        enabled.setChecked(bool(getattr(task, "enabled", True)))
        enabled.stateChanged.connect(lambda _v: self._on_toggle(bool(enabled.isChecked())))
        top.addWidget(enabled)

        run_btn = self._make_action_button(
            icon_name="play",
            bg="#2563eb",
            hover="#1d4ed8",
            tooltip="Run now",
            on_click=lambda: self._on_run(),
        )
        top.addWidget(run_btn)

        edit_btn = self._make_action_button(
            icon_name="edit",
            bg="#f59e0b",
            hover="#d97706",
            tooltip="Edit",
            on_click=lambda: self._on_edit(),
        )
        top.addWidget(edit_btn)

        del_btn = self._make_action_button(
            icon_name="delete",
            bg="#dc2626",
            hover="#b91c1c",
            tooltip="Delete",
            on_click=lambda: self._on_delete(),
        )
        top.addWidget(del_btn)

        layout.addLayout(top)

        schedule = getattr(task, "schedule", {}) or {}
        times_val = schedule.get("times") or schedule.get("time") or ""
        if isinstance(times_val, (list, tuple)):
            times_txt = ", ".join(str(t) for t in times_val if str(t).strip())
        else:
            times_txt = str(times_val).strip()
        sched_txt = f"{schedule.get('type','')} {times_txt}".strip()
        action_name = str(getattr(task, "action_type", "") or "")
        safety = ""
        try:
            params = dict(getattr(task, "action_params", {}) or {})
            if action_name == "clean_folder":
                min_age = float(params.get("min_age_seconds", 604800) or 0)
                days = int(round(min_age / 86400)) if min_age > 0 else 0
                age_text = f"older than {days} day(s)" if days else "any age"
                download_text = "skips active downloads" if bool(params.get("skip_active_downloads", True)) else "does not skip active downloads"
                safety = f" • {age_text}, {download_text}"
        except Exception:
            safety = ""
        meta = QLabel(
            f"⏱ {sched_txt or 'unscheduled'} • ⚙️ {action_name}{safety} • 📁 {getattr(task,'target_path','')}"
        )
        meta.setWordWrap(True)
        self._meta = meta
        layout.addWidget(meta)

        last = getattr(task, "last_run_iso", None) or ""
        if last:
            last_lbl = QLabel(f"Last run: {last}")
            self._last = last_lbl
            layout.addWidget(last_lbl)

        try:
            nr = task.next_run()  # type: ignore[attr-defined]
        except Exception:
            nr = None
        if nr:
            try:
                nr_txt = nr.strftime("%Y-%m-%d %H:%M")
            except Exception:
                nr_txt = str(nr)
            next_lbl = QLabel(f"Next run: {nr_txt}")
            self._next = next_lbl
            layout.addWidget(next_lbl)
        self.refresh_theme()

    def _tinted_icon(self, icon_name: str, size: int, color: str) -> QIcon:
        pm = self.icons.pixmap(icon_name, int(size))
        if pm.isNull():
            return self.icons.icon(icon_name)
        out = QPixmap(pm.size())
        out.fill(Qt.transparent)
        p = QPainter(out)
        try:
            p.drawPixmap(0, 0, pm)
            p.setCompositionMode(QPainter.CompositionMode_SourceIn)
            p.fillRect(out.rect(), QColor(color))
        finally:
            p.end()
        return QIcon(out)

    def _make_action_button(self, *, icon_name: str, bg: str, hover: str, tooltip: str, on_click):
        btn = QPushButton()
        btn.setFixedSize(36, 36)
        btn.setIcon(self._tinted_icon(icon_name, 18, "#ffffff"))
        btn.setIconSize(QSize(18, 18))
        btn.setToolTip(tooltip)
        btn.setStyleSheet(
            f"QPushButton{{background:{bg}; border:1px solid {bg}; border-radius:10px;}}"
            f"QPushButton:hover{{background:{hover}; border:1px solid {hover};}}"
        )
        btn.clicked.connect(on_click)
        return btn

    def refresh_theme(self) -> None:
        t = _ui_theme_tokens()
        try:
            if self._title:
                self._title.setStyleSheet(f"color:{t['text']}; font-weight:700;")
            if self._meta:
                self._meta.setStyleSheet(f"color:{t['muted']};")
            if self._last:
                self._last.setStyleSheet(f"color:{t['subtle']};")
            if self._next:
                self._next.setStyleSheet(f"color:{t['subtle']};")
        except Exception:
            pass


class _EditScheduledTaskDialog(QDialog):
    def __init__(self, parent: QWidget, *, task=None):
        super().__init__(parent)
        self._task = task
        self.setWindowTitle("Scheduled Task")
        self.setModal(True)
        self.setMinimumWidth(520)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(14, 14, 14, 14)
        layout.setSpacing(10)

        form = QFormLayout()
        form.setLabelAlignment(Qt.AlignRight)

        self.title = QLineEdit(str(getattr(task, "title", "")) if task else "")
        self.title.setPlaceholderText("Daily temp cleanup")
        form.addRow("Title:", self.title)

        self.schedule_type = QComboBox()
        self.schedule_type.addItems(["daily", "weekly", "once"])
        if task:
            try:
                cur = str((task.schedule or {}).get("type") or "daily").strip().lower()
                idx = self.schedule_type.findText(cur)
                if idx >= 0:
                    self.schedule_type.setCurrentIndex(idx)
            except Exception:
                pass
        form.addRow("Schedule:", self.schedule_type)

        self.times = QLineEdit()
        self.times.setPlaceholderText("08:00, 16:00")
        if task:
            try:
                sch = (task.schedule or {}) if task else {}
                v = sch.get("times") or sch.get("time") or ""
                if isinstance(v, (list, tuple)):
                    self.times.setText(", ".join(str(x) for x in v))
                else:
                    self.times.setText(str(v))
            except Exception:
                pass
        form.addRow("Times:", self.times)

        # Weekly days (Mon..Sun) = 0..6 (Python weekday)
        self.days_box = QGroupBox()
        days_l = QHBoxLayout(self.days_box)
        days_l.setContentsMargins(0, 0, 0, 0)
        days_l.setSpacing(10)
        self.day_cbs: list[QCheckBox] = []
        for i, name in enumerate(["Mon", "Tue", "Wed", "Thu", "Fri", "Sat", "Sun"]):
            cb = QCheckBox(name)
            self.day_cbs.append(cb)
            days_l.addWidget(cb)
        days_l.addStretch(1)
        if task:
            try:
                days = (task.schedule or {}).get("days") or []
                days_i = {int(d) for d in days}
                for i, cb in enumerate(self.day_cbs):
                    cb.setChecked(i in days_i)
            except Exception:
                pass
        form.addRow("Days:", self.days_box)

        self.once_dt = QDateTimeEdit()
        self.once_dt.setCalendarPopup(True)
        self.once_dt.setDisplayFormat("yyyy-MM-dd HH:mm")
        self.once_dt.setDateTime(QDateTime.currentDateTime())
        if task:
            try:
                s = str((task.schedule or {}).get("datetime") or "")
                if s:
                    dt = QDateTime.fromString(s, Qt.ISODate)
                    if dt.isValid():
                        self.once_dt.setDateTime(dt)
            except Exception:
                pass
        form.addRow("Run once:", self.once_dt)

        def _sync_schedule_visibility():
            st = self.schedule_type.currentText()
            is_weekly = st == "weekly"
            is_once = st == "once"
            self.days_box.setVisible(is_weekly)
            self.once_dt.setVisible(is_once)
            # "Times" is used by daily/weekly, hidden for once.
            self.times.setVisible(not is_once)
            try:
                lab_times = form.labelForField(self.times)
                if lab_times:
                    lab_times.setVisible(not is_once)
            except Exception:
                pass
            try:
                lab_days = form.labelForField(self.days_box)
                if lab_days:
                    lab_days.setVisible(is_weekly)
            except Exception:
                pass
            try:
                lab_once = form.labelForField(self.once_dt)
                if lab_once:
                    lab_once.setVisible(is_once)
            except Exception:
                pass

        self.schedule_type.currentIndexChanged.connect(_sync_schedule_visibility)
        _sync_schedule_visibility()

        self.action_type = QComboBox()
        self.action_type.addItems(
            [
                "clean_folder",
                "delete",
                "copy",
                "move",
                "archive",
                "organize",
                "execute",
                "rename",
                "cloud_sync_upload_only",
                "cloud_sync_download_only",
                "cloud_sync_two_way",
            ]
        )
        if task:
            cur = str(getattr(task, "action_type", "") or "")
            idx = self.action_type.findText(cur)
            if idx >= 0:
                self.action_type.setCurrentIndex(idx)
        form.addRow("Action:", self.action_type)

        self.target = QLineEdit(str(getattr(task, "target_path", "")) if task else "")
        browse = QPushButton("Browse…")
        browse.clicked.connect(self._browse_target)
        trow = QHBoxLayout()
        trow.addWidget(self.target, 1)
        trow.addWidget(browse)
        thost = QWidget()
        thost.setLayout(trow)
        form.addRow("Target path:", thost)

        self.enabled = QCheckBox("Enabled")
        self.enabled.setChecked(bool(getattr(task, "enabled", True)) if task else True)
        form.addRow("", self.enabled)

        layout.addLayout(form)

        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        try:
            ok_btn = buttons.button(QDialogButtonBox.Ok)
            if ok_btn:
                ok_btn.setObjectName("PrimaryButton")
        except Exception:
            pass
        buttons.accepted.connect(self._validate_and_accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)

    def _browse_target(self):
        folder = QFileDialog.getExistingDirectory(self, "Select target folder")
        if folder:
            self.target.setText(folder)

    def _validate_and_accept(self):
        if not self.title.text().strip():
            QMessageBox.warning(self, "Missing title", "Title is required.")
            return
        if not self.target.text().strip():
            QMessageBox.warning(self, "Missing target", "Target path is required.")
            return
        st = self.schedule_type.currentText().strip().lower()
        if st in ("daily", "weekly"):
            if not self.times.text().strip():
                QMessageBox.warning(self, "Missing time", "At least one time is required (e.g., 08:00 or 08:00,16:00).")
                return
        if st == "weekly":
            if not any(cb.isChecked() for cb in self.day_cbs):
                QMessageBox.warning(self, "Missing days", "Select at least one day for weekly schedules.")
                return
        self.accept()

    def task_dict(self) -> dict:
        base = {}
        if self._task:
            base["task_id"] = getattr(self._task, "task_id", "")
        st = self.schedule_type.currentText().strip().lower()
        schedule: dict = {"type": st}
        if st in ("daily", "weekly"):
            raw = self.times.text().strip()
            parts = [p.strip() for p in raw.split(",") if p.strip()]
            if len(parts) <= 1:
                schedule["time"] = parts[0] if parts else raw
            else:
                schedule["times"] = parts
        if st == "weekly":
            schedule["days"] = [i for i, cb in enumerate(self.day_cbs) if cb.isChecked()]
        if st == "once":
            schedule["datetime"] = self.once_dt.dateTime().toString(Qt.ISODate)
        action_type = self.action_type.currentText()
        action_params = {}
        if action_type == "clean_folder":
            action_params = {
                "include_subfolders": True,
                "use_recycle_bin": True,
                "skip_active_downloads": True,
                "min_age_seconds": 604800,
            }
        base.update(
            {
                "title": self.title.text().strip(),
                "schedule": schedule,
                "action_type": action_type,
                "action_params": action_params,
                "target_path": self.target.text().strip(),
                "enabled": bool(self.enabled.isChecked()),
            }
        )
        return base


class _QtAiOpCard(QFrame):
    def __init__(self, *, key: str, title: str, description: str, icon, enabled: bool):
        super().__init__()
        self.setObjectName("PageCard")
        self.key = key
        self._enabled = bool(enabled)
        self._checked = False
        self._icon = icon
        self._icon_lbl: QLabel | None = None
        self._name_lbl: QLabel | None = None
        self._desc_lbl: QLabel | None = None

        layout = QVBoxLayout(self)
        layout.setContentsMargins(14, 12, 14, 12)
        layout.setSpacing(10)

        top = QHBoxLayout()
        top.setSpacing(10)

        icon_lbl = QLabel()
        icon_lbl.setFixedSize(34, 34)
        self._icon_lbl = icon_lbl
        top.addWidget(icon_lbl)

        top.addStretch(1)

        self.check = QCheckBox("")
        self.check.setChecked(False)
        self.check.setEnabled(self._enabled)
        self.check.stateChanged.connect(lambda _v: self._sync_checked())
        top.addWidget(self.check)
        layout.addLayout(top)

        name = QLabel(title)
        self._name_lbl = name
        layout.addWidget(name)

        desc = QLabel(description)
        desc.setWordWrap(True)
        self._desc_lbl = desc
        layout.addWidget(desc, 1)

        if not self._enabled:
            self.setToolTip("Coming soon")

        self.setCursor(Qt.PointingHandCursor)
        self.refresh_theme()

    def mousePressEvent(self, event):  # noqa: N802
        if self._enabled:
            self.check.setChecked(not self.check.isChecked())
        super().mousePressEvent(event)

    def _sync_checked(self):
        self._checked = bool(self.check.isChecked())
        self.refresh_theme()

    def is_checked(self) -> bool:
        return bool(self.check.isChecked())

    def set_checked(self, v: bool):
        if not self._enabled:
            return
        self.check.setChecked(bool(v))

    @staticmethod
    def _tint(pm: QPixmap, color: str) -> QPixmap:
        if pm.isNull():
            return pm
        out = QPixmap(pm.size())
        out.fill(Qt.transparent)
        p = QPainter(out)
        try:
            p.drawPixmap(0, 0, pm)
            p.setCompositionMode(QPainter.CompositionMode_SourceIn)
            p.fillRect(out.rect(), QColor(color))
        finally:
            p.end()
        return out

    def refresh_theme(self) -> None:
        t = _ui_theme_tokens()
        mode = _ui_theme_mode()
        try:
            if self._name_lbl:
                self._name_lbl.setStyleSheet(f"color:{t['text']}; font-size:13px; font-weight:700;")
            if self._desc_lbl:
                self._desc_lbl.setStyleSheet(f"color:{t['muted']};")
            if mode == "light":
                self.check.setStyleSheet(
                    "QCheckBox::indicator{width:18px; height:18px; border-radius:5px; border:1px solid #94a3b8; background:#ffffff;}"
                    "QCheckBox::indicator:checked{background:#0d6efd; border:1px solid #0d6efd;}"
                )
            else:
                self.check.setStyleSheet("")
        except Exception:
            pass
        try:
            if self._icon_lbl and self._icon:
                pm = self._icon.pixmap(34, 34)
                if not pm.isNull() and mode == "light":
                    pm = self._tint(pm, t["icon"])
                self._icon_lbl.setPixmap(pm)
        except Exception:
            pass
        try:
            if not self._enabled:
                bg = "#f3f6fc" if mode == "light" else "#14171c"
                br = "#d8e0ec" if mode == "light" else "#232730"
                self.setStyleSheet(f"QFrame#PageCard{{background:{bg}; border:1px solid {br}; border-radius:14px;}}")
                return
            if self._checked:
                if mode == "light":
                    self.setStyleSheet(
                        "QFrame#PageCard{background:#dbeafe; border:1px solid #93c5fd; border-radius:14px;}"
                    )
                else:
                    self.setStyleSheet(
                        "QFrame#PageCard{background:#1b2433; border:1px solid #1f3a5c; border-radius:14px;}"
                    )
            else:
                self.setStyleSheet("")
        except Exception:
            pass


class _QtAiHubWorker(QObject):
    status = Signal(str)
    progress = Signal(float)  # 0..1
    log = Signal(str)
    finished = Signal(bool, str)  # ok, report_dir
    error = Signal(str)

    def __init__(self, *, backend, target_folder: str, operations: list[str], options: dict):
        super().__init__()
        self.backend = backend
        self.target_folder = target_folder
        self.operations = operations
        self.options = options
        self.cancelled = False

    def cancel(self):
        self.cancelled = True

    def run(self):
        from pathlib import Path
        from datetime import datetime
        import json
        import shutil

        try:
            target = Path(self.target_folder)
            if not target.exists():
                self.error.emit("Target folder does not exist.")
                return

            settings = getattr(self.backend, "settings_manager", None)
            out_base = None
            try:
                out_base = Path(settings.app_folder) / "reports" / "ai_hub"  # type: ignore[attr-defined]
            except Exception:
                out_base = Path.home() / ".fylorra" / "reports" / "ai_hub"
            stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            report_dir = (out_base / stamp)
            report_dir.mkdir(parents=True, exist_ok=True)

            ai = getattr(self.backend, "ai_manager", None)

            apply_changes = bool(self.options.get("apply"))
            include_subfolders = bool(self.options.get("include_subfolders", True))
            filter_key = str(self.options.get("filter_key") or "all")
            use_vision = bool(self.options.get("use_vision"))
            use_ai_docs = bool(self.options.get("use_ai_docs"))

            needs_ai = bool(use_vision) or ("security_scan" in self.operations) or ("content_analysis" in self.operations) or bool(use_ai_docs)
            if needs_ai and not ai:
                self.error.emit("AI Manager is not available in this environment.")
                return

            def _log(msg: str):
                self.log.emit(msg)

            def _set_status(msg: str):
                self.status.emit(msg)

            def _set_progress(p: float):
                self.progress.emit(max(0.0, min(1.0, float(p))))

            # Model loading is handled by the dialog before this worker starts.
            if needs_ai and ai and (not getattr(ai, "is_ready", False)):
                self.error.emit("AI model is not loaded. Load it before starting AI Hub operations.")
                return

            from core.bulk_ai_processor import BulkAIProcessor, ProcessingOptions, ProcessingMode

            def _exts_for_filter(k: str):
                k = (k or "all").lower()
                if k == "images":
                    return BulkAIProcessor.get_image_extensions()
                if k == "videos":
                    return BulkAIProcessor.get_video_extensions()
                if k == "documents":
                    return BulkAIProcessor.get_document_extensions()
                if k == "code":
                    return BulkAIProcessor.get_code_extensions()
                return None

            file_exts = _exts_for_filter(filter_key)

            # Operation loop.
            op_total = len(self.operations)
            for op_idx, op in enumerate(self.operations):
                if self.cancelled:
                    break

                base_p = (op_idx / max(1, op_total))
                _set_progress(0.2 + base_p * 0.8)

                if op == "smart_rename":
                    if not ai:
                        _log("Smart Rename skipped (AI unavailable).")
                        continue
                    _set_status("Smart Rename: scanning files…")
                    proc = BulkAIProcessor(ai)
                    options = ProcessingOptions(
                        include_subfolders=include_subfolders,
                        file_extensions=file_exts,
                        batch_size=50,
                        mode=ProcessingMode.SMART,
                        max_files=None,
                    )
                    files = proc.scan_folder(target, options)
                    _log(f"Smart Rename: found {len(files)} files")
                    if not files:
                        continue

                    from utils.intelligent_rename import sanitize_ai_filename, get_unique_filename
                    from utils.universal_undo import record_bulk_rename
                    from core.semantic_analyzer import SemanticAnalyzer

                    sem = SemanticAnalyzer(ai) if (use_ai_docs and ai and getattr(ai, "is_ready", False)) else None
                    image_exts = {"png", "jpg", "jpeg", "gif", "bmp", "webp", "tiff", "tif"}
                    pairs = []
                    preview = []
                    total = len(files)
                    for i, fp in enumerate(files):
                        if self.cancelled:
                            break
                        if (i % 10) == 0 or i == total - 1:
                            _set_status(f"Smart Rename: analyzing {i+1}/{total}…")
                        used_ai = False
                        suggested = None
                        ext = fp.suffix.lower().lstrip(".")
                        try:
                            if sem and fp.suffix.lower() in {".pdf", ".docx", ".doc"}:
                                res = sem.analyze_document(fp)
                                if res and res.suggested_filename:
                                    suggested = res.suggested_filename
                            if not suggested:
                                if use_vision and ext in image_exts and getattr(ai, "is_ready", False):
                                    used_ai = True
                                suggested = ai.analyze_file_for_rename(fp, use_ai=used_ai)
                                if (not suggested) and used_ai:
                                    suggested = ai.analyze_file_for_rename(fp, use_ai=False)
                        except Exception:
                            suggested = None

                        if not suggested:
                            continue

                        validation = sanitize_ai_filename(str(suggested), preserve_case=True)
                        desired = validation.sanitized_name
                        if not desired or desired.strip() == fp.stem:
                            continue

                        unique_name, _why = get_unique_filename(fp, desired, {"used_ai": used_ai})
                        new_path = fp.with_name(unique_name + fp.suffix)
                        if str(new_path) == str(fp):
                            continue
                        preview.append({"from": str(fp), "to": str(new_path), "used_ai": used_ai})
                        if apply_changes:
                            try:
                                fp.rename(new_path)
                                pairs.append((fp, new_path))
                            except Exception:
                                continue

                    with open(report_dir / "smart_rename.json", "w", encoding="utf-8") as f:
                        json.dump({"apply": apply_changes, "items": preview}, f, indent=2)

                    if apply_changes and pairs:
                        try:
                            record_bulk_rename([(a, b) for (a, b) in pairs], metadata={"ai_hub": True})
                        except Exception:
                            pass
                    _log(f"Smart Rename: {'renamed' if apply_changes else 'previewed'} {len(preview)} files")

                elif op == "auto_categorize":
                    _set_status("Auto-Categorize: scanning…")
                    from core.enhanced_categorizer import EnhancedCategorizer

                    cat = EnhancedCategorizer(ai_manager=ai, use_ai_vision=bool(use_vision))
                    categorized = cat.categorize_folder(
                        target,
                        include_subfolders=include_subfolders,
                        use_ai_vision=bool(use_vision),
                        smart_scope=True,
                        include_other=bool(self.options.get("include_other", True)),
                        use_ai_documents=bool(use_ai_docs),
                    )
                    categorized_ser = {
                        k: ([str(p) for p in v] if isinstance(v, list) else v) for k, v in categorized.items()
                    }
                    counts = {k: len(v) for k, v in categorized_ser.items() if isinstance(v, list)}
                    with open(report_dir / "auto_categorize.json", "w", encoding="utf-8") as f:
                        json.dump(
                            {
                                "version": 1,
                                "generated_at": stamp,
                                "target_folder": str(target),
                                "apply_requested": apply_changes,
                                "options": {
                                    "include_subfolders": include_subfolders,
                                    "smart_scope": True,
                                    "include_other": bool(self.options.get("include_other", True)),
                                    "use_ai_vision": bool(use_vision),
                                    "use_ai_documents": bool(use_ai_docs),
                                    "filter_key": filter_key,
                                },
                                "counts": counts,
                                "categories": categorized_ser,
                            },
                            f,
                            indent=2,
                        )
                    _log(
                        f"Auto-Categorize: found {sum(counts.values())} file placements in {len(counts)} categories"
                        + (" — review results below" if apply_changes else " (preview only — enable Apply changes to apply moves)")
                    )

                elif op == "security_scan":
                    if not ai or not getattr(ai, "is_ready", False):
                        _log("Security Scan skipped (AI model not loaded).")
                        continue
                    _set_status("Security Scan: scanning images…")
                    proc = BulkAIProcessor(ai)
                    options = ProcessingOptions(
                        include_subfolders=include_subfolders,
                        file_extensions=BulkAIProcessor.get_image_extensions(),
                        batch_size=50,
                        mode=ProcessingMode.SMART,
                        max_files=None,
                    )
                    files = proc.scan_folder(target, options)
                    sensitive = []
                    total = len(files)
                    for i, fp in enumerate(files):
                        if self.cancelled:
                            break
                        if (i % 5) == 0 or i == total - 1:
                            _set_status(f"Security Scan: {i+1}/{total}…")
                        try:
                            res = ai.detect_sensitive_content(fp)
                            if res.get("sensitive"):
                                sensitive.append({"file": str(fp), "reason": res.get("reason")})
                        except Exception:
                            continue
                    with open(report_dir / "security_scan.json", "w", encoding="utf-8") as f:
                        json.dump({"matches": sensitive}, f, indent=2)
                    _log(f"Security Scan: found {len(sensitive)} sensitive images")

                elif op == "content_analysis":
                    if not ai or not getattr(ai, "is_ready", False):
                        _log("Content Analysis skipped (AI model not loaded).")
                        continue
                    _set_status("Content Analysis: scanning documents…")
                    from core.semantic_analyzer import SemanticAnalyzer

                    sem = SemanticAnalyzer(ai)
                    proc = BulkAIProcessor(ai)
                    options = ProcessingOptions(
                        include_subfolders=include_subfolders,
                        file_extensions=BulkAIProcessor.get_document_extensions(),
                        batch_size=50,
                        mode=ProcessingMode.SMART,
                        max_files=500,  # keep bounded for first Qt version
                    )
                    files = proc.scan_folder(target, options)
                    out = []
                    total = len(files)
                    for i, fp in enumerate(files):
                        if self.cancelled:
                            break
                        if (i % 5) == 0 or i == total - 1:
                            _set_status(f"Content Analysis: {i+1}/{total}…")
                        try:
                            res = sem.analyze_document(fp)
                            if not res:
                                continue
                            out.append(
                                {
                                    "file": str(fp),
                                    "document_type": res.document_type,
                                    "domain": res.domain,
                                    "confidence": res.confidence,
                                    "explanation": res.explanation,
                                    "suggested_filename": res.suggested_filename,
                                    "suggested_category": res.suggested_category,
                                    "sensitivity": res.sensitivity,
                                }
                            )
                        except Exception:
                            continue
                    with open(report_dir / "content_analysis.json", "w", encoding="utf-8") as f:
                        json.dump({"items": out}, f, indent=2)
                    _log(f"Content Analysis: analyzed {len(out)} documents")

                elif op == "duplicate_detection":
                    _set_status("Duplicate Detection: scanning…")
                    from core.duplicate_finder import find_exact_duplicates

                    def cb(cur: int, total: int, p: Path):
                        # total == 0 means "indeterminate" (counting stage).
                        if total and total > 0:
                            if (cur % 10) == 0 or cur == total:
                                self.progress.emit(max(0.0, min(1.0, float(cur) / float(max(1, total)))))
                            _set_status(f"Duplicate Detection: {cur}/{total}…")
                        else:
                            # Counting files (no known total yet).
                            if (cur % 2000) == 0:
                                _set_status(f"Duplicate Detection: scanning… ({cur} files)")

                    cancelled_ref = lambda: bool(self.cancelled)

                    class _CancelProxy:
                        def is_set(self) -> bool:
                            return cancelled_ref()

                    groups = find_exact_duplicates(
                        target,
                        include_subfolders=include_subfolders,
                        cancel_event=_CancelProxy(),
                        progress_cb=cb,
                    )
                    rep = {
                        "apply": apply_changes,
                        "groups": [
                            {"sha256": g.sha256, "size": g.size, "original": g.original, "duplicates": g.duplicates}
                            for g in groups
                        ],
                    }
                    with open(report_dir / "duplicate_detection.json", "w", encoding="utf-8") as f:
                        json.dump(rep, f, indent=2)
                    _log(
                        f"Duplicate Detection: found {sum(len(g['duplicates']) for g in rep['groups'])} duplicates in {len(rep['groups'])} groups"
                    )
                    if apply_changes:
                        _log("Duplicate Detection: review results below to apply any actions.")

                else:
                    _log(f"{op}: not implemented yet.")

            ok = not self.cancelled
            self.finished.emit(ok, str(report_dir))

        except Exception as e:
            self.error.emit(str(e))


class _QtAiHubRunDialog(QDialog):
    def __init__(self, parent: QWidget, *, target_folder: str, operations: list[str], options: dict):
        super().__init__(parent)
        self.setWindowTitle("AI Hub Run")
        self.setModal(True)
        self.setMinimumWidth(700)
        self.setMinimumHeight(420)
        self.setStyleSheet(_qt_modern_dialog_stylesheet())

        self._target_folder = target_folder
        self._operations = list(operations or [])
        self._options = dict(options or {})
        self._report_dir: str | None = None

        top = parent.window()
        self.backend = getattr(top, "backend", None)
        self.worker = _QtAiHubWorker(backend=self.backend, target_folder=target_folder, operations=operations, options=options)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(14, 14, 14, 14)
        layout.setSpacing(10)

        title = QLabel("Running AI Hub operations")
        title.setObjectName("DialogTitle")
        layout.addWidget(title)

        self.status = QLabel("Starting…")
        self.status.setObjectName("DialogSubtitle")
        layout.addWidget(self.status)

        self.bar = QProgressBar()
        self.bar.setRange(0, 1000)
        self.bar.setValue(0)
        layout.addWidget(self.bar)

        self.tabs = QTabWidget()
        self.tabs.setStyleSheet(
            "QTabWidget::pane{border:0;}"
            "QTabBar::tab{background:#1a1e25; color:#c8ccd6; padding:8px 12px; border-radius:8px; margin-right:6px;}"
            "QTabBar::tab:selected{background:#243040; color:#ffffff;}"
        )

        # Results tab (structured preview + apply).
        results = QWidget()
        rlay = QVBoxLayout(results)
        rlay.setContentsMargins(0, 0, 0, 0)
        rlay.setSpacing(10)

        self.results_summary = QLabel("No results yet.")
        self.results_summary.setStyleSheet("color:#9aa0a9;")
        rlay.addWidget(self.results_summary)

        self.cat_group = QGroupBox("Auto-Categorize")
        self.cat_group.setStyleSheet(
            "QGroupBox{color:#ffffff; font-weight:700; border:1px solid #232730; border-radius:10px; margin-top:12px;}"
            "QGroupBox::title{subcontrol-origin: margin; left: 12px; padding: 0 4px;}"
        )
        cg = QVBoxLayout(self.cat_group)
        cg.setContentsMargins(10, 10, 10, 10)
        cg.setSpacing(8)
        self.cat_tree = QTreeWidget()
        self.cat_tree.setHeaderLabels(["Category", "Count / Sample files"])
        self.cat_tree.setStyleSheet(
            "QTreeWidget{background:#14171c; border:1px solid #232730; border-radius:10px; color:#c8ccd6;}"
            "QTreeWidget::item{padding:4px 2px;}"
        )
        self.cat_tree.setRootIsDecorated(True)
        self.cat_tree.setAlternatingRowColors(False)
        cg.addWidget(self.cat_tree, 1)

        cat_btns = QHBoxLayout()
        self.btn_cat_select_all = QPushButton("Select all")
        self.btn_cat_select_all.clicked.connect(lambda: self._set_all_categories_checked(True))
        self.btn_cat_select_none = QPushButton("Select none")
        self.btn_cat_select_none.clicked.connect(lambda: self._set_all_categories_checked(False))
        cat_btns.addWidget(self.btn_cat_select_all)
        cat_btns.addWidget(self.btn_cat_select_none)
        self.btn_open_categorize_dialog = QPushButton("Open Full Dialog…")
        self.btn_open_categorize_dialog.clicked.connect(self._open_categorize_dialog)
        self.btn_open_categorize_dialog.setEnabled(False)
        cat_btns.addWidget(self.btn_open_categorize_dialog)
        cat_btns.addStretch(1)
        self.btn_apply_categorize = QPushButton("Apply Organization")
        self.btn_apply_categorize.clicked.connect(self._apply_categorize)
        self.btn_apply_categorize.setEnabled(False)
        cat_btns.addWidget(self.btn_apply_categorize)
        cg.addLayout(cat_btns)

        rlay.addWidget(self.cat_group, 1)

        self.dup_group = QGroupBox("Duplicate Detection")
        self.dup_group.setStyleSheet(self.cat_group.styleSheet())
        dg = QVBoxLayout(self.dup_group)
        dg.setContentsMargins(10, 10, 10, 10)
        dg.setSpacing(8)
        self.dup_list = QListWidget()
        self.dup_list.setStyleSheet(
            "QListWidget{background:#14171c; border:1px solid #232730; border-radius:10px; color:#c8ccd6;}"
        )
        dg.addWidget(self.dup_list, 1)
        dup_btns = QHBoxLayout()
        dup_btns.addStretch(1)
        self.btn_review_duplicates = QPushButton("Review / Delete…")
        self.btn_review_duplicates.clicked.connect(self._open_duplicates_dialog)
        self.btn_review_duplicates.setEnabled(False)
        dup_btns.addWidget(self.btn_review_duplicates)
        dg.addLayout(dup_btns)
        rlay.addWidget(self.dup_group, 1)

        self.tabs.addTab(results, "Results")

        # Log tab (existing stream).
        log = QWidget()
        ll = QVBoxLayout(log)
        ll.setContentsMargins(0, 0, 0, 0)
        ll.setSpacing(0)
        self.log_list = QListWidget()
        self.log_list.setStyleSheet(
            "QListWidget{background:#14171c; border:1px solid #232730; border-radius:10px; color:#c8ccd6;}"
        )
        ll.addWidget(self.log_list, 1)
        self.tabs.addTab(log, "Log")

        layout.addWidget(self.tabs, 1)

        row = QHBoxLayout()
        self.btn_open_reports = QPushButton("Open Reports Folder")
        self.btn_open_reports.clicked.connect(self._open_reports)
        self.btn_open_reports.setEnabled(False)
        row.addWidget(self.btn_open_reports)
        row.addStretch(1)
        self.btn_cancel = QPushButton("Cancel")
        self.btn_cancel.clicked.connect(self._cancel)
        row.addWidget(self.btn_cancel)
        layout.addLayout(row)

        self._thread = QThread(self)
        self.worker.moveToThread(self._thread)
        self._thread.started.connect(self.worker.run)
        self.worker.status.connect(self.status.setText)
        self.worker.progress.connect(lambda p: self.bar.setValue(int(p * 1000)))
        self.worker.log.connect(self._append_log)
        self.worker.finished.connect(self._on_finished)
        self.worker.error.connect(self._on_error)
        self.worker.finished.connect(self._thread.quit)
        self.worker.error.connect(self._thread.quit)
        self._thread.finished.connect(self._thread.deleteLater)
        self._thread.start()

    def _set_all_categories_checked(self, checked: bool):
        for i in range(self.cat_tree.topLevelItemCount()):
            it = self.cat_tree.topLevelItem(i)
            if it is None:
                continue
            it.setCheckState(0, Qt.Checked if checked else Qt.Unchecked)

    def _append_log(self, text: str):
        self.log_list.insertItem(0, QListWidgetItem(text))
        while self.log_list.count() > 300:
            self.log_list.takeItem(self.log_list.count() - 1)

    def _cancel(self):
        self.btn_cancel.setEnabled(False)
        try:
            self.worker.cancel()
        except Exception:
            pass
        self.status.setText("Cancelling…")

    def _on_finished(self, ok: bool, report_dir: str):
        self._report_dir = report_dir
        try:
            self._load_results(report_dir)
        except Exception:
            pass
        self.bar.setValue(1000 if ok else self.bar.value())
        self.status.setText(("Complete." if ok else "Cancelled.") + f" Reports: {report_dir}")
        self.btn_cancel.setText("Close")
        self.btn_cancel.setEnabled(True)
        try:
            self.btn_cancel.clicked.disconnect()
        except Exception:
            pass
        self.btn_cancel.clicked.connect(self.accept)

    def _open_reports(self):
        if not self._report_dir:
            return
        try:
            QDesktopServices.openUrl(QUrl.fromLocalFile(self._report_dir))
        except Exception:
            pass

    def _load_results(self, report_dir: str):
        from pathlib import Path
        import json

        base = Path(report_dir)
        self.btn_open_reports.setEnabled(base.exists())
        # Cache for quick “Review/Delete…” without re-reading/parsing.
        self._dup_groups = []

        # Auto-Categorize
        cat_path = base / "auto_categorize.json"
        self.cat_group.setVisible(cat_path.exists())
        self.cat_tree.clear()
        can_apply = False
        if cat_path.exists():
            try:
                data = json.loads(cat_path.read_text(encoding="utf-8"))
            except Exception:
                data = {}
            counts = dict(data.get("counts") or {})
            categories = dict(data.get("categories") or {})
            apply_requested = bool(data.get("apply_requested") or data.get("apply"))
            can_apply = apply_requested and bool(categories)
            self.btn_open_categorize_dialog.setEnabled(True)

            total = int(sum(int(v) for v in counts.values() if isinstance(v, (int, float))))
            self.results_summary.setText(
                f"Auto-Categorize: {total} placements in {len(counts)} categories"
                + ("" if apply_requested else " (preview only — enable Apply changes to apply moves)")
            )

            for cat_name, file_list in sorted(categories.items(), key=lambda kv: len(kv[1]) if isinstance(kv[1], list) else 0, reverse=True):
                if not isinstance(file_list, list):
                    continue
                n = len(file_list)
                top = QTreeWidgetItem([str(cat_name), str(n)])
                top.setFlags(top.flags() | Qt.ItemIsUserCheckable)
                default_checked = cat_name not in {"empty_folders", "ignored_projects"}
                top.setCheckState(0, Qt.Checked if default_checked else Qt.Unchecked)
                # sample children
                for fp in file_list[:10]:
                    try:
                        p = Path(fp)
                        child_txt = f"{p.name}  —  {p.parent}"
                    except Exception:
                        child_txt = str(fp)
                    child = QTreeWidgetItem(["", child_txt])
                    child.setFlags(child.flags() & ~Qt.ItemIsUserCheckable)
                    top.addChild(child)
                self.cat_tree.addTopLevelItem(top)
            self.cat_tree.expandToDepth(0)
            self.cat_tree.resizeColumnToContents(0)

        self.btn_apply_categorize.setEnabled(bool(can_apply))

        # Duplicate Detection
        dup_path = base / "duplicate_detection.json"
        self.dup_group.setVisible(dup_path.exists())
        self.dup_list.clear()
        if dup_path.exists():
            try:
                d = json.loads(dup_path.read_text(encoding="utf-8"))
            except Exception:
                d = {}
            groups = list(d.get("groups") or [])
            self._dup_groups = groups
            dup_total = 0
            for g in groups:
                dup_total += len(g.get("duplicates") or [])
            self.dup_list.addItem(QListWidgetItem(f"{dup_total} duplicate files across {len(groups)} groups"))
            for g in groups[:50]:
                o = str(g.get("original") or "")
                dups = list(g.get("duplicates") or [])
                self.dup_list.addItem(QListWidgetItem(f"{Path(o).name}  →  {len(dups)} duplicates"))
            if len(groups) > 50:
                self.dup_list.addItem(QListWidgetItem(f"…and {len(groups)-50} more groups (see report JSON)."))
            self.btn_review_duplicates.setEnabled(True)
        else:
            self.btn_review_duplicates.setEnabled(False)

        # Prefer showing Results when complete.
        self.tabs.setCurrentIndex(0)

    def _open_categorize_dialog(self):
        if not self._report_dir:
            return
        dlg = _QtAutoCategorizeDialog(self, backend=self.backend, report_dir=self._report_dir)
        dlg.exec()
        try:
            self._load_results(self._report_dir)
        except Exception:
            pass

    def _open_duplicates_dialog(self):
        if not self._report_dir:
            return
        dlg = _QtDuplicateReviewDialog(
            self,
            backend=self.backend,
            report_dir=self._report_dir,
            groups=(getattr(self, "_dup_groups", None) or None),
        )
        dlg.exec()
        try:
            self._load_results(self._report_dir)
        except Exception:
            pass

    def _apply_categorize(self):
        if not self._report_dir:
            return
        selected = set()
        for i in range(self.cat_tree.topLevelItemCount()):
            it = self.cat_tree.topLevelItem(i)
            if not it:
                continue
            if it.checkState(0) == Qt.Checked:
                selected.add(it.text(0))
        if not selected:
            QMessageBox.information(self, "Apply Organization", "No categories selected.")
            return
        dlg = _QtAiHubApplyCategorizeDialog(self, report_dir=self._report_dir, selected_categories=sorted(selected))
        dlg.exec()
        # Refresh results after apply (files moved).
        try:
            self._load_results(self._report_dir)
        except Exception:
            pass

    def _on_error(self, msg: str):
        QMessageBox.critical(self, "AI Hub Failed", msg)
        self.reject()


class _QtAiHubApplyCategorizeWorker(QObject):
    status = Signal(str)
    progress = Signal(float)
    finished = Signal(bool, int, int)  # ok, moved, skipped
    error = Signal(str)

    def __init__(self, *, report_dir: str, selected_categories: list[str]):
        super().__init__()
        self.report_dir = report_dir
        self.selected_categories = selected_categories
        self.cancelled = False

    def cancel(self):
        self.cancelled = True

    def run(self):
        from pathlib import Path
        import json
        import shutil

        try:
            base = Path(self.report_dir)
            cat_path = base / "auto_categorize.json"
            if not cat_path.exists():
                self.error.emit("auto_categorize.json not found.")
                return
            data = json.loads(cat_path.read_text(encoding="utf-8"))
            target = Path(str(data.get("target_folder") or ""))
            categories = dict(data.get("categories") or {})

            from core.enhanced_categorizer import EnhancedCategorizer
            from utils.universal_undo import record_categorize

            cat = EnhancedCategorizer(ai_manager=None, use_ai_vision=False)

            selected = [c for c in self.selected_categories if c in categories]
            files_to_move: list[tuple[str, str]] = []
            for c in selected:
                v = categories.get(c)
                if isinstance(v, list):
                    for fp in v:
                        files_to_move.append((c, str(fp)))

            total = len(files_to_move)
            moved_pairs = []
            moved = 0
            skipped = 0
            for idx, (category, src) in enumerate(files_to_move, start=1):
                if self.cancelled:
                    break
                if total > 0:
                    self.progress.emit(float(idx - 1) / float(total))
                self.status.emit(f"Moving {idx}/{total}…")

                try:
                    src_p = Path(src)
                    if not src_p.exists():
                        skipped += 1
                        continue
                    folder_rel = cat.get_category_folder(category)
                    dest_dir = (target / folder_rel)
                    dest_dir.mkdir(parents=True, exist_ok=True)

                    # Skip if already in correct folder.
                    try:
                        if dest_dir.resolve() in src_p.resolve().parents:
                            skipped += 1
                            continue
                    except Exception:
                        pass

                    out = dest_dir / src_p.name
                    if out.exists():
                        stem = out.stem
                        suf = out.suffix
                        n = 1
                        while out.exists():
                            out = dest_dir / f"{stem}_{n}{suf}"
                            n += 1
                    shutil.move(str(src_p), str(out))
                    moved_pairs.append((src_p, out))
                    moved += 1
                except Exception:
                    skipped += 1
                    continue

            if moved_pairs:
                try:
                    record_categorize(moved_pairs, metadata={"ai_hub": True})
                except Exception:
                    pass

            self.progress.emit(1.0)
            self.finished.emit(not self.cancelled, moved, skipped)
        except Exception as e:
            self.error.emit(str(e))


class _QtDuplicateReviewDialog(QDialog):
    """
    Modern duplicate review UI:
    - Shows duplicate groups
    - Auto-selects deletions by confidence threshold (exact duplicates = 100%)
    - Side-by-side preview for quick validation
    """

    def __init__(self, parent: QWidget, *, backend, report_dir: str, groups: list[dict] | None = None):
        super().__init__(parent)
        self.backend = backend
        self.report_dir = report_dir
        self.setWindowTitle("Duplicate Detection — Review")
        self.setMinimumSize(980, 620)

        outer = QVBoxLayout(self)
        outer.setContentsMargins(14, 14, 14, 14)
        outer.setSpacing(10)

        header = QHBoxLayout()
        title = QLabel("Duplicate Detection")
        title.setStyleSheet("color:#ffffff; font-size:18px; font-weight:800;")
        header.addWidget(title)
        header.addStretch(1)

        self.keep_policy = QComboBox()
        self.keep_policy.addItems(["Keep original (recommended)", "Keep newest", "Keep oldest"])
        self.keep_policy.setFixedWidth(220)
        header.addWidget(self.keep_policy)

        self.thresh = QSlider(Qt.Horizontal)
        self.thresh.setRange(50, 100)
        self.thresh.setValue(95)
        self.thresh.setFixedWidth(170)
        header.addWidget(QLabel("Auto-select ≥"))
        self.thresh_lbl = QLabel("95%")
        self.thresh.valueChanged.connect(lambda v: self.thresh_lbl.setText(f"{int(v)}%"))
        header.addWidget(self.thresh)
        header.addWidget(self.thresh_lbl)

        self.btn_autoselect = QPushButton("Auto-select")
        self.btn_autoselect.clicked.connect(self._auto_select_current_group)
        header.addWidget(self.btn_autoselect)

        outer.addLayout(header)

        split = QSplitter(Qt.Horizontal)

        # Left: groups list
        left = QWidget()
        ll = QVBoxLayout(left)
        ll.setContentsMargins(0, 0, 0, 0)
        ll.setSpacing(8)
        self.group_list = QListWidget()
        self.group_list.setStyleSheet(
            "QListWidget{background:#14171c; border:1px solid #232730; border-radius:12px; color:#c8ccd6; outline:0; padding:6px;}"
            "QListWidget::item{padding:10px 10px; margin:4px 4px; border-radius:10px;}"
            "QListWidget::item:selected{background:#1b2430; border:1px solid #2a3442; color:#e6e9f2;}"
            "QListWidget::item:hover{background:#141c25;}"
        )
        self.group_list.currentRowChanged.connect(self._load_group)
        ll.addWidget(self.group_list, 1)
        split.addWidget(left)

        # Right: files table + preview
        right = QWidget()
        rl = QVBoxLayout(right)
        rl.setContentsMargins(0, 0, 0, 0)
        rl.setSpacing(8)

        self.table = QTableWidget(0, 5)
        self.table.setHorizontalHeaderLabels(["Delete", "File", "Folder", "Size", "Confidence"])
        self.table.horizontalHeader().setStretchLastSection(True)
        self.table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeToContents)
        self.table.horizontalHeader().setSectionResizeMode(1, QHeaderView.Stretch)
        self.table.horizontalHeader().setSectionResizeMode(2, QHeaderView.Stretch)
        self.table.horizontalHeader().setSectionResizeMode(3, QHeaderView.ResizeToContents)
        self.table.horizontalHeader().setSectionResizeMode(4, QHeaderView.ResizeToContents)
        self.table.setSelectionBehavior(QTableWidget.SelectRows)
        self.table.setSelectionMode(QTableWidget.SingleSelection)
        self.table.itemSelectionChanged.connect(self._update_preview)
        self.table.setStyleSheet(
            "QTableWidget{background:#14171c; border:1px solid #232730; border-radius:12px; color:#c8ccd6;}"
            "QHeaderView::section{background:#0f1217; color:#c8ccd6; border:0; padding:6px; font-weight:700;}"
        )
        rl.addWidget(self.table, 2)

        prev_split = QSplitter(Qt.Horizontal)
        self.prev_a = QLabel("Preview")
        self.prev_a.setAlignment(Qt.AlignCenter)
        self.prev_a.setStyleSheet("background:#0f1217; border:1px solid #232730; border-radius:12px; color:#9aa0a9;")
        self.prev_b = QLabel("Compare")
        self.prev_b.setAlignment(Qt.AlignCenter)
        self.prev_b.setStyleSheet("background:#0f1217; border:1px solid #232730; border-radius:12px; color:#9aa0a9;")
        prev_split.addWidget(self.prev_a)
        prev_split.addWidget(self.prev_b)
        prev_split.setSizes([1, 1])
        rl.addWidget(prev_split, 2)

        btn_row = QHBoxLayout()
        self.status = QLabel("Ready.")
        self.status.setStyleSheet("color:#9aa0a9;")
        btn_row.addWidget(self.status)
        btn_row.addStretch(1)
        self.btn_show = QPushButton("Show in Folder")
        self.btn_show.clicked.connect(self._show_selected)
        btn_row.addWidget(self.btn_show)
        self.btn_delete = QPushButton("Delete Selected")
        self.btn_delete.clicked.connect(self._delete_selected)
        self.btn_delete.setStyleSheet("QPushButton{background:#b23b3b; color:#fff; padding:8px 12px; border-radius:10px; font-weight:800;}")
        btn_row.addWidget(self.btn_delete)
        self.btn_close = QPushButton("Close")
        self.btn_close.clicked.connect(self.accept)
        btn_row.addWidget(self.btn_close)
        rl.addLayout(btn_row)

        split.addWidget(right)
        split.setSizes([320, 660])
        outer.addWidget(split, 1)

        self._groups: list[dict] = list(groups or [])
        self._current_files: list[str] = []
        if self._groups:
            self._populate_groups_from_memory()
        else:
            self._load_report()

    def _populate_groups_from_memory(self):
        from pathlib import Path

        self.group_list.clear()
        total_dups = 0
        for g in self._groups:
            total_dups += len(g.get("duplicates") or [])
        for g in self._groups:
            o = str(g.get("original") or "")
            ndup = len(g.get("duplicates") or [])
            size = int(g.get("size") or 0)
            mb = (size * max(1, ndup)) / (1024 * 1024)
            it = QListWidgetItem()
            it.setSizeHint(QSize(10, 56))
            self.group_list.addItem(it)
            self.group_list.setItemWidget(it, _QtDuplicateGroupCard(Path(o).name, ndup, mb))
        self.status.setText(f"{total_dups} duplicate files across {len(self._groups)} groups.")
        if self.group_list.count() > 0:
            self.group_list.setCurrentRow(0)


class _QtDuplicateGroupCard(QFrame):
    def __init__(self, name: str, dup_count: int, mb: float):
        super().__init__()
        self.setObjectName("DupGroupCard")
        self.setStyleSheet(
            "#DupGroupCard{background:transparent; border:0;}"
            "#DupTitle{color:#e6e9f2; font-weight:800;}"
            "#DupSub{color:#9aa0a9;}"
            "#DupPill{background:#1b2430; border:1px solid #2a3442; color:#9cc7ff; padding:2px 8px; border-radius:10px; font-weight:700;}"
        )
        lay = QHBoxLayout(self)
        lay.setContentsMargins(0, 0, 0, 0)
        lay.setSpacing(10)

        left = QVBoxLayout()
        left.setSpacing(2)
        t = QLabel(name)
        t.setObjectName("DupTitle")
        t.setTextInteractionFlags(Qt.TextSelectableByMouse)
        left.addWidget(t)
        s = QLabel(f"{dup_count} duplicate(s)  •  {mb:.1f} MB")
        s.setObjectName("DupSub")
        left.addWidget(s)
        lay.addLayout(left, 1)

        pill = QLabel(f"{dup_count} dup")
        pill.setObjectName("DupPill")
        pill.setAlignment(Qt.AlignCenter)
        lay.addWidget(pill, 0, Qt.AlignRight | Qt.AlignVCenter)

    def _load_report(self):
        from pathlib import Path
        import json

        base = Path(self.report_dir)
        p = base / "duplicate_detection.json"
        if not p.exists():
            QMessageBox.information(self, "Duplicates", "No duplicate report found.")
            return
        try:
            d = json.loads(p.read_text(encoding="utf-8"))
        except Exception:
            d = {}
        raw_groups = list(d.get("groups") or [])
        # Prune missing files so the review reflects current disk state.
        pruned: list[dict] = []
        for g in raw_groups:
            try:
                o = str(g.get("original") or "")
                dups = [str(x) for x in (g.get("duplicates") or [])]
                existing = [x for x in ([o] + dups) if x and Path(x).exists()]
                if len(existing) < 2:
                    continue
                # Keep the reported original if it still exists, else pick the first remaining.
                orig = o if o and Path(o).exists() else existing[0]
                rest = [x for x in existing if x != orig]
                pruned.append({"original": orig, "duplicates": rest, "size": int(g.get("size") or 0), "sha256": g.get("sha256")})
            except Exception:
                continue
        self._groups = pruned
        self.group_list.clear()
        total_dups = 0
        for g in self._groups:
            total_dups += len(g.get("duplicates") or [])
        for i, g in enumerate(self._groups):
            o = str(g.get("original") or "")
            ndup = len(g.get("duplicates") or [])
            size = int(g.get("size") or 0)
            mb = (size * max(1, ndup)) / (1024 * 1024)
            it = QListWidgetItem()
            it.setSizeHint(QSize(10, 56))
            self.group_list.addItem(it)
            self.group_list.setItemWidget(it, _QtDuplicateGroupCard(Path(o).name, ndup, mb))
        self.status.setText(f"{total_dups} duplicate files across {len(self._groups)} groups.")
        if self.group_list.count() > 0:
            self.group_list.setCurrentRow(0)

    def _sorted_keep_policy(self, paths: list[str]) -> tuple[str, list[str]]:
        from pathlib import Path
        import os

        if not paths:
            return "", []
        mode = self.keep_policy.currentText()
        if "newest" in mode.lower():
            paths_sorted = sorted(paths, key=lambda s: (-(Path(s).stat().st_mtime if Path(s).exists() else 0.0), s.lower()))
        elif "oldest" in mode.lower():
            paths_sorted = sorted(paths, key=lambda s: ((Path(s).stat().st_mtime if Path(s).exists() else 0.0), s.lower()))
        else:
            # Keep original (as reported): first entry is "original" already.
            return paths[0], paths[1:]
        keep = paths_sorted[0]
        dups = [p for p in paths_sorted[1:]]
        return keep, dups

    def _load_group(self, idx: int):
        from pathlib import Path

        self.table.setRowCount(0)
        self.prev_a.setText("Preview")
        self.prev_b.setText("Compare")
        self.prev_a.setPixmap(QPixmap())
        self.prev_b.setPixmap(QPixmap())
        self._current_files = []
        if idx < 0 or idx >= len(self._groups):
            return
        g = self._groups[idx]
        o = str(g.get("original") or "")
        dups = [str(x) for x in (g.get("duplicates") or [])]
        paths = [o] + dups
        keep, del_list = self._sorted_keep_policy(paths)
        self._current_files = [keep] + del_list

        threshold = float(self.thresh.value()) / 100.0
        # Exact duplicates are 100% confidence.
        conf = 1.0

        for row, fp in enumerate(self._current_files):
            p = Path(fp)
            self.table.insertRow(row)

            chk = QTableWidgetItem("")
            chk.setFlags(Qt.ItemIsUserCheckable | Qt.ItemIsEnabled | Qt.ItemIsSelectable)
            should_delete = (fp != keep) and (conf >= threshold)
            chk.setCheckState(Qt.Checked if should_delete else Qt.Unchecked)
            self.table.setItem(row, 0, chk)

            name = QTableWidgetItem(p.name)
            name.setFlags(name.flags() ^ Qt.ItemIsEditable)
            self.table.setItem(row, 1, name)

            folder = QTableWidgetItem(str(p.parent))
            folder.setFlags(folder.flags() ^ Qt.ItemIsEditable)
            self.table.setItem(row, 2, folder)

            try:
                sz = p.stat().st_size
            except Exception:
                sz = 0
            size_it = QTableWidgetItem(f"{sz/1024/1024:.2f} MB")
            size_it.setTextAlignment(Qt.AlignRight | Qt.AlignVCenter)
            size_it.setFlags(size_it.flags() ^ Qt.ItemIsEditable)
            self.table.setItem(row, 3, size_it)

            c_it = QTableWidgetItem(f"{int(conf*100)}%")
            c_it.setTextAlignment(Qt.AlignRight | Qt.AlignVCenter)
            c_it.setFlags(c_it.flags() ^ Qt.ItemIsEditable)
            self.table.setItem(row, 4, c_it)

        self.table.selectRow(0)
        self._update_preview()

    def _auto_select_current_group(self):
        threshold = float(self.thresh.value()) / 100.0
        conf = 1.0
        for row in range(self.table.rowCount()):
            it = self.table.item(row, 0)
            if not it:
                continue
            # Keep row 0 (current keep) unselected.
            if row == 0:
                it.setCheckState(Qt.Unchecked)
                continue
            it.setCheckState(Qt.Checked if conf >= threshold else Qt.Unchecked)

    def _selected_paths(self) -> list[str]:
        out: list[str] = []
        for row in range(self.table.rowCount()):
            chk = self.table.item(row, 0)
            if not chk or chk.checkState() != Qt.Checked:
                continue
            try:
                name = self.table.item(row, 1).text()
                folder = self.table.item(row, 2).text()
                out.append(str(Path(folder) / name))
            except Exception:
                continue
        return out

    def _current_selected_path(self) -> str | None:
        rows = self.table.selectionModel().selectedRows()
        if not rows:
            return None
        r = int(rows[0].row())
        try:
            name = self.table.item(r, 1).text()
            folder = self.table.item(r, 2).text()
            return str(Path(folder) / name)
        except Exception:
            return None

    def _show_selected(self):
        p = self._current_selected_path()
        if not p:
            return
        try:
            folder = str(Path(p).parent)
            QDesktopServices.openUrl(QUrl.fromLocalFile(folder))
        except Exception:
            pass

    def _preview_pixmap(self, path: str, *, max_size: QSize) -> QPixmap | None:
        try:
            p = Path(path)
        except Exception:
            return None
        suffix = p.suffix.lower()
        try:
            if suffix in {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}:
                pm = _load_oriented_pixmap(str(p), max_size=max_size, smooth=True)
                if pm is not None and not pm.isNull():
                    return pm
        except Exception:
            pass
        try:
            if suffix == ".pdf":
                import fitz

                doc = fitz.open(str(p))
                try:
                    page = doc.load_page(0)
                    pix = page.get_pixmap(matrix=fitz.Matrix(0.35, 0.35), alpha=True)
                    fmt = QImage.Format_RGBA8888 if pix.alpha else QImage.Format_RGB888
                    img = QImage(pix.samples, pix.width, pix.height, pix.stride, fmt).copy()
                    pm = QPixmap.fromImage(img)
                    if not pm.isNull():
                        return pm.scaled(max_size, Qt.KeepAspectRatio, Qt.SmoothTransformation)
                finally:
                    doc.close()
        except Exception:
            pass
        return None

    def _update_preview(self):
        p = self._current_selected_path()
        if not p:
            return
        max_size = QSize(520, 340)
        pm = self._preview_pixmap(p, max_size=max_size)
        if pm is not None and not pm.isNull():
            self.prev_a.setPixmap(pm)
            self.prev_a.setText("")
        else:
            self.prev_a.setPixmap(QPixmap())
            self.prev_a.setText(Path(p).name)

        # Compare with keep (row 0)
        try:
            keep_name = self.table.item(0, 1).text()
            keep_folder = self.table.item(0, 2).text()
            keep_path = str(Path(keep_folder) / keep_name)
        except Exception:
            keep_path = None
        if keep_path and keep_path != p:
            pm2 = self._preview_pixmap(keep_path, max_size=max_size)
            if pm2 is not None and not pm2.isNull():
                self.prev_b.setPixmap(pm2)
                self.prev_b.setText("")
            else:
                self.prev_b.setPixmap(QPixmap())
                self.prev_b.setText(Path(keep_path).name)

    def _delete_selected(self):
        from pathlib import Path
        from core.file_ops import delete_specific_files

        to_del = self._selected_paths()
        if not to_del:
            QMessageBox.information(self, "Delete", "No files selected for deletion.")
            return
        total_mb = 0.0
        for fp in to_del:
            try:
                total_mb += Path(fp).stat().st_size / 1024 / 1024
            except Exception:
                pass
        if QMessageBox.question(
            self,
            "Confirm Delete",
            f"Delete {len(to_del)} file(s) (~{total_mb:.1f} MB)?\n\n"
            "Files are sent to the Recycle Bin when available (recommended).",
        ) != QMessageBox.Yes:
            return

        res = delete_specific_files([Path(p) for p in to_del], use_recycle_bin=True)
        if not res.ok:
            QMessageBox.critical(self, "Delete Failed", res.message)
            return
        QMessageBox.information(self, "Delete", res.message)
        # Update report JSON to remove deleted paths (so future reviews stay accurate).
        try:
            import json

            base = Path(self.report_dir)
            rp = base / "duplicate_detection.json"
            if rp.exists():
                data = json.loads(rp.read_text(encoding="utf-8"))
                groups = list(data.get("groups") or [])
                deleted_set = {str(Path(x)) for x in to_del}
                new_groups = []
                for g in groups:
                    o = str(g.get("original") or "")
                    dups = [str(x) for x in (g.get("duplicates") or [])]
                    keep_paths = [x for x in ([o] + dups) if x and str(Path(x)) not in deleted_set]
                    if len(keep_paths) < 2:
                        continue
                    orig = o if o and str(Path(o)) in keep_paths else keep_paths[0]
                    rest = [x for x in keep_paths if x != orig]
                    new_groups.append({"sha256": g.get("sha256"), "size": g.get("size"), "original": orig, "duplicates": rest})
                data["groups"] = new_groups
                rp.write_text(json.dumps(data, indent=2), encoding="utf-8")
        except Exception:
            pass
        # Reload report UI.
        self._load_report()


class _QtAiHubApplyCategorizeDialog(QDialog):
    def __init__(self, parent: QWidget, *, report_dir: str, selected_categories: list[str]):
        super().__init__(parent)
        self.setWindowTitle("Apply Auto-Categorize")
        self.setModal(True)
        self.setMinimumWidth(520)
        self.setStyleSheet(_qt_modern_dialog_stylesheet())

        layout = QVBoxLayout(self)
        layout.setContentsMargins(14, 14, 14, 14)
        layout.setSpacing(10)

        title = QLabel("Applying organization…")
        title.setObjectName("DialogTitle")
        layout.addWidget(title)

        self.status = QLabel("Starting…")
        self.status.setObjectName("DialogSubtitle")
        layout.addWidget(self.status)

        self.bar = QProgressBar()
        self.bar.setRange(0, 1000)
        self.bar.setValue(0)
        layout.addWidget(self.bar)

        row = QHBoxLayout()
        row.addStretch(1)
        self.btn_cancel = QPushButton("Cancel")
        self.btn_cancel.clicked.connect(self._cancel)
        row.addWidget(self.btn_cancel)
        layout.addLayout(row)

        self.worker = _QtAiHubApplyCategorizeWorker(report_dir=report_dir, selected_categories=selected_categories)
        self._thread = QThread(self)
        self.worker.moveToThread(self._thread)
        self._thread.started.connect(self.worker.run)
        self.worker.status.connect(self.status.setText)
        self.worker.progress.connect(lambda p: self.bar.setValue(int(p * 1000)))
        self.worker.finished.connect(self._on_finished)
        self.worker.error.connect(self._on_error)
        self.worker.finished.connect(self._thread.quit)
        self.worker.error.connect(self._thread.quit)
        self._thread.finished.connect(self._thread.deleteLater)
        self._thread.start()

    def _cancel(self):
        self.btn_cancel.setEnabled(False)
        try:
            self.worker.cancel()
        except Exception:
            pass
        self.status.setText("Cancelling…")

    def _on_finished(self, ok: bool, moved: int, skipped: int):
        self.bar.setValue(1000 if ok else self.bar.value())
        self.status.setText(f"Done. Moved: {moved}  Skipped: {skipped}")
        self.btn_cancel.setText("Close")
        self.btn_cancel.setEnabled(True)
        try:
            self.btn_cancel.clicked.disconnect()
        except Exception:
            pass
        self.btn_cancel.clicked.connect(self.accept)

    def _on_error(self, msg: str):
        QMessageBox.critical(self, "Apply Failed", msg)
        self.reject()


class _QtAutoCategorizeWorker(QObject):
    status = Signal(str)
    progress = Signal(float)
    finished = Signal(dict, dict)  # report, categorized
    error = Signal(str)

    def __init__(
        self,
        *,
        backend,
        target_folder: str,
        include_subfolders: bool,
        smart_scope: bool,
        include_other: bool,
        use_ai_vision: bool,
        use_ai_docs: bool,
    ):
        super().__init__()
        self.backend = backend
        self.target_folder = target_folder
        self.include_subfolders = include_subfolders
        self.smart_scope = smart_scope
        self.include_other = include_other
        self.use_ai_vision = use_ai_vision
        self.use_ai_docs = use_ai_docs
        self.cancelled = False

    def cancel(self):
        self.cancelled = True

    def run(self):
        from pathlib import Path
        from datetime import datetime
        import json

        try:
            target = Path(self.target_folder)
            if not target.exists():
                self.error.emit("Target folder does not exist.")
                return

            ai = getattr(self.backend, "ai_manager", None) if self.backend else None
            needs_ai = bool(self.use_ai_vision) or bool(self.use_ai_docs)
            if needs_ai and not ai:
                self.error.emit("AI Manager is not available.")
                return
            if needs_ai and ai and (not getattr(ai, "is_ready", False)):
                self.error.emit("AI model is not loaded. Load it before starting Auto-Categorize.")
                return

            from core.enhanced_categorizer import EnhancedCategorizer

            def prog_cb(msg: str, p: float, cur: int, total: int):
                self.status.emit(f"{msg} {cur}/{total}")
                # keep in 0.3..1.0 (0..0.3 reserved for model prep)
                self.progress.emit(0.3 + 0.7 * float(p))

            cat = EnhancedCategorizer(ai_manager=ai, use_ai_vision=bool(self.use_ai_vision))
            categorized = cat.categorize_folder(
                target,
                include_subfolders=bool(self.include_subfolders),
                progress_callback=prog_cb,
                cancel_check=lambda: bool(self.cancelled),
                use_ai_vision=bool(self.use_ai_vision),
                smart_scope=bool(self.smart_scope),
                include_other=bool(self.include_other),
                use_ai_documents=bool(self.use_ai_docs),
            )
            categorized_ser = {k: ([str(p) for p in v] if isinstance(v, list) else v) for k, v in categorized.items()}
            counts = {k: len(v) for k, v in categorized_ser.items() if isinstance(v, list)}

            report = {
                "version": 1,
                "generated_at": datetime.now().isoformat(),
                "target_folder": str(target),
                "apply_requested": False,
                "options": {
                    "include_subfolders": bool(self.include_subfolders),
                    "smart_scope": bool(self.smart_scope),
                    "include_other": bool(self.include_other),
                    "use_ai_vision": bool(self.use_ai_vision),
                    "use_ai_documents": bool(self.use_ai_docs),
                },
                "counts": counts,
                "categories": categorized_ser,
            }
            self.progress.emit(1.0)
            self.finished.emit(report, categorized_ser)
        except Exception as e:
            self.error.emit(str(e))


class _QtAutoCategorizeDialog(QDialog):
    """
    Qt port of the old CustomTkinter Auto-Categorize dialog:
    - Top options (include subfolders, smart scope, include other, AI vision, AI scanned PDFs, move/apply)
    - Category preview list with checkboxes
    - Re-run, Undo, Apply Organization, Export report
    """

    def __init__(
        self,
        parent: QWidget,
        *,
        backend,
        report_dir: str | None = None,
        target_folder: str | None = None,
        autorun: bool = False,
        initial_options: dict | None = None,
    ):
        super().__init__(parent)
        self.setWindowTitle("Auto-Categorize - 51 Categories")
        self.setModal(True)
        self.setMinimumWidth(860)
        self.setMinimumHeight(620)
        self.setStyleSheet(_qt_modern_dialog_stylesheet())

        self.backend = backend
        self._report_dir = report_dir
        self._report: dict = {}
        self._categorized: dict = {}
        self._target_folder: str = ""
        self._autorun = bool(autorun)
        self._initial_options = dict(initial_options or {})

        layout = QVBoxLayout(self)
        layout.setContentsMargins(14, 14, 14, 14)
        layout.setSpacing(10)

        header = QFrame()
        header.setObjectName("DialogHeader")
        header_l = QHBoxLayout(header)
        header_l.setContentsMargins(12, 10, 12, 10)
        header_l.setSpacing(10)

        tcol = QVBoxLayout()
        tcol.setContentsMargins(0, 0, 0, 0)
        tcol.setSpacing(2)
        title = QLabel("Enhanced Categorization - 51 Categories")
        title.setObjectName("DialogTitle")
        tcol.addWidget(title)
        self.subtitle = QLabel("Ready.")
        self.subtitle.setObjectName("DialogSubtitle")
        tcol.addWidget(self.subtitle)
        header_l.addLayout(tcol, 1)

        self.btn_export = QPushButton("Export")
        self.btn_export.setObjectName("SecondaryButton")
        self.btn_export.clicked.connect(self._export_report)
        header_l.addWidget(self.btn_export, 0, Qt.AlignRight)

        layout.addWidget(header)

        opts_box = QGroupBox()
        opts_box.setObjectName("Card")
        opts = QGridLayout(opts_box)
        opts.setContentsMargins(12, 12, 12, 12)
        opts.setHorizontalSpacing(16)
        opts.setVerticalSpacing(8)

        self.cb_subfolders = QCheckBox("Include subfolders")
        self.cb_subfolders.setChecked(True)
        self.cb_smart_scope = QCheckBox("Smart scope (skip projects/caches)")
        self.cb_smart_scope.setChecked(True)
        self.cb_other = QCheckBox("Include 'Other' (move unknown formats)")
        self.cb_other.setChecked(False)
        self.cb_ai_vision = QCheckBox("AI vision (images)")
        self.cb_ai_docs = QCheckBox("AI (scanned PDFs)")
        self.cb_move = QCheckBox("Move files to category folders")
        self.cb_move.setChecked(False)

        opts.addWidget(self.cb_subfolders, 0, 0)
        opts.addWidget(self.cb_smart_scope, 0, 1)
        opts.addWidget(self.cb_other, 1, 0)
        opts.addWidget(self.cb_move, 1, 1)
        opts.addWidget(self.cb_ai_vision, 2, 0)
        opts.addWidget(self.cb_ai_docs, 2, 1)

        # Right-side action row
        right = QVBoxLayout()
        right.setContentsMargins(0, 0, 0, 0)
        right.setSpacing(8)
        self.btn_rerun = QPushButton("Re-run")
        self.btn_rerun.setObjectName("SecondaryButton")
        self.btn_rerun.clicked.connect(self._run)
        right.addWidget(self.btn_rerun)
        self.btn_apply = QPushButton("Apply Organization")
        self.btn_apply.setObjectName("PrimaryButton")
        self.btn_apply.clicked.connect(self._apply)
        right.addWidget(self.btn_apply)
        right.addStretch(1)
        opts.addLayout(right, 0, 2, 3, 1)

        layout.addWidget(opts_box)

        self.bar = QProgressBar()
        self.bar.setRange(0, 1000)
        self.bar.setValue(0)
        self.bar.setObjectName("Progress")
        layout.addWidget(self.bar)

        self.summary = QLabel("No analysis yet.")
        self.summary.setObjectName("DialogSummary")
        layout.addWidget(self.summary)

        self.tree = QTreeWidget()
        self.tree.setHeaderLabels(["Category", "Count / Sample files"])
        self.tree.setObjectName("ResultsTree")
        layout.addWidget(self.tree, 1)

        bottom = QHBoxLayout()
        self.btn_cancel = QPushButton("Close")
        self.btn_cancel.setObjectName("SecondaryButton")
        self.btn_cancel.clicked.connect(self.reject)
        bottom.addWidget(self.btn_cancel)
        self.btn_undo = QPushButton("Undo")
        self.btn_undo.setObjectName("SecondaryButton")
        self.btn_undo.clicked.connect(self._undo)
        bottom.addWidget(self.btn_undo)

        bottom.addStretch(1)
        layout.addLayout(bottom)

        self._refresh_undo_label()
        if report_dir:
            try:
                self._load_from_report_dir(report_dir)
            except Exception:
                pass

        if target_folder:
            self._target_folder = str(target_folder)

        # Apply any initial options provided (e.g., coming from AI Hub selections).
        if self._initial_options:
            try:
                if "include_subfolders" in self._initial_options:
                    self.cb_subfolders.setChecked(bool(self._initial_options["include_subfolders"]))
                if "smart_scope" in self._initial_options:
                    self.cb_smart_scope.setChecked(bool(self._initial_options["smart_scope"]))
                if "include_other" in self._initial_options:
                    self.cb_other.setChecked(bool(self._initial_options["include_other"]))
                if "use_ai_vision" in self._initial_options:
                    self.cb_ai_vision.setChecked(bool(self._initial_options["use_ai_vision"]))
                if "use_ai_documents" in self._initial_options:
                    self.cb_ai_docs.setChecked(bool(self._initial_options["use_ai_documents"]))
                if "move_files" in self._initial_options:
                    self.cb_move.setChecked(bool(self._initial_options["move_files"]))
            except Exception:
                pass

        # Match the old behavior: run immediately with defaults, user can tweak and hit Re-run.
        if self._autorun and self._target_folder:
            from PySide6.QtCore import QTimer

            QTimer.singleShot(50, self._run)

    def _refresh_undo_label(self):
        try:
            from utils.universal_undo import get_undo_manager

            stats = get_undo_manager().get_statistics()
            n = int(stats.get("undoable_transactions") or 0)
            self.btn_undo.setText(f"Undo ({n})" if n else "Undo")
            self.btn_undo.setEnabled(n > 0)
        except Exception:
            self.btn_undo.setEnabled(True)

    def _load_from_report_dir(self, report_dir: str):
        from pathlib import Path
        import json

        base = Path(report_dir)
        p = base / "auto_categorize.json"
        if not p.exists():
            return
        data = json.loads(p.read_text(encoding="utf-8"))
        self._report = data
        self._categorized = dict(data.get("categories") or {})
        self._target_folder = str(data.get("target_folder") or "")
        opts = dict(data.get("options") or {})
        self.cb_subfolders.setChecked(bool(opts.get("include_subfolders", True)))
        self.cb_smart_scope.setChecked(bool(opts.get("smart_scope", True)))
        self.cb_other.setChecked(bool(opts.get("include_other", False)))
        self.cb_ai_vision.setChecked(bool(opts.get("use_ai_vision", False)))
        self.cb_ai_docs.setChecked(bool(opts.get("use_ai_documents", False)))
        self._render_tree()

    def _render_tree(self):
        from pathlib import Path

        self.tree.clear()
        counts = {k: len(v) for k, v in self._categorized.items() if isinstance(v, list)}
        total = sum(counts.values())
        self.summary.setText(f"Found {total} files in {len(counts)} categories")

        for cat_name, file_list in sorted(self._categorized.items(), key=lambda kv: len(kv[1]) if isinstance(kv[1], list) else 0, reverse=True):
            if not isinstance(file_list, list):
                continue
            top = QTreeWidgetItem([str(cat_name), str(len(file_list))])
            top.setFlags(top.flags() | Qt.ItemIsUserCheckable)
            default_checked = cat_name not in {"empty_folders", "ignored_projects"}
            top.setCheckState(0, Qt.Checked if default_checked else Qt.Unchecked)
            for fp in file_list[:12]:
                try:
                    p = Path(fp)
                    child_txt = p.name
                except Exception:
                    child_txt = str(fp)
                top.addChild(QTreeWidgetItem(["", child_txt]))
            self.tree.addTopLevelItem(top)
        self.tree.expandToDepth(0)
        self.tree.resizeColumnToContents(0)

    def _run(self):
        # If we came from AI Hub report, keep same target. Otherwise, try to infer from report or backend selection.
        target = self._target_folder
        if not target:
            QMessageBox.information(self, "Auto-Categorize", "No target folder available. Run from AI Hub with a target folder.")
            return

        needs_loaded_ai = bool(self.cb_ai_vision.isChecked()) or bool(self.cb_ai_docs.isChecked())
        if needs_loaded_ai:
            top = self.parent().window() if self.parent() else None
            ensure = getattr(top, "_ensure_ai_ready", None)
            if callable(ensure) and not ensure(title="Prepare Auto-Categorize", kind="vision"):
                return

        self.btn_rerun.setEnabled(False)
        self.btn_apply.setEnabled(False)
        self.subtitle.setText("Scanning…")
        self.bar.setValue(0)

        self.worker = _QtAutoCategorizeWorker(
            backend=self.backend,
            target_folder=target,
            include_subfolders=self.cb_subfolders.isChecked(),
            smart_scope=self.cb_smart_scope.isChecked(),
            include_other=self.cb_other.isChecked(),
            use_ai_vision=self.cb_ai_vision.isChecked(),
            use_ai_docs=self.cb_ai_docs.isChecked(),
        )
        self._thread = QThread(self)
        self.worker.moveToThread(self._thread)
        self._thread.started.connect(self.worker.run)
        self.worker.status.connect(self.subtitle.setText)
        self.worker.progress.connect(lambda p: self.bar.setValue(int(p * 1000)))
        self.worker.finished.connect(self._on_run_finished)
        self.worker.error.connect(self._on_run_error)
        self.worker.finished.connect(self._thread.quit)
        self.worker.error.connect(self._thread.quit)
        self._thread.finished.connect(self._thread.deleteLater)
        self._thread.start()

    def _on_run_finished(self, report: dict, categorized: dict):
        self._report = report
        self._categorized = categorized
        self._target_folder = str(report.get("target_folder") or self._target_folder)
        self._render_tree()
        self.btn_rerun.setEnabled(True)
        self.btn_apply.setEnabled(True)
        self.subtitle.setText("Analysis complete.")

    def _on_run_error(self, msg: str):
        QMessageBox.critical(self, "Auto-Categorize Failed", msg)
        self.btn_rerun.setEnabled(True)
        self.btn_apply.setEnabled(True)

    def _selected_categories(self) -> list[str]:
        selected = []
        for i in range(self.tree.topLevelItemCount()):
            it = self.tree.topLevelItem(i)
            if it and it.checkState(0) == Qt.Checked:
                selected.append(it.text(0))
        return selected

    def _apply(self):
        if not self.cb_move.isChecked():
            QMessageBox.information(self, "Apply Organization", "Enable 'Move files to category folders' to apply changes.")
            return
        if not self._target_folder or not self._categorized:
            QMessageBox.information(self, "Apply Organization", "No categorized results to apply.")
            return

        from pathlib import Path
        import shutil

        try:
            from core.enhanced_categorizer import EnhancedCategorizer
            from utils.universal_undo import record_categorize
        except Exception:
            QMessageBox.critical(self, "Apply Organization", "Missing required components.")
            return

        target = Path(self._target_folder)
        cat = EnhancedCategorizer(ai_manager=None, use_ai_vision=False)
        chosen = set(self._selected_categories())

        moved_pairs = []
        moved = 0
        skipped = 0
        for category in chosen:
            files = self._categorized.get(category)
            if not isinstance(files, list):
                continue
            if category in {"empty_folders", "ignored_projects"}:
                continue
            folder_rel = cat.get_category_folder(category)
            dest_dir = (target / folder_rel)
            dest_dir.mkdir(parents=True, exist_ok=True)
            for src in files:
                try:
                    sp = Path(src)
                    if not sp.exists():
                        skipped += 1
                        continue
                    try:
                        if dest_dir.resolve() in sp.resolve().parents:
                            skipped += 1
                            continue
                    except Exception:
                        pass
                    out = dest_dir / sp.name
                    if out.exists():
                        stem = out.stem
                        suf = out.suffix
                        n = 1
                        while out.exists():
                            out = dest_dir / f"{stem}_{n}{suf}"
                            n += 1
                    shutil.move(str(sp), str(out))
                    moved_pairs.append((sp, out))
                    moved += 1
                except Exception:
                    skipped += 1
                    continue

        if moved_pairs:
            try:
                record_categorize(moved_pairs, metadata={"qt_dialog": "auto_categorize"})
            except Exception:
                pass
        self._refresh_undo_label()
        QMessageBox.information(self, "Apply Organization", f"Moved: {moved}\nSkipped: {skipped}")

    def _undo(self):
        try:
            from utils.universal_undo import undo_last_operation

            ok, msg, _n = undo_last_operation()
            QMessageBox.information(self, "Undo", msg)
            self._refresh_undo_label()
            if ok and self._report_dir:
                try:
                    self._load_from_report_dir(self._report_dir)
                except Exception:
                    pass
        except Exception as e:
            QMessageBox.critical(self, "Undo", str(e))

    def _export_report(self):
        from pathlib import Path
        import json
        if not self._report:
            QMessageBox.information(self, "Export Report", "No report to export yet.")
            return
        out, _ = QFileDialog.getSaveFileName(self, "Export Auto-Categorize Report", "auto_categorize.json", "JSON (*.json)")
        if not out:
            return
        try:
            Path(out).write_text(json.dumps(self._report, indent=2), encoding="utf-8")
            QMessageBox.information(self, "Export Report", f"Saved: {out}")
        except Exception as e:
            QMessageBox.critical(self, "Export Report", str(e))


def _qt_scan_files(folder: Path, *, include_subfolders: bool, allowed_exts: set[str] | None) -> list[Path]:
    import os

    out: list[Path] = []
    folder = Path(folder)
    if include_subfolders:
        for root, _dirs, fnames in os.walk(folder):
            for name in fnames:
                p = Path(root) / name
                if not p.is_file():
                    continue
                if allowed_exts is not None and p.suffix.lower() not in allowed_exts:
                    continue
                out.append(p)
    else:
        for p in folder.iterdir():
            if not p.is_file():
                continue
            if allowed_exts is not None and p.suffix.lower() not in allowed_exts:
                continue
            out.append(p)
    return out


class _QtSmartRenameWorker(QObject):
    status = Signal(str)
    progress = Signal(float)
    finished = Signal(dict, list)  # report, items
    error = Signal(str)

    def __init__(
        self,
        *,
        backend,
        target_folder: str,
        include_subfolders: bool,
        filter_key: str,
        use_vision: bool,
        use_ai_docs: bool,
    ):
        super().__init__()
        self.backend = backend
        self.target_folder = target_folder
        self.include_subfolders = include_subfolders
        self.filter_key = filter_key
        self.use_vision = use_vision
        self.use_ai_docs = use_ai_docs
        self.cancelled = False

    def cancel(self):
        self.cancelled = True

    def run(self):
        from pathlib import Path
        from datetime import datetime

        try:
            target = Path(self.target_folder)
            if not target.exists():
                self.error.emit("Target folder does not exist.")
                return

            ai = getattr(self.backend, "ai_manager", None) if self.backend else None
            if not ai:
                self.error.emit("AI Manager is not available.")
                return

            needs_ai = bool(self.use_vision) or bool(self.use_ai_docs)
            if needs_ai and (not getattr(ai, "is_ready", False)):
                self.error.emit("AI model is not loaded. Load it before starting Smart Rename.")
                return

            # Determine file filter.
            from core.bulk_ai_processor import BulkAIProcessor

            fk = (self.filter_key or "all").lower()
            allowed: set[str] | None = None
            if fk == "images":
                allowed = set(BulkAIProcessor.get_image_extensions())
            elif fk == "videos":
                allowed = set(BulkAIProcessor.get_video_extensions())
            elif fk == "documents":
                allowed = set(BulkAIProcessor.get_document_extensions())
            elif fk == "code":
                allowed = set(BulkAIProcessor.get_code_extensions())

            files = _qt_scan_files(target, include_subfolders=bool(self.include_subfolders), allowed_exts=allowed)
            total = len(files)
            self.status.emit(f"Found {total} files")
            # If we loaded a model above, progress is already ~0.3. Keep analysis in 0.3..1.0.
            base_prog = 0.02 if not needs_ai else 0.3
            self.progress.emit(base_prog)

            from utils.intelligent_rename import sanitize_ai_filename, get_unique_filename
            from core.semantic_analyzer import SemanticAnalyzer

            sem = SemanticAnalyzer(ai) if (self.use_ai_docs and getattr(ai, "is_ready", False)) else None
            image_exts = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp", ".tiff", ".tif"}

            items = []
            for i, fp in enumerate(files, start=1):
                if self.cancelled:
                    break
                if (i % 10) == 0 or i == total:
                    self.status.emit(f"Analyzing {i}/{total}…")
                    self.progress.emit(base_prog + (1.0 - base_prog) * float(i) / float(max(1, total)))

                used_ai = False
                suggested = None
                try:
                    if sem and fp.suffix.lower() in {".pdf", ".docx", ".doc"}:
                        res = sem.analyze_document(fp)
                        if res and res.suggested_filename:
                            suggested = res.suggested_filename
                    if not suggested:
                        if self.use_vision and fp.suffix.lower() in image_exts and getattr(ai, "is_ready", False):
                            used_ai = True
                        suggested = ai.analyze_file_for_rename(fp, use_ai=used_ai)
                        if (not suggested) and used_ai:
                            suggested = ai.analyze_file_for_rename(fp, use_ai=False)
                except Exception:
                    suggested = None

                if not suggested:
                    continue
                validation = sanitize_ai_filename(str(suggested), preserve_case=True)
                desired = validation.sanitized_name
                if not desired or desired.strip() == fp.stem:
                    continue
                unique_name, _why = get_unique_filename(fp, desired, {"used_ai": used_ai})
                new_path = fp.with_name(unique_name + fp.suffix)
                if str(new_path) == str(fp):
                    continue
                items.append({"from": str(fp), "to": str(new_path), "used_ai": used_ai})

            report = {
                "version": 1,
                "generated_at": datetime.now().isoformat(),
                "target_folder": str(target),
                "options": {
                    "include_subfolders": bool(self.include_subfolders),
                    "filter_key": fk,
                    "use_vision": bool(self.use_vision),
                    "use_ai_docs": bool(self.use_ai_docs),
                },
                "items": items,
            }
            self.finished.emit(report, items)
        except Exception as e:
            self.error.emit(str(e))


class _QtSmartRenameApplyWorker(QObject):
    status = Signal(str)
    progress = Signal(float)
    finished = Signal(bool, int, int)  # ok, renamed, skipped
    error = Signal(str)

    def __init__(self, *, items: list[dict]):
        super().__init__()
        self.items = items
        self.cancelled = False

    def cancel(self):
        self.cancelled = True

    def run(self):
        from pathlib import Path
        try:
            from utils.universal_undo import record_bulk_rename
        except Exception:
            record_bulk_rename = None  # type: ignore

        renamed = 0
        skipped = 0
        pairs = []
        total = len(self.items)
        for i, it in enumerate(self.items, start=1):
            if self.cancelled:
                break
            self.status.emit(f"Renaming {i}/{total}…")
            self.progress.emit(float(i - 1) / float(max(1, total)))
            try:
                src = Path(str(it.get("from") or ""))
                dst = Path(str(it.get("to") or ""))
                if not src.exists():
                    skipped += 1
                    continue
                try:
                    if dst.parent.resolve() != src.parent.resolve():
                        skipped += 1
                        continue
                except Exception:
                    if str(dst.parent) != str(src.parent):
                        skipped += 1
                        continue
                if dst.exists() and dst.resolve() != src.resolve():
                    skipped += 1
                    continue
                src.rename(dst)
                renamed += 1
                pairs.append((src, dst))
            except Exception:
                skipped += 1
                continue

        if record_bulk_rename and pairs:
            try:
                record_bulk_rename([(a, b) for (a, b) in pairs], metadata={"qt_dialog": "smart_rename"})
            except Exception:
                pass

        self.progress.emit(1.0)
        self.finished.emit(not self.cancelled, renamed, skipped)


class _QtSmartRenameDialog(QDialog):
    def __init__(
        self,
        parent: QWidget,
        *,
        backend,
        target_folder: str,
        autorun: bool = True,
        initial_options: dict | None = None,
    ):
        super().__init__(parent)
        self.setWindowTitle("Smart Rename - Fast Cleanup")
        self.setModal(True)
        self.setMinimumWidth(900)
        self.setMinimumHeight(620)
        self.setStyleSheet(_qt_modern_dialog_stylesheet())

        self.backend = backend
        self._target_folder = str(target_folder)
        self._report: dict = {}
        self._items: list[dict] = []
        self._autorun = bool(autorun)
        self._initial_options = dict(initial_options or {})

        layout = QVBoxLayout(self)
        layout.setContentsMargins(14, 14, 14, 14)
        layout.setSpacing(10)

        header = QFrame()
        header.setObjectName("DialogHeader")
        hl = QHBoxLayout(header)
        hl.setContentsMargins(12, 10, 12, 10)
        hl.setSpacing(10)

        tcol = QVBoxLayout()
        tcol.setContentsMargins(0, 0, 0, 0)
        tcol.setSpacing(2)
        title = QLabel("Smart Rename - AI Vision + Rules")
        title.setObjectName("DialogTitle")
        tcol.addWidget(title)
        self.subtitle = QLabel("Ready.")
        self.subtitle.setObjectName("DialogSubtitle")
        tcol.addWidget(self.subtitle)
        hl.addLayout(tcol, 1)
        layout.addWidget(header)

        opts = QGroupBox()
        opts.setObjectName("Card")
        ol = QGridLayout(opts)
        ol.setContentsMargins(12, 12, 12, 12)
        ol.setHorizontalSpacing(12)
        ol.setVerticalSpacing(8)

        ol.addWidget(QLabel("Mode:"), 0, 0)
        self.mode = QComboBox()
        self.mode.addItems(["Fast (rules only)", "Smart (balanced)", "Deep (AI vision)"])
        self.mode.currentTextChanged.connect(self._mode_changed)
        ol.addWidget(self.mode, 0, 1)

        self.cb_subfolders = QCheckBox("Include subfolders (recursive)")
        self.cb_subfolders.setChecked(True)
        ol.addWidget(self.cb_subfolders, 0, 2, 1, 2)

        ol.addWidget(QLabel("Filter:"), 1, 0)
        self.filter = QComboBox()
        self.filter.addItem("All Files", userData="all")
        self.filter.addItem("Images Only", userData="images")
        self.filter.addItem("Videos Only", userData="videos")
        self.filter.addItem("Documents Only", userData="documents")
        self.filter.addItem("Code Files Only", userData="code")
        ol.addWidget(self.filter, 1, 1)

        self.cb_ai_vision = QCheckBox("Use AI vision for images (slower, more accurate)")
        self.cb_ai_docs = QCheckBox("Use AI for scanned PDFs (invoices, statements, etc)")
        ol.addWidget(self.cb_ai_vision, 1, 2, 1, 2)
        ol.addWidget(self.cb_ai_docs, 2, 2, 1, 2)

        self.btn_scan = QPushButton("Scan")
        self.btn_scan.setObjectName("SecondaryButton")
        self.btn_scan.clicked.connect(self._scan)
        ol.addWidget(self.btn_scan, 2, 0, 1, 2)

        layout.addWidget(opts)

        self.bar = QProgressBar()
        self.bar.setRange(0, 1000)
        self.bar.setValue(0)
        layout.addWidget(self.bar)

        self.tree = QTreeWidget()
        self.tree.setObjectName("ResultsTree")
        self.tree.setHeaderLabels(["", "Original", "New Name"])
        self.tree.setColumnWidth(0, 36)
        layout.addWidget(self.tree, 1)

        bottom = QHBoxLayout()
        self.btn_select_all = QPushButton("Select All")
        self.btn_select_all.setObjectName("SecondaryButton")
        self.btn_select_all.clicked.connect(lambda: self._set_all_checked(True))
        bottom.addWidget(self.btn_select_all)
        self.btn_select_none = QPushButton("Select None")
        self.btn_select_none.setObjectName("SecondaryButton")
        self.btn_select_none.clicked.connect(lambda: self._set_all_checked(False))
        bottom.addWidget(self.btn_select_none)

        self.btn_undo = QPushButton("Undo")
        self.btn_undo.setObjectName("SecondaryButton")
        self.btn_undo.clicked.connect(self._undo)
        bottom.addWidget(self.btn_undo)
        bottom.addStretch(1)

        self.btn_cancel = QPushButton("Close")
        self.btn_cancel.setObjectName("SecondaryButton")
        self.btn_cancel.clicked.connect(self.reject)
        bottom.addWidget(self.btn_cancel)

        self.btn_apply = QPushButton("Apply Renames")
        self.btn_apply.setObjectName("PrimaryButton")
        self.btn_apply.clicked.connect(self._apply)
        bottom.addWidget(self.btn_apply)
        layout.addLayout(bottom)

        self._refresh_undo_label()
        self._apply_initial_options()
        if self._autorun:
            from PySide6.QtCore import QTimer

            QTimer.singleShot(50, self._scan)

    def _apply_initial_options(self):
        if not self._initial_options:
            self._mode_changed(self.mode.currentText())
            return
        try:
            if "include_subfolders" in self._initial_options:
                self.cb_subfolders.setChecked(bool(self._initial_options["include_subfolders"]))
            if "filter_key" in self._initial_options:
                fk = str(self._initial_options["filter_key"] or "all")
                for i in range(self.filter.count()):
                    if str(self.filter.itemData(i) or "") == fk:
                        self.filter.setCurrentIndex(i)
                        break
        except Exception:
            pass
        self._mode_changed(self.mode.currentText())

    def _mode_changed(self, text: str):
        t = (text or "").lower()
        if "deep" in t:
            self.cb_ai_vision.setChecked(True)
        elif "fast" in t:
            self.cb_ai_vision.setChecked(False)
            self.cb_ai_docs.setChecked(False)

    def _set_all_checked(self, checked: bool):
        for i in range(self.tree.topLevelItemCount()):
            it = self.tree.topLevelItem(i)
            if it is None:
                continue
            it.setCheckState(0, Qt.Checked if checked else Qt.Unchecked)

    def _scan(self):
        fk = str(self.filter.currentData() or "all")
        needs_loaded_ai = bool(self.cb_ai_vision.isChecked()) or bool(self.cb_ai_docs.isChecked())
        if needs_loaded_ai:
            top = self.parent().window() if self.parent() else None
            ensure = getattr(top, "_ensure_ai_ready", None)
            if callable(ensure) and not ensure(title="Prepare Smart Rename", kind="vision"):
                return

        self.btn_scan.setEnabled(False)
        self.btn_apply.setEnabled(False)
        self.subtitle.setText("Scanning…")
        self.tree.clear()
        self.bar.setValue(0)

        self.worker = _QtSmartRenameWorker(
            backend=self.backend,
            target_folder=self._target_folder,
            include_subfolders=bool(self.cb_subfolders.isChecked()),
            filter_key=fk,
            use_vision=bool(self.cb_ai_vision.isChecked()),
            use_ai_docs=bool(self.cb_ai_docs.isChecked()),
        )
        self._thread = QThread(self)
        self.worker.moveToThread(self._thread)
        self._thread.started.connect(self.worker.run)
        self.worker.status.connect(self.subtitle.setText)
        self.worker.progress.connect(lambda p: self.bar.setValue(int(p * 1000)))
        self.worker.finished.connect(self._on_scan_finished)
        self.worker.error.connect(self._on_scan_error)
        self.worker.finished.connect(self._thread.quit)
        self.worker.error.connect(self._thread.quit)
        self._thread.finished.connect(self._thread.deleteLater)
        self._thread.start()

    def _on_scan_finished(self, report: dict, items: list):
        self._report = report
        self._items = list(items or [])
        self.subtitle.setText(f"Ready to rename {len(self._items)} files")
        self.btn_scan.setEnabled(True)
        self.btn_apply.setEnabled(bool(self._items))
        self.bar.setValue(1000)

        for it in self._items:
            orig = str(it.get("from") or "")
            new = str(it.get("to") or "")
            row = QTreeWidgetItem(["", orig, new])
            row.setFlags(row.flags() | Qt.ItemIsUserCheckable)
            row.setCheckState(0, Qt.Checked)
            self.tree.addTopLevelItem(row)
        self.tree.resizeColumnToContents(0)

    def _on_scan_error(self, msg: str):
        QMessageBox.critical(self, "Smart Rename Failed", msg)
        self.btn_scan.setEnabled(True)
        self.btn_apply.setEnabled(bool(self._items))

    def _selected_items(self) -> list[dict]:
        selected = []
        # Map from tree rows to underlying items by index.
        for i in range(min(self.tree.topLevelItemCount(), len(self._items))):
            w = self.tree.topLevelItem(i)
            if w and w.checkState(0) == Qt.Checked:
                selected.append(self._items[i])
        return selected

    def _apply(self):
        sel = self._selected_items()
        if not sel:
            QMessageBox.information(self, "Apply Renames", "No items selected.")
            return

        self.btn_apply.setEnabled(False)
        self.btn_scan.setEnabled(False)
        self.subtitle.setText("Renaming…")

        self.apply_worker = _QtSmartRenameApplyWorker(items=sel)
        self._apply_thread = QThread(self)
        self.apply_worker.moveToThread(self._apply_thread)
        self._apply_thread.started.connect(self.apply_worker.run)
        self.apply_worker.status.connect(self.subtitle.setText)
        self.apply_worker.progress.connect(lambda p: self.bar.setValue(int(p * 1000)))
        self.apply_worker.finished.connect(self._on_apply_finished)
        self.apply_worker.error.connect(self._on_apply_error)
        self.apply_worker.finished.connect(self._apply_thread.quit)
        self.apply_worker.error.connect(self._apply_thread.quit)
        self._apply_thread.finished.connect(self._apply_thread.deleteLater)
        self._apply_thread.start()

    def _on_apply_finished(self, ok: bool, renamed: int, skipped: int):
        self.bar.setValue(1000 if ok else self.bar.value())
        self.subtitle.setText(f"Smart Rename: renamed {renamed} files (skipped {skipped})")
        self._refresh_undo_label()
        self.btn_scan.setEnabled(True)
        self.btn_apply.setEnabled(False)

    def _on_apply_error(self, msg: str):
        QMessageBox.critical(self, "Apply Renames Failed", msg)
        self.btn_scan.setEnabled(True)
        self.btn_apply.setEnabled(True)

    def _refresh_undo_label(self):
        try:
            from utils.universal_undo import get_undo_manager

            stats = get_undo_manager().get_statistics()
            n = int(stats.get("undoable_transactions") or 0)
            self.btn_undo.setText(f"Undo ({n})" if n else "Undo")
            self.btn_undo.setEnabled(n > 0)
        except Exception:
            self.btn_undo.setEnabled(True)

    def _undo(self):
        try:
            from utils.universal_undo import undo_last_operation

            ok, msg, _n = undo_last_operation()
            QMessageBox.information(self, "Undo", msg)
            self._refresh_undo_label()
        except Exception as e:
            QMessageBox.critical(self, "Undo", str(e))


class _QtContentAnalysisWorker(QObject):
    status = Signal(str)
    progress = Signal(float)
    finished = Signal(dict, list)  # report, items
    error = Signal(str)

    def __init__(self, *, backend, target_folder: str, include_subfolders: bool, max_files: int):
        super().__init__()
        self.backend = backend
        self.target_folder = target_folder
        self.include_subfolders = include_subfolders
        self.max_files = max_files
        self.cancelled = False

    def cancel(self):
        self.cancelled = True

    def run(self):
        from pathlib import Path
        from datetime import datetime

        try:
            target = Path(self.target_folder)
            if not target.exists():
                self.error.emit("Target folder does not exist.")
                return

            ai = getattr(self.backend, "ai_manager", None) if self.backend else None
            if not ai:
                self.error.emit("AI Manager is not available.")
                return

            if not getattr(ai, "is_ready", False):
                self.error.emit("AI model is not loaded. Load it before starting Content Analysis.")
                return

            from core.bulk_ai_processor import BulkAIProcessor, ProcessingOptions, ProcessingMode
            from core.semantic_analyzer import SemanticAnalyzer

            sem = SemanticAnalyzer(ai)
            proc = BulkAIProcessor(ai)
            options = ProcessingOptions(
                include_subfolders=bool(self.include_subfolders),
                file_extensions=BulkAIProcessor.get_document_extensions(),
                batch_size=50,
                mode=ProcessingMode.SMART,
                max_files=int(self.max_files) if self.max_files else None,
            )
            files = proc.scan_folder(target, options)
            total = len(files)
            items = []
            base_prog = 0.3
            self.status.emit(f"Scanning {total} documents…")
            self.progress.emit(base_prog)
            for i, fp in enumerate(files, start=1):
                if self.cancelled:
                    break
                if (i % 5) == 0 or i == total:
                    self.status.emit(f"Analyzing {i}/{total}…")
                    self.progress.emit(base_prog + (1.0 - base_prog) * float(i) / float(max(1, total)))
                try:
                    res = sem.analyze_document(fp)
                    if not res:
                        continue
                    items.append(
                        {
                            "file": str(fp),
                            "document_type": res.document_type,
                            "domain": res.domain,
                            "confidence": res.confidence,
                            "explanation": res.explanation,
                            "suggested_filename": res.suggested_filename,
                            "suggested_category": res.suggested_category,
                            "sensitivity": res.sensitivity,
                            "model_used": res.model_used,
                            "analyzed_at": res.analyzed_at,
                        }
                    )
                except Exception:
                    continue

            report = {
                "version": 1,
                "generated_at": datetime.now().isoformat(),
                "target_folder": str(target),
                "options": {"include_subfolders": bool(self.include_subfolders), "max_files": int(self.max_files)},
                "items": items,
            }
            self.progress.emit(1.0)
            self.finished.emit(report, items)
        except Exception as e:
            self.error.emit(str(e))


class _QtContentAnalysisDialog(QDialog):
    def __init__(
        self,
        parent: QWidget,
        *,
        backend,
        target_folder: str,
        autorun: bool = True,
        initial_options: dict | None = None,
    ):
        super().__init__(parent)
        self.setWindowTitle("Semantic Document Analysis")
        self.setModal(True)
        self.setMinimumWidth(980)
        self.setMinimumHeight(640)
        self.setStyleSheet(_qt_modern_dialog_stylesheet())

        self.backend = backend
        self._target_folder = str(target_folder)
        self._autorun = bool(autorun)
        self._initial_options = dict(initial_options or {})
        self._items: list[dict] = []
        self._report: dict = {}

        layout = QVBoxLayout(self)
        layout.setContentsMargins(14, 14, 14, 14)
        layout.setSpacing(10)

        header = QFrame()
        header.setObjectName("DialogHeader")
        hl = QHBoxLayout(header)
        hl.setContentsMargins(12, 10, 12, 10)
        hl.setSpacing(10)

        tcol = QVBoxLayout()
        tcol.setContentsMargins(0, 0, 0, 0)
        tcol.setSpacing(2)
        title = QLabel("Semantic Document Analysis")
        title.setObjectName("DialogTitle")
        tcol.addWidget(title)
        self.subtitle = QLabel("Ready.")
        self.subtitle.setObjectName("DialogSubtitle")
        tcol.addWidget(self.subtitle)
        hl.addLayout(tcol, 1)

        self.btn_scan = QPushButton("Analyze")
        self.btn_scan.setObjectName("PrimaryButton")
        self.btn_scan.clicked.connect(self._scan)
        hl.addWidget(self.btn_scan)

        layout.addWidget(header)

        opts = QGroupBox()
        opts.setObjectName("Card")
        ol = QHBoxLayout(opts)
        ol.setContentsMargins(12, 12, 12, 12)
        ol.setSpacing(12)
        self.cb_subfolders = QCheckBox("Include subfolders (recursive)")
        self.cb_subfolders.setChecked(True)
        ol.addWidget(self.cb_subfolders)
        ol.addWidget(QLabel("Max files:"))
        self.max_files = QSpinBox()
        self.max_files.setRange(10, 5000)
        self.max_files.setValue(500)
        ol.addWidget(self.max_files)
        ol.addStretch(1)
        layout.addWidget(opts)

        self.bar = QProgressBar()
        self.bar.setRange(0, 1000)
        self.bar.setValue(0)
        layout.addWidget(self.bar)

        self.summary = QLabel("No analysis yet.")
        self.summary.setObjectName("DialogSummary")
        layout.addWidget(self.summary)

        self.tree = QTreeWidget()
        self.tree.setObjectName("ResultsTree")
        self.tree.setHeaderLabels(["", "File", "Type", "Suggested rename", "Suggested category", "Conf"])
        self.tree.setColumnWidth(0, 36)
        layout.addWidget(self.tree, 1)

        bottom = QHBoxLayout()
        self.btn_select_all = QPushButton("Select All")
        self.btn_select_all.setObjectName("SecondaryButton")
        self.btn_select_all.clicked.connect(lambda: self._set_all_checked(True))
        bottom.addWidget(self.btn_select_all)
        self.btn_select_none = QPushButton("Select None")
        self.btn_select_none.setObjectName("SecondaryButton")
        self.btn_select_none.clicked.connect(lambda: self._set_all_checked(False))
        bottom.addWidget(self.btn_select_none)

        self.btn_undo = QPushButton("Undo")
        self.btn_undo.setObjectName("SecondaryButton")
        self.btn_undo.clicked.connect(self._undo)
        bottom.addWidget(self.btn_undo)

        bottom.addStretch(1)
        self.btn_apply_rename = QPushButton("Apply Rename")
        self.btn_apply_rename.setObjectName("PrimaryButton")
        self.btn_apply_rename.clicked.connect(self._apply_rename)
        bottom.addWidget(self.btn_apply_rename)
        self.btn_apply_category = QPushButton("Apply Category")
        self.btn_apply_category.setObjectName("PrimaryButton")
        self.btn_apply_category.clicked.connect(self._apply_category)
        bottom.addWidget(self.btn_apply_category)
        self.btn_close = QPushButton("Close")
        self.btn_close.setObjectName("SecondaryButton")
        self.btn_close.clicked.connect(self.reject)
        bottom.addWidget(self.btn_close)
        layout.addLayout(bottom)

        self._refresh_undo_label()
        try:
            if "include_subfolders" in self._initial_options:
                self.cb_subfolders.setChecked(bool(self._initial_options["include_subfolders"]))
            if "max_files" in self._initial_options:
                self.max_files.setValue(int(self._initial_options["max_files"]))
        except Exception:
            pass

        if self._autorun:
            from PySide6.QtCore import QTimer

            QTimer.singleShot(50, self._scan)

    def _set_all_checked(self, checked: bool):
        for i in range(self.tree.topLevelItemCount()):
            it = self.tree.topLevelItem(i)
            if it:
                it.setCheckState(0, Qt.Checked if checked else Qt.Unchecked)

    def _selected_items(self) -> list[dict]:
        selected = []
        for i in range(min(self.tree.topLevelItemCount(), len(self._items))):
            it = self.tree.topLevelItem(i)
            if it and it.checkState(0) == Qt.Checked:
                selected.append(self._items[i])
        return selected

    def _scan(self):
        top = self.parent().window() if self.parent() else None
        ensure = getattr(top, "_ensure_ai_ready", None)
        if callable(ensure) and not ensure(title="Prepare Content Analysis", kind="vision"):
            return

        self.btn_scan.setEnabled(False)
        self.btn_apply_rename.setEnabled(False)
        self.btn_apply_category.setEnabled(False)
        self.tree.clear()
        self.bar.setValue(0)
        self.subtitle.setText("Analyzing…")

        self.worker = _QtContentAnalysisWorker(
            backend=self.backend,
            target_folder=self._target_folder,
            include_subfolders=bool(self.cb_subfolders.isChecked()),
            max_files=int(self.max_files.value()),
        )
        self._thread = QThread(self)
        self.worker.moveToThread(self._thread)
        self._thread.started.connect(self.worker.run)
        self.worker.status.connect(self.subtitle.setText)
        self.worker.progress.connect(lambda p: self.bar.setValue(int(p * 1000)))
        self.worker.finished.connect(self._on_finished)
        self.worker.error.connect(self._on_error)
        self.worker.finished.connect(self._thread.quit)
        self.worker.error.connect(self._thread.quit)
        self._thread.finished.connect(self._thread.deleteLater)
        self._thread.start()

    def _on_finished(self, report: dict, items: list):
        self._report = report
        self._items = list(items or [])
        self.btn_scan.setEnabled(True)
        self.btn_apply_rename.setEnabled(True)
        self.btn_apply_category.setEnabled(True)
        self.bar.setValue(1000)
        self.subtitle.setText("Analysis complete.")
        self.summary.setText(f"Analyzed {len(self._items)} documents")

        from pathlib import Path

        for it in self._items:
            fp = Path(str(it.get("file") or ""))
            dt = str(it.get("document_type") or "")
            rn = str(it.get("suggested_filename") or "")
            cat = str(it.get("suggested_category") or "")
            conf = it.get("confidence")
            conf_s = f"{float(conf):.2f}" if isinstance(conf, (int, float)) else ""
            row = QTreeWidgetItem(["", fp.name, dt, rn, cat, conf_s])
            row.setFlags(row.flags() | Qt.ItemIsUserCheckable)
            row.setCheckState(0, Qt.Checked)
            self.tree.addTopLevelItem(row)
        self.tree.resizeColumnToContents(0)

    def _on_error(self, msg: str):
        QMessageBox.critical(self, "Content Analysis Failed", msg)
        self.btn_scan.setEnabled(True)

    def _refresh_undo_label(self):
        try:
            from utils.universal_undo import get_undo_manager

            stats = get_undo_manager().get_statistics()
            n = int(stats.get("undoable_transactions") or 0)
            self.btn_undo.setText(f"Undo ({n})" if n else "Undo")
            self.btn_undo.setEnabled(n > 0)
        except Exception:
            self.btn_undo.setEnabled(True)

    def _undo(self):
        try:
            from utils.universal_undo import undo_last_operation

            ok, msg, _n = undo_last_operation()
            QMessageBox.information(self, "Undo", msg)
            self._refresh_undo_label()
        except Exception as e:
            QMessageBox.critical(self, "Undo", str(e))

    def _apply_rename(self):
        from pathlib import Path

        selected = self._selected_items()
        if not selected:
            QMessageBox.information(self, "Apply Rename", "No items selected.")
            return
        from utils.intelligent_rename import sanitize_ai_filename, get_unique_filename
        from utils.universal_undo import record_bulk_rename

        pairs = []
        renamed = 0
        skipped = 0
        for it in selected:
            try:
                src = Path(str(it.get("file") or ""))
                sug = str(it.get("suggested_filename") or "").strip()
                if (not sug) or (not src.exists()):
                    skipped += 1
                    continue
                validation = sanitize_ai_filename(sug, preserve_case=True)
                desired = validation.sanitized_name
                if not desired:
                    skipped += 1
                    continue
                unique, _why = get_unique_filename(src, desired, {"used_ai": True})
                dst = src.with_name(unique + src.suffix)
                if str(dst) == str(src):
                    skipped += 1
                    continue
                src.rename(dst)
                pairs.append((src, dst))
                renamed += 1
            except Exception:
                skipped += 1
                continue
        if pairs:
            try:
                record_bulk_rename([(a, b) for (a, b) in pairs], metadata={"qt_dialog": "content_analysis"})
            except Exception:
                pass
        self._refresh_undo_label()
        QMessageBox.information(self, "Apply Rename", f"Renamed: {renamed}\nSkipped: {skipped}")

    def _apply_category(self):
        from pathlib import Path
        import shutil
        import re

        selected = self._selected_items()
        if not selected:
            QMessageBox.information(self, "Apply Category", "No items selected.")
            return
        base = QFileDialog.getExistingDirectory(self, "Select Base Folder for Categories", self._target_folder)
        if not base:
            return
        base_p = Path(base)
        from utils.universal_undo import record_categorize

        moved_pairs = []
        moved = 0
        skipped = 0
        reserved_names = {
            "con", "prn", "aux", "nul",
            "com1", "com2", "com3", "com4", "com5", "com6", "com7", "com8", "com9",
            "lpt1", "lpt2", "lpt3", "lpt4", "lpt5", "lpt6", "lpt7", "lpt8", "lpt9",
        }
        for it in selected:
            try:
                src = Path(str(it.get("file") or ""))
                cat = str(it.get("suggested_category") or "").strip()
                if (not cat) or (not src.exists()):
                    skipped += 1
                    continue
                cat = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", cat)
                cat = re.sub(r"\s+", " ", cat).strip(" .")
                if not cat:
                    skipped += 1
                    continue
                if cat.split(".", 1)[0].lower() in reserved_names:
                    cat = f"_{cat}"
                cat = cat[:80]
                dest_dir = base_p / cat
                dest_dir.mkdir(parents=True, exist_ok=True)
                out = dest_dir / src.name
                if out.exists():
                    stem = out.stem
                    suf = out.suffix
                    n = 1
                    while out.exists():
                        out = dest_dir / f"{stem}_{n}{suf}"
                        n += 1
                shutil.move(str(src), str(out))
                moved_pairs.append((src, out))
                moved += 1
            except Exception:
                skipped += 1
                continue
        if moved_pairs:
            try:
                record_categorize(moved_pairs, metadata={"qt_dialog": "content_analysis"})
            except Exception:
                pass
        self._refresh_undo_label()
        QMessageBox.information(self, "Apply Category", f"Moved: {moved}\nSkipped: {skipped}")


class _QtAIModelDownloadDialog(QDialog):
    """Download model files without emitting Qt signals from the download thread."""

    def __init__(self, parent: QWidget, *, ai_manager, kind: str | None = None):
        super().__init__(parent)
        self.ai_manager = ai_manager
        k = (kind or "").strip().lower()
        self.kind = k if k in ("vision", "text") else None
        self._events = queue.Queue()
        self._done = False

        self.setWindowTitle("AI Model Download")
        self.setModal(True)
        self.setMinimumWidth(560)
        self.setStyleSheet(_qt_modern_dialog_stylesheet())

        layout = QVBoxLayout(self)
        layout.setContentsMargins(14, 14, 14, 14)
        layout.setSpacing(10)

        kind_label = ""
        if self.kind in ("vision", "text"):
            kind_label = " (Vision)" if self.kind == "vision" else " (Text)"
        title = QLabel(f"Downloading the AI model{kind_label}...")
        title.setObjectName("DialogTitle")
        layout.addWidget(title)

        self.status = QLabel("Starting download...")
        self.status.setObjectName("DialogSubtitle")
        self.status.setWordWrap(True)
        layout.addWidget(self.status)

        self.bar = QProgressBar()
        self.bar.setRange(0, 1000)
        self.bar.setValue(0)
        layout.addWidget(self.bar)

        self.detail = QLabel("Model files are stored in the local Fylorra AI models folder.")
        self.detail.setObjectName("DialogSubtitle")
        self.detail.setWordWrap(True)
        layout.addWidget(self.detail)

        row = QHBoxLayout()
        row.addStretch(1)
        self.btn_close = QPushButton("Close")
        self.btn_close.clicked.connect(self.reject)
        self.btn_close.setEnabled(False)
        self.btn_close.setVisible(False)
        row.addWidget(self.btn_close)
        layout.addLayout(row)

        self._timer = QTimer(self)
        self._timer.setInterval(100)
        self._timer.timeout.connect(self._drain_events)
        self._timer.start()

        self._thread = threading.Thread(target=self._run_download, name="FylorraAIModelDownload", daemon=True)
        self._thread.start()

    def _run_download(self):
        try:
            ai = self.ai_manager
            if not ai:
                self._events.put(("error", "AI Manager is not available."))
                return
            if self.kind:
                try:
                    ai.select_kind(self.kind)
                except Exception:
                    pass

            def dl_cb(message: str, progress: float, downloaded: str = "", speed: str = ""):
                try:
                    pct = max(0.0, min(1.0, float(progress)))
                except Exception:
                    pct = 0.0
                self._events.put(("progress", str(message or "Downloading..."), pct, str(downloaded or ""), str(speed or "")))

            ok = bool(ai.ensure_model_downloaded(dl_cb))
            if ok:
                self._events.put(("finished", True, "Model files downloaded."))
            else:
                err = str(getattr(ai, "load_error", "") or "").strip()
                self._events.put(("finished", False, err or "Model download failed."))
        except Exception as e:
            self._events.put(("error", f"Model download failed: {e}"))

    def _drain_events(self):
        while True:
            try:
                event = self._events.get_nowait()
            except queue.Empty:
                break
            kind = event[0]
            if kind == "progress":
                _kind, message, pct, downloaded, speed = event
                self.status.setText(message)
                self.bar.setValue(int(pct * 1000))
                parts = []
                if downloaded:
                    parts.append(downloaded)
                if speed:
                    parts.append(speed)
                self.detail.setText(" / ".join(parts) if parts else "Downloading model files...")
            elif kind == "finished":
                _kind, ok, message = event
                self._done = True
                self._timer.stop()
                if ok:
                    self.bar.setValue(1000)
                    self.status.setText("Downloaded.")
                    self.detail.setText("The model files are ready. Load the model only when you need AI features now.")
                    QTimer.singleShot(250, self.accept)
                else:
                    self._show_failure(message)
            elif kind == "error":
                self._done = True
                self._timer.stop()
                self._show_failure(event[1])

    def _show_failure(self, message: str):
        self.status.setText("Failed.")
        self.detail.setText((message or "Failed to download the AI model.").strip())
        self.btn_close.setVisible(True)
        self.btn_close.setEnabled(True)
        self.btn_close.setText("Close")
        try:
            self.btn_close.clicked.disconnect()
        except Exception:
            pass
        self.btn_close.clicked.connect(self.reject)

class _QtAIModelLoadDialog(QDialog):
    """Prepare/load model files without touching Qt widgets from worker threads."""

    def __init__(self, parent: QWidget, *, ai_manager, kind: str | None = None, action: str = "prepare"):
        super().__init__(parent)
        self.ai_manager = ai_manager
        k = (kind or "").strip().lower()
        self.kind = k if k in ("vision", "text") else None
        self.action = (action or "prepare").strip().lower()
        if self.action not in ("download", "load", "prepare", "load_unload"):
            self.action = "prepare"
        self._events = queue.Queue()
        self._done = False

        self.setWindowTitle("AI Model")
        self.setModal(True)
        self.setMinimumWidth(560)
        self.setStyleSheet(_qt_modern_dialog_stylesheet())

        layout = QVBoxLayout(self)
        layout.setContentsMargins(14, 14, 14, 14)
        layout.setSpacing(10)

        kind_label = ""
        if self.kind in ("vision", "text"):
            kind_label = " (Vision)" if self.kind == "vision" else " (Text)"
        verb = "Downloading" if self.action == "download" else "Loading"
        title = QLabel(f"{verb} the AI model{kind_label}...")
        title.setObjectName("DialogTitle")
        layout.addWidget(title)

        self.status = QLabel("Starting...")
        self.status.setObjectName("DialogSubtitle")
        self.status.setWordWrap(True)
        layout.addWidget(self.status)

        self.bar = QProgressBar()
        self.bar.setRange(0, 1000)
        self.bar.setValue(0)
        layout.addWidget(self.bar)

        self.detail = QLabel("Preparing local AI runtime...")
        self.detail.setObjectName("DialogSubtitle")
        self.detail.setWordWrap(True)
        layout.addWidget(self.detail)

        row = QHBoxLayout()
        row.addStretch(1)
        self.btn_close = QPushButton("Close")
        self.btn_close.clicked.connect(self.reject)
        self.btn_close.setEnabled(False)
        self.btn_close.setVisible(False)
        row.addWidget(self.btn_close)
        layout.addLayout(row)

        self._timer = QTimer(self)
        self._timer.setInterval(100)
        self._timer.timeout.connect(self._drain_events)
        self._timer.start()

        self._thread = threading.Thread(target=self._run_prepare, name="FylorraAIModelLoad", daemon=True)
        self._thread.start()

    def _run_prepare(self):
        ai = self.ai_manager
        try:
            if not ai:
                self._events.put(("error", "AI Manager is not available."))
                return

            if self.kind:
                try:
                    ai.select_kind(self.kind)
                except Exception:
                    pass

            if getattr(ai, "is_ready", False):
                if not self.kind:
                    self._events.put(("finished", True, "AI model is already loaded."))
                    return
                try:
                    current_kind = ai.get_active_kind()
                except Exception:
                    current_kind = "vision" if getattr(ai, "is_vision_model", False) else "text"
                if current_kind == self.kind:
                    if self.action == "load_unload":
                        self._queue_progress("Unloading AI model...", 1.0, "", "")
                        try:
                            ai.unload_model()
                        except Exception as e:
                            self._events.put(("error", f"Model unload failed: {e}"))
                            return
                        if bool(getattr(ai, "is_ready", False)):
                            self._events.put(("error", "Model unload did not complete."))
                            return
                        self._events.put(("finished", True, "AI model loaded and unloaded cleanly."))
                        return
                    self._events.put(("finished", True, "AI model is already loaded."))
                    return
                try:
                    ai.unload_model()
                except Exception:
                    pass

            def dl_cb(message: str, progress: float, downloaded: str = "", speed: str = ""):
                self._queue_progress(message, 0.0 + 0.4 * self._safe_progress(progress), downloaded, speed)

            try:
                ok = bool(ai.ensure_model_downloaded(dl_cb))
            except Exception as e:
                self._events.put(("error", f"Model download failed: {e}"))
                return
            if not ok:
                err = str(getattr(ai, "load_error", "") or "").strip()
                self._events.put(("error", err or "Model download failed."))
                return

            if self.action == "download":
                self._events.put(("finished", True, "Model files downloaded."))
                return

            def load_cb(message: str, progress: float, downloaded: str = "", speed: str = ""):
                self._queue_progress(message, 0.4 + 0.6 * self._safe_progress(progress), downloaded, speed)

            try:
                ai.load_model(load_cb)
            except Exception as e:
                self._events.put(("error", f"Model load failed: {e}"))
                return

            if bool(getattr(ai, "is_ready", False)):
                if self.action == "load_unload":
                    self._queue_progress("Unloading AI model...", 1.0, "", "")
                    try:
                        ai.unload_model()
                    except Exception as e:
                        self._events.put(("error", f"Model unload failed: {e}"))
                        return
                    if bool(getattr(ai, "is_ready", False)):
                        self._events.put(("error", "Model unload did not complete."))
                        return
                    self._events.put(("finished", True, "AI model loaded and unloaded cleanly."))
                    return
                self._events.put(("finished", True, "AI model is ready."))
                return
            err = str(getattr(ai, "load_error", "") or "").strip()
            self._events.put(("error", err or "Failed to load model."))
        except Exception as e:
            self._events.put(("error", str(e)))

    @staticmethod
    def _safe_progress(value: float) -> float:
        try:
            return max(0.0, min(1.0, float(value)))
        except Exception:
            return 0.0

    def _queue_progress(self, message: str, progress: float, downloaded: str = "", speed: str = ""):
        self._events.put(("progress", str(message or "Working..."), self._safe_progress(progress), str(downloaded or ""), str(speed or "")))

    def _drain_events(self):
        while True:
            try:
                event = self._events.get_nowait()
            except queue.Empty:
                break
            kind = event[0]
            if kind == "progress":
                _kind, message, pct, downloaded, speed = event
                self.status.setText(message)
                self.bar.setValue(int(pct * 1000))
                parts = []
                if downloaded:
                    parts.append(downloaded)
                if speed:
                    parts.append(speed)
                if parts:
                    self.detail.setText(" / ".join(parts))
            elif kind == "finished":
                _kind, ok, message = event
                self._done = True
                self._timer.stop()
                self.bar.setValue(1000 if ok else self.bar.value())
                done_label = "Downloaded." if self.action == "download" else ("Load/unload OK." if self.action == "load_unload" else "Ready.")
                self.status.setText(done_label if ok else "Failed.")
                self.detail.setText(str(message or ""))
                if ok:
                    QTimer.singleShot(250, self.accept)
                else:
                    self._show_failure(message)
            elif kind == "error":
                self._done = True
                self._timer.stop()
                self._show_failure(event[1])

    def _show_failure(self, msg: str):
        clean = (msg or "Failed to prepare the AI model.").strip()
        self.status.setText("Failed.")
        self.detail.setText(clean)
        self.btn_close.setVisible(True)
        self.btn_close.setEnabled(True)
        self.btn_close.setText("Close")
        try:
            self.btn_close.clicked.disconnect()
        except Exception:
            pass
        self.btn_close.clicked.connect(self.reject)

class _QtAICommandPlanWorker(QObject):
    status = Signal(str)
    finished = Signal(object, str)  # plan, rendered_text
    error = Signal(str)

    def __init__(self, *, backend, target_folder: str, instruction: str):
        super().__init__()
        self.backend = backend
        self.target_folder = target_folder
        self.instruction = instruction

    def run(self):
        try:
            from pathlib import Path
            import json
            from core.ai_command import plan_from_nl, _heuristic_plan

            ai = getattr(self.backend, "ai_manager", None)
            # fallback heuristic if AI not available/ready
            plan = None
            if not ai or not getattr(ai, "is_ready", False):
                plan = _heuristic_plan(self.instruction)
                if plan is None:
                    raise RuntimeError("AI model is not ready. Load the model first.")
            else:
                self.status.emit("Planning…")
                plan = plan_from_nl(ai, self.instruction, target_folder=Path(self.target_folder))

            # render
            d = {"intent_summary": plan.intent_summary, "steps": []}
            for i, s in enumerate(plan.steps, start=1):
                d["steps"].append({"step": i, "tool": s.tool, "description": s.description, "destructive": bool(s.destructive), "args": s.args})
            rendered = json.dumps(d, indent=2)
            self.finished.emit(plan, rendered)
        except Exception as e:
            self.error.emit(str(e))


class _QtAICommandPlanDialog(QDialog):
    def __init__(self, parent: QWidget, *, target_folder: str, instruction: str):
        super().__init__(parent)
        self.setWindowTitle("Generate Plan")
        self.setModal(True)
        self.setMinimumWidth(740)
        self.setMinimumHeight(520)

        top = parent.window()
        self.backend = getattr(top, "backend", None)
        self.plan = None
        self.plan_text = ""

        layout = QVBoxLayout(self)
        layout.setContentsMargins(14, 14, 14, 14)
        layout.setSpacing(10)

        self.status = QLabel("Preparing…")
        self.status.setStyleSheet("color:#9aa0a9;")
        layout.addWidget(self.status)

        self.view = QTextEdit()
        self.view.setReadOnly(True)
        layout.addWidget(self.view, 1)

        row = QHBoxLayout()
        row.addStretch(1)
        self.btn_close = QPushButton("Cancel")
        self.btn_close.clicked.connect(self.reject)
        row.addWidget(self.btn_close)
        layout.addLayout(row)

        ai = getattr(self.backend, "ai_manager", None) if self.backend else None
        # If model not ready, load it first.
        if ai and not getattr(ai, "is_ready", False):
            loader = _QtAIModelLoadDialog(self, ai_manager=ai)
            if loader.exec() != QDialog.Accepted:
                self._cancelled = True
                self.status.setText("Cancelled.")
                self.btn_close.setText("Close")
                self.btn_close.clicked.disconnect()
                self.btn_close.clicked.connect(self.reject)
                return

        self.worker = _QtAICommandPlanWorker(backend=self.backend, target_folder=target_folder, instruction=instruction)
        self._thread = QThread(self)
        self.worker.moveToThread(self._thread)
        self._thread.started.connect(self.worker.run)
        self.worker.status.connect(self.status.setText)
        self.worker.finished.connect(self._on_finished)
        self.worker.error.connect(self._on_error)
        self.worker.finished.connect(self._thread.quit)
        self.worker.error.connect(self._thread.quit)
        self._thread.finished.connect(self._thread.deleteLater)
        self._thread.start()

    def _on_finished(self, plan, rendered: str):
        self.plan = plan
        self.plan_text = rendered
        self.status.setText("Plan ready.")
        self.view.setPlainText(rendered)
        self.btn_close.setText("Use Plan")
        try:
            self.btn_close.clicked.disconnect()
        except Exception:
            pass
        self.btn_close.clicked.connect(self.accept)

    def _on_error(self, msg: str):
        QMessageBox.critical(self, "Plan Failed", msg)
        self.reject()


class _QtAICommandRunWorker(QObject):
    status = Signal(str)
    progress = Signal(float)
    log = Signal(str)
    finished = Signal(object)  # report dict
    error = Signal(str)

    def __init__(self, *, backend, target_folder: str, plan):
        super().__init__()
        self.backend = backend
        self.target_folder = target_folder
        self.plan = plan

    def run(self):
        try:
            from pathlib import Path
            from core.ai_command import run_plan

            ai = getattr(self.backend, "ai_manager", None)
            def prog(msg: str, frac: float):
                self.status.emit(msg)
                self.progress.emit(float(frac))

            report = run_plan(self.plan, target_folder=Path(self.target_folder), ai_manager=ai, progress=prog)
            self.finished.emit(report)
        except Exception as e:
            self.error.emit(str(e))


class _QtAICommandRunDialog(QDialog):
    def __init__(self, parent: QWidget, *, target_folder: str, plan):
        super().__init__(parent)
        self.setWindowTitle("Run Plan")
        self.setModal(True)
        self.setMinimumWidth(760)
        self.setMinimumHeight(520)

        top = parent.window()
        self.backend = getattr(top, "backend", None)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(14, 14, 14, 14)
        layout.setSpacing(10)

        self.status = QLabel("Starting…")
        self.status.setStyleSheet("color:#9aa0a9;")
        layout.addWidget(self.status)

        self.bar = QProgressBar()
        self.bar.setRange(0, 1000)
        self.bar.setValue(0)
        layout.addWidget(self.bar)

        self.log_list = QListWidget()
        self.log_list.setStyleSheet(
            "QListWidget{background:#14171c; border:1px solid #232730; border-radius:10px; color:#c8ccd6;}"
        )
        layout.addWidget(self.log_list, 1)

        row = QHBoxLayout()
        row.addStretch(1)
        self.btn_close = QPushButton("Close")
        self.btn_close.setEnabled(False)
        self.btn_close.clicked.connect(self.accept)
        row.addWidget(self.btn_close)
        layout.addLayout(row)

        self.worker = _QtAICommandRunWorker(backend=self.backend, target_folder=target_folder, plan=plan)
        self._thread = QThread(self)
        self.worker.moveToThread(self._thread)
        self._thread.started.connect(self.worker.run)
        self.worker.status.connect(self._on_status)
        self.worker.progress.connect(lambda p: self.bar.setValue(int(max(0.0, min(1.0, p)) * 1000)))
        self.worker.finished.connect(self._on_finished)
        self.worker.error.connect(self._on_error)
        self.worker.finished.connect(self._thread.quit)
        self.worker.error.connect(self._thread.quit)
        self._thread.finished.connect(self._thread.deleteLater)
        self._thread.start()

    def _on_status(self, msg: str):
        self.status.setText(msg)
        self.log_list.insertItem(0, QListWidgetItem(msg))
        while self.log_list.count() > 300:
            self.log_list.takeItem(self.log_list.count() - 1)

    def _on_finished(self, report):
        self.bar.setValue(1000)
        self.status.setText("Complete.")
        self.log_list.insertItem(0, QListWidgetItem("Done."))
        self.btn_close.setEnabled(True)

    def _on_error(self, msg: str):
        QMessageBox.critical(self, "Run Failed", msg)
        self.btn_close.setEnabled(True)


class _QtLibraryIndexWorker(QObject):
    status = Signal(str)
    progress = Signal(float)
    finished = Signal(int)
    error = Signal(str)

    def __init__(
        self,
        *,
        backend,
        folder: str,
        include_subfolders: bool,
        ai_summarize: bool,
        ocr_scanned_pdfs: bool,
        extract_images: bool,
        compute_hashes: bool,
    ):
        super().__init__()
        self.backend = backend
        self.folder = folder
        self.include_subfolders = include_subfolders
        self.ai_summarize = ai_summarize
        self.ocr_scanned_pdfs = ocr_scanned_pdfs
        self.extract_images = extract_images
        self.compute_hashes = compute_hashes
        self._cancelled = False

    def cancel(self):
        self._cancelled = True

    def run(self):
        try:
            from pathlib import Path
            from core.library_index import LibraryIndex

            folder = Path(self.folder)
            lib = LibraryIndex()
            ai = getattr(self.backend, "ai_manager", None)

            def cb(msg: str, p: float):
                if self._cancelled:
                    raise RuntimeError("Cancelled")
                self.status.emit(msg)
                try:
                    self.progress.emit(float(p))
                except Exception:
                    pass

            self.status.emit("Preparing index…")
            n = lib.index_folder(
                folder,
                include_subfolders=bool(self.include_subfolders),
                ai_manager=ai,
                ai_summarize=bool(self.ai_summarize),
                extract_images=bool(self.extract_images),
                compute_hashes=bool(self.compute_hashes),
                ocr_scanned_pdfs=bool(self.ocr_scanned_pdfs),
                progress_cb=cb,
            )
            self.finished.emit(int(n))
        except Exception as e:
            msg = str(e)
            if "Cancelled" in msg:
                self.error.emit("Cancelled.")
            else:
                self.error.emit(msg)


class _QtWritingAssistantWorker(QObject):
    finished = Signal(str)
    error = Signal(str)

    def __init__(self, *, ai_manager, messages: list[dict], temperature: float, max_tokens: int):
        super().__init__()
        self.ai_manager = ai_manager
        self.messages = list(messages or [])
        self.temperature = float(temperature)
        self.max_tokens = int(max_tokens)

    def run(self):
        try:
            ai = self.ai_manager
            if not ai or not getattr(ai, "is_ready", False) or not getattr(ai, "model", None):
                raise RuntimeError("AI model not loaded.")
            response = ai.model.create_chat_completion(
                messages=self.messages,
                temperature=self.temperature,
                max_tokens=self.max_tokens,
            )
            text = (response["choices"][0]["message"]["content"] or "").strip()
            self.finished.emit(text)
        except Exception as e:
            self.error.emit(str(e))


class _QtLibrarySearchWorker(QObject):
    status = Signal(str)
    finished = Signal(object)  # list
    error = Signal(str)

    def __init__(self, *, backend, query: str, rerank: bool, folder: str = ""):
        super().__init__()
        self.backend = backend
        self.query = query
        self.rerank = rerank
        self.folder = folder

    def run(self):
        try:
            from pathlib import Path
            from core.library_index import LibraryIndex
            from core.ai_search import ai_search

            lib = LibraryIndex()
            ai = getattr(self.backend, "ai_manager", None)
            folder = None
            try:
                if str(self.folder or "").strip():
                    folder = Path(str(self.folder).strip())
            except Exception:
                folder = None
            if folder is not None and not folder.exists():
                raise RuntimeError(f"Folder not found: {folder}")
            if folder is not None:
                self.status.emit(f"Preparing search… (scope: {folder})")
            else:
                self.status.emit("Preparing search…")

            def cb(msg: str):
                try:
                    self.status.emit(str(msg))
                except Exception:
                    pass

            results = ai_search(
                lib,
                self.query,
                ai_manager=ai,
                rerank=bool(self.rerank),
                folder=folder,
                progress_cb=cb,
            )
            self.finished.emit(results)
        except Exception as e:
            self.error.emit(str(e))


class _DropFolderLineEdit(QLineEdit):
    """
    QLineEdit that accepts drag-and-drop of folders (or files).
    - Dropping a folder sets the text to that folder.
    - Dropping a file sets the text to its parent folder.
    """

    pathDropped = Signal(str)

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.setAcceptDrops(True)

    def dragEnterEvent(self, e):  # noqa: N802 (Qt naming)
        try:
            md = e.mimeData()
            if md and md.hasUrls():
                e.acceptProposedAction()
                return
        except Exception:
            pass
        try:
            e.ignore()
        except Exception:
            pass

    def dropEvent(self, e):  # noqa: N802 (Qt naming)
        try:
            md = e.mimeData()
            if not md or not md.hasUrls():
                return
            urls = md.urls() or []
            if not urls:
                return
            local = urls[0].toLocalFile()
            if not local:
                return
            from pathlib import Path

            p = Path(local)
            if p.exists() and p.is_file():
                p = p.parent
            if p.exists() and p.is_dir():
                self.setText(str(p))
                try:
                    self.pathDropped.emit(str(p))
                except Exception:
                    pass
                e.acceptProposedAction()
                return
        except Exception:
            pass
        try:
            e.ignore()
        except Exception:
            pass


class _QtAISearchResultCard(QFrame):
    """
    Card-style list item for AI Search results.
    Designed to look premium vs plain text rows.
    """

    _icon_provider = QFileIconProvider()
    _icon_cache: dict[str, QIcon] = {}
    _thumb_cache: dict[str, QPixmap] = {}

    def __init__(self, res):
        super().__init__()
        self._selected = False
        self.res = res
        self.setObjectName("AISearchCard")
        self.setMinimumHeight(66)
        self.setStyleSheet(self._style(False))

        outer = QHBoxLayout(self)
        outer.setContentsMargins(10, 8, 10, 8)
        outer.setSpacing(10)

        path = ""
        name = ""
        ext = ""
        matched_query = ""
        try:
            path = str(res.item.path or "")
            name = str(res.item.name or Path(path).name)
            ext = str(res.item.ext or Path(path).suffix)
            matched_query = str(getattr(res, "matched_query", "") or "")
        except Exception:
            pass

        tokens = self._query_tokens(matched_query)

        # Thumbnail / icon
        thumb_label = QLabel()
        thumb_label.setFixedSize(78, 50)
        thumb_label.setAlignment(Qt.AlignCenter)
        thumb_label.setStyleSheet("background:#0f1217; border:1px solid #232730; border-radius:10px;")
        pm = self._thumbnail(path, ext)
        if pm is not None and not pm.isNull():
            thumb_label.setPixmap(pm)
        else:
            icon = self._file_icon(path, ext)
            thumb_label.setPixmap(icon.pixmap(26, 26))
        outer.addWidget(thumb_label)

        # Text block
        mid = QVBoxLayout()
        mid.setSpacing(2)
        title = QLabel(name)
        title.setStyleSheet("color:#e6e9f2; font-weight:700;")
        if tokens:
            title.setTextFormat(Qt.RichText)
            title.setText(self._highlight(name, tokens))
        title.setTextInteractionFlags(Qt.TextSelectableByMouse)
        mid.addWidget(title)

        folder = ""
        try:
            folder = str(Path(path).parent)
        except Exception:
            folder = ""
        sub = QLabel(folder)
        sub.setStyleSheet("color:#9aa0a9;")
        sub.setTextInteractionFlags(Qt.TextSelectableByMouse)
        mid.addWidget(sub)

        snippet = ""
        try:
            snippet = (res.item.ai_summary or res.item.extracted_text or "").strip()
        except Exception:
            snippet = ""
        snippet = " ".join(snippet.split())
        if len(snippet) > 140:
            snippet = snippet[:140].rstrip() + "…"
        if snippet:
            sn = QLabel(snippet)
            sn.setStyleSheet("color:#c0c5cf;")
            if tokens:
                sn.setTextFormat(Qt.RichText)
                sn.setText(self._highlight(snippet, tokens))
            sn.setWordWrap(True)
            mid.addWidget(sn)

        outer.addLayout(mid, 1)

        # Right badges
        right = QVBoxLayout()
        right.setSpacing(4)
        right.setAlignment(Qt.AlignTop | Qt.AlignRight)
        badge = QLabel((ext or "").lstrip(".").upper()[:8])
        badge.setAlignment(Qt.AlignCenter)
        badge.setStyleSheet(
            "padding:4px 8px; border-radius:10px; background:#1b2430; color:#9cc7ff; border:1px solid #2a3442;"
        )
        right.addWidget(badge)
        outer.addLayout(right)

    @classmethod
    def _file_icon(cls, path: str, ext: str) -> QIcon:
        key = (ext or "").lower().strip() or "__file__"
        if key in cls._icon_cache:
            return cls._icon_cache[key]
        try:
            ico = cls._icon_provider.icon(QFileInfo(path)) if path else QIcon()
            if ico and not ico.isNull():
                cls._icon_cache[key] = ico
                return ico
        except Exception:
            pass
        cls._icon_cache[key] = QApplication.style().standardIcon(QStyle.SP_FileIcon)
        return cls._icon_cache[key]

    @classmethod
    def _thumbnail(cls, path: str, ext: str) -> QPixmap | None:
        p = (path or "").strip()
        if not p:
            return None
        if p in cls._thumb_cache:
            return cls._thumb_cache[p]
        try:
            suffix = (ext or Path(p).suffix).lower().strip()
        except Exception:
            suffix = (ext or "").lower().strip()
        try:
            if suffix in {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}:
                pm = _load_oriented_pixmap(p, max_size=QSize(74, 46), smooth=True)
                if pm is not None and not pm.isNull():
                    cls._thumb_cache[p] = pm
                    return pm
        except Exception:
            pass
        try:
            if suffix == ".pdf":
                import fitz  # PyMuPDF

                doc = fitz.open(p)
                try:
                    page = doc.load_page(0)
                    pix = page.get_pixmap(matrix=fitz.Matrix(0.25, 0.25), alpha=True)
                    fmt = QImage.Format_RGBA8888 if pix.alpha else QImage.Format_RGB888
                    img = QImage(pix.samples, pix.width, pix.height, pix.stride, fmt).copy()
                    pm = QPixmap.fromImage(img)
                    if not pm.isNull():
                        pm = pm.scaled(QSize(74, 46), Qt.KeepAspectRatio, Qt.SmoothTransformation)
                        cls._thumb_cache[p] = pm
                        return pm
                finally:
                    doc.close()
        except Exception:
            pass
        cls._thumb_cache[p] = QPixmap()
        return None

    @staticmethod
    def _query_tokens(matched_query: str) -> list[str]:
        q = (matched_query or "").strip()
        if not q:
            return []
        # MATCH queries might include OR/quotes/wildcards; normalize.
        q = q.replace("\\", " ").replace("/", " ").replace("*", " ").replace('"', " ").replace("'", " ")
        parts = [p.strip() for p in q.replace(" OR ", " ").replace(" AND ", " ").split() if p.strip()]
        toks: list[str] = []
        for p in parts:
            p2 = "".join(ch for ch in p if ch.isalnum() or ch in {"_", "-", "."}).strip("._-")
            if len(p2) < 3:
                continue
            toks.append(p2)
        # Prefer longer tokens first; de-dupe case-insensitive.
        toks.sort(key=lambda s: (-len(s), s.lower()))
        out: list[str] = []
        seen: set[str] = set()
        for t in toks:
            k = t.lower()
            if k in seen:
                continue
            seen.add(k)
            out.append(t)
            if len(out) >= 8:
                break
        return out

    @staticmethod
    def _highlight(text: str, tokens: list[str]) -> str:
        import html
        import re

        raw = text or ""
        if not raw or not tokens:
            return html.escape(raw)
        safe = html.escape(raw)
        # Build a single regex alternation with longest tokens first.
        esc = [re.escape(t) for t in tokens if t]
        if not esc:
            return safe
        pat = re.compile("(" + "|".join(esc) + ")", flags=re.IGNORECASE)

        def repl(m):
            s = m.group(0)
            return f'<span style="background:#1e3a5f; color:#dbe9ff; padding:1px 4px; border-radius:6px;">{html.escape(s)}</span>'

        return pat.sub(repl, safe)

    @staticmethod
    def _style(selected: bool) -> str:
        if selected:
            return "QFrame{background:#1b2330; border:1px solid #2b3b52; border-radius:12px;}"
        return "QFrame{background:#161a21; border:1px solid #232730; border-radius:12px;}"

    def set_selected(self, selected: bool):
        selected = bool(selected)
        if self._selected == selected:
            return
        self._selected = selected
        self.setStyleSheet(self._style(selected))


class _QtRuleGenWorker(QObject):
    status = Signal(str)
    finished = Signal(object)  # RuleGenerationResult
    error = Signal(str)

    def __init__(self, *, backend, text: str, current_folder: str | None = None):
        super().__init__()
        self.backend = backend
        self.text = text
        self.current_folder = current_folder

    def run(self):
        try:
            from core.nl_rule_builder import NaturalLanguageRuleBuilder

            ai = getattr(self.backend, "ai_manager", None)
            if not ai or not getattr(ai, "is_ready", False):
                raise RuntimeError("AI model not loaded.")

            b = NaturalLanguageRuleBuilder(ai)
            ctx = {}
            if self.current_folder:
                ctx["current_folder"] = self.current_folder
            self.status.emit("Generating rule…")
            res = b.generate_rule(self.text, ctx)
            self.finished.emit(res)
        except Exception as e:
            self.error.emit(str(e))


class _QtAiRuleBuilderPanel(QFrame):
    rule_added = Signal()

    EXAMPLES = [
        "Move all PDFs to Documents folder",
        "Copy new images to Backup drive",
        "Organize videos by date",
        "Move invoices with 'paid' in name to Archive folder",
        "Copy work documents to Projects folder",
        "Organize music files by file type",
            "Delete temporary files older than 7 days every night",
    ]

    def __init__(self, *, backend, icons: QtIconLoader, ensure_ai_ready_cb):
        super().__init__()
        self.backend = backend
        self.icons = icons
        self._ensure_ai_ready_cb = ensure_ai_ready_cb
        self._result = None
        self._thread: QThread | None = None
        self._worker: QObject | None = None

        self.setObjectName("PageCard")
        outer = QVBoxLayout(self)
        outer.setContentsMargins(14, 12, 14, 12)
        outer.setSpacing(10)

        header = QHBoxLayout()
        header.setContentsMargins(0, 0, 0, 0)
        header.setSpacing(10)

        title = QLabel("Rule Builder")
        title.setStyleSheet("font-size:14px; font-weight:700; color:#ffffff;")
        subtitle = QLabel("Describe what you want — generate a rule or scheduled task.")
        subtitle.setStyleSheet("color:#9aa0a9;")
        header.addWidget(title)
        header.addSpacing(10)
        header.addWidget(subtitle, 1)

        self.btn_popout = QToolButton()
        self.btn_popout.setText("Pop out")
        self.btn_popout.setToolButtonStyle(Qt.ToolButtonTextOnly)
        self.btn_popout.setToolTip("Open in a separate dialog")
        header.addWidget(self.btn_popout)

        self.btn_hide = QToolButton()
        self.btn_hide.setText("Hide")
        self.btn_hide.setToolButtonStyle(Qt.ToolButtonTextOnly)
        self.btn_hide.setToolTip("Hide builder")
        header.addWidget(self.btn_hide)

        outer.addLayout(header)

        self.input_text = QTextEdit()
        self.input_text.setPlaceholderText("Example: Move all PDFs to Documents folder")
        self.input_text.setFixedHeight(78)
        outer.addWidget(self.input_text)

        ex1 = QHBoxLayout()
        ex1.setSpacing(8)
        for ex in self.EXAMPLES[:3]:
            b = QPushButton(ex)
            b.setToolTip(ex)
            b.clicked.connect(lambda _=False, t=ex: self._use_example(t))
            ex1.addWidget(b)
        outer.addLayout(ex1)

        ex2 = QHBoxLayout()
        ex2.setSpacing(8)
        for ex in self.EXAMPLES[3:6]:
            b = QPushButton(ex)
            b.setToolTip(ex)
            b.clicked.connect(lambda _=False, t=ex: self._use_example(t))
            ex2.addWidget(b)
        outer.addLayout(ex2)

        action_row = QHBoxLayout()
        action_row.setSpacing(10)
        self.btn_generate = QPushButton("Generate")
        self.btn_generate.setObjectName("PrimaryButton")
        self.btn_generate.setIcon(self.icons.icon("brain"))
        self.btn_generate.setIconSize(QSize(18, 18))
        self.btn_generate.clicked.connect(self._generate)
        action_row.addWidget(self.btn_generate)

        self.btn_add = QPushButton("Add")
        self.btn_add.setEnabled(False)
        self.btn_add.setIcon(self.icons.icon("add"))
        self.btn_add.setIconSize(QSize(18, 18))
        self.btn_add.clicked.connect(self._add_rule)
        action_row.addWidget(self.btn_add)
        action_row.addStretch(1)
        outer.addLayout(action_row)

        self.conf_label = QLabel("No result yet.")
        self.conf_label.setStyleSheet("color:#9aa0a9;")
        outer.addWidget(self.conf_label)

        self.rule_view = QTextEdit()
        self.rule_view.setReadOnly(True)
        self.rule_view.setMinimumHeight(220)
        outer.addWidget(self.rule_view, 1)

        mon_row = QHBoxLayout()
        mon_row.setSpacing(10)
        mon_row.addWidget(QLabel("Apply to:"))
        self.monitor_combo = QComboBox()
        self.monitor_combo.setMinimumWidth(420)
        mon_row.addWidget(self.monitor_combo, 1)
        self._populate_monitors()
        outer.addLayout(mon_row)

        self.warn_label = QLabel("")
        self.warn_label.setStyleSheet("color:#FBC02D;")
        self.warn_label.setWordWrap(True)
        outer.addWidget(self.warn_label)

    def _use_example(self, text: str) -> None:
        self.input_text.setPlainText(text)
        self.input_text.setFocus()

    def _populate_monitors(self) -> None:
        self.monitor_combo.clear()
        self.monitor_combo.addItem("➕ Create New Monitor", userData=None)
        try:
            for mid, mon in self.backend.monitor_manager.monitors.items():
                if not hasattr(mon, "path"):
                    continue
                self.monitor_combo.addItem(f"{Path(mon.path).name} ({mon.path})", userData=str(mid))
        except Exception:
            pass

    def _generate(self) -> None:
        user_input = (self.input_text.toPlainText() or "").strip()
        if len(user_input) < 5:
            QMessageBox.warning(self, "Input Required", "Please describe what you want to automate.")
            return

        if callable(self._ensure_ai_ready_cb):
            if not self._ensure_ai_ready_cb(title="AI Required for Rule Builder"):
                return

        self.btn_generate.setEnabled(False)
        self.btn_add.setEnabled(False)
        self.conf_label.setText("Generating…")
        self.warn_label.setText("")
        self.rule_view.setPlainText("")

        worker = _QtRuleGenWorker(backend=self.backend, text=user_input)
        th = QThread(self)
        worker.moveToThread(th)
        th.started.connect(worker.run)
        worker.finished.connect(self._on_generated)
        worker.error.connect(self._on_error)
        worker.finished.connect(th.quit)
        worker.error.connect(th.quit)
        th.finished.connect(th.deleteLater)
        th.finished.connect(worker.deleteLater)
        self._thread = th
        self._worker = worker
        th.start()

    def _on_error(self, msg: str) -> None:
        self.btn_generate.setEnabled(True)
        self.conf_label.setText("Failed.")
        QMessageBox.critical(self, "Rule Generation Failed", msg)

    def _on_generated(self, res) -> None:
        self._result = res
        self.btn_generate.setEnabled(True)

        try:
            conf = float(getattr(res, "confidence", 0.0) or 0.0)
        except Exception:
            conf = 0.0
        interp = getattr(res, "interpretation", "") or ""
        expl = getattr(res, "explanation", "") or ""
        warnings = getattr(res, "warnings", []) or []
        rule = getattr(res, "rule", None)

        if rule:
            if rule.get("schedule"):
                self.conf_label.setText(f"⏰ Scheduled Task • confidence {conf:.0%} • {interp}")
            else:
                self.conf_label.setText(f"🤖 Rule • confidence {conf:.0%} • {interp}")
            try:
                import json

                self.rule_view.setPlainText(json.dumps(rule, indent=2))
            except Exception:
                self.rule_view.setPlainText(str(rule))
            self.btn_add.setEnabled(True)
        else:
            self.conf_label.setText(f"Low confidence ({conf:.0%}) • {interp}")
            self.rule_view.setPlainText(expl or "No rule generated.")
            self.btn_add.setEnabled(False)

        warn_lines = []
        if expl:
            warn_lines.append(expl)
        for w in warnings[:6]:
            warn_lines.append(str(w))
        self.warn_label.setText("\n".join(warn_lines).strip())

    def _add_rule(self) -> None:
        if not self._result or not getattr(self._result, "rule", None):
            return
        rule = dict(self._result.rule)
        prompt = (getattr(self._result, "original_input", "") or "").strip()
        if prompt:
            rule["ai_prompt"] = prompt

        if rule.get("schedule") and rule.get("action_type"):
            if rule.get("action_type") in {"clean_folder", "delete"}:
                ok = QMessageBox.question(
                    self,
                    "Confirm Scheduled Task",
                    "This task may delete files.\n\nCleanup tasks use Recycle Bin/app trash when available, skip active download files, and keep recent files unless the rule explicitly says otherwise.\n\nScheduled tasks run only while Fylorra is open.\n\nCreate it?",
                )
                if ok != QMessageBox.Yes:
                    return
            if not rule.get("target_path"):
                QMessageBox.warning(self, "Missing Target", "The generated task did not include a target_path.")
                return
            try:
                created = bool(self.backend.monitor_manager.add_scheduled_task(rule))
            except Exception as e:
                QMessageBox.critical(self, "Create Task Failed", str(e))
                return
            if not created:
                QMessageBox.critical(self, "Create Task Failed", "Could not create the scheduled task.")
                return
            try:
                self.backend.monitor_manager.save_monitors()
            except Exception:
                pass
            QMessageBox.information(self, "Created", "Scheduled task created.")
            self.rule_added.emit()
            return

        action_type = (rule.get("action_type") or "").strip().lower()
        if action_type in {"delete", "rename", "execute", "archive"}:
            ok = QMessageBox.question(
                self,
                "Confirm Rule",
                f"This rule performs a potentially destructive action ({action_type}).\n\nAdd it anyway?",
            )
            if ok != QMessageBox.Yes:
                return

        monitor_id = self.monitor_combo.currentData()
        if monitor_id:
            mon = self.backend.monitor_manager.monitors.get(str(monitor_id))
            if not mon:
                QMessageBox.critical(self, "Add Rule", "Selected monitor not found.")
                return
            if action_type in {"move", "copy"}:
                dest = (rule.get("action_params") or {}).get("destination", "") or ""
                try:
                    src = Path(mon.path).resolve()
                    dst = Path(dest).resolve() if dest else None
                    if dst and (dst == src or src in dst.parents):
                        QMessageBox.critical(
                            self,
                            "Invalid Rule",
                            "Move/Copy destination must be outside the watched folder to prevent automation loops.",
                        )
                        return
                except Exception:
                    pass
            try:
                problems = self.backend.monitor_manager.validate_monitor_config(mon.path, [rule])
                if problems:
                    QMessageBox.critical(self, "Invalid Rule", "\n".join(problems[:8]))
                    return
            except Exception:
                pass
            try:
                mon.rules.append(rule)
                self.backend.monitor_manager.save_monitors()
            except Exception as e:
                QMessageBox.critical(self, "Add Rule Failed", str(e))
                return
            QMessageBox.information(self, "Added", "Rule added to monitor.")
            self.rule_added.emit()
            return

        folder = QFileDialog.getExistingDirectory(self, "Select folder to monitor")
        if not folder:
            return
        if action_type in {"move", "copy"}:
            dest = (rule.get("action_params") or {}).get("destination", "") or ""
            try:
                src = Path(folder).resolve()
                dst = Path(dest).resolve() if dest else None
                if dst and (dst == src or src in dst.parents):
                    QMessageBox.critical(
                        self,
                        "Invalid Rule",
                        "Move/Copy destination must be outside the watched folder to prevent automation loops.",
                    )
                    return
            except Exception:
                pass
        import uuid

        new_id = str(uuid.uuid4())
        try:
            problems = self.backend.monitor_manager.validate_monitor_config(folder, [rule])
            if problems:
                QMessageBox.critical(self, "Invalid Monitor", "\n".join(problems[:8]))
                return
            ok = bool(
                self.backend.monitor_manager.add_monitor(
                    new_id,
                    folder,
                    [rule],
                    notify_created=True,
                    notify_modified=True,
                    notify_deleted=False,
                    notify_moved=False,
                )
            )
            self.backend.monitor_manager.save_monitors()
        except Exception as e:
            QMessageBox.critical(self, "Create Monitor Failed", str(e))
            return
        if not ok:
            QMessageBox.critical(self, "Create Monitor Failed", "Could not create monitor.")
            return
        QMessageBox.information(self, "Created", f"New monitor created for:\n{folder}")
        self.rule_added.emit()


class _QtAiRuleBuilderDialog(QDialog):
    EXAMPLES = [
        "Move all PDFs to Documents folder",
        "Copy new images to Backup drive",
        "Organize videos by date",
        "Move invoices with 'paid' in name to Archive folder",
        "Copy work documents to Projects folder",
        "Organize music files by file type",
        "Delete all content from temp folder C:\\\\Users\\\\me\\\\AppData\\\\Local\\\\Temp every day at 12:00 AM",
    ]

    def __init__(self, parent: QWidget, *, backend, icons: QtIconLoader, initial_monitor_id: str | None = None):
        super().__init__(parent)
        self.backend = backend
        self.icons = icons
        self._initial_monitor_id = str(initial_monitor_id) if initial_monitor_id else None
        self.result = None

        self.setWindowTitle("AI Rule Builder")
        self.setModal(True)
        self.setMinimumWidth(860)
        self.setMinimumHeight(640)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(18, 16, 18, 16)
        layout.setSpacing(12)

        title = QLabel("Natural Language Rule Builder")
        title.setStyleSheet("font-size:18px; font-weight:800; color:#ffffff;")
        sub = QLabel("Describe what you want in plain English — Fylorra will generate a rule or scheduled task.")
        sub.setStyleSheet("color:#9aa0a9;")
        sub.setWordWrap(True)
        layout.addWidget(title)
        layout.addWidget(sub)

        # Existing rules/tasks preview
        existing_box = QGroupBox("Existing Rules & Tasks")
        existing_layout = QVBoxLayout(existing_box)
        existing_layout.setContentsMargins(12, 10, 12, 10)
        existing_layout.setSpacing(8)
        self.existing_list = QListWidget()
        self.existing_list.setFixedHeight(120)
        self.existing_list.setStyleSheet(
            "QListWidget{background:#14171c; border:1px solid #232730; border-radius:10px; color:#c8ccd6;}"
        )
        existing_layout.addWidget(self.existing_list)
        layout.addWidget(existing_box)
        self._populate_existing()

        # Input
        input_box = QGroupBox("What do you want to automate?")
        input_layout = QVBoxLayout(input_box)
        input_layout.setContentsMargins(12, 10, 12, 10)
        input_layout.setSpacing(8)
        self.input_text = QTextEdit()
        self.input_text.setPlaceholderText("Example: “Move all PDFs to Documents folder”…")
        self.input_text.setFixedHeight(80)
        input_layout.addWidget(self.input_text)

        ex_row = QHBoxLayout()
        ex_row.setSpacing(8)
        for ex in self.EXAMPLES[:3]:
            b = QPushButton(ex)
            b.setToolTip(ex)
            b.clicked.connect(lambda _=False, t=ex: self._use_example(t))
            ex_row.addWidget(b)
        input_layout.addLayout(ex_row)
        ex_row2 = QHBoxLayout()
        ex_row2.setSpacing(8)
        for ex in self.EXAMPLES[3:6]:
            b = QPushButton(ex)
            b.setToolTip(ex)
            b.clicked.connect(lambda _=False, t=ex: self._use_example(t))
            ex_row2.addWidget(b)
        input_layout.addLayout(ex_row2)

        layout.addWidget(input_box)

        # Results
        self.results_box = QGroupBox("Result")
        res_layout = QVBoxLayout(self.results_box)
        res_layout.setContentsMargins(12, 10, 12, 10)
        res_layout.setSpacing(8)

        self.conf_label = QLabel("No result yet.")
        self.conf_label.setStyleSheet("color:#9aa0a9;")
        res_layout.addWidget(self.conf_label)

        self.rule_view = QTextEdit()
        self.rule_view.setReadOnly(True)
        self.rule_view.setMinimumHeight(220)
        res_layout.addWidget(self.rule_view, 1)

        # Monitor selection for event-based rules
        mon_row = QHBoxLayout()
        mon_row.setSpacing(10)
        mon_row.addWidget(QLabel("Apply to:"))
        self.monitor_combo = QComboBox()
        self.monitor_combo.setMinimumWidth(360)
        mon_row.addWidget(self.monitor_combo, 1)
        self._populate_monitors()
        if self._initial_monitor_id:
            try:
                idx = self.monitor_combo.findData(self._initial_monitor_id)
                if idx >= 0:
                    self.monitor_combo.setCurrentIndex(idx)
            except Exception:
                pass
        res_layout.addLayout(mon_row)

        self.warn_label = QLabel("")
        self.warn_label.setStyleSheet("color:#FBC02D;")
        self.warn_label.setWordWrap(True)
        res_layout.addWidget(self.warn_label)

        layout.addWidget(self.results_box, 1)

        # Bottom controls
        row = QHBoxLayout()
        row.addStretch(1)

        self.btn_add = QPushButton("Add This Rule")
        self.btn_add.setObjectName("PrimaryButton")
        self.btn_add.setEnabled(False)
        self.btn_add.clicked.connect(self._add_rule)
        row.addWidget(self.btn_add)

        self.btn_generate = QPushButton("Generate with AI")
        self.btn_generate.setObjectName("PrimaryButton")
        self.btn_generate.setIcon(self.icons.icon("brain"))
        self.btn_generate.setIconSize(QSize(18, 18))
        self.btn_generate.clicked.connect(self._generate)
        row.addWidget(self.btn_generate)

        self.btn_close = QPushButton("Close")
        self.btn_close.clicked.connect(self.reject)
        row.addWidget(self.btn_close)
        layout.addLayout(row)

        # ensure AI ready lazily when generating
        self._thread = None
        self._worker = None

        # Bring in front (Windows/Qt sometimes opens behind).
        try:
            self.setWindowModality(Qt.ApplicationModal)
        except Exception:
            pass
        try:
            self.raise_()
            self.activateWindow()
        except Exception:
            pass

    def _use_example(self, text: str):
        self.input_text.setPlainText(text)

    def _populate_existing(self):
        self.existing_list.clear()
        # scheduled tasks
        try:
            tasks = list(self.backend.monitor_manager.scheduled_tasks.list_tasks() or [])
        except Exception:
            tasks = []
        for t in tasks[:6]:
            try:
                title = getattr(t, "title", "Scheduled Task")
                self.existing_list.addItem(QListWidgetItem(f"⏰ {title}"))
            except Exception:
                pass
        # ai rules
        try:
            for _mid, mon in self.backend.monitor_manager.monitors.items():
                for r in getattr(mon, "rules", []) or []:
                    if _is_ai_rule(r):
                        prompt = r.get("ai_prompt") or ""
                        self.existing_list.addItem(QListWidgetItem(f"🤖 {Path(mon.path).name}: {prompt[:60]}"))
                        if self.existing_list.count() >= 12:
                            return
        except Exception:
            pass
        if self.existing_list.count() == 0:
            self.existing_list.addItem(QListWidgetItem("No AI rules or scheduled tasks yet."))

    def _populate_monitors(self):
        self.monitor_combo.clear()
        self.monitor_combo.addItem("➕ Create New Monitor", userData=None)
        try:
            for mid, mon in self.backend.monitor_manager.monitors.items():
                if not hasattr(mon, "path"):
                    continue
                self.monitor_combo.addItem(f"{Path(mon.path).name} ({mon.path})", userData=str(mid))
        except Exception:
            pass

    def _generate(self):
        user_input = (self.input_text.toPlainText() or "").strip()
        if len(user_input) < 5:
            QMessageBox.warning(self, "Input Required", "Please describe what you want to automate.")
            return

        # Ensure AI model is loaded
        top = self.parent().window()
        try:
            ensure = getattr(top, "_ensure_ai_ready", None)
            if callable(ensure):
                if not ensure(title="AI Required for Rule Builder"):
                    return
        except Exception:
            pass

        self.btn_generate.setEnabled(False)
        self.btn_add.setEnabled(False)
        self.conf_label.setText("Generating…")
        self.rule_view.setPlainText("")
        self.warn_label.setText("")

        worker = _QtRuleGenWorker(backend=self.backend, text=user_input)
        th = QThread(self)
        worker.moveToThread(th)
        th.started.connect(worker.run)
        worker.status.connect(self.conf_label.setText)
        worker.finished.connect(self._on_generated)
        worker.error.connect(self._on_error)
        worker.finished.connect(th.quit)
        worker.error.connect(th.quit)
        th.finished.connect(th.deleteLater)
        th.finished.connect(worker.deleteLater)
        self._thread = th
        self._worker = worker
        th.start()

    def _on_error(self, msg: str):
        self.btn_generate.setEnabled(True)
        QMessageBox.critical(self, "Rule Generation Failed", msg)
        self.conf_label.setText("Failed.")

    def _on_generated(self, res):
        self.result = res
        self.btn_generate.setEnabled(True)

        try:
            conf = float(getattr(res, "confidence", 0.0) or 0.0)
        except Exception:
            conf = 0.0
        interp = getattr(res, "interpretation", "") or ""
        expl = getattr(res, "explanation", "") or ""
        warnings = getattr(res, "warnings", []) or []
        rule = getattr(res, "rule", None)

        if rule:
            if rule.get("schedule"):
                self.conf_label.setText(f"⏰ Scheduled Task • confidence {conf:.0%} • {interp}")
            else:
                self.conf_label.setText(f"🤖 Rule • confidence {conf:.0%} • {interp}")
            # pretty json
            try:
                import json

                self.rule_view.setPlainText(json.dumps(rule, indent=2))
            except Exception:
                self.rule_view.setPlainText(str(rule))
            self.btn_add.setEnabled(True)
        else:
            self.conf_label.setText(f"Low confidence ({conf:.0%}) • {interp}")
            self.rule_view.setPlainText(expl or "No rule generated.")
            self.btn_add.setEnabled(False)

        warn_lines = []
        if expl:
            warn_lines.append(expl)
        for w in warnings[:6]:
            warn_lines.append(str(w))
        self.warn_label.setText("\n".join(warn_lines).strip())

    def _add_rule(self):
        if not self.result or not getattr(self.result, "rule", None):
            return
        rule = dict(self.result.rule)
        prompt = (getattr(self.result, "original_input", "") or "").strip()
        if prompt:
            rule["ai_prompt"] = prompt

        # Scheduled task?
        if rule.get("schedule") and rule.get("action_type"):
            # Require confirmation for destructive actions
            if rule.get("action_type") in {"clean_folder", "delete"}:
                ok = QMessageBox.question(
                    self,
                    "Confirm Scheduled Task",
                    "This task may delete files.\n\nCleanup tasks use Recycle Bin/app trash when available, skip active download files, and keep recent files unless the rule explicitly says otherwise.\n\nScheduled tasks run only while Fylorra is open.\n\nCreate it?",
                )
                if ok != QMessageBox.Yes:
                    return
            # Ensure task has required fields
            if not rule.get("target_path"):
                QMessageBox.warning(self, "Missing Target", "The generated task did not include a target_path.")
                return
            try:
                created = bool(self.backend.monitor_manager.add_scheduled_task(rule))
            except Exception as e:
                QMessageBox.critical(self, "Create Task Failed", str(e))
                return
            if not created:
                QMessageBox.critical(self, "Create Task Failed", "Could not create the scheduled task.")
                return
            try:
                self.backend.monitor_manager.save_monitors()
            except Exception:
                pass
            QMessageBox.information(self, "Created", "Scheduled task created.")
            self.accept()
            return

        # Event-based rule
        action_type = (rule.get("action_type") or "").strip().lower()
        if action_type in {"delete", "rename", "execute", "archive"}:
            ok = QMessageBox.question(
                self,
                "Confirm Rule",
                f"This rule performs a potentially destructive action ({action_type}).\n\nAdd it anyway?",
            )
            if ok != QMessageBox.Yes:
                return

        monitor_id = self.monitor_combo.currentData()
        if monitor_id:
            mon = self.backend.monitor_manager.monitors.get(str(monitor_id))
            if not mon:
                QMessageBox.critical(self, "Add Rule", "Selected monitor not found.")
                return
            # Loop guard: keep generated copy/move rules outside the watched tree.
            if action_type in {"move", "copy"}:
                dest = (rule.get("action_params") or {}).get("destination", "") or ""
                try:
                    src = Path(mon.path).resolve()
                    dst = Path(dest).resolve() if dest else None
                    if dst and (dst == src or src in dst.parents):
                        QMessageBox.critical(
                            self,
                            "Invalid Rule",
                            "Move/Copy destination must be outside the watched folder to prevent automation loops.",
                        )
                        return
                except Exception:
                    pass
            try:
                problems = self.backend.monitor_manager.validate_monitor_config(mon.path, [rule])
                if problems:
                    QMessageBox.critical(self, "Invalid Rule", "\n".join(problems[:8]))
                    return
            except Exception:
                pass
            try:
                mon.rules.append(rule)
                self.backend.monitor_manager.save_monitors()
            except Exception as e:
                QMessageBox.critical(self, "Add Rule Failed", str(e))
                return
            QMessageBox.information(self, "Added", "Rule added to monitor.")
            self.accept()
            return

        # Create new monitor
        folder = QFileDialog.getExistingDirectory(self, "Select folder to monitor")
        if not folder:
            return
        # Loop guard for new monitor
        if action_type in {"move", "copy"}:
            dest = (rule.get("action_params") or {}).get("destination", "") or ""
            try:
                src = Path(folder).resolve()
                dst = Path(dest).resolve() if dest else None
                if dst and (dst == src or src in dst.parents):
                    QMessageBox.critical(
                        self,
                        "Invalid Rule",
                        "Move/Copy destination must be outside the watched folder to prevent automation loops.",
                    )
                    return
            except Exception:
                pass
        try:
            problems = self.backend.monitor_manager.validate_monitor_config(folder, [rule])
            if problems:
                QMessageBox.critical(self, "Invalid Monitor", "\n".join(problems[:8]))
                return
        except Exception:
            pass
        import uuid

        new_id = str(uuid.uuid4())
        try:
            ok = bool(
                self.backend.monitor_manager.add_monitor(
                    new_id,
                    folder,
                    [rule],
                    notify_created=True,
                    notify_modified=True,
                    notify_deleted=False,
                    notify_moved=False,
                )
            )
            self.backend.monitor_manager.save_monitors()
        except Exception as e:
            QMessageBox.critical(self, "Create Monitor Failed", str(e))
            return
        if not ok:
            QMessageBox.critical(self, "Create Monitor Failed", "Could not create monitor.")
            return
        QMessageBox.information(self, "Created", f"New monitor created for:\n{folder}")
        self.accept()


class _QtMediaBatchWorker(QObject):
    status = Signal(str)
    progress = Signal(float)
    finished = Signal(str)
    error = Signal(str)
    file_progress = Signal(str, float)

    def __init__(
        self,
        *,
        folder: str,
        source_subfolder: str | None,
        include_subfolders: bool,
        output_format: str,
        output_subfolder: str,
        output_root: str,
        output_directory: str | None,
        preserve_subfolders: bool,
        overwrite: bool,
        audio_bitrate: str | None,
        preserve_metadata: bool,
        preserve_cover_art: bool,
        video_codec: str | None,
        scale_height: int | None,
        use_gpu: bool,
        audio_codec: str | None,
    ):
        super().__init__()
        self.folder = folder
        self.source_subfolder = source_subfolder
        self.include_subfolders = include_subfolders
        self.output_format = output_format
        self.output_subfolder = output_subfolder
        self.output_root = output_root
        self.output_directory = output_directory
        self.preserve_subfolders = preserve_subfolders
        self.overwrite = overwrite
        self.audio_bitrate = audio_bitrate
        self.preserve_metadata = preserve_metadata
        self.preserve_cover_art = preserve_cover_art
        self.video_codec = video_codec
        self.scale_height = scale_height
        self.use_gpu = use_gpu
        self.audio_codec = audio_codec
        import threading

        self._cancel = threading.Event()

    def cancel(self):
        self._cancel.set()

    def run(self):
        try:
            from pathlib import Path
            from core.media_converter import convert_media_in_folder

            current_box = {"cur": 0, "total": 0}

            def overall_cb(cur: int, total: int, p: Path):
                if total <= 0:
                    return
                current_box["cur"] = cur
                current_box["total"] = total
                rel = p.name
                try:
                    rel = str(p.relative_to(Path(self.folder)))
                except Exception:
                    pass
                self.status.emit(f"Converting {cur}/{total}: {rel}")
                self.progress.emit(max(0.0, min(1.0, float(cur - 1) / float(max(1, total)))))

            def file_cb(p: Path, frac: float):
                frac = max(0.0, min(1.0, float(frac)))
                cur = int(current_box.get("cur") or 1)
                total = int(current_box.get("total") or 1)
                overall = (float(cur - 1) + frac) / float(max(1, total))
                self.progress.emit(max(0.0, min(1.0, overall)))
                try:
                    self.file_progress.emit(str(p), frac)
                except Exception:
                    pass

            res = convert_media_in_folder(
                Path(self.folder),
                source_subfolder=self.source_subfolder,
                include_subfolders=bool(self.include_subfolders),
                output_format=self.output_format,
                output_subfolder=self.output_subfolder,
                output_root=self.output_root,
                output_directory=self.output_directory,
                preserve_structure=bool(self.preserve_subfolders),
                preserve_subfolders=bool(self.preserve_subfolders),
                overwrite=bool(self.overwrite),
                audio_bitrate=self.audio_bitrate,
                preserve_metadata=bool(self.preserve_metadata),
                preserve_cover_art=bool(self.preserve_cover_art),
                cancel_event=self._cancel,
                progress_cb=overall_cb,
                file_progress_cb=file_cb,
                use_gpu=bool(self.use_gpu),
                video_codec=self.video_codec,
                scale_height=self.scale_height,
                audio_codec=self.audio_codec,
            )
            if not res.ok:
                self.error.emit(res.message or "Conversion failed.")
                return
            self.progress.emit(1.0)
            self.finished.emit(res.message or "Batch conversion complete.")
        except Exception as e:
            msg = str(e)
            if "Cancelled" in msg:
                self.error.emit("Cancelled.")
            else:
                self.error.emit(msg)


class _QtMediaSingleWorker(QObject):
    status = Signal(str)
    progress = Signal(float)
    finished = Signal(str)
    error = Signal(str)
    file_progress = Signal(str, float)

    def __init__(
        self,
        *,
        target_folder: str,
        input_path: str,
        output_format: str,
        output_name: str,
        output_root: str,
        output_directory: str | None,
        overwrite: bool,
        audio_bitrate: str | None,
        preserve_metadata: bool,
        preserve_cover_art: bool,
        use_gpu: bool,
        video_codec: str | None,
        scale_height: int | None,
        audio_codec: str | None,
    ):
        super().__init__()
        self.target_folder = target_folder
        self.input_path = input_path
        self.output_format = output_format
        self.output_name = output_name
        self.output_root = output_root
        self.output_directory = output_directory
        self.overwrite = overwrite
        self.audio_bitrate = audio_bitrate
        self.preserve_metadata = preserve_metadata
        self.preserve_cover_art = preserve_cover_art
        self.use_gpu = use_gpu
        self.video_codec = video_codec
        self.scale_height = scale_height
        self.audio_codec = audio_codec
        import threading

        self._cancel = threading.Event()

    def cancel(self):
        self._cancel.set()

    def run(self):
        try:
            from pathlib import Path
            from core.media_tools import convert_media_file

            inp = Path(self.input_path)
            if self.output_root == "target":
                base = Path(self.target_folder) if self.target_folder else inp.parent
            elif self.output_root == "custom":
                dest = Path(str(self.output_directory or "").strip())
                if not str(dest):
                    self.error.emit("Destination folder is required.")
                    return
                if not dest.is_absolute():
                    base_root = Path(self.target_folder) if self.target_folder else inp.parent
                    dest = base_root / dest
                base = dest
            else:
                base = inp.parent
            try:
                base.mkdir(parents=True, exist_ok=True)
            except Exception:
                pass
            out_path = base / f"{self.output_name}.{self.output_format}"

            self.status.emit("Converting…")

            def cb(frac: float):
                self.progress.emit(max(0.0, min(1.0, float(frac))))
                try:
                    self.file_progress.emit(str(inp), max(0.0, min(1.0, float(frac))))
                except Exception:
                    pass

            res = convert_media_file(
                inp,
                output_path=out_path,
                overwrite=bool(self.overwrite),
                audio_bitrate=self.audio_bitrate,
                preserve_metadata=bool(self.preserve_metadata),
                preserve_cover_art=bool(self.preserve_cover_art),
                use_gpu=bool(self.use_gpu),
                video_codec=self.video_codec,
                scale_height=self.scale_height,
                audio_codec=self.audio_codec,
                cancel_event=self._cancel,
                progress_cb=cb,
            )
            if not res.ok:
                self.error.emit(res.message or "Conversion failed.")
                return
            self.progress.emit(1.0)
            self.finished.emit(res.message or f"Saved: {out_path}")
        except Exception as e:
            msg = str(e)
            if "Cancelled" in msg:
                self.error.emit("Cancelled.")
            else:
                self.error.emit(msg)


class _QtImagesWorker(QObject):
    status = Signal(str)
    progress = Signal(float)
    finished = Signal(str)
    error = Signal(str)

    def __init__(self, *, folder: str, include_subfolders: bool, output_format: str, output_subfolder: str, overwrite: bool):
        super().__init__()
        self.folder = folder
        self.include_subfolders = include_subfolders
        self.output_format = output_format
        self.output_subfolder = output_subfolder
        self.overwrite = overwrite
        import threading

        self._cancel = threading.Event()

    def cancel(self):
        self._cancel.set()

    def run(self):
        try:
            from pathlib import Path
            from core.image_converter import convert_images_in_folder

            def cb(cur: int, total: int, p: Path):
                if total <= 0:
                    return
                self.status.emit(f"Converting {cur}/{total}: {p.name}")
                self.progress.emit(max(0.0, min(1.0, float(cur - 1) / float(max(1, total)))))

            res = convert_images_in_folder(
                Path(self.folder),
                include_subfolders=bool(self.include_subfolders),
                output_format=self.output_format,
                output_subfolder=self.output_subfolder,
                overwrite=bool(self.overwrite),
                progress_cb=cb,
                cancel_event=self._cancel,
            )
            if not res.ok:
                self.error.emit(res.message or "Conversion failed.")
                return
            self.progress.emit(1.0)
            self.finished.emit(res.message or "Image conversion complete.")
        except Exception as e:
            msg = str(e)
            if "Cancelled" in msg:
                self.error.emit("Cancelled.")
            else:
                self.error.emit(msg)


class _QtTextExtractWorker(QObject):
    status = Signal(str)
    progress = Signal(float)
    finished = Signal(str)
    error = Signal(str)
    text_ready = Signal(str)
    per_file_ready = Signal(str, str)

    def __init__(
        self,
        *,
        files: list[str],
        ocr_mode: str,
        ocr_engine: str,
        ocr_lang: str,
        max_ocr_pages: int,
        output_format: str,
        output_dir: str | None,
        normalize: bool,
        add_headers: bool,
        use_ai: bool,
        translate: bool,
        translate_lang: str,
        ai_manager=None,
    ):
        super().__init__()
        self.files = list(files)
        self.ocr_mode = ocr_mode
        self.ocr_engine = ocr_engine
        self.ocr_lang = ocr_lang
        self.max_ocr_pages = max_ocr_pages
        self.output_format = output_format
        self.output_dir = output_dir
        self.normalize = normalize
        self.add_headers = add_headers
        self.use_ai = use_ai
        self.translate = translate
        self.translate_lang = translate_lang
        self.ai_manager = ai_manager
        import threading

        self._cancel = threading.Event()

    def cancel(self):
        self._cancel.set()

    def _ai_ready(self) -> bool:
        return bool(self.ai_manager and getattr(self.ai_manager, "is_ready", False) and getattr(self.ai_manager, "model", None))

    def _ai_call(self, prompt: str, text: str) -> str:
        if not self._ai_ready():
            raise RuntimeError("AI model not loaded.")
        response = self.ai_manager.model.create_chat_completion(
            messages=[
                {"role": "system", "content": "You are a precise text processing assistant."},
                {"role": "user", "content": f"{prompt}\n\n{text}"},
            ],
            temperature=0.1,
            max_tokens=1200,
        )
        return (response["choices"][0]["message"]["content"] or "").strip()

    def _ai_cleanup(self, text: str, *, to_markdown: bool) -> str:
        if not text.strip():
            return text
        if len(text) > 12000:
            return text
        if to_markdown:
            prompt = "Clean and format the text into readable Markdown. Preserve meaning. Return only the Markdown."
        else:
            prompt = "Clean the text (fix line breaks, remove repeated headers/footers if obvious). Return only the cleaned text."
        return self._ai_call(prompt, text)

    def _chunk_text(self, text: str, max_chars: int) -> list[str]:
        if len(text) <= max_chars:
            return [text]
        parts: list[str] = []
        cur: list[str] = []
        cur_len = 0
        for line in text.splitlines():
            add_len = len(line) + 1
            if cur_len + add_len > max_chars and cur:
                parts.append("\n".join(cur))
                cur = [line]
                cur_len = len(line)
            else:
                cur.append(line)
                cur_len += add_len
        if cur:
            parts.append("\n".join(cur))
        return parts or [text]

    def _ai_translate(self, text: str, target_lang: str) -> str:
        if not text.strip():
            return text
        parts = self._chunk_text(text, 4000)
        if len(parts) == 1:
            prompt = f"Translate the text to {target_lang}. Return only the translation."
            return self._ai_call(prompt, text)
        out: list[str] = []
        total = len(parts)
        for idx, part in enumerate(parts, start=1):
            prompt = f"Translate the text to {target_lang}. Return only the translation. Part {idx} of {total}."
            out.append(self._ai_call(prompt, part))
        return "\n\n".join([o for o in out if o is not None]).strip()

    def run(self):
        try:
            from core.text_extractor import extract_text_from_file, normalize_text

            total = len(self.files)
            combined: list[str] = []
            out_paths: list[str] = []

            for i, f in enumerate(self.files, start=1):
                if self._cancel.is_set():
                    self.error.emit("Cancelled.")
                    return
                p = Path(f)
                self.status.emit(f"Extracting {i}/{total}: {p.name}")
                res = extract_text_from_file(
                    p,
                    ocr_mode=self.ocr_mode,
                    ocr_engine=self.ocr_engine,
                    ocr_lang=self.ocr_lang,
                    max_ocr_pages=self.max_ocr_pages,
                    ai_manager=self.ai_manager,
                )
                if not res.ok:
                    self.error.emit(res.message or f"Failed: {p.name}")
                    return
                text = res.text or ""
                if self.normalize:
                    text = normalize_text(text)
                if self.use_ai:
                    text = self._ai_cleanup(text, to_markdown="Markdown" in self.output_format)
                if self.translate:
                    text = self._ai_translate(text, self.translate_lang)

                if self.add_headers or total > 1:
                    header = f"# {p.name}" if "Markdown" in self.output_format else f"=== {p.name} ==="
                    combined.append(header)
                combined.append(text)
                try:
                    self.per_file_ready.emit(str(p), text)
                except Exception:
                    pass

                if self.output_dir:
                    out_dir = Path(self.output_dir)
                    out_dir.mkdir(parents=True, exist_ok=True)
                    if "Markdown" in self.output_format:
                        out_path = out_dir / f"{p.stem}.md"
                        out_path.write_text(text, encoding="utf-8", errors="replace")
                    elif "DOCX" in self.output_format:
                        try:
                            import docx
                        except Exception:
                            self.error.emit("DOCX export requires python-docx.")
                            return
                        out_path = out_dir / f"{p.stem}.docx"
                        doc = docx.Document()
                        for line in text.splitlines():
                            doc.add_paragraph(line)
                        doc.save(out_path)
                    else:
                        out_path = out_dir / f"{p.stem}.txt"
                        out_path.write_text(text, encoding="utf-8", errors="replace")
                    out_paths.append(str(out_path))

                self.progress.emit(max(0.0, min(1.0, float(i) / float(max(1, total)))))

            preview = "\n\n".join([c for c in combined if c is not None])
            self.text_ready.emit(preview)
            if out_paths:
                self.finished.emit(f"Extracted {total} file(s). Saved to {self.output_dir}")
            else:
                self.finished.emit(f"Extracted {total} file(s).")
        except Exception as e:
            msg = str(e)
            if "Cancelled" in msg:
                self.error.emit("Cancelled.")
            else:
                self.error.emit(msg)


class _QtOCRBoxesWorker(QObject):
    finished = Signal(list)
    error = Signal(str)

    def __init__(self, *, image):
        super().__init__()
        self.image = image

    def run(self):
        try:
            import numpy as np
            from rapidocr_onnxruntime import RapidOCR

            arr = self.image
            if not isinstance(arr, np.ndarray):
                arr = np.array(arr)
            ocr = RapidOCR()
            res = ocr(arr)
            if isinstance(res, tuple) and res:
                res = res[0]
            boxes = []
            if res:
                for item in res:
                    if isinstance(item, (list, tuple)) and item:
                        boxes.append(item[0])
            self.finished.emit(boxes)
        except Exception as e:
            self.error.emit(str(e))


class _QtOCRCompareDialog(QDialog):
    def __init__(self, parent: QWidget, *, file_path: str, extracted_text: str, ocr_engine: str):
        super().__init__(parent)
        self.setWindowTitle("OCR Comparison")
        self.resize(980, 620)
        self.file_path = str(file_path)
        self.extracted_text = extracted_text or ""
        self.ocr_engine = (ocr_engine or "Auto").strip()
        self._base_pixmap = None
        self._overlay_pixmap = None
        self._ocr_boxes: list = []
        self._box_worker = None
        self._box_thread = None
        self._box_connected = False
        self._closing = False
        self._is_pdf = False
        self._pdf_page = 1
        self._pdf_page_count = 0

        root = QVBoxLayout(self)
        root.setContentsMargins(14, 12, 14, 12)
        root.setSpacing(10)

        header = QHBoxLayout()
        header.setSpacing(10)
        title = QLabel("Compare: original vs extracted text")
        title.setStyleSheet("font-weight:700; color:#ffffff;")
        header.addWidget(title)
        header.addStretch(1)
        self._page_label = QLabel("Page")
        self._page_spin = QSpinBox()
        self._page_spin.setRange(1, 1)
        self._page_spin.setValue(1)
        self._page_spin.setMinimumWidth(260)
        self._page_spin.setMaximumWidth(320)
        self._page_spin.setAlignment(Qt.AlignLeft | Qt.AlignVCenter)
        self._page_spin.setVisible(False)
        self._page_label.setVisible(False)
        self._page_spin.valueChanged.connect(self._on_page_changed)
        header.addWidget(self._page_label)
        header.addWidget(self._page_spin)
        self._box_toggle = QCheckBox("Show OCR boxes")
        self._box_toggle.toggled.connect(self._toggle_boxes)
        header.addWidget(self._box_toggle)
        root.addLayout(header)

        split = QSplitter(Qt.Horizontal)
        split.setChildrenCollapsible(False)

        left = QFrame()
        lv = QVBoxLayout(left)
        lv.setContentsMargins(0, 0, 0, 0)
        lv.setSpacing(8)
        self._img_label = QLabel()
        self._img_label.setAlignment(Qt.AlignCenter)
        self._img_label.setStyleSheet("background:#0f1217; border:1px solid #232730; border-radius:10px;")
        self._img_label.setSizePolicy(QSizePolicy.Ignored, QSizePolicy.Ignored)
        self._img_label.setMinimumSize(1, 1)
        lv.addWidget(self._img_label, 1)
        self._img_note = QLabel("")
        self._img_note.setStyleSheet("color:#9aa0a9;")
        lv.addWidget(self._img_note)
        split.addWidget(left)

        right = QFrame()
        rv = QVBoxLayout(right)
        rv.setContentsMargins(0, 0, 0, 0)
        rv.setSpacing(8)
        self._text_view = QTextEdit()
        self._text_view.setReadOnly(True)
        self._text_view.setPlainText(self.extracted_text)
        rv.addWidget(self._text_view, 1)
        split.addWidget(right)

        split.setSizes([520, 460])
        root.addWidget(split, 1)

        self._load_base_image()
        self._update_image()

    def resizeEvent(self, event):  # noqa: N802
        super().resizeEvent(event)
        self._update_image()

    def closeEvent(self, event):  # noqa: N802
        self._closing = True
        self._stop_box_worker()
        super().closeEvent(event)

    def _load_base_image(self):
        p = Path(self.file_path)
        if not p.exists():
            self._img_note.setText("Original file not found.")
            self._page_spin.setVisible(False)
            self._page_label.setVisible(False)
            return
        ext = p.suffix.lower()
        img = None
        self._is_pdf = False
        self._pdf_page_count = 0
        if ext in {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".gif", ".tif", ".tiff"}:
            reader = QImageReader(str(p))
            reader.setAutoTransform(True)
            img = reader.read()
        elif ext == ".pdf":
            try:
                import fitz  # type: ignore

                with fitz.open(str(p)) as doc:
                    if doc.page_count > 0:
                        self._is_pdf = True
                        self._pdf_page_count = doc.page_count
                        if self._pdf_page < 1 or self._pdf_page > doc.page_count:
                            self._pdf_page = 1
                        page = doc.load_page(self._pdf_page - 1)
                        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                        fmt = QImage.Format_RGBA8888 if pix.alpha else QImage.Format_RGB888
                        img = QImage(pix.samples, pix.width, pix.height, pix.stride, fmt).copy()
            except Exception:
                img = None
        if img is None or img.isNull():
            self._img_note.setText("Preview not available for this file type.")
            self._page_spin.setVisible(False)
            self._page_label.setVisible(False)
            return
        self._base_pixmap = QPixmap.fromImage(img)
        if self._is_pdf and self._pdf_page_count > 0:
            self._page_label.setVisible(True)
            self._page_spin.setVisible(True)
            self._page_spin.blockSignals(True)
            self._page_spin.setRange(1, self._pdf_page_count)
            self._page_spin.setValue(self._pdf_page)
            self._page_spin.setSuffix(f" / {self._pdf_page_count}")
            self._page_spin.blockSignals(False)
            self._img_note.setText(f"{p.name} (page {self._pdf_page}/{self._pdf_page_count})")
        else:
            self._page_spin.setVisible(False)
            self._page_label.setVisible(False)
            self._page_spin.setSuffix("")
            self._img_note.setText(p.name)

    def _update_image(self):
        if self._base_pixmap is None:
            return
        pix = self._overlay_pixmap if (self._box_toggle.isChecked() and self._overlay_pixmap) else self._base_pixmap
        target = self._img_label.size()
        if target.width() < 2 or target.height() < 2:
            return
        scaled = pix.scaled(target, Qt.KeepAspectRatio, Qt.SmoothTransformation)
        self._img_label.setPixmap(scaled)

    def _toggle_boxes(self, checked: bool):
        if not checked:
            self._stop_box_worker()
            self._update_image()
            return
        if self._overlay_pixmap is not None:
            self._update_image()
            return
        if self._base_pixmap is None:
            return
        if self._box_thread and self._box_thread.isRunning():
            return
        try:
            import importlib.util

            if importlib.util.find_spec("rapidocr_onnxruntime") is None:
                self._img_note.setText("RapidOCR not available.")
                return
        except Exception:
            self._img_note.setText("RapidOCR not available.")
            return
        self._img_note.setText("Analyzing OCR boxes…")
        self._start_box_worker()

    def _on_page_changed(self, value: int):
        if not self._is_pdf:
            return
        try:
            self._pdf_page = int(value)
        except Exception:
            self._pdf_page = 1
        self._stop_box_worker()
        self._overlay_pixmap = None
        self._ocr_boxes = []
        self._load_base_image()
        self._update_image()
        if self._box_toggle.isChecked():
            self._toggle_boxes(True)

    def _qimage_to_numpy(self, img: QImage):
        import numpy as np

        try:
            img = img.convertToFormat(QImage.Format_RGB888)
            w = img.width()
            h = img.height()
            ptr = img.constBits()
            ptr.setsize(h * img.bytesPerLine())
            arr = np.frombuffer(ptr, np.uint8).reshape((h, img.bytesPerLine()))
            arr = arr[:, : w * 3].reshape((h, w, 3))
            return arr.copy()
        except Exception:
            pass

        try:
            img = img.convertToFormat(QImage.Format_RGBA8888)
            w = img.width()
            h = img.height()
            ptr = img.constBits()
            ptr.setsize(h * img.bytesPerLine())
            arr = np.frombuffer(ptr, np.uint8).reshape((h, img.bytesPerLine()))
            arr = arr[:, : w * 4].reshape((h, w, 4))
            return arr[:, :, :3].copy()
        except Exception:
            pass

        try:
            import io
            from PIL import Image

            buf = QByteArray()
            qbuf = QBuffer(buf)
            qbuf.open(QIODevice.WriteOnly)
            img.save(qbuf, "PNG")
            qbuf.close()
            pil_img = Image.open(io.BytesIO(bytes(buf))).convert("RGB")
            return np.array(pil_img)
        except Exception as e:
            raise RuntimeError(f"Failed to convert image for OCR boxes: {e}") from e

    def _is_valid_qobj(self, obj) -> bool:
        if obj is None:
            return False
        try:
            import shiboken6

            return bool(shiboken6.isValid(obj))
        except Exception:
            return True

    def _stop_box_worker(self):
        worker = self._box_worker
        thread = self._box_thread
        if self._box_connected and self._is_valid_qobj(worker):
            try:
                worker.finished.disconnect(self._ocr_boxes_ready)
            except Exception:
                pass
            try:
                worker.error.disconnect(self._ocr_boxes_error)
            except Exception:
                pass
        if self._is_valid_qobj(thread):
            try:
                if thread.isRunning():
                    thread.quit()
            except Exception:
                pass
            try:
                thread.wait(1500)
            except Exception:
                pass
        self._box_worker = None
        self._box_thread = None
        self._box_connected = False

    def _start_box_worker(self):
        if self._base_pixmap is None:
            return
        try:
            img = self._base_pixmap.toImage()
            arr = self._qimage_to_numpy(img)
        except Exception as e:
            self._img_note.setText(str(e) or "Failed to prepare image for OCR boxes.")
            return

        worker = _QtOCRBoxesWorker(image=arr)
        thread = QThread(self)
        worker.moveToThread(thread)
        thread.started.connect(worker.run)  # type: ignore[attr-defined]
        worker.finished.connect(self._ocr_boxes_ready)
        worker.error.connect(self._ocr_boxes_error)
        worker.finished.connect(thread.quit)
        worker.error.connect(thread.quit)
        thread.finished.connect(thread.deleteLater)
        thread.finished.connect(worker.deleteLater)
        self._box_worker = worker
        self._box_thread = thread
        self._box_connected = True
        thread.start()

    def _ocr_boxes_ready(self, boxes: list):
        if self._closing:
            return
        self._ocr_boxes = boxes or []
        self._overlay_pixmap = self._draw_boxes()
        if self._is_pdf and self._pdf_page_count > 0:
            self._img_note.setText(f"{Path(self.file_path).name} (page {self._pdf_page}/{self._pdf_page_count})")
        else:
            self._img_note.setText(Path(self.file_path).name)
        self._update_image()
        self._box_connected = False

    def _ocr_boxes_error(self, msg: str):
        if self._closing:
            return
        self._img_note.setText(msg or "Failed to compute OCR boxes.")
        self._box_toggle.setChecked(False)
        self._update_image()
        self._box_connected = False

    def _draw_boxes(self):
        if self._base_pixmap is None:
            return None
        if not self._ocr_boxes:
            return None
        pix = QPixmap(self._base_pixmap)
        painter = QPainter(pix)
        pen = painter.pen()
        pen.setColor(Qt.yellow)
        pen.setWidth(2)
        painter.setPen(pen)
        for box in self._ocr_boxes:
            try:
                pts = box
                if isinstance(pts, (list, tuple)):
                    poly = []
                    for p in pts:
                        if isinstance(p, (list, tuple)) and len(p) >= 2:
                            poly.append(QPointF(float(p[0]), float(p[1])))
                    if len(poly) >= 3:
                        painter.drawPolygon(QPolygonF(poly))
            except Exception:
                continue
        painter.end()
        return pix
class _QtArchiveCreateWorker(QObject):
    status = Signal(str)
    progress = Signal(float)
    finished = Signal(str)
    error = Signal(str)

    def __init__(self, *, folder: str, archive_path: str, fmt: str, overwrite: bool, part_size_mb: int):
        super().__init__()
        self.folder = folder
        self.archive_path = archive_path
        self.fmt = fmt
        self.overwrite = overwrite
        self.part_size_mb = part_size_mb
        import threading

        self._cancel = threading.Event()

    def cancel(self):
        self._cancel.set()

    def run(self):
        try:
            from pathlib import Path
            from core.archive_tools import create_archive

            def cb(cur: int, total: int, p: Path):
                if total <= 0:
                    return
                self.status.emit(f"Archiving {cur}/{total}: {p.name}")
                self.progress.emit(max(0.0, min(1.0, float(cur - 1) / float(max(1, total)))))

            part_bytes = None
            if int(self.part_size_mb) > 0:
                part_bytes = int(self.part_size_mb) * 1024 * 1024
            res = create_archive(
                Path(self.folder),
                archive_path=Path(self.archive_path),
                fmt=self.fmt,
                overwrite=bool(self.overwrite),
                progress_cb=cb,
                cancel_event=self._cancel,
                part_size_bytes=part_bytes,
            )
            if not res.ok:
                self.error.emit(res.message or "Archive failed.")
                return
            self.progress.emit(1.0)
            self.finished.emit(res.message or f"Saved: {self.archive_path}")
        except Exception as e:
            msg = str(e)
            if "Cancelled" in msg:
                self.error.emit("Cancelled.")
            else:
                self.error.emit(msg)


class _QtArchiveExtractWorker(QObject):
    status = Signal(str)
    progress = Signal(float)
    finished = Signal(str)
    error = Signal(str)

    def __init__(self, *, archive_path: str, output_dir: str, overwrite: bool):
        super().__init__()
        self.archive_path = archive_path
        self.output_dir = output_dir
        self.overwrite = overwrite
        import threading

        self._cancel = threading.Event()

    def cancel(self):
        self._cancel.set()

    def run(self):
        try:
            from pathlib import Path
            from core.archive_tools import extract_archive

            def cb(cur: int, total: int, name: str):
                if total <= 0:
                    return
                self.status.emit(f"Extracting {cur}/{total}: {name}")
                self.progress.emit(max(0.0, min(1.0, float(cur - 1) / float(max(1, total)))))

            res = extract_archive(
                Path(self.archive_path),
                output_dir=Path(self.output_dir),
                overwrite=bool(self.overwrite),
                progress_cb=cb,
                cancel_event=self._cancel,
            )
            if not res.ok:
                self.error.emit(res.message or "Extract failed.")
                return
            self.progress.emit(1.0)
            self.finished.emit(res.message or f"Extracted to: {self.output_dir}")
        except Exception as e:
            msg = str(e)
            if "Cancelled" in msg:
                self.error.emit("Cancelled.")
            else:
                self.error.emit(msg)


class _QtPdfMergeWorker(QObject):
    status = Signal(str)
    progress = Signal(float)
    finished = Signal(str)
    error = Signal(str)

    def __init__(self, *, folder: str, include_subfolders: bool, output_pdf: str, overwrite: bool):
        super().__init__()
        self.folder = folder
        self.include_subfolders = include_subfolders
        self.output_pdf = output_pdf
        self.overwrite = overwrite
        import threading

        self._cancel = threading.Event()

    def cancel(self):
        self._cancel.set()

    def run(self):
        try:
            from core.pdf_tools import merge_pdfs

            base = Path(self.folder)
            pattern = "**/*.pdf" if bool(self.include_subfolders) else "*.pdf"
            pdfs = [p for p in base.glob(pattern) if p.is_file()]
            if not pdfs:
                self.error.emit("No PDFs found in the target folder.")
                return

            # Basic progress by file count.
            total = len(pdfs)
            for i, p in enumerate(pdfs, start=1):
                if self._cancel.is_set():
                    self.error.emit("Cancelled.")
                    return
                self.status.emit(f"Queued {i}/{total}: {p.name}")
                self.progress.emit(max(0.0, min(1.0, float(i - 1) / float(max(1, total)))))

            res = merge_pdfs(pdfs, output_pdf=Path(self.output_pdf), overwrite=bool(self.overwrite))
            if not res.ok:
                self.error.emit(res.message or "Merge failed.")
                return
            self.progress.emit(1.0)
            self.finished.emit(res.message or f"Saved: {self.output_pdf}")
        except Exception as e:
            msg = str(e)
            if "Cancelled" in msg:
                self.error.emit("Cancelled.")
            else:
                self.error.emit(msg)


class _QtPdfSplitPagesWorker(QObject):
    status = Signal(str)
    progress = Signal(float)
    finished = Signal(str)
    error = Signal(str)

    def __init__(self, *, input_pdf: str, output_dir: str, page_ranges: str, overwrite: bool):
        super().__init__()
        self.input_pdf = input_pdf
        self.output_dir = output_dir
        self.page_ranges = page_ranges
        self.overwrite = overwrite
        import threading

        self._cancel = threading.Event()

    def cancel(self):
        self._cancel.set()

    def run(self):
        try:
            if self._cancel.is_set():
                self.error.emit("Cancelled.")
                return
            from core.pdf_tools import split_pdf_to_pages

            self.status.emit("Splitting PDF…")
            self.progress.emit(0.0)
            res = split_pdf_to_pages(
                Path(self.input_pdf),
                output_dir=Path(self.output_dir),
                overwrite=bool(self.overwrite),
                page_ranges=(self.page_ranges or "all"),
            )
            if not res.ok:
                self.error.emit(res.message or "Split failed.")
                return
            self.progress.emit(1.0)
            self.finished.emit(res.message or f"Saved to: {self.output_dir}")
        except Exception as e:
            msg = str(e)
            if "Cancelled" in msg:
                self.error.emit("Cancelled.")
            else:
                self.error.emit(msg)


class _QtPdfExtractWorker(QObject):
    status = Signal(str)
    progress = Signal(float)
    finished = Signal(str)
    error = Signal(str)

    def __init__(self, *, input_pdf: str, output_pdf: str, page_ranges: str, overwrite: bool):
        super().__init__()
        self.input_pdf = input_pdf
        self.output_pdf = output_pdf
        self.page_ranges = page_ranges
        self.overwrite = overwrite
        import threading

        self._cancel = threading.Event()

    def cancel(self):
        self._cancel.set()

    def run(self):
        try:
            if self._cancel.is_set():
                self.error.emit("Cancelled.")
                return
            from core.pdf_tools import extract_pages_to_pdf

            self.status.emit("Extracting pages…")
            self.progress.emit(0.0)
            res = extract_pages_to_pdf(
                Path(self.input_pdf),
                output_pdf=Path(self.output_pdf),
                page_ranges=(self.page_ranges or "all"),
                overwrite=bool(self.overwrite),
            )
            if not res.ok:
                self.error.emit(res.message or "Extract failed.")
                return
            self.progress.emit(1.0)
            self.finished.emit(res.message or f"Saved: {self.output_pdf}")
        except Exception as e:
            msg = str(e)
            if "Cancelled" in msg:
                self.error.emit("Cancelled.")
            else:
                self.error.emit(msg)


class _QtPdfRotateWorker(QObject):
    status = Signal(str)
    progress = Signal(float)
    finished = Signal(str)
    error = Signal(str)

    def __init__(
        self, *, input_pdf: str, output_pdf: str, rotation_degrees: int, page_ranges: str, overwrite: bool
    ):
        super().__init__()
        self.input_pdf = input_pdf
        self.output_pdf = output_pdf
        self.rotation_degrees = int(rotation_degrees)
        self.page_ranges = page_ranges
        self.overwrite = overwrite
        import threading

        self._cancel = threading.Event()

    def cancel(self):
        self._cancel.set()

    def run(self):
        try:
            if self._cancel.is_set():
                self.error.emit("Cancelled.")
                return
            from core.pdf_tools import rotate_pdf

            self.status.emit("Rotating…")
            self.progress.emit(0.0)
            res = rotate_pdf(
                Path(self.input_pdf),
                output_pdf=Path(self.output_pdf),
                rotation_degrees=int(self.rotation_degrees),
                page_ranges=(self.page_ranges or "all"),
                overwrite=bool(self.overwrite),
            )
            if not res.ok:
                self.error.emit(res.message or "Rotate failed.")
                return
            self.progress.emit(1.0)
            self.finished.emit(res.message or f"Saved: {self.output_pdf}")
        except Exception as e:
            msg = str(e)
            if "Cancelled" in msg:
                self.error.emit("Cancelled.")
            else:
                self.error.emit(msg)


class _QtPdfSplitBookmarksWorker(QObject):
    status = Signal(str)
    progress = Signal(float)
    finished = Signal(str)
    error = Signal(str)

    def __init__(self, *, input_pdf: str, output_dir: str, overwrite: bool, min_pages: int):
        super().__init__()
        self.input_pdf = input_pdf
        self.output_dir = output_dir
        self.overwrite = overwrite
        self.min_pages = int(min_pages)
        import threading

        self._cancel = threading.Event()

    def cancel(self):
        self._cancel.set()

    def run(self):
        try:
            if self._cancel.is_set():
                self.error.emit("Cancelled.")
                return
            from core.pdf_tools import split_pdf_by_bookmarks

            self.status.emit("Splitting by bookmarks…")
            self.progress.emit(0.0)
            res = split_pdf_by_bookmarks(
                Path(self.input_pdf),
                output_dir=Path(self.output_dir),
                overwrite=bool(self.overwrite),
                min_pages=int(self.min_pages),
            )
            if not res.ok:
                self.error.emit(res.message or "Bookmark split failed.")
                return
            self.progress.emit(1.0)
            self.finished.emit(res.message or f"Saved to: {self.output_dir}")
        except Exception as e:
            msg = str(e)
            if "Cancelled" in msg:
                self.error.emit("Cancelled.")
            else:
                self.error.emit(msg)


class _QtOfficeBatchWorker(QObject):
    status = Signal(str)
    progress = Signal(float)
    finished = Signal(str)
    error = Signal(str)

    def __init__(
        self,
        *,
        folder: str,
        source_sub: str,
        include_subfolders: bool,
        output_subfolder: str,
        output_root: str,
        output_format: str,
        overwrite: bool,
    ):
        super().__init__()
        self.folder = folder
        self.source_sub = source_sub
        self.include_subfolders = include_subfolders
        self.output_subfolder = output_subfolder
        self.output_root = output_root
        self.output_format = output_format
        self.overwrite = overwrite
        import threading

        self._cancel = threading.Event()

    def cancel(self):
        self._cancel.set()

    @staticmethod
    def _allowed_exts_for(fmt: str) -> set[str]:
        fmt = (fmt or "pdf").strip().lower().lstrip(".")
        doc = {".doc", ".docx", ".odt", ".rtf", ".txt", ".html", ".htm"}
        sheet = {".xls", ".xlsx", ".ods", ".csv", ".tsv"}
        pres = {".ppt", ".pptx", ".odp"}
        office = doc | sheet | pres
        if fmt in {"pdf"}:
            return office
        if fmt in {"docx", "odt", "txt", "html"}:
            return doc | {".pdf"}
        if fmt in {"xlsx", "ods", "csv"}:
            return sheet
        if fmt in {"pptx", "odp"}:
            return pres
        return office

    def run(self):
        try:
            from core.lo_converter import LibreOfficeConverter
            from core.pdf_text_extract import pdf_to_txt
            from core.pdf_to_docx import pdf_to_docx_text

            conv = LibreOfficeConverter()
            if not conv.is_available():
                self.error.emit("LibreOffice (soffice) not found. Configure it in Settings.")
                return

            folder = Path(self.folder)
            src_sub = (self.source_sub or "").strip().replace("\\", "/").strip("/")
            src_root = folder / src_sub if src_sub else folder
            if not src_root.exists():
                self.error.emit("Source folder not found.")
                return

            out_fmt = (self.output_format or "pdf").strip().lower().lstrip(".")
            exts = self._allowed_exts_for(out_fmt)

            include_sub = bool(self.include_subfolders)
            pattern = "**/*" if include_sub else "*"

            out_base = folder if (self.output_root or "target").strip().lower() == "target" else src_root
            out_dir = out_base / (self.output_subfolder or "Converted_Office")
            out_dir.mkdir(parents=True, exist_ok=True)

            files = [p for p in src_root.glob(pattern) if p.is_file() and p.suffix.lower() in exts]
            # Avoid re-processing outputs if output is inside the scan tree.
            try:
                out_res = out_dir.resolve()
                files = [p for p in files if not str(p.resolve()).startswith(str(out_res))]
            except Exception:
                pass

            files = sorted(files, key=lambda p: str(p).lower())
            total = len(files)
            if total <= 0:
                self.finished.emit("No matching documents found.")
                self.progress.emit(1.0)
                return

            converted = 0
            skipped = 0
            for i, f in enumerate(files, start=1):
                if self._cancel.is_set():
                    self.error.emit("Cancelled.")
                    return
                self.status.emit(f"Converting {i}/{total}: {f.name}")
                self.progress.emit(max(0.0, min(1.0, float(i - 1) / float(max(1, total)))))

                out_guess = out_dir / (f.stem + "." + out_fmt)
                if out_guess.exists() and not bool(self.overwrite):
                    skipped += 1
                    continue

                out_path, err_msg = conv.convert_to_format_verbose(f, out_dir=out_dir, output_format=out_fmt)
                ok_this = out_path is not None and out_path.exists()

                # LibreOffice sometimes creates empty TXT exports for PDFs. Detect and fallback.
                if ok_this and f.suffix.lower() == ".pdf" and out_fmt == "txt":
                    try:
                        if out_path and out_path.stat().st_size == 0:
                            ok_this = False
                    except Exception:
                        ok_this = False

                if not ok_this and f.suffix.lower() == ".pdf" and out_fmt == "docx":
                    try:
                        fb = out_dir / (f.stem + ".docx")
                        pdf_to_docx_text(f, output_path=fb)
                        if fb.exists():
                            ok_this = True
                    except Exception as e:
                        err_msg = (err_msg or "") + f" | Fallback failed: {e}"

                if not ok_this and f.suffix.lower() == ".pdf" and out_fmt == "txt":
                    try:
                        fb = out_dir / (f.stem + ".txt")
                        pdf_to_txt(f, output_path=fb)
                        if fb.exists():
                            ok_this = True
                    except Exception as e:
                        err_msg = (err_msg or "") + f" | TXT fallback failed: {e}"

                if ok_this:
                    converted += 1
                else:
                    # Don't fail the entire job; log and continue.
                    if err_msg:
                        self.status.emit(f"Failed: {f.name} ({err_msg[:120]})")

            self.progress.emit(1.0)
            self.finished.emit(f"Converted {converted} files to .{out_fmt} (skipped {skipped}).")
        except Exception as e:
            msg = str(e)
            if "Cancelled" in msg:
                self.error.emit("Cancelled.")
            else:
                self.error.emit(msg)


class _QtOfficeSingleWorker(QObject):
    status = Signal(str)
    progress = Signal(float)
    finished = Signal(str)
    error = Signal(str)

    def __init__(self, *, folder: str, input_path: str, output_subfolder: str, output_root: str, output_format: str, overwrite: bool):
        super().__init__()
        self.folder = folder
        self.input_path = input_path
        self.output_subfolder = output_subfolder
        self.output_root = output_root
        self.output_format = output_format
        self.overwrite = overwrite
        import threading

        self._cancel = threading.Event()

    def cancel(self):
        self._cancel.set()

    def run(self):
        try:
            from core.lo_converter import LibreOfficeConverter
            from core.pdf_text_extract import pdf_to_txt
            from core.pdf_to_docx import pdf_to_docx_text

            conv = LibreOfficeConverter()
            if not conv.is_available():
                self.error.emit("LibreOffice (soffice) not found. Configure it in Settings.")
                return

            inp = Path(self.input_path)
            if not inp.exists():
                self.error.emit("Input file not found.")
                return

            out_fmt = (self.output_format or "pdf").strip().lower().lstrip(".")
            target = Path(self.folder)
            out_base = target if (self.output_root or "target").strip().lower() == "target" else inp.parent
            out_dir = out_base / (self.output_subfolder or "Converted_Office")
            out_dir.mkdir(parents=True, exist_ok=True)

            out_guess = out_dir / (inp.stem + "." + out_fmt)
            if out_guess.exists() and not bool(self.overwrite):
                self.progress.emit(1.0)
                self.finished.emit("Output already exists (skipped).")
                return

            self.status.emit(f"Converting: {inp.name}")
            self.progress.emit(0.0)
            out_path, err_msg = conv.convert_to_format_verbose(inp, out_dir=out_dir, output_format=out_fmt)
            ok_this = out_path is not None and out_path.exists()

            if ok_this and inp.suffix.lower() == ".pdf" and out_fmt == "txt":
                try:
                    if out_path and out_path.stat().st_size == 0:
                        ok_this = False
                except Exception:
                    ok_this = False

            if not ok_this and inp.suffix.lower() == ".pdf" and out_fmt == "docx":
                try:
                    fb = out_dir / (inp.stem + ".docx")
                    pdf_to_docx_text(inp, output_path=fb)
                    if fb.exists():
                        ok_this = True
                        out_path = fb
                except Exception as e:
                    err_msg = (err_msg or "") + f" | Fallback failed: {e}"

            if not ok_this and inp.suffix.lower() == ".pdf" and out_fmt == "txt":
                try:
                    fb = out_dir / (inp.stem + ".txt")
                    pdf_to_txt(inp, output_path=fb)
                    if fb.exists():
                        ok_this = True
                        out_path = fb
                except Exception as e:
                    err_msg = (err_msg or "") + f" | TXT fallback failed: {e}"

            if not ok_this:
                self.error.emit(err_msg or "Conversion failed.")
                return

            self.progress.emit(1.0)
            self.finished.emit(f"Saved: {out_path}")
        except Exception as e:
            msg = str(e)
            if "Cancelled" in msg:
                self.error.emit("Cancelled.")
            else:
                self.error.emit(msg)


class _QtMarkdownExportWorker(QObject):
    status = Signal(str)
    progress = Signal(float)
    finished = Signal(str)
    error = Signal(str)

    def __init__(
        self,
        *,
        folder: str,
        markdown_text: str,
        output_subfolder: str,
        output_format: str,
        output_name: str,
        overwrite: bool,
    ):
        super().__init__()
        self.folder = folder
        self.markdown_text = markdown_text
        self.output_subfolder = output_subfolder
        self.output_format = output_format
        self.output_name = output_name
        self.overwrite = overwrite
        import threading

        self._cancel = threading.Event()

    def cancel(self):
        self._cancel.set()

    def run(self):
        try:
            from core.office_tools import markdown_to_pdf, markdown_to_docx

            root = Path(self.folder)
            out_dir = root / (self.output_subfolder or "Converted_Office")
            out_dir.mkdir(parents=True, exist_ok=True)

            fmt = (self.output_format or "pdf").strip().lower()
            name = (self.output_name or "markdown_document").strip()
            name = re.sub(r"[^A-Za-z0-9_.-]+", "_", name).strip("._") or "markdown_document"

            targets: list[str]
            if fmt == "pdf + docx":
                targets = ["pdf", "docx"]
            elif fmt in {"pdf", "docx"}:
                targets = [fmt]
            else:
                self.error.emit(f"Unsupported export format: {self.output_format}")
                return

            outputs: list[str] = []
            total = max(1, len(targets))
            for idx, t in enumerate(targets, start=1):
                if self._cancel.is_set():
                    self.error.emit("Cancelled.")
                    return
                self.status.emit(f"Exporting {idx}/{total}: .{t}")
                self.progress.emit(max(0.0, min(1.0, float(idx - 1) / float(total))))
                out_path = out_dir / f"{name}.{t}"
                if t == "pdf":
                    res = markdown_to_pdf(self.markdown_text, output_path=out_path, overwrite=bool(self.overwrite))
                else:
                    res = markdown_to_docx(self.markdown_text, output_path=out_path, overwrite=bool(self.overwrite))
                if not res.ok:
                    self.error.emit(res.message or f"Failed to export {t}.")
                    return
                if res.output_path:
                    outputs.append(str(res.output_path))

            self.progress.emit(1.0)
            if outputs:
                self.finished.emit("Saved: " + " | ".join(outputs))
            else:
                self.finished.emit("Markdown exported.")
        except Exception as e:
            msg = str(e)
            if "Cancelled" in msg:
                self.error.emit("Cancelled.")
            else:
                self.error.emit(msg)


class _QtLibreOfficeDownloadWorker(QObject):
    status = Signal(str)
    progress = Signal(float)
    finished = Signal(str, str)  # message, file_path
    error = Signal(str)

    def __init__(self, *, dest_dir: Path):
        super().__init__()
        self.dest_dir = Path(dest_dir)
        import threading

        self._cancel = threading.Event()

    def cancel(self):
        self._cancel.set()

    def run(self):
        try:
            import platform as _platform
            from core.libreoffice_downloader import get_latest_stable_version, build_download_url, download_file

            self.status.emit("Checking latest LibreOffice version…")
            v = get_latest_stable_version() or ""
            if not v:
                raise RuntimeError("Could not fetch LibreOffice version list.")

            plat = "windows" if _platform.system().lower().startswith("win") else "macos" if _platform.system().lower() == "darwin" else "linux"
            arch = "x64"
            try:
                mach = _platform.machine().lower()
                if "arm" in mach:
                    arch = "arm"
            except Exception:
                pass

            built = build_download_url(v, platform=plat, arch=arch)
            if not built:
                raise RuntimeError("Unsupported platform for automatic download.")
            url, filename = built

            self.dest_dir.mkdir(parents=True, exist_ok=True)
            dest_path = self.dest_dir / filename

            def cb(frac: float, downloaded: int, total: int, speed_bps: float):
                if self._cancel.is_set():
                    raise RuntimeError("Cancelled.")
                sp = ""
                try:
                    if speed_bps > 0:
                        sp = f" • {speed_bps/1024/1024:.1f} MB/s"
                except Exception:
                    pass
                self.status.emit(f"Downloading {filename} ({int(frac*100)}%){sp}")
                self.progress.emit(max(0.0, min(1.0, float(frac))))

            res = download_file(url, dest_path, cancel_event=self._cancel, progress_cb=cb)
            if not res.ok or not res.path:
                raise RuntimeError(res.message or "Download failed.")
            self.progress.emit(1.0)
            self.finished.emit("Downloaded LibreOffice installer.", res.path)
        except Exception as e:
            msg = str(e)
            if "Cancelled" in msg:
                self.error.emit("Cancelled.")
            else:
                self.error.emit(msg)


class _QtLibreOfficeDownloadDialog(QDialog):
    def __init__(self, parent: QWidget):
        super().__init__(parent)
        self.setWindowTitle("Download LibreOffice")
        self.setModal(True)
        self.setMinimumWidth(640)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(18, 16, 18, 16)
        layout.setSpacing(12)

        title = QLabel("Download LibreOffice")
        title.setStyleSheet("font-size:16px; font-weight:800; color:#ffffff;")
        layout.addWidget(title)

        self.status = QLabel("Preparing…")
        self.status.setStyleSheet("color:#9aa0a9;")
        self.status.setWordWrap(True)
        layout.addWidget(self.status)

        self.bar = QProgressBar()
        self.bar.setRange(0, 1000)
        self.bar.setValue(0)
        layout.addWidget(self.bar)

        row = QHBoxLayout()
        row.addStretch(1)
        self.btn_cancel = QPushButton("Cancel")
        self.btn_cancel.clicked.connect(self._cancel)
        row.addWidget(self.btn_cancel)
        self.btn_open = QPushButton("Open Folder")
        self.btn_open.setEnabled(False)
        self.btn_open.clicked.connect(self._open_folder)
        row.addWidget(self.btn_open)
        layout.addLayout(row)

        try:
            base = Path.home() / ".fylorra" / "downloads"
        except Exception:
            base = Path.cwd()
        self._download_path = ""

        self.worker = _QtLibreOfficeDownloadWorker(dest_dir=base)
        self._thread = QThread(self)
        self.worker.moveToThread(self._thread)
        self._thread.started.connect(self.worker.run)
        self.worker.status.connect(self.status.setText)
        self.worker.progress.connect(lambda p: self.bar.setValue(int(max(0.0, min(1.0, p)) * 1000)))
        self.worker.finished.connect(self._on_done)
        self.worker.error.connect(self._on_error)
        self.worker.finished.connect(self._thread.quit)
        self.worker.error.connect(self._thread.quit)
        self._thread.finished.connect(self._thread.deleteLater)
        self._thread.finished.connect(self.worker.deleteLater)
        self._thread.start()

    def _cancel(self):
        try:
            self.worker.cancel()
        except Exception:
            pass
        self.btn_cancel.setEnabled(False)
        self.status.setText("Cancelling…")

    def _on_done(self, msg: str, path: str):
        self._download_path = path or ""
        self.status.setText(f"{msg}\n\nSaved to:\n{path}\n\nRun the installer to enable Office conversions.")
        self.bar.setValue(1000)
        self.btn_cancel.setText("Close")
        try:
            self.btn_cancel.clicked.disconnect()
        except Exception:
            pass
        self.btn_cancel.clicked.connect(self.accept)
        self.btn_cancel.setEnabled(True)
        self.btn_open.setEnabled(True)

    def _on_error(self, msg: str):
        self.status.setText("Download failed.")
        QMessageBox.critical(self, "LibreOffice Download", msg)
        self.btn_cancel.setText("Close")
        try:
            self.btn_cancel.clicked.disconnect()
        except Exception:
            pass
        self.btn_cancel.clicked.connect(self.reject)
        self.btn_cancel.setEnabled(True)

    def _open_folder(self):
        p = self._download_path
        if not p:
            return
        try:
            import subprocess

            subprocess.Popen(f'explorer /select,"{p}"')
        except Exception:
            try:
                QDesktopServices.openUrl(QUrl.fromLocalFile(str(Path(p).parent)))
            except Exception:
                pass


class _QtFfmpegToolsDownloadWorker(QObject):
    status = Signal(str)
    progress = Signal(float)
    finished = Signal(bool, str)  # ok, bin_dir
    error = Signal(str)

    def run(self):
        import os
        import zipfile
        import tempfile
        from pathlib import Path
        import requests

        try:
            if os.name != "nt":
                raise RuntimeError("FFmpeg tools auto-download is currently implemented for Windows only.")

            url = "https://www.gyan.dev/ffmpeg/builds/ffmpeg-release-essentials.zip"
            self.status.emit("Downloading FFmpeg tools…")
            self.progress.emit(0.01)

            tmp = Path(tempfile.gettempdir()) / "fylorra_ffmpeg_tools.zip"
            with requests.get(url, stream=True, timeout=60) as r:
                r.raise_for_status()
                total = int(r.headers.get("content-length", 0) or 0)
                got = 0
                with open(tmp, "wb") as f:
                    for chunk in r.iter_content(chunk_size=1024 * 256):
                        if not chunk:
                            continue
                        f.write(chunk)
                        got += len(chunk)
                        if total > 0:
                            self.progress.emit(min(0.78, 0.02 + 0.76 * (got / total)))

            self.status.emit("Extracting…")
            self.progress.emit(0.8)

            base = Path.home() / ".fylorra" / "tools" / "ffmpeg"
            base.mkdir(parents=True, exist_ok=True)

            ffmpeg_path = None
            with zipfile.ZipFile(tmp, "r") as z:
                # Extract only the needed executables.
                names = z.namelist()
                wanted = [n for n in names if n.lower().endswith(("/bin/ffmpeg.exe", "/bin/ffprobe.exe", "/bin/ffplay.exe"))]
                if not wanted:
                    raise RuntimeError("Could not find ffmpeg/ffprobe/ffplay inside the bundle.")
                for i, n in enumerate(wanted):
                    z.extract(n, base)
                    self.progress.emit(0.8 + 0.18 * ((i + 1) / max(1, len(wanted))))

            # Find extracted ffmpeg.exe
            found = list(base.rglob("ffmpeg.exe"))
            if found:
                ffmpeg_path = found[0]
            if not ffmpeg_path or not ffmpeg_path.exists():
                raise RuntimeError("Extraction failed (ffmpeg.exe not found).")

            ffprobe_path = ffmpeg_path.with_name("ffprobe.exe")
            ffplay_path = ffmpeg_path.with_name("ffplay.exe")

            from core.tool_manager import ToolManager

            tm = ToolManager()
            tm.set_ffmpeg(str(ffmpeg_path))
            if ffprobe_path.exists():
                tm.set_ffprobe(str(ffprobe_path))
            if ffplay_path.exists():
                tm.set_ffplay(str(ffplay_path))

            self.status.emit("Installed.")
            self.progress.emit(1.0)
            self.finished.emit(True, str(ffmpeg_path.parent))
        except Exception as e:
            self.error.emit(str(e))


class _QtFfmpegToolsDownloadDialog(QDialog):
    def __init__(self, parent: QWidget):
        super().__init__(parent)
        self.setWindowTitle("Download FFmpeg Tools")
        self.setModal(True)
        self.setMinimumWidth(640)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(18, 16, 18, 16)
        layout.setSpacing(12)

        title = QLabel("Download FFmpeg Tools")
        title.setStyleSheet("font-size:16px; font-weight:800; color:#ffffff;")
        layout.addWidget(title)

        self.status = QLabel("Preparing…")
        self.status.setStyleSheet("color:#9aa0a9;")
        self.status.setWordWrap(True)
        layout.addWidget(self.status)

        self.bar = QProgressBar()
        self.bar.setRange(0, 1000)
        self.bar.setValue(0)
        layout.addWidget(self.bar)

        row = QHBoxLayout()
        row.addStretch(1)
        self.btn_cancel = QPushButton("Cancel")
        self.btn_cancel.clicked.connect(self._cancel)
        row.addWidget(self.btn_cancel)
        layout.addLayout(row)

        self.worker = _QtFfmpegToolsDownloadWorker()
        self._thread = QThread(self)
        self.worker.moveToThread(self._thread)
        self._thread.started.connect(self.worker.run)
        self.worker.status.connect(self.status.setText)
        self.worker.progress.connect(lambda p: self.bar.setValue(int(max(0.0, min(1.0, p)) * 1000)))
        self.worker.finished.connect(self._on_done)
        self.worker.error.connect(self._on_error)
        self.worker.finished.connect(self._thread.quit)
        self.worker.error.connect(self._thread.quit)
        self._thread.finished.connect(self._thread.deleteLater)
        self._thread.finished.connect(self.worker.deleteLater)
        self._thread.start()

    def _cancel(self):
        # Simple dialog; cancellation requires restarting download right now.
        self.status.setText("Cancel not supported during download. Close this dialog to abort.")
        self.btn_cancel.setText("Close")
        try:
            self.btn_cancel.clicked.disconnect()
        except Exception:
            pass
        self.btn_cancel.clicked.connect(self.reject)

    def _on_done(self, ok: bool, bin_dir: str):
        self.status.setText(f"Installed FFmpeg tools.\n\nBin:\n{bin_dir}\n\nRestart conversions to use ffprobe/ffplay.")
        self.bar.setValue(1000 if ok else self.bar.value())
        self.btn_cancel.setText("OK")
        try:
            self.btn_cancel.clicked.disconnect()
        except Exception:
            pass
        self.btn_cancel.clicked.connect(self.accept)

    def _on_error(self, msg: str):
        self.status.setText("Download failed.")
        QMessageBox.critical(self, "FFmpeg Tools", msg)
        self.btn_cancel.setText("Close")
        try:
            self.btn_cancel.clicked.disconnect()
        except Exception:
            pass
        self.btn_cancel.clicked.connect(self.reject)


class _QtCloudWorker(QObject):
    status = Signal(str)
    progress = Signal(int, int, str)  # done, total, message
    finished = Signal(object)
    error = Signal(str)

    def __init__(self, *, fn):
        super().__init__()
        self._fn = fn
        import threading

        self._cancel = threading.Event()

    def cancel(self):
        self._cancel.set()

    def _check_cancel(self):
        if self._cancel.is_set():
            raise RuntimeError("Cancelled.")

    def run(self):
        try:
            # Backward-compatible signature support:
            # - fn(status_cb, progress_cb)
            # - fn(status_cb, progress_cb, cancel_cb) where cancel_cb raises if cancelled
            try:
                res = self._fn(self.status.emit, self.progress.emit, self._check_cancel)
            except TypeError:
                res = self._fn(self.status.emit, self.progress.emit)
            self.finished.emit(res)
        except Exception as e:
            self.error.emit(str(e))


class _QtDeviceCodeDialog(QDialog):
    def __init__(self, *, title: str, url: str, code: str, message: str, parent=None):
        super().__init__(parent)
        self.setWindowTitle(title)
        self.setModal(True)
        self.setMinimumWidth(520)
        self.setStyleSheet(_qt_modern_dialog_stylesheet())

        layout = QVBoxLayout(self)
        layout.setContentsMargins(18, 16, 18, 16)
        layout.setSpacing(12)

        header = QFrame()
        header.setObjectName("DialogHeader")
        hl = QVBoxLayout(header)
        hl.setContentsMargins(16, 12, 16, 12)
        hl.setSpacing(6)

        t = QLabel(title)
        t.setObjectName("DialogTitle")
        hl.addWidget(t)

        sub = QLabel("Copy the code, then open the login page.")
        sub.setObjectName("DialogSubtitle")
        hl.addWidget(sub)
        layout.addWidget(header)

        msg = QTextBrowser()
        msg.setPlainText(message.strip() or "Follow the instructions to sign in.")
        msg.setFixedHeight(120)
        layout.addWidget(msg)

        code_row = QHBoxLayout()
        code_row.setSpacing(10)
        self.code_edit = QLineEdit(code or "")
        self.code_edit.setReadOnly(True)
        self.code_edit.setPlaceholderText("Code")
        code_row.addWidget(self.code_edit, 1)
        btn_copy = QPushButton("Copy")
        btn_copy.setObjectName("PrimaryButton")
        btn_copy.clicked.connect(self._copy)
        code_row.addWidget(btn_copy)
        layout.addLayout(code_row)

        btn_row = QHBoxLayout()
        btn_row.setSpacing(10)
        btn_row.addStretch(1)
        btn_open = QPushButton("Open Login Page")
        btn_open.setObjectName("PrimaryButton")
        btn_open.clicked.connect(lambda: self._open(url))
        btn_row.addWidget(btn_open)
        btn_close = QPushButton("Close")
        btn_close.clicked.connect(self.accept)
        btn_row.addWidget(btn_close)
        layout.addLayout(btn_row)

        self._url = url

    def _copy(self):
        try:
            QApplication.clipboard().setText(self.code_edit.text().strip())
        except Exception:
            pass

    def _open(self, url: str):
        u = (url or "").strip()
        if not u:
            return
        try:
            QDesktopServices.openUrl(QUrl(u))
        except Exception:
            pass


class _QtWorkspaceWorker(QObject):
    status = Signal(str)
    log = Signal(str)
    progress = Signal(float)
    finished = Signal(str)
    error = Signal(str)
    review_requested = Signal(str, dict)

    def __init__(self, *, backend, folder: str, include_subfolders: bool, actions: list[str]):
        super().__init__()
        self.backend = backend
        self.folder = folder
        self.include_subfolders = include_subfolders
        self.actions = list(actions)
        import threading

        self._cancel = threading.Event()
        self._review_evt = threading.Event()
        self._review_result: dict | None = None

    def cancel(self):
        self._cancel.set()
        try:
            self._review_evt.set()
        except Exception:
            pass

    def _check_cancel(self):
        if self._cancel.is_set():
            raise RuntimeError("Cancelled.")

    def submit_review_result(self, result: dict) -> None:
        """
        Called from the UI thread via queued connection.
        Expected: {"decision":"apply|skip|cancel", ...}
        """
        try:
            self._review_result = dict(result or {})
        except Exception:
            self._review_result = {"decision": "cancel"}
        try:
            self._review_evt.set()
        except Exception:
            pass

    def _wait_for_review(self) -> dict:
        import time

        self._review_evt.clear()
        self._review_result = None
        while not self._review_evt.is_set():
            self._check_cancel()
            time.sleep(0.05)
        self._check_cancel()
        return dict(self._review_result or {"decision": "cancel"})

    def run(self):
        try:
            from pathlib import Path
            from core.library_index import LibraryIndex
            from core.archive_tools import create_archive
            from core.image_converter import convert_images_in_folder
            from core.media_converter import convert_media_in_folder
            from core.media_tools import convert_media_file
            from core.enhanced_categorizer import EnhancedCategorizer
            from core.semantic_analyzer import SemanticAnalyzer
            from dataclasses import asdict
            import json
            import time
            import shutil
            import zipfile
            import os

            folder = Path(self.folder)
            ai = getattr(self.backend, "ai_manager", None)
            actions = list(self.actions)
            total_steps = max(1, len(actions))

            def set_prog(i: int, frac: float):
                self.progress.emit(max(0.0, min(1.0, (float(i) + float(frac)) / float(total_steps))))

            # Reports folder
            try:
                base = Path.home() / ".fylorra" / "reports" / "workspace" / time.strftime("%Y%m%d_%H%M%S")
                base.mkdir(parents=True, exist_ok=True)
            except Exception:
                base = Path.cwd()

            for idx, act in enumerate(actions):
                self._check_cancel()
                self.status.emit(f"{act}…")
                self.log.emit(f"Start: {act}")
                set_prog(idx, 0.0)

                if act == "index_folder":
                    # Preview/confirm (indexing is non-destructive but can take time).
                    def _estimate_files(limit: int = 50000) -> str | None:
                        try:
                            count = 0
                            for root, _dirs, files in os.walk(folder):
                                count += len(files)
                                if count >= limit:
                                    return f"{limit}+"
                                if not self.include_subfolders:
                                    break
                            return str(count)
                        except Exception:
                            return None

                    self.review_requested.emit(
                        "index",
                        {
                            "folder": str(folder),
                            "estimated_files": _estimate_files(),
                            "note": "Builds a local index to speed up AI Search. No files are moved or modified.",
                        },
                    )
                    decision = self._wait_for_review()
                    if decision.get("decision") == "cancel":
                        raise RuntimeError("Cancelled.")
                    if decision.get("decision") == "skip":
                        self.log.emit("Index skipped.")
                        set_prog(idx, 1.0)
                        continue

                    lib = LibraryIndex()

                    def cb(msg: str, p: float):
                        self._check_cancel()
                        self.status.emit(msg)
                        set_prog(idx, float(p))

                    lib.index_folder(folder, include_subfolders=bool(self.include_subfolders), ai_manager=ai, ai_summarize=False, progress_cb=cb)
                    self.log.emit("Indexed folder.")

                elif act == "zip_folder":
                    out = folder / "Archive.zip"
                    pattern = "**/*" if self.include_subfolders else "*"
                    items: list[dict] = []
                    for p in folder.glob(pattern):
                        if not p.is_file():
                            continue
                        if p.name.lower() == "archive.zip":
                            continue
                        try:
                            arc = str(p.relative_to(folder)).replace("\\", "/")
                        except Exception:
                            arc = p.name
                        items.append({"path": str(p), "arc": arc})
                        if len(items) >= 8000:
                            break
                    note = ""
                    try:
                        total_files = sum(1 for p in folder.glob(pattern) if p.is_file())
                        if total_files > len(items):
                            note = f"Showing first {len(items)} files (of {total_files})."
                    except Exception:
                        pass
                    self.review_requested.emit(
                        "zip",
                        {
                            "folder": str(folder),
                            "output": str(out),
                            "items": items,
                            "note": note,
                        },
                    )
                    decision = self._wait_for_review()
                    if decision.get("decision") == "cancel":
                        raise RuntimeError("Cancelled.")
                    if decision.get("decision") == "skip":
                        self.log.emit("Archive skipped.")
                        set_prog(idx, 1.0)
                        continue

                    selected = set(map(str, decision.get("selected", []) or []))
                    chosen = [it for it in items if (not selected) or str(it.get("path") or "") in selected]
                    if not chosen:
                        self.log.emit("Archive skipped (no files selected).")
                        set_prog(idx, 1.0)
                        continue
                    try:
                        if out.exists():
                            out.unlink()
                    except Exception:
                        pass

                    total = max(1, len(chosen))
                    with zipfile.ZipFile(out, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=6) as zf:
                        for i2, it in enumerate(chosen, start=1):
                            self._check_cancel()
                            p = Path(str(it.get("path") or ""))
                            arc = str(it.get("arc") or p.name)
                            self.status.emit(f"Archiving {i2}/{len(chosen)}: {p.name}")
                            set_prog(idx, float(i2 - 1) / float(total))
                            try:
                                zf.write(str(p), arcname=arc)
                            except Exception:
                                continue
                    self.log.emit(f"Archive created: {out}")

                elif act == "convert_images_webp":
                    # Preview list + allow skipping specific files.
                    in_exts = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}
                    out_dir = folder / "Converted_Images"
                    try:
                        out_dir_resolved = out_dir.resolve()
                    except Exception:
                        out_dir_resolved = out_dir
                    pattern = "**/*" if self.include_subfolders else "*"
                    items: list[dict] = []
                    for p in folder.glob(pattern):
                        if not p.is_file():
                            continue
                        if p.suffix.lower() not in in_exts:
                            continue
                        try:
                            rp = p.resolve()
                            if rp == out_dir_resolved or str(rp).startswith(str(out_dir_resolved)):
                                continue
                        except Exception:
                            pass
                        items.append({"path": str(p), "output_name": p.stem + ".webp"})
                        if len(items) >= 6000:
                            break
                    note = ""
                    try:
                        total_imgs = sum(1 for p in folder.glob(pattern) if p.is_file() and p.suffix.lower() in in_exts)
                        if total_imgs > len(items):
                            note = f"Showing first {len(items)} files (of {total_imgs})."
                    except Exception:
                        pass
                    self.review_requested.emit(
                        "convert_images",
                        {
                            "folder": str(folder),
                            "output": str(out_dir),
                            "items": items,
                            "note": note,
                        },
                    )
                    decision = self._wait_for_review()
                    if decision.get("decision") == "cancel":
                        raise RuntimeError("Cancelled.")
                    if decision.get("decision") == "skip":
                        self.log.emit("Image conversion skipped.")
                        set_prog(idx, 1.0)
                        continue

                    selected = set(map(str, decision.get("selected", []) or []))
                    chosen = [it for it in items if (not selected) or str(it.get("path") or "") in selected]
                    if not chosen:
                        self.log.emit("Image conversion skipped (no files selected).")
                        set_prog(idx, 1.0)
                        continue

                    def cb(cur: int, total: int, p: Path):
                        self._check_cancel()
                        self.status.emit(f"Images {cur}/{total}: {p.name}")
                        set_prog(idx, float(cur - 1) / float(max(1, total)))

                    # Convert only the selected files.
                    try:
                        from PIL import Image, ImageOps  # type: ignore
                    except Exception as e:
                        raise RuntimeError("Image conversion requires Pillow (pip install pillow).") from e
                    out_dir.mkdir(parents=True, exist_ok=True)
                    total = max(1, len(chosen))
                    converted = 0
                    skipped = 0
                    for i2, it in enumerate(chosen, start=1):
                        self._check_cancel()
                        p = Path(str(it.get("path") or ""))
                        cb(i2, total, p)
                        try:
                            rel = p.relative_to(folder)
                        except Exception:
                            rel = Path(p.name)
                        dest_dir = (out_dir / rel.parent) if self.include_subfolders else out_dir
                        dest_dir.mkdir(parents=True, exist_ok=True)
                        out_path = dest_dir / (p.stem + ".webp")
                        if out_path.exists():
                            skipped += 1
                            continue
                        try:
                            with Image.open(p) as im:
                                try:
                                    im = ImageOps.exif_transpose(im)
                                except Exception:
                                    pass
                                im.save(out_path, format="WEBP", quality=92, method=6)
                            converted += 1
                        except Exception:
                            skipped += 1
                    self.log.emit(f"Converted {converted} images (skipped {skipped}).")

                elif act == "convert_media_mp4":
                    in_exts = {
                        ".mp4",
                        ".mkv",
                        ".mov",
                        ".avi",
                        ".webm",
                        ".mp3",
                        ".wav",
                        ".m4a",
                        ".aac",
                        ".flac",
                        ".ogg",
                    }
                    out_dir = folder / "Converted_Media"
                    try:
                        out_dir_resolved = out_dir.resolve()
                    except Exception:
                        out_dir_resolved = out_dir
                    pattern = "**/*" if self.include_subfolders else "*"
                    items: list[dict] = []
                    for p in folder.glob(pattern):
                        if not p.is_file():
                            continue
                        if p.suffix.lower() not in in_exts:
                            continue
                        try:
                            rp = p.resolve()
                            if rp == out_dir_resolved or str(rp).startswith(str(out_dir_resolved)):
                                continue
                        except Exception:
                            pass
                        items.append({"path": str(p), "output_name": p.stem + ".mp4"})
                        if len(items) >= 4000:
                            break
                    note = ""
                    try:
                        total_media = sum(1 for p in folder.glob(pattern) if p.is_file() and p.suffix.lower() in in_exts)
                        if total_media > len(items):
                            note = f"Showing first {len(items)} files (of {total_media})."
                    except Exception:
                        pass
                    self.review_requested.emit(
                        "convert_media",
                        {
                            "folder": str(folder),
                            "output": str(out_dir),
                            "items": items,
                            "note": note,
                        },
                    )
                    decision = self._wait_for_review()
                    if decision.get("decision") == "cancel":
                        raise RuntimeError("Cancelled.")
                    if decision.get("decision") == "skip":
                        self.log.emit("Media conversion skipped.")
                        set_prog(idx, 1.0)
                        continue

                    selected = set(map(str, decision.get("selected", []) or []))
                    chosen = [it for it in items if (not selected) or str(it.get("path") or "") in selected]
                    if not chosen:
                        self.log.emit("Media conversion skipped (no files selected).")
                        set_prog(idx, 1.0)
                        continue

                    def cb(cur: int, total: int, p: Path):
                        self._check_cancel()
                        self.status.emit(f"Media {cur}/{total}: {p.name}")
                        set_prog(idx, float(cur - 1) / float(max(1, total)))

                    out_dir.mkdir(parents=True, exist_ok=True)
                    total = max(1, len(chosen))
                    converted = 0
                    skipped = 0
                    for i2, it in enumerate(chosen, start=1):
                        self._check_cancel()
                        p = Path(str(it.get("path") or ""))
                        cb(i2, total, p)
                        try:
                            rel = p.relative_to(folder)
                        except Exception:
                            rel = Path(p.name)
                        dest_dir = (out_dir / rel.parent) if self.include_subfolders else out_dir
                        dest_dir.mkdir(parents=True, exist_ok=True)
                        out_path = dest_dir / (p.stem + ".mp4")
                        if out_path.exists():
                            skipped += 1
                            continue

                        def _fp(frac: float):
                            self._check_cancel()
                            frac = float(frac or 0.0)
                            frac = max(0.0, min(1.0, frac))
                            set_prog(idx, (float(i2 - 1) + frac) / float(total))

                        r = convert_media_file(
                            p,
                            output_path=out_path,
                            overwrite=False,
                            preserve_metadata=True,
                            preserve_cover_art=True,
                            cancel_event=self._cancel,
                            progress_cb=_fp,
                        )
                        if r.ok:
                            converted += 1
                        else:
                            skipped += 1
                    self.log.emit(f"Converted {converted} media files (skipped {skipped}).")

                elif act.startswith("ai_hub_"):
                    # AI operations: report-only by default
                    if not ai or not getattr(ai, "is_ready", False):
                        raise RuntimeError("AI model not loaded.")

                    if act == "ai_hub_rename":
                        # Report-only rename suggestions (no filesystem changes here).
                        # Avoid importing utils.intelligent_rename (optional deps).
                        files = [p for p in folder.rglob("*") if p.is_file()][:400]
                        suggestions = []
                        for i2, p in enumerate(files, start=1):
                            self._check_cancel()
                            self.status.emit(f"Smart Rename {i2}/{len(files)}: {p.name}")
                            set_prog(idx, float(i2 - 1) / float(max(1, len(files))))
                            try:
                                new = ai.analyze_file_for_rename(p)
                                if new:
                                    suggestions.append({"path": str(p), "suggested": new})
                            except Exception:
                                continue
                        (base / "smart_rename.json").write_text(json.dumps(suggestions, indent=2), encoding="utf-8")
                        self.log.emit(f"Smart Rename report: {base/'smart_rename.json'}")

                        # Ask user to review/apply.
                        self.review_requested.emit(
                            "rename",
                            {
                                "folder": str(folder),
                                "items": suggestions,
                            },
                        )
                        decision = self._wait_for_review()
                        if decision.get("decision") == "cancel":
                            raise RuntimeError("Cancelled.")
                        if decision.get("decision") == "apply":
                            selected = set(map(str, decision.get("selected", []) or []))

                            def _sanitize_name(s: str) -> str:
                                s = (s or "").strip().strip(".")
                                s = re.sub(r"[<>:\"/\\\\|?*\\x00-\\x1f]", "_", s)
                                s = re.sub(r"[\\s_]+", "_", s)
                                return (s[:180] or "file").strip("_")

                            def _unique_path(dest: Path) -> Path:
                                if not dest.exists():
                                    return dest
                                stem = dest.stem
                                suf = dest.suffix
                                for n in range(2, 5000):
                                    cand = dest.with_name(f"{stem}_{n}{suf}")
                                    if not cand.exists():
                                        return cand
                                return dest

                            applied = 0
                            for i2, it in enumerate(suggestions, start=1):
                                self._check_cancel()
                                p = Path(it.get("path", ""))
                                if not p.exists():
                                    continue
                                if selected and str(p) not in selected:
                                    continue
                                sug = str(it.get("suggested") or "").strip()
                                if not sug:
                                    continue
                                self.status.emit(f"Applying rename {i2}/{len(suggestions)}…")
                                set_prog(idx, float(i2 - 1) / float(max(1, len(suggestions))))
                                try:
                                    dest = p.with_name(_sanitize_name(sug) + p.suffix)
                                    dest = _unique_path(dest)
                                    if dest != p:
                                        p.rename(dest)
                                        applied += 1
                                except Exception:
                                    continue
                            self.log.emit(f"Smart Rename applied: {applied} files")

                    elif act == "ai_hub_categorize":
                        cat = EnhancedCategorizer(ai)
                        out_map: dict[str, list[str]] = {}

                        def _progress(msg: str, prog: float, cur: int, total: int):
                            self._check_cancel()
                            self.status.emit(msg)
                            # cat reports 0..1 for folder-wide progress
                            set_prog(idx, float(prog))

                        res = cat.categorize_folder(
                            folder,
                            include_subfolders=bool(self.include_subfolders),
                            progress_callback=_progress,
                            cancel_check=lambda: self._cancel.is_set(),
                            smart_scope=True,
                            include_other=False,
                            use_ai_vision=False,
                            use_ai_documents=False,
                        )
                        for k, paths in (res or {}).items():
                            try:
                                if paths:
                                    out_map[str(k)] = [str(p) for p in paths]
                            except Exception:
                                continue

                        (base / "auto_categorize.json").write_text(json.dumps(out_map, indent=2), encoding="utf-8")
                        self.log.emit(f"Auto-Categorize report: {base/'auto_categorize.json'}")

                        self.review_requested.emit(
                            "categorize",
                            {
                                "folder": str(folder),
                                "items": out_map,
                            },
                        )
                        decision = self._wait_for_review()
                        if decision.get("decision") == "cancel":
                            raise RuntimeError("Cancelled.")
                        if decision.get("decision") == "apply":
                            selected = set(map(str, decision.get("selected", []) or []))
                            moved = 0
                            # items: category_key -> list[path]
                            for cat_key, paths in out_map.items():
                                self._check_cancel()
                                try:
                                    dest_rel = cat.get_category_folder(cat_key)
                                except Exception:
                                    dest_rel = "Other"
                                dest_dir = folder / dest_rel
                                try:
                                    dest_dir.mkdir(parents=True, exist_ok=True)
                                except Exception:
                                    pass
                                for p_str in paths:
                                    self._check_cancel()
                                    if selected and p_str not in selected:
                                        continue
                                    p = Path(p_str)
                                    if not p.exists():
                                        continue
                                    dest = dest_dir / p.name
                                    if dest.exists():
                                        stem = dest.stem
                                        suf = dest.suffix
                                        for n in range(2, 5000):
                                            cand = dest.with_name(f"{stem}_{n}{suf}")
                                            if not cand.exists():
                                                dest = cand
                                                break
                                    try:
                                        shutil.move(str(p), str(dest))
                                        moved += 1
                                    except Exception:
                                        continue
                            self.log.emit(f"Auto-Categorize applied: moved {moved} files")

                    elif act == "ai_hub_security":
                        # scan images only
                        imgs = [p for p in folder.rglob("*") if p.is_file() and p.suffix.lower() in {".png",".jpg",".jpeg",".webp",".bmp"}][:300]
                        out = []
                        for i2, p in enumerate(imgs, start=1):
                            self._check_cancel()
                            self.status.emit(f"Security Scan {i2}/{len(imgs)}: {p.name}")
                            set_prog(idx, float(i2 - 1) / float(max(1, len(imgs))))
                            try:
                                r = ai.detect_sensitive_content(p)
                                out.append({"path": str(p), "result": r})
                            except Exception:
                                continue
                        (base / "security_scan.json").write_text(json.dumps(out, indent=2), encoding="utf-8")
                        self.log.emit(f"Security scan report: {base/'security_scan.json'}")

                    elif act == "ai_hub_content":
                        sa = SemanticAnalyzer(ai)
                        # Bounded semantic analysis over documents only (report-only).
                        doc_exts = {".pdf", ".docx", ".doc", ".txt", ".md", ".rtf", ".csv"}
                        docs: list[Path] = [p for p in folder.rglob("*") if p.is_file() and p.suffix.lower() in doc_exts]
                        docs = docs[:250]
                        rep: list[dict] = []
                        total_docs = max(1, len(docs))
                        for i2, p in enumerate(docs, start=1):
                            self._check_cancel()
                            self.status.emit(f"Content Analysis {i2}/{len(docs)}: {p.name}")
                            set_prog(idx, float(i2 - 1) / float(total_docs))
                            try:
                                r = sa.analyze_document(p, use_cache=True)
                                if r:
                                    rep.append({"path": str(p), "analysis": asdict(r)})
                            except Exception:
                                continue

                        (base / "content_analysis.json").write_text(json.dumps(rep, indent=2), encoding="utf-8")
                        self.log.emit(f"Content analysis report: {base/'content_analysis.json'}")

                else:
                    self.log.emit(f"Skipped unknown action: {act}")

                set_prog(idx, 1.0)
                self.log.emit(f"Done: {act}")

            self.progress.emit(1.0)
            self.finished.emit("Workspace workflow complete.")
        except Exception as e:
            msg = str(e)
            if "Cancelled" in msg:
                self.error.emit("Cancelled.")
            else:
                self.error.emit(msg)
